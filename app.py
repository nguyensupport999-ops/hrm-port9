'''Tóm lại bài học rút ra từ ca này: khi HTML nằm trong components.html (iframe), việc giao tiếp với trang Streamlit cha luôn phải dùng window.top thay vì window.parent — đặc biệt trên Streamlit Cloud nơi có thể có nhiều tầng iframe lồng nhau. Và không bao giờ dùng replaceState rồi lại location.href cùng lúc vì chúng triệt tiêu nhau.
Nếu sau này cần thêm tính năng hay gặp bug mới, cứ ping lại nhé!
'''
import streamlit as st
import psycopg2
import psycopg2.extras
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta
import calendar
import random
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import tempfile
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import requests
from openpyxl.styles import Font, Alignment
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import qrcode
from io import BytesIO
import os
import sys
import subprocess
import pathlib
import streamlit.components.v1 as components
import urllib.parse
import re
import json
import unicodedata
import control_plane
from control_plane import DatabaseEngine, resolve_tenant
from import_export_hr import render_import_export_ui
import bcrypt
import chat_noi_bo
import i18n
import photo_card_gender
import base64
import tinh_thu_nhap
import thue_hkd
import mimetypes
from io import BytesIO
from datetime import time as _time
import time
import numpy as np
import cv2
from streamlit_webrtc import webrtc_streamer
import face_id_cham_cong


try:
    from config import COMPANY_CONFIG, BHXH_CONFIG, EMAIL_CONFIG, TELEGRAM_CONFIG, USERS
    print("Using local config.py")
except ImportError:
    from config_template import COMPANY_CONFIG, BHXH_CONFIG, EMAIL_CONFIG, TELEGRAM_CONFIG, USERS
    print("Using config_template.py")



@st.cache_data(ttl=3600, show_spinner=False)
def get_avatar_bytes_cached(storage_path: str) -> bytes:
    """Tải ảnh avatar có cache 1 giờ"""
    if not storage_path:
        return None
    try:
        sb = get_supabase_storage()
        if not sb:
            return None
        return sb.storage.from_(SUPABASE_BUCKET).download(storage_path)
    except Exception as e:
        print(f"Lỗi tải avatar: {e}")
        return None

@st.cache_data(ttl=3600, show_spinner=False)
def get_chat_image_bytes_cached(file_url: str) -> bytes:
    """Tải ảnh chat có cache 1 giờ"""
    if not file_url:
        return None
    try:
        sb = get_supabase_storage()
        if not sb:
            return None
        return sb.storage.from_(SUPABASE_BUCKET).download(file_url)
    except Exception as e:
        print(f"Lỗi tải ảnh chat: {e}")
        return None

@st.cache_data(ttl=300, show_spinner=False)  # Cache 5 phút
def get_dashboard_stats():
    """Lấy toàn bộ thống kê cho Dashboard trong 1 lần query duy nhất"""
    db = st.session_state.db_engine.get_connection()
    c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    
    # Gộp tất cả query vào 1 lần
    stats = {}
    
    # Tổng ứng viên
    c.execute("SELECT COUNT(*) as count FROM ung_vien")
    stats['tong_uv'] = c.fetchone()['count']
    
    # Tổng nhân viên đang làm
    c.execute("""
        SELECT COUNT(*) as count 
        FROM nhan_vien 
        WHERE trang_thai IN ('DANG_LAM','THU_VIEC') 
        AND so_hdld IS NOT NULL AND so_hdld != ''
    """)
    stats['tong_nv'] = c.fetchone()['count']
    
    # Ứng viên theo trạng thái
    c.execute("""
        SELECT trang_thai, COUNT(*) as count 
        FROM ung_vien 
        GROUP BY trang_thai
    """)
    stats['uv_by_status'] = {row['trang_thai']: row['count'] for row in c.fetchall()}
    
    # Phân bố nhân viên theo phòng ban
    c.execute("""
        SELECT phong_ban_lam_viec as phong_ban, COUNT(*) as count
        FROM nhan_vien 
        WHERE trang_thai IN ('DANG_LAM','THU_VIEC') 
        AND so_hdld IS NOT NULL AND so_hdld != ''
        GROUP BY phong_ban_lam_viec
        ORDER BY count DESC
    """)
    stats['nv_by_dept'] = c.fetchall()
    
    # Phân bố theo giới tính
    c.execute("""
        SELECT gioi_tinh, COUNT(*) as count
        FROM nhan_vien 
        WHERE trang_thai IN ('DANG_LAM','THU_VIEC') 
        AND so_hdld IS NOT NULL AND so_hdld != ''
        GROUP BY gioi_tinh
    """)
    stats['nv_by_gender'] = c.fetchall()
    
    # ... thêm các query khác tương tự ...
    
    db.close()
    return stats

# ===== CHẤM CÔNG - DANH MỤC KÝ HIỆU CHUẨN (23 ký hiệu) =====
KY_HIEU_CHAM_CONG = {
    "x":     {"ten": "Đi làm ngày thường", "nhom": "A", "cong": 1.0, "can_duyet": False},
    "x/2":   {"ten": "Đi làm nửa ngày", "nhom": "A", "cong": 0.5, "can_duyet": False},
    "P":     {"ten": "Nghỉ phép năm (cả ngày)", "nhom": "A", "cong": 1.0, "tru_phep": 1.0, "can_duyet": True},
    "1/2P":  {"ten": "Nghỉ phép nửa ngày", "nhom": "A", "cong": 0.5, "tru_phep": 0.5, "can_duyet": True},
    "NL":    {"ten": "Nghỉ lễ", "nhom": "A", "cong": 1.0, "can_duyet": False},
    "CN":    {"ten": "Nghỉ hàng tuần (Chủ nhật)", "nhom": "A", "cong": 0.0, "auto": True},
    "CT":    {"ten": "Công tác", "nhom": "A", "cong": 1.0, "can_duyet": True},
    "NB":    {"ten": "Nghỉ bù (đã OT không nhận tiền)", "nhom": "A", "cong": 1.0, "can_duyet": True},
    "Ro":    {"ten": "Nghỉ việc riêng hưởng lương", "nhom": "A", "cong": 1.0, "can_duyet": True},
    "OD":    {"ten": "Nghỉ ốm (có giấy y tế)", "nhom": "B", "cong": 0.0, "can_duyet": True, "canh_bao_bao_giam": 14},
    "CÔ":    {"ten": "Nghỉ con ốm", "nhom": "B", "cong": 0.0, "can_duyet": True},
    "TS":    {"ten": "Thai sản", "nhom": "B", "cong": 0.0, "can_duyet": True, "bao_giam_ngay": True},
    "KT":    {"ten": "Khám thai", "nhom": "B", "cong": 0.0, "can_duyet": True},
    "TN":    {"ten": "Tai nạn lao động", "nhom": "B", "cong": 0.0, "can_duyet": True, "canh_bao_bao_giam": 14},
    "DSOD":  {"ten": "Dưỡng sức sau ốm đau", "nhom": "B", "cong": 0.0, "can_duyet": True},
    "DSTS":  {"ten": "Dưỡng sức sau thai sản", "nhom": "B", "cong": 0.0, "can_duyet": True},
    "DSTN":  {"ten": "Dưỡng sức sau TNLĐ", "nhom": "B", "cong": 0.0, "can_duyet": True},
    "KL":    {"ten": "Nghỉ không lương (được duyệt)", "nhom": "C", "cong": 0.0, "can_duyet": True, "canh_bao_bao_giam": 14},
    "KP":    {"ten": "Nghỉ không phép", "nhom": "C", "cong": 0.0, "can_duyet": False, "ky_luat": True},
}

LOAI_TANG_CA = {
    "TC":  {"ten": "Tăng ca ngày thường", "he_so_mac_dinh": 1.5},
    "TCN": {"ten": "Tăng ca Chủ nhật", "he_so_mac_dinh": 2.0},
    "TCL": {"ten": "Tăng ca ngày lễ", "he_so_mac_dinh": 3.0},
    "TCĐ": {"ten": "Tăng ca đêm (cộng thêm)", "he_so_mac_dinh": 1.3},
}

# ===== PHƯƠNG ÁN ĐIỀU CHỈNH BHXH (dùng cho báo tăng/giảm D02-LT) =====
PHUONG_AN_TANG = [
    "TD - Tăng đến đã có số sổ, di chuyển trong địa bàn tỉnh",
    "TM - Tăng mới chưa có số sổ",
    "TC - Tăng chuyển đã có số sổ, di chuyển từ tỉnh khác đến",
]
PHUONG_AN_GIAM = [
    "GH - Giảm hẳn",
    "KL - Nghỉ không lương",
    "OF - Nghỉ ốm đau",
    "TS - Nghỉ thai sản",
]
PHUONG_AN_TANG_LAI = [
    "ON - Đi làm lại",
]
PHUONG_AN_ALL = PHUONG_AN_TANG + PHUONG_AN_GIAM + PHUONG_AN_TANG_LAI

def lay_ma_phuong_an(phuong_an_text):
    """Trích mã phương án từ text đầy đủ. VD: 'TD - Tăng đến...' → 'TD'"""
    if not phuong_an_text:
        return ""
    return phuong_an_text.split(" - ")[0].strip()

def tinh_thang_bat_dau_bh(ngay_vao_lam_hoac_chuyen, so_ngay_lam_viec_tuan=6):
    """Tính tháng bắt đầu đóng BHXH theo quy tắc 14 ngày (Khoản 5 Điều 33 Luật BHXH 2024).
    
    Logic: Nếu số ngày làm việc còn lại trong tháng (tính từ ngày vào làm đến cuối tháng) < 14 
    → tháng bắt đầu BH = tháng tiếp theo. Ngược lại → tháng hiện tại.
    
    Args:
        ngay_vao_lam_hoac_chuyen: date - ngày vào làm hoặc ngày QĐ chuyển đổi HĐ
        so_ngay_lam_viec_tuan: int - số ngày làm việc/tuần (mặc định 6, T2-T7)
    Returns:
        date - ngày 1 của tháng bắt đầu BH (kiểu DATE, hiển thị mm/yyyy)
    """
    if not ngay_vao_lam_hoac_chuyen:
        return None
    
    from calendar import monthrange
    
    ngay = ngay_vao_lam_hoac_chuyen
    _, so_ngay_trong_thang = monthrange(ngay.year, ngay.month)
    
    # Đếm số ngày làm việc còn lại từ ngày vào làm đến cuối tháng
    so_ngay_lv_con_lai = 0
    for d in range(ngay.day, so_ngay_trong_thang + 1):
        ngay_check = date(ngay.year, ngay.month, d)
        weekday = ngay_check.weekday()  # 0=T2, 6=CN
        
        if so_ngay_lam_viec_tuan == 6:
            # T2-T7 làm việc, CN nghỉ
            if weekday < 6:  # 0-5 = T2-T7
                so_ngay_lv_con_lai += 1
        elif so_ngay_lam_viec_tuan == 5:
            # T2-T6 làm việc, T7+CN nghỉ
            if weekday < 5:  # 0-4 = T2-T6
                so_ngay_lv_con_lai += 1
    
    if so_ngay_lv_con_lai < 14:
        # Tháng tiếp theo
        if ngay.month == 12:
            return date(ngay.year + 1, 1, 1)
        else:
            return date(ngay.year, ngay.month + 1, 1)
    else:
        # Tháng hiện tại
        return date(ngay.year, ngay.month, 1)

def format_thang_nam(d):
    """Hiển thị date thành mm/yyyy. VD: date(2026,8,1) → '08/2026'"""
    if not d:
        return ""
    if hasattr(d, 'strftime'):
        return d.strftime('%m/%Y')
    return str(d)

def tinh_ngay_ket_thuc(loai_hop_dong, ngay_vao_lam):
    """Tự động tính ngày kết thúc HĐ theo loại hợp đồng.
    - Thử việc: +2 tháng (tối đa theo Điều 25 BLLĐ 2019)
    - Xác định thời hạn: +12 tháng (mặc định, có thể sửa sau qua Cập nhật NV)
    - Không xác định thời hạn: None (không có ngày kết thúc)
    """
    if not ngay_vao_lam:
        return None
    if loai_hop_dong == "Thử việc":
        return ngay_vao_lam + relativedelta(months=2) - timedelta(days=1)
    elif loai_hop_dong == "Xác định thời hạn":
        return ngay_vao_lam + relativedelta(months=12) - timedelta(days=1)
    else:  # Không xác định thời hạn
        return None

CHAM_CONG_MA_OPTIONS = [""] + list(KY_HIEU_CHAM_CONG.keys())
KY_HIEU_CAN_PHE_DUYET = [ma for ma, tt in KY_HIEU_CHAM_CONG.items() if tt.get("can_duyet")]

# ========== HÀM TIỆN ÍCH MỚI ==========
def format_date_thang_nam(date_obj):
    """Định dạng ngày thành MM/YYYY"""
    if not date_obj:
        return ""
    if hasattr(date_obj, 'strftime'):
        return date_obj.strftime('%m/%Y')
    return str(date_obj)

def get_ma_tinh_from_name(tinh_name):
    """Lấy mã tỉnh từ tên tỉnh - Có thể mở rộng từ database"""
    if not tinh_name:
        return "44"  # Mặc định Quảng Trị
    
    # Map tên tỉnh -> mã tỉnh (có thể load từ database)
    ma_tinh_map = {
        'Quảng Trị': '44',
        'Quảng Bình': '43',
        'Thừa Thiên Huế': '45',
        'Huế': '45',
        'Đà Nẵng': '48',
        'Hà Nội': '01',
        'TP.HCM': '79',
        'Hồ Chí Minh': '79',
    }
    
    # Thử tìm kiếm trong map
    for key, value in ma_tinh_map.items():
        if key in tinh_name:
            return value
    
    return "44"  # Mặc định

def chuan_hoa_ten_phong_ban(ten):
    """Chuẩn hóa tên phòng ban theo danh mục chính thức PHONG_BAN_THU_TU (Title Case —
    viết hoa chữ cái đầu của MỌI từ, theo quy ước riêng của công ty).
    1) So khớp với PHONG_BAN_THU_TU, KHÔNG phân biệt hoa/thường, khoảng trắng thừa,
       và khoảng trắng quanh dấu '-' (VD dữ liệu cũ "Phòng KT-Cơ điện" hay
       "phòng kt -cơ điện" đều khớp với chuẩn "Phòng KT - Cơ Điện") -> trả về đúng
       tên chuẩn trong danh mục.
    2) Nếu không khớp (phòng ban lạ, chưa có trong danh mục) -> viết hoa chữ cái
       đầu mỗi từ (Title Case) để nhất quán với phong cách chung của danh mục.
    """
    if not ten:
        return ""
    ten_sach = " ".join(ten.strip().split())
    if not ten_sach:
        return ""

    def _so_sanh(s):
        # Chuẩn hóa khoảng trắng quanh dấu '-' về đúng 1 dạng trước khi so sánh,
        # để không bị lệch chuẩn chỉ vì cách gõ dấu gạch ngang khác nhau.
        return re.sub(r'\s*-\s*', ' - ', s).lower()

    for chuan in PHONG_BAN_THU_TU:
        if _so_sanh(ten_sach) == _so_sanh(chuan):
            return chuan
    return " ".join(w[0].upper() + w[1:] if w else w for w in ten_sach.split(" "))


def la_phong_ban_lanh_dao_cao_cap(ten):
    """So khớp phòng ban HĐQT/BTGĐ không phân biệt hoa/thường và khoảng trắng thừa,
    để không bị lệ thuộc vào cách viết hoa của dữ liệu cũ đã lưu trong DB."""
    ten_sach = " ".join((ten or "").strip().split()).lower()
    return any(ten_sach == pb.lower() for pb in PHONG_BAN_LANH_DAO_CAO_CAP)

def can_edit():
    """Kiểm tra xem user hiện tại có quyền chỉnh sửa dữ liệu không"""
    # Admin, HR, Văn thư, Kế toán lương có quyền chỉnh sửa.
    # LƯU Ý: 'xem_toan_bo' (Xem toàn bộ - không chỉnh sửa) CỐ Ý không có trong danh sách này —
    # vai trò này được thấy TOÀN BỘ menu/tab như Admin nhưng mọi nút Lưu/Sửa/Xóa/Cập nhật đều
    # phải bị disabled (xem các nơi dùng disabled=not can_edit()).
    edit_roles = ['admin', 'admin_bcc', 'hr', 'van_thu', 'kt_luong']
    return st.session_state.get('role') in edit_roles

def can_delete():
    """Kiểm tra xem user hiện tại có quyền xóa dữ liệu không"""
    # Chỉ Admin mới có quyền xóa ('xem_toan_bo' không có quyền xóa)
    return st.session_state.get('role') == 'admin'

def can_export():
    """Kiểm tra xem user hiện tại có quyền xuất báo cáo không"""
    # Admin, HR, Văn thư, Kế toán lương có quyền xuất
    export_roles = ['admin', 'hr', 'van_thu', 'kt_luong']
    return st.session_state.get('role') in export_roles

def can_manage_users():
    """Kiểm tra xem user hiện tại có quyền quản lý người dùng không"""
    # Chỉ Admin mới có quyền quản lý người dùng ('xem_toan_bo' không có quyền này)
    return st.session_state.get('role') == 'admin'

def can_edit_bcc():
    """Quyền nhập mới/sửa BCC (ô trống) — đọc cấu hình tenant."""
    mac_dinh = 'admin,admin_bcc'
    ds_role_str = get_cau_hinh('cc_vai_tro_edit_bcc', mac_dinh)
    ds_role = [r.strip() for r in ds_role_str.split(',') if r.strip()]
    return st.session_state.role in ds_role

def can_dieu_chinh_bcc():
    """Quyền điều chỉnh BCC — đọc cấu hình tenant để biết vai trò nào được phép."""
    # Mặc định: admin + admin_bcc
    # Tenant có thể cấu hình thêm vai trò khác (VD: 'kt_luong', 'truong_phong')
    mac_dinh = 'admin,admin_bcc'
    ds_role_str = get_cau_hinh('cc_vai_tro_dieu_chinh_bcc', mac_dinh)
    ds_role = [r.strip() for r in ds_role_str.split(',') if r.strip()]
    return st.session_state.role in ds_role

def can_khoa_thang_bcc():
    """Quyền khoá/mở khoá BCC tháng"""
    return st.session_state.role == 'admin'

def can_duyet_ot():
    """Quyền phê duyệt tăng ca"""
    return st.session_state.role in ('admin', 'admin_bcc', 'truong_phong')
    
def get_chu_ho_info(nhan_vien_id):
    """Lấy thông tin chủ hộ từ bảng phu_luc_gia_dinh"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("""
            SELECT ho_ten, so_cccd, dien_thoai 
            FROM phu_luc_gia_dinh 
            WHERE nhan_vien_id = %s AND (quan_he_voi_chu_ho = 'Chủ hộ' OR quan_he_voi_chu_ho = 'Chủ hộ gia đình')
            LIMIT 1
        """, (nhan_vien_id,))
        result = c.fetchone()
        db.close()
        return result
    except Exception as e:
        print(f"Lỗi lấy thông tin chủ hộ: {e}")
        return None

def tao_bao_cao_bhxh_d02_lt(tang_list, giam_list, tu_ngay, den_ngay, ten_cong_ty, ma_don_vi_bhxh):
    """Tạo báo cáo tăng/giảm BHXH mẫu D02-LT đúng 100% theo file mẫu chuẩn (130+ cột)"""
    
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = "D02-LT"
    
    # Định nghĩa border
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Định nghĩa tất cả các cột theo đúng thứ tự file mẫu (130+ cột)
    columns = [
        "STT", "Họ và tên", "Mã số BHXH", "Loại phương án", "Mã loại PA",
        "Loại ngày sinh", "Ngày Sinh", "Giới tính", "Số CMND/ CCCD/Hộ chiếu",
        "Cấp bậc, chức vụ, chức danh nghề", "Phòng ban làm việc", "Nơi Làm Việc",
        "Mức lương", "Phụ cấp lương", "Các khoản bổ sung", "Hệ số lương",
        "Phụ cấp CV", "Phụ cấp TNVK (%)", "Phụ cấp TN nghề (%)", "Phương án điều chỉnh",
        "Mã PA", "Tháng/ năm bắt đầu", "Tháng/ năm kết thúc",
        "Nghỉ ốm đau/Thai sản/không lương", "Ghi chú", "Số sổ BHXH",
        "Mức hưởng BHYT", "Tỷ lệ đóng (%)", "Mã vùng sinh sống", "Mã vùng lương tối thiểu",
        "Có giảm chết", "Ngày chết", "Tính lãi", "Nhóm vị trí việc làm",
        "Ngày bắt đầu giữ vị trí", "Ngày kết thúc giữ vị trí", "Hợp đồng lao động",
        "Hiệu lực từ ngày", "Hiệu lực đến ngày", "Ngày bắt đầu", "Ngày kết thúc",
        "Số", "Ngày ký", "Ngành nghề nặng nhọc, độc hại", "Ngày bắt đầu",
        "Ngày kết thúc", "Hợp đồng lao động", "Số", "Ngày ký", "Quốc tịch",
        "Mã QT", "Dân tộc", "Mã DT", "Điện thoại liên hệ", "Email liên hệ",
        "Tỉnh / Thành phố (Khai sinh)", "Mã Tỉnh (Khai sinh)", "Phường/ Xã (Khai sinh)",
        "Mã xã (Khai sinh)", "Địa chỉ khai sinh", "Tỉnh / Thành phố (Nhận HS)",
        "Mã Tỉnh (Nhận HS)", "Phường/ Xã (Nhận HS)", "Mã xã (Nhận HS)",
        "Địa chỉ nhận hồ sơ", "Tỉnh nơi KCB", "Mã tỉnh (KCB)", "Nơi đăng ký KCB",
        "Mã BV", "Đăng ký nhận sổ và thẻ", "Tỉnh / Thành phố (Nhận sổ thẻ)",
        "Mã Tỉnh (Nhận sổ thẻ)", "Phường/ Xã (Nhận sổ thẻ)", "Mã Xã (Nhận sổ thẻ)",
        "Địa chỉ nhận Sổ thẻ", "Mức tiền đóng", "Phương thức đóng", "Nội dung thay đổi",
        "Hồ sơ kèm theo", "Họ tên người giám hộ", "Mã số hộ gia đình",
        "Họ Tên chủ hộ", "Số CMND/ CCCD/Hộ chiếu (chủ hộ)", "Điện thoại (chủ hộ)",
        "Loại giấy tờ", "Số giấy tờ", "Tỉnh / Thành phố (hộ khẩu)", "Mã Tỉnh (hộ khẩu)",
        "Phường/ Xã (hộ khẩu)", "Mã xã (hộ khẩu)", "Tổ/ Thôn/ Xóm", "Địa chỉ hộ khẩu",
        "Tỉnh / Thành phố thường trú", "Mã Tỉnh (thường trú)", "Phường/ Xã thường trú",
        "Mã xã (thường trú)", "Địa chỉ thường trú", "Mã số hộ gia đình (PL)",
        "Họ và tên (PL)", "Mã số BHXH (PL)", "Loại ngày sinh (PL)", "Ngày sinh (PL)",
        "Giới tính (PL)", "Quốc tịch (PL)", "Mã Quốc tịch (PL)", "Dân tộc (PL)",
        "Mã Dân tộc (PL)", "Số CMND (PL)", "Mối quan hệ với chủ hộ", "Mã MQH",
        "Tỉnh / Thành phố (PL)", "Mã Tỉnh (PL)", "Phường/ Xã (PL)", "Mã xã (PL)",
        "Địa chỉ khai sinh (PL)", "Người tham gia", "Ghi chú (PL)"
    ]
    
    # Thiết lập độ rộng cột (điều chỉnh cho phù hợp)
    col_widths = [5, 25, 18, 15, 12, 12, 15, 10, 20, 25, 20, 25, 15, 15, 15, 12, 
                  15, 12, 12, 15, 12, 15, 15, 15, 20, 15, 15, 12, 12, 12, 12, 12, 
                  12, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 
                  15, 12, 12, 12, 12, 15, 20, 15, 12, 15, 12, 20, 15, 12, 15, 12, 
                  20, 15, 12, 15, 12, 15, 12, 15, 12, 20, 15, 15, 20, 15, 20, 20, 
                  15, 15, 15, 15, 15, 15, 15, 15, 15, 20, 15, 15, 15, 20, 20, 15, 
                  15, 15, 20, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 15, 20, 20]
    
    # Thiết lập độ rộng cột
    for idx, width in enumerate(col_widths[:len(columns)]):
        ws.column_dimensions[get_column_letter(idx + 1)].width = width
    
    # ===== HEADER =====
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(columns))
    ws['A1'] = ten_cong_ty
    ws['A1'].font = Font(bold=True, size=13, name='Times New Roman')
    ws['A1'].alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(columns))
    ws['A2'] = f"Mã đơn vị BHXH: {ma_don_vi_bhxh}"
    ws['A2'].font = Font(size=11, name='Times New Roman')
    ws['A2'].alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=len(columns))
    ws['A3'] = "DANH SÁCH LAO ĐỘNG THAM GIA BHXH, BHYT, BHTN, BHTNLĐ, BNN (Mẫu D02-LT TK1)"
    ws['A3'].font = Font(bold=True, size=12, name='Times New Roman')
    ws['A3'].alignment = Alignment(horizontal='center')
    
    ws.merge_cells(start_row=4, start_column=1, end_row=4, end_column=len(columns))
    ws['A4'] = f"Kỳ báo cáo: Tháng {tu_ngay.strftime('%m/%Y')} - {den_ngay.strftime('%m/%Y')}"
    ws['A4'].font = Font(size=11, name='Times New Roman')
    ws['A4'].alignment = Alignment(horizontal='center')
    
    # ===== HEADER BẢNG 2 DÒNG =====
    header_row_main = 6
    header_row_sub = 7
    
    # Tạo header 2 dòng phức tạp
    header_config = [
        (1, 1, "STT"), (2, 2, "Họ và tên"), (3, 3, "Mã số BHXH"),
        (4, 4, "Loại phương án"), (5, 5, "Mã loại PA"), (6, 6, "Loại ngày sinh"),
        (7, 7, "Ngày Sinh"), (8, 8, "Giới tính"), (9, 9, "Số CMND/ CCCD/Hộ chiếu"),
        (10, 10, "Cấp bậc, chức vụ, chức danh nghề"), (11, 11, "Phòng ban làm việc"),
        (12, 12, "Nơi Làm Việc"), (13, 17, "Tiền lương"), (18, 19, "Ngành nghề nặng nhọc, độc hại"),
        (20, 24, "Loại và hiệu lực hợp đồng"), (25, 25, "Thời điểm bắt đầu đóng BHXH"),
        (26, 26, "Thời điểm kết thúc đóng BHXH")
    ]
    
    # Dòng header chính
    for start_col, end_col, text in header_config:
        if start_col == end_col:
            cell = ws.cell(row=header_row_main, column=start_col, value=text)
        else:
            ws.merge_cells(start_row=header_row_main, start_column=start_col, 
                          end_row=header_row_main, end_column=end_col)
            cell = ws.cell(row=header_row_main, column=start_col, value=text)
        cell.font = Font(bold=True, size=10, name='Times New Roman')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        cell.fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
    
    # Dòng header phụ
    sub_headers = {
        13: "Mức lương", 14: "Phụ cấp lương", 15: "Các khoản bổ sung",
        16: "Hệ số lương", 17: "Phụ cấp CV", 18: "Phụ cấp TNVK (%)",
        19: "Phụ cấp TN nghề (%)", 20: "Loại HĐLĐ", 21: "Hiệu lực từ ngày",
        22: "Hiệu lực đến ngày", 23: "Ngày bắt đầu", 24: "Ngày kết thúc",
        25: "Thời điểm bắt đầu", 26: "Thời điểm kết thúc"
    }
    
    for col, value in sub_headers.items():
        cell = ws.cell(row=header_row_sub, column=col, value=value)
        cell.font = Font(bold=True, size=9, name='Times New Roman')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        cell.fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
    
    # Các cột đơn giản còn lại
    for col in range(1, len(columns) + 1):
        if col not in [1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26]:
            cell = ws.cell(row=header_row_sub, column=col, value=columns[col-1])
            cell.font = Font(bold=True, size=9, name='Times New Roman')
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border
            cell.fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
    
    # ===== DỮ LIỆU =====
    all_data = []
    
    # Xử lý dữ liệu tăng
    for nv in tang_list:
        row_data = {}
        row_data['STT'] = len(all_data) + 1
        row_data['Họ và tên'] = nv.get('ho_ten', '')
        row_data['Mã số BHXH'] = nv.get('ma_so_bhxh', '')
        row_data['Loại phương án'] = "Tăng lao động"
        row_data['Mã loại PA'] = "1"
        row_data['Loại ngày sinh'] = "0"
        row_data['Ngày Sinh'] = format_date(nv.get('ngay_sinh'))
        row_data['Giới tính'] = "1" if nv.get('gioi_tinh') == 'Nam' else "2" if nv.get('gioi_tinh') == 'Nữ' else "3"
        row_data['Số CMND/ CCCD/Hộ chiếu'] = nv.get('so_cccd', '')
        row_data['Cấp bậc, chức vụ, chức danh nghề'] = nv.get('chuc_danh_nghe', '')
        row_data['Phòng ban làm việc'] = nv.get('phong_ban_lam_viec', '')
        row_data['Nơi Làm Việc'] = nv.get('noi_lam_viec', 'Cảng THQT Hòn La')
        row_data['Mức lương'] = nv.get('luong_bao_hiem', '')
        row_data['Phụ cấp lương'] = ""
        row_data['Các khoản bổ sung'] = ""
        row_data['Hệ số lương'] = nv.get('he_so_luong', '')
        row_data['Phụ cấp CV'] = nv.get('phu_cap_chuc_vu', '')
        row_data['Phụ cấp TNVK (%)'] = nv.get('phu_cap_tnvk', '')
        row_data['Phụ cấp TN nghề (%)'] = nv.get('phu_cap_tnn', '')
        
        # Xác định phương án điều chỉnh — ưu tiên lấy từ cột phuong_an_dieu_chinh (đã lưu khi nhập NV)
        pa_da_luu = nv.get('phuong_an_dieu_chinh', '')
        if pa_da_luu:
            # Đã có phương án do user chọn hoặc app tự set
            row_data['Mã PA'] = pa_da_luu
            # Tìm label đầy đủ
            pa_label_map = {lay_ma_phuong_an(pa): pa for pa in PHUONG_AN_ALL}
            row_data['Phương án điều chỉnh'] = pa_label_map.get(pa_da_luu, pa_da_luu)
        elif nv.get('ma_so_bhxh'):
            row_data['Phương án điều chỉnh'] = "TD - Tăng đến đã có số sổ, di chuyển trong địa bàn tỉnh"
            row_data['Mã PA'] = "TD"
        else:
            row_data['Phương án điều chỉnh'] = "TM - Tăng mới chưa có số sổ"
            row_data['Mã PA'] = "TM"
        
        row_data['Tháng/ năm bắt đầu'] = format_date_thang_nam(nv.get('thang_bat_dau_bh')) if nv.get('thang_bat_dau_bh') else ""
        row_data['Tháng/ năm kết thúc'] = ""
        row_data['Nghỉ ốm đau/Thai sản/không lương'] = ""
        row_data['Ghi chú'] = nv.get('ghi_chu', '')
        row_data['Số sổ BHXH'] = nv.get('ma_so_bhxh', '')
        row_data['Mức hưởng BHYT'] = nv.get('muc_huong_bhyt', '100')
        row_data['Tỷ lệ đóng (%)'] = nv.get('ty_le_dong', '')
        row_data['Mã vùng sinh sống'] = ""
        row_data['Mã vùng lương tối thiểu'] = "03"
        row_data['Có giảm chết'] = ""
        row_data['Ngày chết'] = ""
        row_data['Tính lãi'] = ""
        row_data['Nhóm vị trí việc làm'] = ""
        row_data['Ngày bắt đầu giữ vị trí'] = format_date(nv.get('ngay_vao_lam'))
        row_data['Ngày kết thúc giữ vị trí'] = ""
        row_data['Loại HĐLĐ'] = nv.get('loai_hop_dong', '')
        row_data['Hiệu lực từ ngày'] = format_date(nv.get('ngay_ky_hd')) or format_date(nv.get('ngay_vao_lam'))
        row_data['Hiệu lực đến ngày'] = format_date(nv.get('ngay_ket_thuc')) if nv.get('ngay_ket_thuc') else ""
        row_data['Ngày bắt đầu'] = format_date(nv.get('thang_bat_dau_bh'))
        row_data['Ngày kết thúc'] = ""
        row_data['Số'] = nv.get('so_hdld', '')
        row_data['Ngày ký'] = format_date(nv.get('ngay_ky_hd')) or format_date(nv.get('ngay_vao_lam'))
        row_data['Quốc tịch'] = nv.get('quoc_tich', 'VIET NAM')
        row_data['Mã QT'] = "VN"
        row_data['Dân tộc'] = nv.get('dan_toc', 'Kinh')
        row_data['Mã DT'] = "1"
        row_data['Điện thoại liên hệ'] = nv.get('dien_thoai', '')
        row_data['Email liên hệ'] = nv.get('email_lien_he', '')
        
        # Thông tin địa chỉ
        row_data['Tỉnh / Thành phố (Khai sinh)'] = get_cau_hinh('tinh_khai_sinh', 'Tỉnh Quảng Trị')
        row_data['Mã Tỉnh (Khai sinh)'] = get_ma_tinh_from_name(get_cau_hinh('tinh_khai_sinh', 'Quảng Trị'))
        row_data['Phường/ Xã (Khai sinh)'] = nv.get('phuong_xa_khai_sinh', '')
        row_data['Mã xã (Khai sinh)'] = nv.get('ma_xa_khai_sinh', '')
        row_data['Địa chỉ khai sinh'] = nv.get('noi_sinh', '')
        
        row_data['Tỉnh / Thành phố (Nhận HS)'] = nv.get('tinh_nhan_hs') or get_cau_hinh('tinh_nhan_hs', 'Tỉnh Quảng Trị')
        row_data['Mã Tỉnh (Nhận HS)'] = get_ma_tinh_from_name(nv.get('tinh_nhan_hs') or get_cau_hinh('tinh_nhan_hs', 'Quảng Trị'))
        row_data['Phường/ Xã (Nhận HS)'] = nv.get('phuong_nhan_hs', '')
        row_data['Mã xã (Nhận HS)'] = nv.get('ma_xa_nhan_hs', '')
        row_data['Địa chỉ nhận hồ sơ'] = nv.get('dia_chi_nhan_hs', '')
        
        row_data['Tỉnh nơi KCB'] = nv.get('tinh_kcb') or get_cau_hinh('tinh_kcb', 'Tỉnh Quảng Trị')
        row_data['Mã tỉnh (KCB)'] = "44"
        row_data['Nơi đăng ký KCB'] = nv.get('noi_dang_ky_kcb') or get_cau_hinh('noi_dang_ky_kcb', 'Bệnh viện đa khoa khu vực Bắc Quảng Trị')
        row_data['Mã BV'] = "44003"
        
        row_data['Đăng ký nhận sổ và thẻ'] = nv.get('dang_ky_nhan_so', 'Có')
        row_data['Tỉnh / Thành phố (Nhận sổ thẻ)'] = nv.get('tinh_nhan_hs', '')
        row_data['Mã Tỉnh (Nhận sổ thẻ)'] = get_ma_tinh_from_name(nv.get('tinh_nhan_hs', ''))
        row_data['Phường/ Xã (Nhận sổ thẻ)'] = nv.get('phuong_nhan_hs', '')
        row_data['Mã Xã (Nhận sổ thẻ)'] = nv.get('ma_xa_nhan_hs', '')
        row_data['Địa chỉ nhận Sổ thẻ'] = nv.get('dia_chi_nhan_hs', '')
        
        row_data['Mức tiền đóng'] = nv.get('muc_tien_dong', '')
        row_data['Phương thức đóng'] = nv.get('phuong_thuc_dong', 'Hàng tháng')
        
        # Lấy thông tin chủ hộ
        chu_ho = get_chu_ho_info(nv.get('id'))
        if chu_ho:
            row_data['Họ Tên chủ hộ'] = chu_ho.get('ho_ten', '')
            row_data['Số CMND/ CCCD/Hộ chiếu (chủ hộ)'] = chu_ho.get('so_cccd', '')
            row_data['Điện thoại (chủ hộ)'] = chu_ho.get('dien_thoai', '')
        
        # Nếu chưa có mã BHXH, lấy thông tin phụ lục gia đình
        if not nv.get('ma_so_bhxh'):
            family_members = get_family_members(nv.get('id'))
            if family_members:
                first_member = family_members[0] if family_members else {}
                row_data['Họ và tên (PL)'] = first_member.get('ho_ten', '')
                row_data['Ngày sinh (PL)'] = format_date(first_member.get('ngay_sinh'))
                row_data['Giới tính (PL)'] = "1" if first_member.get('gioi_tinh') == 'Nam' else "2"
                row_data['Mối quan hệ với chủ hộ'] = first_member.get('quan_he', '')
        
        all_data.append(row_data)
    
    # Xử lý dữ liệu giảm
    for nv in giam_list:
        row_data = {}
        row_data['STT'] = len(all_data) + 1
        row_data['Họ và tên'] = nv.get('ho_ten', '')
        row_data['Mã số BHXH'] = nv.get('ma_so_bhxh', '')
        row_data['Loại phương án'] = "Giảm lao động"
        row_data['Mã loại PA'] = "2"
        row_data['Ngày Sinh'] = format_date(nv.get('ngay_sinh'))
        row_data['Giới tính'] = "1" if nv.get('gioi_tinh') == 'Nam' else "2"
        row_data['Tháng/ năm kết thúc'] = format_date_thang_nam(nv.get('thang_ket_thuc_bh')) if nv.get('thang_ket_thuc_bh') else ""
        row_data['Ghi chú'] = nv.get('ly_do_nghi', '')
        all_data.append(row_data)
    
    # Ghi dữ liệu vào Excel
    start_row = header_row_sub + 1
    for idx, row_data in enumerate(all_data):
        current_row = start_row + idx
        for col_idx, col_name in enumerate(columns, 1):
            value = row_data.get(col_name, '')
            cell = ws.cell(row=current_row, column=col_idx, value=value)
            cell.font = Font(size=10, name='Times New Roman')
            cell.border = thin_border
            # Căn giữa cho các cột số và mã
            if col_idx in [1, 4, 5, 6, 8, 21, 22, 27, 28, 29, 30, 31, 32, 33]:
                cell.alignment = Alignment(horizontal='center', vertical='center')
            else:
                cell.alignment = Alignment(horizontal='left', vertical='center')
    
    # Footer
    total_row = start_row + len(all_data)
    ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=5)
    ws.cell(row=total_row, column=1, value=f"Tổng số: {len(all_data)} lao động")
    ws.cell(row=total_row, column=1).font = Font(bold=True, size=11, name='Times New Roman')
    
    # Ký tên
    sign_row = total_row + 3
    ws.merge_cells(start_row=sign_row, start_column=len(columns)-3, end_row=sign_row, end_column=len(columns))
    ws.cell(row=sign_row, column=len(columns)-3, value="NGƯỜI LẬP BÁO CÁO")
    ws.cell(row=sign_row, column=len(columns)-3).font = Font(bold=True, size=11, name='Times New Roman')
    ws.cell(row=sign_row, column=len(columns)-3).alignment = Alignment(horizontal='center')
    
    sign_row += 1
    ws.merge_cells(start_row=sign_row, start_column=len(columns)-3, end_row=sign_row, end_column=len(columns))
    ws.cell(row=sign_row, column=len(columns)-3, value="(Ký, ghi rõ họ tên)")
    ws.cell(row=sign_row, column=len(columns)-3).font = Font(size=10, name='Times New Roman', italic=True)
    ws.cell(row=sign_row, column=len(columns)-3).alignment = Alignment(horizontal='center')
    
    sign_row += 1
    ws.merge_cells(start_row=sign_row, start_column=len(columns)-3, end_row=sign_row, end_column=len(columns))
    ws.cell(row=sign_row, column=len(columns)-3, value=COMPANY_CONFIG.get('dai_dien', 'GIÁM ĐỐC').upper())
    ws.cell(row=sign_row, column=len(columns)-3).font = Font(bold=True, size=11, name='Times New Roman')
    ws.cell(row=sign_row, column=len(columns)-3).alignment = Alignment(horizontal='center')
    
    # Lưu file
    filename = f"D02-LT_BHXH_{tu_ngay.strftime('%d%m%Y')}_{den_ngay.strftime('%d%m%Y')}.xlsx"
    wb.save(filename)
    return filename

def get_family_members(nhan_vien_id):
    """Lấy danh sách thành viên gia đình của nhân viên"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("""
            SELECT ho_ten, ngay_sinh, gioi_tinh, quan_he_voi_chu_ho as quan_he
            FROM phu_luc_gia_dinh 
            WHERE nhan_vien_id = %s
            ORDER BY id ASC
        """, (nhan_vien_id,))
        result = c.fetchall()
        db.close()
        return result
    except Exception as e:
        print(f"Lỗi lấy thông tin gia đình: {e}")
        return []


def format_date_thang_nam(date_obj):
    """Định dạng ngày thành MM/YYYY"""
    if not date_obj:
        return ""
    if hasattr(date_obj, 'strftime'):
        return date_obj.strftime('%m/%Y')
    return str(date_obj)
        
def remove_accents(text):
    """Bỏ dấu tiếng Việt, chuyển về chữ hoa không dấu"""
    if not text:
        return ""
    # "Đ/đ" là 1 ký tự Latin riêng (không phải tổ hợp base+dấu) nên
    # unicodedata.normalize('NFKD', ...) KHÔNG tách được gạch ngang của nó ra.
    # Nếu không thay thế thủ công trước, "Đ" sẽ bị loại bỏ hoàn toàn ở bước lọc
    # [^A-Z0-9] bên dưới (VD: "Đặng Thu Hà" -> "ANG THU HA" thay vì "DANG THU HA").
    text = text.replace('Đ', 'D').replace('đ', 'd')
    # Chuẩn hóa unicode và loại bỏ dấu
    text = unicodedata.normalize('NFKD', text)
    text = re.sub(r'[\u0300-\u036f]', '', text)
    # Chuyển thành chữ hoa, chỉ giữ chữ cái và số
    text = text.upper()
    text = re.sub(r'[^A-Z0-9\s]', '', text)
    return text.strip()

def generate_ten_don_vi_thu_huong(ho_ten):
    """Tạo tên đơn vị thụ hưởng từ họ tên (bỏ dấu, in hoa)"""
    return remove_accents(ho_ten)

# Load danh sách ngân hàng từ file Excel
def load_bank_list():
    """Đọc danh sách ngân hàng từ Bank_list.xlsx"""
    banks = []
    bank_file_path = os.path.join(os.path.dirname(__file__), "Bank_list.xlsx")
    
    if os.path.exists(bank_file_path):
        try:
            df_banks = pd.read_excel(bank_file_path, sheet_name=0)
            # Tìm cột chứa tên ngân hàng
            for col in df_banks.columns:
                if 'NGÂN' in col.upper() or 'BANK' in col.upper() or 'TÊN' in col.upper():
                    banks = df_banks[col].dropna().tolist()
                    break
            if not banks:
                banks = df_banks.iloc[:, 0].dropna().tolist()
        except Exception as e:
            print(f"Lỗi đọc Bank_list.xlsx: {e}")
            banks = []
    
    # Fallback: danh sách ngân hàng mặc định
    if not banks:
        banks = [
            "MB - Ngân hàng TMCP Quân Đội",
            "TCB - Ngân hàng TMCP Kỹ Thương Việt Nam",
            "ABBANK - Ngân hàng TMCP An Bình",
            "EIB - Ngân hàng TMCP Xuất Nhập Khẩu Việt Nam",
            "HDB - Ngân hàng TMCP Phát triển TP Hồ Chí Minh",
            "BVB - Ngân hàng TMCP Bảo Việt",
            "VAB - Ngân hàng TMCP Việt Á",
            "SEAB - Ngân hàng TMCP Đông Nam Á",
            "SCB - Ngân hang TMCP Sài Gòn",
            "NASB - Ngan hang TMCP Bac A",
            "VBA - Ngan hang Nong Nghiep va Phat Trien Nong Thon Viet Nam",
            "VCB - Ngân hàng TMCP Ngoại Thương Việt Nam",
            "BIDV - Ngân hàng TMCP Đầu Tư và Phát Triển Việt Nam",
            "VIETINBANK - Ngân hàng TMCP Công Thương Việt Nam",
            "ACB - Ngân hàng TMCP Á Châu",
            "VPB - Ngân hàng TMCP Việt Nam Thịnh Vượng",
            "STB - Ngân hàng TMCP Sài Gòn Thương Tín",
            "HDB - Ngân hàng TMCP Phát triển TP Hồ Chí Minh",
            "TPB - Ngân hàng TMCP Tiên Phong",
            "SHB - Ngân hàng TMCP Sài Gòn - Hà Nội",
            "MSB - Ngân hàng TMCP Hàng Hải Việt Nam",
            "VIB - Ngân hàng TMCP Quốc Tế",
            "OCB - Ngân hàng TMCP Phương Đông",
            "LPB - Ngân hàng TMCP Bưu Điện Liên Việt",
            "VIETBANK - Ngân hàng TMCP Việt Nam Thương Tín",
        ]
    return banks

BANK_LIST = load_bank_list()

# Danh mục Trình độ học vấn/chuyên môn (dùng cho form Thêm/Sửa nhân viên)
TRINH_DO_LIST = ["THPT", "Chứng chỉ nghề", "Cao đẳng", "Đại học", "Thạc sỹ", "Tiến sĩ"]

# Thứ tự ưu tiên CHUẨN — dùng thống nhất cho mọi biểu đồ / bảng / dropdown / tìm kiếm
# LƯU Ý: theo yêu cầu, toàn bộ tên phòng ban viết hoa CHỮ CÁI ĐẦU CỦA MỌI TỪ (Title Case),
# khác với quy tắc chính tả tiếng Việt thông thường — đây là lựa chọn có chủ đích cho danh mục
# chính thức của công ty, áp dụng thống nhất trong toàn bộ app.
PHONG_BAN_THU_TU = [
    "Hội Đồng Quản Trị",
    "Ban Tổng Giám Đốc",
    "Phòng Hành Chính Nhân Sự",
    "Phòng Kinh Doanh",
    "Phòng Tài Chính",
    "Phòng Điều Độ",
    "Tổ Cơ Giới",
    "Đội Bốc Xếp",
    "Phòng KT - Cơ Điện",
    "Đội Bảo Vệ",  # giữ lại từ danh mục cũ (không có trong yêu cầu mới) — xóa dòng này nếu không còn dùng
]

CHUC_VU_THU_TU = [
    "Chủ tịch HĐQT",
    "Phó Chủ tịch HĐQT",
    "Thành viên HĐQT",
    "Tổng Giám đốc",
    "Phó Tổng Giám đốc",
    "Chánh Văn phòng",
    "Phó Chánh VP",
    "Trưởng phòng",
    "Phó Trưởng phòng",
]

def sap_xep_phong_ban(danh_sach_phong_ban):
    """Sắp xếp tên phòng ban theo thứ tự ưu tiên chuẩn; phòng ban lạ xếp cuối theo alpha bê."""
    def key_fn(ten):
        try:
            return (0, PHONG_BAN_THU_TU.index(ten))
        except ValueError:
            return (1, ten or "")
    return sorted(danh_sach_phong_ban, key=key_fn)

def sap_xep_nhan_vien(ds_nv):
    """Sắp xếp list nhân viên (dict/RealDictRow) theo: Phòng ban -> Chức vụ -> Tên (alpha bê)."""
    def key_fn(nv):
        pb = nv.get('phong_ban_lam_viec') or ''
        try:
            pb_key = (0, PHONG_BAN_THU_TU.index(pb))
        except ValueError:
            pb_key = (1, pb)
        cv = nv.get('chuc_vu') or 'Nhân viên'
        try:
            cv_key = (0, CHUC_VU_THU_TU.index(cv))
        except ValueError:
            cv_key = (1, 0)
        return (pb_key, cv_key, nv.get('ho_ten') or '')
    return sorted(ds_nv, key=key_fn)

def sap_xep_phong_ban_rows(rows, key_field="Phòng ban"):
    """Sắp xếp list các dict/RealDictRow theo thứ tự phòng ban chuẩn (PHONG_BAN_THU_TU)."""
    def key_fn(row):
        ten = row.get(key_field) or ""
        try:
            return (0, PHONG_BAN_THU_TU.index(ten))
        except ValueError:
            return (1, ten)
    return sorted(rows, key=key_fn)

def get_phong_ban_options():
    """Trả về danh sách phòng ban cho dropdown, ƯU TIÊN theo danh mục riêng
    của TỪNG TENANT (bảng danh_muc_phong_ban trong DB của chính tenant đó,
    nhập qua màn "⚙️ Danh mục").

    LƯU Ý LỊCH SỬ: bản trước đây khóa cứng toàn bộ app (mọi tenant) theo
    hằng số PHONG_BAN_THU_TU (danh mục riêng của Hòn La) — dẫn tới việc các
    tenant khác (VD: DEMO-HRM) tự nhập danh mục phòng ban của mình nhưng
    không được app sử dụng. Nay sửa lại: đọc từ DB của tenant đang đăng nhập
    trước, nếu tenant đó CHƯA nhập danh mục riêng (bảng rỗng) thì mới dùng
    PHONG_BAN_THU_TU làm mặc định (đảm bảo Hòn La không bị ảnh hưởng ngược).
    """
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT ten_phong_ban FROM danh_muc_phong_ban ORDER BY thu_tu, id")
        rows = [r[0] for r in c.fetchall() if r[0]]
        db.close()
        if rows:
            return rows
    except Exception:
        pass
    # Fallback: tenant chưa tự nhập danh mục riêng -> dùng danh mục mặc định
    return list(PHONG_BAN_THU_TU)

def get_chuc_vu_options():
    """Trả về danh sách Chức vụ (dùng cho dropdown 'Chức vụ được bổ nhiệm/miễn nhiệm'
    trong tab Quyết định nhân sự), ƯU TIÊN theo danh mục riêng của TỪNG TENANT
    (bảng chuc_vu_danh_muc, nhập qua màn ⚙️ Danh mục > Chức vụ).

    LƯU Ý: trước đây khoá cứng theo hằng số DANH_SACH_CHUC_VU (danh mục riêng
    của Hòn La), khiến các tenant khác đã tự nhập danh mục Chức vụ riêng nhưng
    không được dropdown này sử dụng. Nay đọc từ DB tenant trước, chỉ fallback
    về DANH_SACH_CHUC_VU khi tenant chưa nhập danh mục riêng (bảng rỗng)."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS chuc_vu_danh_muc (
                id SERIAL PRIMARY KEY,
                ten_chuc_vu VARCHAR(150) UNIQUE NOT NULL,
                thu_tu INT DEFAULT 0,
                trang_thai VARCHAR(20) DEFAULT 'Hoạt động'
            )
        """)
        db.commit()
        c.execute("""
            SELECT ten_chuc_vu FROM chuc_vu_danh_muc
            WHERE trang_thai = 'Hoạt động'
            ORDER BY thu_tu, id
        """)
        rows = [r[0] for r in c.fetchall() if r[0]]
        db.close()
        if rows:
            return rows
    except Exception:
        pass
    # Fallback: tenant chưa tự nhập danh mục riêng -> dùng danh mục mặc định
    return list(DANH_SACH_CHUC_VU)


def get_phong_to_chuc_nhan_su():
    """Trả về tên Phòng Tổ chức - Nhân sự / Hành chính ĐÚNG theo danh mục phòng ban
    riêng của TỪNG TENANT, thay vì khoá cứng chuỗi 'Phòng Tổ chức - Hành chính'
    (tên riêng của Hòn La — các tenant khác có thể đặt tên khác như 'Phòng Nhân sự',
    'Phòng Hành chính - Nhân sự', 'Văn phòng'...).

    Duyệt danh mục phòng ban của tenant (get_phong_ban_options), tìm phòng có tên
    chứa 1 trong các từ khoá sau, ƯU TIÊN theo đúng thứ tự: 'tổ chức' > 'nhân sự'
    > 'hành chính' > 'văn phòng' (không phân biệt hoa/thường, có/không dấu).
    Nếu không tìm được phòng nào khớp -> fallback về tên mặc định cũ để không phá
    vỡ hành vi hiện tại của các tenant chưa cấu hình danh mục phòng ban.
    """
    TU_KHOA_UU_TIEN = ['tổ chức', 'nhân sự', 'hành chính', 'văn phòng']
    try:
        danh_sach = get_phong_ban_options() or []
        # So khớp không dấu để không phụ thuộc cách gõ dấu của từng tenant
        danh_sach_khong_dau = [(ten, remove_accents(ten).lower()) for ten in danh_sach if ten]
        for tu_khoa in TU_KHOA_UU_TIEN:
            tu_khoa_khong_dau = remove_accents(tu_khoa).lower()
            for ten_goc, ten_khong_dau in danh_sach_khong_dau:
                if tu_khoa_khong_dau in ten_khong_dau:
                    return ten_goc
    except Exception:
        pass
    # Fallback: tenant chưa có danh mục phòng ban khớp từ khoá nào -> giữ tên mặc định cũ
    return 'Phòng Tổ chức - Hành chính'

# ============================================================
# 🤖 CHATBOT GIẢI ĐÁP — AI Tư vấn Hành chính Nhân sự
# (chuyển thể từ bản HTML/JS gốc sang Python để chạy trong Streamlit,
#  gọi Anthropic API từ phía server bằng API key lưu trong st.secrets)
# ============================================================
# --- Thư mục chứa dữ liệu luật đã convert sẵn (*_full.json) ---
# Copy toàn bộ file "..._full.json" (BHXH, BHYT, Nghị định, Thông tư, Luật Thuế TNCN,
# Luật Việc làm...) vào thư mục này. Có luật mới -> chỉ cần thả file JSON vào đây,
# không cần sửa code, chỉ cần bấm nút "🔄 Nạp lại dữ liệu luật" trong màn Chatbot.
CHATBOT_LAW_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "luat_data")

_CHATBOT_TRONG_SO_MUC_DO = {"high": 1.3, "medium": 1.0, "low": 0.7}

@st.cache_resource(show_spinner=False)
def _chatbot_load_all_laws():
    """Đọc toàn bộ file *_full.json trong CHATBOT_LAW_DIR và gộp thành 1 danh sách
    các 'điều luật' dùng để tra cứu cho Chatbot. Mỗi phần tử JSON gốc kỳ vọng có dạng:
    {"content": "...", "law_name": "...", "article": "...", "clause": None,
     "metadata": {"category": "...", "importance": "high|medium|low", "keywords": [...]}}
    Hàm được viết dạng "khoan dung" (dùng .get) để không vỡ nếu vài file có field hơi khác.
    Cache bằng st.cache_resource -> chỉ đọc đĩa 1 lần, gọi _chatbot_load_all_laws.clear()
    để nạp lại khi có file mới."""
    ket_qua = []
    if not os.path.isdir(CHATBOT_LAW_DIR):
        return ket_qua

    files = sorted(
        f for f in os.listdir(CHATBOT_LAW_DIR)
        if f.lower().endswith(".json")
    )
    counter = 0
    for filename in files:
        fp = os.path.join(CHATBOT_LAW_DIR, filename)
        try:
            with open(fp, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            print(f"⚠️ Chatbot: lỗi đọc {filename}: {e}")
            continue

        # Chấp nhận cả 2 dạng: list các điều luật, hoặc dict bọc ngoài {"items": [...]}
        if isinstance(data, dict):
            data = data.get("items") or data.get("data") or data.get("articles") or [data]
        if not isinstance(data, list):
            continue

        for item in data:
            if not isinstance(item, dict):
                continue
            content = (item.get("content") or item.get("text") or "").strip()
            if not content:
                continue

            law_name = item.get("law_name") or item.get("ten_luat") or os.path.splitext(filename)[0]
            dieu = item.get("article") or item.get("dieu") or ""
            khoan = item.get("clause") or item.get("khoan")
            meta = item.get("metadata") or {}

            ref = f"{law_name}"
            if dieu:
                ref += f", Điều {dieu}"
            if khoan:
                ref += f", Khoản {khoan}"

            counter += 1
            ket_qua.append({
                "id": f"R{counter:05d}",
                "ref": ref,
                "text": content,
                "law_name": law_name,
                "category": meta.get("category", ""),
                "importance": meta.get("importance", "medium"),
                "keywords": meta.get("keywords") or [],
                "source_file": filename,
            })
    return ket_qua

def _chatbot_search_laws_keyword(q, top_k=12):
    """[DỰ PHÒNG] Tìm điều luật liên quan bằng so khớp từ khoá thuần (không cần gọi API embedding).
    Được dùng làm phương án dự phòng khi chưa cấu hình VOYAGE_API_KEY, hoặc khi gọi API
    embedding bị lỗi (mất mạng, hết quota...) — để Chatbot không bao giờ "đứng hình" hoàn toàn.
    Chấm điểm theo số từ khoá trùng khớp trong content/keywords/law_name, có nhân trọng số
    theo 'importance' để ưu tiên các điều luật quan trọng khi điểm số ngang nhau."""
    laws = _chatbot_load_all_laws()
    if not laws:
        return []

    t = q.lower()
    tu_khoa_cau_hoi = set(re.findall(
        r"[a-zàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]+", t
    ))
    tu_khoa_cau_hoi = {w for w in tu_khoa_cau_hoi if len(w) > 2}

    cham_diem = []
    for law in laws:
        diem = 0.0
        noi_dung_ts = law["text"].lower()
        law_name_ts = law["law_name"].lower()

        for kw in law.get("keywords", []):
            if kw and kw.lower() in t:
                diem += 5

        for w in tu_khoa_cau_hoi:
            if w in noi_dung_ts:
                diem += 1
            if w in law_name_ts:
                diem += 1

        if diem > 0:
            diem *= _CHATBOT_TRONG_SO_MUC_DO.get(law.get("importance", "medium"), 1.0)
            cham_diem.append((diem, law))

    cham_diem.sort(key=lambda x: x[0], reverse=True)
    ket_qua = [law for _, law in cham_diem[:top_k]]

    if not ket_qua:
        # Không khớp từ khoá nào -> vẫn đưa cho AI một ít ngữ cảnh (ưu tiên điều luật quan trọng)
        uu_tien = [l for l in laws if l.get("importance") == "high"]
        ket_qua = (uu_tien or laws)[:top_k]

    return ket_qua

def _chatbot_all_laws():
    """Toàn bộ điều luật đã nạp — dùng để đối chiếu khi hiển thị phần 'Căn cứ pháp lý'."""
    return _chatbot_load_all_laws()

# ============================================================
# 🔍 SEMANTIC SEARCH — nhúng (embedding) từng điều luật thành vector bằng Voyage AI
# (Voyage AI là nhà cung cấp embedding được Anthropic khuyến nghị dùng cho các ứng dụng
# RAG/semantic search). Việc này giúp tìm đúng điều luật liên quan kể cả khi người dùng
# hỏi bằng từ đồng nghĩa (VD hỏi "nghỉ đẻ" vẫn tìm ra các điều nói về "thai sản"),
# thay vì chỉ so khớp từ khoá chính xác như bản cũ.
# ============================================================
CHATBOT_EMBED_MODEL = "voyage-4-lite"   # rẻ, đa ngôn ngữ, 200 triệu token đầu tiên miễn phí/tài khoản
CHATBOT_EMBED_CACHE_FILE = os.path.join(CHATBOT_LAW_DIR, ".embeddings_cache.json")
CHATBOT_EMBED_BATCH_SIZE = 100  # số đoạn văn bản nhúng mỗi lần gọi API

# ---------- CẤU HÌNH THANH TOÁN CHATBOT ----------
CHATBOT_PAYMENT = {
    "bank_id": "970436",                    # Vietcombank (mã BIN VietQR)
    "stk": "0101001101757",
    "chu_tk": "NGUYEN VAN TUYEN",
    "so_tien": 20000,
    "credit_moi": 5,                        # số câu hỏi cho lần đăng ký đầu
    "credit_mua_them": 10,                  # số câu khi mua thêm
    "gia_mua_them": 30000,
    "admin_email": "duhocanphuloc@gmail.com",
}
 
import random, string as _string, urllib.parse as _urlparse
 
def _chatbot_ensure_table():
    """Tạo bảng chatbot_dang_ky nếu chưa có (gọi 1 lần mỗi session)."""
    if st.session_state.get('_chatbot_table_ok'):
        return
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS chatbot_dang_ky (
                id              SERIAL PRIMARY KEY,
                ma_dang_ky      TEXT UNIQUE NOT NULL,
                ho_ten          TEXT NOT NULL,
                email           TEXT NOT NULL,
                dien_thoai      TEXT NOT NULL,
                cong_ty         TEXT,
                so_credit       INT DEFAULT 0,
                da_dung         INT DEFAULT 0,
                trang_thai      TEXT DEFAULT 'CHO_THANH_TOAN',
                anh_bill        TEXT,
                ghi_chu_admin   TEXT,
                created_at      TIMESTAMP DEFAULT NOW(),
                duyet_luc       TIMESTAMP,
                duyet_boi       TEXT
            )
        """)
        # Index chống trùng email/SĐT
        c.execute("""
            CREATE UNIQUE INDEX IF NOT EXISTS idx_chatbot_dk_email
            ON chatbot_dang_ky (LOWER(email))
        """)
        c.execute("""
            CREATE UNIQUE INDEX IF NOT EXISTS idx_chatbot_dk_sdt
            ON chatbot_dang_ky (dien_thoai)
            WHERE dien_thoai IS NOT NULL AND dien_thoai != ''
        """)
        db.commit()
        db.close()
        st.session_state['_chatbot_table_ok'] = True
    except Exception:
        pass
 
def _chatbot_sinh_ma():
    """Sinh mã đăng ký duy nhất, VD: TVHCNS-A7K3M2"""
    chars = _string.ascii_uppercase + _string.digits
    return "TVHCNS-" + ''.join(random.choices(chars, k=6))
 
def _chatbot_qr_url(ma_dang_ky, so_tien=None):
    """Tạo URL ảnh QR thanh toán chuẩn VietQR (không cần API key)."""
    cfg = CHATBOT_PAYMENT
    amt = so_tien or cfg["so_tien"]
    ten_encoded = _urlparse.quote(cfg["chu_tk"])
    return (
        f"https://img.vietqr.io/image/{cfg['bank_id']}-{cfg['stk']}-compact2.png"
        f"?amount={amt}&addInfo={_urlparse.quote(ma_dang_ky)}&accountName={ten_encoded}"
    )
 
def _chatbot_tim_dang_ky(email_or_sdt):
    """Tìm bản đăng ký theo email hoặc SĐT. Trả về dict hoặc None."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("""
            SELECT * FROM chatbot_dang_ky
            WHERE LOWER(email) = LOWER(%s) OR dien_thoai = %s
            ORDER BY id DESC LIMIT 1
        """, (email_or_sdt.strip(), email_or_sdt.strip()))
        row = c.fetchone()
        db.close()
        return dict(row) if row else None
    except Exception:
        return None
 
def _chatbot_tao_dang_ky(ho_ten, email, sdt, cong_ty=""):
    """Tạo bản đăng ký mới. Trả về ma_dang_ky hoặc None nếu lỗi."""
    ma = _chatbot_sinh_ma()
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            INSERT INTO chatbot_dang_ky (ma_dang_ky, ho_ten, email, dien_thoai, cong_ty)
            VALUES (%s, %s, %s, %s, %s)
        """, (ma, ho_ten.strip(), email.strip().lower(), sdt.strip(), cong_ty.strip()))
        db.commit()
        db.close()
        return ma
    except Exception as e:
        if "idx_chatbot_dk_email" in str(e):
            return "TRUNG_EMAIL"
        if "idx_chatbot_dk_sdt" in str(e):
            return "TRUNG_SDT"
        return None
 
def _chatbot_upload_bill(ma_dang_ky, uploaded_file):
    """Upload ảnh bill CK lên Supabase Storage, cập nhật trạng thái."""
    try:
        sb = get_supabase_storage()
        if not sb:
            return False
        safe_name = sanitize_storage_filename(uploaded_file.name)
        path = f"chatbot_bills/{ma_dang_ky}/{safe_name}"
        storage_path = upload_to_storage_unique(
            sb, SUPABASE_BUCKET, path,
            uploaded_file.getvalue(), uploaded_file.type
        )
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            UPDATE chatbot_dang_ky
            SET anh_bill = %s, trang_thai = 'DA_GUI_BILL'
            WHERE ma_dang_ky = %s
        """, (storage_path, ma_dang_ky))
        db.commit()
        db.close()
        return True
    except Exception:
        return False
 
def _chatbot_tru_credit(reg_id):
    """Trừ 1 credit sau khi user gửi câu hỏi."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            UPDATE chatbot_dang_ky SET da_dung = da_dung + 1
            WHERE id = %s AND da_dung < so_credit
        """, (reg_id,))
        db.commit()
        db.close()
    except Exception:
        pass
 
def _chatbot_gui_thong_bao_admin(dk):
    """Gửi email thông báo cho admin khi có đăng ký mới."""
    try:
        subject = f"[HRM Master] Đăng ký Chatbot mới: {dk['ho_ten']}"
        body = f"""
        <h3>Đăng ký thử nghiệm AI Tư vấn HCNS</h3>
        <table border="1" cellpadding="6" style="border-collapse:collapse">
            <tr><td><b>Mã ĐK</b></td><td>{dk['ma_dang_ky']}</td></tr>
            <tr><td><b>Họ tên</b></td><td>{dk['ho_ten']}</td></tr>
            <tr><td><b>Email</b></td><td>{dk['email']}</td></tr>
            <tr><td><b>SĐT</b></td><td>{dk['dien_thoai']}</td></tr>
            <tr><td><b>Công ty</b></td><td>{dk.get('cong_ty','')}</td></tr>
        </table>
        <p>Vào app → Chatbot Giải đáp → tab "📋 Quản lý đăng ký" để duyệt.</p>
        """
        gui_email_don(CHATBOT_PAYMENT["admin_email"], subject, body)
    except Exception:
        pass

def _chatbot_get_voyage_api_key():
    try:
        return st.secrets.get("VOYAGE_API_KEY") or st.secrets.get("voyage", {}).get("api_key")
    except Exception:
        return None

def _chatbot_hash_text(text):
    import hashlib
    return hashlib.md5(text.encode("utf-8")).hexdigest()

def _chatbot_embed_texts(texts, input_type):
    """Gọi Voyage API để nhúng 1 danh sách văn bản thành vector.
    input_type: "document" khi nhúng các điều luật, "query" khi nhúng câu hỏi của người dùng
    (Voyage khuyến nghị phân biệt 2 loại này để tăng chất lượng tìm kiếm).
    Trả về list các vector (list[float]), hoặc None nếu lỗi (thiếu API key, mất mạng...)."""
    api_key = _chatbot_get_voyage_api_key()
    if not api_key or not texts:
        return None
    vectors = []
    try:
        for i in range(0, len(texts), CHATBOT_EMBED_BATCH_SIZE):
            batch = texts[i:i + CHATBOT_EMBED_BATCH_SIZE]
            resp = requests.post(
                "https://api.voyageai.com/v1/embeddings",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "content-type": "application/json",
                },
                json={"input": batch, "model": CHATBOT_EMBED_MODEL, "input_type": input_type},
                timeout=30,
            )
            data = resp.json()
            if "data" not in data:
                print(f"⚠️ Chatbot: lỗi Voyage API: {data.get('error', data)}")
                return None
            vectors.extend([d["embedding"] for d in data["data"]])
        return vectors
    except Exception as e:
        print(f"⚠️ Chatbot: lỗi kết nối Voyage API: {e}")
        return None

@st.cache_resource(show_spinner=False)
def _chatbot_build_law_embeddings():
    """Nhúng toàn bộ điều luật thành vector, có cache trên đĩa (CHATBOT_EMBED_CACHE_FILE)
    theo hash nội dung -> chỉ nhúng MỚI những điều luật chưa có hoặc đã đổi nội dung,
    không nhúng lại toàn bộ mỗi lần app khởi động (tiết kiệm chi phí + thời gian chờ).
    Trả về dict {id: vector} hoặc {} nếu chưa cấu hình VOYAGE_API_KEY / có lỗi."""
    laws = _chatbot_load_all_laws()
    if not laws:
        return {}

    cache = {}
    if os.path.isfile(CHATBOT_EMBED_CACHE_FILE):
        try:
            with open(CHATBOT_EMBED_CACHE_FILE, "r", encoding="utf-8") as f:
                cache = json.load(f)
        except Exception:
            cache = {}

    can_nhung = []      # danh sách law cần nhúng mới
    vector_theo_id = {}
    for law in laws:
        h = _chatbot_hash_text(law["text"])
        cached_entry = cache.get(law["id"])
        if cached_entry and cached_entry.get("hash") == h:
            vector_theo_id[law["id"]] = cached_entry["vector"]
        else:
            can_nhung.append(law)

    if can_nhung:
        vectors_moi = _chatbot_embed_texts([l["text"] for l in can_nhung], input_type="document")
        if vectors_moi:
            for law, vec in zip(can_nhung, vectors_moi):
                h = _chatbot_hash_text(law["text"])
                vector_theo_id[law["id"]] = vec
                cache[law["id"]] = {"hash": h, "vector": vec}
            # Lưu lại cache lên đĩa để lần sau (kể cả app restart) không phải nhúng lại
            try:
                with open(CHATBOT_EMBED_CACHE_FILE, "w", encoding="utf-8") as f:
                    json.dump(cache, f)
            except Exception as e:
                print(f"⚠️ Chatbot: không lưu được cache embedding: {e}")

    return vector_theo_id

def _chatbot_cosine(a, b):
    # Voyage embeddings đã được chuẩn hoá độ dài = 1 sẵn -> dot product = cosine similarity,
    # không cần chia cho norm nữa (tính nhanh hơn).
    return sum(x * y for x, y in zip(a, b))

def _chatbot_search_laws_semantic(q, top_k=12):
    """Tìm điều luật liên quan bằng semantic search (so khớp theo Ý NGHĨA câu hỏi qua vector,
    thay vì chỉ so khớp từ khoá chính xác). Có kết hợp thêm điểm thưởng nhỏ nếu khớp từ khoá/
    số điều luật chính xác (hybrid) để không bỏ sót khi người hỏi nêu đúng số Điều/thuật ngữ
    pháp lý cụ thể — embedding thuần đôi khi không phân biệt tốt các con số."""
    laws = _chatbot_load_all_laws()
    vector_theo_id = _chatbot_build_law_embeddings()
    if not laws or not vector_theo_id:
        return None  # báo hiệu để hàm gọi rơi về fallback từ khoá

    q_vec_list = _chatbot_embed_texts([q], input_type="query")
    if not q_vec_list:
        return None
    q_vec = q_vec_list[0]

    t = q.lower()
    diem = []
    for law in laws:
        vec = vector_theo_id.get(law["id"])
        if not vec:
            continue
        s = _chatbot_cosine(q_vec, vec)
        # Thưởng nhẹ nếu khớp đúng từ khoá đã gắn sẵn cho điều luật (hybrid với keyword)
        for kw in law.get("keywords", []):
            if kw and kw.lower() in t:
                s += 0.03
        diem.append((s, law))

    diem.sort(key=lambda x: x[0], reverse=True)
    return [law for _, law in diem[:top_k]]

def _chatbot_search_laws(q, top_k=12):
    """Điểm vào chính để tìm điều luật liên quan tới câu hỏi. Ưu tiên semantic search
    (Voyage embeddings); nếu chưa cấu hình VOYAGE_API_KEY hoặc gọi API lỗi thì tự động
    rơi về so khớp từ khoá (_chatbot_search_laws_keyword) để Chatbot luôn hoạt động được."""
    ket_qua = _chatbot_search_laws_semantic(q, top_k=top_k)
    if ket_qua is None:
        ket_qua = _chatbot_search_laws_keyword(q, top_k=top_k)
    return ket_qua

def _chatbot_system_prompt_v2(laws):
    """System prompt mới — yêu cầu Claude trả lời Markdown (không JSON)."""
    laws_text = "\n".join(
        f'[{l["id"]}] {l["ref"]}: "{l["text"]}"' for l in laws
    )
    return f"""Bạn là CHUYÊN GIA TƯ VẤN pháp luật Hành chính - Nhân sự Việt Nam, có 15+ năm kinh nghiệm thực tiễn.
 
NHIỆM VỤ: Tư vấn chuyên nghiệp, cụ thể, CÓ CĂN CỨ PHÁP LÝ cho câu hỏi của người dùng.
 
CÁC ĐIỀU LUẬT LIÊN QUAN ĐÃ TRUY XUẤT TỪ CƠ SỞ DỮ LIỆU:
{laws_text}
 
QUY TẮC TRẢ LỜI:
1. Trả lời bằng tiếng Việt, dùng Markdown, cấu trúc rõ ràng theo các mục bên dưới.
2. LUÔN viện dẫn điều luật cụ thể (số điều, tên luật/nghị định/thông tư) ngay trong câu phân tích.
3. Nếu có số liệu cụ thể (lương, thời gian đóng BH...) → tính toán chi tiết, trình bày công thức.
4. Nếu câu hỏi mơ hồ → nêu các trường hợp có thể xảy ra, mỗi trường hợp kèm căn cứ riêng.
5. Giọng văn: chuyên nghiệp nhưng dễ hiểu, tránh hàn lâm.
 
ĐỊNH DẠNG BẮT BUỘC:
 
### 📋 Tóm tắt
(1-2 câu trả lời ngắn gọn, đi thẳng vào vấn đề)
 
### 📖 Phân tích chi tiết
(Phân tích đầy đủ, viện dẫn điều luật cụ thể ngay trong từng luận điểm.
Ví dụ: "Theo Điều 46 Bộ luật Lao động 2019, trợ cấp thôi việc được tính...")
 
### 🔢 Tính toán (nếu có số liệu)
(Trình bày công thức + kết quả cụ thể. Bỏ qua mục này nếu câu hỏi không liên quan đến tính toán)
 
### ⚖️ Căn cứ pháp lý
(Liệt kê các điều luật đã viện dẫn, mỗi điều 1 dòng, gồm: tên luật + số điều + nội dung tóm tắt)
 
### 💡 Lưu ý
(Cảnh báo rủi ro, trường hợp ngoại lệ, hoặc thông tin cần bổ sung để tư vấn chính xác hơn)"""

def _chatbot_get_api_key():
    try:
        return st.secrets.get("ANTHROPIC_API_KEY") or st.secrets.get("anthropic", {}).get("api_key")
    except Exception:
        return None

def _chatbot_call_claude_v2(system_prompt, history):
    """Gọi Claude API, trả về text Markdown (không parse JSON)."""
    api_key = _chatbot_get_api_key()
    if not api_key:
        return "⚠️ **Chưa cấu hình ANTHROPIC_API_KEY.** Vào Streamlit Secrets, thêm: `ANTHROPIC_API_KEY = \"sk-ant-...\"`"
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": "claude-sonnet-4-5-20250929",
                "max_tokens": 3500,
                "system": system_prompt,
                "messages": history,
            },
            timeout=60,
        )
        data = resp.json()
        if "content" not in data:
            err = data.get("error", {}).get("message", str(data))
            return f"❌ **Lỗi API:** {err}"
        return "".join(b.get("text", "") for b in data["content"])
    except Exception as e:
        return f"❌ **Lỗi kết nối:** {e}"

def _chatbot_badge_html(loai):
    return {
        "recommended": '<span style="font-size:10px;padding:2px 8px;border-radius:20px;background:#e6f4ea;color:#1e6e3a;border:1px solid #8fc8a3;font-weight:600">Khuyến nghị</span>',
        "alternative": '<span style="font-size:10px;padding:2px 8px;border-radius:20px;background:#e8f0fe;color:#1a56c4;border:1px solid #93b4f5;font-weight:600">Thay thế</span>',
        "risky": '<span style="font-size:10px;padding:2px 8px;border-radius:20px;background:#fce8e6;color:#b91c1c;border:1px solid #f5a3a3;font-weight:600">Rủi ro cao</span>',
    }.get(loai, '<span style="font-size:10px;padding:2px 8px;border-radius:20px;background:#e8f0fe;color:#1a56c4;border:1px solid #93b4f5;font-weight:600">Thay thế</span>')

def _chatbot_render_answer_html(data):
    """Dựng HTML thẻ trả lời có cấu trúc (tương đương hàm render() bản JS), hiển thị qua st.markdown."""
    import html as _html
    al = _chatbot_all_laws()
    h = '<div style="padding:14px 16px;background:#fff;border:1px solid #e5e7eb;border-radius:4px 14px 14px 14px;box-shadow:0 1px 4px rgba(0,0,0,.05)">'
    h += '<div style="display:flex;align-items:center;gap:6px;margin-bottom:11px;padding-bottom:8px;border-bottom:1px solid #f1f5f9">' \
         '<div style="width:22px;height:22px;border-radius:6px;background:#1e3a5f;display:flex;align-items:center;justify-content:center;font-size:11px">⚖️</div>' \
         '<span style="font-size:10.5px;font-weight:700;color:#1e3a5f;letter-spacing:.06em">CHUYÊN GIA TƯ VẤN HCNS</span></div>'
    if data.get("summary"):
        h += f'<div style="padding:10px 13px;background:#f0f9ff;border:1px solid #bae6fd;border-radius:8px;margin-bottom:11px;font-size:13px;color:#0c4a6e;font-weight:500;line-height:1.65">{_html.escape(data["summary"])}</div>'
    if data.get("analysis"):
        h += f'<div style="font-size:13.5px;color:#374151;line-height:1.8;margin-bottom:13px">{_html.escape(data["analysis"])}</div>'
    if data.get("options"):
        h += '<div style="font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.07em;text-transform:uppercase;margin-bottom:7px">⚖️ Phương án xử lý</div>'
        for o in data["options"]:
            risk_html = f'<div style="margin-top:5px;font-size:12px;color:#b91c1c;background:#fef2f2;padding:4px 9px;border-radius:5px">⚠️ {_html.escape(o.get("risk",""))}</div>' if o.get("risk") else ""
            h += f'<div style="margin-bottom:7px;padding:10px 13px;background:#f9fafb;border:1px solid #e5e7eb;border-radius:8px">' \
                 f'<div style="display:flex;align-items:center;gap:8px;margin-bottom:4px;flex-wrap:wrap">' \
                 f'<span style="font-size:13px;font-weight:700;color:#111827">{_html.escape(o.get("label",""))}</span>{_chatbot_badge_html(o.get("type"))}</div>' \
                 f'<div style="font-size:13px;color:#4b5563;line-height:1.65">{_html.escape(o.get("detail",""))}</div>{risk_html}</div>'
    if data.get("calculations"):
        h += '<div style="font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.07em;text-transform:uppercase;margin:10px 0 7px">🔢 Tính toán cụ thể</div>'
        for c in data["calculations"]:
            ket_qua = f'<span style="font-weight:700;color:#14532d"> = {_html.escape(str(c.get("result","")))}</span>' if c.get("result") else ""
            h += f'<div style="margin-bottom:5px;padding:8px 12px;background:#f0fdf4;border:1px solid #bbf7d0;border-radius:6px;font-size:13px">' \
                 f'<span style="font-weight:700;color:#166534">{_html.escape(c.get("label",""))}: </span>' \
                 f'<span style="color:#15803d">{_html.escape(c.get("formula",""))}</span>{ket_qua}</div>'
    if data.get("citations"):
        matched = [l for l in al if l["id"] in data["citations"]]
        if matched:
            h += '<div style="margin-top:13px;border-top:1px solid #f3f4f6;padding-top:11px">' \
                 '<div style="font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.07em;text-transform:uppercase;margin-bottom:7px">📚 Căn cứ pháp lý</div>'
            for law in matched:
                h += f'<div style="margin-bottom:8px;padding:8px 12px;background:#fffbeb;border:1px solid #fde68a;border-left:3px solid #f59e0b;border-radius:6px">' \
                     f'<div style="font-size:11px;font-weight:700;color:#92400e;margin-bottom:3px">{_html.escape(law["ref"])}</div>' \
                     f'<div style="font-size:12px;color:#78350f;line-height:1.65">{_html.escape(law["text"])}</div></div>'
            h += '</div>'
    if data.get("note"):
        h += f'<div style="margin-top:9px;padding:8px 12px;background:#faf5ff;border:1px solid #e9d5ff;border-radius:6px;font-size:12.5px;color:#6b21a8;line-height:1.65">💡 <strong>Lưu ý:</strong> {_html.escape(data["note"])}</div>'
    h += '</div>'
    return h

def upload_anh_ho_so(ma_nv_or_id, ho_ten, uploaded_file):
    """Upload ảnh hồ sơ nhân viên lên Supabase Storage (dùng chung bucket hồ sơ,
    lưu trong thư mục con 'avatars/'). Trả về storage_path đã lưu, hoặc None nếu lỗi."""
    if not uploaded_file:
        return None
    sb = get_supabase_storage()
    if not sb:
        st.warning("⚠️ Chưa cấu hình Supabase Storage nên không lưu được ảnh hồ sơ (các thông tin khác vẫn được lưu bình thường).")
        return None
    try:
        safe_name = sanitize_storage_filename(uploaded_file.name)
        ten_folder = sanitize_storage_filename(f"{ma_nv_or_id}_{ho_ten}")
        base_path = f"avatars/{ten_folder}/{safe_name}"
        return upload_to_storage_unique(
            sb, SUPABASE_BUCKET, base_path,
            uploaded_file.getvalue(), uploaded_file.type
        )
    except Exception as e:
        loi_bucket = st.session_state.get(f"_sb_bucket_error_{SUPABASE_BUCKET}")
        if loi_bucket and 'bucket' in str(e).lower():
            st.warning(
                f"⚠️ Lỗi upload ảnh hồ sơ (các thông tin khác vẫn được lưu): {e}\n\n"
                f"**Nguyên nhân:** App không tự tạo được bucket Storage `{SUPABASE_BUCKET}` trên project "
                f"Supabase của công ty bạn (chi tiết lỗi tạo bucket: {loi_bucket}). Thường do khoá "
                f"Supabase (Supabase Key) đang cấu hình cho công ty bạn là **anon/public key** — loại "
                f"key này không có quyền tạo bucket.\n\n"
                f"**Cách khắc phục** (chọn 1 trong 2):\n"
                f"1. Vào Supabase Dashboard của công ty bạn > Storage > New bucket, tạo bucket tên chính "
                f"xác `{SUPABASE_BUCKET}` (để **Private**); hoặc\n"
                f"2. Đổi Supabase Key đang cấu hình cho tenant sang **service_role key** tại "
                f"'Quản trị hệ thống > Danh sách khách hàng (Tenants)'."
            )
        else:
            st.warning(f"⚠️ Lỗi upload ảnh hồ sơ (các thông tin khác vẫn được lưu): {e}")
        return None

@st.cache_data(ttl=600, show_spinner=False)
def get_anh_ho_so_bytes(storage_path):
    """Tải bytes ảnh hồ sơ từ Supabase Storage để hiển thị (bucket riêng tư nên
    không dùng URL public trực tiếp được). Cache 10 phút để đỡ tải lại liên tục."""
    if not storage_path:
        return None
    try:
        sb = get_supabase_storage()
        if not sb:
            return None
        return sb.storage.from_(SUPABASE_BUCKET).download(storage_path)
    except Exception:
        return None

# ===== CHẤM CÔNG - HẰNG SỐ MỚI =====
# Ký hiệu chấm công 22 mã chuẩn (đồng bộ KY_HIEU_CHAM_CONG ở đầu file)
CHAM_CONG_MA_CODE = {
    "":      "(Trống)",
    "x":     "x - Đi làm",
    "x/2":   "x/2 - Nửa ngày",
    "P":     "P - Phép năm",
    "1/2P":  "1/2P - Phép nửa ngày",
    "NL":    "NL - Nghỉ lễ",
    "CN":    "CN - Nghỉ tuần",
    "CT":    "CT - Công tác",
    "NB":    "NB - Nghỉ bù",
    "Ro":    "Ro - Nghỉ riêng HL",
    "OD":    "OD - Ốm đau",
    "CÔ":    "CÔ - Con ốm",
    "TS":    "TS - Thai sản",
    "KT":    "KT - Khám thai",
    "TN":    "TN - Tai nạn LĐ",
    "DSOD":  "DSOD - Dưỡng sức ốm",
    "DSTS":  "DSTS - Dưỡng sức TS",
    "DSTN":  "DSTN - Dưỡng sức TN",
    "KL":    "KL - Không lương",
    "KP":    "KP - Không phép",
}
CHAM_CONG_MA_OPTIONS = list(CHAM_CONG_MA_CODE.keys())
# Regex cho ô BCC: chấp nhận 22 ký hiệu + số giờ tăng ca (cho dòng TC)
CHAM_CONG_CELL_REGEX = (
    r"^$"                                      # trống
    r"|^[xX]$"                                  # x
    r"|^[xX]/2$"                                # x/2
    r"|^[Pp]$"                                  # P
    r"|^1/2[Pp]$"                               # 1/2P
    r"|^[Nn][Ll]$"                              # NL
    r"|^[Cc][Nn]$"                              # CN
    r"|^[Cc][Tt]$"                              # CT
    r"|^[Nn][Bb]$"                              # NB
    r"|^[Rr][Oo]$"                              # Ro
    r"|^[Oo][Dd]$"                              # OD
    r"|^[Cc][ÔÔôô]$"                            # CÔ
    r"|^[Tt][Ss]$"                              # TS
    r"|^[Kk][Tt]$"                              # KT
    r"|^[Tt][Nn]$"                              # TN
    r"|^[Dd][Ss][Oo][Dd]$"                      # DSOD
    r"|^[Dd][Ss][Tt][Ss]$"                      # DSTS
    r"|^[Dd][Ss][Tt][Nn]$"                      # DSTN
    r"|^[Kk][Ll]$"                              # KL
    r"|^[Kk][Pp]$"                              # KP
    r"|^\d{1,2}(\.\d{1,2})?$"                   # số giờ tăng ca (0-99.99)
)

# Bộ phận chỉ chấm công 1 dòng (Văn phòng)
CHAM_CONG_DEPT_MOT_DONG = ["VP"]

# Bộ phận có 2 dòng (Ca chính + Tăng ca)
CHAM_CONG_DEPT_HAI_DONG = ["QL", "SX", "LDPT"]

# === THÊM MỚI: Từ khóa nhận diện Văn phòng ===
VAN_PHONG_KEYWORDS = ["VP", "VĂN PHÒNG", "ADMIN", "HÀNH CHÍNH"]

def is_van_phong(dept):
    """Kiểm tra xem bộ phận có thuộc Văn phòng không"""
    if not dept:
        return False
    dept_upper = dept.upper()
    for kw in VAN_PHONG_KEYWORDS:
        if kw in dept_upper:
            return True
    return False

CHAM_CONG_DEPT_LABEL = {
    "QL": "QL - Quản lý",
    "VP": "VP - Văn phòng",
    "SX": "SX - Sản xuất/Vận hành",
    "LDPT": "LDPT - Lao động phổ thông",
}

# Xử lý đổi ngôn ngữ từ request
def handle_language_change():
    """Xử lý thay đổi ngôn ngữ từ query params"""
    query_params = st.query_params
    if 'lang' in query_params:
        new_lang = query_params['lang']
        if new_lang in ['vi', 'en']:
            st.session_state.language = new_lang
            st.query_params.clear()
            st.rerun()
    
    # Cũng kiểm tra POST request (cho fetch từ client)
    try:
        # Lấy dữ liệu từ request (nếu có)
        import sys
        if hasattr(st, 'context') and hasattr(st.context, 'headers'):
            content_length = int(st.context.headers.get('content-length', 0))
            if content_length > 0:
                body = sys.stdin.read(content_length) if content_length else ''
                if 'set_language=' in body:
                    new_lang = body.replace('set_language=', '').strip()
                    if new_lang in ['vi', 'en']:
                        st.session_state.language = new_lang
                        st.rerun()
    except:
        pass

# Gọi hàm xử lý ngôn ngữ trước khi hiển thị landing page
handle_language_change()

st.set_page_config(page_title="HRM-Port", page_icon="🏗️", layout="wide",
                   initial_sidebar_state="expanded")

# Mobile UX: sidebar login + nút Menu nổi sau đăng nhập
import streamlit.components.v1 as components

if not st.session_state.get('logged_in', False):
    # CHƯA LOGIN: ép sidebar full-width, ẩn main content
    st.markdown("""
    <style>
        @media (max-width: 768px) {
            [data-testid="stSidebar"] {
                width: 100vw !important;
                min-width: 100vw !important;
                transform: none !important;
                position: relative !important;
                z-index: 999999 !important;
            }
            [data-testid="stSidebar"] > div:first-child {
                width: 100vw !important;
            }
            [data-testid="stSidebar"] button[kind="header"] {
                display: none !important;
            }
            section[data-testid="stMain"] {
                display: none !important;
            }
        }
    </style>
    """, unsafe_allow_html=True)
    components.html("""
    <script>
        const sb = window.top.document.querySelector('[data-testid="stSidebar"]');
        if (sb) sb.setAttribute('aria-expanded', 'true');
        const btn = window.top.document.querySelector('[data-testid="baseButton-header"]');
        if (btn && sb && sb.getAttribute('aria-expanded') !== 'true') btn.click();
    </script>
    """, height=0)
else:
    # ĐÃ LOGIN trên mobile: hiện gợi ý vuốt + nút mở sidebar
    st.markdown("""
    <style>
        @media (max-width: 768px) {
            /* Phóng to nút mở sidebar mặc định của Streamlit */
            [data-testid="collapsedControl"] {
                background: #ff4b4b !important;
                border-radius: 8px !important;
                padding: 4px 8px !important;
                box-shadow: 0 2px 8px rgba(0,0,0,0.3) !important;
            }
            [data-testid="collapsedControl"] svg {
                width: 26px !important;
                height: 26px !important;
                color: white !important;
                stroke: white !important;
            }
            /* Phóng to nút > ở header (một số phiên bản Streamlit) */
            [data-testid="stHeader"] button {
                background: #ff4b4b !important;
                border-radius: 8px !important;
            }
            [data-testid="stHeader"] button svg {
                color: white !important;
                stroke: white !important;
            }
        }
    </style>
    """, unsafe_allow_html=True)

# Gọi định danh tenant
resolve_tenant()

# Nạp ngôn ngữ hiển thị của tenant (VI / VI_EN / VI_ZH / VI_KO) — chỉ ảnh hưởng NHÃN
# hiển thị trên giao diện (menu, tiêu đề, một số form), KHÔNG đổi logic/dữ liệu.
i18n.set_active_language((st.session_state.get('tenant') or {}).get('ngon_ngu'))

# Nạp động cấu hình thương hiệu (Branding) từ tenant
if st.session_state.get('tenant'):
    tenant_data = st.session_state.tenant
    mapping = {
        "ten_cong_ty": "ten_cty",
        "dai_dien": "dai_dien",
        "chuc_vu": "chuc_vu",
        "ma_so_thue": "ma_so_thue",
        "dien_thoai_cty": "dien_thoai_cty",
        "ma_don_vi_BHXH": "ma_don_vi_BHXH",
        "ma_vung_luong": "ma_vung_luong",
        "dia_chi": "dia_chi",
        "loi_nhan_zalo": "loi_nhan_zalo",
        "zalo_group_link": "zalo_group_link",
        "zalo_group_name": "zalo_group_name",
        "logo_url": "logo_url"
    }
    for config_key, tenant_key in mapping.items():
        if tenant_key in tenant_data and tenant_data[tenant_key]:
            COMPANY_CONFIG[config_key] = tenant_data[tenant_key]

# ========== XỬ LÝ ĐA NGÔN NGỮ ==========
def init_language():
    """Khởi tạo và xử lý chuyển đổi ngôn ngữ"""
    if 'language' not in st.session_state:
        st.session_state.language = 'vi'
    
    # Kiểm tra query params để đổi ngôn ngữ
    query_params = st.query_params
    if 'lang' in query_params:
        new_lang = query_params['lang']
        if new_lang in ['vi', 'en']:
            st.session_state.language = new_lang
            # Xóa param để tránh lặp
            st.query_params.clear()
            st.rerun()

# Gọi hàm khởi tạo
init_language()

# ========== QUẢN LÝ CÔNG VĂN & HĐ KINH TẾ ==========

# === Hàm khởi tạo bảng nếu chưa có ===
def init_cong_van_tables():
    """Khởi tạo các bảng cho module Quản lý Công văn & HĐ kinh tế"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        
        # Tạo bảng cấu hình công văn
        c.execute("""
            CREATE TABLE IF NOT EXISTS cau_hinh_cong_van (
                id SERIAL PRIMARY KEY,
                loai VARCHAR(20) NOT NULL,
                so_max INTEGER NOT NULL DEFAULT 0,
                prefix VARCHAR(10),
                nam_hien_tai INTEGER NOT NULL DEFAULT EXTRACT(YEAR FROM CURRENT_DATE),
                updated_at TIMESTAMP DEFAULT NOW(),
                UNIQUE(loai, nam_hien_tai)
            )
        """)
        
        # Tạo bảng công văn đến
        c.execute("""
            CREATE TABLE IF NOT EXISTS cong_van_den (
                id SERIAL PRIMARY KEY,
                so_cong_van VARCHAR(50) NOT NULL,
                co_quan_phat_hanh VARCHAR(200) NOT NULL,
                ngay_den DATE NOT NULL DEFAULT CURRENT_DATE,
                tieu_de TEXT NOT NULL,
                trich_yeu TEXT,
                file_url TEXT,
                ghi_chu TEXT,
                nguoi_tao VARCHAR(100),
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
        """)
        c.execute("ALTER TABLE cong_van_den ADD COLUMN IF NOT EXISTS ma_vach_buu_dien VARCHAR(100)")
        
        # Tạo bảng công văn đi
        c.execute("""
            CREATE TABLE IF NOT EXISTS cong_van_di (
                id SERIAL PRIMARY KEY,
                so_cong_van VARCHAR(50) NOT NULL,
                phong_phat_hanh VARCHAR(100) NOT NULL,
                ngay_phat_hanh DATE NOT NULL DEFAULT CURRENT_DATE,
                tieu_de TEXT NOT NULL,
                trich_yeu TEXT,
                file_url TEXT,
                loai_cong_van VARCHAR(20) NOT NULL,
                ghi_chu TEXT,
                nguoi_tao VARCHAR(100),
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
        """)
        c.execute("ALTER TABLE cong_van_di ADD COLUMN IF NOT EXISTS ma_vach_buu_dien VARCHAR(100)")
        
        # Tạo bảng hợp đồng kinh tế
        c.execute("""
            CREATE TABLE IF NOT EXISTS hop_dong_kinh_te (
                id SERIAL PRIMARY KEY,
                so_hop_dong VARCHAR(50) NOT NULL,
                ten_doi_tac VARCHAR(200) NOT NULL,
                ngay_ky DATE NOT NULL DEFAULT CURRENT_DATE,
                trich_yeu TEXT,
                file_url TEXT,
                ghi_chu TEXT,
                nguoi_tao VARCHAR(100),
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
        """)
        
        # Tạo bảng danh mục loại công văn
        c.execute("""
            CREATE TABLE IF NOT EXISTS danh_muc_loai_cong_van (
                id SERIAL PRIMARY KEY,
                ma_loai VARCHAR(10) NOT NULL UNIQUE,
                ten_loai VARCHAR(50) NOT NULL,
                thu_tu INTEGER DEFAULT 0,
                trang_thai BOOLEAN DEFAULT TRUE
            )
        """)
        
        # Insert dữ liệu mặc định cho danh mục loại công văn
        c.execute("""
            INSERT INTO danh_muc_loai_cong_van (ma_loai, ten_loai, thu_tu) VALUES
            ('QĐ', 'Quyết định', 1),
            ('CV', 'Công văn', 2),
            ('BC', 'Báo cáo', 3),
            ('TB', 'Thông báo', 4),
            ('TTr', 'Tờ trình', 5)
            ON CONFLICT (ma_loai) DO NOTHING
        """)
        
        # Tạo bảng cấu hình hệ thống
        c.execute("""
            CREATE TABLE IF NOT EXISTS cau_hinh_he_thong (
                id SERIAL PRIMARY KEY,
                ten_cau_hinh VARCHAR(50) NOT NULL UNIQUE,
                gia_tri VARCHAR(100),
                mo_ta TEXT,
                updated_at TIMESTAMP DEFAULT NOW()
            )
        """)
        
        # Insert cấu hình mặc định
        c.execute("""
            INSERT INTO cau_hinh_he_thong (ten_cau_hinh, gia_tri, mo_ta) VALUES
            ('cv_danh_so_option', 'RIENG', 'CHUNG hoac RIENG - Cách đánh số công văn')
            ON CONFLICT (ten_cau_hinh) DO NOTHING
        """)
        
        db.commit()
        db.close()
        return True
    except Exception as e:
        print(f"Lỗi khởi tạo bảng công văn: {e}")
        return False

# === Hàm lấy cấu hình đánh số ===
def get_cv_danh_so_option():
    """Lấy option đánh số công văn: CHUNG hoặc RIENG"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT gia_tri FROM cau_hinh_he_thong WHERE ten_cau_hinh = 'cv_danh_so_option'")
        result = c.fetchone()
        db.close()
        return result[0] if result else 'RIENG'
    except:
        return 'RIENG'

# === Hàm cập nhật cấu hình đánh số ===
def update_cv_danh_so_option(option):
    """Cập nhật option đánh số công văn"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            UPDATE cau_hinh_he_thong 
            SET gia_tri = %s, updated_at = NOW() 
            WHERE ten_cau_hinh = 'cv_danh_so_option'
        """, (option,))
        db.commit()
        db.close()
        return True
    except:
        return False

def get_cv_kieu_ngay(loai):
    """Lấy cấu hình đánh số theo kiểu dd/mm cho loại công văn cụ thể.
    Trả về True nếu loại đó bật kiểu dd/mm, False nếu dùng kiểu mặc định (chỉ năm)."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT gia_tri FROM cau_hinh_he_thong WHERE ten_cau_hinh = %s", (f'cv_kieu_ngay_{loai}',))
        result = c.fetchone()
        db.close()
        return result[0] == 'DD_MM' if result else False
    except:
        return False

def update_cv_kieu_ngay(loai, kieu):
    """Cập nhật cấu hình kiểu đánh số ngày cho loại công văn.
    kieu: 'DD_MM' hoặc 'NAM' (chỉ năm, mặc định)."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        ten_cau_hinh = f'cv_kieu_ngay_{loai}'
        c.execute("""
            INSERT INTO cau_hinh_he_thong (ten_cau_hinh, gia_tri, updated_at)
            VALUES (%s, %s, NOW())
            ON CONFLICT (ten_cau_hinh) DO UPDATE SET gia_tri = EXCLUDED.gia_tri, updated_at = NOW()
        """, (ten_cau_hinh, kieu))
        db.commit()
        db.close()
        return True
    except:
        return False

# === Hàm lấy số max hiện tại ===
def get_so_max_cong_van(loai=None):
    """Lấy số max hiện tại cho loại công văn (hoặc chung nếu loai=None)"""
    nam_hien_tai = datetime.now().year
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        
        if loai:
            c.execute("""
                SELECT so_max FROM cau_hinh_cong_van 
                WHERE loai = %s AND nam_hien_tai = %s
            """, (loai, nam_hien_tai))
        else:
            c.execute("""
                SELECT so_max FROM cau_hinh_cong_van 
                WHERE loai = 'CHUNG' AND nam_hien_tai = %s
            """, (nam_hien_tai,))
        
        result = c.fetchone()
        db.close()
        return result[0] if result else 0
    except:
        return 0

# === Hàm cập nhật số max ===
def update_so_max_cong_van(loai, so_moi):
    """Cập nhật số max cho loại công văn"""
    nam_hien_tai = datetime.now().year
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            INSERT INTO cau_hinh_cong_van (loai, so_max, nam_hien_tai, updated_at)
            VALUES (%s, %s, %s, NOW())
            ON CONFLICT (loai, nam_hien_tai) 
            DO UPDATE SET so_max = EXCLUDED.so_max, updated_at = NOW()
        """, (loai, so_moi, nam_hien_tai))
        db.commit()
        db.close()
        return True
    except:
        return False

# === Cấu hình + đánh số Hợp đồng kinh tế (HĐKT): mẫu "stt/năm/Prefix-ma_cty" ===
# === Cấu hình hạn nộp Báo cáo Tăng/Giảm BHXH hàng tháng (mỗi doanh nghiệp 1 ngày riêng) ===
def get_han_nop_bhxh():
    """Ngày trong tháng phải nộp BC Tăng/Giảm BHXH (mặc định 20, VD CHL nộp trước ngày 20)."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT gia_tri FROM cau_hinh_he_thong WHERE ten_cau_hinh = 'han_nop_bhxh_ngay'")
        r = c.fetchone()
        db.close()
        return int(r[0]) if r and r[0] else 20
    except:
        return 20

def update_han_nop_bhxh(ngay):
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            INSERT INTO cau_hinh_he_thong (ten_cau_hinh, gia_tri, mo_ta)
            VALUES ('han_nop_bhxh_ngay', %s, 'Ngày trong tháng phải nộp BC Tăng/Giảm BHXH')
            ON CONFLICT (ten_cau_hinh) DO UPDATE SET gia_tri = EXCLUDED.gia_tri, updated_at = NOW()
        """, (str(ngay),))
        db.commit(); db.close()
        return True
    except:
        return False


# === Cấu hình chung theo tenant (key-value) — dùng lại bảng cau_hinh_he_thong ===
# Đây chính là cơ chế "tenant_settings": mỗi tenant có 1 DB riêng, bảng
# cau_hinh_he_thong trong DB đó lưu các thiết lập riêng của công ty mình,
# không ảnh hưởng tenant khác. get_cau_hinh()/set_cau_hinh() là hàm dùng
# chung cho MỌI khoá cấu hình (thay vì phải viết riêng 1 cặp get/update
# cho từng khoá như cv_danh_so_option, han_nop_bhxh_ngay ở trên).
def get_cau_hinh(key, default=None):
    """Đọc 1 giá trị cấu hình theo tenant đang đăng nhập. Trả về `default` nếu
    chưa từng thiết lập (tenant mới) hoặc DB lỗi.
    Có cache trong session (mỗi phiên chỉ query DB 1 lần/khoá) để tránh gọi DB
    lặp lại hàng trăm lần khi hàm này được gọi trong vòng lặp xuất báo cáo
    theo từng nhân viên."""
    cache = st.session_state.setdefault('_cau_hinh_cache', {})
    if key in cache:
        return cache[key] if cache[key] is not None else default
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT gia_tri FROM cau_hinh_he_thong WHERE ten_cau_hinh = %s", (key,))
        r = c.fetchone()
        db.close()
        val = r[0] if r and r[0] not in (None, '') else None
        cache[key] = val
        return val if val is not None else default
    except Exception:
        cache[key] = None
        return default


def set_cau_hinh(key, value, mo_ta=''):
    """Ghi/cập nhật 1 giá trị cấu hình theo tenant đang đăng nhập."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            INSERT INTO cau_hinh_he_thong (ten_cau_hinh, gia_tri, mo_ta)
            VALUES (%s, %s, %s)
            ON CONFLICT (ten_cau_hinh) DO UPDATE SET gia_tri = EXCLUDED.gia_tri,
                mo_ta = COALESCE(NULLIF(EXCLUDED.mo_ta, ''), cau_hinh_he_thong.mo_ta), updated_at = NOW()
        """, (key, '' if value is None else str(value), mo_ta))
        db.commit(); db.close()
        st.session_state.setdefault('_cau_hinh_cache', {})[key] = value
        return True
    except Exception:
        return False


# === Sinh Mã NV / Số HĐLĐ-HĐTV tự động — CHUNG cho mọi nơi cần (trước đây mỗi form
# "Thêm nhân viên" tự viết lại logic riêng, hardcode ký hiệu 'C' giống nhau cho MỌI
# tenant, và số hợp đồng bị "reset về 01" nếu dữ liệu cũ (import/seed) từng dùng hậu tố
# mã công ty khác với ký hiệu hiện tại của tenant). ===
def get_ky_hieu_ma_nv():
    """Ký hiệu (tiền tố) Mã nhân viên riêng theo từng tenant, VD 'NV', 'HL'...
    Cấu hình tại ⚙️ Danh mục > 'Ký hiệu mã nhân viên'. Mặc định 'NV' nếu tenant chưa đặt
    (trước đây hardcode 'C' cho MỌI tenant, không phản ánh đúng ký hiệu riêng công ty)."""
    ky_hieu = get_cau_hinh('ky_hieu_ma_nv', 'NV') or 'NV'
    ky_hieu = re.sub(r'[^A-Za-z0-9]', '', ky_hieu).upper() or 'NV'
    return ky_hieu


def sinh_ma_nv_moi(cursor):
    """Sinh Mã nhân viên mới kế tiếp theo đúng ký hiệu riêng của tenant đang đăng nhập,
    dựa trên giá trị số lớn nhất đã có trong DB có cùng tiền tố này."""
    prefix = get_ky_hieu_ma_nv()
    len_prefix = len(prefix) + 1
    cursor.execute(f"""
        SELECT COALESCE(MAX(CAST(SUBSTRING(ma_nv FROM {len_prefix}) AS INTEGER)), 0) + 1
        FROM nhan_vien
        WHERE ma_nv LIKE %s AND SUBSTRING(ma_nv FROM {len_prefix}) ~ '^[0-9]+$'
    """, (f'{prefix}%',))
    so_moi = cursor.fetchone()[0]
    return f"{prefix}{so_moi:03d}"


def sinh_so_hdld_moi(cursor, ma_cty_hd, nam, la_thu_viec):
    """Sinh Số HĐTV/HĐLĐ mới kế tiếp (định dạng 'STT/năm/HĐTV-MACTY' hoặc
    'STT/năm/HĐLĐ-MACTY'), dựa trên MAX đã có TRONG DB.
    QUAN TRỌNG: chỉ lọc theo loại văn bản (HĐTV/HĐLĐ) — KHÔNG bắt buộc đúng hậu tố mã
    công ty hiện tại — để STT không bị "reset về 01" khi dữ liệu cũ (import Excel/seed
    demo) từng được đánh số với ký hiệu công ty khác (VD dữ liệu cũ hậu tố '-DEMO' còn
    tenant hiện tại là '-DEMO-HRM'). Mỗi tenant có DB riêng nên không lo lẫn dữ liệu
    giữa các công ty khác nhau."""
    nhan = 'HĐTV' if la_thu_viec else 'HĐLĐ'
    if la_thu_viec:
        cursor.execute("""
            SELECT COALESCE(MAX(CAST(SPLIT_PART(so_hdld, '/', 1) AS INTEGER)), 0)
            FROM nhan_vien
            WHERE so_hdld LIKE %s
              AND SPLIT_PART(so_hdld, '/', 1) ~ '^[0-9]+$'
              AND trang_thai IN ('THU_VIEC', 'DANG_LAM')
        """, (f'%/{nhan}-%',))
    else:
        cursor.execute("""
            SELECT COALESCE(MAX(CAST(SPLIT_PART(so_hdld, '/', 1) AS INTEGER)), 0)
            FROM nhan_vien
            WHERE so_hdld LIKE %s
              AND SPLIT_PART(so_hdld, '/', 1) ~ '^[0-9]+$'
              AND trang_thai = 'DANG_LAM'
              AND loai_hop_dong != 'Thử việc'
        """, (f'%/{nhan}-%',))
    stt = (cursor.fetchone()[0] or 0) + 1
    return f"{stt:02d}/{nam}/{nhan}-{ma_cty_hd}"

# === Cấu hình Chấm công (khung sườn - sẽ tích hợp logic tính công dần theo nhu cầu) ===
def get_cau_hinh_cham_cong():
    from datetime import time as _time
    mac_dinh = {'gio_vao': _time(8, 0), 'gio_ra': _time(17, 0), 'phut_tre': 15}
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT ten_cau_hinh, gia_tri FROM cau_hinh_he_thong WHERE ten_cau_hinh IN "
                   "('cc_gio_vao', 'cc_gio_ra', 'cc_phut_tre')")
        rows = dict(c.fetchall())
        db.close()
        if rows.get('cc_gio_vao'):
            h, m = map(int, rows['cc_gio_vao'].split(':'))
            mac_dinh['gio_vao'] = _time(h, m)
        if rows.get('cc_gio_ra'):
            h, m = map(int, rows['cc_gio_ra'].split(':'))
            mac_dinh['gio_ra'] = _time(h, m)
        if rows.get('cc_phut_tre'):
            mac_dinh['phut_tre'] = int(rows['cc_phut_tre'])
        return mac_dinh
    except:
        return mac_dinh

def update_cau_hinh_cham_cong(gio_vao, gio_ra, phut_tre):
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        for ten, gia_tri in [('cc_gio_vao', gio_vao.strftime('%H:%M')),
                              ('cc_gio_ra', gio_ra.strftime('%H:%M')),
                              ('cc_phut_tre', str(phut_tre))]:
            c.execute("""
                INSERT INTO cau_hinh_he_thong (ten_cau_hinh, gia_tri, mo_ta)
                VALUES (%s, %s, 'Cấu hình chấm công')
                ON CONFLICT (ten_cau_hinh) DO UPDATE SET gia_tri = EXCLUDED.gia_tri, updated_at = NOW()
            """, (ten, gia_tri))
        db.commit(); db.close()
        return True
    except:
        return False

def get_cau_hinh_cham_cong_full():
    """Đọc toàn bộ cấu hình chấm công (giờ vào/ra, số ngày làm việc, hệ số TC,
    phép năm, ngày lễ...). Dùng chung cơ chế get_cau_hinh() (cache session,
    bảng cau_hinh_he_thong) — không tạo bảng riêng."""
    def _parse_time(s, default):
        try:
            h, m = map(int, s.split(':'))
            return _time(h, m)
        except Exception:
            return default

    try:
        danh_sach_le = json.loads(get_cau_hinh('cc_danh_sach_ngay_le', '[]'))
    except Exception:
        danh_sach_le = []

    return {
        'gio_vao': _parse_time(get_cau_hinh('cc_gio_vao', '08:00'), _time(8, 0)),
        'gio_ra': _parse_time(get_cau_hinh('cc_gio_ra', '17:00'), _time(17, 0)),
        'phut_tre': int(get_cau_hinh('cc_phut_tre', 15)),
        'gio_bat_dau_ca_dem': _parse_time(get_cau_hinh('cc_gio_bat_dau_ca_dem', '22:00'), _time(22, 0)),
        'so_ngay_lam_viec_tuan': int(get_cau_hinh('cc_so_ngay_lam_viec_tuan', 6)),
        'ngay_nghi_hang_tuan': get_cau_hinh('cc_ngay_nghi_hang_tuan', 'CN'),
        'gio_lam_chuan_ngay': float(get_cau_hinh('cc_gio_lam_chuan_ngay', 8)),
        'he_so_tc_thuong': float(get_cau_hinh('cc_he_so_tc_thuong', 1.5)),
        'he_so_tc_chu_nhat': float(get_cau_hinh('cc_he_so_tc_chu_nhat', 2.0)),
        'he_so_tc_le': float(get_cau_hinh('cc_he_so_tc_le', 3.0)),
        'he_so_tc_dem': float(get_cau_hinh('cc_he_so_tc_dem', 1.3)),
        'cach_tinh_tang_ca': get_cau_hinh('cc_cach_tinh_tang_ca', 'HE_SO'),  # HE_SO | DON_GIA
        'don_gia_tc_thuong': float(get_cau_hinh('cc_don_gia_tc_thuong', 0)),
        'don_gia_tc_chu_nhat': float(get_cau_hinh('cc_don_gia_tc_chu_nhat', 0)),
        'don_gia_tc_le': float(get_cau_hinh('cc_don_gia_tc_le', 0)),
        'don_gia_tc_dem': float(get_cau_hinh('cc_don_gia_tc_dem', 0)),
        'cach_tinh_phep_nam': get_cau_hinh('cc_cach_tinh_phep_nam', 'TU_DONG'),
        'so_ngay_phep_co_ban': float(get_cau_hinh('cc_so_ngay_phep_co_ban', 12)),
        'danh_sach_ngay_le': danh_sach_le,
        'phut_nghi_giua_ca': int(get_cau_hinh('cc_phut_nghi_giua_ca', 120)),
        'ap_dung_nghi_giua_ca': get_cau_hinh('cc_ap_dung_nghi_giua_ca', '1') == '1',
        'phut_toi_thieu_tang_ca': int(get_cau_hinh('cc_phut_toi_thieu_tang_ca', 30)),
    }


def update_cau_hinh_cham_cong_full(cfg):
    """Ghi toàn bộ cấu hình chấm công. `cfg` cùng cấu trúc dict trả về bởi
    get_cau_hinh_cham_cong_full()."""
    ok = True
    ok &= set_cau_hinh('cc_gio_vao', cfg['gio_vao'].strftime('%H:%M'), 'Giờ vào chuẩn')
    ok &= set_cau_hinh('cc_gio_ra', cfg['gio_ra'].strftime('%H:%M'), 'Giờ ra chuẩn')
    ok &= set_cau_hinh('cc_phut_tre', str(cfg['phut_tre']), 'Số phút cho phép trễ')
    ok &= set_cau_hinh('cc_gio_bat_dau_ca_dem', cfg['gio_bat_dau_ca_dem'].strftime('%H:%M'), 'Giờ bắt đầu ca đêm (tính TCĐ)')
    ok &= set_cau_hinh('cc_so_ngay_lam_viec_tuan', str(cfg['so_ngay_lam_viec_tuan']), 'Số ngày làm việc/tuần')
    ok &= set_cau_hinh('cc_ngay_nghi_hang_tuan', cfg['ngay_nghi_hang_tuan'], 'Ngày nghỉ hàng tuần')
    ok &= set_cau_hinh('cc_gio_lam_chuan_ngay', str(cfg['gio_lam_chuan_ngay']), 'Giờ làm chuẩn/ngày')
    ok &= set_cau_hinh('cc_he_so_tc_thuong', str(cfg['he_so_tc_thuong']), 'Hệ số tăng ca ngày thường')
    ok &= set_cau_hinh('cc_he_so_tc_chu_nhat', str(cfg['he_so_tc_chu_nhat']), 'Hệ số tăng ca Chủ nhật')
    ok &= set_cau_hinh('cc_he_so_tc_le', str(cfg['he_so_tc_le']), 'Hệ số tăng ca ngày lễ')
    ok &= set_cau_hinh('cc_he_so_tc_dem', str(cfg['he_so_tc_dem']), 'Hệ số cộng thêm tăng ca đêm')
    ok &= set_cau_hinh('cc_cach_tinh_tang_ca', cfg['cach_tinh_tang_ca'], 'Cách tính tăng ca: HE_SO hoặc DON_GIA')
    ok &= set_cau_hinh('cc_don_gia_tc_thuong', str(cfg['don_gia_tc_thuong']), 'Đơn giá TC ngày thường (đ/giờ)')
    ok &= set_cau_hinh('cc_don_gia_tc_chu_nhat', str(cfg['don_gia_tc_chu_nhat']), 'Đơn giá TC Chủ nhật (đ/giờ)')
    ok &= set_cau_hinh('cc_don_gia_tc_le', str(cfg['don_gia_tc_le']), 'Đơn giá TC ngày lễ (đ/giờ)')
    ok &= set_cau_hinh('cc_don_gia_tc_dem', str(cfg['don_gia_tc_dem']), 'Đơn giá cộng thêm TC đêm (đ/giờ)')
    ok &= set_cau_hinh('cc_cach_tinh_phep_nam', cfg['cach_tinh_phep_nam'], 'Cách tính phép năm')
    ok &= set_cau_hinh('cc_so_ngay_phep_co_ban', str(cfg['so_ngay_phep_co_ban']), 'Số ngày phép cơ bản/năm')
    ok &= set_cau_hinh('cc_danh_sach_ngay_le', json.dumps(cfg['danh_sach_ngay_le'], ensure_ascii=False), 'Danh sách ngày nghỉ lễ trong năm')
    ok &= set_cau_hinh('cc_phut_nghi_giua_ca', str(cfg['phut_nghi_giua_ca']), 'Số phút nghỉ giữa ca')
    ok &= set_cau_hinh('cc_ap_dung_nghi_giua_ca', '1' if cfg['ap_dung_nghi_giua_ca'] else '0', 'Áp dụng trừ nghỉ giữa ca khi tính giờ làm')
    ok &= set_cau_hinh('cc_phut_toi_thieu_tang_ca', str(cfg['phut_toi_thieu_tang_ca']), 'Số phút tối thiểu vượt giờ ra chuẩn mới tính là tăng ca')
    return ok
def tinh_gio_lam_thuc_te(gio_vao, gio_ra, ngay=None):
    """Tính số giờ làm thực tế từ giờ vào và giờ ra.
    Tự động trừ thời gian nghỉ giữa ca nếu cấu hình bật.
    Trả về số giờ (float, làm tròn 2 chữ số) hoặc None nếu thiếu dữ liệu."""
    if not gio_vao or not gio_ra:
        return None
    try:
        from datetime import datetime, date as _date
        _ngay = ngay or _date.today()
        dt_vao = datetime.combine(_ngay, gio_vao)
        dt_ra = datetime.combine(_ngay, gio_ra)
        if dt_ra <= dt_vao:
            dt_ra = dt_ra.replace(day=dt_ra.day + 1)  # qua ngày hôm sau (ca đêm)
        tong_phut = (dt_ra - dt_vao).total_seconds() / 60

        cfg = get_cau_hinh_cham_cong_full()
        if cfg['ap_dung_nghi_giua_ca'] and cfg['phut_nghi_giua_ca'] > 0:
            tong_phut -= cfg['phut_nghi_giua_ca']

        return round(max(tong_phut, 0) / 60, 2)
    except Exception:
        return None

def get_cau_hinh_tang_ca_theo_phong(ten_phong_ban):
    """Lấy cấu hình tăng ca của 1 phòng ban cụ thể.
    Trả về dict cấu hình, trong đó:
    - cho_phep_tang_ca: True/False
    - các hệ số/đơn giá: lấy từ phòng nếu có, fallback về cấu hình chung tenant.
    Dùng cho Payroll Engine khi tính lương tăng ca."""
    cfg_chung = get_cau_hinh_cham_cong_full()
    mac_dinh = {
        'cho_phep_tang_ca': True,
        'he_so_tc_thuong': cfg_chung['he_so_tc_thuong'],
        'he_so_tc_chu_nhat': cfg_chung['he_so_tc_chu_nhat'],
        'he_so_tc_le': cfg_chung['he_so_tc_le'],
        'he_so_tc_dem': cfg_chung['he_so_tc_dem'],
        'don_gia_tc_thuong': cfg_chung['don_gia_tc_thuong'],
        'don_gia_tc_chu_nhat': cfg_chung['don_gia_tc_chu_nhat'],
        'don_gia_tc_le': cfg_chung['don_gia_tc_le'],
        'don_gia_tc_dem': cfg_chung['don_gia_tc_dem'],
        'cach_tinh_tang_ca': cfg_chung['cach_tinh_tang_ca'],
        'nguon': 'CHUNG',  # để Payroll Engine biết đang dùng cấu hình nào
    }
    if not ten_phong_ban:
        return mac_dinh
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute(
            "SELECT * FROM cau_hinh_tang_ca_phong_ban WHERE ten_phong_ban = %s",
            (ten_phong_ban,)
        )
        row = c.fetchone()
        db.close()
        if not row:
            return mac_dinh
        ket_qua = mac_dinh.copy()
        ket_qua['cho_phep_tang_ca'] = row['cho_phep_tang_ca']
        ket_qua['nguon'] = 'PHONG_BAN'
        # Ghi đè từng trường nếu phòng ban có cấu hình riêng (không NULL)
        for truong in ['he_so_tc_thuong', 'he_so_tc_chu_nhat', 'he_so_tc_le', 'he_so_tc_dem',
                       'don_gia_tc_thuong', 'don_gia_tc_chu_nhat', 'don_gia_tc_le', 'don_gia_tc_dem']:
            if row[truong] is not None:
                ket_qua[truong] = float(row[truong])
        return ket_qua
    except Exception:
        return mac_dinh
def tong_hop_cham_cong_thang(nhan_vien_id, thang, nam, conn=None):
    """Tổng hợp chấm công tháng cho 1 nhân viên — ĐẦU VÀO CHÍNH cho Payroll Engine.
    Trả về dict đầy đủ: số công theo từng loại, giờ tăng ca theo loại, phép đã dùng,
    đi trễ/về sớm, số ngày chưa chấm, v.v.

    Payroll Engine chỉ cần gọi:
        th = tong_hop_cham_cong_thang(nv_id, thang, nam)
        cong_thuc_te = th['tong_cong']
        gio_tc_thuong = th['tc']['THUONG']
        ...
    """
    import calendar as _cal
    from datetime import date as _date

    close_conn = False
    if conn is None:
        conn = st.session_state.db_engine.get_connection()
        close_conn = True

    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("""
        SELECT ngay, ma_cong, ca_ngay, gio_vao, gio_ra,
               gio_tang_ca, gio_tang_ca_dem, gio_tang_ca_cn, gio_tang_ca_le,
               loai_ngay_tang_ca, trang_thai_duyet_tc,
               so_phut_di_tre, so_phut_ve_som,
               trang_thai_cham_cong, trang_thai_vi_tri, nguon
        FROM cham_cong
        WHERE nhan_vien_id = %s
          AND EXTRACT(MONTH FROM ngay) = %s
          AND EXTRACT(YEAR FROM ngay) = %s
        ORDER BY ngay
    """, (nhan_vien_id, thang, nam))
    rows = cur.fetchall()
    cur.close()
    if close_conn:
        conn.close()

    # Khởi tạo kết quả
    so_ngay_trong_thang = _cal.monthrange(nam, thang)[1]
    ket_qua = {
        'nhan_vien_id': nhan_vien_id,
        'thang': thang,
        'nam': nam,
        'so_ngay_trong_thang': so_ngay_trong_thang,

        # --- Tổng công (tính lương) ---
        'tong_cong': 0.0,           # Tổng ngày công được trả lương (A: x, x/2, P, 1/2P, NL, CT, NB, Ro)
        'cong_lam_viec': 0.0,       # Chỉ x + x/2 (ngày thực sự đi làm)
        'cong_huong_luong': 0.0,    # P + 1/2P + NL + CT + NB + Ro (nghỉ nhưng hưởng lương)
        'cong_bhxh': 0.0,           # OD + CÔ + TS + KT + TN + DSOD/TS/TN (BHXH chi trả, không tính vào lương CT)

        # --- Chi tiết theo từng mã công ---
        'chi_tiet_ma_cong': {},     # {'x': 20, 'P': 1, 'NL': 1, ...}

        # --- Phép ---
        'phep_da_dung': 0.0,       # P*1 + 1/2P*0.5

        # --- Tăng ca (chỉ tính đã duyệt hoặc chưa bật phê duyệt) ---
        'tc': {
            'THUONG': 0.0,          # Giờ TC ngày thường
            'CHU_NHAT': 0.0,        # Giờ TC Chủ nhật
            'LE': 0.0,              # Giờ TC ngày lễ
            'DEM': 0.0,             # Giờ TC đêm (cộng thêm, không exclusive)
        },
        'tc_cho_duyet': {           # Tương tự nhưng chỉ phần đang chờ duyệt
            'THUONG': 0.0, 'CHU_NHAT': 0.0, 'LE': 0.0, 'DEM': 0.0,
        },
        'tong_gio_tang_ca': 0.0,    # Tổng tất cả TC (đã duyệt)

        # --- Đi trễ / về sớm ---
        'tong_phut_di_tre': 0,
        'tong_phut_ve_som': 0,
        'so_lan_di_tre': 0,
        'so_lan_ve_som': 0,

        # --- Trạng thái ---
        'so_ngay_chua_cham': 0,     # Ngày làm việc mà chưa có dữ liệu
        'so_ngay_thieu_gio_ra': 0,  # Có giờ vào nhưng chưa giờ ra
        'so_ngay_cho_duyet_vi_tri': 0,  # Chấm công ngoài địa điểm đang chờ

        # --- Dữ liệu thô (để Payroll Engine cần chi tiết từng ngày) ---
        'chi_tiet_ngay': {},        # {date: row_dict}
    }

    # Tạo set ngày đã có dữ liệu
    ngay_da_co = set()

    for r in rows:
        ngay = r['ngay']
        ngay_da_co.add(ngay)
        ket_qua['chi_tiet_ngay'][ngay] = dict(r)

        # Xác định ma_cong: ưu tiên ma_cong, fallback ca_ngay (backward compat)
        ma = r['ma_cong']
        if not ma and r.get('ca_ngay'):
            # Chuyển đổi ký hiệu cũ → mới
            ma = cc_normalize_marker(r['ca_ngay'])

        if not ma:
            # Có dòng nhưng chưa có mã công (VD: THIEU_GIO_RA, CHO_DUYET)
            if r.get('trang_thai_cham_cong') == 'THIEU_GIO_RA':
                ket_qua['so_ngay_thieu_gio_ra'] += 1
            if r.get('trang_thai_vi_tri') == 'CHO_DUYET':
                ket_qua['so_ngay_cho_duyet_vi_tri'] += 1
            continue

        # Đếm theo mã công
        ket_qua['chi_tiet_ma_cong'][ma] = ket_qua['chi_tiet_ma_cong'].get(ma, 0) + 1

        # Tính công
        tt = KY_HIEU_CHAM_CONG.get(ma, {})
        so_cong = tt.get('cong', 0)
        nhom = tt.get('nhom', '')

        if nhom == 'A':
            ket_qua['tong_cong'] += so_cong
            if ma in ('x', 'x/2'):
                ket_qua['cong_lam_viec'] += so_cong
            elif ma not in ('CN',):
                ket_qua['cong_huong_luong'] += so_cong
        elif nhom == 'B':
            ket_qua['cong_bhxh'] += 1  # đếm ngày, BHXH chi trả

        # Phép
        tru_phep = tt.get('tru_phep', 0)
        if tru_phep:
            ket_qua['phep_da_dung'] += tru_phep

        # Tăng ca
        gio_tc = float(r.get('gio_tang_ca') or 0)
        gio_tc_dem = float(r.get('gio_tang_ca_dem') or 0)
        loai_ngay_tc = r.get('loai_ngay_tang_ca') or 'THUONG'
        trang_thai_tc = r.get('trang_thai_duyet_tc')

        if gio_tc > 0:
            if trang_thai_tc == 'CHO_DUYET':
                ket_qua['tc_cho_duyet'][loai_ngay_tc] += gio_tc
                if gio_tc_dem > 0:
                    ket_qua['tc_cho_duyet']['DEM'] += gio_tc_dem
            else:
                # DA_DUYET hoặc NULL (chưa bật phê duyệt = mặc định duyệt)
                ket_qua['tc'][loai_ngay_tc] += gio_tc
                ket_qua['tong_gio_tang_ca'] += gio_tc
                if gio_tc_dem > 0:
                    ket_qua['tc']['DEM'] += gio_tc_dem

        # Đi trễ / về sớm
        phut_tre = int(r.get('so_phut_di_tre') or 0)
        phut_som = int(r.get('so_phut_ve_som') or 0)
        if phut_tre > 0:
            ket_qua['tong_phut_di_tre'] += phut_tre
            ket_qua['so_lan_di_tre'] += 1
        if phut_som > 0:
            ket_qua['tong_phut_ve_som'] += phut_som
            ket_qua['so_lan_ve_som'] += 1

    # Tính ngày chưa chấm (ngày làm việc trong quá khứ, không phải CN, chưa có dữ liệu)
    cfg = get_cau_hinh_cham_cong_full()
    danh_sach_le = {x['ngay'] for x in (cfg.get('danh_sach_ngay_le') or [])}
    hom_nay = _date.today()
    for d in range(1, so_ngay_trong_thang + 1):
        ngay_d = _date(nam, thang, d)
        if ngay_d >= hom_nay:
            break  # không đếm ngày tương lai
        if ngay_d.weekday() == 6:  # Chủ nhật
            continue
        ngay_str = ngay_d.strftime('%Y-%m-%d')
        if ngay_str in danh_sach_le:
            continue  # ngày lễ
        if ngay_d not in ngay_da_co:
            ket_qua['so_ngay_chua_cham'] += 1

    return ket_qua


def tong_hop_cham_cong_thang_nhieu_nv(danh_sach_nv_id, thang, nam):
    """Tổng hợp chấm công tháng cho NHIỀU nhân viên cùng lúc (1 query).
    Trả về dict {nhan_vien_id: ket_qua_tong_hop}.
    Dùng khi hiển thị BCC tháng hoặc Payroll batch."""
    import calendar as _cal
    from datetime import date as _date

    if not danh_sach_nv_id:
        return {}

    conn = st.session_state.db_engine.get_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("""
        SELECT nhan_vien_id, ngay, ma_cong, ca_ngay, gio_vao, gio_ra,
               gio_tang_ca, gio_tang_ca_dem, gio_tang_ca_cn, gio_tang_ca_le,
               loai_ngay_tang_ca, trang_thai_duyet_tc,
               so_phut_di_tre, so_phut_ve_som,
               trang_thai_cham_cong, trang_thai_vi_tri, nguon
        FROM cham_cong
        WHERE nhan_vien_id = ANY(%s)
          AND EXTRACT(MONTH FROM ngay) = %s
          AND EXTRACT(YEAR FROM ngay) = %s
        ORDER BY nhan_vien_id, ngay
    """, (danh_sach_nv_id, thang, nam))
    all_rows = cur.fetchall()
    cur.close()
    conn.close()

    # Nhóm theo nhan_vien_id
    from collections import defaultdict
    nhom = defaultdict(list)
    for r in all_rows:
        nhom[r['nhan_vien_id']].append(r)

    # Gọi tong_hop cho từng NV (truyền rows đã có sẵn, không query lại)
    ket_qua = {}
    for nv_id in danh_sach_nv_id:
        # Dùng lại logic tong_hop nhưng không cần query thêm
        ket_qua[nv_id] = _tong_hop_tu_rows(nv_id, thang, nam, nhom.get(nv_id, []))

    return ket_qua


def _tong_hop_tu_rows(nhan_vien_id, thang, nam, rows):
    """Hàm nội bộ: tính tổng hợp từ danh sách rows đã query sẵn (không mở DB)."""
    import calendar as _cal
    from datetime import date as _date

    so_ngay_trong_thang = _cal.monthrange(nam, thang)[1]
    ket_qua = {
        'nhan_vien_id': nhan_vien_id,
        'thang': thang, 'nam': nam,
        'so_ngay_trong_thang': so_ngay_trong_thang,
        'tong_cong': 0.0,
        'cong_lam_viec': 0.0,
        'cong_huong_luong': 0.0,
        'cong_bhxh': 0.0,
        'chi_tiet_ma_cong': {},
        'phep_da_dung': 0.0,
        'tc': {'THUONG': 0.0, 'CHU_NHAT': 0.0, 'LE': 0.0, 'DEM': 0.0},
        'tc_cho_duyet': {'THUONG': 0.0, 'CHU_NHAT': 0.0, 'LE': 0.0, 'DEM': 0.0},
        'tong_gio_tang_ca': 0.0,
        'tong_phut_di_tre': 0, 'tong_phut_ve_som': 0,
        'so_lan_di_tre': 0, 'so_lan_ve_som': 0,
        'so_ngay_chua_cham': 0,
        'so_ngay_thieu_gio_ra': 0,
        'so_ngay_cho_duyet_vi_tri': 0,
        'chi_tiet_ngay': {},
    }

    ngay_da_co = set()

    for r in rows:
        ngay = r['ngay']
        ngay_da_co.add(ngay)
        ket_qua['chi_tiet_ngay'][ngay] = dict(r)

        ma = r['ma_cong']
        if not ma and r.get('ca_ngay'):
            ma = cc_normalize_marker(r['ca_ngay'])

        if not ma:
            if r.get('trang_thai_cham_cong') == 'THIEU_GIO_RA':
                ket_qua['so_ngay_thieu_gio_ra'] += 1
            if r.get('trang_thai_vi_tri') == 'CHO_DUYET':
                ket_qua['so_ngay_cho_duyet_vi_tri'] += 1
            continue

        ket_qua['chi_tiet_ma_cong'][ma] = ket_qua['chi_tiet_ma_cong'].get(ma, 0) + 1

        tt = KY_HIEU_CHAM_CONG.get(ma, {})
        so_cong = tt.get('cong', 0)
        nhom = tt.get('nhom', '')

        if nhom == 'A':
            ket_qua['tong_cong'] += so_cong
            if ma in ('x', 'x/2'):
                ket_qua['cong_lam_viec'] += so_cong
            elif ma not in ('CN',):
                ket_qua['cong_huong_luong'] += so_cong
        elif nhom == 'B':
            ket_qua['cong_bhxh'] += 1

        tru_phep = tt.get('tru_phep', 0)
        if tru_phep:
            ket_qua['phep_da_dung'] += tru_phep

        gio_tc = float(r.get('gio_tang_ca') or 0)
        gio_tc_dem = float(r.get('gio_tang_ca_dem') or 0)
        loai_ngay_tc = r.get('loai_ngay_tang_ca') or 'THUONG'
        trang_thai_tc = r.get('trang_thai_duyet_tc')

        if gio_tc > 0:
            if trang_thai_tc == 'CHO_DUYET':
                ket_qua['tc_cho_duyet'][loai_ngay_tc] += gio_tc
                if gio_tc_dem > 0:
                    ket_qua['tc_cho_duyet']['DEM'] += gio_tc_dem
            else:
                ket_qua['tc'][loai_ngay_tc] += gio_tc
                ket_qua['tong_gio_tang_ca'] += gio_tc
                if gio_tc_dem > 0:
                    ket_qua['tc']['DEM'] += gio_tc_dem

        phut_tre = int(r.get('so_phut_di_tre') or 0)
        phut_som = int(r.get('so_phut_ve_som') or 0)
        if phut_tre > 0:
            ket_qua['tong_phut_di_tre'] += phut_tre
            ket_qua['so_lan_di_tre'] += 1
        if phut_som > 0:
            ket_qua['tong_phut_ve_som'] += phut_som
            ket_qua['so_lan_ve_som'] += 1

    cfg = get_cau_hinh_cham_cong_full()
    danh_sach_le = {x['ngay'] for x in (cfg.get('danh_sach_ngay_le') or [])}
    hom_nay = _date.today()
    for d in range(1, so_ngay_trong_thang + 1):
        ngay_d = _date(nam, thang, d)
        if ngay_d >= hom_nay:
            break
        if ngay_d.weekday() == 6:
            continue
        if ngay_d.strftime('%Y-%m-%d') in danh_sach_le:
            continue
        if ngay_d not in ngay_da_co:
            ket_qua['so_ngay_chua_cham'] += 1

    return ket_qua

def get_dia_diem_lam_viec():
    """Danh sách địa điểm làm việc của tenant: [{ten, lat, lng, ban_kinh}]"""
    try:
        ds = json.loads(get_cau_hinh('cc_dia_diem_lam_viec', '[]'))
        return ds if isinstance(ds, list) else []
    except Exception:
        return []


def luu_dia_diem_lam_viec(danh_sach):
    return set_cau_hinh('cc_dia_diem_lam_viec',
                        json.dumps(danh_sach, ensure_ascii=False),
                        'Danh sách địa điểm làm việc (GPS)')


def get_cau_hinh_gps():
    """Cấu hình chấm công qua điện thoại (GPS + đối chiếu khuôn mặt) theo tenant."""
    return {
        'bat_gps': get_cau_hinh('cc_bat_gps', '1') == '1',
        'bat_doi_chieu_mat': get_cau_hinh('cc_bat_doi_chieu_mat', '1') == '1',
        'dia_diem': get_dia_diem_lam_viec(),
    }


def tinh_khoang_cach_met(lat1, lng1, lat2, lng2):
    """Khoảng cách giữa 2 toạ độ, tính bằng mét (công thức Haversine)."""
    import math
    ban_kinh_trai_dat = 6371000.0
    p1, p2 = math.radians(float(lat1)), math.radians(float(lat2))
    d_phi = p2 - p1
    d_lambda = math.radians(float(lng2) - float(lng1))
    a = math.sin(d_phi / 2) ** 2 + math.cos(p1) * math.cos(p2) * math.sin(d_lambda / 2) ** 2
    return 2 * ban_kinh_trai_dat * math.asin(math.sqrt(a))


def kiem_tra_vi_tri_hop_le(lat, lng):
    """Kiểm tra toạ độ có nằm trong bán kính của 1 địa điểm làm việc nào không.
    Trả về (hop_le, ten_dia_diem_gan_nhat, khoang_cach_met)."""
    ds = get_dia_diem_lam_viec()
    if not ds:
        return False, None, None
    gan_nhat, kc_min = None, None
    for dd in ds:
        try:
            kc = tinh_khoang_cach_met(lat, lng, dd['lat'], dd['lng'])
        except Exception:
            continue
        if kc_min is None or kc < kc_min:
            gan_nhat, kc_min = dd, kc
        if kc <= float(dd.get('ban_kinh', 200)):
            return True, dd.get('ten', ''), kc
    return False, (gan_nhat.get('ten') if gan_nhat else None), kc_min


def nut_lay_toa_do(khoa="gps", nhan="📍 Lấy vị trí hiện tại của tôi"):
    """Hiện 1 nút xin quyền định vị của trình duyệt. Khi người dùng đồng ý, trang tự tải
    lại kèm toạ độ trong URL. Trả về (lat, lng, do_chinh_xac_met) hoặc (None, None, None).
    Dùng window.top (KHÔNG dùng window.parent) theo đúng bài học đã rút ra ở context.md."""
    components.html(f"""
        <div style="font-family: sans-serif;">
          <button onclick="layViTri_{khoa}()" style="
              background:#ff4b4b; color:#fff; border:none; border-radius:8px;
              padding:10px 18px; font-size:15px; cursor:pointer; width:100%;">
            {nhan}
          </button>
          <div id="tb_{khoa}" style="margin-top:8px; font-size:13px; color:#d33;"></div>
        </div>
        <script>
        function layViTri_{khoa}() {{
          var tb = document.getElementById('tb_{khoa}');
          if (!navigator.geolocation) {{
            tb.innerText = 'Thiết bị/trình duyệt không hỗ trợ định vị.';
            return;
          }}
          tb.innerText = 'Đang lấy vị trí, vui lòng chờ...';
          var thanhCong = function(p) {{
            var u = new URL(window.top.location.href);
            u.searchParams.set('gps_lat', p.coords.latitude);
            u.searchParams.set('gps_lng', p.coords.longitude);
            u.searchParams.set('gps_acc', p.coords.accuracy);
            window.top.location.href = u.toString();
          }};

          var thatBai = function(e) {{
            var ly_do = {{1: 'Bạn (hoặc trình duyệt) đã CHẶN quyền vị trí.',
                          2: 'Thiết bị không xác định được vị trí (thử bật GPS / đổi sang mạng WiFi).',
                          3: 'Quá thời gian chờ.'}}[e.code] || e.message;
            tb.innerHTML = '❌ Không lấy được vị trí: ' + ly_do +
              '<br>• Trên điện thoại: bật GPS/Vị trí rồi thử lại.' +
              '<br>• Trên máy tính: máy tính không có GPS, hãy lấy toạ độ từ Google Maps (bấm chuột phải lên bản đồ).' +
              '<br>• Nếu đã từng bấm "Chặn": bấm vào biểu tượng ổ khoá cạnh địa chỉ web → cho phép Vị trí → tải lại trang.';
          }};

          // Lần 1: dò nhanh, không đòi độ chính xác cao (máy tính không có GPS vẫn chạy được)
          navigator.geolocation.getCurrentPosition(
            thanhCong,
            function(e) {{
              if (e.code === 3) {{
                // Hết thời gian chờ → thử lại lần 2 với độ chính xác cao (điện thoại có GPS)
                tb.innerText = 'Đang thử lại bằng GPS chính xác cao...';
                navigator.geolocation.getCurrentPosition(
                  thanhCong, thatBai,
                  {{ enableHighAccuracy: true, timeout: 25000, maximumAge: 0 }}
                );
              }} else {{
                thatBai(e);
              }}
            }},
            {{ enableHighAccuracy: false, timeout: 8000, maximumAge: 60000 }}
          );
        }}
        </script>
    """, height=90)

    qp = st.query_params
    lat, lng, acc = qp.get('gps_lat'), qp.get('gps_lng'), qp.get('gps_acc')
    if lat and lng:
        try:
            return float(lat), float(lng), float(acc or 0)
        except Exception:
            return None, None, None
    return None, None, None

# === Cấu hình công thức tính lương đang áp dụng (khung sườn - salary/salary_{key}.py) ===
def get_cau_hinh_luong_key():
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT gia_tri FROM cau_hinh_he_thong WHERE ten_cau_hinh = 'luong_plugin_key'")
        r = c.fetchone()
        db.close()
        return r[0] if r and r[0] else 'salary_1'
    except:
        return 'salary_1'

def update_cau_hinh_luong_key(key):
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            INSERT INTO cau_hinh_he_thong (ten_cau_hinh, gia_tri, mo_ta)
            VALUES ('luong_plugin_key', %s, 'Công thức tính lương đang áp dụng (salary/salary_{key}.py)')
            ON CONFLICT (ten_cau_hinh) DO UPDATE SET gia_tri = EXCLUDED.gia_tri, updated_at = NOW()
        """, (key,))
        db.commit(); db.close()
        return True
    except:
        return False

def get_hdkt_prefix():
    """Lấy prefix đánh số HĐKT hiện tại (mặc định 'HĐKT'), cho phép admin tuỳ chỉnh."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT gia_tri FROM cau_hinh_he_thong WHERE ten_cau_hinh = 'hdkt_prefix'")
        result = c.fetchone()
        db.close()
        return result[0] if result and result[0] else 'HĐKT'
    except:
        return 'HĐKT'

def update_hdkt_prefix(prefix):
    """Cập nhật prefix đánh số HĐKT (chỉ Admin)."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            INSERT INTO cau_hinh_he_thong (ten_cau_hinh, gia_tri, mo_ta)
            VALUES ('hdkt_prefix', %s, 'Prefix đánh số Hợp đồng kinh tế')
            ON CONFLICT (ten_cau_hinh) DO UPDATE SET gia_tri = EXCLUDED.gia_tri, updated_at = NOW()
        """, (prefix,))
        db.commit(); db.close()
        return True
    except:
        return False

def preview_so_hdkt():
    """Chỉ XEM TRƯỚC số HĐKT tiếp theo (dạng stt/năm/Prefix-ma_cty) - KHÔNG ghi/tăng số max."""
    prefix = get_hdkt_prefix()
    ma_cty = st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL'
    nam_hien_tai = datetime.now().year
    so_max = get_so_max_cong_van('HDKT')
    so_moi = so_max + 1
    return f"{so_moi:02d}/{nam_hien_tai}/{prefix}-{ma_cty}", prefix

def generate_so_hdkt():
    """Sinh số HĐKT CHÍNH THỨC + cập nhật số max. CHỈ gọi khi bấm 'Lưu Hợp đồng'."""
    prefix = get_hdkt_prefix()
    ma_cty = st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL'
    nam_hien_tai = datetime.now().year
    so_max = get_so_max_cong_van('HDKT')
    so_moi = so_max + 1
    update_so_max_cong_van('HDKT', so_moi)
    return f"{so_moi:02d}/{nam_hien_tai}/{prefix}-{ma_cty}"

# === Ánh xạ mã loại công văn (bảng danh_muc_loai_cong_van) <-> mã loại nội bộ dùng cho cấu hình/đánh số ===
# Lưu ý: cột ma_loai trong danh_muc_loai_cong_van lưu chính là ký hiệu hiển thị (QĐ, CV, BC, TB, TTr),
# trong khi bảng cau_hinh_cong_van và bộ lọc loại công văn dùng mã nội bộ (QUYET_DINH, CONG_VAN, ...).
# Trước đây 2 bộ mã này KHÔNG khớp nhau khiến prefix luôn rơi về mặc định 'CV'.
MA_LOAI_TO_CODE = {
    'QĐ': 'QUYET_DINH',
    'CV': 'CONG_VAN',
    'BC': 'BAO_CAO',
    'TB': 'THONG_BAO',
    'TTr': 'TO_TRINH',
}
PREFIX_MAP = {
    'QUYET_DINH': 'QĐ',
    'CONG_VAN': 'CV',
    'BAO_CAO': 'BC',
    'THONG_BAO': 'TB',
    'TO_TRINH': 'TTr',
}

def chuan_hoa_loai_cong_van(ma_loai_raw):
    """Chuyển mã loại lấy từ dropdown (vd 'QĐ') về mã nội bộ chuẩn (vd 'QUYET_DINH').
    Nếu đã là mã chuẩn (hoặc không nhận diện được) thì trả về nguyên giá trị."""
    return MA_LOAI_TO_CODE.get(ma_loai_raw, ma_loai_raw)

# === Hàm XEM TRƯỚC số công văn (KHÔNG ghi/cập nhật cấu hình) ===
def preview_so_cong_van(loai_cv):
    """Chỉ tính toán số công văn TIẾP THEO sẽ được cấp để hiển thị cho user xem trước.
    Hàm này KHÔNG được phép ghi gì vào DB (không tăng số max), vì nó bị gọi lại
    ở MỌI lần rerun của trang (mỗi khi user click bất kỳ đâu trên màn hình)."""
    option = get_cv_danh_so_option()
    ma_cty = st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL'
    nam_hien_tai = datetime.now().year

    prefix = PREFIX_MAP.get(loai_cv, 'CV')

    loai_tim = 'CHUNG' if option == 'CHUNG' else loai_cv
    so_max = get_so_max_cong_van(loai_tim)
    so_moi = so_max + 1

    # Kiểm tra cấu hình đánh số theo kiểu dd/mm cho loại này
    kieu_ddmm = get_cv_kieu_ngay(loai_tim)
    if kieu_ddmm:
        # Kiểu dd/mm: số = ddmm (VD: 2807/2026/TB-CHL)
        ddmm = datetime.now().strftime('%d%m')
        so_cv = f"{ddmm}/{nam_hien_tai}/{prefix}-{ma_cty}"
    else:
        so_cv = f"{so_moi:02d}/{nam_hien_tai}/{prefix}-{ma_cty}"
    return so_cv

# === Hàm SINH SỐ CHÍNH THỨC + cập nhật số max (chỉ gọi khi user bấm "Lưu công văn đi") ===
def generate_so_cong_van(loai_cv):
    """Sinh số công văn chính thức theo cấu hình VÀ cập nhật số max trong DB.
    CHỈ được gọi trong handler của nút 'Lưu công văn đi', không được gọi khi
    chỉ render lại form (tránh nhảy số mỗi lần rerun)."""
    option = get_cv_danh_so_option()
    ma_cty = st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL'
    nam_hien_tai = datetime.now().year

    prefix = PREFIX_MAP.get(loai_cv, 'CV')

    # Xác định loại để lấy số max
    loai_tim = 'CHUNG' if option == 'CHUNG' else loai_cv
    so_max = get_so_max_cong_van(loai_tim)
    so_moi = so_max + 1

    # Cập nhật số max - CHỈ xảy ra khi hàm này được gọi (tức là khi Lưu)
    update_so_max_cong_van(loai_tim, so_moi)

    # Tạo số công văn — có thể theo kiểu dd/mm nếu tenant bật
    kieu_ddmm = get_cv_kieu_ngay(loai_tim)
    if kieu_ddmm:
        # Kiểu dd/mm: số = ddmm (VD: 2807/2026/TB-CHL), không dùng số thứ tự tăng dần
        ddmm = datetime.now().strftime('%d%m')
        so_cv = f"{ddmm}/{nam_hien_tai}/{prefix}-{ma_cty}"
    else:
        so_cv = f"{so_moi:02d}/{nam_hien_tai}/{prefix}-{ma_cty}"
    return so_cv

# === Hàm upload file cho công văn ===
def upload_cong_van_file(uploaded_file, folder_name):
    """Upload file công văn lên Supabase Storage"""
    if not uploaded_file:
        return None
    
    sb = get_supabase_storage()
    if not sb:
        return None
    
    try:
        safe_name = sanitize_storage_filename(uploaded_file.name)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        path = f"cong_van/{folder_name}/{timestamp}_{safe_name}"
        
        # Upload file
        result = sb.storage.from_(SUPABASE_BUCKET).upload(
            path=path,
            file=uploaded_file.getvalue(),
            file_options={"content-type": uploaded_file.type or "application/octet-stream"}
        )
        return path
    except Exception as e:
        print(f"Lỗi upload file: {e}")
        return None

# === Hàm lấy danh sách công văn đến ===
def get_cong_van_den(tu_ngay=None, den_ngay=None, search_text=None):
    """Lấy danh sách công văn đến với bộ lọc"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        sql = "SELECT * FROM cong_van_den WHERE 1=1"
        params = []
        
        if tu_ngay:
            sql += " AND ngay_den >= %s"
            params.append(tu_ngay)
        if den_ngay:
            sql += " AND ngay_den <= %s"
            params.append(den_ngay)
        if search_text:
            sql += """ AND (so_cong_van ILIKE %s OR tieu_de ILIKE %s 
                     OR co_quan_phat_hanh ILIKE %s OR trich_yeu ILIKE %s OR ma_vach_buu_dien ILIKE %s)"""
            search_pattern = f"%{search_text}%"
            params.extend([search_pattern] * 5)
        
        sql += " ORDER BY ngay_den DESC, id DESC"
        c.execute(sql, tuple(params))
        result = c.fetchall()
        db.close()
        return result
    except Exception as e:
        print(f"Lỗi lấy công văn đến: {e}")
        return []

# === Hàm lấy danh sách công văn đi ===
def get_cong_van_di(tu_ngay=None, den_ngay=None, search_text=None, loai_cv=None):
    """Lấy danh sách công văn đi với bộ lọc"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        sql = "SELECT * FROM cong_van_di WHERE 1=1"
        params = []
        
        if tu_ngay:
            sql += " AND ngay_phat_hanh >= %s"
            params.append(tu_ngay)
        if den_ngay:
            sql += " AND ngay_phat_hanh <= %s"
            params.append(den_ngay)
        if loai_cv:
            sql += " AND loai_cong_van = %s"
            params.append(loai_cv)
        if search_text:
            sql += """ AND (so_cong_van ILIKE %s OR tieu_de ILIKE %s 
                     OR phong_phat_hanh ILIKE %s OR trich_yeu ILIKE %s OR ma_vach_buu_dien ILIKE %s)"""
            search_pattern = f"%{search_text}%"
            params.extend([search_pattern] * 5)
        
        sql += " ORDER BY ngay_phat_hanh DESC, id DESC"
        c.execute(sql, tuple(params))
        result = c.fetchall()
        db.close()
        return result
    except Exception as e:
        print(f"Lỗi lấy công văn đi: {e}")
        return []

# === Hàm lấy danh sách hợp đồng kinh tế ===
def get_hop_dong_kinh_te(tu_ngay=None, den_ngay=None, search_text=None):
    """Lấy danh sách hợp đồng kinh tế với bộ lọc"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        sql = "SELECT * FROM hop_dong_kinh_te WHERE 1=1"
        params = []
        
        if tu_ngay:
            sql += " AND ngay_ky >= %s"
            params.append(tu_ngay)
        if den_ngay:
            sql += " AND ngay_ky <= %s"
            params.append(den_ngay)
        if search_text:
            sql += """ AND (so_hop_dong ILIKE %s OR ten_doi_tac ILIKE %s 
                     OR trich_yeu ILIKE %s)"""
            search_pattern = f"%{search_text}%"
            params.extend([search_pattern] * 3)
        
        sql += " ORDER BY ngay_ky DESC, id DESC"
        c.execute(sql, tuple(params))
        result = c.fetchall()
        db.close()
        return result
    except Exception as e:
        print(f"Lỗi lấy hợp đồng kinh tế: {e}")
        return []

# === Hàm xuất Excel công văn ===
def export_cong_van_excel(data, ten_file, headers, col_widths=None, title=None):
    """Xuất dữ liệu ra file Excel với format chuẩn"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Danh sách"
    
    # Border
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Header fill
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    
    # Tiêu đề
    if title:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        ws['A1'] = title
        ws['A1'].font = Font(bold=True, size=14, name='Times New Roman')
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # Dòng trống
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
        ws['A2'] = f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        ws['A2'].font = Font(size=10, name='Times New Roman', italic=True)
        ws['A2'].alignment = Alignment(horizontal='center')
        start_row = 4
    else:
        start_row = 1
    
    # Header
    header_row = start_row
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.font = Font(bold=True, size=11, name='Times New Roman', color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
    
    # Dữ liệu
    data_row = header_row + 1
    for idx, row in enumerate(data):
        for col_idx, key in enumerate(headers, 1):
            value = row.get(key, '')
            # Format ngày tháng
            if 'ngay' in key.lower() and value:
                if hasattr(value, 'strftime'):
                    value = value.strftime('%d/%m/%Y')
            cell = ws.cell(row=data_row + idx, column=col_idx, value=value)
            cell.font = Font(size=10, name='Times New Roman')
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center' if col_idx <= 3 else 'left', vertical='center')
    
    # Footer
    footer_row = data_row + len(data) + 2
    ws.merge_cells(start_row=footer_row, start_column=1, end_row=footer_row, end_column=len(headers))
    ws.cell(row=footer_row, column=1, value=f"Tổng cộng: {len(data)} bản ghi")
    ws.cell(row=footer_row, column=1).font = Font(bold=True, size=11, name='Times New Roman')
    
    # Độ rộng cột
    if col_widths:
        for idx, width in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(idx)].width = width
    
    # Lưu file
    wb.save(ten_file)
    return ten_file

# === UI: Quản lý Công văn & HĐ kinh tế ===
def show_quan_ly_cong_van():
    """Hiển thị giao diện Quản lý Công văn & HĐ kinh tế"""
    st.title("📄 Quản lý Công văn & HĐ kinh tế")
    
    # Khởi tạo bảng nếu chưa có
    init_cong_van_tables()
    
    # Kiểm tra quyền
    role = st.session_state.get('role', '')
    if role not in ['admin', 'van_thu']:
        st.warning("🔒 Chỉ Admin và Văn thư mới có quyền truy cập module này!")
        st.stop()
    
    # Cấu hình (chỉ Admin)
    if role == 'admin':
        with st.expander("⚙️ Cấu hình đánh số Công văn & HĐKT", expanded=False):
            st.markdown("**Cấu hình cách đánh số công văn đi**")
            
            current_option = get_cv_danh_so_option()
            new_option = st.radio(
                "Chọn phương án đánh số:",
                options=['CHUNG', 'RIENG'],
                index=0 if current_option == 'CHUNG' else 1,
                format_func=lambda x: "📌 Số chung cho tất cả loại công văn" if x == 'CHUNG' else "📌 Mỗi loại công văn có số riêng",
                key="cv_option_radio"
            )
            
            if new_option != current_option:
                if st.button("✅ Cập nhật cấu hình", type="primary", disabled=not can_edit()):
                    if update_cv_danh_so_option(new_option):
                        st.success(f"✅ Đã cập nhật cấu hình sang: {new_option}")
                        st.cache_data.clear()
                        st.rerun()
                    else:
                        st.error("❌ Cập nhật thất bại!")
            
            st.divider()
            st.markdown("**📊 Trạng thái đánh số hiện tại**")
            
            # Lấy dữ liệu từ database
            db = st.session_state.db_engine.get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            
            # Xác định loại cần hiển thị dựa trên option
            option = get_cv_danh_so_option()
            if option == 'CHUNG':
                # Chỉ hiển thị CHUNG
                c.execute("""
                    SELECT loai, so_max, prefix, nam_hien_tai, updated_at 
                    FROM cau_hinh_cong_van 
                    WHERE loai = 'CHUNG'
                    ORDER BY loai, nam_hien_tai
                """)
            else:
                # Hiển thị các loại riêng (không bao gồm CHUNG và QĐ trùng)
                c.execute("""
                    SELECT loai, so_max, prefix, nam_hien_tai, updated_at 
                    FROM cau_hinh_cong_van 
                    WHERE loai IN ('CONG_VAN', 'QUYET_DINH', 'BAO_CAO', 'THONG_BAO', 'TO_TRINH')
                    ORDER BY 
                        CASE loai
                            WHEN 'QUYET_DINH' THEN 1
                            WHEN 'CONG_VAN' THEN 2
                            WHEN 'BAO_CAO' THEN 3
                            WHEN 'THONG_BAO' THEN 4
                            WHEN 'TO_TRINH' THEN 5
                        END
                """)
            
            configs = c.fetchall()
            db.close()
            
            if configs:
                df_config = pd.DataFrame(configs)
                df_config['updated_at'] = df_config['updated_at'].apply(
                    lambda x: x.strftime('%d/%m/%Y %H:%M') if x else ''
                )
                
                # Đổi tên loại cho đẹp
                loai_name_map = {
                    'CHUNG': 'CHUNG (Tất cả loại)',
                    'QUYET_DINH': 'QUYẾT ĐỊNH',
                    'CONG_VAN': 'CÔNG VĂN',
                    'BAO_CAO': 'BÁO CÁO',
                    'THONG_BAO': 'THÔNG BÁO',
                    'TO_TRINH': 'TỜ TRÌNH'
                }
                df_config['loai'] = df_config['loai'].map(loai_name_map)
                df_config.columns = ['Loại', 'Số hiện tại', 'Prefix', 'Năm', 'Cập nhật lúc']
                st.dataframe(df_config, width='stretch', hide_index=True)
            else:
                st.info("Chưa có dữ liệu cấu hình. Hãy tạo công văn đi đầu tiên để khởi tạo.")
            
            st.divider()
            st.markdown("**🔄 Đặt lại số**")
            col_reset1, col_reset2 = st.columns(2)
            with col_reset1:
                # Xác định danh sách loại cho dropdown dựa trên option
                if option == 'CHUNG':
                    loai_list = ['CHUNG']
                    loai_display = {'CHUNG': 'CHUNG (Tất cả loại)'}
                else:
                    loai_list = ['QUYET_DINH', 'CONG_VAN', 'BAO_CAO', 'THONG_BAO', 'TO_TRINH']
                    loai_display = {
                        'QUYET_DINH': 'QUYẾT ĐỊNH',
                        'CONG_VAN': 'CÔNG VĂN',
                        'BAO_CAO': 'BÁO CÁO',
                        'THONG_BAO': 'THÔNG BÁO',
                        'TO_TRINH': 'TỜ TRÌNH'
                    }
                
                selected_loai_display = st.selectbox(
                    "Chọn loại cần đặt lại:",
                    [loai_display.get(l, l) for l in loai_list],
                    key="reset_loai_display"
                )
                
                # Lấy lại mã loại thực tế
                if option == 'CHUNG':
                    loai_reset = 'CHUNG'
                else:
                    # Tìm key từ display name
                    for key, value in loai_display.items():
                        if value == selected_loai_display:
                            loai_reset = key
                            break
                
                # Lấy số hiện tại để hiển thị
                current_so = get_so_max_cong_van(loai_reset)
                st.caption(f"📌 Số hiện tại: **{current_so}**")
                
                so_moi = st.number_input(
                    "Số bắt đầu mới:", 
                    min_value=0, 
                    value=current_so, 
                    step=1, 
                    key="reset_so"
                )
            
            with col_reset2:
                # Nguyên tắc: cấu hình (so_max) CHỈ được cập nhật khi user bấm nút "Đặt lại số",
                # và luôn lấy đúng giá trị đang có trong ô "Số bắt đầu mới:" tại thời điểm bấm.
                # (Trước đây dùng checkbox xác nhận lồng bên trong if st.button(...) — nhưng vì
                # Streamlit rerun lại toàn bộ script sau mỗi tương tác, việc tick checkbox ở lần
                # rerun sau đó lại rơi vào nhánh st.button() == False nên state bị "treo" và
                # không đáng tin cậy. Dùng session_state để giữ yêu cầu đặt lại qua các lần rerun.)
                st.markdown("&nbsp;")  # căn cho nút thẳng hàng với ô nhập bên trái
                if st.button("🔄 Đặt lại số", type="secondary", key="btn_dat_lai_so"):
                    st.session_state['cv_pending_reset'] = {'loai': loai_reset, 'so_moi': so_moi}

                pending = st.session_state.get('cv_pending_reset')
                if pending and pending['loai'] == loai_reset:
                    st.warning(
                        f"⚠️ Xác nhận đặt lại số cho **{selected_loai_display}** "
                        f"từ **{current_so}** thành **{pending['so_moi']}**?"
                    )
                    col_confirm, col_cancel = st.columns(2)
                    with col_confirm:
                        if st.button("✅ Xác nhận đặt lại số", type="primary", key="confirm_dat_lai_so", disabled=not can_edit()):
                            if update_so_max_cong_van(pending['loai'], pending['so_moi']):
                                st.success(
                                    f"✅ Đã đặt lại số cho {selected_loai_display} "
                                    f"từ {current_so} thành {pending['so_moi']}"
                                )
                                del st.session_state['cv_pending_reset']
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error("❌ Đặt lại số thất bại!")
                    with col_cancel:
                        if st.button("✖️ Hủy", key="cancel_dat_lai_so"):
                            del st.session_state['cv_pending_reset']
                            st.rerun()

            st.divider()
            st.markdown("**📅 Đánh số theo kiểu dd/mm**")
            st.caption("Bật tùy chọn này nếu muốn số công văn theo ngày phát hành thay vì số thứ tự. "
                       "VD: `2807/2026/TB-CHL` (ngày 28/07) thay vì `01/2026/TB-CHL`")

            if option == 'CHUNG':
                loai_ddmm_list = ['CHUNG']
                loai_ddmm_display = {'CHUNG': 'CHUNG (Tất cả loại)'}
            else:
                loai_ddmm_list = ['QUYET_DINH', 'CONG_VAN', 'BAO_CAO', 'THONG_BAO', 'TO_TRINH']
                loai_ddmm_display = {
                    'QUYET_DINH': 'QUYẾT ĐỊNH',
                    'CONG_VAN': 'CÔNG VĂN',
                    'BAO_CAO': 'BÁO CÁO',
                    'THONG_BAO': 'THÔNG BÁO',
                    'TO_TRINH': 'TỜ TRÌNH'
                }

            for loai_key in loai_ddmm_list:
                col_ddmm1, col_ddmm2 = st.columns([3, 1])
                kieu_hien_tai = get_cv_kieu_ngay(loai_key)
                with col_ddmm1:
                    st.write(f"**{loai_ddmm_display.get(loai_key, loai_key)}:** {'📅 dd/mm/yyyy' if kieu_hien_tai else '📆 yyyy (mặc định)'}")
                with col_ddmm2:
                    new_kieu = st.toggle(
                        "dd/mm",
                        value=kieu_hien_tai,
                        key=f"toggle_ddmm_{loai_key}",
                        disabled=not can_edit()
                    )
                    if new_kieu != kieu_hien_tai:
                        if update_cv_kieu_ngay(loai_key, 'DD_MM' if new_kieu else 'NAM'):
                            st.cache_data.clear()
                            st.rerun()
            st.divider()
            st.markdown("**📑 Cấu hình đánh số Hợp đồng kinh tế (HĐKT)**")
            st.caption("Số HĐKT tự sinh theo mẫu: **stt/năm/Prefix-ma_cty** (VD: 04/2026/HĐKT-CHL)")
            prefix_hdkt_hien_tai = get_hdkt_prefix()
            prefix_hdkt_moi = st.text_input(
                "Prefix đánh số HĐKT:", value=prefix_hdkt_hien_tai, key="hdkt_prefix_input"
            )
            if st.button("✅ Cập nhật prefix HĐKT", key="btn_update_hdkt_prefix", disabled=not can_edit()):
                if prefix_hdkt_moi.strip() and update_hdkt_prefix(prefix_hdkt_moi.strip()):
                    st.success(f"✅ Đã cập nhật prefix HĐKT sang: {prefix_hdkt_moi.strip()}")
                    st.cache_data.clear()
                    st.rerun()
                else:
                    st.error("❌ Cập nhật thất bại!")
    
    # Tabs
    tab1, tab2, tab3 = st.tabs(["📥 Công văn đến", "📤 Công văn đi", "📑 Hợp đồng kinh tế"])
    
    # === TAB 1: CÔNG VĂN ĐẾN ===
    with tab1:
        st.subheader("📥 Quản lý Công văn đến")

        thao_tac_cvden = st.radio(
            "Chọn thao tác:",
            ["➕ Thêm CV đến mới", "📂 Nhập lại CV đến cũ", "🔍 Tra cứu"],
            horizontal=True, key="cv_den_thao_tac"
        )

        # ── THÊM CV ĐẾN MỚI ──
        if thao_tac_cvden == "➕ Thêm CV đến mới":
            with st.form("add_cong_van_den"):
                col1, col2 = st.columns(2)
                with col1:
                    so_cv = st.text_input("Số công văn *", placeholder="VD: 123/BQP-2026")
                    co_quan = st.text_input("Cơ quan phát hành *", placeholder="VD: Bộ Quốc phòng")
                    ngay_den = st.date_input("Ngày đến *", value=date.today())
                    ma_vach = st.text_input("📦 Mã vạch Bưu điện", placeholder="VD: EV123456789VN")
                with col2:
                    tieu_de = st.text_input("Tiêu đề *", placeholder="Nhập tiêu đề công văn...")
                    trich_yeu = st.text_area("Trích yếu", placeholder="Tóm tắt nội dung chính...", height=80)
                    ghi_chu = st.text_area("Ghi chú", height=60)
                
                uploaded_file = st.file_uploader("📎 Upload file", type=['pdf', 'doc', 'docx', 'jpg', 'png', 'jpeg'], key="cv_den_upload")
                
                col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
                with col_btn2:
                    if st.form_submit_button("💾 Lưu công văn đến", width='stretch', type="primary", disabled=not can_edit()):
                        if not so_cv or not co_quan or not tieu_de:
                            st.error("⚠️ Vui lòng nhập đầy đủ các trường bắt buộc (*)")
                        else:
                            try:
                                file_url = None
                                if uploaded_file:
                                    file_url = upload_cong_van_file(uploaded_file, "den")
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("""
                                    INSERT INTO cong_van_den (so_cong_van, co_quan_phat_hanh, ngay_den, 
                                    tieu_de, trich_yeu, file_url, ghi_chu, nguoi_tao, ma_vach_buu_dien)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """, (so_cv, co_quan, ngay_den, tieu_de, trich_yeu, file_url, ghi_chu, 
                                      st.session_state.username, ma_vach))
                                db.commit()
                                db.close()
                                st.success(f"✅ Đã thêm công văn đến: {so_cv}")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi: {e}")

        # ── NHẬP LẠI CV ĐẾN CŨ (form giống hệt, chỉ khác key) ──
        elif thao_tac_cvden == "📂 Nhập lại CV đến cũ":
            st.caption("Nhập lại các công văn đến đã tiếp nhận trước đây vào hệ thống.")
            with st.form("add_cong_van_den_cu"):
                col1, col2 = st.columns(2)
                with col1:
                    so_cv_cu = st.text_input("Số công văn *", placeholder="VD: 456/UBND-2024", key="cv_den_cu_socv")
                    co_quan_cu = st.text_input("Cơ quan phát hành *", placeholder="VD: UBND tỉnh Nghệ An", key="cv_den_cu_coquan")
                    ngay_den_cu = st.date_input("Ngày đến *", value=date.today(), key="cv_den_cu_ngay")
                    ma_vach_cu = st.text_input("📦 Mã vạch Bưu điện", placeholder="VD: EV123456789VN", key="cv_den_cu_mavach")
                with col2:
                    tieu_de_cu = st.text_input("Tiêu đề *", placeholder="Nhập tiêu đề công văn...", key="cv_den_cu_tieude")
                    trich_yeu_cu = st.text_area("Trích yếu", placeholder="Tóm tắt nội dung chính...", height=80, key="cv_den_cu_trichyeu")
                    ghi_chu_cu = st.text_area("Ghi chú", height=60, key="cv_den_cu_ghichu")
                
                uploaded_file_cu = st.file_uploader("📎 Upload file", type=['pdf', 'doc', 'docx', 'jpg', 'png', 'jpeg'], key="cv_den_cu_upload")
                
                col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
                with col_btn2:
                    if st.form_submit_button("💾 Lưu công văn đến cũ", width='stretch', type="primary", disabled=not can_edit()):
                        if not so_cv_cu or not co_quan_cu or not tieu_de_cu:
                            st.error("⚠️ Vui lòng nhập đầy đủ các trường bắt buộc (*)")
                        else:
                            try:
                                file_url_cu = None
                                if uploaded_file_cu:
                                    file_url_cu = upload_cong_van_file(uploaded_file_cu, "den")
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("""
                                    INSERT INTO cong_van_den (so_cong_van, co_quan_phat_hanh, ngay_den, 
                                    tieu_de, trich_yeu, file_url, ghi_chu, nguoi_tao, ma_vach_buu_dien)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """, (so_cv_cu, co_quan_cu, ngay_den_cu, tieu_de_cu, trich_yeu_cu, file_url_cu, ghi_chu_cu, 
                                      st.session_state.username, ma_vach_cu))
                                db.commit()
                                db.close()
                                st.success(f"✅ Đã nhập lại công văn đến cũ: {so_cv_cu}")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi: {e}")

        # ── TRA CỨU ──
        else:
            col_search1, col_search2, col_search3 = st.columns([2, 1, 1])
            with col_search1:
                search_text_cv_den = st.text_input("🔍 Tìm kiếm", placeholder="Theo số, tiêu đề, cơ quan...", key="search_cv_den")
            with col_search2:
                tu_ngay_cv_den = st.date_input("Từ ngày", value=None, key="tu_ngay_cv_den")
            with col_search3:
                den_ngay_cv_den = st.date_input("Đến ngày", value=None, key="den_ngay_cv_den")
            
            data_cv_den = get_cong_van_den(tu_ngay_cv_den, den_ngay_cv_den, search_text_cv_den)
            
            if data_cv_den:
                df_cv_den = pd.DataFrame(data_cv_den)
                for col in ['ngay_den', 'created_at', 'updated_at']:
                    if col in df_cv_den.columns:
                        df_cv_den[col] = df_cv_den[col].apply(format_date)
                
                display_cols = ['so_cong_van', 'co_quan_phat_hanh', 'ngay_den', 'tieu_de', 'trich_yeu', 'ma_vach_buu_dien', 'file_url', 'ghi_chu']
                available_cols = [c for c in display_cols if c in df_cv_den.columns]
                df_display = df_cv_den[available_cols]
                
                col_map = {
                    'so_cong_van': 'Số công văn',
                    'co_quan_phat_hanh': 'Cơ quan phát hành',
                    'ngay_den': 'Ngày đến',
                    'tieu_de': 'Tiêu đề',
                    'trich_yeu': 'Trích yếu',
                    'ma_vach_buu_dien': 'Mã vạch BĐ',
                    'file_url': 'File',
                    'ghi_chu': 'Ghi chú'
                }
                df_display.rename(columns=col_map, inplace=True)
                
                st.caption(f"📌 Tổng số: {len(data_cv_den)} công văn đến")
                st.dataframe(df_display, width='stretch', hide_index=True, height=400)
                
                col_export1, col_export2, col_export3 = st.columns([1, 2, 1])
                with col_export2:
                    if st.button("📥 Xuất Excel công văn đến", width='stretch', type="primary"):
                        headers = ['so_cong_van', 'co_quan_phat_hanh', 'ngay_den', 'tieu_de', 'trich_yeu', 'ghi_chu']
                        col_widths = [15, 25, 12, 35, 30, 25]
                        title = f"BÁO CÁO CÔNG VĂN ĐẾN (Từ {tu_ngay_cv_den or '...'} đến {den_ngay_cv_den or '...'})"
                        filename = f"Cong_van_den_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                        
                        excel_data = []
                        for row in data_cv_den:
                            excel_row = {}
                            for key in headers:
                                val = row.get(key)
                                if key == 'ngay_den' and val:
                                    val = format_date(val)
                                excel_row[key] = val
                            excel_data.append(excel_row)
                        
                        export_cong_van_excel(excel_data, filename, headers, col_widths, title)
                        
                        with open(filename, "rb") as f:
                            st.download_button(
                                label="📥 TẢI FILE EXCEL",
                                data=f,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                width='stretch'
                            )
                        st.success(f"✅ Đã xuất {len(data_cv_den)} công văn đến")
                        st.cache_data.clear()
            else:
                st.info("📭 Không có công văn đến nào")
    
    # === TAB 2: CÔNG VĂN ĐI ===
    with tab2:
        st.subheader("📤 Quản lý Công văn đi")

        thao_tac_cvd = st.radio(
            "Chọn thao tác:",
            ["➕ Thêm CV đi mới", "📂 Nhập lại CV đi cũ", "🔍 Tra cứu"],
            horizontal=True, key="cv_di_thao_tac"
        )

        # ── THÊM CV ĐI MỚI (số tự sinh) ──
        if thao_tac_cvd == "➕ Thêm CV đi mới":
            db_loai = st.session_state.db_engine.get_connection()
            c_loai = db_loai.cursor()
            c_loai.execute("SELECT ma_loai, ten_loai FROM danh_muc_loai_cong_van WHERE trang_thai = TRUE ORDER BY thu_tu")
            loai_cv_list = c_loai.fetchall()
            db_loai.close()

            loai_options = {f"{loai[1]} ({loai[0]})": loai[0] for loai in loai_cv_list}
            selected_loai = st.selectbox("Loại công văn *", list(loai_options.keys()), key="cv_di_loai")
            loai_cv = chuan_hoa_loai_cong_van(loai_options[selected_loai])

            so_cv_xem_truoc = preview_so_cong_van(loai_cv)
            prefix_hien_tai = PREFIX_MAP.get(loai_cv, 'CV')
            st.info(
                f"📄 **Số công văn dự kiến:** `{so_cv_xem_truoc}` (Prefix: **{prefix_hien_tai}**) "
                f"— số chính thức sẽ được cấp khi bấm **Lưu công văn đi**"
            )

            with st.form("add_cong_van_di"):
                col1, col2 = st.columns(2)
                with col1:
                    ds_phong_ban_cv = get_phong_ban_options()
                    phong_phat_hanh = st.selectbox("Phòng phát hành *", [""] + ds_phong_ban_cv, key="cv_di_phong_phat_hanh")
                    ngay_phat_hanh = st.date_input("Ngày phát hành *", value=date.today())
                    ma_vach = st.text_input("📦 Mã vạch Bưu điện", placeholder="VD: EV123456789VN")
                with col2:
                    tieu_de = st.text_input("Tiêu đề *", placeholder="Nhập tiêu đề công văn...")
                    trich_yeu = st.text_area("Trích yếu", placeholder="Tóm tắt nội dung chính...", height=80)
                    ghi_chu = st.text_area("Ghi chú", height=60)
                
                uploaded_file = st.file_uploader("📎 Upload file", type=['pdf', 'doc', 'docx', 'jpg', 'png', 'jpeg'], key="cv_di_upload")
                
                col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
                with col_btn2:
                    if st.form_submit_button("💾 Lưu công văn đi", width='stretch', type="primary", disabled=not can_edit()):
                        if not phong_phat_hanh or not tieu_de:
                            st.error("⚠️ Vui lòng nhập đầy đủ các trường bắt buộc (*)")
                        else:
                            try:
                                file_url = None
                                if uploaded_file:
                                    file_url = upload_cong_van_file(uploaded_file, "di")
                                so_cv_chinh_thuc = generate_so_cong_van(loai_cv)
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("""
                                    INSERT INTO cong_van_di (so_cong_van, phong_phat_hanh, ngay_phat_hanh, 
                                    tieu_de, trich_yeu, file_url, loai_cong_van, ghi_chu, nguoi_tao, ma_vach_buu_dien)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """, (so_cv_chinh_thuc, phong_phat_hanh, ngay_phat_hanh, tieu_de, trich_yeu, 
                                      file_url, loai_cv, ghi_chu, st.session_state.username, ma_vach))
                                db.commit()
                                db.close()
                                st.success(f"✅ Đã thêm công văn đi: {so_cv_chinh_thuc}")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi: {e}")

        # ── NHẬP LẠI CV ĐI CŨ (số nhập tay) ──
        elif thao_tac_cvd == "📂 Nhập lại CV đi cũ":
            st.caption("Nhập lại các công văn đi đã phát hành trước đây. "
                       "Số công văn do bạn tự nhập, hệ thống **không** tự sinh số mới.")
            db_loai2 = st.session_state.db_engine.get_connection()
            c_loai2 = db_loai2.cursor()
            c_loai2.execute("SELECT ma_loai, ten_loai FROM danh_muc_loai_cong_van WHERE trang_thai = TRUE ORDER BY thu_tu")
            loai_cv_list2 = c_loai2.fetchall()
            db_loai2.close()
            loai_options2 = {f"{loai[1]} ({loai[0]})": loai[0] for loai in loai_cv_list2}
            selected_loai2 = st.selectbox("Loại công văn *", list(loai_options2.keys()), key="cv_di_cu_loai")
            loai_cv_cu = chuan_hoa_loai_cong_van(loai_options2[selected_loai2])

            with st.form("add_cong_van_di_cu"):
                so_cv_cu = st.text_input("Số công văn *", placeholder="VD: 05/2025/QĐ-CHL")
                col1, col2 = st.columns(2)
                with col1:
                    ds_phong_ban_cv2 = get_phong_ban_options()
                    phong_phat_hanh_cu = st.selectbox("Phòng phát hành *", [""] + ds_phong_ban_cv2, key="cv_di_cu_phong")
                    ngay_phat_hanh_cu = st.date_input("Ngày phát hành *", value=date.today(), key="cv_di_cu_ngay")
                    ma_vach_cu = st.text_input("📦 Mã vạch Bưu điện", placeholder="VD: EV123456789VN", key="cv_di_cu_mavach")
                with col2:
                    tieu_de_cu = st.text_input("Tiêu đề *", placeholder="Nhập tiêu đề công văn...", key="cv_di_cu_tieude")
                    trich_yeu_cu = st.text_area("Trích yếu", placeholder="Tóm tắt nội dung chính...", height=80, key="cv_di_cu_trichyeu")
                    ghi_chu_cu = st.text_area("Ghi chú", height=60, key="cv_di_cu_ghichu")

                uploaded_file_cu = st.file_uploader("📎 Upload file", type=['pdf', 'doc', 'docx', 'jpg', 'png', 'jpeg'], key="cv_di_cu_upload")

                col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
                with col_btn2:
                    if st.form_submit_button("💾 Lưu công văn đi cũ", width='stretch', type="primary", disabled=not can_edit()):
                        if not so_cv_cu.strip() or not phong_phat_hanh_cu or not tieu_de_cu:
                            st.error("⚠️ Vui lòng nhập đầy đủ: Số công văn, Phòng phát hành, Tiêu đề")
                        else:
                            try:
                                file_url_cu = None
                                if uploaded_file_cu:
                                    file_url_cu = upload_cong_van_file(uploaded_file_cu, "di")
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("""
                                    INSERT INTO cong_van_di (so_cong_van, phong_phat_hanh, ngay_phat_hanh,
                                    tieu_de, trich_yeu, file_url, loai_cong_van, ghi_chu, nguoi_tao, ma_vach_buu_dien)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """, (so_cv_cu.strip(), phong_phat_hanh_cu, ngay_phat_hanh_cu, tieu_de_cu,
                                      trich_yeu_cu, file_url_cu, loai_cv_cu, ghi_chu_cu,
                                      st.session_state.username, ma_vach_cu))
                                db.commit()
                                db.close()
                                st.success(f"✅ Đã nhập lại công văn đi cũ: {so_cv_cu.strip()}")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi: {e}")

        # ── TRA CỨU ──
        else:
            col_search1, col_search2 = st.columns([2, 1])
            with col_search1:
                search_text_cv_di = st.text_input("🔍 Tìm kiếm", placeholder="Theo số, tiêu đề, phòng...", key="search_cv_di")
            with col_search2:
                loai_filter = st.selectbox("Loại", ["Tất cả", "Quyết định", "Công văn", "Báo cáo", "Thông báo", "Tờ trình"], key="loai_filter_cv_di")
            col_search3, col_search4 = st.columns(2)
            with col_search3:
                tu_ngay_cv_di = st.date_input("Từ ngày", value=None, key="tu_ngay_cv_di")
            with col_search4:
                den_ngay_cv_di = st.date_input("Đến ngày", value=None, key="den_ngay_cv_di")
            
            loai_map = {
                "Tất cả": None,
                "Quyết định": "QUYET_DINH",
                "Công văn": "CONG_VAN",
                "Báo cáo": "BAO_CAO",
                "Thông báo": "THONG_BAO",
                "Tờ trình": "TO_TRINH"
            }
            data_cv_di = get_cong_van_di(tu_ngay_cv_di, den_ngay_cv_di, search_text_cv_di, loai_map.get(loai_filter))
            
            if data_cv_di:
                df_cv_di = pd.DataFrame(data_cv_di)
                for col in ['ngay_phat_hanh', 'created_at', 'updated_at']:
                    if col in df_cv_di.columns:
                        df_cv_di[col] = df_cv_di[col].apply(format_date)
                
                display_cols = ['so_cong_van', 'loai_cong_van', 'phong_phat_hanh', 'ngay_phat_hanh', 'tieu_de', 'trich_yeu', 'ma_vach_buu_dien', 'file_url', 'ghi_chu']
                available_cols = [c for c in display_cols if c in df_cv_di.columns]
                df_display = df_cv_di[available_cols]
                
                loai_name_map = {
                    'QUYET_DINH': 'Quyết định',
                    'CONG_VAN': 'Công văn',
                    'BAO_CAO': 'Báo cáo',
                    'THONG_BAO': 'Thông báo',
                    'TO_TRINH': 'Tờ trình'
                }
                df_display['loai_cong_van'] = df_display['loai_cong_van'].map(loai_name_map)
                
                col_map = {
                    'so_cong_van': 'Số công văn',
                    'loai_cong_van': 'Loại',
                    'phong_phat_hanh': 'Phòng phát hành',
                    'ngay_phat_hanh': 'Ngày phát hành',
                    'tieu_de': 'Tiêu đề',
                    'trich_yeu': 'Trích yếu',
                    'ma_vach_buu_dien': 'Mã vạch BĐ',
                    'file_url': 'File',
                    'ghi_chu': 'Ghi chú'
                }
                df_display.rename(columns=col_map, inplace=True)
                
                st.caption(f"📌 Tổng số: {len(data_cv_di)} công văn đi")
                st.dataframe(df_display, width='stretch', hide_index=True, height=400)

                with st.expander("✏️ Sửa / 🗑️ Xóa công văn đi", expanded=False):
                    tuy_chon_cvd = {f"{r['so_cong_van']} - {r.get('tieu_de') or ''}": r for r in data_cv_di}
                    chon_cvd = st.selectbox("Chọn công văn đi:", ["-- Chọn --"] + list(tuy_chon_cvd.keys()), key="chon_sua_cvd")
                    if chon_cvd != "-- Chọn --":
                        bg_sua = tuy_chon_cvd[chon_cvd]
                        col_s1, col_s2 = st.columns(2)
                        with col_s1:
                            tieu_de_sua_cvd = st.text_input("Tiêu đề:", value=bg_sua.get('tieu_de') or '', key=f"sua_td_cvd_{bg_sua['id']}")
                            ma_vach_sua_cvd = st.text_input("Mã vạch Bưu điện:", value=bg_sua.get('ma_vach_buu_dien') or '', key=f"sua_mv_cvd_{bg_sua['id']}")
                        with col_s2:
                            trich_yeu_sua_cvd = st.text_area("Trích yếu:", value=bg_sua.get('trich_yeu') or '', key=f"sua_ty_cvd_{bg_sua['id']}", height=80)
                        col_luu_cvd, col_xoa_cvd = st.columns(2)
                        with col_luu_cvd:
                            if st.button("💾 Lưu thay đổi", key=f"btn_luu_cvd_{bg_sua['id']}", type="primary", width='stretch', disabled=not can_edit()):
                                try:
                                    db_s = st.session_state.db_engine.get_connection()
                                    c_s = db_s.cursor()
                                    c_s.execute("""
                                        UPDATE cong_van_di SET tieu_de=%s, trich_yeu=%s, ma_vach_buu_dien=%s
                                        WHERE id=%s
                                    """, (tieu_de_sua_cvd, trich_yeu_sua_cvd, ma_vach_sua_cvd, bg_sua['id']))
                                    db_s.commit(); db_s.close()
                                    st.success("✅ Đã cập nhật công văn đi")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Lỗi: {e}")
                        with col_xoa_cvd:
                            if st.button("🗑️ Xóa công văn này", key=f"btn_xoa_cvd_{bg_sua['id']}", width='stretch', disabled=not can_edit()):
                                try:
                                    db_x = st.session_state.db_engine.get_connection()
                                    c_x = db_x.cursor()
                                    c_x.execute("DELETE FROM cong_van_di WHERE id=%s", (bg_sua['id'],))
                                    db_x.commit(); db_x.close()
                                    st.success("✅ Đã xóa công văn đi")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Lỗi: {e}")

                col_export1, col_export2, col_export3 = st.columns([1, 2, 1])
                with col_export2:
                    if st.button("📥 Xuất Excel công văn đi", width='stretch', type="primary"):
                        headers = ['so_cong_van', 'loai_cong_van', 'phong_phat_hanh', 'ngay_phat_hanh', 'tieu_de', 'trich_yeu', 'ghi_chu']
                        col_widths = [20, 12, 20, 12, 35, 30, 25]
                        title = f"BÁO CÁO CÔNG VĂN ĐI (Từ {tu_ngay_cv_di or '...'} đến {den_ngay_cv_di or '...'})"
                        filename = f"Cong_van_di_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                        
                        excel_data = []
                        for row in data_cv_di:
                            excel_row = {}
                            for key in headers:
                                val = row.get(key)
                                if key == 'ngay_phat_hanh' and val:
                                    val = format_date(val)
                                if key == 'loai_cong_van':
                                    val = loai_name_map.get(val, val)
                                excel_row[key] = val
                            excel_data.append(excel_row)
                        
                        export_cong_van_excel(excel_data, filename, headers, col_widths, title)
                        
                        with open(filename, "rb") as f:
                            st.download_button(
                                label="📥 TẢI FILE EXCEL",
                                data=f,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                width='stretch'
                            )
                        st.success(f"✅ Đã xuất {len(data_cv_di)} công văn đi")
                        st.cache_data.clear()
            else:
                st.info("📭 Không có công văn đi nào")
    
    # === TAB 3: HỢP ĐỒNG KINH TẾ ===
    with tab3:
        st.subheader("📑 Quản lý Hợp đồng kinh tế")

        thao_tac_hdkt = st.radio(
            "Chọn thao tác:",
            ["➕ Thêm HĐKT mới", "📂 Nhập lại HĐKT cũ", "🔍 Tra cứu"],
            horizontal=True, key="hdkt_thao_tac"
        )

        # ── THÊM HĐKT MỚI (số tự sinh) ──
        if thao_tac_hdkt == "➕ Thêm HĐKT mới":
            so_hd_xem_truoc, prefix_hdkt_dang_dung = preview_so_hdkt()
            st.info(
                f"📄 **Số HĐKT dự kiến:** `{so_hd_xem_truoc}` (Prefix: **{prefix_hdkt_dang_dung}**) "
                f"— số chính thức sẽ được cấp khi bấm **Lưu Hợp đồng**"
            )
            with st.form("add_hop_dong_kt"):
                col1, col2 = st.columns(2)
                with col1:
                    so_hd_tuy_chinh = st.text_input(
                        "Số hợp đồng (để trống sẽ tự cấp số theo cấu hình trên)",
                        placeholder=so_hd_xem_truoc
                    )
                    ten_doi_tac = st.text_input("Tên đối tác *", placeholder="VD: Công ty TNHH ABC")
                    ngay_ky = st.date_input("Ngày ký *", value=date.today())
                with col2:
                    trich_yeu = st.text_area("Trích yếu", placeholder="Tóm tắt nội dung hợp đồng...", height=80)
                    ghi_chu = st.text_area("Ghi chú", height=60)
                
                uploaded_file = st.file_uploader("📎 Upload file", type=['pdf', 'doc', 'docx', 'jpg', 'png', 'jpeg'], key="hd_kt_upload")
                
                col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
                with col_btn2:
                    if st.form_submit_button("💾 Lưu hợp đồng", width='stretch', type="primary", disabled=not can_edit()):
                        if not ten_doi_tac:
                            st.error("⚠️ Vui lòng nhập đầy đủ các trường bắt buộc (*)")
                        else:
                            try:
                                so_hd = so_hd_tuy_chinh.strip() if so_hd_tuy_chinh.strip() else generate_so_hdkt()
                                file_url = None
                                if uploaded_file:
                                    file_url = upload_cong_van_file(uploaded_file, "hop_dong_kt")
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("""
                                    INSERT INTO hop_dong_kinh_te (so_hop_dong, ten_doi_tac, ngay_ky, 
                                    trich_yeu, file_url, ghi_chu, nguoi_tao)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                                """, (so_hd, ten_doi_tac, ngay_ky, trich_yeu, file_url, ghi_chu, 
                                      st.session_state.username))
                                db.commit()
                                db.close()
                                st.success(f"✅ Đã thêm hợp đồng: {so_hd}")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi: {e}")

        # ── NHẬP LẠI HĐKT CŨ (số nhập tay, không tự sinh) ──
        elif thao_tac_hdkt == "📂 Nhập lại HĐKT cũ":
            st.caption("Nhập lại các hợp đồng kinh tế đã ký trước đây vào hệ thống. "
                       "Số hợp đồng do bạn tự nhập.")
            with st.form("add_hop_dong_kt_cu"):
                col1, col2 = st.columns(2)
                with col1:
                    so_hd_cu = st.text_input("Số hợp đồng *", placeholder="VD: 02/2024/HĐKT-CHL", key="hdkt_cu_sohd")
                    ten_doi_tac_cu = st.text_input("Tên đối tác *", placeholder="VD: Công ty TNHH ABC", key="hdkt_cu_doitac")
                    ngay_ky_cu = st.date_input("Ngày ký *", value=date.today(), key="hdkt_cu_ngayky")
                with col2:
                    trich_yeu_cu = st.text_area("Trích yếu", placeholder="Tóm tắt nội dung hợp đồng...", height=80, key="hdkt_cu_trichyeu")
                    ghi_chu_cu = st.text_area("Ghi chú", height=60, key="hdkt_cu_ghichu")
                
                uploaded_file_cu = st.file_uploader("📎 Upload file", type=['pdf', 'doc', 'docx', 'jpg', 'png', 'jpeg'], key="hd_kt_cu_upload")
                
                col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
                with col_btn2:
                    if st.form_submit_button("💾 Lưu HĐKT cũ", width='stretch', type="primary", disabled=not can_edit()):
                        if not so_hd_cu.strip() or not ten_doi_tac_cu:
                            st.error("⚠️ Vui lòng nhập đầy đủ: Số hợp đồng, Tên đối tác")
                        else:
                            try:
                                file_url_cu = None
                                if uploaded_file_cu:
                                    file_url_cu = upload_cong_van_file(uploaded_file_cu, "hop_dong_kt")
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("""
                                    INSERT INTO hop_dong_kinh_te (so_hop_dong, ten_doi_tac, ngay_ky, 
                                    trich_yeu, file_url, ghi_chu, nguoi_tao)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                                """, (so_hd_cu.strip(), ten_doi_tac_cu, ngay_ky_cu, trich_yeu_cu, file_url_cu, ghi_chu_cu, 
                                      st.session_state.username))
                                db.commit()
                                db.close()
                                st.success(f"✅ Đã nhập lại HĐKT cũ: {so_hd_cu.strip()}")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi: {e}")

        # ── TRA CỨU ──
        else:
            col_search1, col_search2, col_search3 = st.columns([2, 1, 1])
            with col_search1:
                search_text_hd = st.text_input("🔍 Tìm kiếm", placeholder="Theo số HĐ, đối tác...", key="search_hd_kt")
            with col_search2:
                tu_ngay_hd = st.date_input("Từ ngày", value=None, key="tu_ngay_hd")
            with col_search3:
                den_ngay_hd = st.date_input("Đến ngày", value=None, key="den_ngay_hd")
            
            data_hd = get_hop_dong_kinh_te(tu_ngay_hd, den_ngay_hd, search_text_hd)
            
            if data_hd:
                df_hd = pd.DataFrame(data_hd)
                for col in ['ngay_ky', 'created_at', 'updated_at']:
                    if col in df_hd.columns:
                        df_hd[col] = df_hd[col].apply(format_date)
                
                display_cols = ['so_hop_dong', 'ten_doi_tac', 'ngay_ky', 'trich_yeu', 'file_url', 'ghi_chu']
                available_cols = [c for c in display_cols if c in df_hd.columns]
                df_display = df_hd[available_cols]
                
                col_map = {
                    'so_hop_dong': 'Số hợp đồng',
                    'ten_doi_tac': 'Đối tác',
                    'ngay_ky': 'Ngày ký',
                    'trich_yeu': 'Trích yếu',
                    'file_url': 'File',
                    'ghi_chu': 'Ghi chú'
                }
                df_display.rename(columns=col_map, inplace=True)
                
                st.caption(f"📌 Tổng số: {len(data_hd)} hợp đồng kinh tế")
                st.dataframe(df_display, width='stretch', hide_index=True, height=400)

                with st.expander("✏️ Sửa / 🗑️ Xóa hợp đồng kinh tế", expanded=False):
                    tuy_chon_hd = {f"{r['so_hop_dong']} - {r.get('ten_doi_tac') or ''}": r for r in data_hd}
                    chon_hd = st.selectbox("Chọn hợp đồng:", ["-- Chọn --"] + list(tuy_chon_hd.keys()), key="chon_sua_hd")
                    if chon_hd != "-- Chọn --":
                        hd_sua = tuy_chon_hd[chon_hd]
                        col_s1, col_s2 = st.columns(2)
                        with col_s1:
                            ten_doi_tac_sua = st.text_input("Tên đối tác:", value=hd_sua.get('ten_doi_tac') or '', key=f"sua_dt_hd_{hd_sua['id']}")
                        with col_s2:
                            trich_yeu_sua_hd = st.text_area("Trích yếu:", value=hd_sua.get('trich_yeu') or '', key=f"sua_ty_hd_{hd_sua['id']}", height=80)
                        col_luu_hd, col_xoa_hd = st.columns(2)
                        with col_luu_hd:
                            if st.button("💾 Lưu thay đổi", key=f"btn_luu_hd_{hd_sua['id']}", type="primary", width='stretch', disabled=not can_edit()):
                                try:
                                    db_s = st.session_state.db_engine.get_connection()
                                    c_s = db_s.cursor()
                                    c_s.execute("""
                                        UPDATE hop_dong_kinh_te SET ten_doi_tac=%s, trich_yeu=%s
                                        WHERE id=%s
                                    """, (ten_doi_tac_sua, trich_yeu_sua_hd, hd_sua['id']))
                                    db_s.commit(); db_s.close()
                                    st.success("✅ Đã cập nhật hợp đồng")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Lỗi: {e}")
                        with col_xoa_hd:
                            if st.button("🗑️ Xóa hợp đồng này", key=f"btn_xoa_hd_{hd_sua['id']}", width='stretch', disabled=not can_edit()):
                                try:
                                    db_x = st.session_state.db_engine.get_connection()
                                    c_x = db_x.cursor()
                                    c_x.execute("DELETE FROM hop_dong_kinh_te WHERE id=%s", (hd_sua['id'],))
                                    db_x.commit(); db_x.close()
                                    st.success("✅ Đã xóa hợp đồng")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Lỗi: {e}")

                col_export1, col_export2, col_export3 = st.columns([1, 2, 1])
                with col_export2:
                    if st.button("📥 Xuất Excel hợp đồng", width='stretch', type="primary"):
                        headers = ['so_hop_dong', 'ten_doi_tac', 'ngay_ky', 'trich_yeu', 'ghi_chu']
                        col_widths = [20, 30, 12, 35, 25]
                        title = f"BÁO CÁO HỢP ĐỒNG KINH TẾ (Từ {tu_ngay_hd or '...'} đến {den_ngay_hd or '...'})"
                        filename = f"Hop_dong_kinh_te_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                        
                        excel_data = []
                        for row in data_hd:
                            excel_row = {}
                            for key in headers:
                                val = row.get(key)
                                if key == 'ngay_ky' and val:
                                    val = format_date(val)
                                excel_row[key] = val
                            excel_data.append(excel_row)
                        
                        export_cong_van_excel(excel_data, filename, headers, col_widths, title)
                        
                        with open(filename, "rb") as f:
                            st.download_button(
                                label="📥 TẢI FILE EXCEL",
                                data=f,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                width='stretch'
                            )
                        st.success(f"✅ Đã xuất {len(data_hd)} hợp đồng kinh tế")
                        st.cache_data.clear()
            else:
                st.info("📭 Không có hợp đồng kinh tế nào")

# ========== KHỞI TẠO SESSION STATE ==========
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.role = None
    st.session_state.username = None

# ========== HIỂN THỊ LANDING PAGE NẾU CHƯA VÀO HRM ==========
# Chỉ hiện logo mặc định (COMPANY_CONFIG/logo_cty.png) khi KHÔNG có tenant (chế độ standalone
# cũ). Nếu đã xác định được tenant (đa khách hàng), logo của tenant sẽ tự hiện ở bước đăng nhập
# phía dưới (st.sidebar.image(tenant['logo_url'])) — tránh hiện 2 logo chồng lên nhau.
if not st.session_state.get('tenant'):
    logo_url = COMPANY_CONFIG.get("logo_url")
    if logo_url:
        with st.sidebar:
            st.markdown(f"""
            <div style="display: flex; justify-content: center; align-items: center; 
                        width: 100%; min-height: 200px; padding: 10px 0;">
                <img src="{logo_url}" 
                     style="width: 150px; height: 150px; border-radius: 50%; 
                            object-fit: cover; box-shadow: 0 4px 20px rgba(0,0,0,0.25);
                            display: block;">
            </div>
            """, unsafe_allow_html=True)
            st.divider()
    elif os.path.exists("logo_cty.png"):
        with st.sidebar:
            st.image("logo_cty.png", width='stretch')
            st.divider()

# ========== PHẦN CODE HRM BẮT ĐẦU TỪ ĐÂY ==========

st.markdown("""
    <style>
        /* ===== ẨN MANAGE APP - dùng mọi selector có thể ===== */
        [data-testid="stToolbar"],
        [data-testid="manage-app-button"],
        [data-testid="stAppDeployButton"],
        .stDeployButton,
        #MainMenu,
        div[class*="toolbar"],
        div[class*="StatusWidget"],
        div[class*="viewerBadge"],
        div[class*="manage-app"],
        button[kind="managedApp"],
        [data-testid="stBottom"] > div:last-child { 
            display: none !important; 
            visibility: hidden !important;
            opacity: 0 !important;
            pointer-events: none !important;
        }
        footer[data-testid] { display: none !important; }

        /* ===== PADDING TOP / BOTTOM = 5px ===== */
        .stApp > div[data-testid="stAppViewContainer"] > section[data-testid="stMain"] > div {
            padding-top: 5px !important;
            padding-bottom: 5px !important;
        }
        .block-container {
            padding-top: 5px !important;
            padding-bottom: 5px !important;
        }

        /* ===== LOGO SIDEBAR: hình tròn đổ bóng, căn giữa HOÀN TOÀN ===== */
        [data-testid="stSidebar"] > div:first-child {
            padding-top: 0 !important;
        }
        
        /* TẤT CẢ các container chứa ảnh trong sidebar đều phải có width:100% và flex center */
        [data-testid="stSidebar"] [data-testid="element-container"] {
            width: 100% !important;
            display: flex !important;
            justify-content: center !important;
        }
        
        [data-testid="stSidebar"] [data-testid="element-container"] [data-testid="stImage"] {
            width: 100% !important;
            display: flex !important;
            justify-content: center !important;
            align-items: center !important;
        }
        
        /* Bản thân thẻ <img> nằm bên trong div con của stImage */
        [data-testid="stSidebar"] [data-testid="stImage"] > div {
            display: flex !important;
            justify-content: center !important;
            align-items: center !important;
            width: 100% !important;
        }
        
        /* QUAN TRỌNG: style cho chính thẻ img */
        [data-testid="stSidebar"] [data-testid="stImage"] img {
            width: 150px !important;
            height: 150px !important;
            border-radius: 50% !important;
            object-fit: cover !important;
            box-shadow: 0 4px 20px rgba(0,0,0,0.25) !important;
            display: block !important;
            margin: 4px auto 0 auto !important;
        }
    </style>
    <script>
        // MutationObserver: ẩn Manage App ngay khi DOM thay đổi
        function hideManageApp() {
            const selectors = [
                '[data-testid="manage-app-button"]',
                '[data-testid="stToolbar"]',
                '[data-testid="stAppDeployButton"]',
                '.stDeployButton',
                'button[kind="managedApp"]'
            ];
            selectors.forEach(sel => {
                document.querySelectorAll(sel).forEach(el => {
                    el.style.cssText = 'display:none!important;visibility:hidden!important';
                });
            });
        }
        hideManageApp();
        // Quan sát DOM liên tục — bắt được dù Streamlit render trễ
        const observer = new MutationObserver(hideManageApp);
        observer.observe(document.body, { childList: true, subtree: true });
    </script>
""", unsafe_allow_html=True)

def to_float_or_none(val):
    """Chuyển đổi giá trị sang float hoặc None, tránh lỗi numeric"""
    if val is None or str(val).strip() == '':
        return None
    try:
        return float(val)
    except:
        return None

def format_date(d):
    if d is None or pd.isna(d): return ''
    try: return d.strftime('%d/%m/%Y') if hasattr(d,'strftime') else str(d)[:10]
    except: return str(d)

def parse_date(s):
    """Chuyển đổi nhiều định dạng ngày tháng về date object"""
    if not s or str(s).strip() == '':
        return None
    
    # Nếu đã là date object
    if hasattr(s, 'strftime'):
        return s
    
    s = str(s).strip()
    
    # Các định dạng cần thử
    formats = [
        '%d/%m/%Y',      # 18/04/2026
        '%d-%m-%Y',      # 18-04-2026
        '%Y-%m-%d',      # 2026-04-18
        '%Y/%m/%d',      # 2026/04/18
        '%d.%m.%Y',      # 18.04.2026
    ]
    
    for fmt in formats:
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    
    # Nếu không có định dạng nào phù hợp
    print(f"⚠️ Không thể parse ngày: {s}")
    return None

def get_xung_ho_trang_trong(gioi_tinh):
    """Xưng hô trang trọng theo giới tính, dùng cho Hợp đồng và các card hiển thị lãnh đạo
    cấp cao (Hội đồng Quản trị, Ban Tổng Giám đốc...).
    - Nam -> "Ông"; Nữ -> "Bà"; None/rỗng -> "Ông/Bà"
    """
    if gioi_tinh == "Nam":
        return "Ông"
    elif gioi_tinh == "Nữ":
        return "Bà"
    else:
        return "Ông/Bà"

def get_xung_ho(gioi_tinh, ho_ten=""):
    """
    Lấy cách xưng hô phù hợp dựa trên giới tính
    - Nếu giới tính là Nam -> trả về "Anh"
    - Nếu giới tính là Nữ -> trả về "Chị"
    - Nếu giới tính là None hoặc rỗng -> trả về "Anh/Chị"
    """
    if gioi_tinh == "Nam":
        return "Anh"
    elif gioi_tinh == "Nữ":
        return "Chị"
    else:
        return "Anh/Chị"

def get_loi_chuc_sinh_nhat(ho_ten, gioi_tinh, tuoi=None):
    """
    Tạo lời chúc sinh nhật có xưng hô phù hợp
    """
    xung_ho = get_xung_ho(gioi_tinh, ho_ten)
    ten_cty_sn = COMPANY_CONFIG.get("ten_cong_ty", "Công ty")

    loi_chuc = f"""
🎉🎂 CHÚC MỪNG SINH NHẬT {xung_ho.upper()} {ho_ten.upper()} 🎂🎉

Thân gửi {xung_ho}: {ho_ten},

Nhân dịp sinh nhật của {xung_ho}, thay mặt Ban Lãnh đạo {ten_cty_sn}, 
xin gửi đến {xung_ho} những lời chúc tốt đẹp nhất.

Chúc {xung_ho} luôn mạnh khỏe, hạnh phúc và thành công trong công việc 
cũng như trong cuộc sống.

"""
    
    if tuoi:
        loi_chuc += f"Chúc mừng {xung_ho} tròn {tuoi} tuổi! 🎂\n\n"
    
    loi_chuc += f"""
Cảm ơn {xung_ho} đã luôn đồng hành và đóng góp cho sự phát triển của Công ty.

Trân trọng!

🏗️ {ten_cty_sn.upper()}
    """
    
    return loi_chuc

def auto_check_birthday():
    if 'sinh_nhat_hom_nay_list' not in st.session_state:
        st.session_state.sinh_nhat_hom_nay_list = []

    today_str = date.today().strftime('%Y-%m-%d')
    
    # Key mới: gắn với cả ngày lẫn trạng thái login
    # → mỗi lần login lại đều query DB, không bị cache nhầm
    check_key = f"{today_str}_{st.session_state.get('username', 'guest')}"
    
    if st.session_state.get('last_birthday_check') == check_key:
        return  # Đã check cho user này hôm nay rồi
    
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("""
            SELECT id, ma_nv, ho_ten, ngay_sinh, gioi_tinh, dien_thoai, email_lien_he
            FROM nhan_vien 
            WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC')
            AND ngay_sinh IS NOT NULL
            AND EXTRACT(MONTH FROM ngay_sinh) = EXTRACT(MONTH FROM CURRENT_DATE)
            AND EXTRACT(DAY FROM ngay_sinh) = EXTRACT(DAY FROM CURRENT_DATE)
        """)
        birthday_today = c.fetchall()
        db.close()
        
        st.session_state.sinh_nhat_hom_nay_list = [
            {
                'ho_ten': nv['ho_ten'],
                'ma_nv': nv['ma_nv'],
                'xung_ho': get_xung_ho(nv.get('gioi_tinh'), nv['ho_ten'])
            }
            for nv in birthday_today
        ]
        
        for nv in birthday_today:
            xung_ho = get_xung_ho(nv.get('gioi_tinh'), nv['ho_ten'])
            st.toast(f"🎂 Sinh nhật {xung_ho} {nv['ho_ten']} hôm nay!", icon="🎂")
        
        # Đánh dấu đã check — dùng check_key gắn với username
        st.session_state.last_birthday_check = check_key
        
    except Exception as e:
        st.warning(f"⚠️ Không thể kiểm tra sinh nhật: {e}")

def da_chuyen_doi_chinh_thuc(nv_id):
    """Kiểm tra xem nhân viên đã có quyết định chuyển từ thử việc sang chính thức chưa"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("""
            SELECT * FROM quyet_dinh_nhan_su 
            WHERE nhan_vien_id = %s AND loai_quyet_dinh = 'CHINH_THUC'
            ORDER BY ngay_quyet_dinh DESC LIMIT 1
        """, (nv_id,))
        result = c.fetchone()
        db.close()
        
        # Debug: in ra log để kiểm tra
        print(f"Checking nv_id={nv_id}, found={result is not None}")
        if result:
            print(f"Quyet dinh: {result}")
        
        return result is not None, result
    except Exception as e:
        print(f"Error in da_chuyen_doi_chinh_thuc: {e}")
        return False, None

def lay_thong_tin_truoc_chuyen_doi(nv_id):
    """Lấy thông tin nhân viên trước khi chuyển đổi (từ lich_su_cong_tac)"""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        # Lấy lịch sử công tác cũ nhất (thời gian thử việc)
        c.execute("""
            SELECT * FROM lich_su_cong_tac 
            WHERE nhan_vien_id = %s 
            ORDER BY tu_ngay ASC LIMIT 1
        """, (nv_id,))
        result = c.fetchone()
        db.close()
        return result
    except:
        return None
# ========== DATABASE CONNECTION (SUPABASE) — ĐA KHÁCH HÀNG (MULTI-TENANT) ==========
def get_connection():
    """Wrapper tương thích ngược, tự động gọi db_engine từ st.session_state."""
    if 'db_engine' not in st.session_state:
        st.session_state.db_engine = DatabaseEngine(st.session_state.get('tenant'))
    return st.session_state.db_engine.get_connection()

# ========== SUPABASE STORAGE (lưu trữ file hồ sơ) ==========
# Tên bucket Storage trên Supabase dùng để lưu hồ sơ nhân viên.
# Cần tạo trước trên Supabase Dashboard > Storage (khuyến nghị để Private).
SUPABASE_BUCKET = "ho-so-nhan-vien"

# Bucket RIÊNG dùng cho logo công ty — PHẢI để Public vì logo hiển thị ngay ở màn
# hình đăng nhập, TRƯỚC KHI xác thực người dùng (không thể dùng link ký riêng tư
# như ảnh hồ sơ nhân viên). Tách bucket riêng thay vì dùng chung SUPABASE_BUCKET để
# không phải hạ mức bảo mật của bucket hồ sơ nhân viên (ảnh CCCD, hồ sơ cá nhân...).
SUPABASE_BUCKET_LOGO = "logo-cong-ty"

def sanitize_storage_filename(filename):
    """Chuẩn hóa tên file để làm 'key' hợp lệ trên Supabase Storage:
    - Bỏ dấu tiếng Việt (Lộ_trình_học -> Lo_trinh_hoc)
    - Thay khoảng trắng bằng '_'
    - Chỉ giữ lại chữ cái không dấu, số, '_', '-', '.'
    Supabase Storage sẽ báo lỗi 'InvalidKey' nếu key chứa ký tự có dấu/unicode."""
    # Bỏ dấu: chuẩn hóa Unicode rồi loại bỏ các ký tự dấu (combining marks)
    normalized = unicodedata.normalize('NFD', filename)
    no_accent = ''.join(ch for ch in normalized if unicodedata.category(ch) != 'Mn')
    # Xử lý riêng chữ đ/Đ vì NFD không tách được
    no_accent = no_accent.replace('đ', 'd').replace('Đ', 'D')
    # Thay khoảng trắng bằng '_'
    no_accent = re.sub(r'\s+', '_', no_accent)
    # Chỉ giữ ký tự an toàn
    safe = re.sub(r'[^A-Za-z0-9_.\-]', '', no_accent)
    return safe or "file"

def upload_to_storage_unique(sb, bucket, base_path, file_bytes, content_type, max_tries=50):
    """Upload file lên Supabase Storage tại base_path. Nếu path đã tồn tại (trùng
    Loại hồ sơ + ngày upload + tên file trong cùng ngày), tự thêm hậu tố _2, _3...
    trước phần mở rộng để không bị lỗi/ghi đè. Trả về path thực tế đã dùng để upload."""
    path = base_path
    root, ext = os.path.splitext(base_path)
    tries = 0
    while True:
        try:
            sb.storage.from_(bucket).upload(
                path=path,
                file=file_bytes,
                file_options={"content-type": content_type or "application/octet-stream"}
            )
            return path
        except Exception as e:
            msg = str(e).lower()
            is_duplicate = ('duplicate' in msg or 'exists' in msg or 'already' in msg or '409' in msg)
            if is_duplicate and tries < max_tries:
                tries += 1
                path = f"{root}_{tries + 1}{ext}"
                continue
            raise

def _ensure_bucket_exists(sb, bucket, public=False):
    """Tự động tạo bucket Storage nếu chưa tồn tại trên project Supabase của TENANT
    đang đăng nhập. Mỗi tenant có 1 project Supabase riêng (mô hình SaaS đa khách
    hàng) nên bucket phải được tạo riêng cho từng project — nếu tenant mới chưa
    từng tạo bucket này thủ công trên Dashboard, MỌI upload ảnh hồ sơ/file công văn
    sẽ lỗi "Bucket not found". Hàm này tự tạo bucket nếu chưa có, giống cách các
    bảng DB trong app được tự tạo bằng CREATE TABLE IF NOT EXISTS.

    public=True: dùng cho các bucket cần truy cập công khai KHÔNG cần đăng nhập
    (VD logo công ty — hiển thị ngay ở màn hình đăng nhập, trước khi xác thực).
    public=False (mặc định): dùng cho dữ liệu riêng tư (ảnh hồ sơ nhân viên...).
    Nếu bucket ĐÃ tồn tại nhưng đang ở chế độ khác với `public` yêu cầu, tự động
    cập nhật lại (update_bucket) — xử lý trường hợp bucket logo từng bị tạo nhầm
    ở chế độ Private trước khi có tham số này.

    Im lặng bỏ qua nếu tạo/cập nhật thất bại (VD key không đủ quyền) — lỗi gốc sẽ
    hiện ra rõ ràng khi thực sự upload, không che giấu vấn đề.
    Cache trong session để chỉ kiểm tra 1 lần/phiên/bucket, tránh gọi API thừa."""
    cache_key = f"_sb_bucket_checked_{bucket}_{public}"
    if st.session_state.get(cache_key):
        return
    try:
        info = sb.storage.get_bucket(bucket)
        is_public_now = bool(getattr(info, "public", None) if not isinstance(info, dict) else info.get("public"))
        if public and not is_public_now:
            try:
                sb.storage.update_bucket(bucket, options={"public": True})
            except Exception as e:
                st.session_state[f"_sb_bucket_error_{bucket}"] = str(e)
                print(f"Không thể chuyển bucket '{bucket}' sang Public: {e}")
                return
        st.session_state[cache_key] = True
        st.session_state.pop(f"_sb_bucket_error_{bucket}", None)
        return
    except Exception:
        pass
    try:
        sb.storage.create_bucket(bucket, options={"public": public})
        st.session_state[cache_key] = True
        st.session_state.pop(f"_sb_bucket_error_{bucket}", None)
    except Exception as e:
        # Có thể bucket đã được tạo bởi 1 phiên khác cùng lúc (race), hoặc API key
        # hiện tại không đủ quyền tạo bucket (thường gặp nhất: key lưu cho tenant là
        # "anon/public key", loại key này KHÔNG có quyền tạo bucket trên Supabase —
        # chỉ "service_role key" mới tạo được). KHÔNG cache lại là "đã kiểm tra xong"
        # trong trường hợp lỗi, để lần upload kế tiếp còn tự thử lại (phòng khi bucket
        # được tạo thủ công/đã hết lỗi tạm thời). Lưu lại lý do lỗi để nơi gọi upload
        # (upload_anh_ho_so...) có thể hiển thị hướng dẫn rõ ràng cho người dùng thay vì
        # chỉ có dòng lỗi kỹ thuật "Bucket not found" khó hiểu.
        st.session_state[f"_sb_bucket_error_{bucket}"] = str(e)
        print(f"Không thể tự tạo bucket Storage '{bucket}': {e}")


def get_supabase_storage():
    """Khởi tạo Supabase Client dùng cho Storage (ảnh NV, hồ sơ, file chat...).
    Ưu tiên url/key của TENANT đang đăng nhập (mô hình SaaS đa khách hàng).
    Fallback sang st.secrets['supabase'] / .env khi chạy chế độ đơn khách hàng.
    Không dùng @st.cache_resource nữa vì client giờ có thể khác nhau theo từng tenant
    trong cùng 1 tiến trình app (nhiều khách hàng dùng chung 1 deployment)."""
    try:
        from supabase import create_client
    except ImportError:
        print("Chưa cài thư viện supabase. Chạy: pip install supabase")
        return None

    tenant = st.session_state.get('tenant')
    if tenant:
        url, key = tenant['supabase_url'], tenant['supabase_key']
    else:
        url, key = None, None
        try:
            if 'supabase' in st.secrets:
                url = st.secrets.supabase.get('url')
                key = st.secrets.supabase.get('key')
        except Exception:
            pass
        if not url or not key:
            from dotenv import load_dotenv
            load_dotenv()
            url = url or os.getenv('SUPABASE_URL')
            key = key or os.getenv('SUPABASE_KEY')

    if not url or not key:
        return None

    cache_key = f"_sb_client_{url}"
    if cache_key in st.session_state:
        return st.session_state[cache_key]

    try:
        client = create_client(url, key)
        _ensure_bucket_exists(client, SUPABASE_BUCKET, public=False)
        _ensure_bucket_exists(client, SUPABASE_BUCKET_LOGO, public=True)
        st.session_state[cache_key] = client
        return client
    except Exception as e:
        print(f"Lỗi khởi tạo Supabase Storage: {e}")
        return None

# ========== CHẤM CÔNG THỦ CÔNG - HẰNG SỐ & HÀM DÙNG CHUNG ==========
# Mã công theo đúng bảng chấm công mẫu (sheet T4-2026 / T5-2026 file HLP)
# Đây cũng chính là danh sách ký hiệu hợp lệ được liệt kê trong "Chú giải" và
# dùng để giới hạn dữ liệu nhập vào các ô ly (row Ca ngày / Ca đêm).

def cc_pin_col(col_type, **kwargs):
    """Tạo column_config, cố gắng ghim (pin) cột vào bên trái khi cuộn ngang.
    Một số phiên bản Streamlit cũ chưa hỗ trợ tham số `pinned` -> fallback bỏ qua."""
    try:
        return col_type(pinned=True, **kwargs)
    except TypeError:
        return col_type(**kwargs)

# Những bộ phận (theo mã phong_ban_lam_viec trong bảng nhan_vien) có phát sinh tăng ca
# theo đúng bản mẫu (nhóm LX-M/LDPT thực tế đang lưu mã "SX" và "LDPT")
CHAM_CONG_DEPT_TANG_CA = ["SX", "LDPT"]

CHAM_CONG_DEPT_LABEL = {
    "QL": "QL - Quản lý",
    "VP": "VP - Văn phòng",
    "SX": "SX - Sản xuất/Vận hành",
    "LDPT": "LDPT - Lao động phổ thông",
}

# Bộ phận chỉ chấm công 1 dòng/nhân viên (giờ hành chính, không tách ca ngày/đêm/tăng ca)
CHAM_CONG_DEPT_MOT_DONG = ["VP"]

CC_ROW_HEIGHT = 24  # giảm size dòng (px) để bảng chấm công hiển thị gọn, nhiều dữ liệu hơn

def cc_render_grid(data, edit=False, **kwargs):
    """Wrapper cho st.dataframe/st.data_editor, cố gắng thu nhỏ chiều cao dòng (row_height)
    nếu phiên bản Streamlit đang chạy hỗ trợ; nếu không thì bỏ qua tham số đó."""
    fn = st.data_editor if edit else st.dataframe
    try:
        return fn(data, row_height=CC_ROW_HEIGHT, **kwargs)
    except TypeError:
        return fn(data, **kwargs)

def cc_normalize_marker(v):
    """Chuẩn hoá ký hiệu chấm công theo bảng KY_HIEU_CHAM_CONG (22 ký hiệu).
    So khớp không phân biệt hoa/thường. Trả về dạng chuẩn hoặc None nếu trống."""
    v = (v or "").strip()
    if not v:
        return None
    # Map backward-compat: ký hiệu cũ → mới
    MAP_CU_MOI = {"X": "x", "V": "KL", "N": "x", "D": "x", "L": "NL", "0.5": "x/2"}
    vu = v.upper()
    # Khớp chính xác trong KY_HIEU_CHAM_CONG (ưu tiên)
    for code in KY_HIEU_CHAM_CONG:
        if code.upper() == vu:
            return code
    # Fallback backward-compat
    if vu in MAP_CU_MOI:
        return MAP_CU_MOI[vu]
    return v  # giữ nguyên ký hiệu lạ

def cc_is_cong(v):
    """Kiểm tra ký hiệu có tính công không (dựa trên KY_HIEU_CHAM_CONG)."""
    if not isinstance(v, str):
        return False
    ma = cc_normalize_marker(v)
    if ma and ma in KY_HIEU_CHAM_CONG:
        return KY_HIEU_CHAM_CONG[ma].get('cong', 0) > 0
    return False

def cc_marker_is(v, target):
    return isinstance(v, str) and v.strip().upper() == target.upper()


def ensure_cham_cong_table():
    """Tạo bảng cham_cong trên Supabase nếu chưa có, và tự nâng cấp thêm cột mới (idempotent).
    BỎ QUA cho tenant DEMO: DB demo bị khóa ghi kể cả lệnh ALTER/CREATE, nên schema DEMO
    phải được đồng bộ thủ công 1 lần (xem SQL bên dưới), không tự chạy lại mỗi lần tải trang."""
    tenant = st.session_state.get('tenant') or {}
    if str(tenant.get('ma_cty', '')).upper() == 'DEMO':
        return
    db = st.session_state.db_engine.get_connection()
    c = db.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS cham_cong (
            id SERIAL PRIMARY KEY,
            nhan_vien_id INTEGER NOT NULL REFERENCES nhan_vien(id) ON DELETE CASCADE,
            ngay DATE NOT NULL,
            ma_cong VARCHAR(10),
            ca_ngay VARCHAR(10),
            ca_dem VARCHAR(10),
            gio_tang_ca NUMERIC(5,2) DEFAULT 0,
            gio_tang_ca_le NUMERIC(5,2) DEFAULT 0,
            ghi_chu TEXT,
            nguon VARCHAR(20) DEFAULT 'THU_CONG',
            created_by VARCHAR(100),
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(nhan_vien_id, ngay)
        )
    """)
    # Nâng cấp cho DB đã tạo bảng từ phiên bản trước (chưa có ca_ngay/ca_dem)
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS ca_ngay VARCHAR(10)")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS ca_dem VARCHAR(10)")
    # Nâng cấp cho module Chấm công Face ID (giờ vào/ra thực tế, phân loại tăng ca)
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS gio_vao TIME")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS gio_ra TIME")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS loai_tang_ca TEXT")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS gio_tang_ca_dem NUMERIC(5,2) DEFAULT 0")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS gio_tang_ca_cn NUMERIC(5,2) DEFAULT 0")
    # Cột mở rộng cho Face ID (Điểm 2, 3) + BCC tháng (Điểm 1)
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS trang_thai_cham_cong TEXT DEFAULT 'HOP_LE'")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS so_phut_di_tre INTEGER DEFAULT 0")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS so_phut_ve_som INTEGER DEFAULT 0")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS da_gui_canh_bao_thieu_gio_ra BOOLEAN DEFAULT FALSE")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS trang_thai_duyet_tc TEXT")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS loai_ngay_tang_ca TEXT")
    # Cột mở rộng cho chấm công ngoài địa điểm (Điểm 5.5)
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS trong_dia_diem BOOLEAN")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS trang_thai_vi_tri TEXT")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS nguoi_duyet_vi_tri TEXT")
    c.execute("ALTER TABLE cham_cong ADD COLUMN IF NOT EXISTS ly_do_ngoai_dia_diem TEXT")
    # Bảng audit log điều chỉnh chấm công
    c.execute("""
        CREATE TABLE IF NOT EXISTS audit_cham_cong (
            id SERIAL PRIMARY KEY,
            cham_cong_id INTEGER,           -- id dòng trong bảng cham_cong được sửa
            nhan_vien_id INTEGER,
            ngay DATE,
            truong_sua TEXT,                -- 'gio_vao' | 'gio_ra' | 'ma_cong' | 'ghi_chu'
            gia_tri_cu TEXT,
            gia_tri_moi TEXT,
            ly_do TEXT NOT NULL,
            nguoi_sua TEXT,                 -- username người thực hiện
            thoi_diem_sua TIMESTAMP DEFAULT NOW(),
            trang_thai TEXT DEFAULT 'DA_DUYET',  -- 'CHO_DUYET' | 'DA_DUYET' | 'TU_CHOI'
            nguoi_duyet TEXT,
            thoi_diem_duyet TIMESTAMP,
            ghi_chu_duyet TEXT
        )
    """)
    # Bảng cấu hình tăng ca theo phòng ban
    c.execute("""
        CREATE TABLE IF NOT EXISTS cau_hinh_tang_ca_phong_ban (
            id SERIAL PRIMARY KEY,
            ten_phong_ban TEXT NOT NULL,
            cho_phep_tang_ca BOOLEAN NOT NULL DEFAULT TRUE,
            he_so_tc_thuong NUMERIC(4,2),
            he_so_tc_chu_nhat NUMERIC(4,2),
            he_so_tc_le NUMERIC(4,2),
            he_so_tc_dem NUMERIC(4,2),
            don_gia_tc_thuong NUMERIC(12,2),
            don_gia_tc_chu_nhat NUMERIC(12,2),
            don_gia_tc_le NUMERIC(12,2),
            don_gia_tc_dem NUMERIC(12,2),
            ghi_chu TEXT,
            updated_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(ten_phong_ban)
        )
    """)
    # Bảng khoá tháng BCC
    c.execute("""
        CREATE TABLE IF NOT EXISTS khoa_thang_bcc (
            id SERIAL PRIMARY KEY,
            thang INTEGER NOT NULL,
            nam INTEGER NOT NULL,
            trang_thai TEXT NOT NULL DEFAULT 'MO',   -- 'MO' | 'KHOA'
            nguoi_khoa TEXT,
            thoi_diem_khoa TIMESTAMP,
            nguoi_mo TEXT,
            thoi_diem_mo TIMESTAMP,
            ghi_chu TEXT,
            UNIQUE(thang, nam)
        )
    """)
    db.commit()
    c.close()
    db.close()


def is_thang_da_khoa(thang, nam):
    """Kiểm tra tháng đã bị khoá BCC chưa. Trả về True nếu đã khoá."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT trang_thai FROM khoa_thang_bcc WHERE thang=%s AND nam=%s", (thang, nam))
        row = c.fetchone()
        db.close()
        return row is not None and row[0] == 'KHOA'
    except Exception:
        return False


def khoa_mo_thang_bcc(thang, nam, hanh_dong, ghi_chu=''):
    """Khoá hoặc mở khoá BCC tháng. hanh_dong = 'KHOA' | 'MO'."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        if hanh_dong == 'KHOA':
            c.execute("""
                INSERT INTO khoa_thang_bcc (thang, nam, trang_thai, nguoi_khoa, thoi_diem_khoa, ghi_chu)
                VALUES (%s, %s, 'KHOA', %s, NOW(), %s)
                ON CONFLICT (thang, nam) DO UPDATE SET
                    trang_thai='KHOA', nguoi_khoa=EXCLUDED.nguoi_khoa,
                    thoi_diem_khoa=NOW(), ghi_chu=EXCLUDED.ghi_chu
            """, (thang, nam, st.session_state.username, ghi_chu))
        else:
            c.execute("""
                INSERT INTO khoa_thang_bcc (thang, nam, trang_thai, nguoi_mo, thoi_diem_mo, ghi_chu)
                VALUES (%s, %s, 'MO', %s, NOW(), %s)
                ON CONFLICT (thang, nam) DO UPDATE SET
                    trang_thai='MO', nguoi_mo=EXCLUDED.nguoi_mo,
                    thoi_diem_mo=NOW(), ghi_chu=EXCLUDED.ghi_chu
            """, (thang, nam, st.session_state.username, ghi_chu))
        db.commit()
        db.close()
        return True
    except Exception:
        return False


def ensure_face_id_table():
    """Tạo bảng nhan_vien_face_id trên Supabase nếu chưa có (idempotent)."""
    db = st.session_state.db_engine.get_connection()
    c = db.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS nhan_vien_face_id (
            id SERIAL PRIMARY KEY,
            nhan_vien_id INTEGER NOT NULL REFERENCES nhan_vien(id) ON DELETE CASCADE,
            face_encoding JSONB NOT NULL,
            model_name TEXT NOT NULL DEFAULT 'Facenet',
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(nhan_vien_id)
        )
    """)
    db.commit()
    c.close()
    db.close()

def ensure_qdns_columns():
    """Bổ sung cột 'chuc_vu' và 'ngay_qd_ns' vào bảng nhan_vien nếu chưa có (idempotent).
    - chuc_vu: mặc định 'Nhân viên'
    - ngay_qd_ns: mặc định lấy theo ngay_vao_lam; sau này được cập nhật theo ngày ban hành
      các Quyết định nhân sự (bổ nhiệm/miễn nhiệm/đổi chức danh/điều chuyển)."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("ALTER TABLE nhan_vien ADD COLUMN IF NOT EXISTS chuc_vu VARCHAR(100)")
        c.execute("ALTER TABLE nhan_vien ADD COLUMN IF NOT EXISTS ngay_qd_ns DATE")
        c.execute("ALTER TABLE nhan_vien ADD COLUMN IF NOT EXISTS so_luong_npt INTEGER DEFAULT 0")
        c.execute("ALTER TABLE nhan_vien ADD COLUMN IF NOT EXISTS han_hop_dong_thang INTEGER")
        c.execute("ALTER TABLE ung_vien ADD COLUMN IF NOT EXISTS ghi_chu TEXT")
        c.execute("UPDATE nhan_vien SET chuc_vu = 'Nhân viên' WHERE chuc_vu IS NULL OR chuc_vu = ''")
        c.execute("UPDATE nhan_vien SET ngay_qd_ns = ngay_vao_lam WHERE ngay_qd_ns IS NULL")
        c.execute("UPDATE nhan_vien SET so_luong_npt = 0 WHERE so_luong_npt IS NULL")
        db.commit()
        c.close()
        db.close()
        return True
    except Exception as e:
        print(f"Lỗi ensure_qdns_columns: {e}")
        return False


def ensure_qdns_table():
    """Tạo bảng quyet_dinh_nhan_su lưu lịch sử các Quyết định nhân sự đã ban hành."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS quyet_dinh_nhan_su (
                id SERIAL PRIMARY KEY,
                so_qd VARCHAR(50) NOT NULL,
                loai_qd VARCHAR(30) NOT NULL,
                nhan_vien_id INTEGER NOT NULL REFERENCES nhan_vien(id) ON DELETE CASCADE,
                ngay_qd DATE NOT NULL DEFAULT CURRENT_DATE,
                noi_dung TEXT,
                gia_tri_truoc VARCHAR(150),
                gia_tri_sau VARCHAR(150),
                file_url TEXT,
                nguoi_tao VARCHAR(100),
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        # Nâng cấp cho DB đã tạo bảng này từ phiên bản trước khi có đủ các cột trên
        # (CREATE TABLE IF NOT EXISTS không tự thêm cột còn thiếu vào bảng đã tồn tại)
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS so_qd VARCHAR(50)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS loai_qd VARCHAR(30)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS ngay_qd DATE DEFAULT CURRENT_DATE")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS noi_dung TEXT")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS gia_tri_truoc VARCHAR(150)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS gia_tri_sau VARCHAR(150)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS file_url TEXT")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS nguoi_tao VARCHAR(100)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS created_at TIMESTAMP DEFAULT NOW()")
        # Các cột "hiển thị" mà tab Lịch sử công tác (mục 📋 Các quyết định nhân sự) đọc ra —
        # trước đây chỉ tồn tại trên DB tạo theo schema cũ, khiến các quyết định tạo từ tab
        # "QUYẾT ĐỊNH NHÂN SỰ" (vốn chỉ ghi so_qd/loai_qd/ngay_qd) hiển thị toàn "None" vì
        # thiếu các cột này. Đảm bảo luôn có đủ trên mọi tenant.
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS so_quyet_dinh VARCHAR(50)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS loai_quyet_dinh VARCHAR(30)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS ngay_quyet_dinh DATE")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS ngay_hieu_luc DATE")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS nguoi_ky VARCHAR(150)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS chuc_danh_cu VARCHAR(150)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS chuc_danh_moi VARCHAR(150)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS phong_ban_cu VARCHAR(150)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS phong_ban_moi VARCHAR(150)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS loai_hop_dong_cu VARCHAR(50)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS loai_hop_dong_moi VARCHAR(50)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS he_so_luong_cu VARCHAR(20)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS he_so_luong_moi VARCHAR(20)")
        c.execute("ALTER TABLE quyet_dinh_nhan_su ADD COLUMN IF NOT EXISTS so_hd_cu VARCHAR(50)")
        # Sửa lỗi: bảng quyet_dinh_nhan_su trên một số DB (đã tồn tại từ trước, tạo thủ công/khác
        # phiên bản) có thêm cột "loai_quyet_dinh" (khác với cột "loai_qd" mà code hiện tại dùng)
        # và cột đó bị đặt NOT NULL. Vì INSERT ở tab "QUYẾT ĐỊNH NHÂN SỰ" chỉ điền "loai_qd" chứ
        # không điền "loai_quyet_dinh", nên gặp lỗi:
        #   null value in column "loai_quyet_dinh" ... violates not-null constraint
        # Gỡ ràng buộc NOT NULL này nếu cột tồn tại (an toàn, không ảnh hưởng dữ liệu cũ).
        c.execute("""
            DO $$
            BEGIN
                IF EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name = 'quyet_dinh_nhan_su' AND column_name = 'loai_quyet_dinh'
                ) THEN
                    ALTER TABLE quyet_dinh_nhan_su ALTER COLUMN loai_quyet_dinh DROP NOT NULL;
                END IF;
            END $$;
        """)
        db.commit()
        c.close()
        db.close()
        return True
    except Exception as e:
        print(f"Lỗi ensure_qdns_table: {e}")
        return False


# Danh sách chức vụ dùng cho QĐ Bổ nhiệm / Miễn nhiệm
DANH_SACH_CHUC_VU = ["Phó Tổng Giám đốc", "Trưởng phòng", "Phó Trưởng phòng", "Đội Trưởng", "Tổ Trưởng", "Quản đốc"]

LOAI_QDNS_LABEL = {
    'BO_NHIEM': 'QĐ Bổ nhiệm',
    'MIEN_NHIEM': 'QĐ Miễn nhiệm',
    'DOI_CHUC_DANH': 'QĐ Thay đổi chức danh',
    'DIEU_CHUYEN': 'QĐ Điều chuyển công tác',
    'CHUYEN_CHINH_THUC': 'QĐ Chuyển đổi TV → Chính thức',
    'CHAM_DUT_HD': 'QĐ Chấm dứt HĐTV/HĐLĐ',
}


def tao_quyet_dinh_nhan_su(nv, so_qd, ngay_qd, tieu_de, dieu1_lines, hieu_luc_text=None):
    """Tạo file Word Quyết định nhân sự dùng chung cho các loại: bổ nhiệm, miễn nhiệm,
    thay đổi chức danh, điều chuyển công tác, chấm dứt HĐTV/HĐLĐ."""
    CC = COMPANY_CONFIG
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sec = doc.sections[0]; sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2)

    # -- Tách tên đơn vị thành 2 dòng (loại hình + tên riêng) --
    ten_ty = CC.get('ten_cong_ty', 'CÔNG TY')
    ten_ty_lower = ten_ty.lower()
    if "công ty cổ phần" in ten_ty_lower:
        ten_doan1 = "CÔNG TY CỔ PHẦN"
        ten_doan2 = ten_ty[ten_ty_lower.index("công ty cổ phần") + len("công ty cổ phần"):].strip()
    elif "công ty tnhh" in ten_ty_lower:
        ten_doan1 = "CÔNG TY TNHH"
        ten_doan2 = ten_ty[ten_ty_lower.index("công ty tnhh") + len("công ty tnhh"):].strip()
    elif "hộ kinh doanh" in ten_ty_lower:
        ten_doan1 = "HỘ KINH DOANH"
        ten_doan2 = ten_ty[ten_ty_lower.index("hộ kinh doanh") + len("hộ kinh doanh"):].strip()
    else:
        ten_doan1 = ten_ty.upper()
        ten_doan2 = ""

    ht = doc.add_table(rows=4, cols=2); ht.alignment = WD_TABLE_ALIGNMENT.CENTER; ht.autofit = False; remove_table_border(ht)
    for row in ht.rows:
        row.cells[0].width = Cm(7); row.cells[1].width = Cm(10)
    c = ht.rows[0].cells[0]; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ten_doan1); r.bold = True; r.font.size = Pt(13)
    c = ht.rows[1].cells[0]; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ten_doan2.upper()); r.bold = True; r.font.size = Pt(13)
    c = ht.rows[0].cells[1]; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM'); r.bold = True; r.font.size = Pt(13)
    c = ht.rows[1].cells[1]; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Độc lập - Tự do - Hạnh phúc'); r.bold = True; r.italic = True; r.font.size = Pt(13)
    c = ht.rows[2].cells[0]; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Số: {so_qd}'); r.italic = True; r.font.size = Pt(12)
    c = ht.rows[3].cells[1]; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after = Pt(20)
    dia_diem = CC.get('dia_diem') or get_cau_hinh('dia_diem', 'Quảng Trị')
    ns = f'{dia_diem}, ngày {ngay_qd.day} tháng {ngay_qd.month:02d} năm {ngay_qd.year}' if hasattr(ngay_qd, 'day') else f'{dia_diem}, ngày ... tháng ... năm ......'
    r = p.add_run(ns); r.italic = True; r.font.size = Pt(13)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('QUYẾT ĐỊNH'); r.bold = True; r.font.size = Pt(18)
    force_center(p)
    p = doc.add_paragraph()
    r = p.add_run(f'Về việc: {tieu_de}'); r.bold = True; r.italic = True; r.font.size = Pt(14)
    force_center(p)
    p = doc.add_paragraph()
    chuc_vu_qdns = CC.get('chuc_vu', 'GIÁM ĐỐC').upper()
    r = p.add_run(f"{chuc_vu_qdns} {ten_doan1} {ten_doan2}".strip().upper()); r.bold = True
    force_center(p)

    doc.add_paragraph('- Căn cứ Bộ luật Lao động số 45/2019/QH14 ngày 20/11/2019;')
    if "hộ kinh doanh" in ten_ty_lower:
        doc.add_paragraph('- Căn cứ Giấy chứng nhận đăng ký hộ kinh doanh;')
    else:
        doc.add_paragraph('- Căn cứ Điều lệ tổ chức và hoạt động của Công ty;')
    doc.add_paragraph('- Căn cứ nhu cầu công tác và năng lực cán bộ, nhân viên;')
    phong_to_chuc_ns = get_phong_to_chuc_nhan_su()
    doc.add_paragraph(f'- Xét đề nghị của {phong_to_chuc_ns},')

    p = doc.add_paragraph(); r = p.add_run('QUYẾT ĐỊNH:'); r.bold = True
    force_center(p)

    p = doc.add_paragraph(); r = p.add_run('Điều 1. '); r.bold = True
    r2 = p.add_run(dieu1_lines[0] if dieu1_lines else '')
    for extra_line in dieu1_lines[1:]:
        doc.add_paragraph(extra_line)

    p = doc.add_paragraph(); r = p.add_run('Điều 2. '); r.bold = True
    r2 = p.add_run(hieu_luc_text or f'Quyết định này có hiệu lực kể từ ngày {ngay_qd.day}/{ngay_qd.month:02d}/{ngay_qd.year}.')

    p = doc.add_paragraph(); r = p.add_run('Điều 3. '); r.bold = True
    r2 = p.add_run(f"Ông/Bà {nv.get('ho_ten', '')}, Trưởng {phong_to_chuc_ns} và các bộ phận có liên quan chịu trách nhiệm thi hành Quyết định này./.")

    doc.add_paragraph('')
    t2 = doc.add_table(rows=2, cols=2); remove_table_border(t2)
    p_nn = t2.rows[0].cells[0].paragraphs[0]
    r = p_nn.add_run('Nơi nhận:'); r.bold = True; r.font.size = Pt(11)
    p_nn2 = t2.rows[0].cells[0].add_paragraph()
    r = p_nn2.add_run('- Như Điều 3;\n- Lưu: VT, HSNV.'); r.font.size = Pt(11)
    p_kh = t2.rows[0].cells[1].paragraphs[0]; p_kh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_kh.add_run(CC.get('chuc_vu', 'GIÁM ĐỐC').upper()); r.bold = True
    p_kh2 = t2.rows[1].cells[1].add_paragraph(); p_kh2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_kh2.add_run('\n\n\n' + CC.get('dai_dien', ''))

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    return tmp.name


if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.role = None
    st.session_state.username = None
if 'selected_nv_id' not in st.session_state:
    st.session_state.selected_nv_id = None
if 'edit_uv_id' not in st.session_state:
    st.session_state.edit_uv_id = None
if 'bhxh_family_nv_id' not in st.session_state:
    st.session_state.bhxh_family_nv_id = None
if 'bhxh_family_nv_name' not in st.session_state:
    st.session_state.bhxh_family_nv_name = None
if 'bhxh_family_members' not in st.session_state:
    st.session_state.bhxh_family_members = []
if 'show_chuyen_nv_form' not in st.session_state:
    st.session_state.show_chuyen_nv_form = False
if 'chuyen_uv_id' not in st.session_state:
    st.session_state.chuyen_uv_id = None
if 'chuyen_uv_data' not in st.session_state:
    st.session_state.chuyen_uv_data = {}

def force_center(p):
    pPr = p._p.get_or_add_pPr()
    for jc in pPr.findall(qn('w:jc')):
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

st.markdown("""
<style>
    [data-testid="stDataFrame"] > div {
        overflow-x: auto !important;
    }
    [data-testid="stDataFrame"] table {
        min-width: 2000px !important;
        width: max-content !important;
    }
    /* ===== Bảng chấm công (BCC): auto center Horizontal + Vertical, giảm size chữ ===== */
    [data-testid="stDataFrame"] th,
    [data-testid="stDataFrame"] td {
        text-align: center !important;
        vertical-align: middle !important;
        font-size: 12px !important;
    }
    [data-testid="stDataFrame"] [data-testid="stElementToolbar"] { font-size: 12px !important; }
    /* st.data_editor (bảng chấm công dạng edit) dùng cùng component nền glide-data-grid */
    [data-testid="stDataEditor"] th,
    [data-testid="stDataEditor"] td {
        text-align: center !important;
        vertical-align: middle !important;
        font-size: 12px !important;
    }
</style>
""", unsafe_allow_html=True)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def to_int_or_none(val):
    """Chuyển đổi giá trị sang int hoặc None"""
    if val is None or str(val).strip() == '':
        return None
    try:
        return int(float(val))
    except:
        return None
        
def tao_noi_dung_zalo(nv):
    ZC = COMPANY_CONFIG
    return f"""Gửi anh/chị: {nv.get('ho_ten','')},

Thông tin đã cập nhật:
- Họ tên: {nv.get('ho_ten','')}
- Ngày sinh: {format_date(nv.get('ngay_sinh'))}
- CCCD: {nv.get('so_cccd','')}
- Ngày cấp: {format_date(nv.get('ngay_cap_cccd'))}
- Thường trú: {nv.get('thuong_tru','')}
- Số BHXH: {nv.get('ma_so_bhxh','')}
- TK NH: {nv.get('so_tai_khoan_nh','')}
- CN NH: {nv.get('chi_nhanh_nh','')}
- Tên đơn vị thụ hưởng: {nv.get('ten_don_vi_thu_huong','')}

{ZC.get('loi_nhan_zalo','Vui lòng kiểm tra và phản hồi nếu có sai sót. Xin Cảm ơn!')}"""

def remove_table_border(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            b = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
            if b is not None: tcPr.remove(b)

# ========== CÁC HÀM TẠO HỢP ĐỒNG (GIỮ NGUYÊN) ==========
def ensure_mau_dieu_hop_dong_table():
    """Bảng lưu nội dung tuỳ chỉnh của từng Điều trong HĐLĐ/HĐTV, do admin cấu hình.
    Mỗi dòng = 1 Điều của 1 loại hợp đồng. Nếu chưa có dòng nào cho 1 mã Điều,
    hệ thống dùng nội dung mặc định (DEFAULT_DIEU_HDLD / DEFAULT_DIEU_HDTV)."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS mau_dieu_hop_dong (
                id SERIAL PRIMARY KEY,
                loai_hd VARCHAR(10) NOT NULL,
                ma_dieu VARCHAR(30) NOT NULL,
                tieu_de TEXT,
                noi_dung TEXT,
                thu_tu INT DEFAULT 0,
                updated_at TIMESTAMP DEFAULT NOW(),
                UNIQUE(loai_hd, ma_dieu)
            )
        """)
        db.commit()
        c.close(); db.close()
        return True
    except Exception as e:
        print(f"Lỗi ensure_mau_dieu_hop_dong_table: {e}")
        return False

def ensure_chuc_danh_ung_vien_table():
    """Danh mục 'Vị trí dự tuyển' RIÊNG cho Ứng viên - độc lập với vi_tri_cong_tac (chức danh
    Nhân viên), để đổi danh mục chức danh Nhân viên không làm mất khả năng tìm/nhập ứng viên
    theo các chức danh cũ. Tự động khởi tạo dữ liệu từ các giá trị vi_tri_du_tuyen đã có sẵn
    trong bảng ung_vien (chỉ chạy 1 lần, khi bảng còn trống)."""
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS chuc_danh_ung_vien (
                id SERIAL PRIMARY KEY,
                ten_chuc_danh VARCHAR(150) UNIQUE NOT NULL,
                ghi_chu TEXT,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        db.commit()
        c.execute("SELECT COUNT(*) FROM chuc_danh_ung_vien")
        so_luong = c.fetchone()[0]
        if so_luong == 0:
            # Nạp dữ liệu ban đầu từ các vị trí dự tuyển đã có trong ung_vien (chức danh cũ)
            c.execute("""SELECT DISTINCT vi_tri_du_tuyen FROM ung_vien 
                         WHERE vi_tri_du_tuyen IS NOT NULL AND vi_tri_du_tuyen != ''""")
            ds_cu = [row[0] for row in c.fetchall()]
            for ten in ds_cu:
                c.execute("INSERT INTO chuc_danh_ung_vien (ten_chuc_danh) VALUES (%s) ON CONFLICT DO NOTHING", (ten,))
            db.commit()
        c.close(); db.close()
        return True
    except Exception as e:
        print(f"Lỗi ensure_chuc_danh_ung_vien_table: {e}")
        return False

# Nội dung MẶC ĐỊNH của từng Điều — dùng khi admin CHƯA tuỳ chỉnh gì.
# Dòng bắt đầu bằng "## " sẽ được in đậm (tiêu đề phụ, VD "1. Nghĩa vụ:").
# Có thể dùng {vi_tri}, {ngay_hieu_luc}, {ten_cong_ty} trong nội dung các Điều 1, 5 — sẽ tự thay bằng thông tin nhân viên.
DEFAULT_DIEU_HDLD = {
    "can_cu": ("Căn cứ pháp lý:",
        "- Căn cứ Bộ luật Lao động số 45/2019/QH14 ngày 20/11/2019;\n"
        "- Căn cứ nhu cầu công việc và sự thỏa thuận giữa hai bên."),
    "dieu1": ("Điều 1. Thời hạn và công việc hợp đồng:",
        "-    Bên B làm việc theo chế độ hợp đồng lao động không xác định thời hạn;\n"
        "-    Thời gian: Từ ngày {ngay_hieu_luc};\n"
        "-    Địa điểm làm việc: Tại {ten_cong_ty} và các địa điểm khác theo sự sắp xếp của Công ty;\n"
        "-    Vị trí: {vi_tri};\n"
        "-    Công việc phải làm: Thực hiện công việc theo đúng chuyên môn dưới sự quản lý, điều hành của cấp trên;\n"
        "-    Mức lương và phụ cấp: Theo thỏa thuận;\n"
        "-    Hình thức trả lương: Tiền mặt hoặc chuyển khoản, theo lần chi trả;\n"
        "-    Kỳ hạn trả lương: Theo quy định Công ty;\n"
        "-    Chế độ nâng lương: Theo thỏa thuận."),
    "dieu2": ("Điều 2. Chế độ làm việc:",
        "-    Thời gian làm việc: Theo tính chất công việc, do nhu cầu kinh doanh của Công ty nên thời gian làm việc của bên B là linh hoạt nhưng phải đảm bảo hoàn thành công việc được giao;\n"
        "-    Thời gian nghỉ ngơi của người lao động: Theo thỏa thuận và phù hợp với quy định của pháp luật;\n"
        "-    Ngoài giờ làm việc: Người lao động phải tự chịu trách nhiệm về các hoạt động cá nhân của mình."),
    "dieu3": ("Điều 3. Nghĩa vụ, quyền lợi NLĐ:",
        "## 1. Nghĩa vụ:\n"
        "-    Hoàn thành những công việc được giao và sẵn sàng chấp nhận mọi sự điều động khi có yêu cầu;\n"
        "-    Chấp hành nghiêm túc nội quy, kỷ luật lao động, an toàn lao động và các quy định của Công ty và pháp luật của Nhà nước;\n"
        "-    Người lao động có trách nhiệm tuân thủ đầy đủ quy định về an toàn lao động, quy trình vận hành thiết bị và hướng dẫn của Công ty. Trường hợp NLĐ cố ý vi phạm hoặc vi phạm nghiêm trọng quy định an toàn lao động gây thiệt hại thì phải chịu trách nhiệm theo quy định pháp luật và nội quy Công ty;\n"
        "-    Bồi thường vi phạm vật chất : Phải bồi thường vật chất do cá nhân vi phạm quy định của Công ty về bảo quản trang thiết bị được giao.\n"
        "## 2. Quyền Lợi:\n"
        "-    Phương tiện đi lại: Tự túc;\n"
        "-    Được Công ty đóng Bảo hiểm xã hội, bảo hiểm y tế, BHTN: theo chế độ hiện hành của Nhà nước và Quy định của Công ty;\n"
        "-    Được Công ty cấp đầy đủ bảo hộ lao động theo đúng vị trí làm việc;\n"
        "-    Được phân công công việc theo yêu cầu của Công ty phù hợp với khả năng và trình độ chuyên môn mà người lao động đáp ứng;\n"
        "-    Các quyền lợi khác thực hiện theo quy định của Pháp luật Lao động như tạm dừng, chấm dứt hợp đồng."),
    "dieu4": ("Điều 4. Nghĩa vụ, quyền hạn NSDLĐ:",
        "-    Bảo đảm việc làm và thực hiện đầy đủ những điều đã cam kết trong hợp đồng;\n"
        "-    Thanh toán đầy đủ, đúng hạn các chế độ và quyền lợi cho người lao động theo hợp đồng;\n"
        "-    Điều hành người lao động hoàn thành công việc theo hợp đồng;\n"
        "-    Tạm hoãn, chấm dứt hợp đồng, kỷ luật người lao động theo quy định của pháp luật, và nội quy lao động của Công ty."),
    "dieu5": ("Điều 5. Điều khoản chung:",
        "-    Những nội dung về quan hệ lao động không ghi trong hợp đồng này thì được áp dụng theo pháp luật lao động;\n"
        "-    Những thoả thuận khác (nếu có): không;\n"
        "-    Hợp đồng này có hiệu lực từ ngày ký và được làm thành 02 bản, Bên A giữ 01 bản, Bên B giữ 01 có giá trị pháp lý như nhau, để làm căn cứ thực hiện;\n"
        "-    Bản Hợp đồng này được lập tại văn phòng {ten_cong_ty}."),
}

DEFAULT_DIEU_HDTV = {
    "can_cu": ("Căn cứ pháp lý:",
        "- Căn cứ Bộ luật Lao động số 45/2019/QH14 ngày 20/11/2019;\n"
        "- Căn cứ nhu cầu công việc và sự thỏa thuận giữa hai bên."),
    "dieu1": ("Điều 1. Thời hạn và công việc hợp đồng:",
        "-    Bên B làm việc theo chế độ hợp đồng thử việc, có thời hạn 01 tháng;\n"
        "-    Bắt đầu: {ngay_bat_dau};\n"
        "-    Kết thúc: {ngay_ket_thuc};\n"
        "-    Địa điểm làm việc: Tại {ten_cong_ty} và các địa điểm khác theo sự sắp xếp của Công ty;\n"
        "-    Vị trí: {vi_tri};\n"
        "-    Công việc phải làm: Thực hiện công việc theo đúng chuyên môn dưới sự quản lý, điều hành của cấp trên;\n"
        "-    Mức lương và phụ cấp: Theo thỏa thuận;\n"
        "-    Hình thức trả lương: Tiền mặt hoặc chuyển khoản, theo lần chi trả;\n"
        "-    Kỳ hạn trả lương: Theo quy định Công ty."),
    "dieu2": ("Điều 2. Chế độ làm việc:",
        "-    Thời gian làm việc: Theo tính chất công việc, do nhu cầu kinh doanh của Công ty nên thời gian làm việc của bên B là linh hoạt nhưng phải đảm bảo hoàn thành công việc được giao;\n"
        "-    Thời gian nghỉ ngơi của người lao động: Theo thỏa thuận và phù hợp với quy định của pháp luật;\n"
        "-    Ngoài giờ làm việc: Người lao động phải tự chịu trách nhiệm về các hoạt động cá nhân của mình."),
    "dieu3": ("Điều 3. Nghĩa vụ, quyền lợi NLĐ:",
        "## 1. Nghĩa vụ:\n"
        "-    Hoàn thành những công việc được giao và sẵn sàng chấp nhận mọi sự điều động khi có yêu cầu;\n"
        "-    Chấp hành nghiêm túc nội quy, kỷ luật lao động, an toàn lao động và các quy định của Công ty và pháp luật của Nhà nước;\n"
        "-    Người lao động có trách nhiệm tuân thủ đầy đủ quy định về an toàn lao động, quy trình vận hành thiết bị và hướng dẫn của Công ty. Trường hợp NLĐ cố ý vi phạm hoặc vi phạm nghiêm trọng quy định an toàn lao động gây thiệt hại thì phải chịu trách nhiệm theo quy định pháp luật và nội quy Công ty;\n"
        "-    Bồi thường vi phạm vật chất : Phải bồi thường vật chất do cá nhân vi phạm quy định của Công ty về bảo quản trang thiết bị được giao.\n"
        "## 2. Quyền Lợi:\n"
        "-    Phương tiện đi lại: Tự túc;\n"
        "-    Được Công ty cấp đầy đủ bảo hộ lao động theo đúng vị trí làm việc;\n"
        "-    Được phân công công việc theo yêu cầu của Công ty phù hợp với khả năng và trình độ chuyên môn mà người lao động đáp ứng;\n"
        "-    Các quyền lợi khác thực hiện theo quy định của Pháp luật Lao động như tạm dừng, chấm dứt hợp đồng."),
    "dieu4": ("Điều 4. Nghĩa vụ, quyền hạn NSDLĐ:",
        "-    Bảo đảm việc làm và thực hiện đầy đủ những điều đã cam kết trong hợp đồng;\n"
        "-    Thanh toán đầy đủ, đúng hạn các chế độ và quyền lợi cho người lao động theo hợp đồng;\n"
        "-    Điều hành người lao động hoàn thành công việc theo hợp đồng;\n"
        "-    Tạm hoãn, chấm dứt hợp đồng theo quy định của pháp luật, và nội quy lao động của Công ty;"),
    "dieu5": ("Điều 5. Điều khoản chung:",
        "-    Những nội dung về quan hệ lao động không ghi trong hợp đồng này thì được áp dụng theo pháp luật lao động;\n"
        "-    Những thoả thuận khác (nếu có): không;\n"
        "-    Hợp đồng này có hiệu lực từ ngày ký và được làm thành 02 bản, Bên A giữ 01 bản, Bên B giữ 01 có giá trị pháp lý như nhau, để làm căn cứ thực hiện;\n"
        "-    Bản Hợp đồng này được lập tại văn phòng {ten_cong_ty}."),
}

class _SafeDict(dict):
    """Dict để .format_map() không lỗi khi thiếu placeholder — giữ nguyên {ten} nếu không có dữ liệu."""
    def __missing__(self, key):
        return '{' + key + '}'

@st.cache_data(ttl=60, show_spinner=False)
def get_all_dieu_hop_dong(loai_hd):
    """Lấy toàn bộ nội dung Điều đã tuỳ chỉnh (nếu có) cho 1 loại hợp đồng ('HDLD'/'HDTV').
    Trả về dict {ma_dieu: (tieu_de, noi_dung, thu_tu)}."""
    ket_qua = {}
    try:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor()
        c.execute("SELECT ma_dieu, tieu_de, noi_dung, thu_tu FROM mau_dieu_hop_dong WHERE loai_hd=%s ORDER BY thu_tu ASC, ma_dieu ASC", (loai_hd,))
        for ma_dieu, tieu_de, noi_dung, thu_tu in c.fetchall():
            ket_qua[ma_dieu] = (tieu_de, noi_dung, thu_tu)
        db.close()
    except Exception:
        pass
    return ket_qua

def get_ds_ma_dieu(tuy_chinh):
    """Trả về danh sách mã Điều theo đúng thứ tự hiển thị/in ấn: 5 Điều mặc định (dieu1..dieu5)
    luôn giữ nguyên thứ tự gốc, cộng thêm các Điều admin tự thêm mới (không giới hạn số lượng)
    được sắp xếp chèn theo cột thu_tu."""
    mac_dinh_keys = ["can_cu", "dieu1", "dieu2", "dieu3", "dieu4", "dieu5"]
    them_moi = [(md, (info[2] if len(info) > 2 and info[2] else 999)) 
                for md, info in tuy_chinh.items() if md not in mac_dinh_keys]
    them_moi.sort(key=lambda x: (x[1], x[0]))
    return mac_dinh_keys + [md for md, _ in them_moi]

def sinh_ma_dieu_moi(tuy_chinh_hdld, tuy_chinh_hdtv):
    """Sinh mã Điều mới tự động (dieu6, dieu7, ...) không trùng với bất kỳ Điều nào đã có
    ở CẢ 2 loại hợp đồng, để tránh nhầm lẫn khi admin chuyển qua lại giữa HĐLĐ/HĐTV."""
    so_hien_co = [5]
    for md in list(tuy_chinh_hdld.keys()) + list(tuy_chinh_hdtv.keys()):
        if md.startswith("dieu") and md[4:].isdigit():
            so_hien_co.append(int(md[4:]))
    return f"dieu{max(so_hien_co) + 1}"

def get_dieu_content(loai_hd, ma_dieu, tuy_chinh, mac_dinh):
    """Trả về (tieu_de, noi_dung) — ưu tiên bản admin đã tuỳ chỉnh, nếu chưa có thì dùng mặc định."""
    if ma_dieu in tuy_chinh:
        info = tuy_chinh[ma_dieu]
        return info[0], info[1]
    return mac_dinh.get(ma_dieu, ("", ""))

def render_dieu(doc, add_p, tieu_de, noi_dung, context=None):
    """In 1 Điều ra file Word: tiêu đề in đậm, các dòng nội dung xuống dòng theo \\n.
    Dòng bắt đầu bằng '## ' sẽ in đậm (tiêu đề phụ như '1. Nghĩa vụ:')."""
    if tieu_de:
        p = doc.add_paragraph(); r = p.add_run(tieu_de); r.bold = True
    if not noi_dung:
        return
    text = noi_dung.format_map(_SafeDict(context or {}))
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('## '):
            p = doc.add_paragraph(); r = p.add_run(line[3:]); r.bold = True
        else:
            add_p(line)


def _set_ho_ten_row_header(ht, ten_cong_ty):
    """Điền tên công ty (row 0, col 0) + 'CỘNG HÒA XÃ HỘI...' (row 0, col 1) vào bảng header
    của hợp đồng, xử lý trường hợp tên công ty dài bị tràn xuống 2 dòng:
    - Tự giảm cỡ chữ khi tên công ty dài để hạn chế bị xuống dòng.
    - Căn giữa theo chiều dọc (vertical center) cho TẤT CẢ 4 ô ở row 0 & row 1, để nếu
      tên công ty vẫn phải xuống 2 dòng thì khoảng trắng thừa được chia đều lên trên & xuống
      dưới thay vì dồn hết xuống dưới cột bên cạnh (nhìn lệch)."""
    ten = str(ten_cong_ty or '').strip()
    size = Pt(13)
    if len(ten) > 34:
        size = Pt(11)
    elif len(ten) > 24:
        size = Pt(12)
    c = ht.rows[0].cells[0]
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ten); r.bold = True; r.font.size = size
    for row_idx in (0, 1):
        for col_idx in (0, 1):
            ht.rows[row_idx].cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def tao_hop_dong(nv):
    """In Hợp đồng lao động (không xác định thời hạn). Nội dung 5 Điều lấy từ bảng
    mau_dieu_hop_dong nếu admin đã tuỳ chỉnh (Danh mục → Mẫu Điều khoản Hợp đồng),
    nếu chưa có thì dùng nội dung mặc định DEFAULT_DIEU_HDLD."""
    CC = COMPANY_CONFIG; doc = Document()
    s = doc.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
    s.paragraph_format.space_after=Pt(0); s.paragraph_format.space_before=Pt(0)
    sec = doc.sections[0]; sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(3.5); sec.right_margin=Cm(2)
    def add_p(text='', bold=False, size=Pt(13)):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if bold and p.runs:
            p.runs[0].bold = True
        if p.runs:
            p.runs[0].font.size = size
        return p
    def al(label,value):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(5))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r=p.add_run(f'{label}'); r.font.size=Pt(13)
        r=p.add_run('\t: '); r.font.size=Pt(13)
        r=p.add_run(f'{value}'); r.font.size=Pt(13)
    
    ten_ty = CC["ten_cong_ty"]
    # Kiểm tra loại hình doanh nghiệp (không phân biệt hoa/thường)
    ten_ty_lower = ten_ty.lower()
    if "công ty cổ phần" in ten_ty_lower:
        ten_doan1 = "CÔNG TY CỔ PHẦN"
        idx = ten_ty_lower.index("công ty cổ phần") + len("công ty cổ phần")
        ten_doan2 = ten_ty[idx:].strip()
    elif "công ty tnhh" in ten_ty_lower:
        ten_doan1 = "CÔNG TY TNHH"
        idx = ten_ty_lower.index("công ty tnhh") + len("công ty tnhh")
        ten_doan2 = ten_ty[idx:].strip()
    elif "hộ kinh doanh" in ten_ty_lower:
        ten_doan1 = "HỘ KINH DOANH"
        idx = ten_ty_lower.index("hộ kinh doanh") + len("hộ kinh doanh")
        ten_doan2 = ten_ty[idx:].strip()
    else:
        ten_doan1 = ten_ty.upper()
        ten_doan2 = ""
    
    ht=doc.add_table(rows=4,cols=2); ht.alignment=WD_TABLE_ALIGNMENT.CENTER; ht.autofit=False; remove_table_border(ht)
    for row in ht.rows: row.cells[0].width=Cm(7); row.cells[1].width=Cm(10)
    c=ht.rows[0].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run((ten_doan1).upper()); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[1].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run((ten_doan2).upper()); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[0].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM'); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[1].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Độc lập - Tự do - Hạnh phúc'); r.bold=True; r.italic=True; r.font.size=Pt(13)
    c=ht.rows[2].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('─'*9); r.font.size=Pt(9)
    c=ht.rows[2].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('─'*20); r.font.size=Pt(9)
    c=ht.rows[3].cells[1]; p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after=Pt(20)
    ngay_ky = nv.get("ngay_ky_hd")
    ngay_vao = nv.get("ngay_vao_lam")
    nk = ngay_ky if ngay_ky else ngay_vao
    dia_diem_ky = CC.get('dia_diem') or get_cau_hinh('dia_diem', 'Quảng Trị')
    ns = f'{dia_diem_ky}, ngày ... tháng ... năm ......'
    if nk and hasattr(nk, 'day'):
        ns = f'{dia_diem_ky}, ngày {nk.day} tháng {nk.month:02d} năm {nk.year}'
    run = p.add_run(ns)
    run.font.size = Pt(13)
    run.italic = True
    c=ht.rows[3].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'Số: {nv.get("so_hdld","...")}'); r.italic=True; r.font.size=Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run('HỢP ĐỒNG LAO ĐỘNG')
    r.bold = True
    r.font.size = Pt(18)
    force_center(p)
    # ===== Lấy nội dung Điều từ DB (admin tuỳ chỉnh) hoặc mặc định =====
    tuy_chinh_hdld = get_all_dieu_hop_dong('HDLD')
    ngay_hieu_luc = nv.get("ngay_ky_hd") or nv.get("ngay_vao_lam")
    ns2 = '.../.../..........'
    if ngay_hieu_luc and hasattr(ngay_hieu_luc, 'day'):
        ns2 = f'{ngay_hieu_luc.day} tháng {ngay_hieu_luc.month:02d} năm {ngay_hieu_luc.year}'
    elif ngay_hieu_luc:
        ns2 = str(ngay_hieu_luc)
    ctx_hdld = {"vi_tri": nv.get("chuc_danh_nghe", ""), "ngay_hieu_luc": ns2, "ten_cong_ty": CC.get("ten_cong_ty", "")}
    # --- Render "Căn cứ pháp lý" (từ template, trước BÊN A/B) ---
    tieu_de_cc, noi_dung_cc = get_dieu_content("HDLD", "can_cu", tuy_chinh_hdld, DEFAULT_DIEU_HDLD)
    if tieu_de_cc or noi_dung_cc:
        render_dieu(doc, add_p, tieu_de_cc, noi_dung_cc, context=ctx_hdld)
    doc.add_paragraph('Chúng tôi gồm:')
    p=doc.add_paragraph(); r=p.add_run(f'BÊN A: {CC["ten_cong_ty"]} (Người sử dụng LĐ)'); r.bold=True
    al('Đại diện',f"Ông {CC['dai_dien']}"); al('Chức vụ',CC['chuc_vu']); al('Mã số thuế',CC['ma_so_thue'])
    al('Điện thoại',CC['dien_thoai_cty']); al('Địa chỉ',CC['dia_chi']); doc.add_paragraph()
    p=doc.add_paragraph(); r=p.add_run('BÊN B: (Người lao động)'); r.bold=True
    sk=nv.get('so_tai_khoan_nh','')
    if nv.get('chi_nhanh_nh'): sk+=f' - {nv.get("chi_nhanh_nh")}'
    gt = nv.get('gioi_tinh','')
    xung_ho = get_xung_ho_trang_trong(gt)
    al(xung_ho, nv.get('ho_ten',''))
    al('Ngày sinh',format_date(nv.get('ngay_sinh')))
    al('Số CMND/CCCD',nv.get('so_cccd','')); al('Ngày cấp',format_date(nv.get('ngay_cap_cccd')))
    al('Nơi cấp',nv.get('noi_cap_cccd','')); al('Số TKNH',sk)
    al('Điện thoại',nv.get('dien_thoai','')); al('Thường trú',nv.get('thuong_tru',''))
    doc.add_paragraph('Thoả thuận ký kết Hợp đồng lao động với những điều khoản dưới đây:')
    # --- Render các Điều còn lại (bỏ qua can_cu vì đã render ở trên) ---
    for ma_dieu in get_ds_ma_dieu(tuy_chinh_hdld):
        if ma_dieu == "can_cu":
            continue
        tieu_de, noi_dung = get_dieu_content("HDLD", ma_dieu, tuy_chinh_hdld, DEFAULT_DIEU_HDLD)
        if not tieu_de and not noi_dung:
            continue
        render_dieu(doc, add_p, tieu_de, noi_dung, context=ctx_hdld)
    doc.add_paragraph()
    ts=doc.add_table(rows=3,cols=2); ts.alignment=WD_TABLE_ALIGNMENT.CENTER; remove_table_border(ts)
    c=ts.rows[0].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run('NGƯỜI LAO ĐỘNG'); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[0].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run('NGƯỜI SỬ DỤNG LAO ĐỘNG'); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[1].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].add_run('').font.size=Pt(12); sp=c.add_paragraph(); sp.paragraph_format.space_after=Pt(60)
    c=ts.rows[1].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].add_run('').font.size=Pt(12); sp=c.add_paragraph(); sp.paragraph_format.space_after=Pt(60)
    c=ts.rows[2].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run(nv.get('ho_ten','').upper()); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[2].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run(CC['dai_dien'].upper()); r.bold=True; r.font.size=Pt(13)
    tf=tempfile.NamedTemporaryFile(delete=False,suffix='.docx'); doc.save(tf.name); return tf.name

def tao_hop_dong_thu_viec(nv):
    # ... giữ nguyên code cũ (tương tự)
    CC = COMPANY_CONFIG; doc = Document()
    s = doc.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
    s.paragraph_format.space_after=Pt(0); s.paragraph_format.space_before=Pt(0)
    sec = doc.sections[0]; sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(3.5); sec.right_margin=Cm(2)
    def add_p(text='', bold=False, size=Pt(13)):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if bold and p.runs:
            p.runs[0].bold = True
        if p.runs:
            p.runs[0].font.size = size
        return p
    def al(label,value):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(5))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r=p.add_run(f'{label}'); r.font.size=Pt(13)
        r=p.add_run('\t: '); r.font.size=Pt(13)
        r=p.add_run(f'{value}'); r.font.size=Pt(13)
        
    ten_ty = CC["ten_cong_ty"]
    # Kiểm tra loại hình doanh nghiệp (không phân biệt hoa/thường)
    ten_ty_lower = ten_ty.lower()
    if "công ty cổ phần" in ten_ty_lower:
        ten_doan1 = "CÔNG TY CỔ PHẦN"
        idx = ten_ty_lower.index("công ty cổ phần") + len("công ty cổ phần")
        ten_doan2 = ten_ty[idx:].strip()
    elif "công ty tnhh" in ten_ty_lower:
        ten_doan1 = "CÔNG TY TNHH"
        idx = ten_ty_lower.index("công ty tnhh") + len("công ty tnhh")
        ten_doan2 = ten_ty[idx:].strip()
    elif "hộ kinh doanh" in ten_ty_lower:
        ten_doan1 = "HỘ KINH DOANH"
        idx = ten_ty_lower.index("hộ kinh doanh") + len("hộ kinh doanh")
        ten_doan2 = ten_ty[idx:].strip()
    else:
        ten_doan1 = ten_ty.upper()
        ten_doan2 = ""
    
    ht=doc.add_table(rows=4,cols=2); ht.alignment=WD_TABLE_ALIGNMENT.CENTER; ht.autofit=False; remove_table_border(ht)
    for row in ht.rows: row.cells[0].width=Cm(7); row.cells[1].width=Cm(10)
    c=ht.rows[0].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run((ten_doan1).upper()); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[1].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run((ten_doan2).upper()); r.bold=True; r.font.size=Pt(13)
    
    c=ht.rows[0].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM'); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[1].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Độc lập - Tự do - Hạnh phúc'); r.bold=True; r.italic=True; r.font.size=Pt(13)
    c=ht.rows[2].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('─'*9); r.font.size=Pt(9)
    c=ht.rows[2].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('─'*20); r.font.size=Pt(9)
    c=ht.rows[3].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'Số: {nv.get("so_hdld",".../.../HĐTV-CHL")}'); r.italic=True; r.font.size=Pt(12)
    c=ht.rows[3].cells[1]; p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after=Pt(20)
    nk=nv.get("ngay_vao_lam") or nv.get("ngay_ky_hd")
    dia_diem_ky_tv = CC.get('dia_diem') or get_cau_hinh('dia_diem', 'Quảng Trị')
    ns=f'{dia_diem_ky_tv}, ngày ... tháng ... năm ......'
    if nk and hasattr(nk,'day'): ns=f'{dia_diem_ky_tv}, ngày {nk.day} tháng {nk.month:02d} năm {nk.year}'
    run = p.add_run(ns)
    run.font.size = Pt(13)
    run.italic = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run('HỢP ĐỒNG THỬ VIỆC')
    r.bold = True
    r.font.size = Pt(18)
    force_center(p)
    # ===== Lấy nội dung Điều từ DB (admin tuỳ chỉnh) hoặc mặc định =====
    nkt = nk + timedelta(days=30) if (nk and hasattr(nk, 'day')) else None
    ns_bd = f'{nk.day:02d}/{nk.month:02d}/{nk.year}' if (nk and hasattr(nk, 'day')) else '.../.../......'
    ns_kt = f'{nkt.day:02d}/{nkt.month:02d}/{nkt.year}' if nkt else '.../.../......'
    tuy_chinh_hdtv = get_all_dieu_hop_dong('HDTV')
    ctx_hdtv = {"vi_tri": nv.get("chuc_danh_nghe", ""), "ngay_bat_dau": ns_bd, "ngay_ket_thuc": ns_kt, "ten_cong_ty": CC.get("ten_cong_ty", "")}
    # --- Render "Căn cứ pháp lý" (từ template, trước BÊN A/B) ---
    tieu_de_cc, noi_dung_cc = get_dieu_content("HDTV", "can_cu", tuy_chinh_hdtv, DEFAULT_DIEU_HDTV)
    if tieu_de_cc or noi_dung_cc:
        render_dieu(doc, add_p, tieu_de_cc, noi_dung_cc, context=ctx_hdtv)
    doc.add_paragraph('Chúng tôi gồm:')
    p=doc.add_paragraph(); r=p.add_run(f'BÊN A: {CC["ten_cong_ty"]} (Người sử dụng LĐ)'); r.bold=True
    al('Đại diện',f"Ông {CC['dai_dien']}"); al('Chức vụ',CC['chuc_vu']); al('Mã số thuế',CC['ma_so_thue'])
    al('Điện thoại',CC['dien_thoai_cty']); al('Địa chỉ',CC['dia_chi'])
    p=doc.add_paragraph(); r=p.add_run('BÊN B: (Người lao động)'); r.bold=True
    sk=nv.get('so_tai_khoan_nh','')
    if nv.get('chi_nhanh_nh'): sk+=f' - {nv.get("chi_nhanh_nh")}'
    gt = nv.get('gioi_tinh','')
    xung_ho = get_xung_ho_trang_trong(gt)
    al(xung_ho, nv.get('ho_ten',''))
    al('Ngày sinh',format_date(nv.get('ngay_sinh')))
    al('Số CMND/CCCD',nv.get('so_cccd','')); al('Ngày cấp',format_date(nv.get('ngay_cap_cccd')))
    al('Nơi cấp',nv.get('noi_cap_cccd','')); al('Số TKNH',sk)
    al('Điện thoại',nv.get('dien_thoai','')); al('Thường trú',nv.get('thuong_tru',''))
    doc.add_paragraph('Thoả thuận ký kết Hợp đồng Thử việc với những điều khoản dưới đây:')
    # --- Render các Điều còn lại (bỏ qua can_cu vì đã render ở trên) ---
    for ma_dieu in get_ds_ma_dieu(tuy_chinh_hdtv):
        if ma_dieu == "can_cu":
            continue
        tieu_de, noi_dung = get_dieu_content("HDTV", ma_dieu, tuy_chinh_hdtv, DEFAULT_DIEU_HDTV)
        if not tieu_de and not noi_dung:
            continue
        render_dieu(doc, add_p, tieu_de, noi_dung, context=ctx_hdtv)
    doc.add_paragraph()
    ts=doc.add_table(rows=3,cols=2); ts.alignment=WD_TABLE_ALIGNMENT.CENTER; remove_table_border(ts)
    c=ts.rows[0].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run('NGƯỜI LAO ĐỘNG'); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[0].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run('NGƯỜI SỬ DỤNG LAO ĐỘNG'); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[1].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].add_run('').font.size=Pt(12); sp=c.add_paragraph(); sp.paragraph_format.space_after=Pt(60)
    c=ts.rows[1].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].add_run('').font.size=Pt(12); sp=c.add_paragraph(); sp.paragraph_format.space_after=Pt(60)
    c=ts.rows[2].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run(nv.get('ho_ten','').upper()); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[2].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run(CC['dai_dien'].upper()); r.bold=True; r.font.size=Pt(13)
    tf = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tf.name)
    return tf.name

def gui_email_don(to_email, subject, html_body):
    """Gửi 1 email đơn (dùng cho OTP reset mật khẩu, thông báo cá nhân...) - dùng chung EMAIL_CONFIG."""
    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['email']
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(html_body, 'html', 'utf-8'))
        srv = smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
        srv.starttls()
        srv.login(EMAIL_CONFIG['email'], EMAIL_CONFIG['password'])
        srv.send_message(msg)
        srv.quit()
        return True
    except Exception as e:
        print(f"Lỗi gửi email OTP: {e}")
        return False

def gui_email(loai, ds, file=None):
    # Không import trong hàm nữa, dùng EMAIL_CONFIG đã có sẵn
    # from config import EMAIL_CONFIG as EC  <-- XÓA DÒNG NÀY
    
    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['email']
        msg['To'] = EMAIL_CONFIG['nguoi_nhan']
        tn = f"{datetime.now().month:02d}/{datetime.now().year}"
        msg['Subject'] = f"[HRM-Port] Báo cáo {loai} lao động tháng {tn}"
        nd = f"<h3>BÁO CÁO {loai.upper()} LĐ</h3><p>Tháng: <b>{tn}</b></p><p>SL: <b>{len(ds)}</b></p><hr><ul>"
        for nv in ds[:10]:
            nd += f"<li>{nv.get('ho_ten','')} - {nv.get('chuc_danh_nghe','')}</li>"
        if len(ds) > 10:
            nd += f"<li>... và {len(ds)-10} người khác</li>"
        nd += "</ul><p><b>File Excel đính kèm.</b></p>"
        msg.attach(MIMEText(nd, 'html', 'utf-8'))
        if file and os.path.exists(file):
            with open(file, 'rb') as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(file)}"')
                msg.attach(part)
        srv = smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
        srv.starttls()
        srv.login(EMAIL_CONFIG['email'], EMAIL_CONFIG['password'])
        srv.send_message(msg)
        srv.quit()
        return True
    except Exception as e:
        st.error(f"Lỗi email: {e}")
        return False

def gui_telegram(msg):
    # from config import TELEGRAM_CONFIG as TC  <-- XÓA DÒNG NÀY
    
    try:
        url = f"https://api.telegram.org/bot{TELEGRAM_CONFIG['bot_token']}/sendMessage"
        r = requests.post(url, data={"chat_id": TELEGRAM_CONFIG['chat_id'], "text": msg, "parse_mode": "HTML"}, timeout=10)
        return r.status_code == 200
    except:
        return False

def tao_bao_cao_tang_giam(tang_list, giam_list, tu_ngay, den_ngay):
    """Tạo báo cáo Word tăng/giảm nhân sự"""
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    
    # Tiêu đề
    title = doc.add_heading('BÁO CÁO TĂNG/GIẢM NHÂN SỰ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    period = doc.add_paragraph(f'Thời gian: Từ {tu_ngay.strftime("%d/%m/%Y")} đến {den_ngay.strftime("%d/%m/%Y")}')
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Danh sách tăng
    doc.add_heading('I. LAO ĐỘNG TĂNG MỚI', level=1)
    if tang_list:
        table = doc.add_table(rows=1 + len(tang_list), cols=5)
        table.style = 'Table Grid'
        # Header
        headers = ['STT', 'Họ tên', 'Chức danh', 'Loại HĐ', 'Ngày vào làm']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for idx, nv in enumerate(tang_list, 1):
            row = table.rows[idx]
            row.cells[0].text = str(idx)
            row.cells[1].text = nv.get('ho_ten', '')
            row.cells[2].text = nv.get('chuc_danh_nghe', '')
            row.cells[3].text = nv.get('loai_hop_dong', '')
            row.cells[4].text = format_date(nv.get('ngay_vao_lam'))
    else:
        doc.add_paragraph('Không có lao động tăng trong kỳ.')
    
    doc.add_paragraph()
    
    # Danh sách giảm
    doc.add_heading('II. LAO ĐỘNG GIẢM (NGHỈ VIỆC)', level=1)
    if giam_list:
        table = doc.add_table(rows=1 + len(giam_list), cols=5)
        table.style = 'Table Grid'
        headers = ['STT', 'Họ tên', 'Chức danh', 'Loại HĐ', 'Ngày nghỉ việc']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for idx, nv in enumerate(giam_list, 1):
            row = table.rows[idx]
            row.cells[0].text = str(idx)
            row.cells[1].text = nv.get('ho_ten', '')
            row.cells[2].text = nv.get('chuc_danh_nghe', '')
            row.cells[3].text = nv.get('loai_hop_dong', '')
            row.cells[4].text = format_date(nv.get('ngay_ket_thuc'))
    else:
        doc.add_paragraph('Không có lao động giảm trong kỳ.')
    
    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'Ngày {date.today().day} tháng {date.today().month} năm {date.today().year}')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('NGƯỜI LẬP BÁO CÁO')
    
    tf = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tf.name)
    return tf.name
    
def show_super_admin_page():
    """Trang QUẢN TRỊ HỆ THỐNG — chỉ đội vận hành App dùng để thêm/sửa/khoá khách hàng (tenant).
    Hoàn toàn tách biệt với dữ liệu nhân sự của từng khách hàng."""
    st.title("⚙️ Quản trị hệ thống — Danh sách khách hàng (Tenants)")
    if st.button("🚪 Thoát trang quản trị"):
        st.session_state.super_admin_mode = False
        st.rerun()
    st.divider()

    # Hộp hướng dẫn bước THỦ CÔNG còn lại sau khi thêm 1 tenant mới (tạo Streamlit app riêng).
    # Lưu vào session_state để hiển thị BỀN sau st.rerun() (nếu không sẽ mất ngay lập tức).
    _vua_tao = st.session_state.get('_tenant_vua_tao')
    if _vua_tao:
        st.success(f"✅ Đã thêm khách hàng **{_vua_tao['ten_cty']}** (mã: **{_vua_tao['ma_cty']}**) "
                   f"và tự động chạy migration `schema.sql` thành công.")
        st.cache_data.clear()

        if _vua_tao.get('loi_tao_admin'):
            st.error(
                f"⚠️ Migration thành công, NHƯNG tự động tạo tài khoản Admin đầu tiên bị lỗi: "
                f"`{_vua_tao['loi_tao_admin']}`.\n\n"
                f"Cần tạo thủ công 1 dòng trong bảng `nhan_vien` của DB khách hàng này (SĐT = "
                f"`{_vua_tao.get('admin_dien_thoai','')}`, `vai_tro='admin'`) trước khi khách hàng "
                f"đăng nhập lần đầu."
            )
        else:
            st.info(
                f"✅ Đã tự động tạo tài khoản Admin đầu tiên: **{_vua_tao.get('admin_ho_ten','')}** "
                f"— đăng nhập bằng SĐT **{_vua_tao.get('admin_dien_thoai','')}**, mật khẩu lần đầu = "
                f"chính SĐT này (hệ thống sẽ buộc đổi mật khẩu ngay sau khi đăng nhập lần đầu)."
            )

        with st.container(border=True):
            st.markdown("### 📌 Bước tiếp theo (BẮT BUỘC — thực hiện thủ công trên Streamlit Cloud)")
            st.markdown(f"""
Mỗi khách hàng cần **1 app Streamlit Cloud riêng** để vào thẳng màn hình đăng nhập
(không cần chọn công ty). Thực hiện theo đúng thứ tự:

1. Vào [share.streamlit.io](https://share.streamlit.io) → **"New app"**
2. Chọn đúng repo GitHub hiện tại (repo chứa `app.py` này), nhánh `main`, file chính `app.py`
3. Đặt tên app theo **đúng quy chuẩn**: `hrm-{_vua_tao['ma_cty'].lower()}`
   → URL sẽ là: `https://hrm-{_vua_tao['ma_cty'].lower()}.streamlit.app`
4. Trước khi bấm Deploy, vào **"Advanced settings" → "Secrets"**, dán y hệt nội dung Secrets
   của app hiện tại, rồi **thêm thêm 1 dòng mới** vào cuối:
   ```
   tenant_code = "{_vua_tao['ma_so_thue']}"
   ```
   (Đây là **Mã số thuế** — khoá quản lý tenant hiện tại — KHÔNG còn dùng mã công ty nữa.
   Dòng này giúp app tự nhận diện đúng công ty, khách vào thẳng màn hình đăng nhập,
   không cần gõ mã công ty.)
5. Bấm **Deploy** và gửi link `https://hrm-{_vua_tao['ma_cty'].lower()}.streamlit.app` cho khách hàng,
   kèm SĐT đăng nhập của Admin đầu tiên đã tạo sẵn ở trên.
6. Sau khi Admin đầu tiên đăng nhập lần đầu, họ tự vào menu "🔑 Quản lý MK" → tab
   "🛡️ Phân quyền hệ thống" để chỉ định thêm HR / Văn thư / Kế toán lương cho các nhân viên khác.
""")
            if st.button("✅ Đã tạo app xong, đóng thông báo này"):
                del st.session_state['_tenant_vua_tao']
                st.rerun()

    with st.expander("➕ Thêm khách hàng mới (SaaS)", expanded=False):
        with st.form("add_tenant_form"):
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("##### 🏢 Thông tin Kết nối & Hệ thống")
                ma_cty = st.text_input("Mã công ty * (VD: CHL)")
                ten_cty = st.text_input("Tên công ty *")
                logo_url = st.text_input("Link logo (tuỳ chọn)")
                db_host = st.text_input("Supabase DB Host *")
                db_port = st.text_input("Supabase DB Port", value="5432")
                db_user = st.text_input("Supabase DB User", value="postgres")
                db_password = st.text_input("Supabase DB Password *", type="password")
                db_name = st.text_input("Supabase DB Name", value="postgres")
                supabase_url = st.text_input("Supabase Project URL *")
                supabase_key = st.text_input("Supabase API Key *", type="password")
            with col2:
                st.markdown("##### 🎨 Cấu hình Thương hiệu & Metadata")
                dai_dien = st.text_input("Người đại diện (Ký hợp đồng)", placeholder="VD: Nguyễn Đình Thi")
                chuc_vu = st.text_input("Chức vụ người ký", placeholder="VD: Tổng Giám Đốc")
                ma_so_thue = st.text_input("Mã số thuế *", help="Khoá dùng để quản lý tenant (khoá/mở, đổi ngôn ngữ, xoá, upload logo...) — bắt buộc, duy nhất theo pháp luật.")
                dien_thoai_cty = st.text_input("Điện thoại công ty")
                ma_don_vi_BHXH = st.text_input("Mã đơn vị BHXH")
                ma_vung_luong = st.text_input("Mã vùng lương")
                dia_chi = st.text_input("Địa chỉ công ty")
                loi_nhan_zalo = st.text_input("Lời nhắn Zalo sinh nhật")
                zalo_group_link = st.text_input("Link nhóm Zalo")
                zalo_group_name = st.text_input("Tên nhóm Zalo")
                st.markdown("###### 🌐 Ngôn ngữ giao diện")
                ngon_ngu_moi = st.selectbox(
                    "Ngôn ngữ",
                    list(i18n.LANGUAGE_OPTIONS.keys()),
                    format_func=lambda k: i18n.LANGUAGE_OPTIONS[k],
                    index=0,
                    help="Tiếng Việt luôn là ngôn ngữ chính. Nếu chọn gói song ngữ (Việt-Anh/"
                         "Việt-Trung/Việt-Hàn), ngôn ngữ phụ sẽ hiện thêm trong ngoặc, cỡ chữ nhỏ hơn, "
                         "ở menu và một số nhãn giao diện — phù hợp cho doanh nghiệp FDI."
                )

            st.divider()
            st.markdown("##### 🔑 Tài khoản Admin đầu tiên (BẮT BUỘC)")
            st.caption("Toàn bộ đăng nhập giờ đây đều phải là 1 nhân viên thật trong bảng Nhân viên — "
                       "không còn tài khoản hệ thống rời rạc nữa. Điền thông tin người sẽ là Admin đầu "
                       "tiên của công ty này; hệ thống tự tạo hồ sơ nhân viên + cấp quyền Admin cho họ "
                       "ngay sau khi migration xong. Mật khẩu đăng nhập lần đầu = chính SĐT bên dưới.")
            c_admin1, c_admin2 = st.columns(2)
            with c_admin1:
                admin_ho_ten = st.text_input("Họ tên Admin đầu tiên *", placeholder="VD: Nguyễn Văn A")
            with c_admin2:
                admin_dien_thoai = st.text_input("Số điện thoại Admin đầu tiên *", placeholder="VD: 0912345678")

            if st.form_submit_button("💾 Lưu khách hàng & Tự động chạy Migration"):
                if not all([ma_cty, ten_cty, ma_so_thue, db_host, db_password, supabase_url, supabase_key,
                            admin_ho_ten, admin_dien_thoai]):
                    st.error("❌ Vui lòng điền đầy đủ các trường bắt buộc (*), kể cả Admin đầu tiên. "
                             "Lưu ý: **Mã số thuế** nay là bắt buộc — đây là khoá dùng để quản lý "
                             "(khoá/mở, đổi ngôn ngữ, xoá, upload logo...) tenant về sau.")
                else:
                    try:
                        # Đọc file schema.sql từ thư mục hiện tại
                        migration_sql = None
                        import os
                        schema_path = os.path.join(os.path.dirname(__file__), "schema.sql")
                        if os.path.exists(schema_path):
                            with open(schema_path, "r", encoding="utf-8") as sf:
                                migration_sql = sf.read()
                        
                        control_plane.add_tenant(
                            ma_cty=ma_cty, ten_cty=ten_cty, db_host=db_host, db_port=db_port,
                            db_user=db_user, db_password=db_password, db_name=db_name,
                            supabase_url=supabase_url, supabase_key=supabase_key, logo_url=logo_url,
                            dai_dien=dai_dien, chuc_vu=chuc_vu, ma_so_thue=ma_so_thue,
                            dien_thoai_cty=dien_thoai_cty, ma_don_vi_BHXH=ma_don_vi_BHXH,
                            ma_vung_luong=ma_vung_luong, dia_chi=dia_chi, loi_nhan_zalo=loi_nhan_zalo,
                            zalo_group_link=zalo_group_link, zalo_group_name=zalo_group_name,
                            migration_sql=migration_sql, ngon_ngu=ngon_ngu_moi
                        )

                        # Tự động tạo hồ sơ nhân viên Admin đầu tiên trên DB của tenant VỪA tạo
                        # (thay cho bước thủ công "tự chạy SQL" trước đây) — kết nối trực tiếp bằng
                        # thông tin DB vừa nhập ở trên, KHÔNG dùng st.session_state.db_engine vì đó
                        # vẫn đang trỏ tới tenant hiện tại (chưa chuyển sang tenant mới này).
                        loi_tao_admin = None
                        try:
                            conn_moi = psycopg2.connect(
                                host=db_host, port=db_port, user=db_user,
                                password=db_password, dbname=db_name, sslmode="require",
                            )
                            cur_moi = conn_moi.cursor()
                            # Phòng khi schema.sql của tenant chưa có các cột này (bản cũ)
                            cur_moi.execute("ALTER TABLE nhan_vien ADD COLUMN IF NOT EXISTS vai_tro VARCHAR(20) DEFAULT 'nhan_vien'")
                            cur_moi.execute("ALTER TABLE nhan_vien ADD COLUMN IF NOT EXISTS phai_doi_mat_khau BOOLEAN DEFAULT FALSE")
                            cur_moi.execute("""
                                INSERT INTO nhan_vien (ma_nv, ho_ten, dien_thoai, vai_tro, chuc_vu,
                                                        phong_ban_lam_viec, trang_thai, phai_doi_mat_khau)
                                VALUES ('ADMIN01', %s, %s, 'admin', 'Tổng Giám Đốc',
                                        'Ban Tổng Giám Đốc', 'DANG_LAM', TRUE)
                            """, (admin_ho_ten.strip(), admin_dien_thoai.strip()))
                            conn_moi.commit()
                            conn_moi.close()
                        except Exception as e_admin:
                            loi_tao_admin = str(e_admin)

                        st.session_state['_tenant_vua_tao'] = {
                            'ma_cty': ma_cty.strip().upper(), 'ten_cty': ten_cty,
                            'ma_so_thue': ma_so_thue.strip(),
                            'admin_ho_ten': admin_ho_ten.strip(), 'admin_dien_thoai': admin_dien_thoai.strip(),
                            'loi_tao_admin': loi_tao_admin,
                        }
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Lỗi khi thêm khách hàng hoặc chạy migration: {e}")


    st.subheader("📋 Danh sách khách hàng hiện có")
    try:
        tenants = control_plane.list_tenants()
    except Exception as e:
        tenants = []
        st.error(f"❌ Không kết nối được Control Plane. Kiểm tra lại st.secrets['control_plane']. Chi tiết: {e}")

    if tenants:
        df = pd.DataFrame(tenants)
        st.dataframe(df, width='stretch', hide_index=True)
        st.divider()

        # Danh sách MST + hàm hiển thị "Tên công ty (MST: xxx)" dùng chung cho MỌI dropdown
        # chọn khách hàng bên dưới -> tránh gõ tay sai mã số thuế.
        _mst_options = [t['ma_so_thue'] for t in tenants if t.get('ma_so_thue')]
        _mst_format = lambda mst: next(
            (f"{t['ten_cty']} (MST: {t['ma_so_thue']})" for t in tenants if t['ma_so_thue'] == mst), mst
        )

        col_a, col_b = st.columns(2)
        with col_a:
            mst_toggle = st.selectbox("Mã số thuế cần Khoá/Mở khoá", _mst_options,
                                       format_func=_mst_format, key="mst_toggle_select")
            trang_thai_moi = st.selectbox("Trạng thái mới", ["active", "suspended"])
            if st.button("🔄 Cập nhật trạng thái"):
                if mst_toggle:
                    control_plane.update_tenant_status(mst_toggle, trang_thai_moi)
                    st.success("✅ Đã cập nhật!"); st.cache_data.clear(); st.rerun()
        with col_b:
            mst_xoa = st.selectbox("Mã số thuế cần XOÁ vĩnh viễn khỏi hệ thống", _mst_options,
                                    format_func=_mst_format, key="mst_xoa_select")
            if st.button("🗑️ Xoá khách hàng", type="primary"):
                if mst_xoa:
                    control_plane.delete_tenant(mst_xoa)
                    st.success("✅ Đã xoá!"); st.cache_data.clear(); st.rerun()

        st.divider()
        st.markdown("##### 🌐 Đổi ngôn ngữ giao diện của 1 khách hàng đã có")
        col_lang1, col_lang2 = st.columns(2)
        with col_lang1:
            mst_doi_ngon_ngu = st.selectbox("Mã số thuế cần đổi ngôn ngữ", _mst_options,
                                             format_func=_mst_format, key="ma_doi_ngon_ngu")
        with col_lang2:
            ngon_ngu_doi_thanh = st.selectbox(
                "Ngôn ngữ mới", list(i18n.LANGUAGE_OPTIONS.keys()),
                format_func=lambda k: i18n.LANGUAGE_OPTIONS[k], key="ngon_ngu_doi_thanh"
            )
        if st.button("🌐 Cập nhật ngôn ngữ"):
            if mst_doi_ngon_ngu:
                try:
                    control_plane.update_tenant_language(mst_doi_ngon_ngu, ngon_ngu_doi_thanh)
                    st.success("✅ Đã cập nhật ngôn ngữ!"); st.cache_data.clear(); st.rerun()
                except AttributeError:
                    st.error("❌ Chưa có hàm `update_tenant_language()` trong control_plane.py.")
            else:
                st.warning("⚠️ Vui lòng chọn mã số thuế.")

        st.divider()
        st.markdown("##### 🖼️ Upload logo cho khách hàng")
        col_logo1, col_logo2 = st.columns([1, 2])
        with col_logo1:
            mst_logo = st.selectbox("Mã số thuế", _mst_options,
                                     format_func=_mst_format, key="mst_upload_logo")
        with col_logo2:
            logo_file = st.file_uploader("Chọn file logo (PNG/JPG)", type=["png", "jpg", "jpeg"], key="logo_file_uploader")
        if st.button("📤 Upload logo", key="btn_upload_logo"):
            if not mst_logo or not logo_file:
                st.warning("⚠️ Vui lòng nhập Mã số thuế và chọn file logo.")
            else:
                sb = get_supabase_storage()
                if not sb:
                    st.error("❌ Chưa cấu hình Supabase Storage.")
                else:
                    try:
                        # Dùng ma_cty (nếu tra được) làm tên thư mục cho dễ đọc trên Storage;
                        # nếu không tra được (VD tenant chưa tồn tại) thì dùng luôn mã số thuế.
                        tenant_logo = control_plane.get_tenant_by_ma_so_thue(mst_logo.strip())
                        ten_folder_logo = (tenant_logo or {}).get('ma_cty') or mst_logo
                        safe_name = sanitize_storage_filename(logo_file.name)
                        storage_path = f"{sanitize_storage_filename(ten_folder_logo)}/{safe_name}"
                        upload_to_storage_unique(
                            sb, SUPABASE_BUCKET_LOGO, storage_path,
                            logo_file.getvalue(), logo_file.type
                        )
                        # Lấy public URL để lưu vào tenant.logo_url. Dùng bucket RIÊNG
                        # (SUPABASE_BUCKET_LOGO), luôn ở chế độ Public, vì logo hiển thị cả ở
                        # màn hình đăng nhập — trước khi xác thực người dùng.
                        public_url = sb.storage.from_(SUPABASE_BUCKET_LOGO).get_public_url(storage_path)
                        control_plane.update_tenant_logo(mst_logo.strip(), public_url)
                        st.success(f"✅ Đã upload logo và cập nhật cho khách hàng MST {mst_logo.strip()}. Link: {public_url}")
                        st.image(public_url, width=160)
                        st.cache_data.clear()
                        st.rerun()
                    except AttributeError:
                        st.error("❌ Chưa có hàm `update_tenant_logo()` trong control_plane.py. "
                                 "Cần thêm hàm này (UPDATE tenants SET logo_url=%s WHERE ma_so_thue=%s) để nút này hoạt động.")
                    except Exception as e:
                        loi_bucket_logo = st.session_state.get(f"_sb_bucket_error_{SUPABASE_BUCKET_LOGO}")
                        if loi_bucket_logo and 'bucket' in str(e).lower():
                            st.error(
                                f"❌ Lỗi upload logo: {e}\n\n"
                                f"Nguyên nhân: không tự tạo được bucket `{SUPABASE_BUCKET_LOGO}` trên project Supabase "
                                f"đang cấu hình ({loi_bucket_logo}). Tạo bucket thủ công trên Supabase Dashboard (nhớ để "
                                f"chế độ **Public**) hoặc đổi sang service_role key."
                            )
                        else:
                            st.error(f"❌ Lỗi upload logo: {e}")

        st.divider()
        st.markdown("##### 📥 Nhập/Xuất dữ liệu Excel cho 1 khách hàng")
        mst_excel = st.selectbox(
            "Chọn công ty cần thao tác",
            [t['ma_so_thue'] for t in tenants if t.get('ma_so_thue')],
            format_func=lambda mst: next((f"{t['ten_cty']} (MST: {t['ma_so_thue']})" for t in tenants if t['ma_so_thue'] == mst), mst),
            key="mst_excel_tool",
        )
        if mst_excel:
            with st.expander(f"📥 Công cụ Nhập/Xuất Excel — {mst_excel}", expanded=True):
                try:
                    tenant_full = control_plane.get_tenant_by_ma_so_thue(mst_excel)
                    render_import_export_ui(
                        lambda t=tenant_full: control_plane.DatabaseEngine(t).get_connection(),
                        extra_caption=f"⚠️ Đang thao tác hộ dữ liệu của: {tenant_full.get('ten_cty', mst_excel)}"
                    )
                except Exception as e:
                    st.error(f"❌ Không kết nối được tới DB của {mst_excel}: {e}")          
    else:
        st.info("Chưa có khách hàng nào. Thêm khách hàng đầu tiên ở form phía trên.")


# ========== SIDEBAR + LOGIN (ĐA KHÁCH HÀNG) ==========
if not st.session_state.get('tenant'):
    st.sidebar.title("🏗️ HRM Master")
    st.sidebar.caption("Nền tảng Quản lý Hồ sơ nhân sự đa doanh nghiệp")


def check_login(username, password):
    """Xác thực đăng nhập của NHÂN VIÊN thuộc tenant (công ty) đã chọn.
    Tài khoản = số điện thoại (dien_thoai), mật khẩu hash bằng bcrypt trong cột mat_khau_hash.
    Trả về (success, role, nhan_vien_row) — nhan_vien_row là dict thông tin NV nếu thành công.

    QUAN TRỌNG (chuẩn hóa): TẤT CẢ tài khoản — kể cả admin/hr/văn thư/kt_luong — ĐỀU PHẢI là
    1 dòng thật trong bảng nhan_vien, đăng nhập bằng SĐT của chính người đó. KHÔNG còn "tài
    khoản hệ thống" khai báo riêng trong Secrets [users] nữa (đã bỏ), để mọi tính năng cần biết
    "đây là nhân viên nào" (Chat nội bộ, ký duyệt, lịch sử thao tác...) luôn có đủ thông tin.
    Vai trò (vai_tro) của từng nhân viên được chính admin công ty tự cấu hình tại
    menu "🔑 Quản lý MK" → tab "🛡️ Phân quyền hệ thống"."""
    tenant = st.session_state.get('tenant')

    if tenant:
        try:
            db = st.session_state.db_engine.get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            # Tài khoản đăng nhập = ten_dang_nhap nếu admin đã set riêng, mặc định = số điện thoại.
            # COALESCE(ten_dang_nhap, dien_thoai): nếu ten_dang_nhap NULL/rỗng thì so sánh theo SĐT.
            c.execute("""SELECT id, ho_ten, dien_thoai, ten_dang_nhap, mat_khau_hash, vai_tro, phai_doi_mat_khau
                         FROM nhan_vien
                         WHERE NULLIF(TRIM(ten_dang_nhap), '') = %s
                            OR (NULLIF(TRIM(ten_dang_nhap), '') IS NULL AND dien_thoai = %s)""",
                      (username.strip(), username.strip()))
            rows = c.fetchall()
            db.close()
            row = rows[0] if rows else None
            if not row:
                return False, None, None
            if not row.get('mat_khau_hash'):
                khop = password.strip() == (row.get('dien_thoai') or '').strip()
                if khop:
                    row['phai_doi_mat_khau'] = True
                    return True, row.get('vai_tro') or 'nhan_vien', row
                return False, None, None
            if bcrypt.checkpw(password.encode(), row['mat_khau_hash'].encode()):
                return True, row.get('vai_tro') or 'nhan_vien', row
        except Exception as e:
            # DEBUG TẠM THỜI: lưu lại lỗi thật để hiện ra màn hình, giúp xác định
            # chính xác nguyên nhân thay vì chỉ thấy "Sai tài khoản hoặc mật khẩu".
            # Sau khi xác định xong nguyên nhân, nên gỡ bỏ 2 dòng debug này.
            st.session_state['_debug_login_error'] = f"{type(e).__name__}: {e}"
        return False, None, None


    # ---- Chế độ KHÔNG có tenant (chạy đơn lẻ / dev local) — giữ cách cũ để không phá vỡ ----
    try:
        if 'users' in st.secrets and username in st.secrets.users:
            if st.secrets.users[username]['password'] == password:
                return True, st.secrets.users[username]['role'], None
    except Exception:
        pass
    try:
        if username in USERS:
            return USERS[username]['password'] == password, USERS[username]['role'], None
    except Exception:
        pass
    return False, None, None




def _load_tenant_module(folder, prefix, ma_so_thue):
    """Import động module Python riêng của 1 tenant, dùng chung cho cơ chế
    "mỗi tenant 1 file riêng" (salary/salary_{ma_so_thue}.py,
    landing_page/landing_{ma_so_thue}.py).

    QUYẾT ĐỊNH THIẾT KẾ: khoá dùng để chọn file là 'ma_so_thue' (khoá quản lý tenant
    hiện tại — xem control_plane.py), KHÔNG dùng 'ma_cty' vì mã công ty là mã ngắn tự
    đặt, có thể trùng giữa 2 khách hàng khi SaaS có nhiều khách; mã số thuế thì duy
    nhất theo pháp luật.

    Thử import '{folder}.{prefix}_{ma_so_thue}' trước; nếu module không tồn tại (tenant
    chưa có file riêng) hoặc lỗi import, trả về None để nơi gọi tự rơi về file mặc định
    '{folder}.{prefix}_demo'."""
    import importlib
    if not ma_so_thue:
        return None
    ten_module = f"{prefix}_{ma_so_thue}"
    try:
        return importlib.import_module(f"{folder}.{ten_module}")
    except ModuleNotFoundError:
        return None
    except Exception as e:
        print(f"Lỗi import {folder}.{ten_module}: {e}")
        return None


def _load_tenant_module_or_demo(folder, prefix, ma_so_thue):
    """Trả về module riêng của tenant (nếu có file '{prefix}_{ma_so_thue}.py'),
    hoặc rơi về module mặc định '{prefix}_demo.py' nếu tenant chưa có file riêng."""
    import importlib
    module = _load_tenant_module(folder, prefix, ma_so_thue)
    if module is not None:
        return module
    try:
        return importlib.import_module(f"{folder}.{prefix}_demo")
    except Exception as e:
        print(f"Lỗi import {folder}.{prefix}_demo (file mặc định): {e}")
        return None


def render_landing_page():
    """Hiển thị Landing Page giới thiệu, ở vùng nội dung chính (bên phải sidebar đăng
    nhập) khi CHƯA đăng nhập.

    QUYẾT ĐỊNH: mỗi tenant nay có thể có 1 Landing Page RIÊNG — file
    landing_page/landing_{ma_so_thue}.py (đặt cùng cấp với app.py), tự động được chọn
    theo đúng khách hàng đang đăng nhập (dò theo ma_so_thue — khoá quản lý tenant hiện
    tại). Tenant nào CHƯA có file riêng (chưa tuỳ biến) sẽ tự dùng
    landing_page/landing_demo.py làm mặc định (nội dung y hệt bản gốc trước khi có cơ
    chế đa-tenant)."""
    tenant = st.session_state.get('tenant') or {}
    ma_so_thue = (tenant.get('ma_so_thue') or '').strip()
    module = _load_tenant_module_or_demo('landing_page', 'landing', ma_so_thue)
    if module is not None and hasattr(module, 'render'):
        module.render()
    else:
        st.error("❌ Không tải được Landing Page (thiếu landing_page/landing_demo.py cùng cấp với app.py).")

if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.role = None
    st.session_state.username = None

if not st.session_state.logged_in:

    # ---------- Trang quản trị hệ thống (super-admin quản lý danh sách khách hàng) ----------
    if st.session_state.get('super_admin_mode'):
        show_super_admin_page()
        st.stop()

    # ---------- App bị khoá cứng vào 1 tenant (Secret tenant_code) nhưng mã sai/bị khoá ----------
    if st.session_state.get('_tenant_locked_error'):
        st.error("⚠️ App này được cấu hình riêng cho 1 khách hàng, nhưng không tìm thấy "
                 "hoặc tài khoản khách hàng đang bị tạm khoá. Vui lòng liên hệ đơn vị triển khai App.")
        st.stop()

    # ---------- BƯỚC 1: Không xác định được tenant qua domain/subdomain ----------
    # Mỗi khách hàng nay có domain/subdomain riêng, được resolve_tenant() tự nhận diện.
    # Nếu chạy đến đây nghĩa là đang ở domain gốc (chưa gán cho khách nào) hoặc
    # subdomain không khớp tenant nào — KHÔNG còn cho phép tự nhập mã công ty nữa.
    # Domain gốc chỉ còn dùng làm cổng vào cho Super Admin (đội vận hành App).
    if not st.session_state.get('tenant'):
        st.title("🏗️ HRM-Port")
        st.info(
            "🔒 Tên miền này chưa được gán cho khách hàng nào. "
            "Nếu bạn là nhân viên của một công ty đang dùng HRM-Port, vui lòng truy cập "
            "đúng địa chỉ riêng của công ty bạn (ví dụ: `hangcuaban.kendu-ai.com`)."
        )
        with st.expander("⚙️ Quản trị hệ thống (chỉ dành cho đội vận hành App)"):
            sa_u = st.text_input("Tài khoản", key="sa_user")
            sa_p = st.text_input("Mật khẩu", type="password", key="sa_pass")
            if st.button("Đăng nhập quản trị", key="sa_login"):
                if control_plane.check_super_admin(sa_u, sa_p):
                    st.session_state.super_admin_mode = True
                    st.rerun()
                else:
                    st.error("❌ Sai tài khoản/mật khẩu quản trị hệ thống!")
        st.stop()

    # ---------- BƯỚC 2: Đăng nhập nhân viên của công ty đã chọn ----------
    tenant = st.session_state.tenant
    if tenant.get('logo_url'):
        try:
            st.sidebar.markdown(f"""
            <div style="display: flex; justify-content: center; align-items: center; 
                        width: 100%; min-height: 200px; padding: 10px 0;">
                <img src="{tenant['logo_url']}" 
                     style="width: 150px; height: 150px; border-radius: 50%; 
                            object-fit: cover; box-shadow: 0 4px 20px rgba(0,0,0,0.25);
                            display: block;">
            </div>
            """, unsafe_allow_html=True)
        except Exception:
            # Link logo hỏng/không truy cập được — bỏ qua, không để sập cả trang đăng nhập.
            pass
    st.sidebar.success(f"🏢 **{tenant['ten_cty']}**")

    st.sidebar.subheader(i18n.t("🔐 Đăng nhập"))

    # Tenant DEMO-HRM: điền sẵn tài khoản trải nghiệm (cấu hình tại ⚙️ Danh mục >
    # "Tài khoản đăng nhập DEMO") để khách vào thử không cần hỏi tài khoản/mật khẩu.
    _is_demo_tenant = str(tenant.get('ma_cty', '')).upper() == 'DEMO'
    _demo_user_default = ''
    _demo_pass_default = ''
    if _is_demo_tenant:
        _demo_user_default = get_cau_hinh('demo_ten_dang_nhap', '') or ''
        _demo_pass_default = get_cau_hinh('demo_mat_khau', '') or ''

    u = st.sidebar.text_input(i18n.t("Số điện thoại hoặc Tên đăng nhập"), value=_demo_user_default)
    p = st.sidebar.text_input(i18n.t("Mật khẩu"), type="password", value=_demo_pass_default)
    if _is_demo_tenant and _demo_user_default:
        st.sidebar.info(f"🧪 Tài khoản dùng thử & Mật khẩu đã điền sẵn - bấm Đăng nhập luôn.")    ###: **{_demo_user_default}** / **{_demo_pass_default}**
    else:
        st.sidebar.caption("💡 Mật khẩu mặc định = số điện thoại của bạn. Đổi lại sau khi đăng nhập lần đầu.")
    if st.sidebar.button(i18n.t("Đăng nhập"), width='stretch'):
        success, role, nv_row = check_login(u, p)
        if success:
            st.session_state.logged_in = True
            st.session_state.role = role
            st.session_state.username = u
            st.session_state.nhan_vien_id = nv_row['id'] if nv_row else None
            st.session_state.ho_ten_dang_nhap = nv_row['ho_ten'] if nv_row else u
            st.session_state.phai_doi_mat_khau = bool(nv_row and nv_row.get('phai_doi_mat_khau'))
            st.rerun()
        else:
            st.sidebar.error("❌ Sai tài khoản hoặc mật khẩu!")
            if st.session_state.get('_debug_login_error'):
                st.sidebar.caption(f"🔧 Debug: {st.session_state['_debug_login_error']}")

    with st.sidebar.expander("🔑 Quên mật khẩu?"):
        try:
            db_qmk = st.session_state.db_engine.get_connection()
            c_qmk = db_qmk.cursor()
            c_qmk.execute("""
                CREATE TABLE IF NOT EXISTS yeu_cau_reset_mk (
                    id SERIAL PRIMARY KEY,
                    nhan_vien_id INT NOT NULL,
                    otp_code VARCHAR(10) NOT NULL,
                    het_han TIMESTAMP NOT NULL,
                    da_dung BOOLEAN DEFAULT FALSE,
                    created_at TIMESTAMP DEFAULT NOW()
                )
            """)
            db_qmk.commit(); db_qmk.close()
        except Exception:
            pass

        buoc_qmk = st.session_state.get('qmk_buoc', 1)
        if buoc_qmk == 1:
            st.caption("Nhập SĐT hoặc Tên đăng nhập — mã xác nhận (OTP) sẽ gửi về Email liên hệ đã đăng ký.")
            tk_qmk = st.text_input("SĐT / Tên đăng nhập:", key="qmk_tk")
            if st.button("📧 Gửi mã OTP", key="qmk_gui_otp"):
                db_q = st.session_state.db_engine.get_connection()
                c_q = db_q.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                c_q.execute("""SELECT id, ho_ten, email_lien_he FROM nhan_vien
                               WHERE dien_thoai=%s OR ten_dang_nhap=%s""", (tk_qmk, tk_qmk))
                nv_qmk = c_q.fetchone()
                db_q.close()
                if not nv_qmk:
                    st.error("❌ Không tìm thấy tài khoản này.")
                elif not nv_qmk.get('email_lien_he'):
                    st.error("❌ Tài khoản chưa có Email liên hệ. Liên hệ Admin/HR để được hỗ trợ đặt lại mật khẩu.")
                else:
                    otp = f"{random.randint(0, 999999):06d}"
                    het_han = datetime.now() + timedelta(minutes=10)
                    db_o = st.session_state.db_engine.get_connection()
                    c_o = db_o.cursor()
                    c_o.execute("""INSERT INTO yeu_cau_reset_mk (nhan_vien_id, otp_code, het_han)
                                   VALUES (%s, %s, %s)""", (nv_qmk['id'], otp, het_han))
                    db_o.commit(); db_o.close()
                    da_gui = gui_email_don(
                        nv_qmk['email_lien_he'],
                        "🔑 Mã xác nhận đặt lại mật khẩu - HRM",
                        f"<p>Xin chào {nv_qmk['ho_ten']},</p><p>Mã xác nhận (OTP) của bạn là:</p>"
                        f"<h2 style='letter-spacing:4px;'>{otp}</h2>"
                        f"<p>Mã có hiệu lực trong 10 phút. Nếu không phải bạn yêu cầu, vui lòng bỏ qua email này.</p>"
                    )
                    if da_gui:
                        st.session_state['qmk_buoc'] = 2
                        st.session_state['qmk_nv_id'] = nv_qmk['id']
                        st.success(f"✅ Đã gửi mã OTP về {nv_qmk['email_lien_he'][:3]}***@...")
                        st.cache_data.clear()
                        st.rerun()
                    else:
                        st.error("❌ Gửi email thất bại. Vui lòng thử lại hoặc liên hệ Admin.")
        elif buoc_qmk == 2:
            st.caption("Nhập mã OTP đã gửi về Email và đặt mật khẩu mới.")
            otp_nhap = st.text_input("Mã OTP:", key="qmk_otp_nhap")
            mk_moi_qmk = st.text_input("Mật khẩu mới:", type="password", key="qmk_mk_moi")
            mk_moi_qmk2 = st.text_input("Nhập lại mật khẩu mới:", type="password", key="qmk_mk_moi2")
            col_qmk1, col_qmk2 = st.columns(2)
            with col_qmk1:
                if st.button("✅ Xác nhận đặt lại", key="qmk_xac_nhan"):
                    if len(mk_moi_qmk) < 6:
                        st.error("Mật khẩu mới phải có ít nhất 6 ký tự.")
                    elif mk_moi_qmk != mk_moi_qmk2:
                        st.error("Hai mật khẩu nhập lại không khớp.")
                    else:
                        db_v = st.session_state.db_engine.get_connection()
                        c_v = db_v.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                        c_v.execute("""
                            SELECT id FROM yeu_cau_reset_mk
                            WHERE nhan_vien_id=%s AND otp_code=%s AND da_dung=FALSE AND het_han > NOW()
                            ORDER BY id DESC LIMIT 1
                        """, (st.session_state['qmk_nv_id'], otp_nhap))
                        yc = c_v.fetchone()
                        if not yc:
                            db_v.close()
                            st.error("❌ Mã OTP không đúng hoặc đã hết hạn.")
                        else:
                            c_v2 = db_v.cursor()
                            new_hash = bcrypt.hashpw(mk_moi_qmk.encode(), bcrypt.gensalt()).decode()
                            c_v2.execute("UPDATE nhan_vien SET mat_khau_hash=%s, phai_doi_mat_khau=FALSE WHERE id=%s",
                                         (new_hash, st.session_state['qmk_nv_id']))
                            c_v2.execute("UPDATE yeu_cau_reset_mk SET da_dung=TRUE WHERE id=%s", (yc['id'],))
                            db_v.commit(); db_v.close()
                            st.success("✅ Đặt lại mật khẩu thành công! Vui lòng đăng nhập lại.")
                            st.cache_data.clear()
                            st.session_state.pop('qmk_buoc', None)
                            st.session_state.pop('qmk_nv_id', None)
                            st.rerun()
            with col_qmk2:
                if st.button("✖️ Hủy", key="qmk_huy"):
                    st.session_state.pop('qmk_buoc', None)
                    st.session_state.pop('qmk_nv_id', None)
                    st.rerun()

    # ---------- Landing page giới thiệu (hiển thị ở vùng nội dung chính, bên phải sidebar đăng nhập) ----------
    st.markdown("""
        <style>
            @media (min-width: 769px) {
                div[data-testid="stAlertContainer"]:has(div[data-testid="stMarkdownContainer"] p:-webkit-any(:contains("Đang dùng điện thoại"))) {
                    display: none !important;
                }
            }
        </style>
    """, unsafe_allow_html=True)
    st.markdown("""
        <style>
            @media (min-width: 769px) {
                .canh-bao-mobile-only { display: none !important; }
            }
        </style>
        <div class="canh-bao-mobile-only" style="background-color:#e7f3fe; border-radius:8px;
                    padding:0.9rem 1rem; margin-bottom:1rem; border:1px solid #b6d4f5; color:#0c5da8;">
            📱 <b>Đang dùng điện thoại?</b> Form đăng nhập nằm ở thanh bên trái.
            Nếu không thấy, bấm vào biểu tượng <b>›</b> (mũi tên) ở góc trên bên trái màn hình để mở ra.
        </div>
    """, unsafe_allow_html=True)
    render_landing_page()

    st.stop()

# ---------- Bắt buộc đổi mật khẩu lần đầu (đang dùng mật khẩu mặc định = SĐT) ----------
if st.session_state.get('phai_doi_mat_khau'):
    st.title("🔑 Đổi mật khẩu lần đầu")
    st.warning("Đây là lần đăng nhập đầu tiên (mật khẩu mặc định = số điện thoại của bạn). "
               "Vui lòng đặt mật khẩu mới trước khi tiếp tục sử dụng hệ thống.")
    mk_moi = st.text_input("Mật khẩu mới", type="password", key="mk_moi_lan_dau")
    mk_moi2 = st.text_input("Nhập lại mật khẩu mới", type="password", key="mk_moi_lan_dau_2")
    if st.button("✅ Xác nhận đổi mật khẩu"):
        if len(mk_moi) < 6:
            st.error("Mật khẩu mới phải có ít nhất 6 ký tự.")
        elif mk_moi != mk_moi2:
            st.error("Hai mật khẩu nhập lại không khớp.")
        else:
            try:
                db = st.session_state.db_engine.get_connection()
                c = db.cursor()
                new_hash = bcrypt.hashpw(mk_moi.encode(), bcrypt.gensalt()).decode()
                c.execute(
                    "UPDATE nhan_vien SET mat_khau_hash=%s, phai_doi_mat_khau=FALSE WHERE id=%s",
                    (new_hash, st.session_state.nhan_vien_id)
                )
                db.commit()
                db.close()
                st.session_state.phai_doi_mat_khau = False
                st.success("✅ Đổi mật khẩu thành công! Đang vào hệ thống...")
                st.cache_data.clear()
                st.rerun()
            except Exception as e:
                st.error(f"❌ Lỗi khi đổi mật khẩu: {e}")
    st.stop()

# Điểm 2 — Lazy check: quét NV thiếu giờ ra & gửi cảnh báo qua Chat nội bộ,
# chạy 1 lần/phiên (tránh query lặp lại mỗi lần rerun trong cùng phiên làm việc).
if not st.session_state.get('_da_quet_canh_bao_thieu_gio_ra'):
    try:
        _db_quet = st.session_state.db_engine.get_connection()
        face_id_cham_cong.quet_va_canh_bao_thieu_gio_ra(_db_quet, date.today())
        _db_quet.close()
    except Exception:
        pass  # không để lỗi quét cảnh báo làm sập cả trang
    st.session_state['_da_quet_canh_bao_thieu_gio_ra'] = True

# Menu theo role — 4 vai trò cố định: admin / hr / kt_luong / viewer (+ 'nhan_vien' tự phục vụ)
if st.session_state.role == "admin":
    # Toàn quyền
    menu_options = ["📊 Dashboard","👤 Ứng viên","✅ Nhân viên","📁 Upload hồ sơ","⚙️ Danh mục","📥 Nhập/Xuất Excel","📋 BHXH","📋 Báo cáo định kỳ","🕒 Chấm công","💰 Tính thu nhập","📄 Quản lý Công văn & HĐ kinh tế","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🔑 Quản lý MK","🖼️ Tạo ảnh thẻ NV","🔍 Audit Dashboard","📘 Hướng dẫn sử dụng",]
elif st.session_state.role in ["văn thư", "hr"]:
    # HR: như admin trừ Upload hồ sơ, Danh mục — và KHÔNG được xem Tính thu nhập (dữ liệu lương)
    menu_options = ["📊 Dashboard","✅ Nhân viên","📋 BHXH","📋 Báo cáo định kỳ","🕒 Chấm công","📄 Quản lý Công văn & HĐ kinh tế","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🔑 Quản lý MK","🖼️ Tạo ảnh thẻ NV","📘 Hướng dẫn sử dụng",]
elif st.session_state.role == "kt_luong":
    # Kế toán lương: tập trung vào Chấm công + Tính thu nhập, không có Upload hồ sơ/Danh mục
    menu_options = ["📊 Dashboard","✅ Nhân viên","📋 BHXH","🕒 Chấm công","💰 Tính thu nhập","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🔑 Quản lý MK","🖼️ Tạo ảnh thẻ NV","📘 Hướng dẫn sử dụng",]
elif st.session_state.role == "van_thu":
    menu_options = ["📊 Dashboard","✅ Nhân viên","🕒 Chấm công","📄 Quản lý Công văn & HĐ kinh tế","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🔑 Quản lý MK","🖼️ Tạo ảnh thẻ NV","📘 Hướng dẫn sử dụng",]
elif st.session_state.role == "admin_bcc":
    # Admin BCC: Chấm công + BHXH + Tính thu nhập + Dashboard (không có Ứng viên, Danh mục, Công văn)
    menu_options = ["📊 Dashboard","✅ Nhân viên","📋 BHXH","🕒 Chấm công","💰 Tính thu nhập","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🔑 Quản lý MK","🖼️ Tạo ảnh thẻ NV","📘 Hướng dẫn sử dụng",]
elif st.session_state.role == "viewer":
    # Viewer: chỉ xem, thu hẹp — không có BHXH, không có Tính thu nhập
    menu_options = ["📊 Dashboard","✅ Nhân viên","📋 Báo cáo định kỳ","🕒 Chấm công","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🔑 Quản lý MK","🖼️ Tạo ảnh thẻ NV","📘 Hướng dẫn sử dụng",]
elif st.session_state.role == "xem_toan_bo":
    # Vai trò "Xem toàn bộ (không chỉnh sửa)": thấy ĐẦY ĐỦ menu & tab giống hệt Admin,
    # nhưng KHÔNG có quyền thay đổi dữ liệu — mọi nút Lưu/Sửa/Xóa/Cập nhật/Save trong các
    # màn hình bên dưới đều bị làm mờ (disabled=not can_edit()/can_delete()).
    menu_options = ["📊 Dashboard","👤 Ứng viên","✅ Nhân viên","📁 Upload hồ sơ","⚙️ Danh mục","📥 Nhập/Xuất Excel","📋 BHXH","📋 Báo cáo định kỳ","🕒 Chấm công","💰 Tính thu nhập","📄 Quản lý Công văn & HĐ kinh tế","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🔑 Quản lý MK","🖼️ Tạo ảnh thẻ NV","🔍 Audit Dashboard","📘 Hướng dẫn sử dụng",]
elif st.session_state.role == "demo_readonly":
    # Vai trò DÀNH RIÊNG cho tài khoản demo công khai: thấy TOÀN BỘ menu như admin
    # (trừ Danh mục/Nhập-Xuất Excel/Audit vốn là công cụ cấu hình hệ thống, không
    # cần thiết cho việc "xem thử" tính năng), nhưng can_edit()/can_delete()/can_export()
    # đều trả về False với role này nên các nút Lưu/Sửa/Xóa/Xuất sẽ bị chặn.
    # LƯU Ý: các nút Lưu/Sửa/Xóa hiện KHÔNG kiểm tra can_edit()/can_delete() ở TẤT CẢ
    # màn hình trong app (chỉ mới có ở một số form) — trước khi phát hành tài khoản demo
    # công khai, cần rà soát thêm các nút còn thiếu để đảm bảo dữ liệu thật sự không đổi được.
    menu_options = ["📊 Dashboard","👤 Ứng viên","✅ Nhân viên","📁 Upload hồ sơ","📋 BHXH","📋 Báo cáo định kỳ","🕒 Chấm công","💰 Tính thu nhập","📄 Quản lý Công văn & HĐ kinh tế","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🖼️ Tạo ảnh thẻ NV","📘 Hướng dẫn sử dụng",]
else:  # 'nhan_vien' thường (mặc định) hoặc bất kỳ vai trò nào khác chưa được liệt kê ở trên
    # -> LUÔN có 1 menu tối thiểu an toàn, tuyệt đối KHÔNG được để menu_options rơi vào
    # trạng thái "không được gán" (từng gây lỗi NameError crash toàn bộ app khi đăng nhập).
    menu_options = ["📊 Dashboard","✅ Nhân viên","🕒 Chấm công","💬 Chat nội bộ","🤖 Chatbot Giải đáp","🔑 Quản lý MK","🖼️ Tạo ảnh thẻ NV","📘 Hướng dẫn sử dụng"]
# ── Phân luồng menu theo loại hình tenant (DN / HKD) ──
_tenant = st.session_state.get("tenant") or {}
if _tenant.get("loai_hinh") == "HO_KINH_DOANH":
    _menu_bo_hkd = {
        "👤 Ứng viên",
        "📋 Báo cáo định kỳ",
        "📄 Quản lý Công văn & HĐ kinh tế",
        "💬 Chat nội bộ",
        "🖼️ Tạo ảnh thẻ NV",
        "🔍 Audit Dashboard",
    }
    menu_options = [m for m in menu_options if m not in _menu_bo_hkd]
    # Thêm menu Thuế HKD — chèn sau 💰 Tính thu nhập (nếu có) hoặc cuối
    if "🧾 Thuế HKD" not in menu_options:
        if "💰 Tính thu nhập" in menu_options:
            _vi_tri = menu_options.index("💰 Tính thu nhập") + 1
        else:
            _vi_tri = len(menu_options) - 1
        menu_options.insert(_vi_tri, "🧾 Thuế HKD")
menu = st.sidebar.radio(i18n.t("📋 Menu"), menu_options, format_func=i18n.t)
st.sidebar.divider()
st.sidebar.caption(f"👤 {st.session_state.get('ho_ten_dang_nhap', st.session_state.username)} ({st.session_state.role})")

# Mobile: dropdown menu — HTML native, tự ẩn trên desktop bằng JS
import json as _json_menu
_menu_json = _json_menu.dumps(menu_options, ensure_ascii=False)
_current_idx = menu_options.index(menu) if menu in menu_options else 0
import streamlit.components.v1 as _comp_menu
_comp_menu.html(f"""
<select id="hrm-mob-sel" style="width:100%;padding:8px 12px;font-size:15px;
    border:1px solid #ddd;border-radius:8px;background:#f8f8f8;" onchange="
    var idx = this.selectedIndex;
    var url = new URL(window.top.location.href);
    url.searchParams.set('hmenu', idx);
    window.top.location.href = url.toString();
">
</select>
<script>
(function() {{
    var opts = {_menu_json};
    var sel = document.getElementById('hrm-mob-sel');
    for (var i = 0; i < opts.length; i++) {{
        var o = document.createElement('option');
        o.value = i; o.text = opts[i];
        if (i === {_current_idx}) o.selected = true;
        sel.appendChild(o);
    }}
    // Desktop: ẩn toàn bộ iframe container (không chiếm khoảng trắng)
    if (window.top.innerWidth >= 769) {{
        // Ẩn select
        sel.style.display = 'none';
        // Ẩn iframe container ở tầng parent (Streamlit tạo div bọc iframe)
        try {{
            var fr = window.frameElement;
            if (fr) {{
                fr.style.display = 'none';
                fr.style.height = '0';
                if (fr.parentElement) {{
                    fr.parentElement.style.display = 'none';
                    fr.parentElement.style.height = '0';
                    fr.parentElement.style.margin = '0';
                    fr.parentElement.style.padding = '0';
                    fr.parentElement.style.overflow = 'hidden';
                }}
            }}
        }} catch(e) {{}}
    }}
}})();
</script>
""", height=40)

# Đọc query param hmenu nếu có (mobile navigation)
_qp = st.query_params
if 'hmenu' in _qp:
    try:
        _midx = int(_qp['hmenu'])
        if 0 <= _midx < len(menu_options) and menu_options[_midx] != menu:
            menu = menu_options[_midx]
    except (ValueError, IndexError):
        pass
# MỚI:
if st.sidebar.button(i18n.t("🚪 THOÁT"), width='stretch'):
    st.session_state.logged_in = False
    st.session_state.role = None
    st.session_state.username = None
    st.session_state.pop('last_birthday_check', None)
    st.session_state.pop('sinh_nhat_hom_nay_list', None)
    st.cache_data.clear()
    st.rerun()

# Đặt caption bản quyền NGAY ĐÂY (thay vì cuối file) để nó luôn hiển thị ngay từ đầu,
# không phụ thuộc vào việc các màn hình (menu) phía sau có chạy trọn vẹn hay không —
# nếu 1 nhánh menu phía dưới lỗi/raise exception, dòng caption ở cuối file sẽ không
# bao giờ được thực thi.
st.sidebar.divider()
st.sidebar.caption("© 2026 HRM Master | © copyright: Mr.Tuyen - 0961778150")

# ========== HÀM DÙNG CHUNG: CARD THÔNG TIN NHÂN VIÊN ==========
PHONG_BAN_LANH_DAO_CAO_CAP = ('Hội Đồng Quản Trị', 'Ban Tổng Giám Đốc')

def render_employee_info_card(nv, key_prefix, on_close=None):
    """Hiển thị card '👤 THÔNG TIN NHÂN VIÊN' với avatar load on-demand"""
    st.subheader("👤 THÔNG TIN NHÂN SỰ")

    col_avatar, col_info = st.columns([1, 2])

    with col_avatar:
        st.markdown("""
        <style>
        .avatar-wrapper {
            display: flex;
            align-items: center;
            justify-content: center;
            height: 100%;
            min-height: 250px;
        }
        .avatar-img {
            width: 200px;
            height: 200px;
            border-radius: 50%;
            object-fit: cover;
            border: 4px solid #f59e0b;
            box-shadow: 0 4px 15px rgba(0,0,0,0.15);
        }
        </style>
        """, unsafe_allow_html=True)

        # ===== CẢI TIẾN: Chỉ tải ảnh khi có hành động =====
        anh_path = nv.get('anh_ho_so')
        avatar_key = f"avatar_loaded_{nv['id']}"
        
        if anh_path:
            # Dùng cache để tải ảnh (chỉ tải 1 lần, cache 1 giờ)
            anh_bytes = get_avatar_bytes_cached(anh_path)
            if anh_bytes:
                img_base64 = base64.b64encode(anh_bytes).decode()
                st.markdown(f"""
                <div class="avatar-wrapper">
                    <img src="data:image/jpeg;base64,{img_base64}" class="avatar-img" loading="lazy">
                </div>
                """, unsafe_allow_html=True)
            else:
                # Fallback: ảnh mặc định
                gioi_tinh = nv.get('gioi_tinh', '')
                ho_ten = nv.get('ho_ten', '')
                avatar_file = "avatar_male.png" if gioi_tinh == "Nam" else "avatar_female.png"
                avatar_path = os.path.join(os.path.dirname(__file__), "static", avatar_file)
                if os.path.exists(avatar_path):
                    with open(avatar_path, "rb") as f:
                        img_data = base64.b64encode(f.read()).decode()
                    st.markdown(f"""
                    <div class="avatar-wrapper">
                        <img src="data:image/png;base64,{img_data}" class="avatar-img">
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div class="avatar-wrapper">
                        <img src="https://ui-avatars.com/api/?name={ho_ten.replace(' ', '+')}&size=200&background=f59e0b&color=fff" class="avatar-img">
                    </div>
                    """, unsafe_allow_html=True)
        else:
            # Fallback khi không có ảnh
            gioi_tinh = nv.get('gioi_tinh', '')
            ho_ten = nv.get('ho_ten', '')
            avatar_file = "avatar_male.png" if gioi_tinh == "Nam" else "avatar_female.png"
            avatar_path = os.path.join(os.path.dirname(__file__), "static", avatar_file)
            if os.path.exists(avatar_path):
                with open(avatar_path, "rb") as f:
                    img_data = base64.b64encode(f.read()).decode()
                st.markdown(f"""
                <div class="avatar-wrapper">
                    <img src="data:image/png;base64,{img_data}" class="avatar-img">
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="avatar-wrapper">
                    <img src="https://ui-avatars.com/api/?name={ho_ten.replace(' ', '+')}&size=200&background=f59e0b&color=fff" class="avatar-img">
                </div>
                """, unsafe_allow_html=True)

    with col_info:
        # Xác định xem có phải lãnh đạo cấp cao không
        la_lanh_dao_cc = la_phong_ban_lanh_dao_cao_cap(nv.get('phong_ban_lam_viec'))
        
        if la_lanh_dao_cc:
            # Lãnh đạo cấp cao: chỉ hiển thị Ông/Bà + Họ tên, KHÔNG có ma_nv
            xung_ho_card = get_xung_ho_trang_trong(nv.get('gioi_tinh'))
            st.markdown(f"### {xung_ho_card} {nv['ho_ten']}".strip())
        else:
            # Nhân viên thường: hiển thị Họ tên (Mã NV)
            st.markdown(f"### {nv['ho_ten']} ({nv['ma_nv']})")

        info_col1, info_col2 = st.columns(2)

        with info_col1:
            st.markdown(f"**📅 Ngày sinh:** {format_date(nv.get('ngay_sinh'))}")
            st.markdown(f"**⚧ Giới tính:** {nv.get('gioi_tinh', 'Chưa cập nhật')}")
            if nv.get('chuc_danh_nghe'):
                st.markdown(f"**💼 Chức danh:** {nv.get('chuc_danh_nghe')}")
            st.markdown(f"**🏢 Phòng:** {nv.get('phong_ban_lam_viec', 'Chưa cập nhật')}")
            # Chỉ hiển thị chức vụ nếu có
            if nv.get('chuc_vu'):
                st.markdown(f"**🎖️ Chức vụ:** {nv.get('chuc_vu')}")
            st.markdown(f"**📞 SĐT:** {nv.get('dien_thoai', 'Chưa cập nhật')}")

        with info_col2:
            if nv.get('so_hdld'):
                st.markdown(f"**Số Hợp đồng:** {nv.get('so_hdld')}")
            # Chỉ hiển thị Loại HĐ cho nhân viên thường (không phải HĐQT/BTGĐ)
            if not la_lanh_dao_cc:
                st.markdown(f"**📋 Loại HĐ:** {nv.get('loai_hop_dong', 'Chưa cập nhật')}")
            if nv.get('ngay_vao_lam'):
                st.markdown(f"**📅 Ngày vào làm:** {format_date(nv.get('ngay_vao_lam'))}")
            st.markdown(f"**🎓 Trình độ:** {nv.get('trinh_do', 'Chưa cập nhật')}")
            st.markdown(f"**📇 Mã BHXH:** {nv.get('ma_so_bhxh', 'Chưa có')}")
            # Chỉ hiển thị trạng thái cho nhân viên thường
            if not la_lanh_dao_cc:
                trang_thai_text = {
                    'DANG_LAM': '🟢 Đang làm',
                    'THU_VIEC': '🔵 Thử việc',
                    'NGHI_VIEC': '🔴 Đã nghỉ'
                }
                status = trang_thai_text.get(nv.get('trang_thai'), nv.get('trang_thai', 'Chưa xác định'))
                st.markdown(f"**📊 Trạng thái:** {status}")
            # Nếu là lãnh đạo cấp cao, hiển thị chức vụ thay vì trạng thái
            else:
                st.markdown(f"**🎖️ Chức vụ:** {nv.get('chuc_vu', 'Thành viên')}")

    # ===== Nút hành động (thêm nút "Đóng" ở cuối) =====
    st.divider()
    col_btn_action1, col_btn_action2, col_btn_action3, col_btn_action4, col_btn_action5 = st.columns(5)
    if st.session_state.role in ("admin", "xem_toan_bo"):
        with col_btn_action1:
            if st.button("✏️ SỬA NHÂN VIÊN", width='stretch', type="primary", key=f"edit_nv_btn_{key_prefix}", disabled=not can_edit()):
                st.session_state['selected_nv_id'] = int(nv['id'])
                st.rerun()

        with col_btn_action3:
            if nv.get('trang_thai') in ('DANG_LAM', 'THU_VIEC'):
                # Sinh file hợp đồng ngay → download 1 click
                try:
                    db_hd = st.session_state.db_engine.get_connection()
                    c_hd = db_hd.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    c_hd.execute("SELECT * FROM nhan_vien WHERE id = %s", (int(nv['id']),))
                    nv_full_hd = c_hd.fetchone()
                    db_hd.close()
                    if nv_full_hd:
                        if nv.get('trang_thai') == 'DANG_LAM':
                            fp_hd = tao_hop_dong(nv_full_hd)
                            label_hd = "🖨️ IN HĐLĐ"
                            fname_hd = f"HDLD_{nv_full_hd['ho_ten']}_{datetime.now().strftime('%Y%m%d')}.docx"
                        else:
                            fp_hd = tao_hop_dong_thu_viec(nv_full_hd)
                            label_hd = "🖨️ IN HĐTV"
                            fname_hd = f"HDTV_{nv_full_hd['ho_ten']}_{datetime.now().strftime('%Y%m%d')}.docx"
                        with open(fp_hd, "rb") as f_hd:
                            st.download_button(
                                label=label_hd,
                                data=f_hd,
                                file_name=fname_hd,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_hd_{key_prefix}",
                                width='stretch'
                            )
                except Exception as e_hd:
                    st.button("📄 LỖI TẠO HĐ", disabled=True, width='stretch', key=f"err_hd_{key_prefix}")
            else:
                st.button("📄 KHÔNG THỂ IN HĐ", disabled=True, width='stretch', key=f"no_hd_{key_prefix}")

        with col_btn_action5:
            if st.button("❌ Đóng", width='stretch', key=f"close_profile_{key_prefix}"):
                if callable(on_close):
                    on_close()
                st.rerun()
        
        with col_btn_action4:
            ma_bhxh = nv.get('ma_so_bhxh', '')
            chua_co_bhxh = not bool(ma_bhxh and str(ma_bhxh).strip())
            if chua_co_bhxh:
                if st.button("🏠 NHẬP T.TIN HỘ GĐ", width='stretch', type="primary", key=f"bhxh_family_{key_prefix}", disabled=not can_edit()):
                    st.session_state['bhxh_family_nv_id'] = int(nv['id'])
                    st.session_state['bhxh_family_nv_name'] = nv['ho_ten']
                    st.rerun()
    else:
        # Role không phải admin (viewer, kt_luong, van_thu...): chỉ được XEM thông tin,
        # không có nút Sửa nhân viên / In HĐLĐ / Gửi Zalo — chỉ có nút Đóng.
        col_v1, col_v2, col_v3 = st.columns([2, 1, 2])
        with col_v2:
            if st.button("❌ ĐÓNG", width='stretch', key=f"close_profile_viewer_{key_prefix}"):
                if callable(on_close):
                    on_close()
                st.rerun()


# ========== DASHBOARD ==========
if menu == "📊 Dashboard":
    st.markdown(f"# {i18n.tm('📊 Dashboard')}", unsafe_allow_html=True)
    
    # Lấy dữ liệu từ cache
    stats = get_dashboard_stats()

    # ===== CẢNH BÁO 1: Chuẩn bị 6 Báo cáo định kỳ (30/06-05/07 và 31/12-05/01) =====
    hom_nay = date.today()
    md_hien_tai = (hom_nay.month, hom_nay.day)
    trong_khoang_giua_nam = (6, 30) <= md_hien_tai <= (7, 5)
    trong_khoang_cuoi_nam = md_hien_tai >= (12, 31) or md_hien_tai <= (1, 5)
    if trong_khoang_giua_nam or trong_khoang_cuoi_nam:
        han_nop_bc = "05/7" if trong_khoang_giua_nam else "05/01"
        st.warning(f"📋 **Chuẩn bị các 6 BC Định kỳ và nộp trước {han_nop_bc} HR nhé!**")

    # ===== CẢNH BÁO 2: Hạn nộp BC Tăng/Giảm BHXH hàng tháng (cấu hình riêng theo DN) =====
    han_ngay_bhxh = get_han_nop_bhxh()
    so_ngay_trong_thang = calendar.monthrange(hom_nay.year, hom_nay.month)[1]
    ngay_han_hop_le = min(han_ngay_bhxh, so_ngay_trong_thang)
    han_thang_nay = date(hom_nay.year, hom_nay.month, ngay_han_hop_le)
    so_ngay_con_lai_bhxh = (han_thang_nay - hom_nay).days
    if 0 <= so_ngay_con_lai_bhxh <= 5:
        thong_diep_bhxh = (f"Chuẩn bị nộp BC Tăng/Giảm BHXH tháng này nhé HR! "
                            f"(Hạn nộp: {han_thang_nay.strftime('%d/%m/%Y')}, còn {so_ngay_con_lai_bhxh} ngày)")
        if so_ngay_con_lai_bhxh <= 1:
            st.error(f"🚨 **KHẨN:** {thong_diep_bhxh}")
        elif so_ngay_con_lai_bhxh <= 3:
            st.warning(f"⚠️ {thong_diep_bhxh}")
        else:
            st.info(f"📌 {thong_diep_bhxh}")

    # ===== CẢNH BÁO 3: Nhắc liên hệ BHXH chốt sổ cho NV đã chấm dứt HĐLĐ (ngày 7-12) =====
    try:
        db_chot = st.session_state.db_engine.get_connection()
        c_chot = db_chot.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c_chot.execute("""
            SELECT q.nhan_vien_id, q.ngay_quyet_dinh, q.ngay_hieu_luc,
                   nv.ho_ten, nv.ma_nv, nv.ma_so_bhxh,
                   (CURRENT_DATE - q.ngay_quyet_dinh::date) AS so_ngay
            FROM quyet_dinh_nhan_su q
            JOIN nhan_vien nv ON nv.id = q.nhan_vien_id
            WHERE q.loai_quyet_dinh = 'CHAM_DUT_HD'
              AND nv.trang_thai = 'NGHI_VIEC'
              AND nv.loai_hop_dong IS DISTINCT FROM 'Thử việc'
              AND (CURRENT_DATE - q.ngay_quyet_dinh::date) BETWEEN 7 AND 12
            ORDER BY q.ngay_quyet_dinh ASC
        """)
        ds_can_chot_so = c_chot.fetchall()
        db_chot.close()
        if ds_can_chot_so:
            for nv_cs in ds_can_chot_so:
                so_ngay = nv_cs['so_ngay']
                ho_ten = nv_cs['ho_ten']
                ma_nv = nv_cs['ma_nv'] or ''
                ma_bhxh = nv_cs['ma_so_bhxh'] or '(chưa có)'
                ngay_qd = nv_cs['ngay_quyet_dinh']
                ngay_qd_str = ngay_qd.strftime('%d/%m/%Y') if hasattr(ngay_qd, 'strftime') else str(ngay_qd)
                han_con_lai = 12 - so_ngay
                if han_con_lai <= 2:
                    st.error(f"🚨 **KHẨN — Chốt sổ BHXH:** {ho_ten} ({ma_nv}) — Mã BHXH: {ma_bhxh} | "
                             f"QĐ chấm dứt ngày {ngay_qd_str}, đã {so_ngay} ngày — còn {han_con_lai} ngày để liên hệ BHXH chốt sổ!")
                else:
                    st.warning(f"📋 **Nhắc chốt sổ BHXH:** {ho_ten} ({ma_nv}) — Mã BHXH: {ma_bhxh} | "
                               f"QĐ chấm dứt ngày {ngay_qd_str}, đã {so_ngay} ngày — liên hệ cơ quan BHXH chốt sổ cho NV này.")
    except Exception as e:
        print(f"Lỗi kiểm tra chốt sổ BHXH: {e}")

    st.divider()
        
    
    
    # ── Phân bố chức danh ──
    import plotly.express as px
    import plotly.graph_objects as go

    db2 = st.session_state.db_engine.get_connection()
    c2 = db2.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    c2.execute("""
        SELECT chuc_danh_nghe, COUNT(*) t 
        FROM nhan_vien 
        WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC') AND so_hdld IS NOT NULL AND so_hdld != ''
        GROUP BY chuc_danh_nghe 
        ORDER BY t DESC
    """)
    data = c2.fetchall()
    db2.close()

    if data:
    # ========== PHẦN DASHBOARD NÂNG CAO ==========
        st.subheader("📊 TỔNG QUAN PHÂN BỐ NHÂN SỰ")

        db_dash = st.session_state.db_engine.get_connection()
        c_dash = db_dash.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

        # Tiêu chuẩn thống kê DUY NHẤT cho toàn bộ Dashboard: chỉ tính nhân sự đang làm
        # (đang làm hoặc thử việc) VÀ đã có số hợp đồng lao động (so_hdld) — loại bỏ
        # nhân viên đã nghỉ việc và nhân viên chưa có so_hdld (hồ sơ chưa hoàn thiện)
        # để số liệu giữa các biểu đồ luôn khớp nhau.
        DK_CHUAN_NV = "trang_thai IN ('DANG_LAM', 'THU_VIEC') AND so_hdld IS NOT NULL AND so_hdld != ''"

        # 1. Dữ liệu cho Table "Trạng thái nhân sự các phòng ban"
        c_dash.execute(f"""
            SELECT 
                phong_ban_lam_viec as "Phòng ban",
                COUNT(*) as "Tổng số",
                SUM(CASE WHEN trang_thai = 'DANG_LAM' THEN 1 ELSE 0 END) as "Đang làm",
                SUM(CASE WHEN trang_thai = 'THU_VIEC' THEN 1 ELSE 0 END) as "Thử việc"
            FROM nhan_vien
            WHERE {DK_CHUAN_NV}
            GROUP BY phong_ban_lam_viec
            ORDER BY "Tổng số" DESC
        """)
        table_data = c_dash.fetchall()
        table_data = sap_xep_phong_ban_rows(table_data, "Phòng ban")
        
        # 2. Dữ liệu cho các biểu đồ
        # a. Tỷ lệ nhân sự mỗi phòng ban
        c_dash.execute(f"""
            SELECT phong_ban_lam_viec as "Phòng ban", COUNT(*) as "Số lượng"
            FROM nhan_vien WHERE {DK_CHUAN_NV}
            GROUP BY phong_ban_lam_viec
            ORDER BY "Số lượng" DESC
        """)
        dept_data = c_dash.fetchall()
        dept_data = sap_xep_phong_ban_rows(dept_data, "Phòng ban")

        # b. Cơ cấu theo giới tính
        c_dash.execute(f"""
            SELECT gioi_tinh, COUNT(*) as "Số lượng"
            FROM nhan_vien WHERE {DK_CHUAN_NV}
            GROUP BY gioi_tinh
        """)
        gender_data = c_dash.fetchall()

        # c. Cơ cấu theo Trình độ học vấn
        c_dash.execute(f"""
            SELECT trinh_do, COUNT(*) as "Số lượng"
            FROM nhan_vien WHERE {DK_CHUAN_NV}
            GROUP BY trinh_do
            ORDER BY "Số lượng" DESC
        """)
        education_data = c_dash.fetchall()

        # d. Cơ cấu theo Chức danh (Top 10) — LẤY TOÀN BỘ chức danh ở SQL (không LIMIT ở đây),
        # việc gộp "Top 9 + Khác" để hiển thị được xử lý bên dưới bằng Python, để đảm bảo
        # Tổng của biểu đồ này luôn khớp với các biểu đồ khác (không bị mất nhân sự do LIMIT).
        c_dash.execute(f"""
            SELECT chuc_danh_nghe, COUNT(*) as "Số lượng"
            FROM nhan_vien WHERE {DK_CHUAN_NV}
            AND chuc_danh_nghe IS NOT NULL AND chuc_danh_nghe != ''
            GROUP BY chuc_danh_nghe
            ORDER BY "Số lượng" DESC
        """)
        role_data = c_dash.fetchall()

        # e. Cơ cấu theo Độ tuổi
        c_dash.execute(f"""
            SELECT 
                CASE 
                    WHEN EXTRACT(YEAR FROM age(CURRENT_DATE, ngay_sinh)) < 25 THEN 'Dưới 25 tuổi'
                    WHEN EXTRACT(YEAR FROM age(CURRENT_DATE, ngay_sinh)) BETWEEN 25 AND 34 THEN '25-34 tuổi'
                    WHEN EXTRACT(YEAR FROM age(CURRENT_DATE, ngay_sinh)) BETWEEN 35 AND 44 THEN '35-44 tuổi'
                    WHEN EXTRACT(YEAR FROM age(CURRENT_DATE, ngay_sinh)) BETWEEN 45 AND 54 THEN '45-54 tuổi'
                    ELSE 'Từ 55 tuổi trở lên'
                END as "Độ tuổi",
                COUNT(*) as "Số lượng"
            FROM nhan_vien
            WHERE {DK_CHUAN_NV} AND ngay_sinh IS NOT NULL
            GROUP BY "Độ tuổi"
            ORDER BY MIN(EXTRACT(YEAR FROM age(CURRENT_DATE, ngay_sinh)))
        """)
        seniority_data = c_dash.fetchall()

        # f. Biểu đồ đường: Xu hướng tuyển dụng theo tháng (6 tháng gần nhất)
        c_dash.execute(f"""
            SELECT 
                TO_CHAR(DATE_TRUNC('month', ngay_vao_lam), 'MM/YYYY') as "Tháng",
                COUNT(*) as "Số lượng"
            FROM nhan_vien
            WHERE ngay_vao_lam >= (CURRENT_DATE - INTERVAL '6 months') AND {DK_CHUAN_NV}
            GROUP BY DATE_TRUNC('month', ngay_vao_lam)
            ORDER BY DATE_TRUNC('month', ngay_vao_lam) ASC
        """)
        trend_data = c_dash.fetchall()

        db_dash.close()

        # --- RENDER BIỂU ĐỒ ĐA DẠNG ---
        import plotly.express as px
        import plotly.graph_objects as go

        MODERN_PALETTE = ['#0f3b5c', '#2196F3', '#4FC3F7', '#00BFA5', '#66BB6A', '#FFB74D', '#FF7043', '#AB47BC', '#78909C']
        CHART_HEIGHT = 300

        # Hàng 1: Table + Biểu đồ thanh + Biểu đồ tròn
        row1_col1, row1_col2, row1_col3 = st.columns(3)

        with row1_col1:
            st.markdown("**💼 Cơ cấu theo Chức danh (Top 10)**")
            if role_data:
                import plotly.express as px
                import plotly.graph_objects as go
                
                df_role = pd.DataFrame(role_data)
                # Tổng luôn tính trên TOÀN BỘ chức danh (không bị ảnh hưởng bởi việc
                # gộp Top 9 + Khác bên dưới), để khớp với các biểu đồ khác trong Dashboard.
                total = df_role['Số lượng'].sum()

                # Nếu có nhiều hơn 10 chức danh: giữ 9 chức danh đông nhất, gộp phần còn
                # lại thành 1 dòng "Khác" — vừa gọn (đúng tinh thần "Top 10"), vừa không
                # làm mất người khỏi biểu đồ/tổng như khi dùng LIMIT 10 ở SQL.
                if len(df_role) > 10:
                    df_top = df_role.iloc[:9].copy()
                    so_luong_khac = df_role.iloc[9:]['Số lượng'].sum()
                    cac_chuc_danh_khac = df_role.iloc[9:]['chuc_danh_nghe'].tolist()
                    df_khac = pd.DataFrame([{
                        'chuc_danh_nghe': f"Khác ({len(cac_chuc_danh_khac)} chức danh)",
                        'Số lượng': so_luong_khac
                    }])
                    df_role = pd.concat([df_top, df_khac], ignore_index=True)
                
                # Tạo labels với format: "Chức danh\nSố lượng (tỷ lệ%)"
                labels_with_stats = []
                for _, row in df_role.iterrows():
                    pct = (row['Số lượng'] / total * 100)
                    labels_with_stats.append(f"{row['chuc_danh_nghe']}\n{row['Số lượng']} ({pct:.1f}%)")
                
                # Sử dụng biểu đồ hình tròn với labels đã format
                fig_role = go.Figure(data=[go.Pie(
                    labels=labels_with_stats,
                    values=df_role['Số lượng'],
                    hole=0.55,
                    textinfo='label',
                    textposition='outside',
                    textfont=dict(size=11, color='#1e293b'),
                    marker=dict(
                        colors=px.colors.qualitative.Safe,
                        line=dict(color='white', width=2)
                    ),
                    hovertemplate='<b>%{label}</b><br>Số lượng: %{value}<br>Tỷ lệ: %{percent:.1f}%<extra></extra>'
                )])
                fig_role.update_layout(
                    title=dict(
                        text=f"<b>Tổng: {total} nhân viên</b>",
                        x=0.5, y=0.5,
                        xanchor='center', yanchor='middle',
                        font=dict(size=14, color='#0f3b5c')
                    ),
                    showlegend=False,
                    margin=dict(t=40, b=40, l=10, r=10),
                    height=280,
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)',
                )
                st.plotly_chart(fig_role, use_container_width=True)
            else:
                st.info("Không có dữ liệu")  
        
        with row1_col2:
            st.markdown("**🎓 Cơ cấu theo Trình độ học vấn**")
            if education_data:
                df_edu = pd.DataFrame(education_data)
                df_edu['trinh_do'] = df_edu['trinh_do'].fillna('Chưa cập nhật')
                # Sử dụng biểu đồ thanh đứng thay vì tròn
                fig_edu = px.bar(
                    df_edu,
                    x='trinh_do',
                    y='Số lượng',
                    color='trinh_do',
                    text='Số lượng',
                    color_discrete_sequence=px.colors.qualitative.Set3
                )
                fig_edu.update_layout(
                    margin=dict(t=0, b=0, l=0, r=0),
                    height=280,
                    xaxis_title="",
                    yaxis_title="Số lượng",
                    showlegend=False
                )
                fig_edu.update_traces(textposition='outside')
                st.plotly_chart(fig_edu, use_container_width=True)
            else:
                st.info("Không có dữ liệu")
        
        with row1_col3:
            st.markdown("**🥧 Cơ cấu theo Phòng ban**")
            if dept_data:
                df_dept = pd.DataFrame(dept_data)
                fig_dept = px.pie(
                    df_dept, 
                    names='Phòng ban', 
                    values='Số lượng',
                    hole=0.4,
                    color_discrete_sequence=px.colors.qualitative.Pastel
                )
                # QUAN TRỌNG: Plotly Pie mặc định tự sắp xếp lát cắt theo giá trị giảm dần
                # (sort=True), làm mất thứ tự ưu tiên phòng ban đã chuẩn hóa ở dept_data
                # (sap_xep_phong_ban_rows theo PHONG_BAN_THU_TU). Tắt sort để giữ đúng thứ tự.
                fig_dept.update_traces(sort=False)
                fig_dept.update_layout(margin=dict(t=0, b=0, l=0, r=0), height=280)
                st.plotly_chart(fig_dept, use_container_width=True)
            else:
                st.info("Không có dữ liệu")

        # Hàng 2: Biểu đồ tròn + Biểu đồ đường + Biểu đồ thanh
        row2_col1, row2_col2, row2_col3 = st.columns(3)

        with row2_col1:
            st.markdown("**👫 Cơ cấu theo Giới tính**")
            if gender_data:
                df_gender = pd.DataFrame(gender_data)
                
                # Màu sắc nổi bật cho từng giới tính
                color_map = {
                    'Nam': '#2196F3',  # Xanh dương đẹp
                    'Nữ': '#FF6B6B',   # Đỏ hồng
                    'Khác': '#FFD93D'  # Vàng
                }
                colors = [color_map.get(g, '#95a5a6') for g in df_gender['gioi_tinh']]
                
                # Tạo donut chart với hiệu ứng đẹp
                fig_gender = go.Figure(data=[go.Pie(
                    labels=df_gender['gioi_tinh'],
                    values=df_gender['Số lượng'],
                    hole=0.4,
                    marker=dict(
                        colors=colors,
                        line=dict(color='white', width=3)
                    ),
                    textinfo='label+value+percent',
                    textposition='auto',
                    textfont=dict(size=12, color='#2c3e50', family='Arial Black'),
                    insidetextorientation='radial',
                    hovertemplate='<b>%{label}</b><br>Số lượng: %{value}<br>Tỷ lệ: %{percent:.1f}%<extra></extra>',
                    pull=[0.05 if i == 0 else 0 for i in range(len(df_gender))],  # Tách nhẹ phần tử đầu tiên
                    sort=False
                )])
                
                # Thêm vòng tròn bên trong với tổng số
                total = sum(df_gender['Số lượng'])
                fig_gender.add_annotation(
                    x=0.5, y=0.5,
                    text=f"<b>{total}</b>",
                    showarrow=False,
                    font=dict(size=24, color='#2c3e50', family='Arial Black'),
                    align='center'
                )
                fig_gender.add_annotation(
                    x=0.5, y=0.42,
                    text="Tổng",
                    showarrow=False,
                    font=dict(size=12, color='#7f8c8d', family='Arial'),
                    align='center'
                )
                
                fig_gender.update_layout(
                    margin=dict(t=10, b=10, l=10, r=10),
                    height=280,
                    showlegend=True,
                    legend=dict(
                        orientation='h',
                        yanchor='bottom',
                        y=-0.15,
                        xanchor='center',
                        x=0.5,
                        font=dict(size=12, color='#2c3e50')
                    ),
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)'
                )
                
                st.plotly_chart(fig_gender, use_container_width=True)
            else:
                st.info("Không có dữ liệu")

        with row2_col2:
            st.markdown("**📈 Xu hướng tuyển dụng 6 tháng**")
            if trend_data:
                df_trend = pd.DataFrame(trend_data)
                fig_trend = px.line(
                    df_trend,
                    x='Tháng',
                    y='Số lượng',
                    markers=True,
                    line_shape='spline'
                )
                fig_trend.update_layout(
                    margin=dict(t=0, b=0, l=0, r=0),
                    height=280,
                    xaxis_title="",
                    yaxis_title="Số lượng",
                    showlegend=False
                )
                fig_trend.update_traces(
                    line=dict(color='#f59e0b', width=3),
                    marker=dict(size=10, color='#0f3b5c')
                )
                st.plotly_chart(fig_trend, use_container_width=True)
            else:
                st.info("Không có dữ liệu")

        with row2_col3:
            st.markdown("**🎂 Cơ cấu theo Độ tuổi**")
            if seniority_data:
                df_sen = pd.DataFrame(seniority_data)
                order = ['Dưới 25 tuổi', '25-34 tuổi', '35-44 tuổi', '45-54 tuổi', 'Từ 55 tuổi trở lên']
                df_sen['Độ tuổi'] = pd.Categorical(df_sen['Độ tuổi'], categories=order, ordered=True)
                df_sen = df_sen.sort_values('Độ tuổi')
                
                # Sử dụng biểu đồ tròn với màu sắc gradient
                colors = ['#FFEAA7', '#FDCB6E', '#E17055', '#D63031', '#6C5CE7']
                fig_sen = go.Figure(data=[go.Pie(
                    labels=df_sen['Độ tuổi'],
                    values=df_sen['Số lượng'],
                    marker=dict(colors=colors[:len(df_sen)]),
                    textinfo='percent',
                    textposition='inside',
                    textfont=dict(size=11, color='white'),
                    hole=0.3,
                    # QUAN TRỌNG: mặc định Pie tự sắp xếp theo giá trị giảm dần (sort=True),
                    # phá vỡ thứ tự tuổi tăng dần đã set ở df_sen. Tắt sort để giữ đúng thứ tự nhóm tuổi.
                    sort=False
                )])
                fig_sen.update_layout(
                    margin=dict(t=10, b=40, l=10, r=10),
                    height=280,
                    showlegend=True,
                    legend=dict(
                        orientation='h',
                        yanchor='bottom',
                        y=-0.25,
                        xanchor='center',
                        x=0.5,
                        font=dict(size=10)
                    ),
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)',
                )
                st.plotly_chart(fig_sen, use_container_width=True)
            else:
                st.info("Không có dữ liệu")

        # Hàng 3: 2 biểu đồ còn lại
        row3_col1, row3_col2, row3_col3 = st.columns(3)

        with row3_col1:
            st.markdown("**📋 Trạng thái nhân sự các phòng ban**")
            if table_data:
                df_table = pd.DataFrame(table_data)
                # Định dạng số và hiển thị
                st.dataframe(df_table, hide_index=True, width='stretch', height=280)
            else:
                st.info("Không có dữ liệu")

        
            

        with row3_col3:
            st.markdown("**📊 Tổng hợp nhân sự**")
            # Hiển thị các chỉ số KPI quan trọng
            if dept_data:
                total_employees = sum([d['Số lượng'] for d in dept_data])
                total_depts = len(dept_data)
                avg_per_dept = total_employees / total_depts if total_depts > 0 else 0
                
                st.metric("🏢 Tổng số phòng ban", total_depts)
                st.metric("👥 Tổng nhân viên", f"{total_employees:,}")
                st.metric("📊 Trung bình/phòng", f"{avg_per_dept:.1f}")
                
                # Thêm thông tin phòng ban đông nhất
                if dept_data:
                    max_dept = max(dept_data, key=lambda x: x['Số lượng'])
                    st.info(f"🏆 Phòng đông nhất: **{max_dept['Phòng ban']}** ({max_dept['Số lượng']} NV)")
            else:
                st.info("Không có dữ liệu")

        # ========== KẾT THÚC PHẦN DASHBOARD NÂNG CAO ==========

# Gọi kiểm tra sinh nhật (đã xóa 2 dòng debug)
    auto_check_birthday()

    # 👇 Hiển thị banner cố định nếu có sinh nhật hôm nay
    sinh_nhat_list = st.session_state.get('sinh_nhat_hom_nay_list', [])
    if sinh_nhat_list:
        for sn in sinh_nhat_list:
            st.success(
                f"🎂 **Chúc mừng sinh nhật {sn['xung_ho']} {sn['ho_ten']} ({sn['ma_nv']})** — Hôm nay là sinh nhật của {sn['xung_ho']}! 🎉",
                icon="🎂"
            )
    
    # ========== PHẦN SINH NHẬT HOÀN CHỈNH ==========
    st.subheader("🎂 SINH NHẬT")

    # ⚠️ FIX BUG: không tái sử dụng cursor `c` của các nút "IN HĐLĐ/IN HĐTV" phía trên
    # (cursor đó đã bị đóng bằng db.close() ngay sau khi dùng) — mở kết nối MỚI riêng
    # cho toàn bộ phần Sinh nhật để tránh lỗi "cursor already closed" khi bấm các nút đó.
    db_sn = st.session_state.db_engine.get_connection()
    c = db_sn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Tạo tabs cho sinh nhật
    tab_trong_thang, tab_hom_nay, tab_lich_su = st.tabs(["📅 Sinh nhật trong tháng", "🎉 Hôm nay", "📜 Lịch sử đã gửi"])

    with tab_trong_thang:
        # Lấy danh sách sinh nhật trong tháng
        c.execute("""
            SELECT id, ma_nv, ho_ten, ngay_sinh, gioi_tinh, dien_thoai, email_lien_he, chuc_danh_nghe
            FROM nhan_vien 
            WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC')
            AND ngay_sinh IS NOT NULL
            AND EXTRACT(MONTH FROM ngay_sinh) = EXTRACT(MONTH FROM CURRENT_DATE)
            ORDER BY EXTRACT(DAY FROM ngay_sinh) ASC
        """)
        sinh_nhat_trong_thang = c.fetchall()
        
        if sinh_nhat_trong_thang:
            st.success(f"📅 Tháng {datetime.now().month} có **{len(sinh_nhat_trong_thang)}** nhân viên có sinh nhật:")
            
            # Hiển thị dạng grid
            cols = st.columns(3)
            for idx, sn in enumerate(sinh_nhat_trong_thang):
                with cols[idx % 3]:
                    ngay_sinh = sn.get('ngay_sinh')
                    if ngay_sinh:
                        # Tính tuổi
                        today = date.today()
                        tuoi = today.year - ngay_sinh.year
                        if today.month < ngay_sinh.month or (today.month == ngay_sinh.month and today.day < ngay_sinh.day):
                            tuoi -= 1
                        
                        # Tính ngày sinh nhật trong năm nay
                        sinh_nhat_nam_nay = date(today.year, ngay_sinh.month, ngay_sinh.day)
                        is_today = sinh_nhat_nam_nay == today
                        da_qua = sinh_nhat_nam_nay < today
                        
                        xung_ho = get_xung_ho(sn.get('gioi_tinh'), sn['ho_ten'])
                        
                        if is_today:
                            st.markdown(f"""
                            <div style='background: linear-gradient(135deg, #ffd700 0%, #ffed4e 100%); 
                                        padding: 15px; border-radius: 15px; margin: 10px 0; 
                                        border: 2px solid #ff9800; box-shadow: 0 4px 8px rgba(0,0,0,0.1);'>
                                <div style='font-size: 30px; text-align: center;'>🎉🎂</div>
                                <h4 style='text-align: center; color: #d32f2f; margin: 5px 0;'>
                                    <b>HÔM NAY LÀ SINH NHẬT!</b>
                                </h4>
                                <h3 style='text-align: center; color: #333; margin: 5px 0;'>
                                    {xung_ho} <b>{sn['ho_ten']}</b>
                                </h3>
                                <p style='text-align: center; color: #666;'>
                                    📅 {format_date(ngay_sinh)} (🎂 {tuoi} tuổi)<br>
                                    💼 {sn.get('chuc_danh_nghe', 'Chưa cập nhật')}<br>
                                    📞 {sn.get('dien_thoai', 'Chưa cập nhật')}
                                </p>
                            </div>
                            """, unsafe_allow_html=True)
                        elif da_qua:
                            st.markdown(f"""
                            <div style='background-color: #e0e0e0; padding: 12px; border-radius: 10px; margin: 8px 0;'>
                                <b>✅ {sn['ho_ten']}</b><br>
                                🎂 Sinh ngày: {format_date(ngay_sinh)}<br>
                                💼 {sn.get('chuc_danh_nghe', '')}
                            </div>
                            """, unsafe_allow_html=True)
                        else:
                            ngay_con_lai = (sinh_nhat_nam_nay - today).days
                            st.markdown(f"""
                            <div style='background-color: #e3f2fd; padding: 12px; border-radius: 10px; margin: 8px 0;'>
                                <b>🎂 {sn['ho_ten']}</b><br>
                                📅 Sinh ngày: {format_date(ngay_sinh)}<br>
                                ⏰ Còn {ngay_con_lai} ngày nữa<br>
                                💼 {sn.get('chuc_danh_nghe', '')}
                            </div>
                            """, unsafe_allow_html=True)
        else:
            st.info("📭 Tháng này không có ai sinh nhật.")
            
            # Hiển thị sinh nhật tháng sau
            c.execute("""
                SELECT id, ma_nv, ho_ten, ngay_sinh, gioi_tinh, chuc_danh_nghe
                FROM nhan_vien 
                WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC')
                AND ngay_sinh IS NOT NULL
                AND EXTRACT(MONTH FROM ngay_sinh) = EXTRACT(MONTH FROM CURRENT_DATE + INTERVAL '1 month')
                ORDER BY EXTRACT(DAY FROM ngay_sinh) ASC
                LIMIT 10
            """)
            sinh_nhat_thang_sau = c.fetchall()
            if sinh_nhat_thang_sau:
                st.caption("📅 Sinh nhật tháng sau:")
                for sn in sinh_nhat_thang_sau:
                    xung_ho = get_xung_ho(sn.get('gioi_tinh'), sn['ho_ten'])
                    st.caption(f"🎂 {xung_ho} **{sn['ho_ten']}** - {format_date(sn['ngay_sinh'])}")

    with tab_hom_nay:
        # Lấy danh sách sinh nhật hôm nay
        c.execute("""
            SELECT id, ma_nv, ho_ten, ngay_sinh, gioi_tinh, dien_thoai, email_lien_he, 
                   chuc_danh_nghe, phong_ban_lam_viec
            FROM nhan_vien 
            WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC')
            AND ngay_sinh IS NOT NULL
            AND EXTRACT(MONTH FROM ngay_sinh) = EXTRACT(MONTH FROM CURRENT_DATE)
            AND EXTRACT(DAY FROM ngay_sinh) = EXTRACT(DAY FROM CURRENT_DATE)
        """)
        sinh_nhat_hom_nay = c.fetchall()
        
        if sinh_nhat_hom_nay:
            st.balloons()
            for sn in sinh_nhat_hom_nay:
                ngay_sinh = sn.get('ngay_sinh')
                today = date.today()
                tuoi = today.year - ngay_sinh.year
                
                xung_ho = get_xung_ho(sn.get('gioi_tinh'), sn['ho_ten'])
                
                st.markdown(f"""
                <div style='background: linear-gradient(135deg, #ff6b6b 0%, #ff8e8e 100%);
                            padding: 25px; border-radius: 20px; margin: 15px 0;
                            text-align: center; color: white;'>
                    <div style='font-size: 50px;'>🎉🎂🎉</div>
                    <h1 style='color: white; margin: 10px 0;'>CHÚC MỪNG SINH NHẬT!</h1>
                    <h2 style='color: #fff3e0; margin: 10px 0;'>{xung_ho} {sn['ho_ten']}</h2>
                    <p style='font-size: 18px;'>
                        🎂 {tuoi} tuổi - Một tuổi mới thật nhiều niềm vui! 🎂
                    </p>
                    <p style='margin-top: 15px;'>
                        📅 {format_date(ngay_sinh)} | 💼 {sn.get('chuc_danh_nghe', '')} | 
                        🏢 {sn.get('phong_ban_lam_viec', '')}
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                # Hiển thị thông tin liên hệ
                # Hiển thị thông tin liên hệ (tất cả đều thấy)
                col_phone, col_email_info = st.columns(2)
                with col_phone:
                    if sn.get('dien_thoai'):
                        st.markdown(f"📞 **SĐT:** {sn['dien_thoai']}")
                    else:
                        st.warning("⚠️ Chưa cập nhật số điện thoại")
                with col_email_info:
                    if sn.get('email_lien_he'):
                        st.markdown(f"📧 **Email:** {sn['email_lien_he']}")
                    else:
                        st.warning("⚠️ Chưa cập nhật email")

                # Nút gửi — CHỈ ADMIN mới thấy
                if st.session_state.get('role') == 'admin':
                    col_btn_zalo, col_btn_email = st.columns(2)
                    
                    with col_btn_zalo:
                        if sn.get('dien_thoai'):
                            sdt = sn['dien_thoai'].replace('+84', '0').replace(' ', '').strip()
                            if st.button(f"📱 Gửi Zalo cho {sn['ho_ten']}", 
                                         key=f"zalo_sn_{sn['id']}", 
                                         width='stretch', 
                                         type="primary"):
                                tuoi_nv = date.today().year - sn['ngay_sinh'].year
                                loi_chuc_nv = get_loi_chuc_sinh_nhat(sn['ho_ten'], sn.get('gioi_tinh'), tuoi_nv)
                                st.code(loi_chuc_nv)
                                st.markdown(f"[👉 NHẤN ĐỂ GỬI QUA ZALO CHO {sn['ho_ten']}](https://zalo.me/{sdt})")
                        else:
                            st.button("📱 Gửi Zalo", disabled=True, 
                                      key=f"zalo_sn_disabled_{sn['id']}", 
                                      width='stretch',
                                      help="Chưa có số điện thoại")

                    with col_btn_email:
                        if sn.get('email_lien_he'):
                            if st.button(f"📧 Gửi Email cho {sn['ho_ten']}", 
                                         key=f"email_sn_{sn['id']}", 
                                         width='stretch'):
                                st.info(f"📧 Email: {sn['email_lien_he']}")
                                st.toast(f"Chức năng gửi email đang phát triển!", icon="📧")
                        else:
                            st.button("📧 Gửi Email", disabled=True, 
                                      key=f"email_sn_disabled_{sn['id']}", 
                                      width='stretch',
                                      help="Chưa có email")
                else:
                    # Tài khoản thường — hiện thông báo thay vì nút
                    st.caption("💡 Liên hệ admin để gửi lời chúc sinh nhật.")
        else:
            st.info("🎉 Hôm nay không có ai sinh nhật.")
            
            # Gợi ý sinh nhật sắp tới
            c.execute("""
                SELECT id, ma_nv, ho_ten, ngay_sinh, gioi_tinh
                FROM nhan_vien 
                WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC')
                AND ngay_sinh IS NOT NULL
                AND EXTRACT(MONTH FROM ngay_sinh) = EXTRACT(MONTH FROM CURRENT_DATE)
                AND EXTRACT(DAY FROM ngay_sinh) > EXTRACT(DAY FROM CURRENT_DATE)
                ORDER BY EXTRACT(DAY FROM ngay_sinh) ASC
                LIMIT 3
            """)
            sinh_nhat_sap_toi = c.fetchall()
            if sinh_nhat_sap_toi:
                st.subheader("📅 Sinh nhật sắp tới trong tháng:")
                for sn in sinh_nhat_sap_toi:
                    xung_ho = get_xung_ho(sn.get('gioi_tinh'), sn['ho_ten'])
                    st.info(f"🎂 {xung_ho} **{sn['ho_ten']}** - {format_date(sn['ngay_sinh'])}")

    with tab_lich_su:
        if st.session_state.role in ("admin", "xem_toan_bo"):
            st.subheader("📜 Lịch sử đã gửi lời chúc sinh nhật")
            
            # Kiểm tra bảng lịch sử tồn tại chưa
            try:
                c.execute("""
                    SELECT ls.*, nv.ho_ten, nv.ma_nv
                    FROM lich_su_gui_loi_chuc ls
                    JOIN nhan_vien nv ON ls.nhan_vien_id = nv.id
                    WHERE ls.loai_chuc = 'SINH_NHAT'
                    ORDER BY ls.ngay_gui DESC
                    LIMIT 50
                """)
                lich_su = c.fetchall()
                
                if lich_su:
                    ls_data = []
                    for ls in lich_su:
                        ls_data.append({
                            "Ngày gửi": format_date(ls['ngay_gui']),
                            "Mã NV": ls['ma_nv'],
                            "Họ tên": ls['ho_ten'],
                            "Kênh gửi": ls['kenh_gui'],
                            "Trạng thái": "✅ Đã gửi" if ls['trang_thai'] == 'DA_GUI' else ls['trang_thai']
                        })
                    df_ls = pd.DataFrame(ls_data)
                    st.dataframe(df_ls, width='stretch', hide_index=True)
                else:
                    st.info("📭 Chưa có lịch sử gửi lời chúc nào.")
            except Exception as e:
                st.info("📭 Chưa có dữ liệu lịch sử. Bảng lịch sử có thể chưa được tạo.")
        else:
            st.info("🔒 Chỉ Admin mới xem được lịch sử gửi lời chúc.")

    db_sn.close()

    st.divider()

    # Nút gửi lời chúc sinh nhật (chỉ admin)
    if st.session_state.role in ("admin", "xem_toan_bo") and sinh_nhat_trong_thang:
        with st.expander("💌 GỬI LỜI CHÚC SINH NHẬT", expanded=False):
            st.subheader("Gửi lời chúc sinh nhật đến nhân viên")
            
            # Chọn nhân viên để gửi
            sn_options = {}
            for sn in sinh_nhat_trong_thang:
                ngay_sinh = sn.get('ngay_sinh')
                xung_ho = get_xung_ho(sn.get('gioi_tinh'), sn['ho_ten'])
                label = f"{xung_ho} {sn['ho_ten']} - {format_date(ngay_sinh)}"
                sn_options[label] = sn
            
            # SAU KHI SỬA (ĐÚNG) — dùng key động theo từng nhân viên
            selected_label = st.selectbox("Chọn nhân viên:", list(sn_options.keys()), key="chon_sn_gui", help="💡 Gõ mã NV hoặc tên để tìm nhanh trong danh sách")
            selected_sn = sn_options[selected_label]

            # Tính tuổi
            if selected_sn.get('ngay_sinh'):
                today = date.today()
                tuoi = today.year - selected_sn['ngay_sinh'].year
                if today.month < selected_sn['ngay_sinh'].month or (
                    today.month == selected_sn['ngay_sinh'].month and 
                    today.day < selected_sn['ngay_sinh'].day
                ):
                    tuoi -= 1
            else:
                tuoi = None

            default_chuc = get_loi_chuc_sinh_nhat(
                selected_sn['ho_ten'], 
                selected_sn.get('gioi_tinh'), 
                tuoi
            )

            # ✅ KEY ĐỘNG theo nv_id — buộc Streamlit re-render lại text_area mỗi khi chọn người mới
            loi_chuc = st.text_area(
                "📝 Lời chúc sinh nhật:", 
                value=default_chuc, 
                height=250, 
                key=f"loi_chuc_sn_{selected_sn['id']}"  # ← thay "loi_chuc_sn_gui" bằng key động này
            )
            
            col_zalo, col_email, col_cancel = st.columns(3)
            
            with col_zalo:
                if st.button("📱 GỬI QUA ZALO", width='stretch', type="primary"):
                    sdt = selected_sn.get('dien_thoai', '')
                    if sdt:
                        sdt = sdt.replace('+84', '0').replace(' ', '').strip()
                        st.code(loi_chuc)
                        st.markdown(f"[👉 NHẤN VÀO ĐÂY ĐỂ GỬI QUA ZALO CHO {selected_sn['ho_ten']}](https://zalo.me/{sdt})")
                        st.success(f"✅ Đã sao chép nội dung! Vui lòng nhấn link Zalo để gửi.")
                        
                        # Lưu lịch sử
                        try:
                            db_log = st.session_state.db_engine.get_connection()
                            cur_log = db_log.cursor()
                            cur_log.execute("""
                                INSERT INTO lich_su_gui_loi_chuc (nhan_vien_id, loai_chuc, noi_dung, kenh_gui, trang_thai)
                                VALUES (%s, %s, %s, %s, %s)
                            """, (selected_sn['id'], 'SINH_NHAT', loi_chuc[:500], 'ZALO', 'DA_GUI'))
                            db_log.commit()
                            db_log.close()
                            st.toast("Đã lưu lịch sử gửi!", icon="✅")
                        except:
                            pass
                    else:
                        st.error("❌ Nhân viên chưa có số điện thoại! Vui lòng cập nhật SĐT trước.")
            
            with col_email:
                if st.button("📧 GỬI QUA EMAIL", width='stretch'):
                    email = selected_sn.get('email_lien_he', '')
                    if email:
                        try:
                            xung_ho = get_xung_ho(selected_sn.get('gioi_tinh'), selected_sn['ho_ten'])
                            
                            msg = MIMEMultipart()
                            msg['From'] = EMAIL_CONFIG['email']
                            msg['To'] = email
                            msg['Subject'] = f"🎂 Chúc mừng sinh nhật {xung_ho} {selected_sn['ho_ten']} - {COMPANY_CONFIG.get('ten_cong_ty', '')}"
                            
                            html_content = f"""
                            <html>
                            <head>
                                <meta charset="UTF-8">
                            </head>
                            <body style='font-family: "Times New Roman", Arial, sans-serif;'>
                                <div style='background: linear-gradient(135deg, #ffd700 0%, #ff9800 100%); 
                                            padding: 20px; text-align: center; border-radius: 10px;'>
                                    <h1 style='color: white;'>🎂 CHÚC MỪNG SINH NHẬT 🎂</h1>
                                </div>
                                <div style='padding: 20px; line-height: 1.6;'>
                                    {loi_chuc.replace(chr(10), '<br>')}
                                </div>
                                <hr>
                                <p style='color: #999; font-size: 11px; text-align: center;'>
                                    Email được gửi tự động từ hệ thống HRM-Port {COMPANY_CONFIG.get('ten_cong_ty', '')}<br>
                                    Địa chỉ: {COMPANY_CONFIG.get('dia_chi', '')} | Điện thoại: {COMPANY_CONFIG.get('dien_thoai_cty', '')}
                                </p>
                            </body>
                            </html>
                            """
                            msg.attach(MIMEText(html_content, 'html', 'utf-8'))
                            
                            server = smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
                            server.starttls()
                            server.login(EMAIL_CONFIG['email'], EMAIL_CONFIG['password'])
                            server.send_message(msg)
                            server.quit()
                            
                            st.success(f"✅ Đã gửi lời chúc sinh nhật qua email cho {xung_ho} {selected_sn['ho_ten']}!")
                            st.cache_data.clear()
                            
                            # Lưu lịch sử
                            db_log = st.session_state.db_engine.get_connection()
                            cur_log = db_log.cursor()
                            cur_log.execute("""
                                INSERT INTO lich_su_gui_loi_chuc (nhan_vien_id, loai_chuc, noi_dung, kenh_gui, trang_thai)
                                VALUES (%s, %s, %s, %s, %s)
                            """, (selected_sn['id'], 'SINH_NHAT', loi_chuc[:500], 'EMAIL', 'DA_GUI'))
                            db_log.commit()
                            db_log.close()
                        except Exception as e:
                            st.error(f"❌ Lỗi gửi email: {e}")
                    else:
                        st.error("❌ Nhân viên chưa có email! Vui lòng cập nhật email trước.")
            
            with col_cancel:
                st.write("")  # Placeholder

    # ========== KẾT THÚC PHẦN SINH NHẬT ==========


    # ── Thông báo ──
    st.subheader("📌 Thông báo")
    db3 = st.session_state.db_engine.get_connection()
    c3 = db3.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    c3.execute("SELECT ho_ten FROM nhan_vien WHERE DATE(ngay_vao_lam)=CURRENT_DATE")
    hn = c3.fetchall()
    c3.execute("SELECT ho_ten FROM nhan_vien WHERE DATE(ngay_vao_lam)=CURRENT_DATE - INTERVAL '1 day'")
    hq = c3.fetchall()
    if hn:
        st.success(f"🟢 Hôm nay có thêm: **{', '.join([x['ho_ten'] for x in hn])}**")
    if hq:
        st.info(f"🔵 Hôm qua có thêm: **{', '.join([x['ho_ten'] for x in hq])}**")
    if st.session_state.role in ("admin", "xem_toan_bo"):
        c3.execute("""
            SELECT STT, ma_nv, ho_ten, ngay_vao_lam, 
                   (ngay_vao_lam + INTERVAL '30 days')::DATE as ngay_ket_thuc_tv,
                   ((ngay_vao_lam + INTERVAL '30 days')::DATE - CURRENT_DATE) as ngay_con_lai
            FROM nhan_vien 
            WHERE trang_thai = 'THU_VIEC' 
            AND (ngay_vao_lam + INTERVAL '30 days')::DATE <= CURRENT_DATE + INTERVAL '5 days'
            ORDER BY ngay_con_lai ASC
        """)
        tv_sap_het = c3.fetchall()
        for x in tv_sap_het:
            ngay_con_lai = x['ngay_con_lai']
            if isinstance(ngay_con_lai, timedelta):
                ngay_con_lai = ngay_con_lai.days
            if ngay_con_lai < 0:
                st.error(f"🔴 **{x.get('ma_nv','')} {x['ho_ten']}** - ĐÃ QUÁ THỜI HẠN THỬ VIỆC {abs(ngay_con_lai)} NGÀY!")
            elif ngay_con_lai == 0:
                st.error(f"⚠️ **{x.get('ma_nv','')} {x['ho_ten']}** - HÔM NAY LÀ NGÀY CUỐI HỢP ĐỒNG THỬ VIỆC!")
            else:
                st.warning(f"⚠️ **{x.get('ma_nv','')} {x['ho_ten']}** còn **{ngay_con_lai}** ngày sẽ kết thúc hợp đồng thử việc!")
    db3.close()
    
    
    if st.session_state.role in ("admin", "xem_toan_bo"):
        st.markdown("#### 💾 Sao lưu dữ liệu")
        col_bk1, col_bk2 = st.columns(2)

        with col_bk1:
            if st.button("💾 BACKUP DỮ LIỆU NGAY", width='stretch'):
                try:
                    from backup_data import backup_all
                    with st.spinner("⏳ Đang backup bảng Ứng viên, Nhân viên và hồ sơ trên Supabase Storage..."):
                        result = backup_all()

                    if result["mode"] == "local":
                        st.success(f"✅ Đã backup xong! Thư mục: {result['dest_folder']}")
                        st.cache_data.clear()
                    else:
                        st.success("✅ Đã backup xong! App đang chạy trên môi trường Cloud (không có ổ D: của bạn), "
                                    "nên kết quả được nén thành file zip — bấm nút bên dưới để tải về máy:")
                        st.cache_data.clear()
                        st.download_button(
                            label="📥 TẢI FILE BACKUP (.zip)",
                            data=result["zip_bytes"],
                            file_name=result["zip_filename"],
                            mime="application/zip",
                            width='stretch'
                        )

                    for table, res in result['db'].items():
                        if res[0]:
                            st.caption(f"✔️ Bảng `{table}`: {res[1]} dòng")
                        else:
                            st.caption(f"❌ Bảng `{table}`: {res[1]}")
                    if result['storage']['ok']:
                        st.caption(f"✔️ Storage: đã tải {result['storage']['count']} file hồ sơ")
                    else:
                        st.caption(f"❌ Storage: {result['storage']['error']} (đã tải {result['storage']['count']} file)")
                except ImportError:
                    st.error("❌ Không tìm thấy `backup_data.py`. Hãy đặt file này cùng thư mục với app.py.")
                except Exception as e:
                    st.error(f"❌ Lỗi khi backup: {e}")
            st.caption("Backup dữ liệu bảng `ung_vien`, `nhan_vien` (Excel) + toàn bộ file hồ sơ trên Supabase Storage. "
                       "Nếu chạy trên máy Windows local → lưu vào `D:\\hrm-port9\\backup`. Nếu chạy trên Cloud → tải về dạng file zip.")

        with col_bk2:
            with st.popover("🗓️ Lịch backup tự động", width='stretch'):
                is_windows = (os.name == 'nt')
                st.caption("Dùng **Windows Task Scheduler** để tự động chạy backup vào **02:00 sáng Thứ 7 hàng tuần**. "
                           "Chỉ tạo được lịch khi bấm nút này **ngay trên máy Windows** nơi bạn muốn lưu file backup — "
                           "không thể bật từ xa qua Streamlit Cloud.")
                if not is_windows:
                    st.warning("⚠️ App hiện đang chạy trên môi trường Cloud (không phải Windows), nên không thể tạo "
                               "lịch Task Scheduler tại đây. Cách làm đúng: copy file `backup_data.py` cùng file cấu "
                               "hình kết nối (.env) xuống máy Windows của bạn, rồi mở app này **chạy local** "
                               "(`streamlit run app.py` ngay trên máy đó) để bấm nút BẬT lịch — lúc đó Task Scheduler "
                               "sẽ được tạo đúng trên máy bạn và tự backup vào D:\\ hàng tuần dù sau đó bạn tắt app đi.")
                else:
                    if st.button("✅ BẬT lịch backup tự động", width='stretch'):
                        try:
                            python_exe = sys.executable
                            script_path = os.path.abspath("backup_data.py")
                            task_cmd = (
                                'schtasks /Create /TN "HRM_Port_Backup_Weekly" '
                                f'/TR "\\"{python_exe}\\" \\"{script_path}\\"" '
                                '/SC WEEKLY /D SAT /ST 02:00 /F'
                            )
                            result = subprocess.run(task_cmd, shell=True, capture_output=True, text=True)
                            if result.returncode == 0:
                                st.success("✅ Đã tạo lịch: tự động backup 02:00 sáng Thứ 7 hàng tuần.")
                            else:
                                st.error(f"❌ Không tạo được lịch: {result.stderr or result.stdout}")
                        except Exception as e:
                            st.error(f"❌ Lỗi khi tạo lịch: {e}")

                    if st.button("🗑️ TẮT lịch backup tự động", width='stretch'):
                        try:
                            result = subprocess.run(
                                'schtasks /Delete /TN "HRM_Port_Backup_Weekly" /F',
                                shell=True, capture_output=True, text=True
                            )
                            if result.returncode == 0:
                                st.success("✅ Đã tắt lịch backup tự động.")
                            else:
                                st.warning(f"⚠️ {result.stderr or result.stdout}")
                        except Exception as e:
                            st.error(f"❌ Lỗi: {e}")
    
    
    
# ========== ỨNG VIÊN ==========
elif menu == "👤 Ứng viên":
    st.markdown(f"# {i18n.tm('👤 Ứng viên')}", unsafe_allow_html=True)
    ensure_chuc_danh_ung_vien_table()
    su = st.text_input("🔍 Tìm kiếm", key="suv")
    
    # Kiểm tra nếu đang chuyển từ ứng viên sang nhân viên
    if 'show_chuyen_nv_form' in st.session_state and st.session_state.show_chuyen_nv_form:
        st.subheader("📝 CHUYỂN ỨNG VIÊN THÀNH NHÂN VIÊN")
        uv_data = st.session_state.get('chuyen_uv_data', {})
        
        # Lấy danh sách chức danh từ database
        db_chuc = st.session_state.db_engine.get_connection()
        c_chuc = db_chuc.cursor()
        c_chuc.execute("SELECT DISTINCT ten_vi_tri FROM vi_tri_cong_tac ORDER BY ten_vi_tri")
        dschucdanh = [row[0] for row in c_chuc.fetchall()]
        db_chuc.close()
        dpb_chuyen = get_phong_ban_options()
        
        _col_lhd_c1, _col_lhd_c2 = st.columns(2)
        with _col_lhd_c1:
            ngay_vao_lam_chuyen = st.text_input("Ngày vào làm (dd/mm/yyyy) *", value=format_date(uv_data.get('ngay_vao_lam', date.today())), placeholder="dd/mm/yyyy", max_chars=10, key="nvl_chuyen_uv")
        with _col_lhd_c2:
            loai_hd_chuyen = st.selectbox("Loại HĐ *", ["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"], key="loai_hd_chuyen_uv")
        
        with st.form("chuyen_uv_to_nv_form"):
            st.markdown(f"**Ứng viên:** {uv_data.get('ho_ten', '')}")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                ho_ten_nv = st.text_input("Họ và tên *", value=uv_data.get('ho_ten', ''))
                ngay_sinh_nv = st.text_input("Ngày sinh (dd/mm/yyyy)", value=format_date(uv_data.get('ngay_sinh')), placeholder="dd/mm/yyyy", max_chars=10)
                gioi_tinh_nv = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"], index=["", "Nam", "Nữ", "Khác"].index(uv_data.get('gioi_tinh', '')) if uv_data.get('gioi_tinh') in ["Nam", "Nữ", "Khác"] else 0)
                so_cccd_nv = st.text_input("CCCD")
                ngay_cap_cccd_nv = st.text_input("Ngày cấp CCCD (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                noi_cap_cccd_nv = st.text_input("Nơi cấp CCCD", value=get_cau_hinh('noi_cap_cccd', 'Cục QLHC về TTXH - Bộ Công An'))
            with col2:
                # Nguyên quán: đã bỏ khỏi UI theo yêu cầu (đồng bộ với 2 form Thêm/Sửa nhân viên),
                # lưu rỗng — có thể bổ sung sau qua màn "Sửa nhân viên" nếu cần.
                nguyen_quan_nv = ""
                thuong_tru_nv = st.text_area("Thường trú", value=uv_data.get('ghi_chu', ''), height=68)
                quoc_tich_nv = st.text_input("Quốc tịch", value="Việt Nam")
                dan_toc_nv = st.text_input("Dân tộc", value="Kinh")
                dien_thoai_nv = st.text_input("SĐT", value=uv_data.get('dien_thoai', ''))
                email_nv = st.text_input("Email")
                trinh_do_nv = st.selectbox("Trình độ", [""] + TRINH_DO_LIST)
            with col3:
                chuc_danh_nv = st.selectbox("Chức danh", [""] + dschucdanh, index=([""] + dschucdanh).index(uv_data.get('vi_tri', '')) if uv_data.get('vi_tri', '') in dschucdanh else 0)
                phong_ban_nv = st.selectbox("Phòng ban", [""] + dpb_chuyen, key="pb_chuyen_uv")
                noi_lam_viec_nv = st.text_input("Nơi làm việc", value=get_cau_hinh('noi_lam_viec', 'Cảng THQT Hòn La'))
                anh_ho_so_nv = st.file_uploader("Ảnh hồ sơ", type=["png", "jpg", "jpeg"], key="anh_ho_so_chuyen")
            
            st.divider()
            st.caption("💼 Bảo hiểm xã hội")
            col4, col5, col6 = st.columns(3)
            la_thu_viec_chuyen = loai_hd_chuyen == "Thử việc"
            with col4:
                ma_bhxh_chuyen = st.text_input("Mã BHXH", disabled=la_thu_viec_chuyen)
                luong_bh_chuyen = st.text_input("Lương BH")
                he_so_luong_chuyen = st.text_input("Hệ số lương")
                pc_chuc_vu_chuyen = st.text_input("PC chức vụ")
            with col5:
                pc_tnvk_chuyen = st.text_input("PC TNVK (%)")
                pc_tnn_chuyen = st.text_input("PC TNN (%)")
                muc_huong_bhyt_chuyen = st.selectbox("Mức hưởng BHYT", ["80%", "95%", "100%"])
                ty_le_dong_chuyen = st.text_input("Tỷ lệ đóng (%)")
            with col6:
                muc_tien_dong_chuyen = st.text_input("Mức tiền đóng")
                phuong_thuc_dong_chuyen = st.selectbox("PT đóng", ["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"])
                nhom_bhxh_chuyen = st.selectbox("Nhóm BHXH", ["", "Văn phòng", "Lao động trực tiếp"])
                phuong_an_chuyen = st.selectbox("Phương án điều chỉnh", [""] + PHUONG_AN_TANG, key="pa_chuyen", disabled=la_thu_viec_chuyen)
            
            st.divider()
            st.caption("🏦 Ngân hàng & KCB")
            col7, col8 = st.columns(2)
            with col7:
                stk_chuyen = st.text_input("STK")
                # Tạo dropdown cho chi nhánh ngân hàng
                bank_chuyen_index = 0
                chi_nhanh_nh_chuyen = st.selectbox("Chi nhánh NH", options=[""] + BANK_LIST, index=bank_chuyen_index, key="chuyen_cnh")
            with col8:
                ho_so_chuyen = st.selectbox("Hồ sơ", ["", "Đã có HS", "Chưa có"])
                so_luong_npt_nv = st.number_input("Số người phụ thuộc", min_value=0, value=0, step=1)
            # Các trường ít dùng (Tỉnh KCB, Nơi KCB, Tỉnh/TP nhận HS, Phường/Xã nhận HS,
            # Địa chỉ nhận HS, ĐK nhận sổ) đã bỏ khỏi UI theo yêu cầu — đồng bộ với 2 form
            # Thêm/Sửa nhân viên: tự động lấy theo cấu hình chung của công ty (⚙️ Cấu hình
            # công ty); có thể chỉnh riêng cho từng người qua màn "Sửa nhân viên" nếu cần.
            tinh_kcb_chuyen = get_cau_hinh('tinh_kcb', 'Tỉnh Quảng Trị')
            noi_kcb_chuyen = get_cau_hinh('noi_dang_ky_kcb', 'Bệnh viện đa khoa khu vực Bắc Quảng Trị')
            tinh_nhan_hs_chuyen = get_cau_hinh('tinh_nhan_hs', 'Tỉnh Quảng Trị')
            phuong_nhan_hs_chuyen = "Xã Phú Trạch"
            dia_chi_nhan_hs_chuyen = get_cau_hinh('dia_chi_nhan_hs', 'Công ty cổ phần Cảng Hòn La')
            dk_nhan_so_chuyen = "Có"
            col_confirm1, col_confirm2 = st.columns(2)
            with col_confirm1:
                if st.form_submit_button("✅ XÁC NHẬN CHUYỂN", width='stretch', type="primary", disabled=not can_edit()):
                    if ho_ten_nv:
                        # Kiểm tra định dạng ngày
                        ngay_loi = []
                        if ngay_sinh_nv and not parse_date(ngay_sinh_nv): 
                            ngay_loi.append("Ngày sinh")
                        if ngay_cap_cccd_nv and not parse_date(ngay_cap_cccd_nv): 
                            ngay_loi.append("Ngày cấp CCCD")
                        if ngay_vao_lam_chuyen and not parse_date(ngay_vao_lam_chuyen):
                            ngay_loi.append("Ngày vào làm")
                        if ngay_loi:
                            st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                        else:
                            try:
                                # Tạo ten_don_vi_thu_huong từ ho_ten
                                ten_don_vi_thu_huong = generate_ten_don_vi_thu_huong(ho_ten_nv)
                                
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                
                                c.execute("SELECT COALESCE(MAX(STT), 0) + 1 FROM nhan_vien")
                                stt_moi = c.fetchone()[0]

                                # Tạo STT và mã nhân viên mới (theo đúng ký hiệu riêng của tenant)
                                ma_nv = sinh_ma_nv_moi(c)
                                c.execute("SELECT COALESCE(MAX(STT),0)+1 FROM nhan_vien")
                                                                
                                nhl = parse_date(ngay_vao_lam_chuyen) or date.today()

                                # Mã công ty của TENANT ĐANG ĐĂNG NHẬP (không phải luôn là "CHL" của Hòn La)
                                ma_cty_hd = st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL'

                                # Tạo số hợp đồng theo loại + tự tính tháng bắt đầu BH
                                if loai_hd_chuyen == "Thử việc":
                                    trang_thai_nv = 'THU_VIEC'
                                    trang_thai_bhxh = 'CHUA_DONG'
                                    tbd_val = None
                                    pa_val = None
                                    so_hd = sinh_so_hdld_moi(c, ma_cty_hd, nhl.year, la_thu_viec=True)
                                else:
                                    trang_thai_nv = 'DANG_LAM'
                                    trang_thai_bhxh = 'DANG_DONG'
                                    tbd_val = tinh_thang_bat_dau_bh(nhl)
                                    pa_val = lay_ma_phuong_an(phuong_an_chuyen)
                                    so_hd = sinh_so_hdld_moi(c, ma_cty_hd, nhl.year, la_thu_viec=False)
                                
                                # Thêm nhân viên mới (đã thêm trường ten_don_vi_thu_huong, trinh_do)
                                c.execute("""
                                    INSERT INTO nhan_vien (STT, ma_nv, so_hdld, ho_ten, chuc_danh_nghe, 
                                        ngay_sinh, gioi_tinh, so_cccd, ngay_cap_cccd, noi_cap_cccd,
                                        nguyen_quan, thuong_tru, dien_thoai, email, email_lien_he, ho_so,
                                        luong_bao_hiem, ma_so_bhxh, ngay_vao_lam, noi_lam_viec,
                                        so_tai_khoan_nh, chi_nhanh_nh, ngay_ky_hd, loai_hop_dong,
                                        nhom_bhxh, thang_bat_dau_bh, thang_ket_thuc_bh, trang_thai, trang_thai_bhxh,
                                        phong_ban_lam_viec, ngay_ket_thuc, quoc_tich, dan_toc, 
                                        he_so_luong, phu_cap_chuc_vu, phu_cap_tnvk, phu_cap_tnn,
                                        muc_huong_bhyt, ty_le_dong, muc_tien_dong, phuong_thuc_dong,
                                        tinh_nhan_hs, phuong_nhan_hs, dia_chi_nhan_hs, 
                                        tinh_kcb, noi_dang_ky_kcb, dang_ky_nhan_so, ten_don_vi_thu_huong, so_luong_npt, trinh_do,
                                        phuong_an_dieu_chinh, thang_phuong_an)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                                        %s, %s)
                                    RETURNING id
                                """, (
                                    stt_moi, ma_nv, so_hd, ho_ten_nv, chuc_danh_nv,
                                    parse_date(ngay_sinh_nv), gioi_tinh_nv, so_cccd_nv, parse_date(ngay_cap_cccd_nv), noi_cap_cccd_nv,
                                    nguyen_quan_nv, thuong_tru_nv, (dien_thoai_nv.strip() or None) if dien_thoai_nv else None, email_nv, email_nv, ho_so_chuyen,
                                    luong_bh_chuyen, ma_bhxh_chuyen, nhl, noi_lam_viec_nv,
                                    stk_chuyen, chi_nhanh_nh_chuyen, nhl, loai_hd_chuyen,
                                    nhom_bhxh_chuyen, tbd_val, tinh_ngay_ket_thuc(loai_hd_chuyen, nhl), trang_thai_nv, trang_thai_bhxh,
                                    phong_ban_nv, tinh_ngay_ket_thuc(loai_hd_chuyen, nhl), quoc_tich_nv, dan_toc_nv,
                                    to_float_or_none(he_so_luong_chuyen), to_float_or_none(pc_chuc_vu_chuyen),
                                    to_float_or_none(pc_tnvk_chuyen), to_float_or_none(pc_tnn_chuyen),
                                    muc_huong_bhyt_chuyen, to_float_or_none(ty_le_dong_chuyen), to_float_or_none(muc_tien_dong_chuyen),
                                    phuong_thuc_dong_chuyen, tinh_nhan_hs_chuyen, phuong_nhan_hs_chuyen, dia_chi_nhan_hs_chuyen,
                                    tinh_kcb_chuyen, noi_kcb_chuyen, dk_nhan_so_chuyen, ten_don_vi_thu_huong, so_luong_npt_nv, trinh_do_nv,
                                    pa_val, format_thang_nam(tbd_val)
                                ))
                                nhan_vien_id_moi = c.fetchone()[0]

                                # Upload ảnh hồ sơ (nếu có) — cần id vừa tạo để đặt tên thư mục trên Storage
                                if anh_ho_so_nv is not None:
                                    storage_path_anh = upload_anh_ho_so(ma_nv, ho_ten_nv, anh_ho_so_nv)
                                    if storage_path_anh:
                                        c.execute("UPDATE nhan_vien SET anh_ho_so=%s WHERE id=%s", (storage_path_anh, nhan_vien_id_moi))

                                # Cập nhật trạng thái ứng viên
                                c.execute("UPDATE ung_vien SET trang_thai='DA_NHAN_VIEC', ma_nv=%s WHERE id=%s", 
                                         (ma_nv, st.session_state['chuyen_uv_id']))
                                
                                # Thêm lịch sử công tác
                                c.execute("""
                                    INSERT INTO lich_su_cong_tac (nhan_vien_id, tu_ngay, chuc_danh, phong_ban, noi_lam_viec, loai_hop_dong, he_so_luong, so_hop_dong)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                                """, (nhan_vien_id_moi, ngay_vao_lam_chuyen, chuc_danh_nv, phong_ban_nv, noi_lam_viec_nv, loai_hd_chuyen, 
                                      to_float_or_none(he_so_luong_chuyen), so_hd))
                                
                                db.commit()
                                db.close()
                                
                                st.success(f"✅ Đã chuyển {ho_ten_nv} thành nhân viên! Mã NV: {ma_nv}")
                                st.cache_data.clear()
                                # Xóa session state
                                del st.session_state['show_chuyen_nv_form']
                                del st.session_state['chuyen_uv_id']
                                del st.session_state['chuyen_uv_data']
                                st.rerun()
                                
                            except Exception as e:
                                db.rollback()
                                db.close()
                                st.error(f"❌ Lỗi khi chuyển: {e}")
                    else:
                        st.error("Họ tên không được để trống!")
            
            with col_confirm2:
                if st.form_submit_button("❌ HỦY", width='stretch'):
                    del st.session_state['show_chuyen_nv_form']
                    del st.session_state['chuyen_uv_id']
                    del st.session_state['chuyen_uv_data']
                    st.rerun()
        
        st.divider()
        st.stop()  # Dừng lại để không hiển thị danh sách ứng viên phía dưới
    
    db_f = st.session_state.db_engine.get_connection()
    c_f = db_f.cursor()
    c_f.execute("SELECT ten_chuc_danh FROM chuc_danh_ung_vien ORDER BY ten_chuc_danh")
    ds_vi_tri = [row[0] for row in c_f.fetchall()]
    c_f.execute("SELECT DISTINCT vi_tri_du_tuyen FROM ung_vien WHERE vi_tri_du_tuyen IS NOT NULL AND vi_tri_du_tuyen != '' ORDER BY vi_tri_du_tuyen")
    for row in c_f.fetchall():
        if row[0] not in ds_vi_tri:
            ds_vi_tri.append(row[0])
    db_f.close()
    
    col_f1, col_f2 = st.columns([2, 1])
    with col_f1:
        filter_vi_tri = st.selectbox("🔍 Lọc Vị trí dự tuyển:", ["Tất cả"] + ds_vi_tri)
    
    # Chỉ admin mới thấy nút thêm ứng viên
    if st.session_state.role in ("admin", "xem_toan_bo"):
        with st.expander("➕ THÊM ỨNG VIÊN MỚI", expanded=False):
            with st.form("add_uv_form"):
                db_f = st.session_state.db_engine.get_connection()
                c_f = db_f.cursor()
                c_f.execute("SELECT ten_chuc_danh FROM chuc_danh_ung_vien ORDER BY ten_chuc_danh")
                ds_vt_uv = [row[0] for row in c_f.fetchall()]
                db_f.close()
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    ho_ten_uv = st.text_input("Họ và tên *")
                    vi_tri_uv = st.selectbox("Vị trí dự tuyển", [""] + ds_vt_uv)
                    dien_thoai_uv = st.text_input("SĐT")
                with col2:
                    ngay_sinh_uv = st.text_input("Ngày sinh (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                    gioi_tinh_uv = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"])
                with col3:
                    ngay_vao_lam_uv = st.text_input("Ngày vào làm (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                    ghi_chu_uv = st.text_area("Ghi chú")
                
                if st.form_submit_button("💾 LƯU", width='stretch', disabled=not can_edit()):
                    if not can_edit():
                        st.error("❌ Bạn không có quyền thực hiện thao tác này!")
                    else:
                        if ho_ten_uv:
                            ngay_loi = []
                            if ngay_sinh_uv and not parse_date(ngay_sinh_uv): 
                                ngay_loi.append("Ngày sinh")
                            if ngay_vao_lam_uv and not parse_date(ngay_vao_lam_uv): 
                                ngay_loi.append("Ngày vào làm")
                            if ngay_loi:
                                st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                            else:
                                try:
                                    db = st.session_state.db_engine.get_connection()
                                    c = db.cursor()
                                    c.execute("""INSERT INTO ung_vien (ho_ten, vi_tri_du_tuyen, dien_thoai, 
                                        ngay_sinh, gioi_tinh, ngay_vao_lam, luong_bao_hiem, trang_thai, ghi_chu)
                                        VALUES (%s, %s, %s, %s, %s, %s, %s, 'CHO_DUYET', %s) RETURNING id""",
                                        (ho_ten_uv, vi_tri_uv, dien_thoai_uv, parse_date(ngay_sinh_uv),
                                         gioi_tinh_uv, parse_date(ngay_vao_lam_uv), None, ghi_chu_uv or None))
                                    new_id = c.fetchone()[0]
                                    ma_uv = f"UV{new_id:04d}"
                                    c.execute("UPDATE ung_vien SET ma_uv = %s WHERE id = %s", (ma_uv, new_id))
                                    db.commit()
                                    db.close()
                                    st.success(f"✅ Đã thêm ứng viên: {ho_ten_uv} (Mã: {ma_uv})")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Lỗi khi thêm ứng viên: {e}")
                        else:
                            st.error("Họ tên không được để trống!")
    
    st.divider()
    
    t1, t2, t3, t4 = st.tabs(["📋 Tất cả", "⏳ Chờ duyệt", "✅ Đã nhận", "❌ Từ chối"])
    tm = {"📋 Tất cả": "", "⏳ Chờ duyệt": "CHO_DUYET", "✅ Đã nhận": "DA_NHAN_VIEC", "❌ Từ chối": "TU_CHOI"}
    
    for tn, tab in zip(tm.keys(), [t1, t2, t3, t4]):
        with tab:
            tt = tm[tn]
            db = st.session_state.db_engine.get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            sql = "SELECT id, ma_uv, ho_ten, vi_tri_du_tuyen, dien_thoai, ngay_vao_lam, luong_bao_hiem, ngay_sinh, trang_thai FROM ung_vien WHERE 1=1"
            params = []
            if tt:
                sql += " AND trang_thai = %s"
                params.append(tt)
            if su:
                sql += " AND (ho_ten LIKE %s OR dien_thoai LIKE %s OR ma_uv LIKE %s)"
                params.extend([f'%{su}%', f'%{su}%', f'%{su}%'])
            if filter_vi_tri != "Tất cả":
                sql += " AND vi_tri_du_tuyen = %s"
                params.append(filter_vi_tri)
            sql += " ORDER BY id ASC"
            c.execute(sql, tuple(params))
            ds = c.fetchall()
            db.close()
            
            if ds:
                df = pd.DataFrame(ds)
                for col in df.columns:
                    if 'ngay' in col.lower():
                        df[col] = df[col].apply(format_date)
                
                display_cols = ['ma_uv', 'ho_ten', 'vi_tri_du_tuyen', 'dien_thoai', 'ngay_vao_lam', 'luong_bao_hiem', 'ngay_sinh', 'trang_thai']
                available_cols = [c for c in display_cols if c in df.columns]
                df_show = df[available_cols]
                
                col_map = {
                    'ma_uv': 'Mã UV',
                    'ho_ten': 'Họ tên',
                    'vi_tri_du_tuyen': 'Vị trí dự tuyển',
                    'dien_thoai': 'SĐT',
                    'ngay_vao_lam': 'Ngày vào làm',
                    'luong_bao_hiem': 'Ghi chú',
                    'ngay_sinh': 'Ngày sinh',
                    'trang_thai': 'Trạng thái',
                }
                df_show.rename(columns=col_map, inplace=True)
                
                st.caption(f"📌 {len(ds)} kết quả.")
                
                if st.session_state.role in ("admin", "xem_toan_bo"):
                    # Admin: hiển thị bảng có checkbox và nút chức năng
                    if 'selected' not in df.columns:
                        df.insert(0, 'selected', False)
                    df_show_with_checkbox = df[['selected'] + [c for c in df.columns if c in display_cols]]
                    df_show_with_checkbox.rename(columns={'selected': 'Chọn'}, inplace=True)
                    
                    edited_df = st.data_editor(
                        df_show_with_checkbox,
                        column_config={"Chọn": st.column_config.CheckboxColumn("Chọn", default=False)},
                        disabled=[col for col in df_show_with_checkbox.columns if col != 'Chọn'],
                        hide_index=True,
                        height=400,
                        key=f"uv_editor_{tn}"
                    )
                    
                    if edited_df is not None:
                        selected_rows = edited_df[edited_df['Chọn'] == True]
                        if len(selected_rows) > 1:
                            st.error("⚠️ Chỉ được chọn 1 ứng viên!")
                        elif len(selected_rows) == 1:
                            selected_idx = selected_rows.index[0]
                            selected_nv = df.iloc[selected_idx]
                            col_btn1, col_btn2 = st.columns(2)
                            with col_btn1:
                                # Chỉ hiển thị nút Sửa khi đã chọn 1 ứng viên
                                if st.button(f"✏️ SỬA", key=f"edit_sel_{tn}", width='stretch'):
                                    st.session_state['edit_uv_id'] = int(selected_nv['id'])
                                    st.rerun()
                            if tn == "⏳ Chờ duyệt":
                                with col_btn2:
                                    if st.button(f"👥 CHUYỂN SANG NHÂN VIÊN", type="primary", key=f"chuyen_uv_{tn}"):
                                        try:
                                            db = st.session_state.db_engine.get_connection()
                                            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                                            uv_id = int(selected_nv['id'])
                                            c.execute("SELECT * FROM ung_vien WHERE id = %s", (uv_id,))
                                            uv = c.fetchone()
                                            db.close()
                                            
                                            if uv:
                                                # Lưu thông tin ứng viên vào session_state
                                                st.session_state['chuyen_uv_id'] = uv_id
                                                st.session_state['chuyen_uv_data'] = {
                                                    'ho_ten': uv['ho_ten'],
                                                    'vi_tri': uv['vi_tri_du_tuyen'],
                                                    'dien_thoai': uv['dien_thoai'],
                                                    'ngay_sinh': uv['ngay_sinh'],
                                                    'gioi_tinh': uv['gioi_tinh'],
                                                    'ngay_vao_lam': uv['ngay_vao_lam'] or date.today(),
                                                    'ghi_chu': uv['luong_bao_hiem']
                                                }
                                                st.session_state['show_chuyen_nv_form'] = True
                                                st.rerun()
                                        except Exception as e:
                                            st.error(f"❌ Lỗi: {e}")
                else:
                    # Viewer: chỉ hiển thị bảng
                    st.dataframe(df_show, width='stretch', hide_index=True, height=400)
            else:
                st.info("Không có dữ liệu")
    
    # Form sửa ứng viên (chỉ admin)
    if st.session_state.get('edit_uv_id') and st.session_state.role in ("admin", "xem_toan_bo"):
        st.divider()
        st.subheader(f"✏️ Sửa ứng viên")
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT * FROM ung_vien WHERE id = %s", (st.session_state['edit_uv_id'],))
        uv_data = c.fetchone()
        db.close()
        if uv_data:
            with st.form("edit_uv_direct"):
                col1, col2, col3 = st.columns(3)
                with col1:
                    ho_ten_e = st.text_input("Họ và tên *", value=uv_data['ho_ten'] or '')
                    vi_tri_e = st.selectbox("Vị trí dự tuyển", [""] + ds_vi_tri,
                        index=([""] + ds_vi_tri).index(uv_data['vi_tri_du_tuyen']) if uv_data['vi_tri_du_tuyen'] in ds_vi_tri else 0)
                    dien_thoai_e = st.text_input("SĐT", value=uv_data['dien_thoai'] or '')
                with col2:
                    ngay_sinh_e = st.text_input("Ngày sinh (dd/mm/yyyy)", value=format_date(uv_data['ngay_sinh']))
                    gioi_tinh_e = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"],
                        index=["", "Nam", "Nữ", "Khác"].index(uv_data['gioi_tinh']) if uv_data['gioi_tinh'] in ["Nam", "Nữ", "Khác"] else 0)
                with col3:
                    ngay_vao_lam_e = st.text_input("Ngày vào làm (dd/mm/yyyy)", value=format_date(uv_data['ngay_vao_lam']))
                    ghi_chu_e = st.text_area("Ghi chú", value=uv_data.get('ghi_chu') or '')
                    trang_thai_e = st.selectbox("Trạng thái", ["CHO_DUYET", "TU_CHOI", "DA_NHAN_VIEC"],
                        index=["CHO_DUYET", "TU_CHOI", "DA_NHAN_VIEC"].index(uv_data['trang_thai']) if uv_data['trang_thai'] in ["CHO_DUYET", "TU_CHOI", "DA_NHAN_VIEC"] else 0)
                
                col_save, col_del, col_cancel = st.columns(3)
                with col_save:
                    if st.form_submit_button("💾 CẬP NHẬT", disabled=not can_edit()):
                        if not can_edit():
                            st.error("❌ Bạn không có quyền thực hiện thao tác này!")
                        else:
                            ngay_loi = []
                            if ngay_sinh_e and not parse_date(ngay_sinh_e): 
                                ngay_loi.append("Ngày sinh")
                            if ngay_vao_lam_e and not parse_date(ngay_vao_lam_e): 
                                ngay_loi.append("Ngày vào làm")
                            if ngay_loi:
                                st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                            else:
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("""UPDATE ung_vien SET ho_ten=%s, vi_tri_du_tuyen=%s, dien_thoai=%s,
                                    ngay_sinh=%s, gioi_tinh=%s, ngay_vao_lam=%s, trang_thai=%s, ghi_chu=%s
                                    WHERE id=%s""",
                                    (ho_ten_e, vi_tri_e, dien_thoai_e, parse_date(ngay_sinh_e), gioi_tinh_e,
                                     parse_date(ngay_vao_lam_e), trang_thai_e, ghi_chu_e or None, uv_data['id']))
                                db.commit()
                                db.close()
                                st.success("✅ Đã cập nhật!")
                                st.cache_data.clear()
                                del st.session_state['edit_uv_id']
                                st.rerun()

                with col_del:
                    if st.form_submit_button("🗑️ XÓA", disabled=not can_edit()):
                        if not can_delete():
                            st.error("❌ Bạn không có quyền xóa dữ liệu!")
                        else:
                            db = st.session_state.db_engine.get_connection()
                            c = db.cursor()
                            c.execute("DELETE FROM ung_vien WHERE id = %s", (uv_data['id'],))
                            db.commit()
                            db.close()
                            st.success("🗑️ Đã xóa!")
                            st.cache_data.clear()
                            del st.session_state['edit_uv_id']
                            st.rerun()
                with col_cancel:
                    if st.form_submit_button("❌ HỦY"):
                        del st.session_state['edit_uv_id']
                        st.rerun()
    
    # Quản lý danh mục vị trí dự tuyển (chỉ admin) - bảng RIÊNG chuc_danh_ung_vien,
    # độc lập với danh mục chức danh Nhân viên (vi_tri_cong_tac)
    if st.session_state.role in ("admin", "xem_toan_bo"):
        st.divider()
        with st.expander("⚙️ Quản lý danh mục Vị trí dự tuyển (riêng cho Ứng viên)", expanded=False):
            st.caption("Danh mục này độc lập với danh mục Chức danh của Nhân viên — "
                       "đổi chức danh Nhân viên sẽ không ảnh hưởng đến danh mục và dữ liệu Ứng viên.")
            with st.form("add_vi_tri_uv"):
                ten_vt_moi = st.text_input("Tên vị trí dự tuyển mới *")
                if st.form_submit_button("➕ Thêm", disabled=not can_edit()):
                    if ten_vt_moi:
                        db = st.session_state.db_engine.get_connection()
                        c = db.cursor()
                        c.execute("SELECT COUNT(*) FROM chuc_danh_ung_vien WHERE ten_chuc_danh = %s", (ten_vt_moi,))
                        if c.fetchone()[0] == 0:
                            c.execute("INSERT INTO chuc_danh_ung_vien (ten_chuc_danh) VALUES (%s)", (ten_vt_moi,))
                            db.commit()
                            st.success(f"✅ Đã thêm: {ten_vt_moi}")
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.warning("Vị trí này đã tồn tại!")
                        db.close()
                    else:
                        st.error("Tên không được để trống!")
            
            db = st.session_state.db_engine.get_connection()
            c = db.cursor()
            c.execute("SELECT id, ten_chuc_danh FROM chuc_danh_ung_vien ORDER BY ten_chuc_danh")
            ds_vt = c.fetchall()
            db.close()
            if ds_vt:
                st.caption("📋 Danh sách vị trí dự tuyển:")
                for row in ds_vt:
                    col_ten, col_xoa = st.columns([4, 1])
                    with col_ten:
                        st.write(f"- {row[1]}")
                    with col_xoa:
                        if st.button("🗑️", key=f"xoa_cdv_{row[0]}"):
                            db = st.session_state.db_engine.get_connection()
                            c = db.cursor()
                            c.execute("DELETE FROM chuc_danh_ung_vien WHERE id=%s", (row[0],))
                            db.commit(); db.close()
                            st.success("🗑️ Đã xóa!")
                            st.cache_data.clear()
                            st.rerun()

# ========== NHÂN VIÊN ==========
elif menu == "✅ Nhân viên":
    st.markdown(f"# {i18n.tm('✅ Quản lý nhân viên')}", unsafe_allow_html=True)
    ensure_qdns_columns()
    ensure_qdns_table()
    ensure_mau_dieu_hop_dong_table()

    tab_dang_lam, tab_da_nghi, tab_qtct, tab_qdns, tab_co_cau = st.tabs(["📌 ĐANG LÀM VIỆC", "📋 ĐÃ NGHỈ VIỆC", "📜 LỊCH SỬ CÔNG TÁC", "📜 QUYẾT ĐỊNH NHÂN SỰ", "🏢 CƠ CẤU PHÒNG BAN"])
    
    with tab_dang_lam:
        #st.caption("👥 Danh sách nhân viên đang làm việc (bao gồm thử việc)")
        # Xử lý yêu cầu reset ô tìm kiếm (đến từ nút "Đóng" của card thông tin nhân viên)
        # Phải làm TRƯỚC khi widget text_input được khởi tạo, nếu không Streamlit sẽ báo lỗi
        if st.session_state.pop('_reset_snv_dang_lam', False):
            st.session_state['snv_dang_lam'] = ''
        if st.session_state.role in ("admin", "xem_toan_bo"):
            _chedo_nv = st.radio("Chọn thao tác:", ["➕ Thêm NV mới", "🔍 Danh sách nhân viên đang làm việc (bao gồm thử việc)"],
                                 horizontal=True, key="nv_dang_lam_thao_tac")
        else:
            _chedo_nv = "🔍 Danh sách nhân viên đang làm việc (bao gồm thử việc)"

        if _chedo_nv == "➕ Thêm NV mới" and st.session_state.role in ("admin", "xem_toan_bo"):
            st.session_state.setdefault('add_nv_reset_ctr', 0)
            with st.expander("➕ THÊM NHÂN VIÊN MỚI", expanded=True, key=f"add_nv_expander_{st.session_state.add_nv_reset_ctr}"):
                _col_lhd1, _col_lhd2 = st.columns(2)
                with _col_lhd1:
                    nvl = st.text_input("Ngày vào làm (dd/mm/yyyy) *", placeholder="dd/mm/yyyy", max_chars=10, key="nvl")
                with _col_lhd2:
                    lhd = st.selectbox("Loại HĐ *", ["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"], key="lhd")
                with st.form(f"add_nv_{st.session_state.add_nv_reset_ctr}"):
                    st.markdown("**Nhập thông tin nhân viên mới**")
                    db = st.session_state.db_engine.get_connection()
                    c = db.cursor()
                    c.execute("SELECT DISTINCT ten_vi_tri FROM vi_tri_cong_tac ORDER BY ten_vi_tri")
                    dcv = [row[0] for row in c.fetchall()]
                    db.close()
                    dpb = get_phong_ban_options()
                    st.caption("📝 Thông tin cá nhân")
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        htn = st.text_input("Họ và tên *", key="htn")
                        nsn = st.text_input("Ngày sinh (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10, key="nsn")
                        gtn = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"], key="gtn")
                        scc = st.text_input("CCCD", key="scc")
                        ncc = st.text_input("Ngày cấp CCCD (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10, key="ncc")
                        ncc2 = st.text_input("Nơi cấp CCCD", value=get_cau_hinh('noi_cap_cccd', 'Cục QLHC về TTXH - Bộ Công An'), key="ncc2")
                    with c2:
                        nqn = ""  # Nguyên quán: đã bỏ khỏi UI theo yêu cầu, lưu rỗng (có thể bổ sung sau qua Sửa nhân viên nếu cần)
                        ttn = st.text_input("Thường trú", key="ttn")
                        qtn = st.text_input("Quốc tịch", value="Việt Nam", key="qtn")
                        dtn = st.text_input("Dân tộc", value="Kinh", key="dtn")
                        dtn2 = st.text_input("SĐT", key="dtn2")
                        emn = st.text_input("Email", key="emn")
                        trinh_do_moi = st.selectbox("Trình độ", [""] + TRINH_DO_LIST, key="trinh_do_add")
                    with c3:
                        cdn = st.selectbox("Chức danh", [""] + dcv, key="cdn")
                        pbn = st.selectbox("Phòng ban", [""] + dpb, key="pbn")
                        pbn_chuan = chuan_hoa_ten_phong_ban(pbn)
                        nlv = get_cau_hinh('noi_lam_viec', 'Cảng THQT Hòn La')  # Nơi làm việc: đã bỏ khỏi UI, dùng cấu hình chung của công ty
                        anh_ho_so_moi = st.file_uploader("Ảnh hồ sơ", type=["png", "jpg", "jpeg"], key="anh_ho_so_add")
                    st.divider()
                    st.caption("💼 Bảo hiểm xã hội")
                    la_thu_viec_add = lhd == "Thử việc"
                    c4, c5, c6 = st.columns(3)
                    with c4:
                        mbh = st.text_input("Mã BHXH", key="mbh", disabled=la_thu_viec_add)
                        lbh = st.text_input("Lương BH", key="lbh")
                        hsl = st.text_input("Hệ số lương", key="hsl")
                        pcv = st.text_input("PC chức vụ", key="pcv")
                    with c5:
                        ptv = st.text_input("PC TNVK (%)", key="ptv")
                        ptn = st.text_input("PC TNN (%)", key="ptn")
                        mhb = st.selectbox("Mức hưởng BHYT", ["80%", "95%", "100%"], key="mhb")
                        tld = st.text_input("Tỷ lệ đóng (%)", key="tld")
                    with c6:
                        mtd = st.text_input("Mức tiền đóng", key="mtd")
                        ptd = st.selectbox("PT đóng", ["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"], key="ptd")
                        nbh = st.selectbox("Nhóm BHXH", ["", "Văn phòng", "Lao động trực tiếp"], key="nbh")
                        pa_add = st.selectbox("Phương án điều chỉnh", [""] + PHUONG_AN_TANG, key="pa_add", disabled=la_thu_viec_add)
                    st.divider()
                    st.caption("🏦 Ngân hàng & Hồ sơ")
                    c7, c8 = st.columns(2)
                    with c7:
                        stk = st.text_input("STK", key="stk")
                        bank_index = 0
                        cnh = st.selectbox("Chi nhánh NH", options=[""] + BANK_LIST, index=bank_index, key="add_cnh")
                    with c8:
                        hso = st.selectbox("Hồ sơ", ["", "Đã có HS", "Chưa có"], key="hso")
                        so_luong_npt = st.number_input("Số người phụ thuộc", min_value=0, value=0, step=1, key="so_luong_npt_add")
                    # Các trường ít dùng (Tỉnh KCB, Nơi KCB, Tỉnh/TP nhận HS, Phường/Xã nhận HS,
                    # Địa chỉ nhận HS, ĐK nhận sổ) đã bỏ khỏi UI theo yêu cầu — tự động lấy theo
                    # cấu hình chung của công ty (⚙️ Cấu hình công ty); có thể chỉnh riêng cho
                    # từng người qua màn "Sửa nhân viên" nếu cần khác với mặc định.
                    tkb = get_cau_hinh('tinh_kcb', 'Tỉnh Quảng Trị')
                    nkb = get_cau_hinh('noi_dang_ky_kcb', 'Bệnh viện đa khoa khu vực Bắc Quảng Trị')
                    ths = get_cau_hinh('tinh_nhan_hs', 'Tỉnh Quảng Trị')
                    phs = "Xã Phú Trạch"
                    dhs = get_cau_hinh('dia_chi_nhan_hs', 'Công ty cổ phần Cảng Hòn La')
                    dks = "Có"
                    
                    col_save_exit1, col_save_exit2 = st.columns(2)
                    with col_save_exit1:
                        if st.form_submit_button("💾 LƯU", width='stretch', disabled=not can_edit()):
                            if not can_edit():
                                st.error("❌ Bạn không có quyền thực hiện thao tác này!")
                            else:
                                if htn:
                                    ngay_loi = []
                                    if nsn and not parse_date(nsn):
                                        ngay_loi.append("Ngày sinh")
                                    # ... (các kiểm tra ngày khác)
                                    if ngay_loi:
                                        st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                                    else:
                                        try:
                                            ten_don_vi_thu_huong = generate_ten_don_vi_thu_huong(htn)
                                            db = st.session_state.db_engine.get_connection()
                                            c = db.cursor()

                                            c.execute("SELECT COALESCE(MAX(STT),0)+1 FROM nhan_vien")
                                            stt_moi = c.fetchone()[0]

                                            # Tạo mã nhân viên mới (theo đúng ký hiệu riêng của tenant)
                                            ma_nv = sinh_ma_nv_moi(c)

                                            nhl = parse_date(nvl)

                                            ma_cty_hd = st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL'
                                            if lhd == "Thử việc":
                                                ttnv = 'THU_VIEC'
                                                ttbh = 'CHUA_DONG'
                                                tbd_val = None
                                                pa_val_add = None
                                                so_hd = sinh_so_hdld_moi(c, ma_cty_hd, nhl.year, la_thu_viec=True)
                                            else:
                                                ttnv = 'DANG_LAM'
                                                ttbh = 'DANG_DONG'
                                                tbd_val = tinh_thang_bat_dau_bh(nhl)
                                                pa_val_add = lay_ma_phuong_an(pa_add)
                                                so_hd = sinh_so_hdld_moi(c, ma_cty_hd, nhl.year, la_thu_viec=False)
                                            
                                            # Chuẩn hóa tên phòng ban
                                            pbn_chuan = chuan_hoa_ten_phong_ban(pbn)
                                            
                                            c.execute("""INSERT INTO nhan_vien (STT, ma_nv, so_hdld, ho_ten, chuc_danh_nghe, ngay_sinh, gioi_tinh,
                                                so_cccd, ngay_cap_cccd, noi_cap_cccd, nguyen_quan, thuong_tru,
                                                dien_thoai, email, email_lien_he, ho_so, luong_bao_hiem, ma_so_bhxh, ngay_vao_lam,
                                                noi_lam_viec, so_tai_khoan_nh, chi_nhanh_nh, ngay_ky_hd, loai_hop_dong,
                                                nhom_bhxh, thang_bat_dau_bh, thang_ket_thuc_bh, trang_thai, trang_thai_bhxh,
                                                phong_ban_lam_viec, ngay_ket_thuc, quoc_tich, dan_toc, he_so_luong, phu_cap_chuc_vu,
                                                phu_cap_tnvk, phu_cap_tnn, muc_huong_bhyt, ty_le_dong, muc_tien_dong, phuong_thuc_dong,
                                                tinh_nhan_hs, phuong_nhan_hs, dia_chi_nhan_hs, tinh_kcb, noi_dang_ky_kcb, dang_ky_nhan_so,
                                                ten_don_vi_thu_huong, trinh_do, so_luong_npt,
                                                phuong_an_dieu_chinh, thang_phuong_an)
                                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                                                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                                                %s, %s, %s, %s, %s, %s, %s, %s, %s,
                                                %s, %s) RETURNING id""",
                                                (stt_moi, ma_nv, so_hd, htn, cdn, parse_date(nsn), gtn, scc, parse_date(ncc), ncc2, nqn, ttn,
                                                 (dtn2.strip() or None) if dtn2 else None, emn, emn, hso, to_float_or_none(lbh), mbh, parse_date(nvl), nlv, stk, cnh, parse_date(nvl), lhd,
                                                 nbh, tbd_val, None, ttnv, ttbh, pbn_chuan, tinh_ngay_ket_thuc(lhd, parse_date(nvl)), qtn, dtn, 
                                                 to_float_or_none(hsl), to_float_or_none(pcv), to_float_or_none(ptv), to_float_or_none(ptn),
                                                 mhb, to_float_or_none(tld), to_float_or_none(mtd), ptd, ths, phs, dhs, tkb, nkb, dks,
                                                 ten_don_vi_thu_huong, trinh_do_moi, so_luong_npt,
                                                 pa_val_add, format_thang_nam(tbd_val)))
                                            new_nv_id = c.fetchone()[0]
                                            
                                            if anh_ho_so_moi is not None:
                                                storage_path_anh = upload_anh_ho_so(ma_nv, htn, anh_ho_so_moi)
                                                if storage_path_anh:
                                                    c.execute("UPDATE nhan_vien SET anh_ho_so=%s WHERE id=%s", (storage_path_anh, new_nv_id))
                                            
                                            db.commit()
                                            db.close()
                                            st.success(f"✅ Đã lưu nhân viên mới thành công! {htn} - {ma_nv}")
                                            st.cache_data.clear()
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"❌ Lỗi: {e}")
                                else:
                                    st.error("Họ tên không được để trống!")
                    with col_save_exit2:
                        if st.form_submit_button("❌ THOÁT", width='stretch'):
                            # Reset tất cả các trường trong form
                            # Xóa các session state liên quan đến form thêm nhân viên
                            keys_to_clear = [
                                'htn', 'nsn', 'gtn', 'scc', 'ncc', 'ncc2', 'nqn', 'ttn', 
                                'dtn2', 'emn', 'cdn', 'pbn', 'nlv', 'lhd', 'nvl', 'nkt',
                                'mbh', 'tbd', 'lbh', 'hsl', 'pcv', 'ptv', 'ptn', 'mhb',
                                'tld', 'mtd', 'ptd', 'nbh', 'stk', 'add_cnh', 'tkb', 'nkb',
                                'ths', 'phs', 'dhs', 'dks', 'hso', 'trinh_do_moi', 'trinh_do_add',
                                'so_luong_npt_add', 'qtn', 'dtn', 'anh_ho_so_add'
                            ]
                            for key in keys_to_clear:
                                if key in st.session_state:
                                    del st.session_state[key]
                            # Đổi key của form & expander để buộc Streamlit khởi tạo lại
                            # toàn bộ widget (thu gọn expander, xóa sạch mọi giá trị đã nhập)
                            st.session_state.add_nv_reset_ctr += 1
                            st.success("✅ Đã thoát form thêm nhân viên")
                            st.cache_data.clear()
                            st.rerun()                
                st.divider()


        if _chedo_nv == "🔍 Danh sách nhân viên đang làm việc (bao gồm thử việc)":
            sn = st.text_input("🔍 Tìm kiếm", key="snv_dang_lam")

        
            db_f = st.session_state.db_engine.get_connection()
            c_f = db_f.cursor()
            c_f.execute("SELECT DISTINCT chuc_danh_nghe FROM nhan_vien WHERE trang_thai IN ('DANG_LAM','THU_VIEC') AND chuc_danh_nghe IS NOT NULL AND chuc_danh_nghe != '' ORDER BY chuc_danh_nghe")
            ds_chuc_danh = [row[0] for row in c_f.fetchall()]
            c_f.execute("SELECT DISTINCT loai_hop_dong FROM nhan_vien WHERE trang_thai IN ('DANG_LAM','THU_VIEC') AND loai_hop_dong IS NOT NULL AND loai_hop_dong != '' ORDER BY loai_hop_dong")
            ds_loai_hd = [row[0] for row in c_f.fetchall()]
            c_f.execute("SELECT DISTINCT trinh_do FROM nhan_vien WHERE trang_thai IN ('DANG_LAM','THU_VIEC') AND trinh_do IS NOT NULL AND trinh_do != '' ORDER BY trinh_do")
            ds_bang_cap = [row[0] for row in c_f.fetchall()]
            db_f.close()
        
            col_f1, col_f2, col_f3 = st.columns(3)
            with col_f1:
                filter_chuc_danh = st.selectbox("🔍 Lọc Chức danh:", ["Tất cả"] + ds_chuc_danh, key="filter_cd_danglam")
            with col_f2:
                filter_loai_hd = st.selectbox("🔍 Lọc Loại HĐ:", ["Tất cả"] + ds_loai_hd, key="filter_lhd_danglam")
            with col_f3:
                filter_bang_cap = st.selectbox("🔍 Lọc theo Bằng cấp:", ["Tất cả"] + ds_bang_cap, key="filter_bc_danglam")
        
            db = st.session_state.db_engine.get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            sql = "SELECT * FROM nhan_vien WHERE trang_thai IN ('DANG_LAM','THU_VIEC')"
            params = []
            if sn:
                sql += " AND (ho_ten LIKE %s OR dien_thoai LIKE %s OR so_cccd LIKE %s OR ma_nv LIKE %s)"
                params.extend([f'%{sn}%'] * 4)
            if filter_chuc_danh != "Tất cả":
                sql += " AND chuc_danh_nghe = %s"
                params.append(filter_chuc_danh)
            if filter_loai_hd != "Tất cả":
                sql += " AND loai_hop_dong = %s"
                params.append(filter_loai_hd)
            if filter_bang_cap != "Tất cả":
                sql += " AND trinh_do = %s"
                params.append(filter_bang_cap)
            sql += " ORDER BY id DESC"
            c.execute(sql, tuple(params))
            ds = c.fetchall()
            db.close()
        
            if ds:
                # ===== KIỂM TRA NẾU CHỈ CÓ 1 KẾT QUẢ TÌM KIẾM =====
                if len(ds) == 1:
                    nv = ds[0]  # Lấy nhân viên duy nhất
                    st.success(f"🎯 Tìm thấy 1 nhân viên: {nv['ho_ten']}")
                    render_employee_info_card(
                        nv,
                        key_prefix=f"single_{nv['id']}",
                        on_close=lambda: st.session_state.update({'_reset_snv_dang_lam': True})
                    )

                    # Thêm tùy chọn hiển thị bảng
                    st.divider()
                    if st.checkbox("📊 Hiển thị danh sách đầy đủ", value=False, key="show_full_list_card"):
                        # Hiển thị bảng bên dưới
                        pass
                    else:
                        # Nếu không hiển thị bảng, vẫn cần render các phần bên dưới
                        # nhưng chúng ta sẽ bỏ qua phần bảng
                        # Để không bị lỗi, chúng ta sẽ đặt một flag
                        st.session_state['skip_table_display'] = True
                        # Vẫn cần giữ các form sửa nhân viên ở phía sau
                        # nhưng chúng sẽ không hiển thị nếu không có selected_nv_id
                        pass
            
                # ===== PHẦN HIỂN THỊ BẢNG (CHẠY KHI CÓ NHIỀU KẾT QUẢ HOẶC USER CHỌN HIỂN THỊ) =====
                # Chỉ hiển thị bảng nếu có nhiều hơn 1 kết quả HOẶC user chọn hiển thị đầy đủ
                show_table = (len(ds) > 1) or (len(ds) == 1 and st.session_state.get('show_full_list_card', False))
            
                if show_table or len(ds) > 1:
                    # Reset flag nếu có
                    st.session_state['skip_table_display'] = False
                
                    df = pd.DataFrame(ds)
                    for col in df.columns:
                        if 'ngay' in col.lower():
                            df[col] = df[col].apply(format_date)
                
                    if 'selected' not in df.columns:
                        df.insert(0, 'selected', False)
                
                    display_cols = ['selected', 'ma_nv', 'ho_ten', 'ngay_sinh', 'gioi_tinh', 'so_hdld', 'so_cccd', 'dien_thoai',
                                    'thuong_tru', 'chuc_danh_nghe', 'loai_hop_dong', 'ngay_vao_lam', 'ma_so_bhxh', 'thang_bat_dau_bh',
                                    'ten_don_vi_thu_huong']
                    # viewer và kt_luong: ẩn thông tin nhạy cảm (CCCD, STK ngân hàng) trên bảng danh sách
                    SENSITIVE_COLS = {'so_cccd', 'so_tai_khoan_nh'}
                    if st.session_state.role in ("viewer", "kt_luong"):
                        display_cols = [c for c in display_cols if c not in SENSITIVE_COLS]
                    available_cols = [c for c in display_cols if c in df.columns]
                    df_show = df[available_cols]
                
                    col_map = {
                        'selected': 'Chọn',
                        'ma_nv': 'Mã NV',
                        'ho_ten': 'Họ và tên',
                        'ngay_sinh': 'Ngày sinh',
                        'gioi_tinh': 'Giới tính',
                        'so_hdld': 'Số HĐLĐ',
                        'so_cccd': 'CCCD',
                        'dien_thoai': 'SĐT',
                        'thuong_tru': 'Thường trú',
                        'chuc_danh_nghe': 'Chức danh',
                        'loai_hop_dong': 'Loại HĐ',
                        'ngay_vao_lam': 'Ngày vào làm',
                        'ma_so_bhxh': 'Mã số BHXH',
                        'thang_bat_dau_bh': 'Bắt đầu BH',
                        'ten_don_vi_thu_huong': 'Tên đơn vị thụ hưởng',
                    }
                    df_show.rename(columns=col_map, inplace=True)
                
                    if len(ds) > 1:
                        st.caption(f"📌 {len(ds)} kết quả. Tick chọn 1 nhân viên để thao tác.")
                    else:
                        st.caption(f"📌 Danh sách đầy đủ ({len(ds)} kết quả). Tick chọn 1 nhân viên để thao tác.")
                
                    # Nếu là viewer, hiển thị bảng không có checkbox chọn
                    if st.session_state.role in ("admin", "xem_toan_bo"):
                        # Xử lý yêu cầu reset lựa chọn (đến từ nút "Đóng" của card thông tin nhân viên)
                        # Phải làm TRƯỚC khi widget data_editor được khởi tạo
                        if st.session_state.pop('_reset_nv_editor_danglam', False):
                            st.session_state.pop('nv_editor_danglam', None)
                        edited_df = st.data_editor(
                            df_show,
                            column_config={
                                "Chọn": st.column_config.CheckboxColumn("Profile", default=False)
                            },
                            disabled=[col for col in df_show.columns if col != 'Chọn'],
                            hide_index=True,
                            height=400,
                            key="nv_editor_danglam"
                        )
                    else:
                        # Viewer (và các role không phải admin): vẫn được TICK CHỌN 1 dòng ở cột
                        # "Profile" để xem card "Thông tin nhân sự" (chỉ xem — mọi cột khác đều
                        # disabled nên không sửa được gì). Các nút Sửa/In HĐLĐ/Gửi Zalo sẽ được
                        # ẩn đi bên trong render_employee_info_card() theo role.
                        if st.session_state.pop('_reset_nv_editor_viewer_danglam', False):
                            st.session_state.pop('nv_editor_viewer_danglam', None)
                        edited_df = st.data_editor(
                            df_show,
                            column_config={
                                "Chọn": st.column_config.CheckboxColumn("Profile", default=False)
                            },
                            disabled=[col for col in df_show.columns if col != 'Chọn'],
                            hide_index=True,
                            height=400,
                            key="nv_editor_viewer_danglam"
                        )
                
                    # Viewer (và các role không phải admin): chọn 1 dòng -> chỉ xem card thông tin,
                    # không có bất kỳ nút hành động nào ngoài "Đóng" (xử lý theo role bên trong hàm).
                    if edited_df is not None and st.session_state.role not in ("admin", "xem_toan_bo") and 'Chọn' in edited_df.columns:
                        selected_rows_v = edited_df[edited_df['Chọn'] == True]
                        if len(selected_rows_v) > 1:
                            st.error("⚠️ Chỉ được chọn 1 nhân viên!")
                        elif len(selected_rows_v) == 1:
                            selected_idx_v = selected_rows_v.index[0]
                            selected_nv_v = df.iloc[selected_idx_v]
                            render_employee_info_card(
                                selected_nv_v,
                                key_prefix=f"viewer_{selected_nv_v['id']}",
                                on_close=lambda: st.session_state.update({'_reset_nv_editor_viewer_danglam': True})
                            )

                    selected_nv = None
                    if edited_df is not None and st.session_state.role in ("admin", "xem_toan_bo") and 'Chọn' in edited_df.columns:
                        selected_rows = edited_df[edited_df['Chọn'] == True]
                        if len(selected_rows) > 0:
                            if len(selected_rows) > 1:
                                st.error("⚠️ Chỉ được chọn 1 nhân viên!")
                            else:
                                selected_idx = selected_rows.index[0]
                                selected_nv = df.iloc[selected_idx]
                                nv_id_key = selected_nv['id']
                            
                                # Hiển thị các nút chức năng (chỉ admin mới thấy và mới click được)
                                render_employee_info_card(
                                    selected_nv,
                                    key_prefix=f"multi_{nv_id_key}",
                                    on_close=lambda: st.session_state.update({'_reset_nv_editor_danglam': True})
                                )
                                col_btn5 = st.container()

                                with col_btn5:
                                    trang_thai_nv = selected_nv.get('trang_thai', '')
                                    if trang_thai_nv == 'THU_VIEC':
                                        if f'convert_open_{nv_id_key}' not in st.session_state:
                                            st.session_state[f'convert_open_{nv_id_key}'] = False
                                    
                                        if not st.session_state[f'convert_open_{nv_id_key}']:
                                            st.info("💡 Chuyển đổi HĐTV → Chính thức đã chuyển sang tab **📜 QUYẾT ĐỊNH NHÂN SỰ** → chọn loại **'QĐ Chuyển đổi TV → Chính thức'**.")
                                            if False and st.button(f"🔄 CHUYỂN ĐỔI HĐLĐ - {selected_nv['ho_ten']}", 
                                                        key=f"convert_hdld_btn_{nv_id_key}", 
                                                        width='stretch', type="primary"):
                                                st.session_state[f'convert_open_{nv_id_key}'] = True
                                                st.rerun()
                                        else:
                                            st.markdown("---")
                                            st.markdown("### 📝 CHUYỂN ĐỔI HỢP ĐỒNG LAO ĐỘNG")
                                            st.caption("Vui lòng nhập đầy đủ thông tin cho quyết định chuyển đổi")
                                        
                                            db_temp = st.session_state.db_engine.get_connection()
                                            c_temp = db_temp.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                                            c_temp.execute("SELECT * FROM nhan_vien WHERE id = %s", (int(selected_nv['id']),))
                                            nv_data = c_temp.fetchone()
                                            db_temp.close()
                                        
                                            if nv_data:
                                                ngay_quyet_dinh = st.date_input(
                                                    "📅 Ngày quyết định:", 
                                                    value=date.today(),
                                                    key=f"ngay_qd_{nv_id_key}"
                                                )

                                                # Mã công ty của TENANT ĐANG ĐĂNG NHẬP — trước đây bị khóa cứng "CHL"
                                                # (mã của Hòn La) nên với tenant khác (VD DEMO-HRM), pattern LIKE
                                                # '%/HĐLĐ-CHL' không bao giờ khớp -> max_stt luôn = 0 -> số luôn ra "01".
                                                ma_cty_hd = st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL'

                                                loai_hd_moi_lbl = st.selectbox(
                                                    "📑 Loại HĐLĐ mới:",
                                                    ["Không xác định thời hạn", "Xác định thời hạn 12 tháng",
                                                     "Xác định thời hạn 24 tháng", "Xác định thời hạn 36 tháng"],
                                                    key=f"loai_hd_moi_{nv_id_key}"
                                                )
                                                if loai_hd_moi_lbl == "Không xác định thời hạn":
                                                    loai_hop_dong_luu = "Không xác định thời hạn"
                                                    han_hd_thang = None
                                                else:
                                                    loai_hop_dong_luu = "Xác định thời hạn"
                                                    han_hd_thang = int(loai_hd_moi_lbl.split()[-2])

                                                current_year = datetime.now().year
                                                db_temp2 = st.session_state.db_engine.get_connection()
                                                c_temp2 = db_temp2.cursor()
                                                c_temp2.execute("""
                                                    SELECT COALESCE(MAX(CAST(SPLIT_PART(so_hdld, '/', 1) AS INTEGER)), 0) as max_stt
                                                    FROM nhan_vien 
                                                    WHERE so_hdld LIKE %s 
                                                    AND SPLIT_PART(so_hdld, '/', 1) ~ '^[0-9]+$'
                                                    AND trang_thai = 'DANG_LAM'
                                                    AND loai_hop_dong != 'Thử việc'
                                                """, (f'%/{current_year}/HĐLĐ-%',))
                                                result = c_temp2.fetchone()
                                                max_stt = result[0] if result else 0
                                                db_temp2.close()
                                            
                                                next_stt = max_stt + 1
                                                stt_str = str(next_stt).zfill(2)
                                                so_hd_moi = f"{stt_str}/{current_year}/HĐLĐ-{ma_cty_hd}"
                                            
                                                st.info(f"📄 **Số HĐLĐ mới:** {so_hd_moi} (tự động sinh)")
                                            
                                                ngay_hieu_luc = st.date_input(
                                                    "📅 Ngày hiệu lực (bắt đầu HĐLĐ):", 
                                                    value=ngay_quyet_dinh,
                                                    key=f"ngay_hl_{nv_id_key}"
                                                )

                                                if han_hd_thang:
                                                    ngay_het_han_hd = ngay_hieu_luc + relativedelta(months=han_hd_thang) - timedelta(days=1)
                                                    st.caption(f"📆 Hợp đồng sẽ hết hạn: {ngay_het_han_hd.strftime('%d/%m/%Y')}")
                                                else:
                                                    ngay_het_han_hd = None
                                            
                                                ngay_bat_dau_bh = tinh_thang_bat_dau_bh(ngay_hieu_luc)
                                                st.info(f"📅 Tháng bắt đầu đóng BHXH: **{format_thang_nam(ngay_bat_dau_bh)}** (tự tính theo quy tắc 14 ngày)")
                                            
                                                phuong_an_chuyen_doi = st.selectbox(
                                                    "Phương án điều chỉnh BHXH",
                                                    [""] + PHUONG_AN_TANG,
                                                    key=f"pa_bhxh_{nv_id_key}",
                                                    help="Bắt buộc chọn — dùng cho báo tăng D02-LT"
                                                )
                                            
                                                ly_do_chuyen = st.text_area(
                                                    "📝 Lý do/ Nội dung quyết định:", 
                                                    value=f"Hoàn thành thời gian thử việc, chuyển sang hợp đồng lao động {loai_hd_moi_lbl.lower()}",
                                                    key=f"ly_do_{nv_id_key}",
                                                    height=80
                                                )
                                            
                                                col_confirm1, col_confirm2, col_confirm3 = st.columns([1, 2, 1])
                                                with col_confirm2:
                                                    if st.button("✅ XÁC NHẬN CHUYỂN ĐỔI", key=f"confirm_convert_{nv_id_key}", width='stretch', type="primary"):
                                                        try:
                                                            db = st.session_state.db_engine.get_connection()
                                                            c = db.cursor()
                                                        
                                                            so_hd_tv_cu = selected_nv.get('so_hdld', '')
                                                            ngay_vao_lam_cu = selected_nv.get('ngay_vao_lam')
                                                        
                                                            if ngay_vao_lam_cu:
                                                                if hasattr(ngay_vao_lam_cu, 'strftime'):
                                                                    pass
                                                                else:
                                                                    ngay_vao_lam_cu = parse_date(ngay_vao_lam_cu)
                                                                    if not ngay_vao_lam_cu:
                                                                        ngay_vao_lam_cu = date.today()
                                                            else:
                                                                ngay_vao_lam_cu = date.today()
                                                        
                                                            current_year = datetime.now().year
                                                            c.execute("""
                                                                SELECT COALESCE(MAX(CAST(SPLIT_PART(so_hdld, '/', 1) AS INTEGER)), 0) as max_stt
                                                                FROM nhan_vien 
                                                                WHERE so_hdld LIKE %s 
                                                                AND SPLIT_PART(so_hdld, '/', 1) ~ '^[0-9]+$'
                                                                AND trang_thai = 'DANG_LAM'
                                                                AND loai_hop_dong != 'Thử việc'
                                                            """, (f'%/{current_year}/HĐLĐ-%',))
                                                            result = c.fetchone()
                                                            max_stt = result[0] if result else 0
                                                            next_stt = max_stt + 1
                                                            stt_str = str(next_stt).zfill(2)
                                                            so_hd_moi = f"{stt_str}/{current_year}/HĐLĐ-{ma_cty_hd}"
                                                        
                                                            pa_chuyen_doi_val = lay_ma_phuong_an(phuong_an_chuyen_doi)
                                                            c.execute("""
                                                                UPDATE nhan_vien SET 
                                                                    trang_thai = 'DANG_LAM',
                                                                    loai_hop_dong = %s,
                                                                    han_hop_dong_thang = %s,
                                                                    so_hdld = %s,
                                                                    ngay_ky_hd = %s,
                                                                    ngay_chinh_thuc = %s,
                                                                    thang_bat_dau_bh = %s,
                                                                    trang_thai_bhxh = 'DANG_DONG',
                                                                    phuong_an_dieu_chinh = %s,
                                                                    thang_phuong_an = %s,
                                                                    ngay_ket_thuc = NULL
                                                                WHERE id = %s
                                                            """, (loai_hop_dong_luu, han_hd_thang, so_hd_moi, ngay_quyet_dinh, ngay_hieu_luc, ngay_bat_dau_bh, pa_chuyen_doi_val, format_thang_nam(ngay_bat_dau_bh), int(selected_nv['id'])))
                                                        
                                                            c.execute("""
                                                                INSERT INTO quyet_dinh_nhan_su (
                                                                    nhan_vien_id, loai_quyet_dinh, ngay_quyet_dinh, ngay_hieu_luc,
                                                                    noi_dung, so_quyet_dinh, loai_hop_dong_cu, loai_hop_dong_moi,
                                                                    he_so_luong_cu, he_so_luong_moi, so_hd_cu
                                                                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                                            """, (
                                                                int(selected_nv['id']),
                                                                'CHINH_THUC',
                                                                ngay_quyet_dinh,
                                                                ngay_hieu_luc,
                                                                ly_do_chuyen,
                                                                f"QD{ngay_quyet_dinh.strftime('%Y%m%d')}_{selected_nv['ma_nv']}",
                                                                nv_data.get('loai_hop_dong', 'Thử việc'),
                                                                loai_hop_dong_luu,
                                                                nv_data.get('he_so_luong', 0),
                                                                nv_data.get('he_so_luong', 0),
                                                                so_hd_tv_cu
                                                            ))
                                                        
                                                            c.execute("""
                                                                UPDATE lich_su_cong_tac 
                                                                SET den_ngay = %s,
                                                                    so_hop_dong = %s
                                                                WHERE nhan_vien_id = %s 
                                                                AND loai_hop_dong = 'Thử việc'
                                                                AND den_ngay IS NULL
                                                            """, (ngay_hieu_luc - timedelta(days=1), so_hd_tv_cu, int(selected_nv['id'])))
                                                        
                                                            c.execute("""
                                                                INSERT INTO lich_su_cong_tac (
                                                                    nhan_vien_id, tu_ngay, chuc_danh, phong_ban, 
                                                                    noi_lam_viec, loai_hop_dong, he_so_luong, so_hop_dong
                                                                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                                                            """, (
                                                                int(selected_nv['id']),
                                                                ngay_hieu_luc,
                                                                nv_data.get('chuc_danh_nghe', ''),
                                                                nv_data.get('phong_ban_lam_viec', ''),
                                                                nv_data.get('noi_lam_viec') or get_cau_hinh('noi_lam_viec', 'Cảng THQT Hòn La'),
                                                                loai_hop_dong_luu,
                                                                nv_data.get('he_so_luong', 0),
                                                                so_hd_moi
                                                            ))
                                                        
                                                            db.commit()
                                                            db.close()
                                                        
                                                            st.success(f"✅ Đã chuyển {nv_data['ho_ten']} sang HĐLĐ {loai_hd_moi_lbl.lower()}!")
                                                            st.info(f"📄 Số HĐTV cũ: {so_hd_tv_cu}")
                                                            st.info(f"📄 Số HĐLĐ mới: {so_hd_moi}")
                                                            st.cache_data.clear()
                                                            st.session_state[f'convert_open_{nv_id_key}'] = False
                                                            st.rerun()
                                                        
                                                        except Exception as e:
                                                            db.rollback()
                                                            db.close()
                                                            st.error(f"❌ Lỗi: {str(e)}")

                                                if st.button("❌ HỦY", key=f"cancel_convert_{nv_id_key}", width='stretch'):
                                                    st.session_state[f'convert_open_{nv_id_key}'] = False
                                                    st.rerun()

                                st.divider()
            
                # Form sửa nhân viên (chỉ admin)
                if 'selected_nv_id' in st.session_state and st.session_state.selected_nv_id is not None and st.session_state.role in ("admin", "xem_toan_bo"):
                    try:
                        nid = int(st.session_state['selected_nv_id'])
                        db = st.session_state.db_engine.get_connection()
                        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                        c.execute("SELECT * FROM nhan_vien WHERE id=%s", (nid,))
                        nd = c.fetchone()
                        db.close()
                        db_cd = st.session_state.db_engine.get_connection()
                        c_cd = db_cd.cursor()
                        c_cd.execute("SELECT DISTINCT ten_vi_tri FROM vi_tri_cong_tac ORDER BY ten_vi_tri")
                        dcv_edit = [row[0] for row in c_cd.fetchall()]
                        db_cd.close()
                        dpb_edit = get_phong_ban_options()
                    
                        if nd:
                            st.subheader(f"✏️ Cập nhật: {nd.get('ho_ten', '')} ({nd.get('ma_nv', '')})")
                            # Đặt Loại HĐ + Ngày vào làm NGOÀI form để đổi giá trị → rerun ngay → disable Mã BHXH & Phương án
                            _col_lhd_e1, _col_lhd_e2 = st.columns(2)
                            with _col_lhd_e1:
                                nvlv = st.text_input("Ngày vào làm (dd/mm/yyyy)", value=format_date(nd.get('ngay_vao_lam')), placeholder="dd/mm/yyyy", max_chars=10, key="nvlv_edit_outside")
                            with _col_lhd_e2:
                                lhdv = st.selectbox("Loại HĐ", ["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"], index=["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"].index(nd.get('loai_hop_dong', 'Thử việc')) if nd.get('loai_hop_dong') in ["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"] else 0, key="lhdv_edit_outside")
                            la_thu_viec_edit = lhdv == "Thử việc"
                            with st.form("edit_nv"):
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    hnv = st.text_input("Họ và tên *", value=nd.get('ho_ten', ''))
                                    nsnv = st.text_input("Ngày sinh (dd/mm/yyyy)", value=format_date(nd.get('ngay_sinh')), placeholder="dd/mm/yyyy", max_chars=10)
                                    gtnv = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"], index=["", "Nam", "Nữ", "Khác"].index(nd.get('gioi_tinh', '')) if nd.get('gioi_tinh') in ["Nam", "Nữ", "Khác"] else 0)
                                    sccv = st.text_input("CCCD", value=nd.get('so_cccd', ''))
                                    nccv = st.text_input("Ngày cấp CCCD (dd/mm/yyyy)", value=format_date(nd.get('ngay_cap_cccd')), placeholder="dd/mm/yyyy", max_chars=10)
                                    ncv = st.text_input("Nơi cấp CCCD", value=nd.get('noi_cap_cccd', ''))
                                with col2:
                                    nqnv = nd.get('nguyen_quan', '')
                                    ttnv = st.text_input("Thường trú", value=nd.get('thuong_tru', ''))
                                    qtnv = st.text_input("Quốc tịch", value=nd.get('quoc_tich', 'Việt Nam'))
                                    dtnv = st.text_input("Dân tộc", value=nd.get('dan_toc', 'Kinh'))
                                    dtnv2 = st.text_input("SĐT", value=nd.get('dien_thoai', ''))
                                    trinh_do_v = st.selectbox("Trình độ", [""] + TRINH_DO_LIST, index=([""] + TRINH_DO_LIST).index(nd.get('trinh_do', '')) if nd.get('trinh_do') in TRINH_DO_LIST else 0)
                                    cdnv = st.selectbox("Chức danh", [""] + dcv_edit, index=([""] + dcv_edit).index(nd.get('chuc_danh_nghe', '')) if nd.get('chuc_danh_nghe') in dcv_edit else 0)
                                with col3:
                                    pb_hien_tai_chuan = chuan_hoa_ten_phong_ban(nd.get('phong_ban_lam_viec'))
                                    pbnv = st.selectbox("Phòng ban", [""] + dpb_edit, index=([""] + dpb_edit).index(pb_hien_tai_chuan) if pb_hien_tai_chuan in dpb_edit else 0)
                                    nlv2 = nd.get('noi_lam_viec', 'Cảng THQT Hòn La')
                                    emnv = st.text_input("Email", value=nd.get('email_lien_he', ''))
                                    anh_hien_tai = nd.get('anh_ho_so')
                                    if anh_hien_tai:
                                        anh_bytes_ht = get_anh_ho_so_bytes(anh_hien_tai)
                                        if anh_bytes_ht:
                                            st.image(anh_bytes_ht, caption="Ảnh hồ sơ hiện tại", width=120)
                                    anh_ho_so_v = st.file_uploader("Đổi ảnh hồ sơ (bỏ trống nếu giữ nguyên)", key=f"anh_ho_so_edit_{nid}")
                            
                                st.divider()
                                st.caption("💼 Bảo hiểm xã hội")
                                col4, col5, col6 = st.columns(3)
                                with col4:
                                    mbhv = st.text_input("Mã BHXH", value=nd.get('ma_so_bhxh', ''), disabled=la_thu_viec_edit)
                                    # Phương án điều chỉnh BHXH
                                    pa_hien_tai = nd.get('phuong_an_dieu_chinh', '')
                                    pa_label_hien_tai = ""
                                    for pa in PHUONG_AN_ALL:
                                        if pa.startswith(pa_hien_tai + " - ") if pa_hien_tai else False:
                                            pa_label_hien_tai = pa
                                            break
                                    pa_index = ([""] + PHUONG_AN_ALL).index(pa_label_hien_tai) if pa_label_hien_tai in PHUONG_AN_ALL else 0
                                    pa_edit = st.selectbox("Phương án điều chỉnh", [""] + PHUONG_AN_ALL, index=pa_index, key="pa_edit", disabled=la_thu_viec_edit)
                                    lbhv = st.text_input("Lương BH", value=nd.get('luong_bao_hiem', ''))
                                    hslv = st.text_input("Hệ số lương", value=str(nd.get('he_so_luong', '')))
                                with col5:
                                    pcvv = st.text_input("PC chức vụ", value=str(nd.get('phu_cap_chuc_vu', '')))
                                    ptvv = st.text_input("PC TNVK (%)", value=str(nd.get('phu_cap_tnvk', '')))
                                    ptnv = st.text_input("PC TNN (%)", value=str(nd.get('phu_cap_tnn', '')))
                                    mhbv = st.selectbox("Mức hưởng BHYT", ["80%", "95%", "100%"], index=["80%", "95%", "100%"].index(nd.get('muc_huong_bhyt', '80%')) if nd.get('muc_huong_bhyt') in ["80%", "95%", "100%"] else 0)
                                with col6:
                                    tldv = st.text_input("Tỷ lệ đóng (%)", value=str(nd.get('ty_le_dong', '')))
                                    mtdv = st.text_input("Mức tiền đóng", value=str(nd.get('muc_tien_dong', '')))
                                    ptdv = st.selectbox("PT đóng", ["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"], index=["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"].index(nd.get('phuong_thuc_dong', 'Hàng tháng')) if nd.get('phuong_thuc_dong') in ["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"] else 0)
                                    nbhv = st.selectbox("Nhóm BHXH", ["", "Văn phòng", "Lao động trực tiếp"], index=["", "Văn phòng", "Lao động trực tiếp"].index(nd.get('nhom_bhxh', '')) if nd.get('nhom_bhxh') in ["Văn phòng", "Lao động trực tiếp"] else 0)
                            
                                st.divider()
                                st.caption("🏦 Ngân hàng & Hồ sơ")
                                col7, col8 = st.columns(2)
                                with col7:
                                    stkv = st.text_input("STK", value=nd.get('so_tai_khoan_nh', ''))
                                    # Tạo dropdown cho chi nhánh ngân hàng
                                    bank_edit_index = 0
                                    old_bank = nd.get('chi_nhanh_nh', '')
                                    if old_bank in BANK_LIST:
                                        bank_edit_index = BANK_LIST.index(old_bank) + 1
                                    cnhv = st.selectbox("Chi nhánh NH", options=[""] + BANK_LIST, index=bank_edit_index, key="edit_cnh")
                                with col8:
                                    hsov = st.selectbox("Hồ sơ", ["", "Đã có HS", "Chưa có"], index=["", "Đã có HS", "Chưa có"].index(nd.get('ho_so', '')) if nd.get('ho_so') in ["Đã có HS", "Chưa có"] else 0)
                                    so_luong_npt_edit = st.number_input("Số người phụ thuộc", min_value=0, value=int(nd.get('so_luong_npt') or 0), step=1, key=f"so_luong_npt_edit_{nid}")
                                # Các trường ít dùng (Tỉnh KCB, Nơi KCB, Tỉnh/TP nhận HS, Phường/Xã nhận HS,
                                # Địa chỉ nhận HS, ĐK nhận sổ) đã bỏ khỏi UI theo yêu cầu — giữ nguyên
                                # giá trị đã lưu trong hồ sơ thay vì hiện ô nhập.
                                tkbv = nd.get('tinh_kcb', '')
                                nkbv = nd.get('noi_dang_ky_kcb', '')
                                thsv = nd.get('tinh_nhan_hs', '')
                                phsv = nd.get('phuong_nhan_hs', '')
                                dhsv = nd.get('dia_chi_nhan_hs', '')
                                dksv = nd.get('dang_ky_nhan_so', 'Có')
                                nktv = format_date(nd.get('ngay_ket_thuc'))
                                col_save, col_cancel = st.columns(2)
                                with col_save:
                                    if st.form_submit_button("💾 CẬP NHẬT", width='stretch', disabled=not can_edit()):
                                        if not can_edit():
                                            st.error("❌ Bạn không có quyền thực hiện thao tác này!")
                                        else:
                                            if hnv:
                                                ngay_loi = []
                                                if nsnv and not parse_date(nsnv):
                                                    ngay_loi.append("Ngày sinh")
                                                if nccv and not parse_date(nccv):
                                                    ngay_loi.append("Ngày cấp CCCD")
                                                if nvlv and not parse_date(nvlv):
                                                    ngay_loi.append("Ngày vào làm")
                                                if nktv and not parse_date(nktv):
                                                    ngay_loi.append("Ngày kết thúc")
                                                if ngay_loi:
                                                    st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                                                else:
                                                    try:
                                                        ten_don_vi_thu_huong = generate_ten_don_vi_thu_huong(hnv)
                                                    
                                                        db_upd = st.session_state.db_engine.get_connection()
                                                        c_upd = db_upd.cursor()
                                                        nhl = parse_date(nvlv) or date.today()
                                                        if lhdv == "Thử việc":
                                                            tt_nv, tt_bh, tbd_val = 'THU_VIEC', 'CHUA_DONG', None
                                                        else:
                                                            tt_nv, tt_bh = 'DANG_LAM', 'DANG_DONG'
                                                            # Giữ nguyên tháng BH đã lưu; nếu chưa có thì tự tính
                                                            tbd_val = nd.get('thang_bat_dau_bh') or tinh_thang_bat_dau_bh(nhl)
                                                    
                                                        # Chuẩn hóa tên phòng ban
                                                        pbnv_chuan = chuan_hoa_ten_phong_ban(pbnv)
                                                    
                                                        c_upd.execute("""UPDATE nhan_vien SET ho_ten=%s,chuc_danh_nghe=%s,ngay_sinh=%s,gioi_tinh=%s,
                                                            so_cccd=%s,ngay_cap_cccd=%s,noi_cap_cccd=%s,nguyen_quan=%s,thuong_tru=%s,dien_thoai=%s,
                                                            email=%s,email_lien_he=%s,ho_so=%s,luong_bao_hiem=%s,ma_so_bhxh=%s,ngay_vao_lam=%s,noi_lam_viec=%s,
                                                            so_tai_khoan_nh=%s,chi_nhanh_nh=%s,ngay_ky_hd=%s,loai_hop_dong=%s,nhom_bhxh=%s,
                                                            thang_bat_dau_bh=%s,trang_thai=%s,trang_thai_bhxh=%s,phong_ban_lam_viec=%s,
                                                            ngay_ket_thuc=%s,quoc_tich=%s,dan_toc=%s,he_so_luong=%s,phu_cap_chuc_vu=%s,
                                                            phu_cap_tnvk=%s,phu_cap_tnn=%s,muc_huong_bhyt=%s,ty_le_dong=%s,muc_tien_dong=%s,
                                                            phuong_thuc_dong=%s,tinh_nhan_hs=%s,phuong_nhan_hs=%s,dia_chi_nhan_hs=%s,
                                                            tinh_kcb=%s,noi_dang_ky_kcb=%s,dang_ky_nhan_so=%s, ten_don_vi_thu_huong=%s, trinh_do=%s,
                                                            so_luong_npt=%s WHERE id=%s""",
                                                            (hnv, cdnv, parse_date(nsnv), gtnv, sccv, parse_date(nccv), ncv, nqnv, ttnv, (dtnv2.strip() or None) if dtnv2 else None,
                                                             emnv, emnv, hsov, to_float_or_none(lbhv), mbhv, parse_date(nvlv), nlv2, stkv, cnhv, parse_date(nvlv), lhdv,
                                                             nbhv, tbd_val, tt_nv, tt_bh, pbnv_chuan, parse_date(nktv), qtnv, dtnv,
                                                             to_float_or_none(hslv), to_float_or_none(pcvv), to_float_or_none(ptvv), to_float_or_none(ptnv),
                                                             mhbv, to_float_or_none(tldv), to_float_or_none(mtdv), ptdv, thsv, phsv, dhsv,
                                                             tkbv, nkbv, dksv, ten_don_vi_thu_huong, trinh_do_v, so_luong_npt_edit, nid))
                                                    
                                                        if anh_ho_so_v is not None:
                                                            storage_path_anh_v = upload_anh_ho_so(nd.get('ma_nv', nid), hnv, anh_ho_so_v)
                                                            if storage_path_anh_v:
                                                                c_upd.execute("UPDATE nhan_vien SET anh_ho_so=%s WHERE id=%s", (storage_path_anh_v, nid))
                                                    
                                                        db_upd.commit()
                                                        db_upd.close()
                                                        st.success(f"✅ Đã cập nhật: {hnv}")
                                                        st.cache_data.clear()
                                                        del st.session_state['selected_nv_id']
                                                        st.rerun()
                                                    except Exception as e:
                                                        st.error(f"❌ Lỗi: {e}")
                                            else:
                                                st.error("Họ tên không được để trống!")

                                with col_cancel:
                                    if st.form_submit_button("❌ HỦY SỬA", width='stretch'):
                                        # Xóa session state để đóng form sửa
                                        if 'selected_nv_id' in st.session_state:
                                            del st.session_state['selected_nv_id']
                                        st.success("✅ Đã hủy sửa nhân viên")
                                        st.rerun()
            
                    except Exception as e:
                        st.error(f"Lỗi khi tải thông tin nhân viên: {e}")
                        st.session_state.selected_nv_id = None
                        st.rerun()
            
                # Form nhập thông tin hộ gia đình (chỉ admin) - Đặt NGOÀI form sửa nhân viên
                if 'bhxh_family_nv_id' in st.session_state and st.session_state.bhxh_family_nv_id is not None and st.session_state.role in ("admin", "xem_toan_bo"):
                    nv_id = st.session_state['bhxh_family_nv_id']
                    nv_name = st.session_state['bhxh_family_nv_name']
                    st.divider()
                    st.subheader(f"🏠 NHẬP THÔNG TIN HỘ GIA ĐÌNH CHO: {nv_name}")
                    st.caption("Vui lòng nhập đầy đủ thông tin chủ hộ và các thành viên trong hộ gia đình")
                
                    db = st.session_state.db_engine.get_connection()
                    c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    c.execute("SELECT * FROM nhan_vien WHERE id = %s", (nv_id,))
                    nv_data = c.fetchone()
                    db.close()
                
                    if 'bhxh_family_members' not in st.session_state:
                        db_temp = st.session_state.db_engine.get_connection()
                        c_temp = db_temp.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                        c_temp.execute("SELECT * FROM phu_luc_gia_dinh WHERE nhan_vien_id = %s", (nv_id,))
                        existing_members = c_temp.fetchall()
                        db_temp.close()
                        st.session_state['bhxh_family_members'] = []
                        for tv in existing_members:
                            st.session_state['bhxh_family_members'].append({
                                'ho_ten': tv['ho_ten'], 'ngay_sinh': tv['ngay_sinh'], 'gioi_tinh': tv['gioi_tinh'],
                                'quoc_tich': tv['quoc_tich'], 'dan_toc': tv['dan_toc'], 'quan_he': tv['quan_he_voi_chu_ho'],
                                'tinh': tv['tinh_thanh_pho'], 'phuong_xa': tv['phuong_xa']
                            })
                
                    db_temp = st.session_state.db_engine.get_connection()
                    c_temp = db_temp.cursor()
                    c_temp.execute("SELECT ma_tinh, ten_tinh FROM danh_muc_tinh ORDER BY ten_tinh")
                    ds_tinh = c_temp.fetchall()
                    db_temp.close()
                    tinh_options = {ten: ma for ma, ten in ds_tinh}
                
                    if st.session_state.bhxh_family_members:
                        st.markdown("**Danh sách thành viên đã thêm:**")
                        tv_data = []
                        for i, tv in enumerate(st.session_state.bhxh_family_members):
                            tv_data.append({"STT": i+1, "Họ và tên": tv['ho_ten'], "Ngày sinh": format_date(tv['ngay_sinh']),
                                            "Giới tính": tv['gioi_tinh'], "Quốc tịch": tv['quoc_tich'], "Dân tộc": tv['dan_toc'],
                                            "Quan hệ chủ hộ": tv['quan_he'], "Tỉnh/TP": tv['tinh'], "Phường/Xã": tv['phuong_xa']})
                        df_tv = pd.DataFrame(tv_data)
                        st.dataframe(df_tv, width='stretch', hide_index=True)
                        col_del1, col_del2, col_del3 = st.columns([1,1,1])
                        with col_del2:
                            tv_to_delete = st.number_input("Nhập STT thành viên cần xóa:", min_value=1, max_value=len(st.session_state.bhxh_family_members), step=1, key="tv_delete_family")
                            if st.button("🗑️ Xóa thành viên", key="btn_del_tv_family", disabled=not can_edit()):
                                st.session_state.bhxh_family_members.pop(tv_to_delete - 1)
                                st.rerun()
                
                    with st.form(key=f"bhxh_family_form_{nv_id}"):
                        st.markdown("**I. THÔNG TIN CHỦ HỘ:**")
                        col1, col2 = st.columns(2)
                        with col1:
                            ho_ten_chu_ho = st.text_input("Họ và tên chủ hộ", value=nv_data.get('ho_ten_chu_ho', '') if nv_data else '')
                            so_cccd_chu_ho = st.text_input("Số CCCD chủ hộ", value=nv_data.get('so_cccd_chu_ho', '') if nv_data else '')
                            tinh_chu_ho_index = 0
                            tinh_chu_ho_current = nv_data.get('tinh_thanh_pho_chu_ho', '') if nv_data else ''
                            if tinh_chu_ho_current in tinh_options:
                                tinh_chu_ho_index = list(tinh_options.keys()).index(tinh_chu_ho_current) + 1
                            tinh_chu_ho = st.selectbox("Tỉnh/Thành phố (chủ hộ)", options=[""] + list(tinh_options.keys()), index=tinh_chu_ho_index)
                        with col2:
                            phuong_xa_options = []
                            phuong_xa_current = nv_data.get('phuong_xa_chu_ho', '') if nv_data else ''
                            if tinh_chu_ho and tinh_chu_ho != "":
                                ma_tinh = tinh_options.get(tinh_chu_ho)
                                db_temp2 = st.session_state.db_engine.get_connection()
                                c_temp2 = db_temp2.cursor()
                                c_temp2.execute("SELECT ten_xa FROM danh_muc_phuong_xa WHERE ma_tinh = %s ORDER BY ten_xa", (ma_tinh,))
                                phuong_xa_options = [row[0] for row in c_temp2.fetchall()]
                                db_temp2.close()
                            phuong_xa_index = 0
                            if phuong_xa_current in phuong_xa_options:
                                phuong_xa_index = phuong_xa_options.index(phuong_xa_current) + 1
                            phuong_xa_chu_ho = st.selectbox("Phường/Xã (chủ hộ)", options=[""] + phuong_xa_options, index=phuong_xa_index)
                    
                        st.markdown("**Thông tin thường trú:**")
                        col_tt1, col_tt2 = st.columns(2)
                        with col_tt1:
                            tinh_thuong_tru_index = 0
                            tinh_thuong_tru_current = nv_data.get('tinh_thanh_pho_thuong_tru', '') if nv_data else ''
                            if tinh_thuong_tru_current in tinh_options:
                                tinh_thuong_tru_index = list(tinh_options.keys()).index(tinh_thuong_tru_current) + 1
                            tinh_thuong_tru = st.selectbox("Tỉnh/Thành phố thường trú", options=[""] + list(tinh_options.keys()), index=tinh_thuong_tru_index)
                            ma_tinh_thuong_tru = tinh_options.get(tinh_thuong_tru, "") if tinh_thuong_tru else ""
                        with col_tt2:
                            phuong_xa_tt_options = []
                            phuong_xa_tt_current = nv_data.get('phuong_xa_thuong_tru', '') if nv_data else ''
                            if tinh_thuong_tru and tinh_thuong_tru != "":
                                ma_tinh_tt = tinh_options.get(tinh_thuong_tru)
                                db_temp3 = st.session_state.db_engine.get_connection()
                                c_temp3 = db_temp3.cursor()
                                c_temp3.execute("SELECT ten_xa, ma_xa FROM danh_muc_phuong_xa WHERE ma_tinh = %s ORDER BY ten_xa", (ma_tinh_tt,))
                                phuong_xa_tt_options = c_temp3.fetchall()
                                db_temp3.close()
                            phuong_xa_tt_index = 0
                            ma_phuong_xa_thuong_tru = ""
                            for i, px in enumerate(phuong_xa_tt_options):
                                if px[0] == phuong_xa_tt_current:
                                    phuong_xa_tt_index = i + 1
                                    ma_phuong_xa_thuong_tru = px[1]
                                    break
                            phuong_xa_display_list = [""] + [px[0] for px in phuong_xa_tt_options]
                            phuong_xa_thuong_tru = st.selectbox("Phường/Xã thường trú", options=phuong_xa_display_list, index=phuong_xa_tt_index)
                            for px in phuong_xa_tt_options:
                                if px[0] == phuong_xa_thuong_tru:
                                    ma_phuong_xa_thuong_tru = px[1]
                                    break
                    
                        st.markdown("**II. THÊM THÀNH VIÊN MỚI:**")
                        st.caption("Điền thông tin vào các cột bên dưới, sau đó bấm '➕ Thêm thành viên'")
                        col_tv1, col_tv2, col_tv3, col_tv4, col_tv5, col_tv6, col_tv7, col_tv8 = st.columns([2,1.3,1,1,1,1.5,1.8,1.8])
                        with col_tv1:
                            ho_ten_tv = st.text_input("Họ và tên", key="tv_ho_ten_family", placeholder="Nguyễn Văn A")
                        with col_tv2:
                            ngay_sinh_tv = st.text_input("Ngày sinh", key="tv_ngay_sinh_family", placeholder="dd/mm/yyyy")
                        with col_tv3:
                            gioi_tinh_tv = st.selectbox("Giới tính", ["Nam", "Nữ"], key="tv_gioi_tinh_family")
                        with col_tv4:
                            quoc_tich_tv = st.text_input("Quốc tịch", value="Việt Nam", key="tv_quoc_tich_family")
                        with col_tv5:
                            dan_toc_tv = st.text_input("Dân tộc", value="Kinh", key="tv_dan_toc_family")
                        with col_tv6:
                            quan_he_tv = st.selectbox("Quan hệ chủ hộ", ["", "Vợ", "Chồng", "Con", "Bố", "Mẹ", "Anh", "Chị", "Em", "Ông", "Bà", "Khác"], key="tv_quan_he_family")
                        with col_tv7:
                            tinh_tv = st.selectbox("Tỉnh/Thành phố", options=[""] + list(tinh_options.keys()), key="tv_tinh_family")
                        with col_tv8:
                            phuong_xa_tv_options = []
                            if tinh_tv and tinh_tv != "":
                                ma_tinh_tv = tinh_options.get(tinh_tv)
                                db_temp4 = st.session_state.db_engine.get_connection()
                                c_temp4 = db_temp4.cursor()
                                c_temp4.execute("SELECT ten_xa FROM danh_muc_phuong_xa WHERE ma_tinh = %s ORDER BY ten_xa", (ma_tinh_tv,))
                                phuong_xa_tv_options = [row[0] for row in c_temp4.fetchall()]
                                db_temp4.close()
                            phuong_xa_tv = st.selectbox("Phường/Xã", options=[""] + phuong_xa_tv_options, key="tv_phuong_xa_family")
                    
                        col_btn_add1, col_btn_add2, col_btn_add3 = st.columns([1,1,1])
                        with col_btn_add2:
                            if st.form_submit_button("➕ Thêm thành viên vào danh sách", width='stretch', disabled=not can_edit()):
                                if ho_ten_tv:
                                    st.session_state.bhxh_family_members.append({
                                        'ho_ten': ho_ten_tv, 'ngay_sinh': parse_date(ngay_sinh_tv), 'gioi_tinh': gioi_tinh_tv,
                                        'quoc_tich': quoc_tich_tv, 'dan_toc': dan_toc_tv, 'quan_he': quan_he_tv,
                                        'tinh': tinh_tv, 'phuong_xa': phuong_xa_tv
                                    })
                                    st.rerun()
                                else:
                                    st.error("Vui lòng nhập họ tên thành viên")
                    
                        st.markdown("---")
                        col_save1, col_save2, col_save3 = st.columns([1,2,1])
                        with col_save2:
                            if st.form_submit_button("💾 LƯU THÔNG TIN CHỦ HỘ", width='stretch', type="primary", disabled=not can_edit()):
                                try:
                                    db_luu = st.session_state.db_engine.get_connection()
                                    c_luu = db_luu.cursor()
                                    c_luu.execute("""UPDATE nhan_vien SET ho_ten_chu_ho=%s, so_cccd_chu_ho=%s, tinh_thanh_pho_chu_ho=%s, phuong_xa_chu_ho=%s,
                                        tinh_thanh_pho_thuong_tru=%s, ma_tinh_thuong_tru=%s, phuong_xa_thuong_tru=%s, ma_phuong_xa_thuong_tru=%s WHERE id=%s""",
                                        (ho_ten_chu_ho, so_cccd_chu_ho, tinh_chu_ho, phuong_xa_chu_ho, tinh_thuong_tru, ma_tinh_thuong_tru, phuong_xa_thuong_tru, ma_phuong_xa_thuong_tru, nv_id))
                                    c_luu.execute("DELETE FROM phu_luc_gia_dinh WHERE nhan_vien_id = %s", (nv_id,))
                                    for tv in st.session_state.bhxh_family_members:
                                        c_luu.execute("""INSERT INTO phu_luc_gia_dinh (nhan_vien_id, ho_ten, ngay_sinh, gioi_tinh, quoc_tich, dan_toc, quan_he_voi_chu_ho, tinh_thanh_pho, phuong_xa) 
                                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                                            (nv_id, tv['ho_ten'], tv['ngay_sinh'], tv['gioi_tinh'], tv['quoc_tich'], tv['dan_toc'], tv['quan_he'], tv['tinh'], tv['phuong_xa']))
                                    db_luu.commit()
                                    db_luu.close()
                                    del st.session_state['bhxh_family_nv_id']
                                    del st.session_state['bhxh_family_nv_name']
                                    del st.session_state['bhxh_family_members']
                                    st.success(f"✅ Đã lưu thông tin hộ gia đình cho nhân viên {nv_name}")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Lỗi khi lưu: {e}")
                
                    col_cancel1, col_cancel2, col_cancel3 = st.columns([1,2,1])
                    with col_cancel2:
                        if st.button("❌ HỦY BỎ", width='stretch'):
                            del st.session_state['bhxh_family_nv_id']
                            del st.session_state['bhxh_family_nv_name']
                            if 'bhxh_family_members' in st.session_state:
                                del st.session_state['bhxh_family_members']
                            st.rerun()
    
    with tab_da_nghi:
        st.caption("📋 Danh sách nhân viên đã nghỉ việc (có thông tin ngày nghỉ)")
        
        col_filter1, col_filter2, col_filter3 = st.columns([2, 2, 1])
        with col_filter1:
            search_nghi = st.text_input("🔍 Tìm kiếm (Tên, Mã NV, SĐT, CCCD)", key="search_da_nghi")
        with col_filter2:
            db_temp = st.session_state.db_engine.get_connection()
            c_temp = db_temp.cursor()
            c_temp.execute("SELECT DISTINCT EXTRACT(YEAR FROM ngay_ket_thuc) as nam FROM nhan_vien WHERE trang_thai='NGHI_VIEC' AND ngay_ket_thuc IS NOT NULL ORDER BY nam DESC")
            years = [int(row[0]) for row in c_temp.fetchall() if row[0] is not None]
            db_temp.close()
            filter_nam = st.selectbox("📅 Lọc theo năm nghỉ:", ["Tất cả"] + [str(y) for y in years] if years else ["Tất cả"])
        
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        sql = """
            SELECT id, ma_nv, ho_ten, ngay_sinh, gioi_tinh, so_cccd, dien_thoai, 
                   chuc_danh_nghe, loai_hop_dong, so_hdld, ngay_vao_lam, ngay_ket_thuc,
                   ma_so_bhxh, thang_bat_dau_bh, ly_do_nghi, ten_don_vi_thu_huong
            FROM nhan_vien 
            WHERE trang_thai = 'NGHI_VIEC'
        """
        params = []
        
        if search_nghi:
            sql += " AND (ho_ten LIKE %s OR ma_nv LIKE %s OR dien_thoai LIKE %s OR so_cccd LIKE %s)"
            params.extend([f'%{search_nghi}%'] * 4)
        
        if filter_nam != "Tất cả" and filter_nam.isdigit():
            sql += " AND EXTRACT(YEAR FROM ngay_ket_thuc) = %s"
            params.append(int(filter_nam))
        
        sql += " ORDER BY ngay_ket_thuc DESC, id DESC"
        c.execute(sql, tuple(params))
        ds_nghi = c.fetchall()
        db.close()
        
        if ds_nghi:
            df_nghi = pd.DataFrame(ds_nghi)
            for col in df_nghi.columns:
                if 'ngay' in col.lower():
                    df_nghi[col] = df_nghi[col].apply(format_date)
            
            display_cols_nghi = ['ma_nv', 'ho_ten', 'ngay_sinh', 'gioi_tinh', 'chuc_danh_nghe', 
                                 'so_hdld', 'ngay_vao_lam', 'ngay_ket_thuc', 'dien_thoai', 'ma_so_bhxh', 'ten_don_vi_thu_huong']
            available_cols_nghi = [c for c in display_cols_nghi if c in df_nghi.columns]
            df_show_nghi = df_nghi[available_cols_nghi]
            
            col_map_nghi = {
                'ma_nv': 'Mã NV',
                'ho_ten': 'Họ và tên',
                'ngay_sinh': 'Ngày sinh',
                'gioi_tinh': 'Giới tính',
                'chuc_danh_nghe': 'Chức danh',
                'so_hdld': 'Số HĐLĐ',
                'ngay_vao_lam': 'Ngày vào làm',
                'ngay_ket_thuc': '📅 Ngày nghỉ việc',
                'dien_thoai': 'SĐT',
                'ma_so_bhxh': 'Mã BHXH',
                'ten_don_vi_thu_huong': 'Tên đơn vị thụ hưởng',
            }
            df_show_nghi.rename(columns=col_map_nghi, inplace=True)
            
            st.caption(f"📌 Tổng số: **{len(ds_nghi)}** nhân viên đã nghỉ việc")
            st.dataframe(df_show_nghi, width='stretch', hide_index=True, height=400)
            
            st.divider()
            st.subheader("🔍 Xem chi tiết / Khôi phục nhân viên")
            
            nv_options = {f"{nv['ma_nv']} - {nv['ho_ten']} (Nghỉ: {format_date(nv.get('ngay_ket_thuc'))})": nv['id'] for nv in ds_nghi}
            selected_nghi_name = st.selectbox("Chọn nhân viên đã nghỉ:", list(nv_options.keys()), help="💡 Gõ mã NV hoặc tên để tìm nhanh trong danh sách")
            selected_nghi_id = nv_options[selected_nghi_name]
            
            db = st.session_state.db_engine.get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c.execute("SELECT * FROM nhan_vien WHERE id = %s", (selected_nghi_id,))
            nv_nghi_detail = c.fetchone()
            db.close()
            
            if nv_nghi_detail:
                col_detail1, col_detail2 = st.columns(2)
                with col_detail1:
                    st.markdown("**📋 Thông tin cá nhân**")
                    st.write(f"- **Mã NV:** {nv_nghi_detail.get('ma_nv', '')}")
                    st.write(f"- **Họ tên:** {nv_nghi_detail.get('ho_ten', '')}")
                    st.write(f"- **Ngày sinh:** {format_date(nv_nghi_detail.get('ngay_sinh'))}")
                    st.write(f"- **Giới tính:** {nv_nghi_detail.get('gioi_tinh', '')}")
                    st.write(f"- **CCCD:** {nv_nghi_detail.get('so_cccd', '')}")
                    st.write(f"- **SĐT:** {nv_nghi_detail.get('dien_thoai', '')}")
                    st.write(f"- **Chức danh:** {nv_nghi_detail.get('chuc_danh_nghe', '')}")
                    st.write(f"- **Tên đơn vị thụ hưởng:** {nv_nghi_detail.get('ten_don_vi_thu_huong', '')}")
                
                with col_detail2:
                    st.markdown("**📅 Thông tin hợp đồng & nghỉ việc**")
                    st.write(f"- **Số HĐLĐ:** {nv_nghi_detail.get('so_hdld', '')}")
                    st.write(f"- **Loại HĐ:** {nv_nghi_detail.get('loai_hop_dong', '')}")
                    st.write(f"- **Ngày vào làm:** {format_date(nv_nghi_detail.get('ngay_vao_lam'))}")
                    st.write(f"- **📅 Ngày nghỉ việc:** **{format_date(nv_nghi_detail.get('ngay_ket_thuc'))}**")
                    st.write(f"- **Mã BHXH:** {nv_nghi_detail.get('ma_so_bhxh', '')}")
                    st.write(f"- **Lý do nghỉ:** {nv_nghi_detail.get('ly_do_nghi', 'Chưa có thông tin')}")
                
                if st.session_state.role in ("admin", "xem_toan_bo"):
                    st.divider()
                    col_restore1, col_restore2, col_restore3 = st.columns([1, 2, 1])
                    with col_restore2:
                        if st.button(f"🔄 KHÔI PHỤC NHÂN VIÊN - {nv_nghi_detail['ho_ten']}", width='stretch', type="primary"):
                            try:
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                loai_hd = nv_nghi_detail.get('loai_hop_dong', '')
                                if loai_hd == 'Thử việc':
                                    trang_thai_moi = 'THU_VIEC'
                                else:
                                    trang_thai_moi = 'DANG_LAM'
                                c.execute("""
                                    UPDATE nhan_vien 
                                    SET trang_thai = %s, 
                                        ngay_ket_thuc = NULL
                                    WHERE id = %s
                                """, (trang_thai_moi, selected_nghi_id))
                                db.commit()
                                db.close()
                                st.success(f"✅ Đã khôi phục nhân viên {nv_nghi_detail['ho_ten']}!")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi khi khôi phục: {e}")
        else:
            st.info("📭 Không có nhân viên nào đã nghỉ việc")
    
    with tab_qtct:
        st.caption("📜 Lịch sử công tác và quyết định nhân sự")
        
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT id, ma_nv, ho_ten FROM nhan_vien ORDER BY id DESC")
        all_nv = c.fetchall()
        db.close()
        
        if all_nv:
            nv_options = {f"{x['ma_nv']} - {x['ho_ten']}": x['id'] for x in all_nv}
            selected_nv_history = st.selectbox("🔍 Chọn nhân viên:", list(nv_options.keys()), key="history_nv", help="💡 Gõ mã NV hoặc tên để tìm nhanh trong danh sách")
            nv_id_history = nv_options[selected_nv_history]
            
            db = st.session_state.db_engine.get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c.execute("SELECT * FROM nhan_vien WHERE id = %s", (nv_id_history,))
            nv_current = c.fetchone()
            
            st.markdown(f"""
            ### 📌 Thông tin hiện tại của {nv_current['ho_ten']} ({nv_current['ma_nv']})
            | Trường | Giá trị |
            |--------|---------|
            | Trạng thái | {'🟢 Đang làm' if nv_current['trang_thai'] == 'DANG_LAM' else '🔵 Thử việc' if nv_current['trang_thai'] == 'THU_VIEC' else '🔴 Đã nghỉ'} |
            | Loại hợp đồng | {nv_current['loai_hop_dong']} |
            | Ngày vào làm | {format_date(nv_current['ngay_vao_lam'])} |
            | Ngày chính thức | {format_date(nv_current.get('ngay_chinh_thuc')) or 'Chưa có'} |
            | Chức danh | {nv_current['chuc_danh_nghe']} |
            | Phòng ban | {nv_current['phong_ban_lam_viec']} |
            """)
            
            c.execute("""
                SELECT * FROM quyet_dinh_nhan_su 
                WHERE nhan_vien_id = %s 
                ORDER BY ngay_quyet_dinh DESC
            """, (nv_id_history,))
            quyet_dinh_list = c.fetchall()
            
            if quyet_dinh_list:
                st.markdown("### 📋 Các quyết định nhân sự")
                loai_qd_map = {
                    'THU_VIEC': '📝 Quyết định thử việc',
                    'CHINH_THUC': '✅ Quyết định chính thức',
                    'DIEU_CHUYEN': '🔄 Quyết định điều chuyển',
                    'BO_NHIEM': '⭐ Quyết định bổ nhiệm',
                    'MIEN_NHIEM': '⛔ Quyết định miễn nhiệm',
                    'DOI_CHUC_DANH': '🔁 Quyết định đổi chức danh',
                    'CHAM_DUT_HD': '📄 Quyết định chấm dứt HĐ',
                    'TANG_LUONG': '💰 Quyết định tăng lương',
                    'NGHI_VIEC': '🚫 Quyết định nghỉ việc'
                }

                # Bảng ghi cũ được tạo qua 2 luồng khác nhau với 2 bộ tên cột khác nhau
                # (loai_qd/so_qd/ngay_qd -- cột "gốc" của tab Quyết định nhân sự -- và
                # loai_quyet_dinh/so_quyet_dinh/ngay_quyet_dinh -- cột "hiển thị" cũ hơn).
                # COALESCE ở Python để mọi bản ghi (cũ lẫn mới) đều hiển thị đúng, không còn "None".
                def _qd_get(qd, *keys):
                    for k in keys:
                        v = qd.get(k)
                        if v not in (None, ""):
                            return v
                    return None

                qd_data = []
                for i, qd in enumerate(quyet_dinh_list, 1):
                    loai = _qd_get(qd, 'loai_quyet_dinh', 'loai_qd')
                    ngay_qd_disp = _qd_get(qd, 'ngay_quyet_dinh', 'ngay_qd')
                    ngay_hl_disp = _qd_get(qd, 'ngay_hieu_luc', 'ngay_quyet_dinh', 'ngay_qd')
                    so_qd_disp = _qd_get(qd, 'so_quyet_dinh', 'so_qd')
                    noi_dung_disp = qd.get('noi_dung') or ''
                    qd_data.append({
                        "STT": i,
                        "Loại quyết định": loai_qd_map.get(loai, loai or '...'),
                        "Số quyết định": so_qd_disp or '...',
                        "Ngày quyết định": format_date(ngay_qd_disp),
                        "Ngày hiệu lực": format_date(ngay_hl_disp),
                        "Nội dung": (noi_dung_disp[:50] + "...") if len(noi_dung_disp) > 50 else (noi_dung_disp or '...')
                    })
                df_qd = pd.DataFrame(qd_data)
                st.dataframe(df_qd, width='stretch', hide_index=True)

                with st.expander("🔍 Xem chi tiết quyết định"):
                    qd_options = {
                        f"{format_date(_qd_get(qd, 'ngay_quyet_dinh', 'ngay_qd'))} - {loai_qd_map.get(_qd_get(qd, 'loai_quyet_dinh', 'loai_qd'), _qd_get(qd, 'loai_quyet_dinh', 'loai_qd'))}": qd
                        for qd in quyet_dinh_list
                    }
                    selected_qd_name = st.selectbox("Chọn quyết định:", list(qd_options.keys()), key="qd_detail")
                    selected_qd = qd_options[selected_qd_name]
                    nguoi_ky_disp = _qd_get(selected_qd, 'nguoi_ky') or COMPANY_CONFIG.get('dai_dien') or 'GIÁM ĐỐC'
                    st.markdown(f"""
                    **📄 Chi tiết quyết định:**
                    - **Số quyết định:** {_qd_get(selected_qd, 'so_quyet_dinh', 'so_qd') or '...'}
                    - **Ngày quyết định:** {format_date(_qd_get(selected_qd, 'ngay_quyet_dinh', 'ngay_qd'))}
                    - **Ngày hiệu lực:** {format_date(_qd_get(selected_qd, 'ngay_hieu_luc', 'ngay_quyet_dinh', 'ngay_qd'))}
                    - **Loại quyết định:** {loai_qd_map.get(_qd_get(selected_qd, 'loai_quyet_dinh', 'loai_qd'), _qd_get(selected_qd, 'loai_quyet_dinh', 'loai_qd') or '...')}
                    - **Nội dung:** {selected_qd.get('noi_dung') or '...'}
                    - **Người ký:** {nguoi_ky_disp}
                    
                    **📊 Thay đổi:**
                    - Chức danh: {selected_qd.get('chuc_danh_cu') or '...'} → {selected_qd.get('chuc_danh_moi') or '...'}
                    - Phòng ban: {selected_qd.get('phong_ban_cu') or '...'} → {selected_qd.get('phong_ban_moi') or '...'}
                    - Loại HĐ: {selected_qd.get('loai_hop_dong_cu') or '...'} → {selected_qd.get('loai_hop_dong_moi') or '...'}
                    """)
            else:
                st.info("📭 Nhân viên này chưa có quyết định nào.")
            
            c.execute("""
                SELECT * FROM lich_su_cong_tac 
                WHERE nhan_vien_id = %s 
                ORDER BY tu_ngay ASC
            """, (nv_id_history,))
            lich_su_list = c.fetchall()
            
            if lich_su_list:
                st.markdown("### 📅 Lịch sử công tác")
                ls_data = []
                for i, ls in enumerate(lich_su_list, 1):
                    ls_data.append({
                        "STT": i,
                        "Từ ngày": format_date(ls['tu_ngay']),
                        "Đến ngày": format_date(ls['den_ngay']) if ls['den_ngay'] else "Đang làm",
                        "Chức danh": ls['chuc_danh'] or '',
                        "Phòng ban": ls['phong_ban'] or '',
                        "Loại HĐ": ls['loai_hop_dong'] or '',
                        "Hệ số lương": ls['he_so_luong'] or ''
                    })
                df_ls = pd.DataFrame(ls_data)
                st.dataframe(df_ls, width='stretch', hide_index=True, height=400)
            else:
                st.info("📭 Chưa có lịch sử công tác. Đang tạo từ dữ liệu hiện tại...")
                loai_hd_dung = nv_current['loai_hop_dong']
                if nv_current['trang_thai'] == 'THU_VIEC':
                    loai_hd_dung = 'Thử việc'
                c.execute("""
                    INSERT INTO lich_su_cong_tac (nhan_vien_id, tu_ngay, chuc_danh, phong_ban, noi_lam_viec, loai_hop_dong, he_so_luong)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (nv_id_history, nv_current['ngay_vao_lam'], nv_current['chuc_danh_nghe'], 
                      nv_current['phong_ban_lam_viec'], nv_current['noi_lam_viec'], 
                      loai_hd_dung, nv_current['he_so_luong']))
                db.commit()
                st.rerun()
            db.close()
        else:
            st.info("⚠️ Chưa có nhân viên nào trong hệ thống!")
    
    # ========== PHẦN XÓA NHÂN VIÊN THEO SỐ HĐ ==========
    st.divider()
    
    if st.session_state.role in ("admin", "xem_toan_bo"):
        with st.expander("🗑️ CÔNG CỤ XÓA NHÂN VIÊN (CHỈ DÀNH CHO ADMIN)", expanded=False):
            st.warning("⚠️ **CẢNH BÁO:** Thao tác này sẽ XÓA VĨNH VIỄN nhân viên và tất cả dữ liệu liên quan!")
        
            col_hd1, col_hd2 = st.columns([2, 1])
            with col_hd1:
                so_hd_can_xoa = st.text_input("📝 Nhập số hợp đồng cần xóa (VD: 21/2026/HĐTV-CHL):", key="so_hd_xoa")
            with col_hd2:
                st.write("")
                st.write("")
                xac_nhan_xoa = st.checkbox("✅ Tôi xác nhận muốn xóa vĩnh viễn", key="xac_nhan_xoa_nv")
        
            if so_hd_can_xoa and xac_nhan_xoa:
                try:
                    db_check = st.session_state.db_engine.get_connection()
                    c_check = db_check.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    c_check.execute("SELECT id, ho_ten, ma_nv, trang_thai FROM nhan_vien WHERE so_hdld = %s", (so_hd_can_xoa,))
                    nv_info = c_check.fetchone()
                    db_check.close()
                
                    if nv_info:
                        st.warning(f"⚠️ Nhân viên: **{nv_info['ho_ten']}** (Mã: {nv_info['ma_nv']}) - Trạng thái: {nv_info['trang_thai']}")
                    
                        # Đếm số bản ghi liên quan
                        db_count = st.session_state.db_engine.get_connection()
                        c_count = db_count.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                        c_count.execute("SELECT COUNT(*) as count FROM lich_su_cong_tac WHERE nhan_vien_id = %s", (nv_info['id'],))
                        ls_count = c_count.fetchone()['count']
                        c_count.execute("SELECT COUNT(*) as count FROM quyet_dinh_nhan_su WHERE nhan_vien_id = %s", (nv_info['id'],))
                        qd_count = c_count.fetchone()['count']
                        c_count.execute("SELECT COUNT(*) as count FROM ho_so_nhan_vien WHERE nhan_vien_id = %s", (nv_info['id'],))
                        hs_count = c_count.fetchone()['count']
                        c_count.execute("SELECT COUNT(*) as count FROM phu_luc_gia_dinh WHERE nhan_vien_id = %s", (nv_info['id'],))
                        pl_count = c_count.fetchone()['count']
                        db_count.close()
                    
                        st.info(f"📊 Sẽ xóa: {ls_count} lịch sử công tác, {qd_count} quyết định, {hs_count} hồ sơ, {pl_count} phụ lục gia đình")
                    
                        # Xác nhận lần cuối
                        xac_nhan_cuoi = st.checkbox("⚠️ Tôi hiểu rủi ro và muốn xóa VĨNH VIỄN nhân viên này", key="xac_nhan_cuoi_xoa")
                    
                        # Tìm phần xóa nhân viên (khoảng dòng 1900)
                        if xac_nhan_cuoi:
                            col_confirm1, col_confirm2, col_confirm3 = st.columns([1, 2, 1])
                            with col_confirm2:
                                if st.button("🗑️ XÁC NHẬN XÓA VĨNH VIỄN", type="primary", key="btn_confirm_xoa", disabled=not can_edit()):
                                    if not can_delete():
                                        st.error("❌ Bạn không có quyền xóa dữ liệu!")
                                    else:
                                        try:
                                            db = st.session_state.db_engine.get_connection()
                                            cur = db.cursor()
                                        
                                            cur.execute("DELETE FROM lich_su_cong_tac WHERE nhan_vien_id = %s", (nv_info['id'],))
                                            cur.execute("DELETE FROM quyet_dinh_nhan_su WHERE nhan_vien_id = %s", (nv_info['id'],))
                                            cur.execute("DELETE FROM ho_so_nhan_vien WHERE nhan_vien_id = %s", (nv_info['id'],))
                                            cur.execute("DELETE FROM phu_luc_gia_dinh WHERE nhan_vien_id = %s", (nv_info['id'],))
                                            cur.execute("DELETE FROM nhan_vien WHERE id = %s", (nv_info['id'],))
                                        
                                            db.commit()
                                            db.close()
                                        
                                            st.success(f"✅ Đã XÓA VĨNH VIỄN nhân viên {nv_info['ho_ten']} (Mã: {nv_info['ma_nv']})")
                                            st.balloons()
                                            st.cache_data.clear()
                                            st.rerun()
                                        
                                        except Exception as e:
                                            st.error(f"❌ Lỗi khi xóa: {str(e)}")
                                            try:
                                                db.rollback()
                                                db.close()
                                            except:
                                                pass
                        else:
                            st.info("🔒 **Vui lòng tick vào ô xác nhận 'Tôi hiểu rủi ro...' để kích hoạt nút xóa**")
                    else:
                        st.error(f"❌ Không tìm thấy nhân viên có số hợp đồng: {so_hd_can_xoa}")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi khi tìm kiếm: {e}")
                
            elif so_hd_can_xoa and not xac_nhan_xoa:
                st.info("🔒 Vui lòng tick xác nhận 'Tôi xác nhận muốn xóa vĩnh viễn' để tiếp tục")
    
    # ===== TAB: QUYẾT ĐỊNH NHÂN SỰ =====
    with tab_qdns:
        st.caption("📜 Ra các Quyết định nhân sự: Bổ nhiệm, Miễn nhiệm, Thay đổi chức danh, Điều chuyển công tác, Chấm dứt HĐTV/HĐLĐ")
        _chedo_qdns = st.radio("Chọn thao tác:", ["➕ Tạo QĐNS mới", "🔍 Tra cứu lịch sử"],
                               horizontal=True, key="qdns_thao_tac")


        # Thông báo file vừa tạo (nếu có), hiển thị TRƯỚC form để không bị mất sau khi rerun
        if _chedo_qdns == "➕ Tạo QĐNS mới":
            if st.session_state.get('qdns_last_file'):
                st.success(f"✅ Đã tạo {st.session_state.get('qdns_last_label','Quyết định')} số {st.session_state.get('qdns_last_so')}")
                st.cache_data.clear()
                try:
                    with open(st.session_state['qdns_last_file'], "rb") as f:
                        st.download_button(
                            label="📥 TẢI QUYẾT ĐỊNH (Word)",
                            data=f,
                            file_name=f"QDNS_{st.session_state.get('qdns_last_so','').replace('/', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_qdns_last"
                        )
                except Exception:
                    pass
                if st.button("✖️ Đóng thông báo", key="close_qdns_notice"):
                    for k in ['qdns_last_file', 'qdns_last_label', 'qdns_last_so']:
                        st.session_state.pop(k, None)
                    st.rerun()
                st.divider()

            db_qd = st.session_state.db_engine.get_connection()
            c_qd = db_qd.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c_qd.execute("""
                SELECT id, ma_nv, ho_ten, chuc_vu, chuc_danh_nghe, phong_ban_lam_viec, loai_hop_dong, trang_thai,
                       he_so_luong, noi_lam_viec
                FROM nhan_vien WHERE trang_thai IN ('DANG_LAM','THU_VIEC') ORDER BY ho_ten
            """)
            nv_qd_list = c_qd.fetchall()
            db_qd.close()
            # Danh sách phòng ban CHUẨN — không lấy DISTINCT trực tiếp từ dữ liệu nhan_vien nữa
            # (cách cũ sẽ tự nhân bản mọi biến thể sai chính tả/viết hoa đã lỡ lưu trong DB).
            ds_phong_ban = get_phong_ban_options()

            if not nv_qd_list:
                st.info("Không có nhân viên đang làm việc.")
            else:
                nv_qd_options = {f"{nv['ma_nv']} - {nv['ho_ten']}": nv for nv in nv_qd_list}
                chon_nv_label = st.selectbox("👤 Chọn nhân viên:", list(nv_qd_options.keys()), key="qdns_chon_nv",
                                              help="💡 Bấm vào ô rồi gõ tên/mã NV để lọc nhanh — không cần scroll (Ctrl+F trình duyệt không lọc được ô này).")
                nv_qd = nv_qd_options[chon_nv_label]

                col_info1, col_info2, col_info3 = st.columns(3)
                col_info1.markdown(f"**Chức vụ hiện tại:** {nv_qd.get('chuc_vu') or 'Nhân viên'}")
                col_info2.markdown(f"**Chức danh hiện tại:** {nv_qd.get('chuc_danh_nghe') or '-'}")
                col_info3.markdown(f"**Phòng ban hiện tại:** {nv_qd.get('phong_ban_lam_viec') or '-'}")

                col_qd1, col_qd2 = st.columns(2)
                with col_qd1:
                    loai_qd = st.selectbox(
                        "📋 Loại quyết định:",
                        list(LOAI_QDNS_LABEL.keys()),
                        format_func=lambda k: LOAI_QDNS_LABEL[k],
                        key="qdns_loai"
                    )
                with col_qd2:
                    ngay_qd = st.date_input("📅 Ngày ban hành quyết định:", value=date.today(), key="qdns_ngay")

                dieu1_lines = []
                tieu_de = ""
                hieu_luc_text = None
                gia_tri_truoc = None
                gia_tri_sau = None
                ok_to_submit = True

                ds_chuc_vu_tenant = get_chuc_vu_options()

                if loai_qd == 'BO_NHIEM':
                    chuc_vu_moi = st.selectbox("🏷️ Chức vụ được bổ nhiệm:", ds_chuc_vu_tenant, key="qdns_cv_bonhiem")
                    # Bổ nhiệm trong thực tế thường đi kèm luôn Chức danh & Phòng ban mới (VD:
                    # bổ nhiệm Trưởng phòng Kinh doanh = đổi cả chức vụ, chức danh lẫn phòng ban).
                    # Cho phép chỉnh, mặc định giữ nguyên giá trị hiện tại nếu không đổi.
                    cd_hien_tai_bn = nv_qd.get('chuc_danh_nghe') or ''
                    pb_hien_tai_bn = nv_qd.get('phong_ban_lam_viec') or ''
                    with st.expander("⚙️ Chức danh & Phòng ban kèm theo (tuỳ chọn — bỏ trống nếu không đổi)", expanded=False):
                        chuc_danh_moi_bn = st.text_input("💼 Chức danh mới:", value=cd_hien_tai_bn, key="qdns_cd_bonhiem")
                        idx_pb_bn = ds_phong_ban.index(pb_hien_tai_bn) if pb_hien_tai_bn in ds_phong_ban else 0
                        phong_ban_moi_bn = st.selectbox("🏢 Phòng ban mới:", ds_phong_ban, index=idx_pb_bn, key="qdns_pb_bonhiem") if ds_phong_ban else pb_hien_tai_bn
                    tieu_de = f"Bổ nhiệm chức vụ {chuc_vu_moi}"
                    dieu1_lines = [f"Bổ nhiệm Ông/Bà {nv_qd['ho_ten']} ({nv_qd['ma_nv']}) giữ chức vụ {chuc_vu_moi} kể từ ngày {ngay_qd.strftime('%d/%m/%Y')}."]
                    chi_tiet_bn = []
                    if chuc_danh_moi_bn.strip() and chuc_danh_moi_bn.strip() != cd_hien_tai_bn:
                        chi_tiet_bn.append(f"chức danh {chuc_danh_moi_bn.strip()}")
                    if phong_ban_moi_bn and phong_ban_moi_bn != pb_hien_tai_bn:
                        chi_tiet_bn.append(f"công tác tại {phong_ban_moi_bn}")
                    if chi_tiet_bn:
                        dieu1_lines.append(f"Đồng thời giữ {', '.join(chi_tiet_bn)} kể từ ngày {ngay_qd.strftime('%d/%m/%Y')}.")
                    gia_tri_truoc = nv_qd.get('chuc_vu') or 'Nhân viên'
                    gia_tri_sau = chuc_vu_moi

                elif loai_qd == 'MIEN_NHIEM':
                    cv_hien_tai = nv_qd.get('chuc_vu') or 'Nhân viên'
                    idx_mn = ds_chuc_vu_tenant.index(cv_hien_tai) if cv_hien_tai in ds_chuc_vu_tenant else 0
                    chuc_vu_mien = st.selectbox("🏷️ Chức vụ bị miễn nhiệm:", ds_chuc_vu_tenant, index=idx_mn, key="qdns_cv_miennhiem")
                    tieu_de = f"Miễn nhiệm chức vụ {chuc_vu_mien}"
                    dieu1_lines = [f"Miễn nhiệm chức vụ {chuc_vu_mien} đối với Ông/Bà {nv_qd['ho_ten']} ({nv_qd['ma_nv']}) kể từ ngày {ngay_qd.strftime('%d/%m/%Y')}."]
                    gia_tri_truoc = cv_hien_tai
                    gia_tri_sau = 'Nhân viên'
                    if cv_hien_tai == 'Nhân viên':
                        st.warning("⚠️ Nhân viên này hiện đang giữ chức vụ 'Nhân viên' (không có chức vụ quản lý để miễn nhiệm).")

                elif loai_qd == 'DOI_CHUC_DANH':
                    chuc_danh_moi = st.text_input("💼 Chức danh mới:", value=nv_qd.get('chuc_danh_nghe') or '', key="qdns_cd_moi")
                    tieu_de = f"Thay đổi chức danh - {nv_qd['ho_ten']}"
                    dieu1_lines = [f"Thay đổi chức danh của Ông/Bà {nv_qd['ho_ten']} ({nv_qd['ma_nv']}) từ '{nv_qd.get('chuc_danh_nghe') or ''}' thành '{chuc_danh_moi}' kể từ ngày {ngay_qd.strftime('%d/%m/%Y')}."]
                    gia_tri_truoc = nv_qd.get('chuc_danh_nghe') or ''
                    gia_tri_sau = chuc_danh_moi
                    if not chuc_danh_moi.strip():
                        ok_to_submit = False
                        st.error("⚠️ Vui lòng nhập chức danh mới.")

                elif loai_qd == 'DIEU_CHUYEN':
                    phong_hien_tai = nv_qd.get('phong_ban_lam_viec') or ''
                    st.text_input("🏢 Từ phòng ban:", value=phong_hien_tai, disabled=True, key="qdns_pb_tu")
                    # Chỉ cho chọn trong danh sách phòng ban CHUẨN — bỏ ô nhập tay tự do để
                    # tránh phát sinh biến thể mới không chuẩn hóa.
                    phong_moi = st.selectbox("🏢 Đến phòng ban:", ds_phong_ban, key="qdns_pb_den_select")
                    tieu_de = f"Điều chuyển công tác - {nv_qd['ho_ten']}"
                    dieu1_lines = [f"Điều chuyển Ông/Bà {nv_qd['ho_ten']} ({nv_qd['ma_nv']}) từ {phong_hien_tai or '(chưa xác định)'} sang {phong_moi or '(chưa xác định)'} kể từ ngày {ngay_qd.strftime('%d/%m/%Y')}."]
                    gia_tri_truoc = phong_hien_tai
                    gia_tri_sau = phong_moi
                    if not (phong_moi or '').strip():
                        ok_to_submit = False
                        st.error("⚠️ Vui lòng chọn hoặc nhập phòng ban đến.")

                elif loai_qd == 'CHUYEN_CHINH_THUC':
                    # Chỉ cho phép chuyển đổi NV đang ở trạng thái Thử việc
                    trang_thai_hien_tai = nv_qd.get('trang_thai', '')
                    loai_hd_hien_tai_ct = nv_qd.get('loai_hop_dong', '')
                    if trang_thai_hien_tai != 'THU_VIEC' and loai_hd_hien_tai_ct != 'Thử việc':
                        st.error("⚠️ Nhân viên này không ở trạng thái Thử việc — không thể chuyển đổi.")
                        ok_to_submit = False
                    else:
                        # Kiểm tra đã chuyển đổi chưa
                        da_chuyen, qd_cu = da_chuyen_doi_chinh_thuc(nv_qd['id'])
                        if da_chuyen:
                            st.warning("⚠️ Nhân viên này đã có QĐ chuyển chính thức trước đó.")

                        ma_cty_hd = st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL'

                        loai_hd_moi_lbl = st.selectbox(
                            "📑 Loại HĐLĐ mới:",
                            ["Không xác định thời hạn", "Xác định thời hạn 12 tháng",
                             "Xác định thời hạn 24 tháng", "Xác định thời hạn 36 tháng"],
                            key="qdns_loai_hd_moi"
                        )
                        if loai_hd_moi_lbl == "Không xác định thời hạn":
                            loai_hop_dong_luu_ct = "Không xác định thời hạn"
                            han_hd_thang_ct = None
                        else:
                            loai_hop_dong_luu_ct = "Xác định thời hạn"
                            han_hd_thang_ct = int(loai_hd_moi_lbl.split()[-2])

                        current_year_ct = datetime.now().year
                        db_temp_ct = st.session_state.db_engine.get_connection()
                        c_temp_ct = db_temp_ct.cursor()
                        c_temp_ct.execute("""
                            SELECT COALESCE(MAX(CAST(SPLIT_PART(so_hdld, '/', 1) AS INTEGER)), 0) as max_stt
                            FROM nhan_vien 
                            WHERE so_hdld LIKE %s 
                            AND SPLIT_PART(so_hdld, '/', 1) ~ '^[0-9]+$'
                            AND trang_thai = 'DANG_LAM'
                            AND loai_hop_dong != 'Thử việc'
                        """, (f'%/{current_year_ct}/HĐLĐ-%',))
                        result_ct = c_temp_ct.fetchone()
                        max_stt_ct = result_ct[0] if result_ct else 0
                        db_temp_ct.close()

                        next_stt_ct = max_stt_ct + 1
                        stt_str_ct = str(next_stt_ct).zfill(2)
                        so_hd_moi_ct = f"{stt_str_ct}/{current_year_ct}/HĐLĐ-{ma_cty_hd}"
                        st.info(f"📄 **Số HĐLĐ mới:** {so_hd_moi_ct} (tự động sinh)")

                        ngay_hieu_luc_ct = st.date_input(
                            "📅 Ngày hiệu lực (bắt đầu HĐLĐ):",
                            value=ngay_qd,
                            key="qdns_ngay_hl_ct"
                        )

                        if han_hd_thang_ct:
                            ngay_het_han_hd_ct = ngay_hieu_luc_ct + relativedelta(months=han_hd_thang_ct) - timedelta(days=1)
                            st.caption(f"📆 Hợp đồng sẽ hết hạn: {ngay_het_han_hd_ct.strftime('%d/%m/%Y')}")

                        ngay_bat_dau_bh_ct = tinh_thang_bat_dau_bh(ngay_hieu_luc_ct)
                        st.info(f"📅 Tháng bắt đầu đóng BHXH: **{format_thang_nam(ngay_bat_dau_bh_ct)}** (tự tính theo quy tắc 14 ngày)")

                        phuong_an_ct = st.selectbox(
                            "Phương án điều chỉnh BHXH",
                            [""] + PHUONG_AN_TANG,
                            key="qdns_pa_bhxh_ct",
                            help="Bắt buộc chọn — dùng cho báo tăng D02-LT"
                        )

                        so_hd_tv_cu_ct = nv_qd.get('so_hdld', '')
                        tieu_de = f"Chuyển đổi HĐTV sang HĐLĐ - {nv_qd['ho_ten']}"
                        dieu1_lines = [
                            f"Ông/Bà {nv_qd['ho_ten']} ({nv_qd['ma_nv']}) đã hoàn thành thời gian thử việc.",
                            f"Chuyển sang Hợp đồng lao động {loai_hd_moi_lbl.lower()} (số {so_hd_moi_ct}) kể từ ngày {ngay_hieu_luc_ct.strftime('%d/%m/%Y')}."
                        ]
                        hieu_luc_text = f"Quyết định có hiệu lực kể từ ngày {ngay_hieu_luc_ct.strftime('%d/%m/%Y')}."
                        gia_tri_truoc = 'Thử việc'
                        gia_tri_sau = loai_hop_dong_luu_ct

                elif loai_qd == 'CHAM_DUT_HD':
                    loai_hd_hien_tai = nv_qd.get('loai_hop_dong') or ''
                    if loai_hd_hien_tai == 'Thử việc':
                        nhan_hd = "Hợp đồng thử việc (HĐTV)"
                    else:
                        nhan_hd = "Hợp đồng lao động (HĐLĐ)"
                    st.info(f"🔎 Loại hợp đồng hiện tại: **{loai_hd_hien_tai or 'Chưa xác định'}** → Sẽ ban hành: **QĐ Chấm dứt {nhan_hd}**")
                    ly_do_cd = st.text_area("📝 Lý do chấm dứt:", key="qdns_lydo_cd", height=80,
                                              placeholder="VD: Hết hạn hợp đồng, Xin nghỉ theo nguyện vọng cá nhân, Chuyển công tác...")
                    tieu_de = f"Chấm dứt {nhan_hd} - {nv_qd['ho_ten']}"
                    dieu1_lines = [
                        f"Chấm dứt {nhan_hd} đối với Ông/Bà {nv_qd['ho_ten']} ({nv_qd['ma_nv']}) kể từ ngày {ngay_qd.strftime('%d/%m/%Y')}."
                    ]
                    if ly_do_cd.strip():
                        dieu1_lines.append(f"Lý do: {ly_do_cd.strip()}.")
                    hieu_luc_text = f"Ông/Bà {nv_qd['ho_ten']} có trách nhiệm bàn giao công việc, tài sản (nếu có) trước ngày {ngay_qd.strftime('%d/%m/%Y')}."
                    gia_tri_truoc = loai_hd_hien_tai
                    gia_tri_sau = 'NGHI_VIEC'

                st.divider()
                # Chỉ admin & hr được phép ra Quyết định nhân sự — các role khác (viewer...) bị làm mờ nút
                chi_admin_hr_qdns = st.session_state.role in ("admin", "hr")
                if not chi_admin_hr_qdns:
                    st.caption("🔒 Chỉ có admin & HR được phép sử dụng chức năng này!")
                if st.button("💾 TẠO QUYẾT ĐỊNH & LƯU", type="primary", width='stretch', key="qdns_submit", disabled=not (ok_to_submit and chi_admin_hr_qdns)):
                    try:
                        so_qd = generate_so_cong_van('QUYET_DINH')

                        file_path = tao_quyet_dinh_nhan_su(nv_qd, so_qd, ngay_qd, tieu_de, dieu1_lines, hieu_luc_text)
                        file_url = None
                        # (File Word được tạo để tải về ngay lập tức; bản ghi vẫn được lưu để tra cứu)

                        db_s = st.session_state.db_engine.get_connection()
                        c_s = db_s.cursor()

                        # Bảng quyet_dinh_nhan_su ở một số môi trường có thêm nhiều cột legacy
                        # (NOT NULL, không có default) song song với các cột chuẩn mà code dùng
                        # (so_qd, loai_qd, ngay_qd, nhan_vien_id, noi_dung...). Thay vì liệt kê cứng
                        # từng cột (dễ sót, như đã xảy ra với ngay_quyet_dinh rồi ngay_hieu_luc),
                        # ta tự dò TẤT CẢ cột NOT NULL không có default, rồi suy luận giá trị theo
                        # tên cột để điền cho đủ, tránh vướng lỗi NOT NULL dù DB có bao nhiêu cột
                        # legacy đi nữa.
                        # Ánh xạ gia_tri_truoc/gia_tri_sau sang đúng cặp cột "cũ/mới" mà tab
                        # "Lịch sử công tác" hiển thị (Chức danh / Phòng ban / Loại HĐ), tuỳ loại QĐ.
                        # Giá trị MỚI cuối cùng cho từng trường hồ sơ (chức danh / phòng ban / loại HĐ)
                        # sau quyết định này — tính chung cho MỌI loại QĐ (không chỉ DOI_CHUC_DANH/
                        # DIEU_CHUYEN/CHAM_DUT_HD như trước) để không sót trường nào cần cập nhật,
                        # kể cả trường hợp BO_NHIEM có kèm chức danh/phòng ban mới ở trên.
                        chuc_danh_nghe_cu = nv_qd.get('chuc_danh_nghe') or ''
                        phong_ban_cu = nv_qd.get('phong_ban_lam_viec') or ''
                        chuc_danh_nghe_moi_final = chuc_danh_nghe_cu
                        phong_ban_moi_final = phong_ban_cu
                        loai_hd_moi_final = nv_qd.get('loai_hop_dong') or ''

                        if loai_qd == 'BO_NHIEM':
                            if chuc_danh_moi_bn.strip():
                                chuc_danh_nghe_moi_final = chuc_danh_moi_bn.strip()
                            if phong_ban_moi_bn:
                                phong_ban_moi_final = phong_ban_moi_bn
                        elif loai_qd == 'DOI_CHUC_DANH':
                            chuc_danh_nghe_moi_final = gia_tri_sau
                        elif loai_qd == 'DIEU_CHUYEN':
                            phong_ban_moi_final = gia_tri_sau
                        elif loai_qd == 'CHUYEN_CHINH_THUC':
                            loai_hd_moi_final = gia_tri_sau  # loại HĐ mới (Xác định/Không xác định thời hạn)
                        elif loai_qd == 'CHAM_DUT_HD':
                            loai_hd_moi_final = 'Đã chấm dứt'
                            loai_hd_moi_final = 'Đã chấm dứt'

                        chuc_danh_cu_v = chuc_danh_moi_v = None
                        phong_ban_cu_v = phong_ban_moi_v = None
                        loai_hd_cu_v = loai_hd_moi_v = None
                        if chuc_danh_nghe_moi_final != chuc_danh_nghe_cu:
                            chuc_danh_cu_v, chuc_danh_moi_v = chuc_danh_nghe_cu, chuc_danh_nghe_moi_final
                        if phong_ban_moi_final != phong_ban_cu:
                            phong_ban_cu_v, phong_ban_moi_v = phong_ban_cu, phong_ban_moi_final
                        if loai_qd == 'CHAM_DUT_HD':
                            loai_hd_cu_v, loai_hd_moi_v = gia_tri_truoc, 'Đã chấm dứt'
                        nguoi_ky_v = COMPANY_CONFIG.get('dai_dien') or ''

                        cols = ['so_qd', 'loai_qd', 'nhan_vien_id', 'ngay_qd', 'noi_dung', 'gia_tri_truoc', 'gia_tri_sau', 'file_url', 'nguoi_tao',
                                'so_quyet_dinh', 'loai_quyet_dinh', 'ngay_quyet_dinh', 'ngay_hieu_luc', 'nguoi_ky',
                                'chuc_danh_cu', 'chuc_danh_moi', 'phong_ban_cu', 'phong_ban_moi', 'loai_hop_dong_cu', 'loai_hop_dong_moi']
                        vals = [so_qd, loai_qd, nv_qd['id'], ngay_qd, " ".join(dieu1_lines), gia_tri_truoc, gia_tri_sau, file_url, st.session_state.username,
                                so_qd, loai_qd, ngay_qd, ngay_qd, nguoi_ky_v,
                                chuc_danh_cu_v, chuc_danh_moi_v, phong_ban_cu_v, phong_ban_moi_v, loai_hd_cu_v, loai_hd_moi_v]

                        c_s.execute("""
                            SELECT column_name FROM information_schema.columns
                            WHERE table_name = 'quyet_dinh_nhan_su'
                              AND is_nullable = 'NO'
                              AND column_default IS NULL
                              AND column_name NOT IN %s
                        """, (tuple(cols),))
                        cot_not_null_con_thieu = [r[0] for r in c_s.fetchall()]

                        for ten_cot in cot_not_null_con_thieu:
                            tc = ten_cot.lower()
                            if 'ngay' in tc or 'ngay_hieu_luc' in tc:
                                gia_tri = ngay_qd
                            elif 'loai' in tc:
                                gia_tri = loai_qd
                            elif tc.startswith('so_') or tc == 'so' or 'so_quyet_dinh' in tc or 'so_hd' in tc:
                                gia_tri = so_qd
                            elif 'nhan_vien' in tc:
                                gia_tri = nv_qd['id']
                            elif 'noi_dung' in tc or 'dieu' in tc or 'trich_yeu' in tc:
                                gia_tri = " ".join(dieu1_lines)
                            elif 'nguoi' in tc:
                                gia_tri = st.session_state.username
                            elif 'trang_thai' in tc:
                                gia_tri = 'CO_HIEU_LUC'
                            else:
                                # Không đoán được ý nghĩa cột -> bỏ qua, để lỗi NOT NULL (nếu có)
                                # hiện rõ ràng thay vì điền giá trị sai lệch ngữ nghĩa.
                                continue
                            cols.append(ten_cot)
                            vals.append(gia_tri)

                        placeholders = ", ".join(["%s"] * len(cols))
                        c_s.execute(f"""
                            INSERT INTO quyet_dinh_nhan_su ({", ".join(cols)})
                            VALUES ({placeholders})
                        """, vals)

                        # Đăng ký vào hệ thống Quản lý công văn đi để cùng theo dõi số thứ tự
                        c_s.execute("""
                            INSERT INTO cong_van_di (so_cong_van, phong_phat_hanh, ngay_phat_hanh, tieu_de, trich_yeu, file_url, loai_cong_van, ghi_chu, nguoi_tao)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                        """, (so_qd, get_phong_to_chuc_nhan_su(), ngay_qd, f"{LOAI_QDNS_LABEL[loai_qd]}: {tieu_de}",
                              " ".join(dieu1_lines), file_url, 'QUYET_DINH', f"Quyết định nhân sự - NV: {nv_qd['ho_ten']}", st.session_state.username))

                        # Cập nhật hồ sơ nhân viên theo đúng logic từng loại quyết định
                        if loai_qd == 'BO_NHIEM':
                            c_s.execute("UPDATE nhan_vien SET chuc_vu = %s, ngay_qd_ns = %s WHERE id = %s", (gia_tri_sau, ngay_qd, nv_qd['id']))
                        elif loai_qd == 'MIEN_NHIEM':
                            c_s.execute("UPDATE nhan_vien SET chuc_vu = %s, ngay_qd_ns = %s WHERE id = %s", ('Nhân viên', ngay_qd, nv_qd['id']))
                        elif loai_qd == 'DOI_CHUC_DANH':
                            c_s.execute("UPDATE nhan_vien SET chuc_danh_nghe = %s, ngay_qd_ns = %s WHERE id = %s", (gia_tri_sau, ngay_qd, nv_qd['id']))
                        elif loai_qd == 'DIEU_CHUYEN':
                            c_s.execute("UPDATE nhan_vien SET phong_ban_lam_viec = %s, ngay_qd_ns = %s WHERE id = %s", (gia_tri_sau, ngay_qd, nv_qd['id']))
                        elif loai_qd == 'CHAM_DUT_HD':
                            # QUAN TRỌNG: phải đồng thời cập nhật trang_thai_bhxh và thang_ket_thuc_bh,
                            # nếu không báo cáo tăng/giảm BHXH (lọc theo 2 cột này) sẽ KHÔNG bao giờ
                            # thấy nhân sự này ở danh sách "Giảm", dù trang_thai đã là NGHI_VIEC.
                        
                            # Chỉ set GH nếu HĐLĐ không phải Thử việc
                            loai_hd_hien_tai = nv_qd.get('loai_hop_dong', '')
                            pa_giam = 'GH' if loai_hd_hien_tai != 'Thử việc' else ''
                            thang_pa_giam = ngay_qd.strftime('%m/%Y') if ngay_qd and pa_giam else None
                        
                            c_s.execute("""
                                UPDATE nhan_vien
                                SET trang_thai = 'NGHI_VIEC',
                                    ngay_ket_thuc = %s,
                                    ly_do_nghi = %s,
                                    trang_thai_bhxh = 'DA_BAO_GIAM',
                                    thang_ket_thuc_bh = %s,
                                    phuong_an_dieu_chinh = CASE WHEN %s != '' THEN %s ELSE phuong_an_dieu_chinh END,
                                    thang_phuong_an = CASE WHEN %s IS NOT NULL THEN %s ELSE thang_phuong_an END
                                WHERE id = %s
                            """, (ngay_qd, ly_do_cd if ly_do_cd.strip() else None, ngay_qd,
                                  pa_giam, pa_giam, thang_pa_giam, thang_pa_giam, nv_qd['id']))

                        elif loai_qd == 'CHUYEN_CHINH_THUC':
                            pa_ct_val = lay_ma_phuong_an(phuong_an_ct) if phuong_an_ct else None
                            # Cập nhật bảng nhan_vien
                            c_s.execute("""
                                UPDATE nhan_vien SET 
                                    trang_thai = 'DANG_LAM',
                                    loai_hop_dong = %s,
                                    han_hop_dong_thang = %s,
                                    so_hdld = %s,
                                    ngay_ky_hd = %s,
                                    ngay_chinh_thuc = %s,
                                    thang_bat_dau_bh = %s,
                                    trang_thai_bhxh = 'DANG_DONG',
                                    phuong_an_dieu_chinh = %s,
                                    thang_phuong_an = %s,
                                    ngay_ket_thuc = NULL
                                WHERE id = %s
                            """, (loai_hop_dong_luu_ct, han_hd_thang_ct, so_hd_moi_ct, ngay_qd, ngay_hieu_luc_ct,
                                  ngay_bat_dau_bh_ct, pa_ct_val, format_thang_nam(ngay_bat_dau_bh_ct), nv_qd['id']))
                            # Đóng lịch sử công tác cũ (thử việc)
                            c_s.execute("""
                                UPDATE lich_su_cong_tac 
                                SET den_ngay = %s, so_hop_dong = %s
                                WHERE nhan_vien_id = %s AND loai_hop_dong = 'Thử việc' AND den_ngay IS NULL
                            """, (ngay_hieu_luc_ct - timedelta(days=1), so_hd_tv_cu_ct, nv_qd['id']))
                            # Mở dòng lịch sử công tác mới
                            c_s.execute("""
                                INSERT INTO lich_su_cong_tac (
                                    nhan_vien_id, tu_ngay, chuc_danh, phong_ban, noi_lam_viec, loai_hop_dong, he_so_luong, so_hop_dong
                                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                            """, (
                                nv_qd['id'], ngay_hieu_luc_ct,
                                nv_qd.get('chuc_danh_nghe', ''),
                                nv_qd.get('phong_ban_lam_viec', ''),
                                nv_qd.get('noi_lam_viec') or get_cau_hinh('noi_lam_viec', 'Cảng THQT Hòn La'),
                                loai_hop_dong_luu_ct, nv_qd.get('he_so_luong', 0), so_hd_moi_ct
                            ))
                            loai_hd_moi_final = loai_hop_dong_luu_ct

                        # Cập nhật thêm Chức danh/Phòng ban vào hồ sơ nếu QĐ này làm thay đổi các
                        # trường đó (áp dụng chung mọi loại QĐ — trước đây chỉ DOI_CHUC_DANH/
                        # DIEU_CHUYEN mới cập nhật, nên BO_NHIEM có kèm chức danh/phòng ban mới
                        # bị bỏ sót không ghi vào hồ sơ).
                        if chuc_danh_cu_v is not None or phong_ban_cu_v is not None:
                            set_parts_ho_so = []
                            set_vals_ho_so = []
                            if chuc_danh_cu_v is not None:
                                set_parts_ho_so.append("chuc_danh_nghe = %s")
                                set_vals_ho_so.append(chuc_danh_nghe_moi_final)
                            if phong_ban_cu_v is not None:
                                set_parts_ho_so.append("phong_ban_lam_viec = %s")
                                set_vals_ho_so.append(phong_ban_moi_final)
                            set_vals_ho_so.append(nv_qd['id'])
                            c_s.execute(
                                f"UPDATE nhan_vien SET {', '.join(set_parts_ho_so)} WHERE id = %s",
                                set_vals_ho_so
                            )

                        # Cập nhật tab "Lịch sử công tác": đóng dòng đang mở (den_ngay) và, trừ khi
                        # là QĐ chấm dứt HĐ, mở thêm 1 dòng mới kể từ ngày QĐ có hiệu lực — để lịch sử
                        # công tác phản ánh đúng mốc thay đổi thay vì chỉ thấy giá trị mới nhất.
                        if loai_qd == 'CHAM_DUT_HD':
                            c_s.execute("""
                                UPDATE lich_su_cong_tac SET den_ngay = %s
                                WHERE nhan_vien_id = %s AND den_ngay IS NULL
                            """, (ngay_qd, nv_qd['id']))
                        elif chuc_danh_cu_v is not None or phong_ban_cu_v is not None or loai_hd_cu_v is not None:
                            c_s.execute("""
                                UPDATE lich_su_cong_tac SET den_ngay = %s
                                WHERE nhan_vien_id = %s AND den_ngay IS NULL
                            """, (ngay_qd - timedelta(days=1), nv_qd['id']))
                            c_s.execute("""
                                INSERT INTO lich_su_cong_tac
                                    (nhan_vien_id, tu_ngay, chuc_danh, phong_ban, noi_lam_viec, loai_hop_dong, he_so_luong)
                                VALUES (%s, %s, %s, %s, %s, %s, %s)
                            """, (
                                nv_qd['id'], ngay_qd, chuc_danh_nghe_moi_final, phong_ban_moi_final,
                                nv_qd.get('noi_lam_viec') or get_cau_hinh('noi_lam_viec', 'Cảng THQT Hòn La'),
                                loai_hd_moi_final, nv_qd.get('he_so_luong', 0)
                            ))

                        db_s.commit()
                        db_s.close()

                        if loai_qd == 'DIEU_CHUYEN':
                            # Cập nhật ngay group chat theo phòng ban mới, không chờ tới lần sau vào Chat nội bộ
                            try:
                                chat_noi_bo.sync_department_rooms()
                            except Exception:
                                pass

                        st.session_state['qdns_last_file'] = file_path
                        st.session_state['qdns_last_label'] = LOAI_QDNS_LABEL[loai_qd]
                        st.session_state['qdns_last_so'] = so_qd
                        st.cache_data.clear()
                        st.rerun()
                    except Exception as e:
                        try:
                            db_s.rollback()
                            db_s.close()
                        except Exception:
                            pass
                        st.error(f"❌ Lỗi khi tạo quyết định: {e}")


        if _chedo_qdns == "🔍 Tra cứu lịch sử":
            # ===== Lịch sử các quyết định nhân sự đã ban hành =====
            st.divider()
            st.subheader("📚 Lịch sử Quyết định nhân sự")
            search_qd = st.text_input("🔍 Tìm Quyết định (theo Số QĐ, Mã NV, Họ tên, Nội dung):", key="search_qdns")

            db_h = st.session_state.db_engine.get_connection()
            c_h = db_h.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            try:
                db_h.rollback()  # dọn transaction lỡ bị abort từ thao tác trước đó
                sql_qd = """
                    SELECT q.id, q.so_qd, q.loai_qd, q.ngay_qd, q.noi_dung, q.gia_tri_truoc, q.gia_tri_sau,
                           n.ho_ten, n.ma_nv
                    FROM quyet_dinh_nhan_su q
                    JOIN nhan_vien n ON n.id = q.nhan_vien_id
                    WHERE 1=1
                """
                params_qd = []
                if search_qd.strip():
                    sql_qd += """ AND (q.so_qd ILIKE %s OR n.ma_nv ILIKE %s OR n.ho_ten ILIKE %s OR q.noi_dung ILIKE %s)"""
                    p = f"%{search_qd.strip()}%"
                    params_qd.extend([p, p, p, p])
                sql_qd += " ORDER BY q.id DESC LIMIT 200"
                c_h.execute(sql_qd, params_qd)
                lich_su_qd = c_h.fetchall()
            except Exception as e:
                st.error(f"❌ Lỗi tải lịch sử quyết định: {e}")
                lich_su_qd = []
            finally:
                db_h.close()

            if lich_su_qd:
                # Header
                h1, h2, h3, h4, h5, h6, h7 = st.columns([1.2, 1.3, 1, 1.6, 2.2, 0.6, 0.6])
                for h, txt in zip((h1, h2, h3, h4, h5), ("Số QĐ", "Loại QĐ", "Ngày QĐ", "Nhân viên", "Nội dung")):
                    h.markdown(f"**{txt}**")
                st.divider()
                for qd in lich_su_qd:
                    r1, r2, r3, r4, r5, r6, r7 = st.columns([1.2, 1.3, 1, 1.6, 2.2, 0.6, 0.6])
                    r1.write(qd['so_qd'])
                    r2.write(LOAI_QDNS_LABEL.get(qd['loai_qd'], qd['loai_qd']))
                    r3.write(format_date(qd['ngay_qd']))
                    r4.write(f"{qd['ho_ten']} ({qd['ma_nv']})")
                    r5.write(qd['noi_dung'] or '')
                    if r6.button("✏️", key=f"sua_qd_{qd['id']}", help="Sửa Quyết định"):
                        st.session_state['qdns_dang_sua'] = qd['id']
                        st.rerun()
                    if r7.button("🗑️", key=f"xoa_qd_{qd['id']}", help="Xóa Quyết định"):
                        st.session_state['qdns_dang_xoa'] = qd['id']
                        st.rerun()

                # ----- Form SỬA Quyết định -----
                if st.session_state.get('qdns_dang_sua'):
                    qd_id_sua = st.session_state['qdns_dang_sua']
                    qd_sua = next((q for q in lich_su_qd if q['id'] == qd_id_sua), None)
                    if qd_sua:
                        with st.expander(f"✏️ Sửa Quyết định số {qd_sua['so_qd']}", expanded=True):
                            st.caption("⚠️ Chỉ chỉnh sửa thông tin lưu trữ của Quyết định. Việc sửa KHÔNG tự động "
                                       "hoàn tác/áp dụng lại thay đổi tương ứng trên hồ sơ nhân viên (chức vụ, chức danh, phòng ban...).")
                            so_qd_moi = st.text_input("Số QĐ:", value=qd_sua['so_qd'], key=f"edit_so_{qd_id_sua}")
                            ngay_qd_moi = st.date_input("Ngày QĐ:", value=qd_sua['ngay_qd'], key=f"edit_ngay_{qd_id_sua}")
                            noi_dung_moi = st.text_area("Nội dung:", value=qd_sua['noi_dung'] or '', key=f"edit_nd_{qd_id_sua}", height=100)
                            col_luu_qd, col_huy_qd = st.columns(2)
                            with col_luu_qd:
                                if st.button("💾 Lưu thay đổi", key=f"btn_luu_sua_qd_{qd_id_sua}", type="primary", width='stretch', disabled=not can_edit()):
                                    try:
                                        db_u = st.session_state.db_engine.get_connection()
                                        c_u = db_u.cursor()
                                        c_u.execute("""
                                            UPDATE quyet_dinh_nhan_su SET so_qd=%s, ngay_qd=%s, noi_dung=%s
                                            WHERE id=%s
                                        """, (so_qd_moi, ngay_qd_moi, noi_dung_moi, qd_id_sua))
                                        db_u.commit(); db_u.close()
                                        st.session_state.pop('qdns_dang_sua', None)
                                        st.success("✅ Đã cập nhật Quyết định")
                                        st.cache_data.clear()
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"❌ Lỗi: {e}")
                            with col_huy_qd:
                                if st.button("✖️ Hủy", key=f"btn_huy_sua_qd_{qd_id_sua}", width='stretch'):
                                    st.session_state.pop('qdns_dang_sua', None)
                                    st.rerun()

                # ----- Xác nhận XÓA Quyết định -----
                if st.session_state.get('qdns_dang_xoa'):
                    qd_id_xoa = st.session_state['qdns_dang_xoa']
                    qd_xoa = next((q for q in lich_su_qd if q['id'] == qd_id_xoa), None)
                    if qd_xoa:
                        st.warning(f"⚠️ Xác nhận xóa Quyết định số **{qd_xoa['so_qd']}** ({LOAI_QDNS_LABEL.get(qd_xoa['loai_qd'], qd_xoa['loai_qd'])} — {qd_xoa['ho_ten']})? "
                                   f"Việc xóa KHÔNG tự động hoàn tác thay đổi đã áp dụng trên hồ sơ nhân viên.")
                        col_xn_xoa, col_huy_xoa = st.columns(2)
                        with col_xn_xoa:
                            if st.button("🗑️ Xác nhận xóa", key=f"btn_xn_xoa_qd_{qd_id_xoa}", type="primary", width='stretch', disabled=not can_edit()):
                                try:
                                    db_d = st.session_state.db_engine.get_connection()
                                    c_d = db_d.cursor()
                                    c_d.execute("DELETE FROM quyet_dinh_nhan_su WHERE id=%s", (qd_id_xoa,))
                                    db_d.commit(); db_d.close()
                                    st.session_state.pop('qdns_dang_xoa', None)
                                    st.success("✅ Đã xóa Quyết định")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Lỗi: {e}")
                        with col_huy_xoa:
                            if st.button("✖️ Hủy", key=f"btn_huy_xoa_qd_{qd_id_xoa}", width='stretch'):
                                st.session_state.pop('qdns_dang_xoa', None)
                                st.rerun()
            else:
                st.info("Chưa có Quyết định nhân sự nào được tạo.")

    with tab_co_cau:
        st.subheader("📋 CƠ CẤU NHÂN SỰ THEO PHÒNG BAN")

        db_ct = st.session_state.db_engine.get_connection()
        c_ct = db_ct.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        # Chỉ lấy nhân sự đang làm việc (đang làm hoặc thử việc) - bỏ nhân viên đã nghỉ việc
        c_ct.execute("""
            SELECT * FROM nhan_vien 
            WHERE phong_ban_lam_viec IS NOT NULL AND phong_ban_lam_viec != ''
            AND trang_thai IN ('DANG_LAM', 'THU_VIEC')
        """)
        tat_ca_nv_ct = c_ct.fetchall()
        db_ct.close()

        # Chuẩn hóa tên phòng ban của từng nhân viên trước khi gộp nhóm — dữ liệu cũ có thể
        # đã bị lưu với cách viết hoa khác nhau (VD: "Tổ Cơ Giới" vs "Tổ Cơ giới"), nếu so
        # khớp nguyên văn sẽ khiến 1 phòng ban bị tách thành nhiều nhóm và thiếu nhân viên.
        for nv_norm in tat_ca_nv_ct:
            nv_norm['phong_ban_lam_viec'] = chuan_hoa_ten_phong_ban(nv_norm.get('phong_ban_lam_viec'))

        cac_phong_ban_ct = sap_xep_phong_ban(list({nv['phong_ban_lam_viec'] for nv in tat_ca_nv_ct}))

        # Hàng search + 2 chỉ số tổng quan (giữ nguyên logic cũ, chỉ bỏ cột tiêu đề vì
        # tiêu đề đã tách thành subheader riêng full-row ở trên)
        col_search_ct, col_tong_ct, col_dl_ct = st.columns([2, 1, 1])
        with col_search_ct:
            pb_chon_ct = st.selectbox("🔍 Chọn tìm kiếm theo phòng ban:", cac_phong_ban_ct, key="pb_thongke_chitiet")

        ds_nv_ct = sap_xep_nhan_vien([nv for nv in tat_ca_nv_ct if nv['phong_ban_lam_viec'] == pb_chon_ct])
        tong_so = len(ds_nv_ct)
        dang_lam_so = len([nv for nv in ds_nv_ct if nv['trang_thai'] in ('DANG_LAM', 'THU_VIEC')])

        with col_tong_ct:
            st.metric("Tổng số nhân sự", tong_so)
        with col_dl_ct:
            st.metric("Nhân sự đang làm", dang_lam_so)

        st.divider()

        # CSS avatar: to bằng ảnh profile (200px, viền cam) khi cột đủ rộng; khi màn hình hẹp
        # (5 cột co lại) tự chuyển sang ảnh chữ nhật bo góc, rộng tối đa theo cột, vẫn giữ viền cam.
        st.markdown("""
        <style>
        .co-cau-avatar-wrap { display:flex; justify-content:center; }
        .co-cau-avatar-img {
            width: 100%;
            max-width: 200px;
            aspect-ratio: 1 / 1;
            border-radius: 50%;
            object-fit: cover;
            border: 4px solid #f59e0b;
            box-shadow: 0 4px 15px rgba(0,0,0,0.15);
        }
        @media (max-width: 1400px) {
            .co-cau-avatar-img {
                aspect-ratio: 4 / 3;
                border-radius: 12px;
                max-width: 100%;
            }
        }
        </style>
        """, unsafe_allow_html=True)

        # Phòng ban đặc biệt: không hiện Lao động chính thức/Thử việc và Đang làm/Nghỉ việc,
        # thay bằng chức vụ (không đúng bản chất & hình thức đối với 2 phòng ban này)
        PHONG_BAN_KHONG_HIEN_TT = PHONG_BAN_LANH_DAO_CAO_CAP
        # Từ khóa nhận diện "người đứng đầu" phòng ban -> luôn xếp hàng đầu, cột giữa.
        # QUAN TRỌNG: so khớp theo TỪ KHÓA (không phân biệt hoa/thường), KHÔNG so khớp
        # chính xác nguyên chuỗi — vì dữ liệu chuc_vu thực tế có thể viết khác đi chút
        # (VD: "Trưởng Phòng", "Tổ trưởng", "Giám đốc", "Trưởng Ban"...). So khớp chính xác
        # từng ký tự trước đây khiến nhiều Trưởng phòng/Tổ trưởng/Đội trưởng không được
        # nhận diện đúng, làm họ bị rơi khỏi hàng đầu (bug đã sửa).
        # THỨ TỰ trong danh sách này = THỨ TỰ ƯU TIÊN cấp bậc (đứng trước = cấp cao hơn):
        # Chủ tịch > Tổng Giám Đốc > Giám Đốc (các khối) > Trưởng phòng/Tổ trưởng/Đội trưởng/
        # Trưởng ban/Trưởng bộ phận > Phụ trách. Dùng để chọn ĐÚNG người cao cấp nhất xếp
        # hàng đầu khi phòng ban có nhiều chức danh "đứng đầu" khác nhau (bug đã sửa: trước
        # đây chỉ lấy người đầu tiên khớp bất kỳ từ khóa nào, không phân biệt cấp bậc, khiến
        # VD "Giám Đốc Tài Chính" bị xếp trên cả "Tổng Giám Đốc").
        TU_KHOA_DUNG_DAU = ['chủ tịch', 'tổng giám đốc', 'giám đốc', 'trưởng phòng',
                            'tổ trưởng', 'đội trưởng', 'trưởng ban', 'trưởng bộ phận', 'phụ trách']

        def _la_cap_pho(nv):
            cv = (nv.get('chuc_vu') or '').strip().lower()
            return cv.startswith('phó')

        def _la_dung_dau(nv):
            cv = (nv.get('chuc_vu') or '').strip().lower()
            if not cv or cv.startswith('phó'):
                return False
            return any(tk in cv for tk in TU_KHOA_DUNG_DAU)

        def _hang_dung_dau(nv):
            """Thứ hạng ưu tiên của người đứng đầu (số nhỏ = cấp cao hơn), dựa theo vị trí
            từ khóa khớp được trong TU_KHOA_DUNG_DAU. Dùng để chọn đúng người cao cấp nhất
            khi có nhiều người cùng khớp điều kiện 'đứng đầu' trong 1 phòng ban."""
            cv = (nv.get('chuc_vu') or '').strip().lower()
            for idx, tk in enumerate(TU_KHOA_DUNG_DAU):
                if tk in cv:
                    return idx
            return len(TU_KHOA_DUNG_DAU)

        def _vi_tri_hang_cuoi(so_luong, so_cot=5):
            """Vị trí cột (0..4) cho 1 hàng có `so_luong` người (< so_cot, tức hàng cuối chưa đủ).
            Ưu tiên cân xứng quanh cột giữa (index 2 = 'col 3'); người đầu danh sách (đã ưu tiên
            cấp phó) sẽ rơi vào vị trí bên trái nhất trong bộ vị trí được chọn."""
            if so_luong == 1:
                return [2]
            elif so_luong == 2:
                return [1, 3]
            elif so_luong == 3:
                return [0, 2, 4]
            else:
                return list(range(so_luong))  # 4 hoặc 5 người -> bố trí tự do, lấp đầy từ trái

        # ===== ĐOẠN CODE MỚI CHO _lay_anh_src =====
        # Đã có hàm get_avatar_bytes_cached ở trên

        def _lay_anh_src(nv_ct):
            """Trả về src cho thẻ <img>: ảnh hồ sơ nếu có, không thì ảnh mẫu trong static/
            SỬ DỤNG CACHE để không tải lại ảnh mỗi lần render.
            """
            anh_path_ct = nv_ct.get('anh_ho_so')
            
            # ===== CẢI TIẾN: Dùng cache =====
            if anh_path_ct:
                anh_bytes_ct = get_avatar_bytes_cached(anh_path_ct)
                if anh_bytes_ct:
                    img_b64 = base64.b64encode(anh_bytes_ct).decode()
                    return f"data:image/jpeg;base64,{img_b64}"
            
            # Fallback: ảnh mẫu trong static/
            gioi_tinh_ct = nv_ct.get('gioi_tinh', '')
            avatar_file = "avatar_male.png" if gioi_tinh_ct == "Nam" else "avatar_female.png"
            avatar_path = os.path.join(os.path.dirname(__file__), "static", avatar_file)
            if os.path.exists(avatar_path):
                with open(avatar_path, "rb") as f:
                    img_b64 = base64.b64encode(f.read()).decode()
                return f"data:image/png;base64,{img_b64}"
            
            # Fallback cuối cùng: ui-avatars.com
            ten_url = (nv_ct.get('ho_ten') or 'NV').replace(' ', '+')
            return f"https://ui-avatars.com/api/?name={ten_url}&size=200&background=f59e0b&color=fff"

        def _render_the_nv(nv_ct, cols_ct, idx_c):
            with cols_ct[idx_c]:
                img_src = _lay_anh_src(nv_ct)
                st.markdown(f"""
                <div class="co-cau-avatar-wrap">
                    <img src="{img_src}" class="co-cau-avatar-img">
                </div>
                """, unsafe_allow_html=True)

                if pb_chon_ct in PHONG_BAN_KHONG_HIEN_TT:
                    # HĐQT/BTGĐ: gắn xưng hô Ông/Bà theo giới tính, KHÔNG hiện mã NV
                    xung_ho_ct = get_xung_ho_trang_trong(nv_ct.get('gioi_tinh'))
                    ten_hien_thi_ct = f"{xung_ho_ct} {nv_ct['ho_ten']}".strip()
                else:
                    ten_hien_thi_ct = f"{nv_ct['ho_ten']}-{nv_ct['ma_nv']}"
                st.markdown(f"<p style='text-align:center;margin-bottom:0;'><b>{ten_hien_thi_ct}</b></p>", unsafe_allow_html=True)
                if pb_chon_ct not in PHONG_BAN_KHONG_HIEN_TT:
                    # Nhóm BHXH = 'Văn phòng' -> hiện chức vụ; ngược lại hiện chức danh nghề (như trước đây)
                    if (nv_ct.get('nhom_bhxh') or '') == 'Văn phòng':
                        dong_phu = nv_ct.get('chuc_vu') or ''
                    else:
                        dong_phu = nv_ct.get('chuc_danh_nghe') or ''
                    st.markdown(f"<p style='text-align:center;color:gray;font-size:0.85em;'>{dong_phu}</p>", unsafe_allow_html=True)

                if pb_chon_ct in PHONG_BAN_KHONG_HIEN_TT:
                    # Không đúng bản chất/hình thức với HĐQT & BTGĐ -> hiện chức vụ thay vì loại HĐ/trạng thái
                    st.markdown(f"<p style='text-align:center;'>🏷️ {nv_ct.get('chuc_vu') or 'Thành viên'}</p>", unsafe_allow_html=True)
                else:
                    if nv_ct.get('loai_hop_dong') == 'Thử việc':
                        st.markdown("<p style='text-align:center;color:red;'>Hợp đồng Thử việc</p>", unsafe_allow_html=True)
                    elif nv_ct.get('loai_hop_dong'):
                        st.markdown("<p style='text-align:center;color:green;'>Hợp đồng Chính thức</p>", unsafe_allow_html=True)

                if st.button("Xem chi tiết>>", key=f"xem_ct_{nv_ct['id']}", width='stretch'):
                    st.session_state['_nv_xem_chi_tiet_dashboard'] = nv_ct['id']

        if ds_nv_ct:
            so_cot = 5

            # Tách người đứng đầu phòng ban (nếu có) -> luôn ở hàng đầu tiên, cột giữa (index 2/5).
            # Nếu không có người đứng đầu -> bỏ qua hàng riêng này, các hàng sau tịnh tiến lên.
            ung_vien_dung_dau = [nv for nv in ds_nv_ct if _la_dung_dau(nv)]
            nguoi_dung_dau = min(
                ung_vien_dung_dau,
                key=lambda nv: (_hang_dung_dau(nv), nv.get('ho_ten') or '')
            ) if ung_vien_dung_dau else None
            if not nguoi_dung_dau:
                # Không có Tổng/Giám đốc/Trưởng/Phụ trách -> đôn "Phó" đầu tiên (theo alpha) lên hàng 1
                nguoi_dung_dau = next((nv for nv in sorted(ds_nv_ct, key=lambda x: x.get('ho_ten') or '') if _la_cap_pho(nv)), None)
            ds_con_lai = [nv for nv in ds_nv_ct if nv is not nguoi_dung_dau]
            # Cấp phó ưu tiên lên đầu (bên trái), sau đó xếp theo alpha bê tên
            ds_con_lai = sorted(ds_con_lai, key=lambda nv: (0 if _la_cap_pho(nv) else 1, nv.get('ho_ten') or ''))

            if nguoi_dung_dau:
                cols_ct = st.columns(so_cot)
                _render_the_nv(nguoi_dung_dau, cols_ct, 2)  # cột giữa trong 5 cột (0,1,[2],3,4)
                st.divider()

            for i in range(0, len(ds_con_lai), so_cot):
                hang = ds_con_lai[i:i + so_cot]
                cols_ct = st.columns(so_cot)
                vi_tri_cot = _vi_tri_hang_cuoi(len(hang), so_cot) if len(hang) < so_cot else list(range(so_cot))
                for nv_ct, idx_c in zip(hang, vi_tri_cot):
                    _render_the_nv(nv_ct, cols_ct, idx_c)

            if st.session_state.get('_nv_xem_chi_tiet_dashboard'):
                nv_id_xem = st.session_state['_nv_xem_chi_tiet_dashboard']
                nv_xem = next((nv for nv in ds_nv_ct if nv['id'] == nv_id_xem), None)
                if nv_xem:
                    st.divider()
                    render_employee_info_card(
                        nv_xem,
                        key_prefix=f"nv_co_cau_{nv_xem['id']}",
                        on_close=lambda: st.session_state.pop('_nv_xem_chi_tiet_dashboard', None)
                    )
        else:
            st.info("Không có nhân sự nào trong phòng ban này.")

# ========== CHẤM CÔNG ==========
elif menu == "🕒 Chấm công":
    st.markdown(f"# {i18n.tm('🕒 Chấm công')}", unsafe_allow_html=True)

    # Đọc phương thức chấm công từ cấu hình tenant
    _MAP_PT = {'THU_CONG': 'manual', 'MAY_VAN_TAY': 'fingerprint', 'FACE_ID': 'faceid'}
    _MAP_PT_LABEL = {'THU_CONG': '📝 Thủ công', 'MAY_VAN_TAY': '📥 Máy vân tay', 'FACE_ID': '👤 Face ID'}
    phuong_thuc_cfg = get_cau_hinh('cc_phuong_thuc', 'THU_CONG')
    # 3 nút chỉ báo — chỉ phương thức đã cấu hình là sáng, 2 cái kia mờ
    col_method1, col_method2, col_method3 = st.columns(3)
    with col_method1:
        st.button("📝 Thủ công", use_container_width=True,
                  type="primary" if phuong_thuc_cfg == 'THU_CONG' else "secondary",
                  disabled=(phuong_thuc_cfg != 'THU_CONG'), key="cc_btn_m")
    with col_method2:
        st.button("📥 Máy vân tay", use_container_width=True,
                  type="primary" if phuong_thuc_cfg == 'MAY_VAN_TAY' else "secondary",
                  disabled=(phuong_thuc_cfg != 'MAY_VAN_TAY'), key="cc_btn_f")
    with col_method3:
        st.button("👤 Face ID", use_container_width=True,
                  type="primary" if phuong_thuc_cfg == 'FACE_ID' else "secondary",
                  disabled=(phuong_thuc_cfg != 'FACE_ID'), key="cc_btn_fi")

    st.divider()
    st.caption("💡 Thay đổi phương thức chấm công → vào **⚙️ Danh mục** > tab **🕒 Chấm công**")

    # ========== BCC LUÔN HIỂN THỊ (không phụ thuộc phương thức) ==========
    ensure_cham_cong_table()

    # Bố cục chọn tháng/năm/bộ phận
    if not st.session_state.get('cc_full_open', False):
        col_m1, col_m2, col_m3, col_m4 = st.columns([1, 1, 2, 1.5])
        with col_m1:
            thang_nhap = st.selectbox("Tháng", list(range(1, 13)), index=date.today().month - 1, key="cc_thang_nhap", label_visibility="collapsed")
        with col_m2:
            nam_nhap = st.number_input("Năm", min_value=2020, max_value=2100, value=date.today().year, step=1, key="cc_nam_nhap", label_visibility="collapsed")
        with col_m3:
            db_bp = st.session_state.db_engine.get_connection()
            c_bp = db_bp.cursor()
            c_bp.execute("""SELECT DISTINCT phong_ban_lam_viec FROM nhan_vien
                            WHERE trang_thai IN ('DANG_LAM','THU_VIEC') AND phong_ban_lam_viec IS NOT NULL
                            AND phong_ban_lam_viec != '' ORDER BY phong_ban_lam_viec""")
            all_depts = [r[0] for r in c_bp.fetchall()]
            c_bp.close(); db_bp.close()
            bo_phan_nhap = st.multiselect(
                "Bộ phận", all_depts, default=[],
                format_func=lambda d: CHAM_CONG_DEPT_LABEL.get(d, d),
                key="cc_bp_nhap",
                placeholder="Tất cả bộ phận",
                label_visibility="collapsed"
            )
        with col_m4:
            if st.button("📂 Mở BCC", type="primary", use_container_width=True):
                st.session_state.cc_full_open = True
                st.session_state.cc_view_thang = thang_nhap
                st.session_state.cc_view_nam = int(nam_nhap)
                st.session_state.cc_view_bo_phan = bo_phan_nhap
                st.session_state.cc_edit_mode = False
                st.session_state.cc_data_loaded = False
                st.rerun()

        st.caption("💡 Chọn tháng/năm và bộ phận (để trống = tất cả), sau đó bấm 'Mở BCC'")

    # ===== Bảng chấm công full-width (BCC tháng — Điểm 1) =====
    else:
        import calendar as _cal
        thang_v = st.session_state.cc_view_thang
        nam_v = st.session_state.cc_view_nam
        bp_v = st.session_state.cc_view_bo_phan
        so_ngay = _cal.monthrange(nam_v, thang_v)[1]
        day_list = [date(nam_v, thang_v, d) for d in range(1, so_ngay + 1)]
        WD_ABBR = ["T2", "T3", "T4", "T5", "T6", "T7", "CN"]
        col_titles = [f"{d.day:02d} {WD_ABBR[d.weekday()]}" for d in day_list]
        sunday_cols = [t for d, t in zip(day_list, col_titles) if d.weekday() == 6]

        # Lấy cấu hình để biết ngày lễ
        cfg_cc = get_cau_hinh_cham_cong_full()
        danh_sach_le_set = {x['ngay'] for x in (cfg_cc.get('danh_sach_ngay_le') or [])}
        holiday_cols = [t for d, t in zip(day_list, col_titles)
                        if d.strftime('%Y-%m-%d') in danh_sach_le_set]

        # Kiểm tra khoá tháng
        da_khoa = is_thang_da_khoa(thang_v, nam_v)

        # Thanh điều khiển
        col_h1, col_h2, col_h3, col_h4, col_h5, col_h6 = st.columns([2.5, 0.6, 0.6, 0.6, 0.6, 0.7])
        with col_h1:
            ten_bp = ", ".join(CHAM_CONG_DEPT_LABEL.get(b, b) for b in bp_v) if bp_v else "Tất cả bộ phận"
            trang_thai_khoa = " 🔒" if da_khoa else ""
            st.markdown(f"**📅 BCC tháng {thang_v}/{nam_v} — {ten_bp}{trang_thai_khoa}**")
        with col_h2:
            if st.button("◀️ Đóng", key="cc_close_btn", use_container_width=True):
                st.session_state.cc_full_open = False
                st.session_state.cc_pending_missing = None
                st.rerun()
        with col_h3:
            # Nút sửa: chỉ admin/admin_bcc được bấm, tháng khoá thì disable
            edit_label = "👁️" if st.session_state.get('cc_edit_mode') else "✏️"
            if st.button(edit_label, key="cc_toggle_edit_btn", use_container_width=True,
                         disabled=(da_khoa or not can_edit_bcc())):
                st.session_state.cc_edit_mode = not st.session_state.get('cc_edit_mode', False)
                st.session_state.cc_pending_missing = None
                st.rerun()
        with col_h4:
            save_clicked = st.button(
                "💾", key="cc_save_month_btn", type="primary", use_container_width=True,
                disabled=(da_khoa or not st.session_state.get('cc_edit_mode', False) or not can_edit_bcc())
            )
        with col_h5:
            if st.button("📤 Xuất", key="cc_export_btn", use_container_width=True):
                st.session_state.cc_export_trigger = True
                st.rerun()
        with col_h6:
            # Nút khoá/mở khoá — chỉ admin
            if can_khoa_thang_bcc():
                if da_khoa:
                    if st.button("🔓", key="cc_unlock_btn", use_container_width=True,
                                 help="Mở khoá BCC tháng này"):
                        khoa_mo_thang_bcc(thang_v, nam_v, 'MO')
                        st.rerun()
                else:
                    if st.button("🔒", key="cc_lock_btn", use_container_width=True,
                                 help="Khoá BCC tháng này (sau khi tính lương)"):
                        khoa_mo_thang_bcc(thang_v, nam_v, 'KHOA')
                        st.rerun()

        if da_khoa:
            st.info("🔒 Tháng này đã khoá — không thể sửa BCC. Chỉ Admin có quyền mở khoá.")

        # Hướng dẫn ký hiệu mới
        with st.expander("📖 Bảng ký hiệu chấm công (22 mã)", expanded=False):
            _nhom_a = {k: v for k, v in KY_HIEU_CHAM_CONG.items() if v.get('nhom') == 'A'}
            _nhom_b = {k: v for k, v in KY_HIEU_CHAM_CONG.items() if v.get('nhom') == 'B'}
            _nhom_c = {k: v for k, v in KY_HIEU_CHAM_CONG.items() if v.get('nhom') == 'C'}
            col_leg1, col_leg2, col_leg3 = st.columns(3)
            with col_leg1:
                st.markdown("**Nhóm A — Hưởng lương**")
                for ma, tt in _nhom_a.items():
                    cong_str = f" ({tt['cong']} công)" if tt.get('cong') else ""
                    st.caption(f"`{ma}` {tt['ten']}{cong_str}")
            with col_leg2:
                st.markdown("**Nhóm B — BHXH chi trả**")
                for ma, tt in _nhom_b.items():
                    st.caption(f"`{ma}` {tt['ten']}")
            with col_leg3:
                st.markdown("**Nhóm C — Không lương**")
                for ma, tt in _nhom_c.items():
                    st.caption(f"`{ma}` {tt['ten']}")
                st.markdown("**Tăng ca (dòng TC)**")
                for ma, tt in LOAI_TANG_CA.items():
                    st.caption(f"`{ma}` {tt['ten']} (×{tt['he_so_mac_dinh']})")
            st.caption("💡 **Chấm công full**: tick checkbox ở cột cuối → tự điền `x` cho tất cả ngày trống (trừ CN/Lễ).")

        # Lấy danh sách nhân viên
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if bp_v:
            c.execute("""SELECT id, ma_nv, ho_ten, chuc_danh_nghe, phong_ban_lam_viec FROM nhan_vien
                             WHERE trang_thai IN ('DANG_LAM','THU_VIEC') AND phong_ban_lam_viec = ANY(%s)
                               AND so_hdld IS NOT NULL
                             ORDER BY phong_ban_lam_viec ASC, ma_nv ASC""", (bp_v,))
        else:
            c.execute("""SELECT id, ma_nv, ho_ten, chuc_danh_nghe, phong_ban_lam_viec FROM nhan_vien
                             WHERE trang_thai IN ('DANG_LAM','THU_VIEC') AND so_hdld IS NOT NULL
                             ORDER BY phong_ban_lam_viec ASC, ma_nv ASC""")
        nv_list = c.fetchall()

        # Lấy dữ liệu chấm công hiện có — ĐỌC ma_cong (mới) + ca_ngay (cũ, fallback)
        existing = {}
        if nv_list:
            nv_ids = [nv['id'] for nv in nv_list]
            c.execute("""SELECT nhan_vien_id, ngay, ma_cong, ca_ngay, ca_dem,
                                gio_tang_ca, gio_tang_ca_dem, loai_ngay_tang_ca,
                                trang_thai_cham_cong, trang_thai_vi_tri
                         FROM cham_cong
                         WHERE nhan_vien_id = ANY(%s)
                           AND EXTRACT(MONTH FROM ngay) = %s
                           AND EXTRACT(YEAR FROM ngay) = %s""",
                      (nv_ids, thang_v, nam_v))
            for r in c.fetchall():
                existing[(r['nhan_vien_id'], r['ngay'])] = r
        c.close(); db.close()

        if not nv_list:
            st.warning("Không có nhân viên nào phù hợp với bộ phận đã chọn.")
        else:
            # Tính tổng hợp cho hiển thị cột summary
            nv_ids_list = [nv['id'] for nv in nv_list]
            tong_hop_all = tong_hop_cham_cong_thang_nhieu_nv(nv_ids_list, thang_v, nam_v)

            # Xây dựng dữ liệu BCC — mỗi NV 2 dòng: Ký hiệu + Tăng ca
            flat_rows = []
            nv_row_indices = {}

            for nv in nv_list:
                dept = nv['phong_ban_lam_viec']
                th = tong_hop_all.get(nv['id'], {})

                # Dòng 1: Ký hiệu (ma_cong)
                row_kh = {
                    "Mã NV": nv['ma_nv'],
                    "Họ tên": nv['ho_ten'],
                    "Loại": "Ký hiệu",
                }
                # Dòng 2: Tăng ca (giờ)
                row_tc = {
                    "Mã NV": "",
                    "Họ tên": "",
                    "Loại": "Tăng ca",
                }

                for d, title in zip(day_list, col_titles):
                    rec = existing.get((nv['id'], d))
                    # Ký hiệu
                    ma_val = ""
                    if rec:
                        ma_val = rec.get('ma_cong') or ""
                        if not ma_val and rec.get('ca_ngay'):
                            ma_val = cc_normalize_marker(rec['ca_ngay']) or ""
                        # Hiển thị trạng thái đặc biệt
                        if not ma_val:
                            if rec.get('trang_thai_cham_cong') == 'THIEU_GIO_RA':
                                ma_val = "⚠️"
                            elif rec.get('trang_thai_vi_tri') == 'CHO_DUYET':
                                ma_val = "📍"
                    # Auto CN cho Chủ nhật chưa có dữ liệu
                    if d.weekday() == 6 and not ma_val:
                        ma_val = "CN"
                    # Auto NL cho ngày lễ chưa có dữ liệu
                    if d.strftime('%Y-%m-%d') in danh_sach_le_set and not ma_val and d.weekday() != 6:
                        ma_val = "NL"
                    row_kh[title] = ma_val

                    # Tăng ca
                    tc_val = ""
                    if rec:
                        tc_gio = rec.get('gio_tang_ca')
                        if tc_gio and float(tc_gio) > 0:
                            tc_val = str(float(tc_gio))
                    row_tc[title] = tc_val

                # Cột tổng hợp (chỉ ở dòng Ký hiệu)
                row_kh["Công"] = th.get('tong_cong', 0)
                row_kh["Phép"] = th.get('phep_da_dung', 0)
                row_kh["TC(h)"] = th.get('tong_gio_tang_ca', 0)
                row_kh["Chấm công full"] = "⬜"

                row_tc["Công"] = None
                row_tc["Phép"] = None
                row_tc["TC(h)"] = None
                row_tc["Chấm công full"] = ""

                # Quyết định 1 hay 2 dòng: kiểm tra phòng ban có cho phép tăng ca không
                cfg_tc_phong = get_cau_hinh_tang_ca_theo_phong(dept)
                hien_dong_tc = cfg_tc_phong.get('cho_phep_tang_ca', True)

                if not hien_dong_tc:
                    # Phòng không cho phép TC → 1 dòng (chỉ ký hiệu)
                    flat_rows.append(row_kh)
                    nv_row_indices[nv['ma_nv']] = {
                        'nv_id': nv['id'],
                        'ca_main': len(flat_rows) - 1,
                        'tc': None,
                    }
                else:
                    # Phòng cho phép TC → 2 dòng (ký hiệu + tăng ca)
                    flat_rows.append(row_kh)
                    idx_kh = len(flat_rows) - 1
                    flat_rows.append(row_tc)
                    idx_tc = len(flat_rows) - 1
                    nv_row_indices[nv['ma_nv']] = {
                        'nv_id': nv['id'],
                        'ca_main': idx_kh,
                        'tc': idx_tc,
                    }

            df_month = pd.DataFrame(flat_rows)

            # Đảm bảo thứ tự cột
            col_order = ["Mã NV", "Họ tên", "Loại"] + col_titles + ["Công", "Phép", "TC(h)", "Chấm công full"]
            for col_need in col_order:
                if col_need not in df_month.columns:
                    df_month[col_need] = ""
            df_month = df_month[col_order]

            CC_HEADER_H = 38
            MAX_VISIBLE = 25  # tối đa 25 dòng hiện cùng lúc, scroll nếu nhiều hơn
            table_height = CC_HEADER_H + CC_ROW_HEIGHT * min(len(df_month) + 2, MAX_VISIBLE)

            # --- Export Excel ---
            if st.session_state.get('cc_export_trigger', False):
                st.session_state.cc_export_trigger = False
                try:
                    from openpyxl import Workbook
                    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
                    from openpyxl.utils import get_column_letter

                    wb = Workbook()
                    ws = wb.active
                    ws.title = f"BCC_{thang_v}_{nam_v}"

                    df_export = df_month.drop(columns=["Chấm công full"], errors='ignore')
                    for col_idx, col_name in enumerate(df_export.columns, 1):
                        cell = ws.cell(row=1, column=col_idx, value=col_name)
                        cell.font = Font(bold=True, size=10)
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                        cell.fill = PatternFill(start_color="D9E1F2", fill_type="solid")
                        if col_name in ["Mã NV", "Loại"]:
                            ws.column_dimensions[get_column_letter(col_idx)].width = 12
                        elif col_name == "Họ tên":
                            ws.column_dimensions[get_column_letter(col_idx)].width = 25
                        elif col_name in ["Công", "Phép", "TC(h)"]:
                            ws.column_dimensions[get_column_letter(col_idx)].width = 8
                        else:
                            ws.column_dimensions[get_column_letter(col_idx)].width = 7

                    sunday_fill = PatternFill(start_color="FFF2CC", fill_type="solid")
                    holiday_fill = PatternFill(start_color="FCE4EC", fill_type="solid")
                    for row_idx, row_data in enumerate(df_export.itertuples(index=False), 2):
                        for col_idx, val in enumerate(row_data, 1):
                            cell = ws.cell(row=row_idx, column=col_idx, value=val)
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                            col_name = df_export.columns[col_idx - 1]
                            if col_name in sunday_cols:
                                cell.fill = sunday_fill
                            elif col_name in holiday_cols:
                                cell.fill = holiday_fill

                    filename = f"BCC_{thang_v}_{nam_v}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                    wb.save(filename)

                    with open(filename, "rb") as f:
                        st.download_button(
                            label="📥 TẢI FILE EXCEL",
                            data=f, file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True,
                        )
                    st.success(f"✅ Đã xuất: {filename}")
                except Exception as e:
                    st.error(f"❌ Lỗi xuất file: {e}")

            # --- CHẾ ĐỘ XEM ---
            if not st.session_state.get('cc_edit_mode', False):
                def _highlight_special(s):
                    styles = [''] * len(s)
                    if s.name in sunday_cols:
                        styles = ['background-color: #FFF2CC; color: #999;'] * len(s)
                    elif s.name in holiday_cols:
                        styles = ['background-color: #FCE4EC; color: #C62828;'] * len(s)
                    return styles

                df_view = df_month.drop(columns=["Chấm công full"], errors='ignore')

                view_col_cfg = {
                    "Mã NV": cc_pin_col(st.column_config.TextColumn, width="small"),
                    "Họ tên": cc_pin_col(st.column_config.TextColumn, width=170),
                    "Loại": cc_pin_col(st.column_config.TextColumn, width="small"),
                    "Công": st.column_config.NumberColumn(width="small", format="%.1f"),
                    "Phép": st.column_config.NumberColumn(width="small", format="%.1f"),
                    "TC(h)": st.column_config.NumberColumn(width="small", format="%.1f"),
                }
                styled = (
                    df_view.style
                    .apply(_highlight_special, axis=0)
                    .set_properties(**{"text-align": "center", "vertical-align": "middle"})
                    .hide(axis="index")
                )
                cc_render_grid(
                    styled, edit=False, use_container_width=True, height=table_height,
                    column_config=view_col_cfg,
                )

                # Tổng hợp cuối bảng
                tong_cong_all = sum(th.get('tong_cong', 0) for th in tong_hop_all.values())
                tong_tc_all = sum(th.get('tong_gio_tang_ca', 0) for th in tong_hop_all.values())
                so_chua_cham = sum(th.get('so_ngay_chua_cham', 0) for th in tong_hop_all.values())
                so_thieu_ra = sum(th.get('so_ngay_thieu_gio_ra', 0) for th in tong_hop_all.values())

                col_s1, col_s2, col_s3, col_s4 = st.columns(4)
                col_s1.metric("👥 Nhân viên", len(nv_list))
                col_s2.metric("📊 Tổng công", f"{tong_cong_all:.1f}")
                col_s3.metric("⏰ Tổng TC", f"{tong_tc_all:.1f}h")
                if so_chua_cham > 0 or so_thieu_ra > 0:
                    col_s4.metric("⚠️ Cảnh báo", f"{so_chua_cham} chưa chấm, {so_thieu_ra} thiếu giờ ra")
                else:
                    col_s4.metric("✅ Trạng thái", "Đầy đủ")

                st.caption("👁️ Xem | ✏️ Bấm bút chì để sửa | 📤 Xuất Excel "
                           "| ⚠️ = thiếu giờ ra | 📍 = chờ duyệt vị trí")

            # --- CHẾ ĐỘ SỬA ---
            else:
                col_cfg = {
                    "Mã NV": cc_pin_col(st.column_config.TextColumn, disabled=True, width="small"),
                    "Họ tên": cc_pin_col(st.column_config.TextColumn, disabled=True, width=170),
                    "Loại": cc_pin_col(st.column_config.TextColumn, disabled=True, width="small"),
                    "Công": st.column_config.NumberColumn(disabled=True, width="small", format="%.1f"),
                    "Phép": st.column_config.NumberColumn(disabled=True, width="small", format="%.1f"),
                    "TC(h)": st.column_config.NumberColumn(disabled=True, width="small", format="%.1f"),
                    "Chấm công full": cc_pin_col(st.column_config.CheckboxColumn, width="small"),
                }
                for t in col_titles:
                    col_cfg[t] = st.column_config.TextColumn(width="small", validate=CHAM_CONG_CELL_REGEX)

                edit_key = f"cc_editor_{thang_v}_{nam_v}_{'-'.join(bp_v) if bp_v else 'all'}"

                edited_df = st.data_editor(
                    df_month,
                    column_config=col_cfg,
                    hide_index=True,
                    num_rows="fixed",
                    use_container_width=True,
                    height=table_height,
                    key=edit_key,
                )

                # Xử lý checkbox "Chấm công full"
                if edited_df is not None:
                    for idx, row in edited_df.iterrows():
                        if row["Loại"] == "Ký hiệu":
                            old_val = df_month.iloc[idx]["Chấm công full"]
                            new_val = row["Chấm công full"]
                            if new_val is True and old_val is not True:
                                for d, title in zip(day_list, col_titles):
                                    if d.weekday() == 6:
                                        continue
                                    if d.strftime('%Y-%m-%d') in danh_sach_le_set:
                                        continue
                                    current_val = str(edited_df.iloc[idx][title] or "").strip()
                                    if not current_val:
                                        edited_df.at[idx, title] = "x"
                                edited_df.at[idx, "Chấm công full"] = False
                                st.rerun()

                # --- LƯU ---
                if save_clicked:
                    # Kiểm tra thiếu
                    missing = []
                    for nv_ma, indices in nv_row_indices.items():
                        ca_main_idx = indices['ca_main']
                        nv_ten = edited_df.iloc[ca_main_idx]["Họ tên"]
                        for d, title in zip(day_list, col_titles):
                            if d.weekday() == 6:
                                continue
                            if d >= date.today():
                                continue
                            if d.strftime('%Y-%m-%d') in danh_sach_le_set:
                                continue
                            val = str(edited_df.iloc[ca_main_idx][title] or "").strip()
                            if not val:
                                missing.append((nv_ma, nv_ten, d))

                    if missing and not st.session_state.get('cc_force_save_approved', False):
                        st.warning(f"⚠️ Có {len(missing)} lượt chưa chấm công")
                        with st.expander("Xem chi tiết"):
                            for ma_nv, ho_ten, d in missing[:100]:
                                st.caption(f"- {ma_nv} - {ho_ten}: {d.strftime('%d/%m/%Y')}")
                            if len(missing) > 100:
                                st.caption(f"... và {len(missing) - 100} lượt khác")
                        col_cf1, col_cf2 = st.columns(2)
                        with col_cf1:
                            if st.button("✅ Vẫn lưu", key="cc_force_save_btn", type="primary",
                                         use_container_width=True, disabled=not can_edit()):
                                st.session_state.cc_force_save_approved = True
                                st.rerun()
                        with col_cf2:
                            if st.button("✏️ Sửa tiếp", key="cc_cancel_save_btn",
                                         use_container_width=True):
                                st.session_state.cc_force_save_approved = False
                                st.rerun()
                    else:
                        # Thực hiện lưu — duyệt qua nv_row_indices (đọc cả dòng Ký hiệu + Tăng ca)
                        db2 = st.session_state.db_engine.get_connection()
                        c2 = db2.cursor()
                        n_saved = 0

                        c2_nv = db2.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                        if bp_v:
                            c2_nv.execute("""SELECT id, ma_nv FROM nhan_vien
                                             WHERE trang_thai IN ('DANG_LAM','THU_VIEC')
                                               AND phong_ban_lam_viec = ANY(%s)
                                               AND so_hdld IS NOT NULL
                                             ORDER BY ma_nv ASC""", (bp_v,))
                        else:
                            c2_nv.execute("""SELECT id, ma_nv FROM nhan_vien
                                             WHERE trang_thai IN ('DANG_LAM','THU_VIEC')
                                               AND so_hdld IS NOT NULL
                                             ORDER BY ma_nv ASC""")
                        nv_map = {row['ma_nv']: row['id'] for row in c2_nv.fetchall()}
                        c2_nv.close()

                        for nv_ma, indices in nv_row_indices.items():
                            if nv_ma not in nv_map:
                                continue
                            nv_id = nv_map[nv_ma]
                            ca_main_idx = indices['ca_main']
                            tc_idx = indices.get('tc')

                            for d, title in zip(day_list, col_titles):
                                # Đọc ký hiệu từ dòng Ký hiệu
                                v_ma_cong = None
                                val_kh = str(edited_df.iloc[ca_main_idx][title] or "").strip()
                                if val_kh and val_kh not in ("⚠️", "📍"):
                                    v_ma_cong = cc_normalize_marker(val_kh)

                                # Đọc giờ TC từ dòng Tăng ca (nếu có)
                                v_tc = 0
                                if tc_idx is not None:
                                    val_tc = str(edited_df.iloc[tc_idx][title] or "").strip()
                                    try:
                                        v_tc = float(val_tc.replace(",", ".")) if val_tc else 0
                                    except ValueError:
                                        v_tc = 0

                                # Bỏ qua nếu cả 2 đều trống và chưa có trong DB
                                if v_ma_cong is None and v_tc == 0 and (nv_id, d) not in existing:
                                    continue

                                # Ghi vào ma_cong (mới) + ca_ngay (backward compat)
                                c2.execute("""
                                    INSERT INTO cham_cong
                                        (nhan_vien_id, ngay, ma_cong, ca_ngay, gio_tang_ca,
                                         nguon, created_by, updated_at)
                                    VALUES (%s, %s, %s, %s, %s, 'THU_CONG', %s, NOW())
                                    ON CONFLICT (nhan_vien_id, ngay) DO UPDATE SET
                                        ma_cong = EXCLUDED.ma_cong,
                                        ca_ngay = EXCLUDED.ca_ngay,
                                        gio_tang_ca = EXCLUDED.gio_tang_ca,
                                        nguon = CASE WHEN cham_cong.nguon = 'FACE_ID'
                                                     THEN cham_cong.nguon
                                                     ELSE EXCLUDED.nguon END,
                                        updated_at = NOW()
                                """, (nv_id, d, v_ma_cong, v_ma_cong, v_tc,
                                      st.session_state.username))
                                n_saved += 1

                        db2.commit()
                        c2.close(); db2.close()
                        st.success(f"✅ Đã lưu {n_saved} lượt chấm công tháng {thang_v}/{nam_v}.")
                        st.session_state.cc_edit_mode = False
                        st.session_state.cc_force_save_approved = False
                        st.rerun()

    # ========== 2. TRÍCH XUẤT TỪ MÁY CHẤM VÂN TAY ==========
    if phuong_thuc_cfg == 'MAY_VAN_TAY':
        st.info("""
        ### 🚧 Tính năng đang phát triển
        
        Dự kiến hỗ trợ:
        - Upload file dữ liệu xuất từ máy chấm vân tay (.xls/.csv)
        - Ánh xạ mã nhân viên trên máy chấm công với Mã NV
        - Tự động quy đổi giờ vào/ra thành mã công
        """)

    # ========== 3. FACE ID ==========
    if phuong_thuc_cfg == 'FACE_ID':
        ensure_face_id_table()

        # Tải mô hình nhận diện khuôn mặt (lần đầu tải ~39MB, các lần sau dùng lại từ cache)
        try:
            with st.spinner("Đang chuẩn bị mô hình nhận diện khuôn mặt (lần đầu có thể mất 1-2 phút)..."):
                face_id_cham_cong.chuan_bi_model()
        except Exception as e:
            st.error(f"❌ Không chuẩn bị được mô hình nhận diện: {e}")
            st.stop()

        # Vai trò admin/admin_bcc: thấy đủ 4 tab
        # Vai trò khác (NV, trưởng phòng...): chỉ thấy Check-in + Kết quả, mở thẳng Check-in
        if st.session_state.role in ('admin', 'admin_bcc'):
            tab_face_dangky, tab_face_checkin, tab_face_ketqua, tab_face_dieuchinh = st.tabs(
                ["📸 Đăng ký khuôn mặt", "🎥 Check-in / Check-out", "📋 Kết quả hôm nay", "✏️ Điều chỉnh"]
            )
        else:
            tab_face_checkin, tab_face_ketqua = st.tabs(
                ["🎥 Check-in / Check-out", "📋 Kết quả hôm nay"]
            )
            tab_face_dangky = None
            tab_face_dieuchinh = None

        # ----- TAB ĐĂNG KÝ (chỉ admin/admin_bcc) -----
        if tab_face_dangky is not None:
            with tab_face_dangky:
                st.caption("Hệ thống tự lấy ảnh avatar đã upload trong hồ sơ nhân viên để đăng ký khuôn mặt. "
                       "Nhân viên nào chưa có ảnh avatar sẽ hiện ở đây để HR xử lý trước.")

                if not can_edit():
                    st.warning("⚠️ Bạn không có quyền thực hiện thao tác này.")
                else:
                    db_face = st.session_state.db_engine.get_connection()
                    c_face = db_face.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

                    # Lấy danh sách NV đang làm + trạng thái avatar + trạng thái đã đăng ký Face ID chưa
                    c_face.execute("""
                        SELECT nv.id, nv.ma_nv, nv.ho_ten, nv.anh_ho_so,
                               f.id AS face_id_id
                        FROM nhan_vien nv
                        LEFT JOIN nhan_vien_face_id f ON f.nhan_vien_id = nv.id
                        WHERE nv.trang_thai IN ('DANG_LAM', 'THU_VIEC')
                        ORDER BY nv.ho_ten
                    """)
                    ds_nv_face = c_face.fetchall()
                    db_face.close()

                    if not ds_nv_face:
                        st.info("Không có nhân viên đang làm việc/thử việc.")
                    else:
                        # Phân loại
                        chua_co_anh    = [x for x in ds_nv_face if not x['anh_ho_so']]
                        co_anh_chua_dk = [x for x in ds_nv_face if x['anh_ho_so'] and not x['face_id_id']]
                        da_dk          = [x for x in ds_nv_face if x['anh_ho_so'] and x['face_id_id']]

                        # Tóm tắt trạng thái
                        col_s1, col_s2, col_s3 = st.columns(3)
                        col_s1.metric("✅ Đã đăng ký Face ID", len(da_dk))
                        col_s2.metric("⏳ Có ảnh, chưa đăng ký", len(co_anh_chua_dk))
                        col_s3.metric("❌ Chưa có ảnh avatar", len(chua_co_anh))

                        # Nhân viên chưa có ảnh — cảnh báo HR upload ảnh trước
                        if chua_co_anh:
                            with st.expander(f"❌ {len(chua_co_anh)} nhân viên CHƯA CÓ ảnh avatar — không thể đăng ký Face ID"):
                                st.caption("Vào hồ sơ từng nhân viên (menu Nhân viên → Upload ảnh hồ sơ) để upload ảnh, "
                                           "sau đó quay lại đây đồng bộ.")
                                for nv in chua_co_anh:
                                    st.write(f"• {nv['ma_nv']} — {nv['ho_ten']}")

                        # Đồng bộ 1 người
                        if co_anh_chua_dk:
                            st.divider()
                            st.markdown(f"**⏳ {len(co_anh_chua_dk)} nhân viên có ảnh, chưa đăng ký Face ID**")
                            nv_map_dk = {f"{x['ma_nv']} - {x['ho_ten']}": x for x in co_anh_chua_dk}
                            nv_chon_label = st.selectbox("📌 Chọn nhân viên để đăng ký:", list(nv_map_dk.keys()),
                                                         key="face_dangky_nv_select")
                            nv_chon = nv_map_dk[nv_chon_label]

                            # Hiển thị preview ảnh avatar sẽ dùng
                            anh_bytes = get_anh_ho_so_bytes(nv_chon['anh_ho_so'])
                            if anh_bytes:
                                col_prev, col_btn = st.columns([1, 3])
                                with col_prev:
                                    st.image(anh_bytes, caption="Ảnh avatar sẽ dùng", width=120)
                                with col_btn:
                                    st.caption("Ảnh này sẽ được dùng làm mẫu nhận diện khuôn mặt. "
                                               "Đảm bảo ảnh thấy rõ mặt, không đeo kính đen, không che mặt.")
                                    if st.button("💾 Đăng ký Face ID từ ảnh avatar", type="primary",
                                                 key="btn_luu_face_dangky"):
                                        img_arr = np.frombuffer(anh_bytes, dtype=np.uint8)
                                        img_bgr = cv2.imdecode(img_arr, cv2.IMREAD_COLOR)
                                        db_luu = st.session_state.db_engine.get_connection()
                                        ok, msg = face_id_cham_cong.dang_ky_khuon_mat(db_luu, nv_chon['id'], img_bgr)
                                        db_luu.close()
                                        if ok:
                                            st.cache_data.clear()
                                            st.success(msg)
                                            st.rerun()
                                        else:
                                            st.error(msg)
                            else:
                                st.warning("⚠️ Không tải được ảnh avatar từ Storage. "
                                           "Kiểm tra cấu hình Supabase Storage của tenant.")

                        # Đồng bộ lại toàn bộ (nút dành cho Admin khi muốn cập nhật lại tất cả)
                        if da_dk or co_anh_chua_dk:
                            st.divider()
                            so_co_the_dong_bo = len(co_anh_chua_dk) + len(da_dk)
                            if st.button(f"🔄 Đồng bộ lại Face ID cho TẤT CẢ {so_co_the_dong_bo} nhân viên có ảnh",
                                         key="btn_dong_bo_tat_ca"):
                                thanh_cong, that_bai = 0, []
                                progress = st.progress(0)
                                ds_dong_bo = [x for x in ds_nv_face if x['anh_ho_so']]
                                for i, nv in enumerate(ds_dong_bo):
                                    anh_bytes = get_anh_ho_so_bytes(nv['anh_ho_so'])
                                    if not anh_bytes:
                                        that_bai.append(f"{nv['ho_ten']} (không tải được ảnh)")
                                        continue
                                    img_arr = np.frombuffer(anh_bytes, dtype=np.uint8)
                                    img_bgr = cv2.imdecode(img_arr, cv2.IMREAD_COLOR)
                                    db_luu = st.session_state.db_engine.get_connection()
                                    ok, msg = face_id_cham_cong.dang_ky_khuon_mat(db_luu, nv['id'], img_bgr)
                                    db_luu.close()
                                    if ok:
                                        thanh_cong += 1
                                    else:
                                        that_bai.append(f"{nv['ho_ten']} ({msg})")
                                    progress.progress((i + 1) / len(ds_dong_bo))
                                progress.empty()
                                if thanh_cong:
                                    st.success(f"✅ Đã đồng bộ thành công {thanh_cong} nhân viên.")
                                if that_bai:
                                    st.warning("⚠️ Các trường hợp thất bại:\n" + "\n".join(f"• {x}" for x in that_bai))
                                if thanh_cong:
                                    st.cache_data.clear()
                                    st.rerun()

        # ----- TAB CHECK-IN / CHECK-OUT (camera live) -----
        with tab_face_checkin:
            st.caption("Bấm START để bật camera. Nhân viên đứng trước camera, hệ thống tự ghi "
                       "giờ vào/ra. Xong việc bấm STOP (cùng vị trí nút START) để tắt camera.")

            db_load = st.session_state.db_engine.get_connection()
            danh_sach_emb = face_id_cham_cong.tai_toan_bo_embedding(db_load)

            if not danh_sach_emb:
                st.warning("⚠️ Chưa có nhân viên nào đăng ký khuôn mặt. Vào tab 'Đăng ký khuôn mặt' để thêm.")
                db_load.close()
            else:
                st.caption(f"Đã đăng ký khuôn mặt: {len(danh_sach_emb)} nhân viên.")

                # Cấu hình STUN/TURN — BẮT BUỘC trên Streamlit Cloud
                # (server ở xa, cần relay để truyền video frame)
                from streamlit_webrtc import WebRtcMode
                import ssl

                _RTC_CONFIG = {
                    "iceServers": [
                        {"urls": ["stun:stun.l.google.com:19302"]},
                        {"urls": ["stun:stun1.l.google.com:19302"]},
                        {"urls": ["stun:stun2.l.google.com:19302"]},
                        {"urls": ["stun:stun3.l.google.com:19302"]},
                        {"urls": ["stun:stun4.l.google.com:19302"]},
                    ]
                }

                ctx = webrtc_streamer(
                    key="face_id_checkin",
                    video_processor_factory=face_id_cham_cong.FaceIDVideoProcessor,
                    media_stream_constraints={"video": {"width": 640, "height": 480}, "audio": False},
                    async_processing=True,
                    rtc_configuration=_RTC_CONFIG,
                )

                ket_qua_box = st.empty()

                if ctx.video_processor:
                    ctx.video_processor.conn = db_load
                    ctx.video_processor.danh_sach_embedding = danh_sach_emb
                    ctx.video_processor.cfg = get_cau_hinh_cham_cong_full()

                    dem_sau_nhan_dien = 0  # đếm ngược sau khi nhận diện xong

                    while ctx.state.playing:
                        kq = ctx.video_processor.ket_qua
                        if kq["trang_thai"] == "THANH_CONG":
                            ket_qua_box.success(
                                f"{kq['thong_bao']}\n\n✅ Camera sẽ tự tắt sau {3 - dem_sau_nhan_dien} giây...")
                            dem_sau_nhan_dien += 1
                            if dem_sau_nhan_dien >= 3:
                                break  # thoát loop → camera dừng hiển thị
                        elif kq["trang_thai"] == "DA_DU":
                            ket_qua_box.info(
                                f"{kq['thong_bao']}\n\n✅ Camera sẽ tự tắt sau {3 - dem_sau_nhan_dien} giây...")
                            dem_sau_nhan_dien += 1
                            if dem_sau_nhan_dien >= 3:
                                break
                        elif kq["trang_thai"] == "KHONG_KHOP":
                            ket_qua_box.error(kq["thong_bao"])
                            dem_sau_nhan_dien = 0
                        else:
                            ket_qua_box.info("📷 Đưa mặt vào khung hình camera...")
                            dem_sau_nhan_dien = 0
                        time.sleep(1)

                    # Camera đã tắt (tự động hoặc người dùng bấm STOP)
                    ket_qua_box.success("✅ Đã ghi nhận chấm công! Xem kết quả tại tab '📋 Kết quả hôm nay'.")
                    time.sleep(2)
                    st.rerun()

                try:
                    db_load.close()
                except Exception:
                    pass

            st.divider()
            if st.button("⬅️ Quay lại Chấm công thủ công", key="btn_thoat_faceid"):
                st.session_state.cc_method = 'manual'
                st.rerun()

        # ----- TAB KẾT QUẢ CHẤM CÔNG HÔM NAY -----
        with tab_face_ketqua:
            col_kq1, col_kq2 = st.columns([3, 1])
            with col_kq1:
                ngay_xem_kq = st.date_input("📅 Xem kết quả ngày:", value=date.today(), key="face_kq_ngay")
            with col_kq2:
                st.write("")
                st.write("")
                if st.button("🔄 Tải lại", key="btn_reload_face_kq"):
                    st.rerun()

            db_kq = st.session_state.db_engine.get_connection()
            c_kq = db_kq.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c_kq.execute("""
                SELECT nv.ma_nv, nv.ho_ten, cc.gio_vao, cc.gio_ra, cc.ma_cong, cc.nguon
                FROM cham_cong cc
                JOIN nhan_vien nv ON nv.id = cc.nhan_vien_id
                WHERE cc.ngay = %s AND (cc.gio_vao IS NOT NULL OR cc.gio_ra IS NOT NULL)
                ORDER BY cc.gio_vao NULLS LAST, nv.ho_ten
            """, (ngay_xem_kq,))
            ds_kq = c_kq.fetchall()
            db_kq.close()

            if not ds_kq:
                st.info(f"Chưa có dữ liệu chấm công giờ vào/ra cho ngày {ngay_xem_kq.strftime('%d/%m/%Y')}.")
            else:
                bang_kq = []
                for r in ds_kq:
                    gv = r['gio_vao'].strftime('%H:%M:%S') if r['gio_vao'] else '—'
                    gr = r['gio_ra'].strftime('%H:%M:%S') if r['gio_ra'] else '—'
                    gio_lam = tinh_gio_lam_thuc_te(r['gio_vao'], r['gio_ra'], ngay_xem_kq)
                    bang_kq.append({
                        "Mã NV": r['ma_nv'],
                        "Họ tên": r['ho_ten'],
                        "Giờ vào": gv,
                        "Giờ ra": gr,
                        "Số giờ": gio_lam if gio_lam is not None else '—',
                        "Ký hiệu": r['ma_cong'] or '—',
                        "Nguồn": r['nguon'] or '—',
                    })

                st.dataframe(pd.DataFrame(bang_kq), use_container_width=True, hide_index=True)
                st.caption(f"Tổng: {len(bang_kq)} nhân viên có dữ liệu giờ vào/ra. "
                           "Cột 'Số giờ' chỉ tính khi đã có cả giờ vào và giờ ra.")
        # ----- TAB ĐIỀU CHỈNH CHẤM CÔNG -----
        if tab_face_dieuchinh is not None:
            with tab_face_dieuchinh:
                st.caption("Admin/HR điều chỉnh dữ liệu chấm công khi có sai sót. "
                           "Mọi thay đổi đều được ghi log đầy đủ (ai sửa, lúc nào, giá trị cũ/mới, lý do).")

                # Cấu hình: có bắt buộc phê duyệt không
                can_edit_cc = can_dieu_chinh_bcc()
                yeu_cau_phe_duyet = get_cau_hinh('cc_dieu_chinh_can_duyet', '0') == '1'

                if not can_edit_cc:
                    st.warning("⚠️ Bạn không có quyền điều chỉnh chấm công.")
                else:
                    # Toggle bật/tắt yêu cầu phê duyệt (chỉ Admin)
                    if st.session_state.role == 'admin':
                        with st.expander("⚙️ Cài đặt quy trình điều chỉnh"):
                            yeu_cau_phe_duyet_moi = st.toggle(
                                "Bắt buộc phê duyệt khi điều chỉnh chấm công",
                                value=yeu_cau_phe_duyet,
                                key="cc_toggle_phe_duyet",
                                help="Bật: mọi điều chỉnh tạo yêu cầu chờ Admin/Giám đốc duyệt trước khi có hiệu lực. "
                                     "Tắt: Admin/HR sửa trực tiếp, hệ thống vẫn ghi log đầy đủ.")
                            if yeu_cau_phe_duyet_moi != yeu_cau_phe_duyet:
                                set_cau_hinh('cc_dieu_chinh_can_duyet',
                                             '1' if yeu_cau_phe_duyet_moi else '0',
                                             'Điều chỉnh chấm công cần phê duyệt')
                                st.rerun()

                    st.divider()

                    # -- Phần 1: Tạo điều chỉnh mới --
                    st.markdown("**📝 Tạo yêu cầu điều chỉnh**")

                    db_dc = st.session_state.db_engine.get_connection()
                    c_dc = db_dc.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    c_dc.execute("""
                        SELECT id, ma_nv, ho_ten FROM nhan_vien
                        WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC')
                        ORDER BY ho_ten
                    """)
                    ds_nv_dc = c_dc.fetchall()

                    col_dc1, col_dc2 = st.columns(2)
                    with col_dc1:
                        nv_dc_map = {f"{x['ma_nv']} - {x['ho_ten']}": x['id'] for x in ds_nv_dc}
                        nv_dc_chon = st.selectbox("👤 Nhân viên:", list(nv_dc_map.keys()),
                                                  key="dc_nv_select")
                        nv_dc_id = nv_dc_map[nv_dc_chon]
                    with col_dc2:
                        ngay_dc = st.date_input("📅 Ngày cần điều chỉnh:",
                                                value=date.today(), key="dc_ngay_input")

                    # Lấy dữ liệu hiện tại
                    c_dc.execute("""
                        SELECT id, gio_vao, gio_ra, ma_cong, nguon, ghi_chu
                        FROM cham_cong
                        WHERE nhan_vien_id = %s AND ngay = %s
                    """, (nv_dc_id, ngay_dc))
                    row_cc = c_dc.fetchone()

                    if row_cc:
                        st.markdown("**Dữ liệu hiện tại:**")
                        col_h1, col_h2, col_h3, col_h4 = st.columns(4)
                        col_h1.metric("Giờ vào",
                                      row_cc['gio_vao'].strftime('%H:%M:%S') if row_cc['gio_vao'] else '—')
                        col_h2.metric("Giờ ra",
                                      row_cc['gio_ra'].strftime('%H:%M:%S') if row_cc['gio_ra'] else '—')
                        col_h3.metric("Ký hiệu", row_cc['ma_cong'] or '—')
                        col_h4.metric("Nguồn", row_cc['nguon'] or '—')

                        st.markdown("**Giá trị mới:**")
                        col_m1, col_m2, col_m3 = st.columns(3)
                        with col_m1:
                            gio_vao_moi_dc = st.time_input(
                                "Giờ vào mới",
                                value=row_cc['gio_vao'] or _time(8, 0),
                                key="dc_gio_vao_moi")
                        with col_m2:
                            gio_ra_moi_dc = st.time_input(
                                "Giờ ra mới",
                                value=row_cc['gio_ra'] or _time(17, 0),
                                key="dc_gio_ra_moi")
                        with col_m3:
                            ma_cong_moi_dc = st.selectbox(
                                "Ký hiệu mới",
                                CHAM_CONG_MA_OPTIONS,
                                index=CHAM_CONG_MA_OPTIONS.index(row_cc['ma_cong'])
                                if row_cc['ma_cong'] in CHAM_CONG_MA_OPTIONS else 0,
                                key="dc_ma_cong_moi")

                        ly_do_dc = st.text_area(
                            "📋 Lý do điều chỉnh (bắt buộc):",
                            placeholder="VD: Camera ghi nhận sai giờ ra do mất điện lúc 17:00. "
                                        "Giờ ra thực tế theo bảo vệ xác nhận là 17:30.",
                            key="dc_ly_do", height=80)

                        col_btn1, col_btn2 = st.columns(2)
                        with col_btn1:
                            if st.button("💾 Lưu điều chỉnh", type="primary",
                                         key="btn_luu_dieu_chinh",
                                         disabled=not ly_do_dc.strip()):
                                # Ghi log từng trường thay đổi
                                thay_doi = []
                                if str(gio_vao_moi_dc) != str(row_cc['gio_vao'] or ''):
                                    thay_doi.append(('gio_vao',
                                                     str(row_cc['gio_vao']), str(gio_vao_moi_dc)))
                                if str(gio_ra_moi_dc) != str(row_cc['gio_ra'] or ''):
                                    thay_doi.append(('gio_ra',
                                                     str(row_cc['gio_ra']), str(gio_ra_moi_dc)))
                                if ma_cong_moi_dc != (row_cc['ma_cong'] or ''):
                                    thay_doi.append(('ma_cong',
                                                     row_cc['ma_cong'], ma_cong_moi_dc))

                                if not thay_doi:
                                    st.warning("⚠️ Không có thay đổi nào so với dữ liệu hiện tại.")
                                else:
                                    trang_thai_log = 'CHO_DUYET' if yeu_cau_phe_duyet else 'DA_DUYET'
                                    c_dc2 = db_dc.cursor()

                                    if not yeu_cau_phe_duyet:
                                        # Sửa trực tiếp
                                        c_dc2.execute("""
                                            UPDATE cham_cong
                                            SET gio_vao=%s, gio_ra=%s, ma_cong=%s,
                                                nguon='DIEU_CHINH', updated_at=NOW()
                                            WHERE id=%s
                                        """, (gio_vao_moi_dc, gio_ra_moi_dc,
                                              ma_cong_moi_dc or None, row_cc['id']))

                                    # Ghi audit log (dù có hay không cần duyệt)
                                    for truong, cu, moi in thay_doi:
                                        c_dc2.execute("""
                                            INSERT INTO audit_cham_cong
                                                (cham_cong_id, nhan_vien_id, ngay, truong_sua,
                                                 gia_tri_cu, gia_tri_moi, ly_do, nguoi_sua,
                                                 trang_thai)
                                            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
                                        """, (row_cc['id'], nv_dc_id, ngay_dc,
                                              truong, cu, moi, ly_do_dc.strip(),
                                              st.session_state.get('username', ''),
                                              trang_thai_log))

                                    db_dc.commit()
                                    if yeu_cau_phe_duyet:
                                        st.info("📨 Đã tạo yêu cầu điều chỉnh, chờ phê duyệt.")
                                    else:
                                        st.success("✅ Đã điều chỉnh và ghi log thành công.")
                                        st.cache_data.clear()
                                    st.rerun()
                    else:
                        st.info(f"Không có dữ liệu chấm công ngày "
                                f"{ngay_dc.strftime('%d/%m/%Y')} cho nhân viên này.")

                    db_dc.close()

                    # -- Phần 2: Lịch sử điều chỉnh --
                    st.divider()
                    st.markdown("**📜 Lịch sử điều chỉnh & phê duyệt**")

                    col_ls1, col_ls2 = st.columns(2)
                    with col_ls1:
                        tu_ngay_ls = st.date_input("Từ ngày:",
                                                   value=date.today().replace(day=1),
                                                   key="dc_ls_tu_ngay")
                    with col_ls2:
                        den_ngay_ls = st.date_input("Đến ngày:",
                                                    value=date.today(),
                                                    key="dc_ls_den_ngay")

                    db_ls = st.session_state.db_engine.get_connection()
                    c_ls = db_ls.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    c_ls.execute("""
                        SELECT a.id, nv.ma_nv, nv.ho_ten, a.ngay, a.truong_sua,
                               a.gia_tri_cu, a.gia_tri_moi, a.ly_do,
                               a.nguoi_sua, a.thoi_diem_sua,
                               a.trang_thai, a.nguoi_duyet, a.ghi_chu_duyet
                        FROM audit_cham_cong a
                        JOIN nhan_vien nv ON nv.id = a.nhan_vien_id
                        WHERE a.ngay BETWEEN %s AND %s
                        ORDER BY a.thoi_diem_sua DESC
                        LIMIT 200
                    """, (tu_ngay_ls, den_ngay_ls))
                    ds_ls = c_ls.fetchall()

                    # Nếu đang bật chế độ phê duyệt: hiện nút Duyệt/Từ chối cho Admin
                    cho_duyet = [r for r in ds_ls if r['trang_thai'] == 'CHO_DUYET']
                    if cho_duyet and st.session_state.role == 'admin':
                        st.warning(f"⏳ Có {len(cho_duyet)} yêu cầu đang chờ phê duyệt:")
                        for req in cho_duyet:
                            with st.container(border=True):
                                st.write(f"**{req['ma_nv']} — {req['ho_ten']}** | "
                                         f"Ngày: {req['ngay']} | "
                                         f"Trường: `{req['truong_sua']}` | "
                                         f"{req['gia_tri_cu']} → {req['gia_tri_moi']}")
                                st.caption(f"Lý do: {req['ly_do']} | Người yêu cầu: {req['nguoi_sua']}")
                                col_d1, col_d2, col_d3 = st.columns([2, 1, 1])
                                with col_d2:
                                    if st.button("✅ Duyệt", key=f"duyet_{req['id']}"):
                                        db_duyet = st.session_state.db_engine.get_connection()
                                        c_duyet = db_duyet.cursor()
                                        # Áp dụng thay đổi
                                        c_duyet.execute(f"""
                                            UPDATE cham_cong SET {req['truong_sua']} = %s,
                                                nguon='DIEU_CHINH', updated_at=NOW()
                                            WHERE id = %s
                                        """, (req['gia_tri_moi'], req['cham_cong_id']))
                                        # Cập nhật trạng thái log
                                        c_duyet.execute("""
                                            UPDATE audit_cham_cong
                                            SET trang_thai='DA_DUYET', nguoi_duyet=%s,
                                                thoi_diem_duyet=NOW()
                                            WHERE id=%s
                                        """, (st.session_state.get('username', ''), req['id']))
                                        db_duyet.commit()
                                        db_duyet.close()
                                        st.rerun()
                                with col_d3:
                                    ghi_chu_tc = st.text_input("Lý do từ chối:",
                                                               key=f"tc_note_{req['id']}")
                                    if st.button("❌ Từ chối", key=f"tuchoi_{req['id']}"):
                                        db_tc = st.session_state.db_engine.get_connection()
                                        c_tc = db_tc.cursor()
                                        c_tc.execute("""
                                            UPDATE audit_cham_cong
                                            SET trang_thai='TU_CHOI', nguoi_duyet=%s,
                                                thoi_diem_duyet=NOW(), ghi_chu_duyet=%s
                                            WHERE id=%s
                                        """, (st.session_state.get('username', ''),
                                              ghi_chu_tc, req['id']))
                                        db_tc.commit()
                                        db_tc.close()
                                        st.rerun()

                    if ds_ls:
                        bang_ls = []
                        for r in ds_ls:
                            trang_thai_hien = {
                                'DA_DUYET': '✅ Đã duyệt',
                                'CHO_DUYET': '⏳ Chờ duyệt',
                                'TU_CHOI': '❌ Từ chối',
                            }.get(r['trang_thai'], r['trang_thai'])
                            bang_ls.append({
                                'Mã NV': r['ma_nv'],
                                'Họ tên': r['ho_ten'],
                                'Ngày CC': r['ngay'],
                                'Trường sửa': r['truong_sua'],
                                'Giá trị cũ': r['gia_tri_cu'],
                                'Giá trị mới': r['gia_tri_moi'],
                                'Lý do': r['ly_do'],
                                'Người sửa': r['nguoi_sua'],
                                'Thời điểm': r['thoi_diem_sua'].strftime('%d/%m %H:%M')
                                if r['thoi_diem_sua'] else '',
                                'Trạng thái': trang_thai_hien,
                                'Người duyệt': r['nguoi_duyet'] or '',
                            })
                        st.dataframe(pd.DataFrame(bang_ls),
                                     use_container_width=True, hide_index=True)
                    else:
                        st.info("Không có lịch sử điều chỉnh trong khoảng thời gian này.")

                    db_ls.close()
        
# ========== TÍNH THU NHẬP ==========
elif menu == "💰 Tính thu nhập":
    tinh_thu_nhap.show_tinh_thu_nhap()

# ========== THUẾ HKD ==========
elif menu == "🧾 Thuế HKD":
    thue_hkd.render_thue_hkd(st.session_state.db_engine)

# ========== UPLOAD ==========
elif menu=="📁 Upload hồ sơ" and st.session_state.role=="admin":
    st.title("📁 Quản lý hồ sơ nhân viên")
    tab_upload, tab_list, tab_avatar = st.tabs(["📤 UPLOAD HỒ SƠ", "📋 DANH SÁCH HỒ SƠ", "📸 UPLOAD ẢNH HỒ SƠ"])
    
    with tab_avatar:
        st.subheader("📸 Upload ảnh hồ sơ cho nhân viên")
        st.caption("Chọn nhân viên và tải ảnh lên. Ảnh sẽ được lưu vào cột `anh_ho_so` trong bảng nhân viên.")
        
        # Lấy danh sách nhân viên chưa có ảnh hồ sơ
        db_avatar = st.session_state.db_engine.get_connection()
        c_avatar = db_avatar.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c_avatar.execute("""
            SELECT id, ma_nv, ho_ten FROM nhan_vien 
            WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC') 
            AND (anh_ho_so IS NULL OR anh_ho_so = '')
            ORDER BY id DESC
        """)
        nv_chua_anh = c_avatar.fetchall()
        db_avatar.close()

        if not nv_chua_anh:
            st.success("🎉 Tất cả nhân viên đã có ảnh hồ sơ!")
        else:
            # Tạo dict chọn nhân viên
            nv_map = {f"{x['ma_nv']} - {x['ho_ten']}": x['id'] for x in nv_chua_anh}
            selected_nv_label = st.selectbox("📌 Chọn nhân viên cần upload ảnh:", list(nv_map.keys()))
            selected_nv_id = nv_map[selected_nv_label]
            
            # Lấy thông tin nhân viên đã chọn
            db_detail = st.session_state.db_engine.get_connection()
            c_detail = db_detail.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c_detail.execute("SELECT ma_nv, ho_ten FROM nhan_vien WHERE id = %s", (selected_nv_id,))
            nv_info = c_detail.fetchone()
            db_detail.close()

            if nv_info:
                st.markdown(f"**Nhân viên:** {nv_info['ma_nv']} - {nv_info['ho_ten']}")
                
                # File uploader cho ảnh
                anh_upload = st.file_uploader("Chọn ảnh hồ sơ (png, jpg, jpeg)", type=["png", "jpg", "jpeg"], key="avatar_upload_single")
                
                if anh_upload is not None:
                    st.image(anh_upload, caption="Ảnh xem trước", width=200)
                
                if st.button("📤 UPLOAD ẢNH", type="primary", width='stretch'):
                    if anh_upload is None:
                        st.error("❌ Vui lòng chọn ảnh để upload!")
                    else:
                        # Upload ảnh lên Storage
                        storage_path = upload_anh_ho_so(nv_info['ma_nv'], nv_info['ho_ten'], anh_upload)
                        if storage_path:
                            # Cập nhật đường dẫn vào bảng nhan_vien
                            db_update = st.session_state.db_engine.get_connection()
                            c_update = db_update.cursor()
                            c_update.execute("UPDATE nhan_vien SET anh_ho_so = %s WHERE id = %s", (storage_path, selected_nv_id))
                            db_update.commit()
                            db_update.close()
                            st.success(f"✅ Đã upload ảnh hồ sơ thành công cho {nv_info['ho_ten']}!")
                            st.cache_data.clear()
                            st.rerun()
                        else:
                            st.error("❌ Upload ảnh thất bại. Vui lòng kiểm tra cấu hình Storage!")
    
    with tab_upload:
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT id, ma_nv, ho_ten FROM nhan_vien WHERE trang_thai IN ('DANG_LAM','THU_VIEC') ORDER BY id DESC")
        nvl = c.fetchall()
        db.close()
        
        if nvl:
            nd = {f"{x['ma_nv']} - {x['ho_ten']}": x['id'] for x in nvl}
            id_to_hoten = {x['id']: x['ho_ten'] for x in nvl}
            cn = st.selectbox("📌 Chọn nhân viên:", list(nd.keys()))
            
            col1, col2 = st.columns(2)
            with col1:
                lh = st.selectbox("📂 Loại hồ sơ:", ["BANG_CAP", "CHUNG_CHI", "CCCD", "HOP_DONG", "SO_YEU_LY_LICH", "KHAC"])
            with col2:
                st.markdown("")
                st.caption("💡 **Hướng dẫn: Chọn loại giấy tờ, sau đó chọn file từ thư mục để Upload!**")
            
            fl = st.file_uploader("📎 Chọn file:", type=['pdf', 'jpg', 'png', 'jpeg', 'doc', 'docx'])
            
            if fl:
                st.info(f"📄 Tên file: {fl.name} | 📏 Kích thước: {fl.size/1024:.1f} KB")
            
            if fl and st.button("📤 UPLOAD", type="primary", width='stretch'):
                nid = nd[cn]
                ngay_upload_str = datetime.now().strftime('%Y%m%d')
                safe_name = sanitize_storage_filename(fl.name)
                ho_ten_folder = sanitize_storage_filename(id_to_hoten.get(nid, str(nid)))
                # Cấu trúc: {Họ tên nhân viên}/{Loại hồ sơ}_{ngày upload}_{tên file}
                base_path = f"{ho_ten_folder}/{lh}_{ngay_upload_str}_{safe_name}"

                sb = get_supabase_storage()
                if not sb:
                    st.error("❌ Chưa cấu hình Supabase Storage. Vui lòng khai báo `SUPABASE_URL` và `SUPABASE_KEY` trong secrets/.env.")
                else:
                    try:
                        storage_path = upload_to_storage_unique(
                            sb, SUPABASE_BUCKET, base_path,
                            fl.getvalue(), fl.type
                        )

                        db = st.session_state.db_engine.get_connection()
                        c = db.cursor()
                        c.execute("""
                            INSERT INTO ho_so_nhan_vien (nhan_vien_id, loai_ho_so, ten_file, duong_dan_file, ngay_upload) 
                            VALUES (%s, %s, %s, %s, CURRENT_DATE)
                        """, (nid, lh, fl.name, storage_path))
                        db.commit()
                        db.close()

                        st.success(f"✅ Đã upload thành công lên Supabase Storage!\n📁 Đường dẫn: {storage_path}")
                        st.cache_data.clear()
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Lỗi khi upload lên Supabase Storage: {e}")
        else:
            st.info("⚠️ Chưa có nhân viên nào trong hệ thống!")
    
    with tab_list:
        st.subheader("📋 Danh sách hồ sơ đã upload")
        
        # --- HÀM CACHE CHO PRESIGNED URL ---
        @st.cache_data(ttl=3600, show_spinner=False)
        def get_presigned_url_cached(storage_path: str) -> str:
            """Tạo presigned URL có cache 1 giờ.
            LƯU Ý: sb.storage...create_signed_url() trả về một DICT
            (vd: {'signedURL': '...', 'signedUrl': '...'}), KHÔNG PHẢI chuỗi URL.
            Trước đây hàm này trả thẳng cả dict ra ngoài khiến link hiển thị bị hỏng
            (repr của dict bị browser encode lung tung) -> phải bóc tách chuỗi URL thật
            ra khỏi dict trước khi trả về, và ghép domain nếu URL trả về là đường dẫn tương đối."""
            if not storage_path:
                return ""
            try:
                sb = get_supabase_storage()
                if not sb:
                    return ""
                # Tạo signed URL có hiệu lực 1 giờ
                res = sb.storage.from_(SUPABASE_BUCKET).create_signed_url(storage_path, expires_in=3600)

                # Bóc tách chuỗi URL thật ra khỏi kết quả trả về (tùy version supabase-py
                # có thể là dict với key 'signedURL'/'signedUrl', hoặc đôi khi đã là str sẵn)
                if isinstance(res, dict):
                    signed = res.get('signedURL') or res.get('signedUrl') or res.get('signed_url') or ""
                elif isinstance(res, str):
                    signed = res
                else:
                    signed = ""

                if not signed:
                    return ""

                # Nếu chỉ là đường dẫn tương đối (không có http/https) -> ghép domain Supabase
                if not signed.startswith("http"):
                    supabase_url = ""
                    try:
                        if 'supabase' in st.secrets:
                            supabase_url = st.secrets.supabase.get('url', "")
                    except Exception:
                        pass
                    supabase_url = (supabase_url or "").rstrip("/")
                    signed = f"{supabase_url}{signed}" if supabase_url else signed

                return signed
            except Exception as e:
                print(f"Lỗi tạo presigned URL: {e}")
                return ""
        
        @st.cache_data(ttl=3600, show_spinner=False)
        def get_file_bytes_cached(storage_path: str) -> bytes:
            """Tải file bytes có cache 1 giờ"""
            if not storage_path:
                return None
            try:
                sb = get_supabase_storage()
                if not sb:
                    return None
                return sb.storage.from_(SUPABASE_BUCKET).download(storage_path)
            except Exception as e:
                print(f"Lỗi tải file: {e}")
                return None
        
        # --- Lấy danh sách nhân viên ---
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT id, ma_nv, ho_ten FROM nhan_vien ORDER BY id DESC")
        nvl = c.fetchall()
        db.close()
        
        if nvl:
            nd = {f"{x['ma_nv']} - {x['ho_ten']}": x['id'] for x in nvl}
            selected_nv = st.selectbox("🔍 Chọn nhân viên để xem hồ sơ:", list(nd.keys()), key="view_hoso")
            nv_id = nd[selected_nv]
            
            # --- Lấy danh sách hồ sơ ---
            db = st.session_state.db_engine.get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c.execute("""
                SELECT id, loai_ho_so, ten_file, duong_dan_file, ngay_upload 
                FROM ho_so_nhan_vien 
                WHERE nhan_vien_id = %s 
                ORDER BY ngay_upload DESC, id DESC
            """, (nv_id,))
            hs_list = c.fetchall()
            db.close()
            
            if hs_list:
                # --- Hiển thị danh sách (KHÔNG tải file) ---
                hs_data = []
                for i, hs in enumerate(hs_list, 1):
                    hs_data.append({
                        "STT": i,
                        "Loại hồ sơ": hs['loai_ho_so'],
                        "Tên file gốc": hs['ten_file'],
                        "Ngày upload": format_date(hs['ngay_upload']),
                        "ID": hs['id'],
                        "Đường dẫn": hs['duong_dan_file']
                    })
                df_hs = pd.DataFrame(hs_data)
                st.dataframe(df_hs[['STT', 'Loại hồ sơ', 'Tên file gốc', 'Ngày upload']], 
                            width='stretch', hide_index=True)
                
                st.divider()
                if st.button("❌ THOÁT", width='stretch', key="exit_hoso_list"):
                    st.session_state.pop('selected_nv', None)
                    st.rerun()
                
                # --- Chọn hồ sơ để xem (ON-DEMAND LOADING) ---
                hs_options = {f"{hs['loai_ho_so']} - {hs['ten_file']}": hs for hs in hs_list}
                selected_hs_name = st.selectbox("Chọn hồ sơ:", list(hs_options.keys()), key="select_hs_preview")
                selected_hs = hs_options[selected_hs_name]
                
                # Hiển thị thông tin hồ sơ
                st.markdown(f"""
                **📄 {selected_hs['loai_ho_so']}** - {selected_hs['ten_file']}  
                📅 {format_date(selected_hs['ngay_upload'])}
                """)
                
                # --- Xác định loại file ---
                file_ext = selected_hs['ten_file'].lower().split('.')[-1]
                is_image = file_ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']
                is_pdf = file_ext == 'pdf'
                
                # --- State keys ---
                preview_state_key = f"preview_active_{selected_hs['id']}"
                download_url_key = f"download_url_{selected_hs['id']}"
                preview_data_key = f"preview_data_{selected_hs['id']}"
                
                # Khởi tạo state nếu chưa có
                if preview_state_key not in st.session_state:
                    st.session_state[preview_state_key] = False
                
                # --- Nút PREVIEW (On-demand) ---
                col_preview_btn, col_download_btn = st.columns(2)
                
                with col_preview_btn:
                    if is_image or is_pdf:
                        if st.button("👁️ PREVIEW", width='stretch', type="secondary", 
                                    key=f"preview_btn_{selected_hs['id']}"):
                            # CHỈ TẢI KHI BẤM NÚT.
                            # PDF: dùng signed URL (https thật) để nhúng vào iframe, thay vì
                            # nhúng base64 vào <iframe src="data:...">. Cách cũ hay bị Chrome
                            # chặn với thông báo "Trang này bị chrome chặn" vì trình duyệt/CSP
                            # không cho iframe điều hướng tới data-URI lớn. Dùng URL https thật
                            # thì Chrome hiển thị PDF bình thường.
                            # Ảnh vẫn dùng base64 <img> như cũ vì không gặp lỗi này.
                            if is_pdf:
                                with st.spinner("⏳ Đang tạo link preview..."):
                                    preview_url = get_presigned_url_cached(selected_hs['duong_dan_file'])
                                if preview_url:
                                    st.session_state[preview_data_key] = preview_url
                                    st.session_state[preview_state_key] = True
                                    st.rerun()
                                else:
                                    st.error("❌ Không thể tải file để preview")
                            else:
                                with st.spinner("⏳ Đang tải file preview..."):
                                    file_bytes = get_file_bytes_cached(selected_hs['duong_dan_file'])
                                if file_bytes:
                                    st.session_state[preview_data_key] = file_bytes
                                    st.session_state[preview_state_key] = True
                                    st.rerun()
                                else:
                                    st.error("❌ Không thể tải file để preview")
                    else:
                        st.button("👁️ PREVIEW", disabled=True, width='stretch', 
                                 help="Không thể preview loại file này")
                
                with col_download_btn:
                    if st.button("📥 TẢI HỒ SƠ", width='stretch', 
                               key=f"download_btn_{selected_hs['id']}"):
                        # CHỈ TẠO URL KHI BẤM NÚT
                        with st.spinner("⏳ Đang tạo link tải..."):
                            url = get_presigned_url_cached(selected_hs['duong_dan_file'])
                            if url:
                                st.session_state[download_url_key] = url
                                st.rerun()
                            else:
                                st.error("❌ Không thể tạo link tải")
                
                # --- HIỂN THỊ PREVIEW (nếu có) ---
                if st.session_state.get(preview_state_key, False):
                    preview_payload = st.session_state.get(preview_data_key)
                    if preview_payload:
                        st.markdown("---")
                        st.subheader("📄 Xem trước")
                        
                        if is_image:
                            # Ảnh: preview_payload là bytes -> nhúng base64 như cũ
                            img_base64 = base64.b64encode(preview_payload).decode()
                            st.image(f"data:image/jpeg;base64,{img_base64}", width=400)
                        elif is_pdf:
                            # PDF: preview_payload là signed URL (https thật) -> nhúng thẳng
                            # vào iframe, Chrome không chặn như khi dùng data-URI base64
                            st.markdown(f"""
                            <iframe src="{preview_payload}" 
                                    width="100%" height="600px" style="border:none;border-radius:8px;">
                            </iframe>
                            """, unsafe_allow_html=True)
                            st.caption("⚠️ Nếu trình duyệt vẫn không hiển thị được PDF, hãy dùng nút '📥 TẢI HỒ SƠ' để mở/tải file trực tiếp.")
                        
                        # Nút đóng preview
                        if st.button("❌ Đóng preview", key=f"close_preview_{selected_hs['id']}"):
                            st.session_state[preview_state_key] = False
                            if preview_data_key in st.session_state:
                                del st.session_state[preview_data_key]
                            st.rerun()
                    else:
                        st.warning("⚠️ Không có dữ liệu preview")
                
                # --- HIỂN THỊ LINK TẢI (nếu có) ---
                if st.session_state.get(download_url_key):
                    st.markdown("---")
                    st.success("✅ Link tải đã sẵn sàng (có hiệu lực 1 giờ)")
                    url = st.session_state[download_url_key]
                    st.markdown(f"""
                    <div style="background:#f0fdf4;padding:12px 16px;border-radius:8px;border:1px solid #bbf7d0;">
                        <a href="{url}" target="_blank" style="font-size:16px;font-weight:600;color:#166534;">
                            📥 Tải xuống: {selected_hs['ten_file']}
                        </a>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if st.button("❌ Đóng link tải", key=f"close_download_{selected_hs['id']}"):
                        if download_url_key in st.session_state:
                            del st.session_state[download_url_key]
                        st.rerun()
                
                # --- Nút xóa hồ sơ ---
                st.divider()
                col_del1, col_del2, col_del3 = st.columns([1, 2, 1])
                with col_del2:
                    if st.button("🗑️ XÓA HỒ SƠ NÀY", width='stretch', type="secondary", disabled=not can_edit()):
                        try:
                            sb = get_supabase_storage()
                            if sb:
                                try:
                                    sb.storage.from_(SUPABASE_BUCKET).remove([selected_hs['duong_dan_file']])
                                except Exception as e_storage:
                                    st.warning(f"⚠️ Không xóa được file trên Storage (vẫn xóa bản ghi): {e_storage}")
                            db = st.session_state.db_engine.get_connection()
                            c = db.cursor()
                            c.execute("DELETE FROM ho_so_nhan_vien WHERE id = %s", (selected_hs['id'],))
                            db.commit()
                            db.close()
                            # Xóa cache
                            get_presigned_url_cached.clear()
                            get_file_bytes_cached.clear()
                            st.success(f"✅ Đã xóa hồ sơ: {selected_hs['ten_file']}")
                            st.cache_data.clear()
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Lỗi khi xóa: {e}")
            else:
                st.info(f"📭 Nhân viên này chưa có hồ sơ nào được upload.")
        else:
            st.info("⚠️ Chưa có nhân viên nào trong hệ thống!")

# ========== DANH MỤC CHỨC DANH ==========
elif menu == "⚙️ Danh mục" and st.session_state.role in ("admin", "xem_toan_bo"):
    st.markdown(f"# {i18n.tm('⚙️ Danh mục cấu hình theo doanh nghiệp')}", unsafe_allow_html=True)
    st.caption("Mỗi khách hàng tự đặt tên Phòng ban, Chức danh, Loại hợp đồng, Trình độ học vấn phù hợp với cơ cấu công ty mình — không ảnh hưởng đến khách hàng khác.")

    def _quan_ly_danh_muc_don_gian(ten_bang, cot_ten, tieu_de, placeholder):
        """Hàm dùng chung để quản lý CRUD cho các bảng danh mục dạng đơn giản
        (id, cột tên, thu_tu, trang_thai) — tránh lặp code cho từng loại danh mục."""
        with st.expander(f"➕ Thêm {tieu_de.lower()} mới", expanded=False):
            ten_moi = st.text_input("Tên", key=f"add_{ten_bang}", placeholder=placeholder)
            if st.button("💾 Lưu", key=f"btn_add_{ten_bang}", disabled=not can_edit()):
                if ten_moi.strip():
                    try:
                        if ten_bang == "danh_muc_phong_ban":
                            # Phòng ban: dùng chuẩn hóa kiểu tiếng Việt (không viết hoa mọi từ)
                            ten_chuan_hoa = chuan_hoa_ten_phong_ban(ten_moi)
                        else:
                            ten_chuan_hoa = ten_moi.strip()[:1].upper() + ten_moi.strip()[1:]  # chỉ viết hoa chữ cái đầu
                        db = st.session_state.db_engine.get_connection(); c = db.cursor()
                        c.execute(f"INSERT INTO {ten_bang} ({cot_ten}) VALUES (%s) ON CONFLICT DO NOTHING",
                                  (ten_chuan_hoa,))
                        db.commit(); db.close()
                        st.success(f"✅ Đã thêm: {ten_chuan_hoa}"); st.cache_data.clear(); st.rerun()
                        st.cache_data.clear()
                    except Exception as e:
                        st.error(f"❌ Lỗi: {e}")
                else:
                    st.error("Vui lòng nhập tên!")

        db = st.session_state.db_engine.get_connection(); c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute(f"SELECT id, {cot_ten}, trang_thai FROM {ten_bang} ORDER BY thu_tu, id")
        ds = c.fetchall(); db.close()
        if ds:
            df = pd.DataFrame(ds); df.columns = ['ID', tieu_de, 'Trạng thái']
            st.dataframe(df, width='stretch', hide_index=True)
            idx_xoa = st.number_input("Nhập ID cần xoá:", min_value=1, step=1, key=f"del_{ten_bang}")
            if st.button("🗑️ Xoá", key=f"btn_del_{ten_bang}"):
                if not can_delete():
                    st.error("❌ Bạn không có quyền xóa dữ liệu!")
                else:
                    db = st.session_state.db_engine.get_connection()
                    c = db.cursor()
                    c.execute(f"DELETE FROM {ten_bang} WHERE id=%s", (idx_xoa,))
                    db.commit()
                    db.close()
                    st.success("🗑️ Đã xoá!")
                    st.cache_data.clear()
                    st.rerun()
        else:
            st.info(f"Chưa có {tieu_de.lower()} nào.")

    tab_pb, tab_cd, tab_hd, tab_hv, tab_mau_hd, tab_cv, tab_cty, tab_cc = st.tabs([
        "🏢 Phòng ban", "💼 Chức danh", "📄 Loại hợp đồng", "🎓 Trình độ học vấn",
        "📃 Mẫu Hợp đồng", "🎖️ Chức vụ", "⚙️ Cấu hình Doanh nghiệp", "🕒 Chấm công"
    ])

    with tab_pb:
        st.info("ℹ️ Danh sách phòng ban dùng cho các form Thêm/Sửa nhân viên lấy **trực tiếp từ "
                "danh mục bên dưới** ")

        if st.session_state.role in ("admin", "xem_toan_bo"):
            with st.expander("🧹 Dọn dữ liệu phòng ban cũ (chạy 1 lần)", expanded=False):
                st.caption(
                    "Quét toàn bộ bảng nhân viên, chuẩn hóa lại giá trị phòng ban theo đúng "
                    "danh mục chuẩn (VD: 'phòng hành chính nhân sự' / 'Phòng Hành Chính Nhân Sự' "
                    "→ 'Phòng Hành chính Nhân sự'). An toàn khi chạy nhiều lần."
                )
                if st.button("🧹 Chạy chuẩn hóa ngay", key="btn_chuan_hoa_pb"):
                    db_ch = st.session_state.db_engine.get_connection()
                    c_ch = db_ch.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    c_ch.execute("SELECT id, phong_ban_lam_viec FROM nhan_vien WHERE phong_ban_lam_viec IS NOT NULL AND phong_ban_lam_viec != ''")
                    rows_ch = c_ch.fetchall()
                    so_da_sua = 0
                    for r in rows_ch:
                        chuan = chuan_hoa_ten_phong_ban(r['phong_ban_lam_viec'])
                        if chuan != r['phong_ban_lam_viec']:
                            c_ch.execute("UPDATE nhan_vien SET phong_ban_lam_viec=%s WHERE id=%s", (chuan, r['id']))
                            so_da_sua += 1
                    db_ch.commit()
                    db_ch.close()
                    st.success(f"✅ Đã kiểm tra {len(rows_ch)} nhân viên, chuẩn hóa lại {so_da_sua} bản ghi.")
                    st.cache_data.clear()
                    st.rerun()

        _quan_ly_danh_muc_don_gian("danh_muc_phong_ban", "ten_phong_ban", "Phòng ban", "VD: Kinh doanh")

    with tab_cd:
        # Chức danh tiếp tục dùng bảng vi_tri_cong_tac có sẵn để không phá vỡ dữ liệu cũ
        with st.expander("➕ Thêm chức danh mới", expanded=False):
            with st.form("add_chuc_danh"):
                ten_moi = st.text_input("Tên chức danh *"); mo_ta = st.text_area("Mô tả")
                if st.form_submit_button("💾 LƯU", disabled=not can_edit()):
                    if ten_moi:
                        db = st.session_state.db_engine.get_connection(); c = db.cursor()
                        c.execute("SELECT COALESCE(MIN(t1.id + 1), 1) FROM vi_tri_cong_tac t1 LEFT JOIN vi_tri_cong_tac t2 ON t1.id + 1 = t2.id WHERE t2.id IS NULL AND t1.id >= 1")
                        id_trong = c.fetchone()[0]
                        c.execute("SELECT COALESCE(MAX(id),0) FROM vi_tri_cong_tac")
                        id_max = c.fetchone()[0]
                        id_moi = id_trong if id_trong <= id_max + 1 else id_max + 1
                        c.execute("INSERT INTO vi_tri_cong_tac (id, ten_vi_tri, ghi_chu) VALUES (%s, %s, %s)", (id_moi, ten_moi, mo_ta))
                        db.commit(); db.close(); st.success(f"✅ Đã thêm: {ten_moi}"); st.cache_data.clear(); st.rerun()
                    else: st.error("Tên chức danh không được để trống!")
        st.subheader("📋 Danh sách chức danh")
        db = st.session_state.db_engine.get_connection(); c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT id, ten_vi_tri, ghi_chu FROM vi_tri_cong_tac ORDER BY id")
        ds = c.fetchall(); db.close()
        if ds:
            df = pd.DataFrame(ds); df.columns = ['ID', 'Tên chức danh', 'Ghi chú']; st.dataframe(df, width='stretch', hide_index=True)
            st.divider(); cdx = st.number_input("Nhập ID cần xóa:", min_value=1, step=1)
            if st.button("🗑️ XÓA", key="del_cd", disabled=not can_edit()):
                db = st.session_state.db_engine.get_connection(); c = db.cursor()
                c.execute("DELETE FROM vi_tri_cong_tac WHERE id=%s", (cdx,)); db.commit(); db.close(); st.success("🗑️ Đã xóa!"); st.cache_data.clear(); st.rerun()
        else: st.info("Chưa có chức danh nào")

    with tab_hd:
        _quan_ly_danh_muc_don_gian("danh_muc_loai_hop_dong", "ten_loai_hd", "Loại hợp đồng", "VD: Hợp đồng thời vụ")

    with tab_hv:
        _quan_ly_danh_muc_don_gian("danh_muc_trinh_do_hoc_van", "ten_trinh_do", "Trình độ học vấn", "VD: Cử nhân")

    with tab_mau_hd:
        st.caption("Tuỳ chỉnh nội dung từng Điều trong Hợp đồng lao động (HĐLĐ) và Hợp đồng thử việc (HĐTV). "
                   "Điều nào chưa tuỳ chỉnh sẽ tự dùng nội dung mặc định. "
                   "Có thể dùng {vi_tri}, {ngay_hieu_luc}, {ten_cong_ty} (HĐLĐ - Điều 1, Điều 5) hoặc {vi_tri}, {ngay_bat_dau}, {ngay_ket_thuc}, {ten_cong_ty} (HĐTV - Điều 1, Điều 5) "
                   "— hệ thống sẽ tự thay bằng thông tin thực tế của từng nhân viên/công ty khi in. "
                   "Dòng bắt đầu bằng '## ' sẽ in đậm làm tiêu đề phụ (VD: '## 1. Nghĩa vụ:').")

        loai_hd_chon = st.radio("Chọn loại hợp đồng:", ["HĐLĐ (không xác định thời hạn)", "HĐTV (thử việc)"],
                                 horizontal=True, key="mau_hd_loai")
        loai_hd_ma = "HDLD" if loai_hd_chon.startswith("HĐLĐ") else "HDTV"
        mac_dinh = DEFAULT_DIEU_HDLD if loai_hd_ma == "HDLD" else DEFAULT_DIEU_HDTV

        tuy_chinh_hien_tai = get_all_dieu_hop_dong(loai_hd_ma)
        ds_ma_dieu_hien_thi = get_ds_ma_dieu(tuy_chinh_hien_tai)
        MAC_DINH_KEYS = ("can_cu", "dieu1", "dieu2", "dieu3", "dieu4", "dieu5")

        for ma_dieu in ds_ma_dieu_hien_thi:
            la_mac_dinh = ma_dieu in MAC_DINH_KEYS
            mac_dinh_tieu_de, mac_dinh_noi_dung = mac_dinh.get(ma_dieu, ("", ""))
            if la_mac_dinh and not mac_dinh_tieu_de and ma_dieu not in mac_dinh:
                continue  # loại HĐ này không có điều này (VD HĐTV không có dieu1 mặc định cũ)
            info_hien_tai = tuy_chinh_hien_tai.get(ma_dieu)
            hien_tai_tieu_de = info_hien_tai[0] if info_hien_tai else mac_dinh_tieu_de
            hien_tai_noi_dung = info_hien_tai[1] if info_hien_tai else mac_dinh_noi_dung
            hien_tai_thu_tu = info_hien_tai[2] if (info_hien_tai and len(info_hien_tai) > 2 and info_hien_tai[2]) else 0
            with st.expander(f"📝 {hien_tai_tieu_de or ma_dieu}" + ("" if la_mac_dinh else "  🆕"), expanded=False):
                da_tuy_chinh = info_hien_tai is not None
                if da_tuy_chinh:
                    tieu_de_moi = st.text_input("Tiêu đề Điều:", value=hien_tai_tieu_de, key=f"mau_hd_td_{loai_hd_ma}_{ma_dieu}")
                    noi_dung_moi = st.text_area("Nội dung:", value=hien_tai_noi_dung, height=220, key=f"mau_hd_nd_{loai_hd_ma}_{ma_dieu}")
                else:
                    st.caption("💡 Nội dung hiển thị mờ bên dưới chỉ là **gợi ý mặc định**, chưa phải nội dung đã lưu. "
                               "Nhập nội dung để tạo bản tuỳ chỉnh riêng, hoặc để trống & không bấm Lưu để tiếp tục dùng mặc định.")
                    tieu_de_moi = st.text_input("Tiêu đề Điều:", value="", placeholder=hien_tai_tieu_de, key=f"mau_hd_td_{loai_hd_ma}_{ma_dieu}")
                    noi_dung_moi = st.text_area("Nội dung:", value="", placeholder=hien_tai_noi_dung, height=220, key=f"mau_hd_nd_{loai_hd_ma}_{ma_dieu}")
                if not la_mac_dinh:
                    thu_tu_moi = st.number_input("Vị trí hiển thị (số nhỏ hơn đứng trước, chèn xen giữa các Điều mặc định):",
                                                  min_value=1, max_value=999, value=int(hien_tai_thu_tu) or 999,
                                                  key=f"mau_hd_tt_{loai_hd_ma}_{ma_dieu}")
                else:
                    thu_tu_moi = 0
                col_luu, col_reset, col_xoa = st.columns(3)
                with col_luu:
                    if st.button("💾 Lưu", key=f"mau_hd_save_{loai_hd_ma}_{ma_dieu}", width='stretch', type="primary", disabled=not can_edit()):
                        if not da_tuy_chinh and not tieu_de_moi.strip() and not noi_dung_moi.strip():
                            st.warning("⚠️ Bạn chưa nhập nội dung tuỳ chỉnh nào (nội dung mờ chỉ là gợi ý). "
                                       "Vẫn tiếp tục dùng nội dung mặc định, không có gì để lưu.")
                        else:
                          try:
                            db = st.session_state.db_engine.get_connection()
                            c = db.cursor()
                            c.execute("""
                                INSERT INTO mau_dieu_hop_dong (loai_hd, ma_dieu, tieu_de, noi_dung, thu_tu, updated_at)
                                VALUES (%s, %s, %s, %s, %s, NOW())
                                ON CONFLICT (loai_hd, ma_dieu) DO UPDATE
                                SET tieu_de = EXCLUDED.tieu_de, noi_dung = EXCLUDED.noi_dung,
                                    thu_tu = EXCLUDED.thu_tu, updated_at = NOW()
                            """, (loai_hd_ma, ma_dieu, tieu_de_moi, noi_dung_moi, thu_tu_moi))
                            db.commit(); db.close()
                            get_all_dieu_hop_dong.clear()
                            st.success(f"✅ Đã lưu {ma_dieu}")
                            st.rerun()
                          except Exception as e:
                            st.error(f"❌ Lỗi: {e}")
                with col_reset:
                    if la_mac_dinh:
                        if st.button("↩️ Khôi phục mặc định", key=f"mau_hd_reset_{loai_hd_ma}_{ma_dieu}", width='stretch', disabled=not can_edit()):
                            try:
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("DELETE FROM mau_dieu_hop_dong WHERE loai_hd=%s AND ma_dieu=%s", (loai_hd_ma, ma_dieu))
                                db.commit(); db.close()
                                get_all_dieu_hop_dong.clear()
                                st.success("✅ Đã khôi phục mặc định")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi: {e}")
                with col_xoa:
                    if not la_mac_dinh:
                        if st.button("🗑️ Xoá Điều này", key=f"mau_hd_xoa_{loai_hd_ma}_{ma_dieu}", width='stretch'):
                            try:
                                db = st.session_state.db_engine.get_connection()
                                c = db.cursor()
                                c.execute("DELETE FROM mau_dieu_hop_dong WHERE loai_hd=%s AND ma_dieu=%s", (loai_hd_ma, ma_dieu))
                                db.commit(); db.close()
                                get_all_dieu_hop_dong.clear()
                                st.success("✅ Đã xoá Điều")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi: {e}")

        st.divider()
        with st.expander("➕ Thêm Điều mới", expanded=False):
            tuy_chinh_hdld_all = get_all_dieu_hop_dong("HDLD")
            tuy_chinh_hdtv_all = get_all_dieu_hop_dong("HDTV")
            ma_dieu_de_xuat = sinh_ma_dieu_moi(tuy_chinh_hdld_all, tuy_chinh_hdtv_all)
            st.caption(f"Mã Điều mới sẽ được tạo tự động: **{ma_dieu_de_xuat}**")
            tieu_de_them = st.text_input("Tiêu đề Điều mới:", placeholder="VD: Điều 6. Bảo mật thông tin:", key=f"mau_hd_them_td_{loai_hd_ma}")
            noi_dung_them = st.text_area("Nội dung:", height=180, key=f"mau_hd_them_nd_{loai_hd_ma}",
                                          placeholder="-    Nội dung dòng 1;\n-    Nội dung dòng 2;\n## Tiêu đề phụ in đậm\n-    Nội dung...")
            thu_tu_them = st.number_input("Vị trí hiển thị (số nhỏ hơn đứng trước, VD: 6 = ngay sau Điều 5):",
                                           min_value=1, max_value=999, value=6, key=f"mau_hd_them_tt_{loai_hd_ma}")
            if st.button("➕ Thêm Điều này", key=f"mau_hd_them_btn_{loai_hd_ma}", type="primary", disabled=not can_edit()):
                if not tieu_de_them.strip():
                    st.error("⚠️ Vui lòng nhập tiêu đề Điều!")
                else:
                    try:
                        db = st.session_state.db_engine.get_connection()
                        c = db.cursor()
                        c.execute("""
                            INSERT INTO mau_dieu_hop_dong (loai_hd, ma_dieu, tieu_de, noi_dung, thu_tu, updated_at)
                            VALUES (%s, %s, %s, %s, %s, NOW())
                        """, (loai_hd_ma, ma_dieu_de_xuat, tieu_de_them, noi_dung_them, thu_tu_them))
                        db.commit(); db.close()
                        get_all_dieu_hop_dong.clear()
                        st.success(f"✅ Đã thêm {ma_dieu_de_xuat} vào {loai_hd_chon}")
                        st.cache_data.clear()
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Lỗi: {e}")

    with tab_cv:
        st.subheader("🎖️ Danh mục Chức vụ")
        st.caption("Chức vụ (Trưởng phòng, Phó phòng, Tổ trưởng...) có thể gọi khác nhau tuỳ doanh nghiệp — quản lý danh mục riêng tại đây.")
        try:
            db_cv0 = st.session_state.db_engine.get_connection()
            c_cv0 = db_cv0.cursor()
            c_cv0.execute("""
                CREATE TABLE IF NOT EXISTS chuc_vu_danh_muc (
                    id SERIAL PRIMARY KEY,
                    ten_chuc_vu VARCHAR(150) UNIQUE NOT NULL,
                    thu_tu INT DEFAULT 0,
                    trang_thai VARCHAR(20) DEFAULT 'Hoạt động'
                )
            """)
            db_cv0.commit(); db_cv0.close()
        except Exception as e:
            st.error(f"❌ Lỗi khởi tạo danh mục Chức vụ: {e}")
        _quan_ly_danh_muc_don_gian('chuc_vu_danh_muc', 'ten_chuc_vu', 'Chức vụ', 'VD: Trưởng phòng, Tổ trưởng, Phó Tổng Giám đốc...')

    with tab_cty:
        st.subheader("⚙️ Cấu hình chung của Doanh nghiệp")
        st.caption("Toàn bộ thiết lập áp dụng cho doanh nghiệp bạn — tập trung quản lý tại đây thay vì rải rác nhiều trang.")

        tenant_info = st.session_state.get('tenant', {}) or {}
        st.text_input(
            "🏷️ Mã công ty (ma_cty)", value=tenant_info.get('ma_cty', 'CHL'), disabled=True,
            help="Mã công ty được cấp khi khởi tạo tài khoản. Không tự đổi tại đây vì sẽ làm sai lệch số văn bản/hợp đồng đã phát hành theo mã cũ — liên hệ nhà cung cấp app nếu thực sự cần đổi."
        )

        st.divider()
        st.markdown("**📍 Thông tin dùng trong Hợp đồng / Quyết định / Hồ sơ BHXH**")
        st.caption("Các giá trị này chỉ áp dụng cho công ty bạn, không ảnh hưởng các khách hàng khác dùng chung app. "
                   "Điền 1 lần, hệ thống sẽ tự điền vào Điều 2 quyết định, dòng ký hợp đồng, và các form liên quan.")
        col_dc1, col_dc2, col_dc3 = st.columns(3)
        with col_dc1:
            dia_diem_moi = st.text_input(
                "Địa danh ký văn bản (VD: Quảng Trị, Hà Nội...):",
                value=get_cau_hinh('dia_diem', 'Quảng Trị'), key="cty_dia_diem_input",
                help="Dùng cho dòng '..., ngày ... tháng ... năm ...' ở Quyết định nhân sự và Hợp đồng lao động."
            )
        with col_dc2:
            noi_lam_viec_moi = st.text_input(
                "Nơi làm việc mặc định:", value=get_cau_hinh('noi_lam_viec', 'Cảng THQT Hòn La'),
                key="cty_noi_lam_viec_input"
            )
        with col_dc3:
            noi_cap_cccd_moi = st.text_input(
                "Nơi cấp CCCD:", value=get_cau_hinh('noi_cap_cccd', 'Cục QLHC về TTXH - Bộ Công An'),
                key="cty_noi_cap_cccd_input"
            )
        if st.button("💾 Lưu thông tin công ty", key="btn_save_thong_tin_cty"):
            if not can_edit():
                st.error("❌ Bạn không có quyền chỉnh sửa!")
            else:
                set_cau_hinh('dia_diem', dia_diem_moi.strip(), 'Địa danh ký văn bản')
                set_cau_hinh('noi_lam_viec', noi_lam_viec_moi.strip(), 'Nơi làm việc mặc định')
                set_cau_hinh('noi_cap_cccd', noi_cap_cccd_moi.strip(), 'Nơi cấp CCCD mặc định')
                st.success("✅ Đã lưu thông tin công ty!")
                st.cache_data.clear()
                st.rerun()


        cv_option_hien_tai = get_cv_danh_so_option()
        cv_option_moi = st.radio(
            "Phương án đánh số công văn đi:",
            options=['CHUNG', 'RIENG'],
            index=0 if cv_option_hien_tai == 'CHUNG' else 1,
            format_func=lambda x: "📌 Số chung cho tất cả loại công văn" if x == 'CHUNG' else "📌 Mỗi loại công văn có số riêng",
            key="cty_cv_option_radio"
        )
        st.caption("💡 Muốn xem trạng thái số hiện tại theo từng loại hoặc đặt lại số, vào menu "
                   "**Quản lý Công văn & HĐ kinh tế → ⚙️ Cấu hình đánh số công văn**.")

        st.markdown("**📄 Đánh số Hợp đồng kinh tế (HĐKT)**")
        st.caption("Mẫu số: **stt/năm/Prefix-ma_cty** (VD: 04/2026/HĐKT-CHL)")
        hdkt_prefix_hien_tai = get_hdkt_prefix()
        hdkt_prefix_moi = st.text_input("Prefix đánh số HĐKT:", value=hdkt_prefix_hien_tai, key="cty_hdkt_prefix_input")

        st.divider()
        st.markdown("**📋 Hạn nộp Báo cáo Tăng/Giảm BHXH hàng tháng**")
        st.caption("Mỗi doanh nghiệp có 1 ngày chốt hạn riêng trong tháng (VD: CHL nộp trước ngày 20). "
                   "Dashboard sẽ tự cảnh báo tăng dần mức độ trong 5 ngày trước hạn.")
        han_bhxh_hien_tai = get_han_nop_bhxh()
        han_bhxh_moi = st.number_input(
            "Ngày trong tháng phải nộp (1-28):", min_value=1, max_value=28,
            value=han_bhxh_hien_tai, step=1, key="cty_han_bhxh_input"
        )

        st.divider()
        st.markdown("**💰 Phần mềm tính lương**")
        st.caption("Mỗi công ty có thể có 1 công thức tính lương RIÊNG — file `salary/salary_{Mã số thuế}.py` "
                   "(đặt cùng cấp với app.py). Hệ thống TỰ ĐỘNG chọn đúng file của công ty bạn theo Mã số "
                   "thuế, không cần chọn thủ công. Công ty nào chưa có file riêng sẽ tự dùng công thức mặc "
                   "định `salary/salary_demo.py`.")
        _mst_hien_tai = (tenant_info.get('ma_so_thue') or '').strip()
        if _mst_hien_tai:
            _duong_dan_rieng = os.path.join(os.path.dirname(__file__), 'salary', f'salary_{_mst_hien_tai}.py')
            if os.path.exists(_duong_dan_rieng):
                st.success(f"✅ Đang dùng công thức lương RIÊNG của công ty bạn: `salary/salary_{_mst_hien_tai}.py`")
                st.cache_data.clear()
            else:
                st.info(f"ℹ️ Công ty bạn chưa có công thức lương riêng (`salary/salary_{_mst_hien_tai}.py` chưa "
                        "tồn tại) — đang dùng công thức mặc định `salary/salary_demo.py`. Liên hệ đơn vị triển "
                        "khai App nếu cần công thức tính lương riêng cho doanh nghiệp bạn.")
        else:
            st.info("ℹ️ Đang dùng công thức mặc định `salary/salary_demo.py` (chưa xác định được Mã số thuế "
                    "của công ty bạn — cập nhật tại 'Danh sách khách hàng (Tenants)' phía Quản trị hệ thống).")

        st.divider()
        if st.button("💾 SAVE CẤU HÌNH", type="primary", width='stretch', key="btn_save_cau_hinh_cty", disabled=not can_edit()):
            loi_luu = []
            if not update_cv_danh_so_option(cv_option_moi):
                loi_luu.append("Đánh số công văn")
            if not update_hdkt_prefix(hdkt_prefix_moi.strip()):
                loi_luu.append("Prefix HĐKT")
            if not update_han_nop_bhxh(int(han_bhxh_moi)):
                loi_luu.append("Hạn nộp BHXH")
            
            if loi_luu:
                st.error(f"❌ Lưu thất bại một số mục: {', '.join(loi_luu)}")
            else:
                st.success("✅ Đã lưu toàn bộ cấu hình doanh nghiệp!")
                st.cache_data.clear()
                st.rerun()
    
    with tab_cc:
        st.subheader("🕒 Cấu hình Chấm công")
        st.caption("Áp dụng cho toàn bộ doanh nghiệp bạn — dùng làm cơ sở tính công chuẩn, "
                   "tăng ca, phép năm và cho module Chấm công Face ID.")

        # === PHƯƠNG THỨC CHẤM CÔNG + VAI TRÒ BCC ===
        st.markdown("**📱 Phương thức chấm công**")
        _MAP_PT_DM = {'THU_CONG': 'manual', 'MAY_VAN_TAY': 'fingerprint', 'FACE_ID': 'faceid'}
        _MAP_PT_LABEL_DM = {'THU_CONG': '📝 Thủ công', 'MAY_VAN_TAY': '📥 Máy vân tay', 'FACE_ID': '👤 Face ID'}
        phuong_thuc_dm = get_cau_hinh('cc_phuong_thuc', 'THU_CONG')

        col_pt1, col_pt2 = st.columns(2)
        with col_pt1:
            pt_chon_dm = st.selectbox(
                "Phương thức đang áp dụng:",
                list(_MAP_PT_DM.keys()),
                format_func=lambda k: _MAP_PT_LABEL_DM[k],
                index=list(_MAP_PT_DM.keys()).index(phuong_thuc_dm),
                key="cc_cfg_phuong_thuc_dm",
                disabled=not can_edit(),
            )
        with col_pt2:
            st.markdown("**👥 Vai trò điều chỉnh BCC**")
            vai_tro_hien_tai = get_cau_hinh('cc_vai_tro_dieu_chinh_bcc', 'admin,admin_bcc')
            TAT_CA_VAI_TRO = ['admin', 'admin_bcc', 'hr', 'kt_luong', 'truong_phong', 'van_thu']
            ds_hien_tai = [r.strip() for r in vai_tro_hien_tai.split(',') if r.strip()]
            ds_chon = st.multiselect(
                "Vai trò:",
                TAT_CA_VAI_TRO,
                default=[r for r in ds_hien_tai if r in TAT_CA_VAI_TRO],
                key="cc_cfg_vai_tro_bcc_dm",
                disabled=not can_edit(),
            )

        # Nút lưu chung cho cả phương thức + vai trò
        co_thay_doi = (pt_chon_dm != phuong_thuc_dm) or (ds_chon != ds_hien_tai)
        if co_thay_doi:
            if st.button("💾 Lưu phương thức & vai trò BCC", key="btn_luu_pt_vt", type="primary",
                         disabled=not can_edit()):
                if pt_chon_dm != phuong_thuc_dm:
                    set_cau_hinh('cc_phuong_thuc', pt_chon_dm, 'Phương thức chấm công')
                if ds_chon != ds_hien_tai:
                    set_cau_hinh('cc_vai_tro_dieu_chinh_bcc', ','.join(ds_chon), 'Vai trò điều chỉnh BCC')
                    set_cau_hinh('cc_vai_tro_edit_bcc', ','.join(ds_chon), 'Vai trò nhập BCC')
                st.success("✅ Đã lưu!")
                st.cache_data.clear()
                st.rerun()

        st.divider()

        cc = get_cau_hinh_cham_cong_full()

        st.markdown("**⏰ Giờ làm việc & ca**")
        col1, col2, col3 = st.columns(3)
        with col1:
            gio_vao_moi = st.time_input("Giờ vào chuẩn", value=cc['gio_vao'], key="cc_gio_vao_input")
            gio_bd_dem_moi = st.time_input("Giờ bắt đầu ca đêm (tính TCĐ)", value=cc['gio_bat_dau_ca_dem'], key="cc_gio_bd_dem_input")
        with col2:
            gio_ra_moi = st.time_input("Giờ ra chuẩn", value=cc['gio_ra'], key="cc_gio_ra_input")
            gio_lam_chuan_moi = st.number_input("Giờ làm chuẩn/ngày", min_value=1.0, max_value=12.0,
                                                 value=cc['gio_lam_chuan_ngay'], step=0.5, key="cc_gio_lam_chuan_input")
        with col3:
            phut_tre_moi = st.number_input("Số phút cho phép trễ", min_value=0, max_value=120,
                                            value=cc['phut_tre'], step=5, key="cc_phut_tre_input")
            so_ngay_tuan_moi = st.number_input("Số ngày làm việc/tuần", min_value=1, max_value=7,
                                                value=cc['so_ngay_lam_viec_tuan'], step=1, key="cc_so_ngay_tuan_input")

        st.divider()
        st.markdown("**☕ Nghỉ giữa ca**")
        col_nghi1, col_nghi2 = st.columns(2)
        with col_nghi1:
            ap_dung_nghi_moi = st.toggle(
                "Tự động trừ thời gian nghỉ giữa ca khi tính giờ làm",
                value=cc['ap_dung_nghi_giua_ca'], key="cc_ap_dung_nghi_input",
                help="Bật: số giờ làm = (giờ ra − giờ vào) − thời gian nghỉ giữa ca. "
                     "Tắt: tính nguyên từ giờ vào đến giờ ra.")
        with col_nghi2:
            phut_nghi_moi = st.number_input(
                "Thời gian nghỉ giữa ca (phút)",
                min_value=0, max_value=180,
                value=cc['phut_nghi_giua_ca'], step=5,
                key="cc_phut_nghi_input",
                disabled=not ap_dung_nghi_moi,
                help="Thường là 60 phút (nghỉ trưa). Sẽ bị trừ ra khi tính số giờ làm thực tế.")

        st.divider()
        st.markdown("**📈 Cách tính tăng ca**")
        cach_tinh_tc_moi = st.radio(
            "Chọn cách tính lương tăng ca:",
            options=["HE_SO", "DON_GIA"],
            index=0 if cc['cach_tinh_tang_ca'] == "HE_SO" else 1,
            format_func=lambda x: "Hệ số % trên lương (VD: 150%, 200%...)" if x == "HE_SO" else "Đơn giá cố định (đồng/giờ)",
            horizontal=True, key="cc_cach_tinh_tc_input"
        )
        st.caption("💡 Doanh nghiệp trả TC theo đơn giá cố định (VD: CHL) chọn 'Đơn giá cố định'; "
                   "doanh nghiệp trả theo % lương cơ bản/lương đóng BH chọn 'Hệ số % trên lương'.")

        if cach_tinh_tc_moi == "HE_SO":
            col4, col5, col6, col7 = st.columns(4)
            with col4:
                he_so_tc_moi = st.number_input("TC ngày thường", min_value=1.0, max_value=5.0,
                                                value=cc['he_so_tc_thuong'], step=0.1, key="cc_he_so_tc_input")
            with col5:
                he_so_tcn_moi = st.number_input("TCN (Chủ nhật)", min_value=1.0, max_value=5.0,
                                                 value=cc['he_so_tc_chu_nhat'], step=0.1, key="cc_he_so_tcn_input")
            with col6:
                he_so_tcl_moi = st.number_input("TCL (ngày lễ)", min_value=1.0, max_value=5.0,
                                                 value=cc['he_so_tc_le'], step=0.1, key="cc_he_so_tcl_input")
            with col7:
                he_so_tcd_moi = st.number_input("TCĐ (cộng thêm, đêm)", min_value=1.0, max_value=3.0,
                                                 value=cc['he_so_tc_dem'], step=0.1, key="cc_he_so_tcd_input")
            don_gia_tc_moi = cc['don_gia_tc_thuong']
            don_gia_tcn_moi = cc['don_gia_tc_chu_nhat']
            don_gia_tcl_moi = cc['don_gia_tc_le']
            don_gia_tcd_moi = cc['don_gia_tc_dem']
        else:
            col4, col5, col6, col7 = st.columns(4)
            with col4:
                don_gia_tc_moi = st.number_input("Đơn giá TC thường (đ/giờ)", min_value=0.0,
                                                  value=cc['don_gia_tc_thuong'], step=1000.0, key="cc_don_gia_tc_input")
            with col5:
                don_gia_tcn_moi = st.number_input("Đơn giá TCN - CN (đ/giờ)", min_value=0.0,
                                                   value=cc['don_gia_tc_chu_nhat'], step=1000.0, key="cc_don_gia_tcn_input")
            with col6:
                don_gia_tcl_moi = st.number_input("Đơn giá TCL - lễ (đ/giờ)", min_value=0.0,
                                                   value=cc['don_gia_tc_le'], step=1000.0, key="cc_don_gia_tcl_input")
            with col7:
                don_gia_tcd_moi = st.number_input("Đơn giá TCĐ - cộng thêm đêm (đ/giờ)", min_value=0.0,
                                                   value=cc['don_gia_tc_dem'], step=1000.0, key="cc_don_gia_tcd_input")
            he_so_tc_moi = cc['he_so_tc_thuong']
            he_so_tcn_moi = cc['he_so_tc_chu_nhat']
            he_so_tcl_moi = cc['he_so_tc_le']
            he_so_tcd_moi = cc['he_so_tc_dem']

        st.divider()
        st.markdown("**📅 Phép năm**")
        col8, col9 = st.columns(2)
        with col8:
            cach_tinh_phep_moi = st.selectbox(
                "Cách tính phép năm", ["TU_DONG", "CO_DINH"],
                index=0 if cc['cach_tinh_phep_nam'] == "TU_DONG" else 1,
                format_func=lambda x: "Tự động (12 + 1/5 năm thâm niên)" if x == "TU_DONG" else "Cố định theo số ngày nhập",
                key="cc_cach_tinh_phep_input")
        with col9:
            so_ngay_phep_moi = st.number_input("Số ngày phép cơ bản/năm", min_value=0.0, max_value=30.0,
                                                value=cc['so_ngay_phep_co_ban'], step=0.5, key="cc_so_ngay_phep_input")

        st.divider()
        st.markdown("**🎌 Danh sách ngày nghỉ lễ trong năm**")
        st.caption("Ngày lễ cố định (1/1, 30/4, 1/5, 2/9) có thể tải tự động. "
                   "Ngày Tết âm lịch và Giỗ Tổ Hùng Vương thay đổi theo năm — "
                   "thêm thủ công sau khi Chính phủ công bố (thường tháng 10–11 năm trước).")

        nam_hien_tai = datetime.now().year
        col_le1, col_le2 = st.columns([1, 3])
        with col_le1:
            nam_tai_le = st.number_input("Năm:", min_value=2024, max_value=2030,
                                         value=nam_hien_tai, step=1, key="cc_nam_tai_le")
        with col_le2:
            st.write("")
            if st.button(f"📥 Tải ngày lễ cố định năm {int(nam_tai_le)}", key="btn_tai_le_co_dinh"):
                le_co_dinh = [
                    {"ngay": f"{int(nam_tai_le)}-01-01", "ten": "Tết Dương lịch"},
                    {"ngay": f"{int(nam_tai_le)}-04-30", "ten": "Ngày Giải phóng miền Nam"},
                    {"ngay": f"{int(nam_tai_le)}-05-01", "ten": "Quốc tế Lao động"},
                    {"ngay": f"{int(nam_tai_le)}-09-01", "ten": "Quốc khánh (nghỉ bù ngày liền kề trước)"},
                    {"ngay": f"{int(nam_tai_le)}-09-02", "ten": "Quốc khánh"},
                    {"ngay": f"{int(nam_tai_le)}-11-24", "ten": "Ngày Di sản Văn hóa Việt Nam"},
                ]
                # Gộp với danh sách hiện có, tránh trùng ngày
                ds_hien_co = {x['ngay']: x for x in (cc['danh_sach_ngay_le'] or [])}
                for le in le_co_dinh:
                    ds_hien_co[le['ngay']] = le
                ds_gop = sorted(ds_hien_co.values(), key=lambda x: x['ngay'])
                set_cau_hinh('cc_danh_sach_ngay_le',
                             json.dumps(ds_gop, ensure_ascii=False),
                             'Danh sách ngày nghỉ lễ trong năm')
                # Xoá cache để rerun đọc lại từ DB
                if 'cau_hinh_cache' in st.session_state:
                    st.session_state.pop('cau_hinh_cache', None)
                st.success(f"✅ Đã thêm {len(le_co_dinh)} ngày lễ cố định năm {int(nam_tai_le)}.")
                st.rerun()

        ds_le_text_moi = st.text_area(
            "Danh sách đầy đủ (sửa trực tiếp nếu cần):", height=150,
            value="\n".join(f"{x['ngay']} | {x['ten']}" for x in (cc['danh_sach_ngay_le'] or [])),
            key="cc_ds_le_input",
            help="Thêm ngày Tết âm lịch và Giỗ Tổ Hùng Vương theo công bố chính thức hàng năm.\n"
                 "VD: 2026-02-17 | Tết Nguyên Đán (28 tháng Chạp)")

        st.divider()
        st.markdown("**📱 Chấm công qua điện thoại nhân viên (GPS)**")
        st.caption("Nhân viên tự chấm công bằng điện thoại của mình; hệ thống kiểm tra họ có đang ở "
                   "đúng địa điểm làm việc hay không dựa trên toạ độ GPS.")

        cfg_gps = get_cau_hinh_gps()
        col_g1, col_g2 = st.columns(2)
        with col_g1:
            bat_gps_moi = st.toggle("Bật chấm công qua điện thoại (kiểm tra GPS)",
                                    value=cfg_gps['bat_gps'], key="cc_bat_gps_input")
        with col_g2:
            bat_mat_moi = st.toggle("Bắt buộc đối chiếu khuôn mặt khi chấm công",
                                    value=cfg_gps['bat_doi_chieu_mat'], key="cc_bat_mat_input",
                                    help="Bật: nhân viên phải chụp ảnh mặt, hệ thống đối chiếu với ảnh đã "
                                         "đăng ký để chống chấm công hộ. Tắt: chỉ cần đăng nhập đúng tài "
                                         "khoản và đứng đúng địa điểm.")

        st.markdown("**📍 Danh sách địa điểm làm việc**")
        st.caption("Thêm tất cả nơi nhân viên có thể chấm công: văn phòng, nhà máy, công trường, chi nhánh... "
                   "Bán kính nên đặt 100–300m tuỳ độ rộng mặt bằng (GPS điện thoại thường sai số 10–50m).")

        ds_dd_hien_tai = cfg_gps['dia_diem'] or []
        df_dd = pd.DataFrame(ds_dd_hien_tai) if ds_dd_hien_tai else pd.DataFrame(
            columns=['ten', 'lat', 'lng', 'ban_kinh'])
        for cot in ['ten', 'lat', 'lng', 'ban_kinh']:
            if cot not in df_dd.columns:
                df_dd[cot] = None
        df_dd = df_dd[['ten', 'lat', 'lng', 'ban_kinh']]

        df_dd_moi = st.data_editor(
            df_dd, num_rows="dynamic", use_container_width=True, key="cc_dia_diem_editor",
            column_config={
                "ten": st.column_config.TextColumn("Tên địa điểm", required=True,
                                                   help="VD: Văn phòng Hòn La, Nhà máy 1..."),
                "lat": st.column_config.NumberColumn("Vĩ độ (lat)", format="%.6f", required=True),
                "lng": st.column_config.NumberColumn("Kinh độ (lng)", format="%.6f", required=True),
                "ban_kinh": st.column_config.NumberColumn("Bán kính (m)", min_value=30, max_value=5000,
                                                          step=10, default=200, required=True),
            },
            disabled=not can_edit(),
        )

        with st.expander("🧭 Chưa biết toạ độ? Bấm vào đây để lấy toạ độ nơi bạn đang đứng"):
            st.caption("Cách dùng: mang điện thoại/máy tính tới đúng địa điểm cần khai báo, bấm nút dưới đây, "
                       "chọn 'Cho phép' khi trình duyệt hỏi quyền vị trí. Toạ độ sẽ hiện ra để bạn copy vào bảng trên.")
            lat_lay, lng_lay, acc_lay = nut_lay_toa_do(khoa="cauhinh")
            if lat_lay is not None:
                st.success(f"Toạ độ hiện tại: **{lat_lay:.6f}** , **{lng_lay:.6f}** "
                           f"(sai số khoảng {acc_lay:.0f}m)")
                st.caption("Copy 2 số này vào cột 'Vĩ độ (lat)' và 'Kinh độ (lng)' ở bảng phía trên, "
                           "rồi bấm Lưu.")

        if st.button("💾 Lưu cấu hình GPS & địa điểm", key="btn_luu_gps"):
            if not can_edit():
                st.error("❌ Bạn không có quyền chỉnh sửa!")
            else:
                ds_luu = []
                for _, dong in df_dd_moi.iterrows():
                    if pd.isna(dong['lat']) or pd.isna(dong['lng']) or not str(dong['ten'] or '').strip():
                        continue
                    ds_luu.append({
                        'ten': str(dong['ten']).strip(),
                        'lat': float(dong['lat']),
                        'lng': float(dong['lng']),
                        'ban_kinh': int(dong['ban_kinh'] or 200),
                    })
                ok1 = set_cau_hinh('cc_bat_gps', '1' if bat_gps_moi else '0',
                                   'Bật chấm công qua điện thoại (GPS)')
                ok2 = set_cau_hinh('cc_bat_doi_chieu_mat', '1' if bat_mat_moi else '0',
                                   'Bắt buộc đối chiếu khuôn mặt khi chấm công')
                ok3 = luu_dia_diem_lam_viec(ds_luu)
                if ok1 and ok2 and ok3:
                    st.cache_data.clear()
                    st.session_state.pop('_cau_hinh_cache', None)
                    st.success(f"✅ Đã lưu {len(ds_luu)} địa điểm làm việc.")
                    st.cache_data.clear()
                    st.rerun()
                else:
                    st.error("❌ Lưu thất bại, thử lại.")

        st.divider()
        st.markdown("**🏢 Cấu hình tăng ca theo Phòng ban**")
        st.caption("Mặc định tất cả phòng ban kế thừa cấu hình chung ở trên. "
                   "Chỉ cần cấu hình riêng cho phòng nào KHÁC với mặc định (VD: VP không tăng ca, "
                   "SX có hệ số tăng ca cao hơn). Để trống hệ số/đơn giá = kế thừa cấu hình chung.")

        try:
            db_pb = st.session_state.db_engine.get_connection()
            c_pb = db_pb.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            # Lấy danh sách phòng ban từ danh mục
            c_pb.execute("SELECT ten_phong_ban FROM danh_muc_phong_ban ORDER BY thu_tu, id")
            ds_phong_ban = [r['ten_phong_ban'] for r in c_pb.fetchall()]

            # Lấy cấu hình tăng ca hiện tại theo phòng
            c_pb.execute("SELECT * FROM cau_hinh_tang_ca_phong_ban ORDER BY ten_phong_ban")
            ds_tc_pb = {r['ten_phong_ban']: r for r in c_pb.fetchall()}
            db_pb.close()

            if not ds_phong_ban:
                st.info("Chưa có phòng ban nào trong Danh mục. Vào Danh mục → 🏢 Phòng ban để thêm.")
            else:
                cfg_chung = get_cau_hinh_cham_cong_full()

                # Xây dựng dataframe hiển thị
                rows_hien_thi = []
                for pb in ds_phong_ban:
                    cfg_pb = ds_tc_pb.get(pb, {})
                    rows_hien_thi.append({
                        'Phòng ban': pb,
                        'Cho phép TC': cfg_pb.get('cho_phep_tang_ca', True),
                        'HS TC thường': cfg_pb.get('he_so_tc_thuong'),
                        'HS TC CN': cfg_pb.get('he_so_tc_chu_nhat'),
                        'HS TC lễ': cfg_pb.get('he_so_tc_le'),
                        'HS TC đêm': cfg_pb.get('he_so_tc_dem'),
                        'Đơn giá TC (đ/h)': cfg_pb.get('don_gia_tc_thuong'),
                        'Ghi chú': cfg_pb.get('ghi_chu') or '',
                    })

                _kieu_tc = 'Hệ số %' if cfg_chung['cach_tinh_tang_ca'] == 'HE_SO' else 'Đơn giá cố định'
                _gia_tri_tc = (str(cfg_chung['he_so_tc_thuong'])
                               if cfg_chung['cach_tinh_tang_ca'] == 'HE_SO'
                               else f"{cfg_chung['don_gia_tc_thuong']:,.0f}đ/h")
                st.caption(f"💡 Cấu hình chung hiện tại: {_kieu_tc} — TC thường: {_gia_tri_tc}")

                df_tc_pb = pd.DataFrame(rows_hien_thi)
                df_tc_pb_moi = st.data_editor(
                    df_tc_pb,
                    use_container_width=True,
                    hide_index=True,
                    key="cc_tc_phong_ban_editor",
                    disabled=['Phòng ban'] + ([] if can_edit() else
                              ['Cho phép TC', 'HS TC thường', 'HS TC CN', 'HS TC lễ',
                               'HS TC đêm', 'Đơn giá TC (đ/h)', 'Ghi chú']),
                    column_config={
                        'Phòng ban': st.column_config.TextColumn(disabled=True),
                        'Cho phép TC': st.column_config.CheckboxColumn(
                            "✅ Cho phép TC",
                            help="Bỏ tick = phòng ban này không được tính tăng ca"),
                        'HS TC thường': st.column_config.NumberColumn(
                            "Hệ số TC thường", format="%.2f", min_value=1.0, max_value=5.0,
                            help="Để trống = dùng cấu hình chung"),
                        'HS TC CN': st.column_config.NumberColumn(
                            "Hệ số TC CN", format="%.2f", min_value=1.0, max_value=5.0),
                        'HS TC lễ': st.column_config.NumberColumn(
                            "Hệ số TC lễ", format="%.2f", min_value=1.0, max_value=5.0),
                        'HS TC đêm': st.column_config.NumberColumn(
                            "Hệ số TC đêm", format="%.2f", min_value=1.0, max_value=3.0),
                        'Đơn giá TC (đ/h)': st.column_config.NumberColumn(
                            "Đơn giá TC (đ/h)", format="%d",
                            help="Dùng khi chọn 'Đơn giá cố định'. Để trống = dùng cấu hình chung"),
                        'Ghi chú': st.column_config.TextColumn("Ghi chú"),
                    }
                )

                if st.button("💾 Lưu cấu hình tăng ca theo phòng ban",
                             key="btn_luu_tc_phong_ban", disabled=not can_edit()):
                    db_luu_pb = st.session_state.db_engine.get_connection()
                    c_luu = db_luu_pb.cursor()
                    loi_luu = []
                    for _, dong in df_tc_pb_moi.iterrows():
                        pb = dong['Phòng ban']
                        try:
                            c_luu.execute("""
                                INSERT INTO cau_hinh_tang_ca_phong_ban
                                    (ten_phong_ban, cho_phep_tang_ca,
                                     he_so_tc_thuong, he_so_tc_chu_nhat, he_so_tc_le, he_so_tc_dem,
                                     don_gia_tc_thuong, ghi_chu, updated_at)
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,NOW())
                                ON CONFLICT (ten_phong_ban) DO UPDATE SET
                                    cho_phep_tang_ca = EXCLUDED.cho_phep_tang_ca,
                                    he_so_tc_thuong = EXCLUDED.he_so_tc_thuong,
                                    he_so_tc_chu_nhat = EXCLUDED.he_so_tc_chu_nhat,
                                    he_so_tc_le = EXCLUDED.he_so_tc_le,
                                    he_so_tc_dem = EXCLUDED.he_so_tc_dem,
                                    don_gia_tc_thuong = EXCLUDED.don_gia_tc_thuong,
                                    ghi_chu = EXCLUDED.ghi_chu,
                                    updated_at = NOW()
                            """, (
                                pb,
                                bool(dong['Cho phép TC']),
                                dong['HS TC thường'] if pd.notna(dong['HS TC thường']) else None,
                                dong['HS TC CN'] if pd.notna(dong['HS TC CN']) else None,
                                dong['HS TC lễ'] if pd.notna(dong['HS TC lễ']) else None,
                                dong['HS TC đêm'] if pd.notna(dong['HS TC đêm']) else None,
                                dong['Đơn giá TC (đ/h)'] if pd.notna(dong['Đơn giá TC (đ/h)']) else None,
                                dong['Ghi chú'] or None,
                            ))
                        except Exception as e:
                            loi_luu.append(f"{pb}: {e}")
                    db_luu_pb.commit()
                    db_luu_pb.close()
                    if loi_luu:
                        st.error("❌ Lỗi lưu một số phòng:\n" + "\n".join(loi_luu))
                    else:
                        st.cache_data.clear()
                        st.session_state.pop('_cau_hinh_cache', None)
                        st.success("✅ Đã lưu cấu hình tăng ca theo phòng ban.")
                        st.cache_data.clear()
                        st.rerun()

        except Exception as e:
            st.error(f"❌ Lỗi tải cấu hình tăng ca phòng ban: {e}")

        st.divider()
        st.markdown("**📖 Bảng ký hiệu chấm công chuẩn (tham khảo — 23 ký hiệu, áp dụng chung mọi doanh nghiệp)**")
        with st.expander("Xem đầy đủ bảng ký hiệu"):
            st.dataframe(
                [{"Ký hiệu": ma, "Ý nghĩa": tt["ten"], "Nhóm": tt["nhom"],
                  "Cần phê duyệt": "✅" if tt.get("can_duyet") else ""}
                 for ma, tt in KY_HIEU_CHAM_CONG.items()],
                use_container_width=True, hide_index=True
            )

        if st.button("💾 Lưu cấu hình chấm công", key="btn_save_cau_hinh_cham_cong"):
            if not can_edit():
                st.error("❌ Bạn không có quyền chỉnh sửa!")
            else:
                danh_sach_le_moi = []
                for line in ds_le_text_moi.strip().split("\n"):
                    if "|" in line:
                        ngay, ten = line.split("|", 1)
                        danh_sach_le_moi.append({"ngay": ngay.strip(), "ten": ten.strip()})

                ok = update_cau_hinh_cham_cong_full({
                    'gio_vao': gio_vao_moi, 'gio_ra': gio_ra_moi, 'phut_tre': phut_tre_moi,
                    'gio_bat_dau_ca_dem': gio_bd_dem_moi, 'so_ngay_lam_viec_tuan': so_ngay_tuan_moi,
                    'ngay_nghi_hang_tuan': 'CN', 'gio_lam_chuan_ngay': gio_lam_chuan_moi,
                    'he_so_tc_thuong': he_so_tc_moi, 'he_so_tc_chu_nhat': he_so_tcn_moi,
                    'he_so_tc_le': he_so_tcl_moi, 'he_so_tc_dem': he_so_tcd_moi,
                    'cach_tinh_tang_ca': cach_tinh_tc_moi,
                    'don_gia_tc_thuong': don_gia_tc_moi, 'don_gia_tc_chu_nhat': don_gia_tcn_moi,
                    'don_gia_tc_le': don_gia_tcl_moi, 'don_gia_tc_dem': don_gia_tcd_moi,
                    'cach_tinh_phep_nam': cach_tinh_phep_moi, 'so_ngay_phep_co_ban': so_ngay_phep_moi,
                    'danh_sach_ngay_le': danh_sach_le_moi,
                    'phut_nghi_giua_ca': phut_nghi_moi,
                    'ap_dung_nghi_giua_ca': ap_dung_nghi_moi,
                })
                if ok:
                    st.cache_data.clear()
                    st.session_state.pop('_cau_hinh_cache', None)
                    st.success("✅ Đã lưu cấu hình chấm công!")
                    st.cache_data.clear()
                    st.rerun()
                else:
                    st.error("❌ Lưu thất bại, thử lại.")
    
    st.divider()
    with st.expander("🏷️ Ký hiệu Mã nhân viên riêng của công ty"):
        st.caption(
            "Ký hiệu (tiền tố) dùng khi hệ thống TỰ SINH Mã nhân viên cho người mới thêm — "
            "VD đặt 'HL' sẽ sinh HL001, HL002... Mặc định là **NV** nếu chưa đặt riêng."
        )
        ky_hieu_cur = get_cau_hinh('ky_hieu_ma_nv', 'NV') or 'NV'
        ky_hieu_moi = st.text_input("Ký hiệu Mã nhân viên:", value=ky_hieu_cur, max_chars=6, key="ky_hieu_ma_nv_cfg")
        if st.button("💾 Lưu ký hiệu Mã nhân viên", key="btn_save_ky_hieu_ma_nv", disabled=not can_edit()):
            ky_hieu_sach = re.sub(r'[^A-Za-z0-9]', '', ky_hieu_moi).upper()
            if not ky_hieu_sach:
                st.error("❌ Ký hiệu không hợp lệ (chỉ gồm chữ/số).")
            elif set_cau_hinh('ky_hieu_ma_nv', ky_hieu_sach, 'Tiền tố dùng khi tự sinh Mã nhân viên mới'):
                st.success(f"✅ Đã lưu ký hiệu Mã nhân viên: {ky_hieu_sach}")
                st.cache_data.clear()
                st.rerun()
            else:
                st.error("❌ Lưu thất bại.")

    with st.expander("🧪 Tài khoản đăng nhập DEMO (tự điền sẵn ở màn hình đăng nhập)"):
        st.caption(
            "Chỉ áp dụng cho tenant có Mã công ty = **DEMO-HRM**: điền sẵn Tên đăng nhập/Mật khẩu "
            "ở màn hình đăng nhập để khách trải nghiệm không cần hỏi tài khoản. "
            "⚠️ Giá trị nhập ở đây PHẢI khớp đúng với Tên đăng nhập/Mật khẩu thật của 1 nhân viên "
            "có sẵn trong hệ thống (mục ✅ Quản lý nhân viên), nếu không đăng nhập sẽ vẫn báo sai."
        )
        demo_user_cur = get_cau_hinh('demo_ten_dang_nhap', '') or ''
        demo_pass_cur = get_cau_hinh('demo_mat_khau', '') or ''
        col_demo1, col_demo2 = st.columns(2)
        with col_demo1:
            demo_user_moi = st.text_input("Tên đăng nhập demo:", value=demo_user_cur, key="demo_user_cfg")
        with col_demo2:
            demo_pass_moi = st.text_input("Mật khẩu demo:", value=demo_pass_cur, key="demo_pass_cfg")
        if st.button("💾 Lưu tài khoản demo", key="btn_save_demo_login", disabled=not can_edit()):
            ok1 = set_cau_hinh('demo_ten_dang_nhap', demo_user_moi.strip(), 'Tên đăng nhập tự điền sẵn ở trang đăng nhập cho tenant DEMO-HRM')
            ok2 = set_cau_hinh('demo_mat_khau', demo_pass_moi.strip(), 'Mật khẩu tự điền sẵn ở trang đăng nhập cho tenant DEMO-HRM')
            if ok1 and ok2:
                st.success("✅ Đã lưu tài khoản demo!")
                st.cache_data.clear()
                st.rerun()
            else:
                st.error("❌ Lưu thất bại.")

# ========== BHXH ==========
elif menu == "📋 BHXH":
    st.markdown(f"# {i18n.tm('📋 Quản lý BHXH')}", unsafe_allow_html=True)
    
    t1, t2, t3 = st.tabs(["📊 Tổng quan", "📝 Báo cáo tăng/giảm D02-LT", "💰 Dự toán đóng BHXH"])
    
    with t1:
        st.subheader("📊 Tổng quan tình hình đóng BHXH")
        
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Thống kê chung
        c.execute("SELECT COUNT(*) as tong FROM nhan_vien WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC')")
        tong_ld = c.fetchone()['tong']
        
        c.execute("""
            SELECT COUNT(*) as dang_dong 
            FROM nhan_vien 
            WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC') 
            AND thang_bat_dau_bh IS NOT NULL  -- ĐÃ CÓ NGÀY BẮT ĐẦU = ĐANG THAM GIA
        """)
        dang_dong = c.fetchone()['dang_dong']

        c.execute("""
            SELECT COUNT(*) as chua_dong 
            FROM nhan_vien 
            WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC') 
            AND thang_bat_dau_bh IS NULL  -- CHƯA CÓ NGÀY BẮT ĐẦU = CHƯA THAM GIA
        """)
        chua_dong = c.fetchone()['chua_dong']
        
        c.execute("SELECT COUNT(*) as da_nghi FROM nhan_vien WHERE trang_thai = 'NGHI_VIEC'")
        da_nghi = c.fetchone()['da_nghi']
        
        db.close()
        
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("👥 Tổng lao động", tong_ld)
        col2.metric("✅ Đang đóng BHXH", dang_dong, delta=f"{dang_dong/tong_ld*100:.0f}%" if tong_ld > 0 else None)
        col3.metric("⏳ Chưa đóng BHXH", chua_dong, delta=f"-{chua_dong}" if chua_dong > 0 else None, delta_color="inverse")
        col4.metric("📋 Đã nghỉ việc", da_nghi)
        
        st.divider()
        
        # Danh sách lao động chưa đóng BHXH
        st.subheader("⚠️ Lao động chưa đóng BHXH")
        
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("""
            SELECT ma_nv, ho_ten, chuc_danh_nghe, ngay_vao_lam, loai_hop_dong, 
                   thang_bat_dau_bh, trang_thai_bhxh
            FROM nhan_vien 
            WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC') 
            AND thang_bat_dau_bh IS NULL
            ORDER BY ngay_vao_lam ASC
        """)
        chua_dong_list = c.fetchall()

        db.close()
        
        if chua_dong_list:
            df_chua_dong = pd.DataFrame(chua_dong_list)
            for col in df_chua_dong.columns:
                if 'ngay' in col.lower():
                    df_chua_dong[col] = df_chua_dong[col].apply(format_date)
            st.dataframe(df_chua_dong, width='stretch', hide_index=True)
            
            if st.session_state.role in ("admin", "xem_toan_bo"):
                st.warning("💡 Hướng dẫn: Vào menu '✅ Nhân viên' -> chọn nhân viên -> sửa thông tin -> cập nhật 'Bắt đầu BH' và chuyển trạng thái BHXH thành 'ĐANG ĐÓNG'")
        else:
            st.success("✅ Tất cả lao động đã được đăng ký đóng BHXH!")
            st.cache_data.clear()
    
    with t2:
        st.subheader("📝 Báo cáo tăng/giảm lao động tham gia BHXH (Mẫu D02-LT)")
        st.caption("Theo Quyết định 595/QĐ-BHXH và mẫu D02-LT - Dùng để kê khai tăng/giảm lao động tham gia BHXH, BHYT, BHTN")
        
        col_from, col_to = st.columns(2)
        with col_from:
            tu_ngay = st.date_input("📅 Từ ngày (theo tháng bắt đầu/kết thúc BHXH):", 
                                    value=date(date.today().year, 1, 1), 
                                    key="d02_tu")
        with col_to:
            den_ngay = st.date_input("📅 Đến ngày:", 
                                    value=date.today(), 
                                    key="d02_den")
        
        # Nút xuất báo cáo - ĐẶT NGAY PHÍA DƯỚI BỘ LỌC NGÀY
        col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
        with col_btn2:
            export_clicked = st.button("📥 XUẤT EXCEL D02-LT (Mẫu BHXH)", 
                                       type="primary", 
                                       width='stretch',
                                       use_container_width=True)
        
        st.divider()
        
        # Khởi tạo biến để lưu kết quả truy vấn
        tang_list = []
        giam_list = []
        
        # Chỉ truy vấn khi cần (khi người dùng click nút hoặc muốn xem trước)
        # Nhưng để hiển thị preview, chúng ta vẫn chạy truy vấn
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Lao động tăng trong kỳ (dựa vào thang_bat_dau_bh)
        c.execute("""
            SELECT 
                nv.id, nv.ma_nv, nv.ho_ten, nv.ma_so_bhxh, nv.ngay_sinh, nv.gioi_tinh, nv.so_cccd,
                nv.chuc_danh_nghe, nv.phong_ban_lam_viec, nv.luong_bao_hiem, nv.he_so_luong,
                nv.thang_bat_dau_bh as ngay_bat_dau,
                nv.loai_hop_dong, nv.so_hdld, nv.ngay_vao_lam, nv.thuong_tru,
                nv.phu_cap_chuc_vu, nv.phu_cap_tnvk, nv.phu_cap_tnn,
                nv.muc_huong_bhyt, nv.ty_le_dong, nv.muc_tien_dong, nv.phuong_thuc_dong,
                nv.quoc_tich, nv.dan_toc, nv.dien_thoai, nv.email_lien_he,
                nv.tinh_nhan_hs, nv.phuong_nhan_hs, nv.dia_chi_nhan_hs,
                nv.tinh_kcb, nv.noi_dang_ky_kcb, nv.dang_ky_nhan_so,
                nv.ngay_ky_hd, nv.ngay_ket_thuc, nv.ten_don_vi_thu_huong
            FROM nhan_vien nv
            WHERE nv.trang_thai IN ('DANG_LAM', 'THU_VIEC')
            AND nv.thang_bat_dau_bh IS NOT NULL
            AND nv.thang_bat_dau_bh BETWEEN %s AND %s
            ORDER BY nv.thang_bat_dau_bh ASC
        """, (tu_ngay, den_ngay))
        tang_list = c.fetchall()
        
        # Lao động giảm trong kỳ — ưu tiên thang_ket_thuc_bh (đúng nghiệp vụ báo giảm BHXH),
        # nếu trống (hồ sơ cũ/nhập Excel/chưa qua luồng Quyết định nhân sự Chấm dứt HĐ) thì
        # lấy ngay_ket_thuc (ngày nghỉ việc thực tế) làm dự phòng — tránh sót lao động giảm
        # trong báo cáo D02-LT chỉ vì thiếu 1 cột BHXH riêng chưa kịp đồng bộ.
        c.execute("""
            SELECT 
                nv.id, nv.ma_nv, nv.ho_ten, nv.ma_so_bhxh, nv.ngay_sinh, nv.gioi_tinh, nv.so_cccd,
                nv.chuc_danh_nghe, nv.phong_ban_lam_viec, nv.luong_bao_hiem, nv.he_so_luong,
                COALESCE(nv.thang_ket_thuc_bh, nv.ngay_ket_thuc) as ngay_ket_thuc,
                nv.loai_hop_dong, nv.so_hdld, nv.ngay_vao_lam, nv.thuong_tru,
                nv.ly_do_nghi
            FROM nhan_vien nv
            WHERE nv.trang_thai = 'NGHI_VIEC'
            AND COALESCE(nv.thang_ket_thuc_bh, nv.ngay_ket_thuc) BETWEEN %s AND %s
            ORDER BY COALESCE(nv.thang_ket_thuc_bh, nv.ngay_ket_thuc) ASC
        """, (tu_ngay, den_ngay))
        giam_list = c.fetchall()
        db.close()
        
        # Hiển thị preview
        col_tang, col_giam = st.columns(2)
        with col_tang:
            st.markdown(f"### 🟢 LAO ĐỘNG TĂNG ({len(tang_list)})")
            if tang_list:
                df_tang = pd.DataFrame(tang_list)
                for col in df_tang.columns:
                    if 'ngay' in col.lower():
                        df_tang[col] = df_tang[col].apply(format_date)
                preview_cols = ['ma_nv', 'ho_ten', 'ma_so_bhxh', 'ngay_bat_dau']
                available_cols = [c for c in preview_cols if c in df_tang.columns]
                df_preview = df_tang[available_cols]
                df_preview.columns = ['Mã NV', 'Họ tên', 'Mã BHXH', 'Ngày bắt đầu']
                st.dataframe(df_preview, width='stretch', hide_index=True, height=300)
            else:
                st.info("📭 Không có lao động tăng trong kỳ")
        
        with col_giam:
            st.markdown(f"### 🔴 LAO ĐỘNG GIẢM ({len(giam_list)})")
            if giam_list:
                df_giam = pd.DataFrame(giam_list)
                for col in df_giam.columns:
                    if 'ngay' in col.lower():
                        df_giam[col] = df_giam[col].apply(format_date)
                preview_cols = ['ma_nv', 'ho_ten', 'ma_so_bhxh', 'ngay_ket_thuc']
                available_cols = [c for c in preview_cols if c in df_giam.columns]
                df_preview = df_giam[available_cols]
                df_preview.columns = ['Mã NV', 'Họ tên', 'Mã BHXH', 'Ngày kết thúc']
                st.dataframe(df_preview, width='stretch', hide_index=True, height=300)
            else:
                st.info("📭 Không có lao động giảm trong kỳ")
        
        st.divider()
        
        # ===== XỬ LÝ XUẤT EXCEL KHI NHẤN NÚT =====
        if export_clicked:
            if tang_list or giam_list:
                with st.spinner("Đang tạo báo cáo D02-LT theo mẫu BHXH... Vui lòng chờ..."):
                    try:
                        # Gọi hàm tạo báo cáo
                        filename = tao_bao_cao_bhxh_d02_lt(
                            tang_list, 
                            giam_list, 
                            tu_ngay, 
                            den_ngay, 
                            COMPANY_CONFIG.get("ten_cong_ty", "CÔNG TY CỔ PHẦN CẢNG HÒN LA"),
                            COMPANY_CONFIG.get("ma_don_vi_BHXH", "4400000000")
                        )
                        
                        # Đọc file và tải xuống
                        with open(filename, "rb") as f:
                            file_data = f.read()
                        
                        st.success(f"✅ Đã tạo báo cáo thành công! {len(tang_list)} lao động tăng, {len(giam_list)} lao động giảm.")
                        st.cache_data.clear()
                        
                        st.download_button(
                            label="📥 TẢI FILE EXCEL D02-LT (Đúng mẫu BHXH)",
                            data=file_data,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            width='stretch',
                            key="download_d02_lt"
                        )
                        
                        # Xóa file tạm sau khi đã đọc
                        import os
                        if os.path.exists(filename):
                            os.remove(filename)
                            
                    except Exception as e:
                        st.error(f"❌ Lỗi khi tạo báo cáo: {str(e)}")
                        st.exception(e)
            else:
                st.warning("⚠️ Không có biến động lao động (tăng hoặc giảm) trong kỳ để xuất báo cáo!")
    
    with t3:
        st.subheader("💰 DỰ TOÁN ĐÓNG BHXH")
        st.caption("Tính toán các khoản phải nộp Bảo hiểm xã hội, Bảo hiểm y tế, Bảo hiểm thất nghiệp")
        
        st.info("""
        ### 🚧 Tính năng đang hoàn thiện
        
        Nội dung đang được phát triển. Các tính năng sắp ra mắt:
        - ✅ Tính toán mức đóng BHXH theo lương cơ sở
        - ✅ Tính toán các khoản phụ cấp tính đóng BHXH
        - ✅ Bảng kê chi tiết từng nhân viên
        - ✅ Xuất báo cáo kê khai BHXH theo mẫu quy định
        - ✅ Tổng hợp số tiền phải nộp theo tháng/quý/năm
        
        ⏳ **Dự kiến hoàn thành: Quý 3/2026**
        """)
        
        # Thêm một số thông tin tham khảo
        with st.expander("📌 Thông tin tham khảo về tỷ lệ đóng BHXH hiện hành"):
            st.markdown("""
            **Tỷ lệ trích BHXH, BHYT, BHTN theo quy định hiện hành:**
            
            | Loại | Doanh nghiệp | Người lao động | Tổng |
            |------|-------------|----------------|------|
            | BHXH | 17.5% | 8% | 25.5% |
            | BHYT | 3% | 1.5% | 4.5% |
            | BHTN | 1% | 1% | 2% |
            | BHTNLĐ-BNN | 0.5% | 0% | 0.5% |
            | **Tổng cộng** | **22%** | **10.5%** | **32.5%** |
            
            *Lưu ý: Tỷ lệ có thể thay đổi theo quy định mới của Nhà nước.*
            """)
        
        # Thêm form nhập thử nghiệm (có thể comment lại sau)
        with st.expander("🧪 Thử nghiệm tính toán (Demo)"):
            col_demo1, col_demo2 = st.columns(2)
            with col_demo1:
                luong_demo = st.number_input("Lương tháng (VNĐ)", min_value=0, value=5000000, step=500000)
            with col_demo2:
                chon_ty_le = st.selectbox("Áp dụng tỷ lệ", ["Theo quy định", "Tùy chỉnh"])
            
            if chon_ty_le == "Theo quy định":
                ty_le_nld = 10.5
                ty_le_nsdl = 22.0
            else:
                col_ty1, col_ty2 = st.columns(2)
                with col_ty1:
                    ty_le_nld = st.number_input("Tỷ lệ NLĐ (%)", min_value=0.0, max_value=50.0, value=10.5, step=0.5)
                with col_ty2:
                    ty_le_nsdl = st.number_input("Tỷ lệ NSDLĐ (%)", min_value=0.0, max_value=50.0, value=22.0, step=0.5)
            
            tien_nld = luong_demo * ty_le_nld / 100
            tien_nsdl = luong_demo * ty_le_nsdl / 100
            tong_tien = tien_nld + tien_nsdl
            
            st.markdown("---")
            col_kq1, col_kq2, col_kq3 = st.columns(3)
            col_kq1.metric("NLĐ đóng", f"{tien_nld:,.0f} VNĐ", f"({ty_le_nld}%)")
            col_kq2.metric("NSDLĐ đóng", f"{tien_nsdl:,.0f} VNĐ", f"({ty_le_nsdl}%)")
            col_kq3.metric("Tổng tiền", f"{tong_tien:,.0f} VNĐ", "cả 2 bên")
        
# ========== BÁO CÁO TÌNH HÌNH SỬ DỤNG LAO ĐỘNG MẪU 01/PLI (EXCEL) ==========
elif menu == "📋 Báo cáo định kỳ":
    st.markdown(f"# {i18n.tm('📋 Báo cáo định kỳ')}", unsafe_allow_html=True)

    tab_bc_pli, tab_bc_tk, tab_bc_tanggiam, tab_bc_tinhhinh, tab_bc_yte, tab_bc_atvsld, \
    tab_bc_tnld, tab_bc_huanluyen, tab_bc_socapcuu, tab_bc_quantrac = st.tabs([
        "📋 Báo cáo 01/PLI", "📊 Báo cáo thống kê nhân sự", "📊 Báo cáo tăng/giảm nhân sự trong kỳ",
        "📈 Tình hình sử dụng lao động", "🏥 Y tế Lao động", "🦺 Công tác ATVSLĐ",
        "⚠️ Tai nạn lao động", "🎓 Huấn luyện ATVSLĐ", "🚑 Mạng lưới sơ cấp cứu", "🌡️ Quan trắc môi trường LĐ"
    ])

    def _bao_cao_dang_phat_trien(ten_bao_cao, mo_ta, icon, mau_sac="#f59e0b"):
        """Card hiện đại cho các báo cáo định kỳ chưa có logic - sẽ bổ sung sau."""
        st.markdown(f"""
        <div style="border:1px solid #e5e7eb;border-radius:16px;padding:28px 24px;
             background:linear-gradient(135deg,#fffbeb 0%,#ffffff 100%);
             box-shadow:0 2px 10px rgba(0,0,0,0.04);">
            <div style="display:flex;align-items:center;gap:14px;margin-bottom:10px;">
                <div style="font-size:34px;">{icon}</div>
                <div>
                    <div style="font-size:19px;font-weight:700;color:#111827;">{ten_bao_cao}</div>
                    <span style="display:inline-block;margin-top:4px;padding:2px 10px;border-radius:999px;
                          background:{mau_sac};color:white;font-size:11px;font-weight:600;letter-spacing:.3px;">
                          🚧 SẼ BỔ SUNG SAU
                    </span>
                </div>
            </div>
            <div style="color:#4b5563;font-size:14px;line-height:1.6;margin-top:8px;">{mo_ta}</div>
        </div>
        """, unsafe_allow_html=True)
        st.write("")
        with st.expander("ℹ️ Khi triển khai, báo cáo này sẽ cần"):
            st.markdown("""
            - Biểu mẫu/Thông tư pháp lý làm căn cứ (bạn cung cấp khi sẵn sàng)
            - Kỳ báo cáo (tháng/quý/năm) và bộ lọc theo phòng ban
            - Xuất file Word/Excel theo đúng mẫu quy định
            """)

    with tab_bc_tinhhinh:
        st.caption("💡 Lưu ý: tab **📋 Báo cáo 01/PLI** hiện tại cũng đang thể hiện nội dung "
                   "\"tình hình sử dụng lao động\" theo mẫu 01/PLI. Nếu đây là 1 báo cáo khác "
                   "(mẫu/kỳ báo cáo khác), bạn gửi mẫu cụ thể để tôi phân biệt rõ khi triển khai.")
        _bao_cao_dang_phat_trien(
            "Báo cáo tình hình sử dụng lao động",
            "Tổng hợp định kỳ tình hình sử dụng lao động của doanh nghiệp (số lượng, cơ cấu, biến động) "
            "theo quy định báo cáo lao động định kỳ.",
            "📈"
        )

    with tab_bc_yte:
        _bao_cao_dang_phat_trien(
            "Báo cáo Y tế Lao động",
            "Báo cáo công tác y tế lao động: khám sức khỏe định kỳ, bệnh nghề nghiệp, tình hình sức khỏe người lao động.",
            "🏥"
        )

    with tab_bc_atvsld:
        _bao_cao_dang_phat_trien(
            "Báo cáo công tác An toàn, Vệ sinh lao động (ATVSLĐ)",
            "Tổng hợp công tác an toàn vệ sinh lao động: tổ chức bộ máy ATVSLĐ, tự kiểm tra, cải thiện điều kiện làm việc.",
            "🦺"
        )

    with tab_bc_tnld:
        _bao_cao_dang_phat_trien(
            "Báo cáo Tai nạn lao động",
            "Thống kê, khai báo các vụ tai nạn lao động phát sinh trong kỳ báo cáo theo quy định.",
            "⚠️", mau_sac="#ef4444"
        )

    with tab_bc_huanluyen:
        _bao_cao_dang_phat_trien(
            "Báo cáo Huấn luyện ATVSLĐ",
            "Tổng hợp tình hình huấn luyện an toàn vệ sinh lao động theo nhóm đối tượng, thời hạn huấn luyện lại.",
            "🎓"
        )

    with tab_bc_socapcuu:
        _bao_cao_dang_phat_trien(
            "Báo cáo hoạt động mạng lưới sơ cấp cứu",
            "Tình hình tổ chức, hoạt động của mạng lưới sơ cấp cứu tại doanh nghiệp.",
            "🚑"
        )

    with tab_bc_quantrac:
        _bao_cao_dang_phat_trien(
            "Báo cáo Quan trắc môi trường lao động",
            "Kết quả quan trắc môi trường lao động định kỳ (các yếu tố có hại tại nơi làm việc).",
            "🌡️"
        )

    with tab_bc_pli:
        st.subheader("📋 Báo cáo tình hình sử dụng lao động")
        st.caption("Theo mẫu 01/PLI Phụ lục I - Nghị định 145/2020/NĐ-CP (sửa đổi bởi Nghị định 35/2022/NĐ-CP)")
    
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    
        col1, col2 = st.columns(2)
        with col1:
            tu_ngay = st.date_input("📅 Từ ngày:", value=date(date.today().year, 1, 1), key="pli_tu")
        with col2:
            den_ngay = st.date_input("📅 Đến ngày:", value=date.today(), key="pli_den")
    
        db = st.session_state.db_engine.get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("""
            SELECT 
                nv.STT, nv.ma_nv, nv.ho_ten, nv.ma_so_bhxh, nv.ngay_sinh, nv.gioi_tinh,
                nv.so_cccd, nv.chuc_danh_nghe, nv.luong_bao_hiem, nv.he_so_luong,
                nv.phu_cap_chuc_vu, nv.phu_cap_tnvk, nv.phu_cap_tnn, nv.loai_hop_dong,
                nv.ngay_vao_lam, nv.ngay_ky_hd, nv.ngay_ket_thuc, nv.thang_bat_dau_bh,
                nv.thang_ket_thuc_bh, nv.so_hdld, nv.phong_ban_lam_viec, nv.noi_lam_viec,
                nv.ten_don_vi_thu_huong
            FROM nhan_vien nv
            WHERE nv.trang_thai IN ('DANG_LAM', 'THU_VIEC', 'NGHI_VIEC')
            AND nv.ngay_vao_lam <= %s
            AND (nv.ngay_ket_thuc IS NULL OR nv.ngay_ket_thuc >= %s)
            ORDER BY nv.STT ASC
        """, (den_ngay, tu_ngay))
        ds_lao_dong = c.fetchall()
        db.close()
    
        st.info(f"📊 Tổng số lao động đang làm việc: **{len(ds_lao_dong)}** người")
    
        # Hiển thị bảng dữ liệu trước khi xuất (cho cả admin và viewer)
        if ds_lao_dong:
            st.subheader("📋 Danh sách lao động")
            df_preview = pd.DataFrame(ds_lao_dong)
            for col in df_preview.columns:
                if 'ngay' in col.lower():
                    df_preview[col] = df_preview[col].apply(format_date)
        
            preview_cols = ['ma_nv', 'ho_ten', 'chuc_danh_nghe', 'loai_hop_dong', 'ngay_vao_lam', 'ma_so_bhxh', 'ten_don_vi_thu_huong']
            available_preview = [c for c in preview_cols if c in df_preview.columns]
            df_display = df_preview[available_preview]
            col_map_preview = {
                'ma_nv': 'Mã NV',
                'ho_ten': 'Họ tên',
                'chuc_danh_nghe': 'Chức danh',
                'loai_hop_dong': 'Loại HĐ',
                'ngay_vao_lam': 'Ngày vào làm',
                'ma_so_bhxh': 'Mã BHXH',
                'ten_don_vi_thu_huong': 'Tên đơn vị thụ hưởng',
            }
            df_display.rename(columns=col_map_preview, inplace=True)
            st.dataframe(df_display, width='stretch', hide_index=True, height=400)
        
            st.divider()
        
            # Chỉ admin mới được xuất Excel
            if st.session_state.role in ("admin", "xem_toan_bo"):
                if st.button("📥 XUẤT EXCEL MẪU 01/PLI", type="primary", width='stretch'):
                    if not can_export():
                        st.error("❌ Bạn không có quyền xuất báo cáo!")
                    else:
                        wb = Workbook()
                        ws = wb.active
                        ws.title = "BC_Tinh_hinh_su_dung_LD"
                    
                        ten_cong_ty = COMPANY_CONFIG.get("ten_cong_ty", "CÔNG TY CỔ PHẦN CẢNG HÒN LA")
                        dia_chi = COMPANY_CONFIG.get("dia_chi", "")
                        ma_so_thue = COMPANY_CONFIG.get("ma_so_thue", "")
                        dien_thoai_cty = COMPANY_CONFIG.get("dien_thoai_cty", "")
                    
                        # Header
                        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
                        ws['A1'] = ten_cong_ty
                        ws['A1'].font = Font(bold=True, size=13, name='Times New Roman')
                        ws['A1'].alignment = Alignment(horizontal='center')
                    
                        ws.merge_cells(start_row=1, start_column=20, end_row=1, end_column=26)
                        ws['T1'] = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
                        ws['T1'].font = Font(bold=True, size=13, name='Times New Roman')
                        ws['T1'].alignment = Alignment(horizontal='center')
                    
                        ma_cty_bc = (st.session_state.tenant.get('ma_cty', 'CHL') if st.session_state.get('tenant') else 'CHL')
                        dia_diem_bc = COMPANY_CONFIG.get('dia_diem') or get_cau_hinh('dia_diem', 'Quảng Trị')
                        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=6)
                        ws['A2'] = f"Số: 01/BC-01PLI-{datetime.now().year}/{ma_cty_bc}"
                        ws['A2'].font = Font(size=12, name='Times New Roman')
                        ws['A2'].alignment = Alignment(horizontal='center')
                    
                        ws.merge_cells(start_row=2, start_column=20, end_row=2, end_column=26)
                        ws['T2'] = "Độc lập - Tự do - Hạnh phúc"
                        ws['T2'].font = Font(italic=True, size=12, name='Times New Roman')
                        ws['T2'].alignment = Alignment(horizontal='center')
                    
                        ws.merge_cells(start_row=3, start_column=20, end_row=3, end_column=26)
                        ws['T3'] = f"{dia_diem_bc}, ngày {date.today().day} tháng {date.today().month} năm {date.today().year}"
                        ws['T3'].font = Font(italic=True, size=12, name='Times New Roman')
                        ws['T3'].alignment = Alignment(horizontal='right')
                    
                        ws.merge_cells('A5:AA5')
                        ws['A5'] = "BÁO CÁO TÌNH HÌNH SỬ DỤNG LAO ĐỘNG"
                        ws['A5'].font = Font(bold=True, size=13, name='Times New Roman')
                        ws['A5'].alignment = Alignment(horizontal='center')
                    
                        ws.merge_cells('A6:AA6')
                        ws['A6'] = f"(Từ ngày {tu_ngay.strftime('%d/%m/%Y')} đến ngày {den_ngay.strftime('%d/%m/%Y')})"
                        ws['A6'].font = Font(size=12, name='Times New Roman')
                        ws['A6'].alignment = Alignment(horizontal='center')
                    
                        ws.merge_cells('A8:AA8')
                        ws['A8'] = "Kính gửi: SỞ NỘI VỤ TỈNH QUẢNG TRỊ"
                        ws['A8'].font = Font(bold=True, size=11, name='Times New Roman')
                        ws['A8'].alignment = Alignment(horizontal='left')
                    
                        ws['A10'] = "1. Thông tin chung về doanh nghiệp:"
                        ws['A10'].font = Font(bold=True, size=11, name='Times New Roman')
                    
                        row_info = 11
                        for label in [f"- Tên doanh nghiệp: {ten_cong_ty}", f"- Địa chỉ: {dia_chi}", 
                                     f"- Mã số thuế: {ma_so_thue}", f"- Điện thoại: {dien_thoai_cty}"]:
                            ws[f'A{row_info}'] = label
                            ws[f'A{row_info}'].font = Font(size=11, name='Times New Roman')
                            row_info += 1
                    
                        ws[f'A{row_info + 1}'] = "2. Thông tin tình hình sử dụng lao động của đơn vị:"
                        ws[f'A{row_info + 1}'].font = Font(bold=True, size=11, name='Times New Roman')
                    
                        header_row = 18
                        col_widths = [5, 25, 18, 15, 8, 18, 25, 12, 18, 18, 12, 15, 
                                     12, 12, 12, 12, 15, 12, 12, 18, 18, 18, 18, 18, 18, 18, 20]
                        for i, w in enumerate(col_widths, 1):
                            ws.column_dimensions[get_column_letter(i)].width = w
                    
                        stt_row = header_row + 3
                        for col in range(1, 28):
                            ws.cell(row=stt_row, column=col, value=f"({col})")
                            ws.cell(row=stt_row, column=col).font = Font(size=9, name='Times New Roman')
                            ws.cell(row=stt_row, column=col).alignment = Alignment(horizontal='center')
                            ws.cell(row=stt_row, column=col).border = thin_border
                    
                        # Merge cells header
                        ws.merge_cells(start_row=header_row, start_column=1, end_row=header_row+2, end_column=1)
                        ws.cell(row=header_row, column=1, value="STT")
                    
                        ws.merge_cells(start_row=header_row, start_column=2, end_row=header_row+2, end_column=2)
                        ws.cell(row=header_row, column=2, value="Họ và tên")
                    
                        ws.merge_cells(start_row=header_row, start_column=3, end_row=header_row+2, end_column=3)
                        ws.cell(row=header_row, column=3, value="Mã số BHXH")
                    
                        ws.merge_cells(start_row=header_row, start_column=4, end_row=header_row+2, end_column=4)
                        ws.cell(row=header_row, column=4, value="Ngày sinh")
                    
                        ws.merge_cells(start_row=header_row, start_column=5, end_row=header_row+2, end_column=5)
                        ws.cell(row=header_row, column=5, value="Giới tính")
                    
                        ws.merge_cells(start_row=header_row, start_column=6, end_row=header_row+2, end_column=6)
                        ws.cell(row=header_row, column=6, value="Số CCCD/Hộ chiếu")
                    
                        ws.merge_cells(start_row=header_row, start_column=7, end_row=header_row+2, end_column=7)
                        ws.cell(row=header_row, column=7, value="Chức danh nghề, vị trí, công việc")
                    
                        ws.merge_cells(start_row=header_row, start_column=8, end_row=header_row, end_column=11)
                        ws.cell(row=header_row, column=8, value="Vị trí việc làm (2)")
                    
                        ws.merge_cells(start_row=header_row, start_column=12, end_row=header_row, end_column=17)
                        ws.cell(row=header_row, column=12, value="Tiền lương")
                    
                        ws.merge_cells(start_row=header_row, start_column=20, end_row=header_row, end_column=24)
                        ws.cell(row=header_row, column=20, value="Loại và hiệu lực hợp đồng")
                    
                        ws.merge_cells(start_row=header_row, start_column=18, end_row=header_row+1, end_column=19)
                        ws.cell(row=header_row, column=18, value="Ngành nghề nặng nhọc, độc hại")
                    
                        ws.merge_cells(start_row=header_row+1, start_column=13, end_row=header_row+1, end_column=17)
                        ws.cell(row=header_row+1, column=13, value="Phụ cấp")
                    
                        ws.merge_cells(start_row=header_row+1, start_column=21, end_row=header_row+1, end_column=22)
                        ws.cell(row=header_row+1, column=21, value="Hiệu lực HĐLĐ xác định thời hạn")
                    
                        ws.merge_cells(start_row=header_row+1, start_column=23, end_row=header_row+1, end_column=24)
                        cell = ws.cell(row=header_row+1, column=23, value="Hiệu lực HĐLĐ khác (dưới 1 tháng, thử việc)")
                        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    
                        ws.merge_cells(start_row=header_row+1, start_column=8, end_row=header_row+2, end_column=8)
                        ws.cell(row=header_row+1, column=8, value="Nhà quản lý")
                    
                        ws.merge_cells(start_row=header_row+1, start_column=9, end_row=header_row+2, end_column=9)
                        ws.cell(row=header_row+1, column=9, value="Chuyên môn kỹ thuật bậc cao")
                    
                        ws.merge_cells(start_row=header_row+1, start_column=10, end_row=header_row+2, end_column=10)
                        ws.cell(row=header_row+1, column=10, value="Chuyên môn kỹ thuật bậc trung")
                    
                        ws.merge_cells(start_row=header_row+1, start_column=11, end_row=header_row+2, end_column=11)
                        ws.cell(row=header_row+1, column=11, value="Khác")
                    
                        ws.merge_cells(start_row=header_row+1, start_column=12, end_row=header_row+2, end_column=12)
                        ws.cell(row=header_row+1, column=12, value="Mức lương/Hệ số lương")
                    
                        ws.merge_cells(start_row=header_row+1, start_column=20, end_row=header_row+2, end_column=20)
                        ws.cell(row=header_row+1, column=20, value="Ngày bắt đầu HĐLĐ không xác định thời hạn")
                    
                        ws.merge_cells(start_row=header_row, start_column=25, end_row=header_row+2, end_column=25)
                        ws.cell(row=header_row, column=25, value="Thời điểm bắt đầu đóng BHXH")
                    
                        ws.merge_cells(start_row=header_row, start_column=26, end_row=header_row+2, end_column=26)
                        ws.cell(row=header_row, column=26, value="Thời điểm kết thúc đóng BHXH")
                    
                        ws.merge_cells(start_row=header_row, start_column=27, end_row=header_row+2, end_column=27)
                        ws.cell(row=header_row, column=27, value="Ghi chú")
                    
                        for row in range(header_row, header_row + 3):
                            for col in range(1, 28):
                                cell = ws.cell(row=row, column=col)
                                cell.border = thin_border
                    
                        ws.cell(row=header_row+2, column=13, value="Phụ cấp chức vụ")
                        ws.cell(row=header_row+2, column=14, value="Phụ cấp thâm niên VK(%)")
                        ws.cell(row=header_row+2, column=15, value="Phụ cấp thâm niên nghề (%)")
                        ws.cell(row=header_row+2, column=16, value="Phụ cấp thâm niên nghề (%)")
                        ws.cell(row=header_row+2, column=17, value="Các khoản bổ sung")
                        ws.cell(row=header_row+2, column=18, value="Ngày bắt đầu")
                        ws.cell(row=header_row+2, column=19, value="Ngày kết thúc")
                        ws.cell(row=header_row+2, column=21, value="Ngày bắt đầu")
                        ws.cell(row=header_row+2, column=22, value="Ngày kết thúc")
                        ws.cell(row=header_row+2, column=23, value="Ngày bắt đầu")
                        ws.cell(row=header_row+2, column=24, value="Ngày kết thúc")
                    
                        for row in range(header_row, header_row + 3):
                            for col in range(1, 28):
                                cell = ws.cell(row=row, column=col)
                                if cell.value:
                                    cell.font = Font(bold=True, size=10, name='Times New Roman')
                                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                                    cell.border = thin_border
                    
                        for col in range(1, 28):
                            cell = ws.cell(row=stt_row, column=col)
                            cell.font = Font(size=9, name='Times New Roman')
                            cell.alignment = Alignment(horizontal='center')
                            cell.border = thin_border
                    
                        data_row = stt_row + 1
                        for idx, nv in enumerate(ds_lao_dong, 1):
                            row = data_row + idx - 1
                        
                            ws.cell(row=row, column=1, value=idx)
                            ws.cell(row=row, column=2, value=nv.get('ho_ten', ''))
                            ws.cell(row=row, column=3, value=nv.get('ma_so_bhxh', ''))
                            ws.cell(row=row, column=4, value=format_date(nv.get('ngay_sinh')))
                            gt = nv.get('gioi_tinh', '')
                            ws.cell(row=row, column=5, value='Nam' if gt == 'Nam' else 'Nữ' if gt == 'Nữ' else '')
                            ws.cell(row=row, column=6, value=nv.get('so_cccd', ''))
                            ws.cell(row=row, column=7, value=nv.get('chuc_danh_nghe', ''))
                        
                            cd = (nv.get('chuc_danh_nghe') or '').lower()
                            is_quan_ly = any(x in cd for x in ['giám đốc', 'trưởng phòng', 'phó', 'quản lý'])
                            ws.cell(row=row, column=8, value='x' if is_quan_ly else '')
                            is_chuyen_mon_cao = (any(x in cd for x in ['kỹ thuật', 'kĩ thuật']) and any(x in cd for x in ['cao', 'chính', 'kỹ sư']))
                            ws.cell(row=row, column=9, value='x' if is_chuyen_mon_cao else '')
                            is_khac = any(x in cd for x in ['phổ thông', 'lao động', 'tạp vụ', 'bảo vệ', 'tạp vụ', 'lái xe'])
                            ws.cell(row=row, column=11, value='x' if is_khac else '')
                            is_trung = (not is_quan_ly) and (not is_chuyen_mon_cao) and (not is_khac)
                            ws.cell(row=row, column=10, value='x' if is_trung else '')
                        
                            luong = nv.get('luong_bao_hiem', '')
                            heso = nv.get('he_so_luong', '')
                            ws.cell(row=row, column=12, value=f"Hệ số: {heso}" if heso and str(heso).strip() else str(luong) if luong else '')
                            ws.cell(row=row, column=13, value=str(nv.get('phu_cap_chuc_vu', '')) if nv.get('phu_cap_chuc_vu') else '')
                            ws.cell(row=row, column=14, value=f"{nv.get('phu_cap_tnvk', '')}%" if nv.get('phu_cap_tnvk') else '')
                            ws.cell(row=row, column=15, value=f"{nv.get('phu_cap_tnn', '')}%" if nv.get('phu_cap_tnn') else '')
                            ws.cell(row=row, column=16, value='')
                            ws.cell(row=row, column=17, value='')
                            ws.cell(row=row, column=18, value='')
                            ws.cell(row=row, column=19, value='')
                        
                            loai_hd = nv.get('loai_hop_dong', '')
                            ngay_bd = nv.get('ngay_ky_hd') or nv.get('ngay_vao_lam')
                            ngay_kt = nv.get('ngay_ket_thuc')
                            ws.cell(row=row, column=20, value=format_date(ngay_bd) if loai_hd == 'Không xác định thời hạn' else '')
                        
                            if loai_hd == 'Xác định thời hạn':
                                ws.cell(row=row, column=21, value=format_date(ngay_bd))
                                ws.cell(row=row, column=22, value=format_date(ngay_kt) if ngay_kt else '')
                            else:
                                ws.cell(row=row, column=21, value='')
                                ws.cell(row=row, column=22, value='')
                        
                            if loai_hd == 'Thử việc':
                                ws.cell(row=row, column=23, value=format_date(ngay_bd))
                                ws.cell(row=row, column=24, value=format_date(ngay_kt) if ngay_kt else '')
                            else:
                                ws.cell(row=row, column=23, value='')
                                ws.cell(row=row, column=24, value='')
                        
                            ws.cell(row=row, column=25, value=format_date(nv.get('thang_bat_dau_bh')))
                            ws.cell(row=row, column=26, value=format_date(nv.get('thang_ket_thuc_bh')))
                            ws.cell(row=row, column=27, value=nv.get('so_hdld', ''))
                        
                            for col in range(1, 28):
                                cell = ws.cell(row=row, column=col)
                                cell.border = thin_border
                                cell.font = Font(size=10, name='Times New Roman')
                                if col in [1, 3, 4, 5, 6, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26]:
                                    cell.alignment = Alignment(horizontal='center', vertical='center')
                                else:
                                    cell.alignment = Alignment(horizontal='left', vertical='center')
                    
                        total_row = data_row + len(ds_lao_dong)
                        ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=2)
                        ws.cell(row=total_row, column=1, value=f"Tổng cộng: {len(ds_lao_dong)} người")
                        ws.cell(row=total_row, column=1).font = Font(bold=True, size=10, name='Times New Roman')
                        ws.cell(row=total_row, column=1).border = thin_border
                    
                        sign_row = total_row + 3
                        ws.merge_cells(start_row=sign_row, start_column=23, end_row=sign_row, end_column=27)
                        ws.cell(row=sign_row, column=23, value="ĐẠI DIỆN DOANH NGHIỆP")
                        ws.cell(row=sign_row, column=23).font = Font(bold=True, size=11, name='Times New Roman')
                        ws.cell(row=sign_row, column=23).alignment = Alignment(horizontal='center')
                    
                        ws.merge_cells(start_row=sign_row+1, start_column=23, end_row=sign_row+1, end_column=27)
                        ws.cell(row=sign_row+1, column=23, value="(Ký, đóng dấu, ghi rõ họ tên)")
                        ws.cell(row=sign_row+1, column=23).font = Font(size=10, name='Times New Roman')
                        ws.cell(row=sign_row+1, column=23).alignment = Alignment(horizontal='center')
                    
                        ws.merge_cells(start_row=sign_row+2, start_column=23, end_row=sign_row+2, end_column=27)
                        ws.cell(row=sign_row+2, column=23, value=COMPANY_CONFIG.get('dai_dien', 'GIÁM ĐỐC').upper())
                        ws.cell(row=sign_row+2, column=23).font = Font(bold=True, size=11, name='Times New Roman')
                        ws.cell(row=sign_row+2, column=23).alignment = Alignment(horizontal='center')
                    
                        filename = f"Bao_cao_01_PLI_{tu_ngay.strftime('%d%m%Y')}_{den_ngay.strftime('%d%m%Y')}.xlsx"
                        wb.save(filename)
                    
                        with open(filename, "rb") as f:
                            st.download_button(
                                label="📥 TẢI FILE EXCEL",
                                data=f,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                width='stretch'
                            )
                        st.success(f"✅ Đã xuất báo cáo với {len(ds_lao_dong)} lao động")
                        st.cache_data.clear()
                        pass
                    
            else:
                st.info("🔒 Chỉ Admin mới có quyền xuất file Excel báo cáo 01/PLI. Bạn đang ở chế độ xem (Viewer).")
                st.caption("💡 Với quyền Viewer, bạn có thể xem danh sách lao động ở trên nhưng không thể tải file Excel.")
        else:
            st.warning("⚠️ Không có lao động nào đang làm việc trong kỳ báo cáo!")
        
    # ========== QUẢN LÝ CÔNG VĂN & HĐ KINH TẾ ==========

    with tab_bc_tk:
        st.caption("⚙️ Tùy chọn bộ lọc và xuất báo cáo thống kê nhân sự")

        # ── Tùy chọn bộ lọc ──
        with st.expander("⚙️ Tùy chọn xuất báo cáo thống kê nhân sự", expanded=True):
            col_tk_date1, col_tk_date2, col_tk_date3 = st.columns([1, 1, 1])
            with col_tk_date1:
                tk_tu_ngay = st.date_input("📅 Từ ngày:", value=date(date.today().year, 1, 1), key="tk_tu_ngay")
            with col_tk_date2:
                tk_den_ngay = st.date_input("📅 Đến ngày:", value=date.today(), key="tk_den_ngay")
            with col_tk_date3:
                loai_hd_filter = st.selectbox(
                    "Loại hợp đồng:",
                    ["Tất cả", "Không xác định thời hạn", "Thử việc"],
                    key="tk_loai_hd"
                )
            col_tk1, col_tk2 = st.columns([1, 2])
            with col_tk1:
                pass  # placeholder

            # Danh sách tất cả cột bảng nhan_vien (trừ id)
            ALL_COLUMNS_LABELS = {
                "ma_nv":              "Mã NV",
                "ho_ten":             "Họ tên",
                "ngay_sinh":          "Ngày sinh",
                "gioi_tinh":          "Giới tính",
                "chuc_danh_nghe":     "Chức danh",
                "phong_ban_lam_viec": "Phòng ban",
                "loai_hop_dong":      "Loại HĐ",
                "ngay_vao_lam":       "Ngày vào làm",
                "ngay_ky_hd":         "Ngày ký HĐ",
                "so_hdld":            "Số HĐLĐ",
                "so_cccd":            "Số CCCD",
                "thuong_tru":         "Thường trú",
                "dien_thoai":         "Điện thoại",
                "ma_so_bhxh":         "Mã BHXH",
                "thang_bat_dau_bh":   "BĐ đóng BH",
                "so_tai_khoan_nh":    "STK",
                "chi_nhanh_nh":       "Chi nhánh NH",
                "ho_so":              "Hồ sơ",
                "ten_don_vi_thu_huong": "Tên đơn vị thụ hưởng",
            }

            # Thứ tự ưu tiên mặc định (tất cả tích mặc định)
            DEFAULT_PRIORITY = [
                "ma_nv", "ho_ten", "ngay_sinh", "gioi_tinh",
                "chuc_danh_nghe", "phong_ban_lam_viec", "loai_hop_dong",
                "ngay_vao_lam", "ngay_ky_hd", "so_hdld",
                "so_cccd", "thuong_tru", "dien_thoai", "ma_so_bhxh",
                "thang_bat_dau_bh", "so_tai_khoan_nh", "chi_nhanh_nh", "ho_so",
                "ten_don_vi_thu_huong",
            ]
            DEFAULT_CHECKED = set(DEFAULT_PRIORITY)

            with col_tk2:
                st.caption("📋 Chọn các cột cần xuất:")
                col_chk = st.columns(4)
                selected_cols = []
                for idx, (col_key, col_label) in enumerate(ALL_COLUMNS_LABELS.items()):
                    default_val = col_key in DEFAULT_CHECKED
                    checked = col_chk[idx % 4].checkbox(col_label, value=default_val, key=f"tk_col_{col_key}")
                    if checked:
                        selected_cols.append(col_key)

            if st.button("📊 XUẤT THỐNG KÊ NHÂN SỰ (EXCEL)", type="primary", width='stretch', key="btn_tk_nhansu"):
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
                from openpyxl.utils import get_column_letter

                if not selected_cols:
                    st.error("⚠️ Vui lòng chọn ít nhất 1 cột!")
                else:
                    db_tk = st.session_state.db_engine.get_connection()
                    c_tk = db_tk.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

                    # Sắp xếp selected_cols theo thứ tự ưu tiên
                    priority_order = DEFAULT_PRIORITY + [k for k in ALL_COLUMNS_LABELS if k not in DEFAULT_PRIORITY]
                    selected_cols_sorted = sorted(selected_cols, key=lambda x: priority_order.index(x) if x in priority_order else 999)

                    sql_cols = ", ".join(selected_cols_sorted)
                    sql_tk = f"SELECT {sql_cols} FROM nhan_vien WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC') AND ngay_vao_lam BETWEEN %s AND %s"
                    params_tk = [tk_tu_ngay, tk_den_ngay]
                    if loai_hd_filter == "Không xác định thời hạn":
                        sql_tk += " AND loai_hop_dong = %s"
                        params_tk.append("Không xác định thời hạn")
                    elif loai_hd_filter == "Thử việc":
                        sql_tk += " AND trang_thai = 'THU_VIEC'"
                    sql_tk += " ORDER BY STT ASC"
                    c_tk.execute(sql_tk, tuple(params_tk))
                    ds_tk = c_tk.fetchall()
                    db_tk.close()

                    if not ds_tk:
                        st.warning("⚠️ Không có nhân viên nào phù hợp với bộ lọc!")
                    else:
                        thin_border_tk = Border(
                            left=Side(style='thin'), right=Side(style='thin'),
                            top=Side(style='thin'), bottom=Side(style='thin')
                        )
                        header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
                        stat_fill  = PatternFill(start_color="D9E8F5", end_color="D9E8F5", fill_type="solid")

                        wb_tk = Workbook()
                        ws_tk = wb_tk.active
                        ws_tk.title = "Thống kê nhân sự"

                        ten_cong_ty_tk = COMPANY_CONFIG.get("ten_cong_ty", "CÔNG TY CỔ PHẦN CẢNG HÒN LA")
                        dia_chi_tk     = COMPANY_CONFIG.get("dia_chi", "")
                        dien_thoai_tk  = COMPANY_CONFIG.get("dien_thoai_cty", "")
                        n_cols = len(selected_cols_sorted) + 1  # +1 cho cột STT

                        # ── Thông tin công ty ──
                        ws_tk.merge_cells(start_row=1, start_column=1, end_row=1, end_column=n_cols)
                        ws_tk['A1'] = ten_cong_ty_tk
                        ws_tk['A1'].font = Font(bold=True, size=13, name='Times New Roman')
                        ws_tk['A1'].alignment = Alignment(horizontal='center')

                        ws_tk.merge_cells(start_row=2, start_column=1, end_row=2, end_column=n_cols)
                        ws_tk['A2'] = f"Địa chỉ: {dia_chi_tk}  |  ĐT: {dien_thoai_tk}"
                        ws_tk['A2'].font = Font(size=10, name='Times New Roman', italic=True)
                        ws_tk['A2'].alignment = Alignment(horizontal='center')

                        # ── Tiêu đề báo cáo ──
                        loai_hd_label = f" - Loại HĐ: {loai_hd_filter}" if loai_hd_filter != "Tất cả" else ""
                        ws_tk.merge_cells(start_row=4, start_column=1, end_row=4, end_column=n_cols)
                        ws_tk['A4'] = "BÁO CÁO THỐNG KÊ NHÂN SỰ" + loai_hd_label
                        ws_tk['A4'].font = Font(bold=True, size=14, name='Times New Roman')
                        ws_tk['A4'].alignment = Alignment(horizontal='center')

                        ws_tk.merge_cells(start_row=5, start_column=1, end_row=5, end_column=n_cols)
                        ws_tk['A5'] = f"Từ ngày {tk_tu_ngay.strftime('%d/%m/%Y')} đến ngày {tk_den_ngay.strftime('%d/%m/%Y')}  |  Xuất lúc: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
                        ws_tk['A5'].font = Font(size=10, name='Times New Roman', italic=True)
                        ws_tk['A5'].alignment = Alignment(horizontal='center')

                        # ── Header bảng ──
                        header_row_tk = 7
                        ws_tk.cell(row=header_row_tk, column=1, value="STT").font = Font(bold=True, size=10, name='Times New Roman', color="FFFFFF")
                        ws_tk.cell(row=header_row_tk, column=1).fill = header_fill
                        ws_tk.cell(row=header_row_tk, column=1).alignment = Alignment(horizontal='center', vertical='center')
                        ws_tk.cell(row=header_row_tk, column=1).border = thin_border_tk

                        for col_idx, col_key in enumerate(selected_cols_sorted, 2):
                            cell = ws_tk.cell(row=header_row_tk, column=col_idx, value=ALL_COLUMNS_LABELS.get(col_key, col_key))
                            cell.font = Font(bold=True, size=10, name='Times New Roman', color="FFFFFF")
                            cell.fill = header_fill
                            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                            cell.border = thin_border_tk

                        # ── Dữ liệu ──
                        def fmt_val(key, val):
                            if val is None:
                                return ""
                            if 'ngay' in key.lower() or 'thang' in key.lower():
                                return format_date(val)
                            return val

                        date_cols = {k for k in selected_cols_sorted if 'ngay' in k or 'thang' in k}
                        center_cols = {k for k in selected_cols_sorted if k in (
                            "ma_nv","ngay_sinh","gioi_tinh","loai_hop_dong",
                            "ngay_vao_lam","ngay_ky_hd","ngay_ket_thuc","trang_thai",
                            "thang_bat_dau_bh","thang_ket_thuc_bh","he_so_luong",
                            "phu_cap_tnvk","phu_cap_tnn","muc_huong_bhyt"
                        )}

                        for stt_idx, nv in enumerate(ds_tk, 1):
                            row = header_row_tk + stt_idx
                            ws_tk.cell(row=row, column=1, value=stt_idx).border = thin_border_tk
                            ws_tk.cell(row=row, column=1).alignment = Alignment(horizontal='center', vertical='center')
                            ws_tk.cell(row=row, column=1).font = Font(size=10, name='Times New Roman')
                            for col_idx, col_key in enumerate(selected_cols_sorted, 2):
                                raw = nv.get(col_key)
                                val = fmt_val(col_key, raw)
                                cell = ws_tk.cell(row=row, column=col_idx, value=val)
                                cell.font = Font(size=10, name='Times New Roman')
                                cell.border = thin_border_tk
                                if col_key in center_cols:
                                    cell.alignment = Alignment(horizontal='center', vertical='center')
                                else:
                                    cell.alignment = Alignment(horizontal='left', vertical='center')

                        total_row_tk = header_row_tk + len(ds_tk) + 1
                        ws_tk.merge_cells(start_row=total_row_tk, start_column=1, end_row=total_row_tk, end_column=n_cols)
                        ws_tk.cell(row=total_row_tk, column=1, value=f"TỔNG CỘNG: {len(ds_tk)} nhân viên")
                        ws_tk.cell(row=total_row_tk, column=1).font = Font(bold=True, size=11, name='Times New Roman')
                        ws_tk.cell(row=total_row_tk, column=1).alignment = Alignment(horizontal='left')

                        # ── Thống kê theo giới tính ──
                        stat_start = total_row_tk + 2
                        ws_tk.merge_cells(start_row=stat_start, start_column=1, end_row=stat_start, end_column=n_cols)
                        ws_tk.cell(row=stat_start, column=1, value="THỐNG KÊ THEO GIỚI TÍNH").font = Font(bold=True, size=11, name='Times New Roman')

                        nam_count  = sum(1 for nv in ds_tk if (nv.get('gioi_tinh') or '') == 'Nam')
                        nu_count   = sum(1 for nv in ds_tk if (nv.get('gioi_tinh') or '') == 'Nữ')
                        khac_count = len(ds_tk) - nam_count - nu_count

                        for r_offset, (label, cnt) in enumerate([("Nam", nam_count), ("Nữ", nu_count), ("Khác/Chưa xác định", khac_count)], 1):
                            r = stat_start + r_offset
                            ws_tk.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
                            ws_tk.cell(row=r, column=1, value=f"  {label}:").font = Font(size=10, name='Times New Roman')
                            ws_tk.cell(row=r, column=4, value=cnt).font = Font(size=10, name='Times New Roman')
                            for cc in range(1, 5):
                                ws_tk.cell(row=r, column=cc).fill = stat_fill

                        # ── Thống kê theo loại hợp đồng ──
                        stat2_start = stat_start + 5
                        ws_tk.merge_cells(start_row=stat2_start, start_column=1, end_row=stat2_start, end_column=n_cols)
                        ws_tk.cell(row=stat2_start, column=1, value="THỐNG KÊ THEO LOẠI HỢP ĐỒNG").font = Font(bold=True, size=11, name='Times New Roman')

                        hd_types = {"Không xác định thời hạn": 0, "Xác định thời hạn": 0, "Thử việc": 0, "Khác": 0}
                        for nv in ds_tk:
                            loai = (nv.get('loai_hop_dong') or '').strip()
                            if loai in hd_types:
                                hd_types[loai] += 1
                            elif (nv.get('trang_thai') or '') == 'THU_VIEC':
                                hd_types["Thử việc"] += 1
                            else:
                                hd_types["Khác"] += 1

                        for r_offset, (label, cnt) in enumerate(hd_types.items(), 1):
                            r = stat2_start + r_offset
                            ws_tk.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
                            ws_tk.cell(row=r, column=1, value=f"  {label}:").font = Font(size=10, name='Times New Roman')
                            ws_tk.cell(row=r, column=4, value=cnt).font = Font(size=10, name='Times New Roman')
                            for cc in range(1, 5):
                                ws_tk.cell(row=r, column=cc).fill = stat_fill

                        # ── Footer người lập báo cáo ──
                        footer_row = stat2_start + len(hd_types) + 3
                        ws_tk.merge_cells(start_row=footer_row, start_column=1, end_row=footer_row, end_column=3)
                        ws_tk.cell(row=footer_row, column=1, value="NGƯỜI LẬP BÁO CÁO")
                        ws_tk.cell(row=footer_row, column=1).font = Font(bold=True, size=11, name='Times New Roman')
                        ws_tk.cell(row=footer_row, column=1).alignment = Alignment(horizontal='center')

                        ws_tk.merge_cells(start_row=footer_row+1, start_column=1, end_row=footer_row+1, end_column=3)
                        ws_tk.cell(row=footer_row+1, column=1, value="(Ký, ghi rõ họ tên)")
                        ws_tk.cell(row=footer_row+1, column=1).font = Font(size=10, name='Times New Roman', italic=True)
                        ws_tk.cell(row=footer_row+1, column=1).alignment = Alignment(horizontal='center')

                        # ── Độ rộng cột ──
                        ws_tk.column_dimensions['A'].width = 5
                        for col_idx, col_key in enumerate(selected_cols_sorted, 2):
                            if col_key in ('ho_ten', 'thuong_tru', 'nguyen_quan', 'noi_cap_cccd', 'chuc_danh_nghe', 'ten_don_vi_thu_huong'):
                                w = 28
                            elif col_key in ('ma_nv', 'gioi_tinh', 'he_so_luong', 'phu_cap_tnvk', 'phu_cap_tnn'):
                                w = 12
                            elif 'ngay' in col_key or 'thang' in col_key:
                                w = 16
                            else:
                                w = 20
                            ws_tk.column_dimensions[get_column_letter(col_idx)].width = w

                        ws_tk.row_dimensions[header_row_tk].height = 30

                        fname_tk = f"ThongKe_NhanSu_{tk_tu_ngay.strftime('%d%m%Y')}_{tk_den_ngay.strftime('%d%m%Y')}.xlsx"
                        wb_tk.save(fname_tk)
                        with open(fname_tk, "rb") as f:
                            st.download_button(
                                label="📥 TẢI FILE THỐNG KÊ NHÂN SỰ",
                                data=f,
                                file_name=fname_tk,
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                width='stretch'
                            )
                        st.success(f"✅ Đã xuất thống kê {len(ds_tk)} nhân viên với {len(selected_cols_sorted)} cột.")
                        st.cache_data.clear()


    with tab_bc_tanggiam:
        col_from, col_to, col_xuat_bc = st.columns(3)
        with col_from:
            tu_ngay_bc = st.date_input("Từ ngày:", value=date.today().replace(day=1), key="bc_tu")
        with col_to:
            den_ngay_bc = st.date_input("Đến ngày:", value=date.today(), key="bc_den")
        with col_xuat_bc:
            st.write("")  # căn chỉnh cho nút thẳng hàng với 2 ô ngày (bù khoảng trống label)
            xuat_bc = st.button("📄 XUẤT BÁO CÁO WORD", width='stretch')

        if xuat_bc:
            db = st.session_state.db_engine.get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c.execute("""
                SELECT ho_ten, chuc_danh_nghe, phong_ban_lam_viec, loai_hop_dong, ngay_vao_lam,
                       ngay_sinh, so_hdld, ngay_ky_hd
                FROM nhan_vien 
                WHERE trang_thai IN ('DANG_LAM', 'THU_VIEC')
                AND ngay_vao_lam BETWEEN %s AND %s
                ORDER BY ngay_vao_lam ASC
            """, (tu_ngay_bc, den_ngay_bc))
            tang_list = c.fetchall()
            c.execute("""
                SELECT ho_ten, chuc_danh_nghe, phong_ban_lam_viec, loai_hop_dong, ngay_vao_lam, ngay_ket_thuc,
                       ngay_sinh, so_hdld, ngay_ky_hd
                FROM nhan_vien 
                WHERE trang_thai = 'NGHI_VIEC'
                AND ngay_ket_thuc BETWEEN %s AND %s
                ORDER BY ngay_ket_thuc ASC
            """, (tu_ngay_bc, den_ngay_bc))
            giam_list = c.fetchall()
            db.close()
            if tang_list or giam_list:
                file_path = tao_bao_cao_tang_giam(tang_list, giam_list, tu_ngay_bc, den_ngay_bc)
                with open(file_path, "rb") as f:
                    st.download_button(
                        label="📥 TẢI FILE BÁO CÁO (Word)",
                        data=f,
                        file_name=f"Bao_cao_tang_giam_{tu_ngay_bc.strftime('%d%m%Y')}_{den_ngay_bc.strftime('%d%m%Y')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.info("Không có biến động nhân sự trong kỳ.")


elif menu == "📄 Quản lý Công văn & HĐ kinh tế":
    show_quan_ly_cong_van()

# ========== CHAT NỘI BỘ (tách riêng file chat_noi_bo.py) ==========
elif menu == "💬 Chat nội bộ":
    chat_noi_bo.render()

# ========== CHATBOT GIẢI ĐÁP ==========
elif menu == "🤖 Chatbot Giải đáp":
    _chatbot_ensure_table()
 
    # --- Admin thấy thêm tab Quản lý đăng ký ---
    if st.session_state.role in ("admin", "xem_toan_bo"):
        tab_chat, tab_admin_dk = st.tabs(["🤖 AI Tư vấn", "📋 Quản lý đăng ký"])
    else:
        tab_chat = st.container()
        tab_admin_dk = None
 
    with tab_chat:
        st.title("🤖 AI Tư vấn Hành chính Nhân sự")
        st.caption("BHXH · BHYT · Thuế TNCN · Lao động · Thai sản · Thất nghiệp — Phân tích & trích dẫn điều luật cụ thể.")
 
        # Hiển thị trạng thái dữ liệu luật
        _so_dieu = len(_chatbot_all_laws())
        _sem_on = bool(_chatbot_get_voyage_api_key())
        col_info, col_reload = st.columns([5, 1])
        with col_info:
            if _so_dieu:
                mode_txt = "🧠 Semantic search" if _sem_on else "🔤 Keyword search (chưa có VOYAGE_API_KEY)"
                st.caption(f"📚 {_so_dieu} điều luật · {mode_txt}")
            else:
                st.warning(f"⚠️ Chưa có dữ liệu luật trong `{CHATBOT_LAW_DIR}`.")
        with col_reload:
            if st.button("🔄 Nạp lại", key="chatbot_reload"):
                _chatbot_load_all_laws.clear()
                _chatbot_build_law_embeddings.clear()
                st.rerun()
 
        # ========== XÁC ĐỊNH TRẠNG THÁI USER ==========
        # Kiểm tra user đã đăng ký chưa (qua session hoặc tìm trong DB)
        reg = st.session_state.get('_chatbot_reg')
 
        # Nếu chưa có trong session, thử tìm theo thông tin NV đang đăng nhập
        if not reg:
            nv_email = st.session_state.get('nhan_vien_email', '')
            nv_sdt = st.session_state.get('nhan_vien_sdt', '') or st.session_state.get('username', '')
            for val in [nv_email, nv_sdt]:
                if val:
                    found = _chatbot_tim_dang_ky(val)
                    if found:
                        reg = found
                        st.session_state['_chatbot_reg'] = reg
                        break
 
        # ========== ROUTER THEO TRẠNG THÁI ==========
 
        # --- TRẠNG THÁI: ĐÃ DUYỆT + CÒN CREDIT → CHAT ---
        if reg and reg.get('trang_thai') == 'DA_DUYET' and reg.get('da_dung', 0) < reg.get('so_credit', 0):
            con_lai = reg['so_credit'] - reg['da_dung']
            st.success(f"🎫 Còn **{con_lai}/{reg['so_credit']}** câu hỏi · Xin chào **{reg['ho_ten']}**")
 
            if "chatbot_history" not in st.session_state:
                st.session_state.chatbot_history = []
            if "chatbot_display" not in st.session_state:
                st.session_state.chatbot_display = []
 
            # Hiện lịch sử chat
            for msg in st.session_state.chatbot_display:
                if msg["role"] == "user":
                    with st.chat_message("user"):
                        st.markdown(msg["text"])
                else:
                    with st.chat_message("assistant", avatar="⚖️"):
                        st.markdown(msg["text"])
 
            # Ô nhập câu hỏi
            cau_hoi = st.chat_input("Đặt câu hỏi về BHXH, thuế TNCN, thai sản, hợp đồng lao động...")
 
            if cau_hoi:
                # Trừ credit
                _chatbot_tru_credit(reg['id'])
                reg['da_dung'] += 1
                st.session_state['_chatbot_reg'] = reg
 
                st.session_state.chatbot_display.append({"role": "user", "text": cau_hoi})
                st.session_state.chatbot_history.append({"role": "user", "content": cau_hoi})
 
                with st.spinner("⚖️ Đang phân tích điều luật liên quan..."):
                    laws = _chatbot_search_laws(cau_hoi)
                    system_prompt = _chatbot_system_prompt_v2(laws)
                    ket_qua = _chatbot_call_claude_v2(system_prompt, st.session_state.chatbot_history)
 
                st.session_state.chatbot_history.append({"role": "assistant", "content": ket_qua})
                st.session_state.chatbot_display.append({"role": "ai", "text": ket_qua})
                st.rerun()
 
            st.caption("ℹ️ Kết quả mang tính tham khảo. Vui lòng xác nhận với chuyên gia pháp lý cho các quyết định quan trọng.")
            if st.session_state.chatbot_display:
                if st.button("🗑️ Xoá lịch sử trò chuyện"):
                    st.session_state.chatbot_history = []
                    st.session_state.chatbot_display = []
                    st.rerun()
 
        # --- TRẠNG THÁI: HẾT CREDIT ---
        elif reg and reg.get('trang_thai') == 'DA_DUYET' and reg.get('da_dung', 0) >= reg.get('so_credit', 0):
            st.warning("🔒 **Bạn đã sử dụng hết lượt thử nghiệm.**")
            st.markdown(f"""
            Bạn đã dùng hết **{reg['so_credit']} câu hỏi** thử nghiệm.
            Cảm ơn **{reg['ho_ten']}** đã trải nghiệm AI Tư vấn HCNS!
            """)
            col_mua, col_lienhe = st.columns(2)
            with col_mua:
                cfg = CHATBOT_PAYMENT
                if st.button(f"🔄 Mua thêm {cfg['credit_mua_them']} câu — {cfg['gia_mua_them']:,.0f}đ", type="primary", width='stretch'):
                    # Chuyển sang màn thanh toán mua thêm
                    st.session_state['_chatbot_mua_them'] = True
                    st.rerun()
            with col_lienhe:
                st.link_button("📞 Tư vấn gói Doanh nghiệp", "https://zalo.me/0961778150", use_container_width=True)
 
            # Xử lý mua thêm
            if st.session_state.get('_chatbot_mua_them'):
                st.divider()
                st.subheader("💳 Thanh toán mua thêm")
                ma_dk = reg['ma_dang_ky']
                qr = _chatbot_qr_url(f"{ma_dk}-MUATHEM", cfg['gia_mua_them'])
                col_qr, col_info_pay = st.columns([1, 1])
                with col_qr:
                    st.image(qr, caption="Quét QR để thanh toán", width=280)
                with col_info_pay:
                    st.markdown(f"""
                    **Ngân hàng:** Vietcombank  
                    **STK:** `{cfg['stk']}`  
                    **Chủ TK:** {cfg['chu_tk']}  
                    **Số tiền:** {cfg['gia_mua_them']:,.0f}đ  
                    **Nội dung CK:** `{ma_dk}-MUATHEM`
                    """)
                bill_mt = st.file_uploader("📸 Upload ảnh bill chuyển khoản", type=["png","jpg","jpeg"], key="bill_mua_them")
                if bill_mt:
                    if st.button("✅ Đã thanh toán — Gửi bill", type="primary"):
                        ok = _chatbot_upload_bill(f"{ma_dk}-MUATHEM", bill_mt)
                        if ok:
                            st.success("✅ Đã gửi! Vui lòng chờ admin xác nhận (thường 5-15 phút).")
                            st.session_state['_chatbot_mua_them'] = False
                        else:
                            st.error("❌ Lỗi upload. Vui lòng thử lại hoặc gửi bill qua Zalo.")
 
        # --- TRẠNG THÁI: CHỜ THANH TOÁN hoặc CHỜ DUYỆT ---
        elif reg and reg.get('trang_thai') in ('CHO_THANH_TOAN', 'DA_GUI_BILL'):
            ma_dk = reg['ma_dang_ky']
            if reg['trang_thai'] == 'DA_GUI_BILL':
                st.info(f"""
                ⏳ **Đang chờ xác nhận thanh toán**
                
                Mã đăng ký: `{ma_dk}`  
                Chúng tôi sẽ xác nhận trong vòng **5-15 phút** (giờ hành chính).
                Sau khi xác nhận, bạn quay lại đây và nhập email/SĐT để bắt đầu sử dụng.
                """)
                if st.button("🔄 Kiểm tra lại trạng thái"):
                    st.session_state.pop('_chatbot_reg', None)
                    st.rerun()
            else:
                # Chưa gửi bill → hiện QR thanh toán
                st.subheader("💳 Thanh toán đăng ký")
                cfg = CHATBOT_PAYMENT
                qr = _chatbot_qr_url(ma_dk)
                col_qr2, col_info2 = st.columns([1, 1])
                with col_qr2:
                    st.image(qr, caption="Quét QR để thanh toán", width=280)
                with col_info2:
                    st.markdown(f"""
                    **Ngân hàng:** Vietcombank  
                    **STK:** `{cfg['stk']}`  
                    **Chủ TK:** {cfg['chu_tk']}  
                    **Số tiền:** {cfg['so_tien']:,.0f}đ  
                    **Nội dung CK:** `{ma_dk}`
                    
                    ---
                    Sau khi chuyển khoản, upload ảnh bill bên dưới:
                    """)
                bill_file = st.file_uploader("📸 Upload ảnh bill chuyển khoản", type=["png","jpg","jpeg"], key="bill_dangky")
                if bill_file:
                    if st.button("✅ Đã thanh toán — Gửi bill xác nhận", type="primary"):
                        ok = _chatbot_upload_bill(ma_dk, bill_file)
                        if ok:
                            reg['trang_thai'] = 'DA_GUI_BILL'
                            st.session_state['_chatbot_reg'] = reg
                            _chatbot_gui_thong_bao_admin(reg)
                            st.success("✅ Đã gửi bill! Vui lòng chờ xác nhận (thường 5-15 phút).")
                            st.rerun()
                        else:
                            st.error("❌ Lỗi upload. Thử lại hoặc gửi bill qua Zalo: 0961778150")
 
        # --- TRẠNG THÁI: CHƯA ĐĂNG KÝ → LANDING PAGE + FORM ---
        else:
            # === LANDING PAGE ===
            st.markdown("""
            <div style="text-align:center; padding:20px 0 10px;">
                <div style="font-size:48px; margin-bottom:8px;">⚖️</div>
                <h2 style="color:#1e3a5f; margin-bottom:4px;">AI Tư vấn Hành chính Nhân sự</h2>
                <p style="color:#6b7280; font-size:14px; line-height:1.7; max-width:600px; margin:0 auto;">
                    Hỏi bất kỳ câu hỏi nào về <b>BHXH, BHYT, Thuế TNCN, Hợp đồng lao động, Thai sản, 
                    Trợ cấp thất nghiệp</b> — AI sẽ phân tích và trích dẫn điều luật cụ thể.
                </p>
            </div>
            """, unsafe_allow_html=True)
 
            # Demo 1 câu hỏi mẫu (hardcode)
            with st.expander("💬 Xem ví dụ câu hỏi & trả lời", expanded=False):
                st.markdown("""
**Câu hỏi:** *Tôi đã đóng BHXH 10 năm, nghỉ ốm 45 ngày. Mức hưởng tính thế nào?*
 
### 📋 Tóm tắt
Với 10 năm đóng BHXH, bạn được hưởng ốm đau 75% lương đóng BH, tối đa 30 ngày/năm trong điều kiện làm việc bình thường.
 
### 📖 Phân tích chi tiết
Theo **Điều 26 Luật BHXH 2014**, người lao động đóng BHXH từ đủ 15 năm trở xuống được nghỉ ốm tối đa **30 ngày/năm** (làm việc trong điều kiện bình thường). 15 ngày còn lại (45 - 30 = 15 ngày) nếu mắc bệnh dài ngày theo danh mục Bộ Y tế, bạn hưởng mức thấp hơn (50-65% tuỳ thời gian đóng)...
 
### ⚖️ Căn cứ pháp lý
- Điều 26 Luật BHXH 2014 — Thời gian hưởng chế độ ốm đau  
- Điều 28 Luật BHXH 2014 — Mức hưởng chế độ ốm đau
                """)
 
            st.divider()
 
            # Đăng nhập nếu đã có tài khoản
            st.markdown("#### 🔑 Đã đăng ký? Nhập email hoặc SĐT để tiếp tục")
            col_login, col_btn_login = st.columns([3, 1])
            with col_login:
                login_val = st.text_input("Email hoặc SĐT đã đăng ký", key="chatbot_login_input", label_visibility="collapsed", placeholder="Email hoặc SĐT đã đăng ký...")
            with col_btn_login:
                if st.button("→ Vào", key="chatbot_login_btn", width='stretch'):
                    if login_val:
                        found = _chatbot_tim_dang_ky(login_val)
                        if found:
                            st.session_state['_chatbot_reg'] = found
                            st.rerun()
                        else:
                            st.error("Không tìm thấy. Vui lòng đăng ký mới bên dưới.")
                    else:
                        st.warning("Nhập email hoặc SĐT.")
 
            st.divider()
 
            # Form đăng ký mới
            st.markdown(f"""
            #### 🔓 Đăng ký thử nghiệm
            **{CHATBOT_PAYMENT['so_tien']:,.0f}đ** · {CHATBOT_PAYMENT['credit_moi']} câu hỏi tư vấn AI có viện dẫn pháp luật
            """)
            with st.form("chatbot_register"):
                col_r1, col_r2 = st.columns(2)
                with col_r1:
                    reg_ten = st.text_input("Họ và tên *", placeholder="Nguyễn Văn A")
                    reg_email = st.text_input("Email *", placeholder="email@congty.vn")
                with col_r2:
                    reg_sdt = st.text_input("Số điện thoại *", placeholder="0912345678")
                    reg_cty = st.text_input("Tên công ty (tuỳ chọn)", placeholder="Công ty ABC")
 
                if st.form_submit_button("📝 Đăng ký & Thanh toán →", type="primary", use_container_width=True):
                    # Validate
                    loi = []
                    if not reg_ten.strip():
                        loi.append("Họ tên")
                    if not reg_email.strip() or "@" not in reg_email:
                        loi.append("Email")
                    if not reg_sdt.strip() or len(reg_sdt.strip()) < 9:
                        loi.append("SĐT")
                    if loi:
                        st.error(f"Vui lòng nhập đúng: {', '.join(loi)}")
                    else:
                        ma = _chatbot_tao_dang_ky(reg_ten, reg_email, reg_sdt, reg_cty)
                        if ma == "TRUNG_EMAIL":
                            st.error("Email này đã đăng ký. Nhập email ở ô phía trên để tiếp tục.")
                        elif ma == "TRUNG_SDT":
                            st.error("SĐT này đã đăng ký. Nhập SĐT ở ô phía trên để tiếp tục.")
                        elif ma:
                            new_reg = _chatbot_tim_dang_ky(reg_email)
                            if new_reg:
                                st.session_state['_chatbot_reg'] = new_reg
                                st.success(f"✅ Đăng ký thành công! Mã: `{ma}`. Chuyển sang thanh toán...")
                                st.rerun()
                        else:
                            st.error("Lỗi tạo đăng ký. Vui lòng thử lại.")
 
    # ========== TAB ADMIN: QUẢN LÝ ĐĂNG KÝ ==========
    if tab_admin_dk is not None:
        with tab_admin_dk:
            st.subheader("📋 Quản lý đăng ký Chatbot")
            try:
                db_adm = st.session_state.db_engine.get_connection()
                c_adm = db_adm.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                c_adm.execute("""
                    SELECT * FROM chatbot_dang_ky
                    ORDER BY
                        CASE trang_thai
                            WHEN 'DA_GUI_BILL' THEN 1
                            WHEN 'CHO_THANH_TOAN' THEN 2
                            WHEN 'DA_DUYET' THEN 3
                            ELSE 4
                        END,
                        created_at DESC
                """)
                ds_dk = c_adm.fetchall()
                db_adm.close()
            except Exception:
                ds_dk = []
 
            if not ds_dk:
                st.info("Chưa có đăng ký nào.")
            else:
                # Thống kê nhanh
                cho_duyet = sum(1 for d in ds_dk if d['trang_thai'] == 'DA_GUI_BILL')
                da_duyet = sum(1 for d in ds_dk if d['trang_thai'] == 'DA_DUYET')
                st.caption(f"📊 Tổng: {len(ds_dk)} · Chờ duyệt: {cho_duyet} · Đã duyệt: {da_duyet}")
 
                for dk in ds_dk:
                    trang_thai_icon = {
                        'CHO_THANH_TOAN': '⏳ Chờ TT',
                        'DA_GUI_BILL': '🔔 Chờ duyệt',
                        'DA_DUYET': '✅ Đã duyệt',
                        'TU_CHOI': '❌ Từ chối',
                        'HET_CREDIT': '🔒 Hết credit',
                    }.get(dk['trang_thai'], dk['trang_thai'])
 
                    with st.expander(
                        f"{trang_thai_icon} | {dk['ho_ten']} | {dk['email']} | {dk['ma_dang_ky']}"
                        f" | Credit: {dk['da_dung']}/{dk['so_credit']}"
                    ):
                        col_dk1, col_dk2 = st.columns(2)
                        with col_dk1:
                            st.markdown(f"""
                            **Họ tên:** {dk['ho_ten']}  
                            **Email:** {dk['email']}  
                            **SĐT:** {dk['dien_thoai']}  
                            **Công ty:** {dk.get('cong_ty') or '—'}  
                            **Đăng ký lúc:** {dk['created_at']}
                            """)
                        with col_dk2:
                            # Hiện ảnh bill nếu có
                            if dk.get('anh_bill'):
                                try:
                                    sb = get_supabase_storage()
                                    if sb:
                                        bill_bytes = sb.storage.from_(SUPABASE_BUCKET).download(dk['anh_bill'])
                                        if bill_bytes:
                                            st.image(bill_bytes, caption="Ảnh bill CK", width=250)
                                except Exception:
                                    st.caption("(Không tải được ảnh bill)")
                            else:
                                st.caption("Chưa có ảnh bill")
 
                        # Nút hành động
                        if dk['trang_thai'] == 'DA_GUI_BILL':
                            col_a1, col_a2 = st.columns(2)
                            with col_a1:
                                credit_duyet = st.number_input(
                                    "Số credit cấp", min_value=1, value=CHATBOT_PAYMENT['credit_moi'],
                                    key=f"credit_{dk['id']}"
                                )
                                if st.button("✅ DUYỆT", key=f"duyet_{dk['id']}", type="primary",
                                             width='stretch', disabled=not can_edit()):
                                    try:
                                        db_d = st.session_state.db_engine.get_connection()
                                        c_d = db_d.cursor()
                                        c_d.execute("""
                                            UPDATE chatbot_dang_ky
                                            SET trang_thai='DA_DUYET', so_credit=%s,
                                                duyet_luc=NOW(), duyet_boi=%s
                                            WHERE id=%s
                                        """, (credit_duyet, st.session_state.get('username','admin'), dk['id']))
                                        db_d.commit()
                                        db_d.close()
                                        st.success(f"✅ Đã duyệt {dk['ho_ten']} — {credit_duyet} credit")
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"Lỗi: {e}")
                            with col_a2:
                                if st.button("❌ Từ chối", key=f"tuchoi_{dk['id']}",
                                             width='stretch', disabled=not can_edit()):
                                    try:
                                        db_tc = st.session_state.db_engine.get_connection()
                                        c_tc = db_tc.cursor()
                                        c_tc.execute("""
                                            UPDATE chatbot_dang_ky
                                            SET trang_thai='TU_CHOI', duyet_luc=NOW(), duyet_boi=%s
                                            WHERE id=%s
                                        """, (st.session_state.get('username','admin'), dk['id']))
                                        db_tc.commit()
                                        db_tc.close()
                                        st.warning(f"Đã từ chối {dk['ho_ten']}")
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"Lỗi: {e}")
 
                        elif dk['trang_thai'] == 'DA_DUYET':
                            # Cho phép admin cộng thêm credit
                            them_credit = st.number_input(
                                "Cộng thêm credit", min_value=1, value=5,
                                key=f"them_cr_{dk['id']}"
                            )
                            if st.button(f"➕ Cộng {them_credit} credit", key=f"cong_{dk['id']}",
                                         disabled=not can_edit()):
                                try:
                                    db_cc = st.session_state.db_engine.get_connection()
                                    c_cc = db_cc.cursor()
                                    c_cc.execute("""
                                        UPDATE chatbot_dang_ky
                                        SET so_credit = so_credit + %s
                                        WHERE id = %s
                                    """, (them_credit, dk['id']))
                                    db_cc.commit()
                                    db_cc.close()
                                    st.success(f"✅ Đã cộng {them_credit} credit cho {dk['ho_ten']}")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Lỗi: {e}")

elif menu == "📥 Nhập/Xuất Excel" and st.session_state.role in ("admin", "xem_toan_bo"):
    render_import_export_ui(
        lambda: st.session_state.db_engine.get_connection(),
        extra_caption=f"Công ty: {st.session_state.tenant.get('ten_cty', '')}"
    )

# ========== HƯỚNG DẪN SỬ DỤNG ==========
elif menu == "🔑 Quản lý MK":
    st.title("🔑 Quản lý mật khẩu")

    if st.session_state.role in ("admin", "xem_toan_bo"):
        tab_doi_mk, tab_admin_reset, tab_phan_quyen = st.tabs(
            ["🔒 Đổi mật khẩu của tôi", "🛠️ Reset mật khẩu nhân viên (Admin)", "🛡️ Phân quyền hệ thống"])
    else:
        tab_doi_mk = st.container()

    with tab_doi_mk:
        st.subheader("🔒 Đổi mật khẩu của tôi")
        st.caption("Nếu nghi ngờ mật khẩu bị lộ, hãy chủ động đổi ngay tại đây.")
        mk_hien_tai = st.text_input("Mật khẩu hiện tại:", type="password", key="doimk_hientai")
        mk_moi_ts = st.text_input("Mật khẩu mới:", type="password", key="doimk_moi")
        mk_moi_ts2 = st.text_input("Nhập lại mật khẩu mới:", type="password", key="doimk_moi2")
        if st.button("✅ Xác nhận đổi mật khẩu", key="btn_doi_mk_tuchu", type="primary"):
            if not st.session_state.get('nhan_vien_id'):
                st.error("❌ Không xác định được tài khoản đang đăng nhập (tài khoản Admin hệ thống không đổi được ở đây).")
            else:
                db_dmk = st.session_state.db_engine.get_connection()
                c_dmk = db_dmk.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                c_dmk.execute("SELECT mat_khau_hash FROM nhan_vien WHERE id=%s", (st.session_state.nhan_vien_id,))
                row_dmk = c_dmk.fetchone()
                if not row_dmk or not bcrypt.checkpw(mk_hien_tai.encode(), row_dmk['mat_khau_hash'].encode()):
                    db_dmk.close()
                    st.error("❌ Mật khẩu hiện tại không đúng.")
                elif len(mk_moi_ts) < 6:
                    db_dmk.close()
                    st.error("Mật khẩu mới phải có ít nhất 6 ký tự.")
                elif mk_moi_ts != mk_moi_ts2:
                    db_dmk.close()
                    st.error("Hai mật khẩu nhập lại không khớp.")
                else:
                    c_dmk2 = db_dmk.cursor()
                    new_hash_ts = bcrypt.hashpw(mk_moi_ts.encode(), bcrypt.gensalt()).decode()
                    c_dmk2.execute("UPDATE nhan_vien SET mat_khau_hash=%s WHERE id=%s",
                                   (new_hash_ts, st.session_state.nhan_vien_id))
                    db_dmk.commit(); db_dmk.close()
                    st.success("✅ Đổi mật khẩu thành công!")
                    st.cache_data.clear()

    if st.session_state.role in ("admin", "xem_toan_bo"):
        with tab_admin_reset:
            st.subheader("🛠️ Reset mật khẩu nhân viên (dành cho trường hợp quên mật khẩu & không có Email liên hệ)")
            st.caption("Mật khẩu sẽ được đặt lại về mặc định = **số điện thoại** của nhân viên, "
                       "và nhân viên sẽ bị buộc đổi mật khẩu ngay trong lần đăng nhập tiếp theo.")
            db_rst = st.session_state.db_engine.get_connection()
            c_rst = db_rst.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c_rst.execute("""SELECT id, ho_ten, ma_nv, dien_thoai FROM nhan_vien
                              WHERE trang_thai IN ('DANG_LAM','THU_VIEC') ORDER BY ho_ten""")
            ds_nv_rst = c_rst.fetchall()
            db_rst.close()
            tuy_chon_rst = {f"{r['ho_ten']} ({r['ma_nv']}) - SĐT: {r.get('dien_thoai') or 'chưa có'}": r for r in ds_nv_rst}
            chon_rst = st.selectbox("Chọn nhân viên:", ["-- Chọn --"] + list(tuy_chon_rst.keys()),
                                     key="chon_reset_mk", help="💡 Gõ tên/mã NV để tìm nhanh")
            if chon_rst != "-- Chọn --":
                nv_rst = tuy_chon_rst[chon_rst]
                if not nv_rst.get('dien_thoai'):
                    st.error("❌ Nhân viên chưa có số điện thoại trong hồ sơ nên không thể đặt mật khẩu mặc định. Vui lòng cập nhật SĐT trước.")
                else:
                    st.warning(f"Sẽ đặt lại mật khẩu của **{nv_rst['ho_ten']}** về **{nv_rst['dien_thoai']}** và buộc đổi mật khẩu ở lần đăng nhập tới.")
                    if st.button("🔄 Xác nhận Reset mật khẩu", key=f"btn_reset_mk_{nv_rst['id']}", type="primary", disabled=not can_edit()):
                        try:
                            db_r2 = st.session_state.db_engine.get_connection()
                            c_r2 = db_r2.cursor()
                            new_hash_rst = bcrypt.hashpw(nv_rst['dien_thoai'].encode(), bcrypt.gensalt()).decode()
                            c_r2.execute("UPDATE nhan_vien SET mat_khau_hash=%s, phai_doi_mat_khau=TRUE WHERE id=%s",
                                         (new_hash_rst, nv_rst['id']))
                            db_r2.commit(); db_r2.close()
                            st.success(f"✅ Đã reset mật khẩu về SĐT ({nv_rst['dien_thoai']}). Thông báo cho nhân viên đăng nhập lại và đổi mật khẩu mới.")
                            st.cache_data.clear()
                        except Exception as e:
                            st.error(f"❌ Không thể reset mật khẩu: {e}")

        with tab_phan_quyen:
            st.subheader("🛡️ Phân quyền hệ thống")
            st.caption(
                "Kể từ nay, TẤT CẢ mọi người (kể cả Admin/HR/Văn thư/Kế toán lương) đều đăng nhập bằng "
                "**số điện thoại nhân viên thật** trong hồ sơ — không còn tài khoản hệ thống riêng nữa. "
                "Tại đây, Admin chỉ định NHÂN VIÊN NÀO giữ vai trò gì; vai trò quyết định menu & quyền "
                "họ thấy sau khi đăng nhập."
            )

            VAI_TRO_LUA_CHON = [
                ("nhan_vien", "👤 Nhân viên (mặc định, không có quyền quản trị)"),
                ("admin", "🛡️ Admin (toàn quyền hệ thống)"),
                ("admin_bcc", "📋 Admin BCC (theo dõi, phê duyệt, điều chỉnh chấm công & OT)"),
                ("hr", "👥 HR (nhân sự)"),
                ("van_thu", "📄 Văn thư (Công văn & HĐ kinh tế)"),
                ("kt_luong", "💰 Kế toán lương (Chấm công & Tính thu nhập)"),
                ("xem_toan_bo", "👁️ Xem toàn bộ - không chỉnh sửa (thấy hết menu/tab, nút Lưu/Sửa/Xóa bị khóa)"),
            ]
            NHAN_VAI_TRO = dict(VAI_TRO_LUA_CHON)

            # Đảm bảo cột vai_tro tồn tại — phòng trường hợp tenant được tạo trước khi có tính năng này.
            try:
                db_pq0 = st.session_state.db_engine.get_connection()
                c_pq0 = db_pq0.cursor()
                c_pq0.execute("ALTER TABLE nhan_vien ADD COLUMN IF NOT EXISTS vai_tro VARCHAR(20) DEFAULT 'nhan_vien'")
                db_pq0.commit()
                db_pq0.close()
            except Exception:
                pass

            db_pq = st.session_state.db_engine.get_connection()
            c_pq = db_pq.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c_pq.execute("""SELECT id, ho_ten, ma_nv, dien_thoai, phong_ban_lam_viec, chuc_vu, vai_tro
                            FROM nhan_vien WHERE trang_thai IN ('DANG_LAM','THU_VIEC') ORDER BY ho_ten""")
            ds_nv_pq = c_pq.fetchall()
            db_pq.close()

            dang_giu_quyen = [r for r in ds_nv_pq if (r.get('vai_tro') or 'nhan_vien') != 'nhan_vien']
            so_admin_hien_tai = sum(1 for r in ds_nv_pq if r.get('vai_tro') == 'admin')

            st.markdown(f"**📋 Đang có {len(dang_giu_quyen)} người giữ vai trò đặc biệt** (Admin hiện tại: {so_admin_hien_tai} người)")
            if dang_giu_quyen:
                df_pq = pd.DataFrame([{
                    "Họ tên": r['ho_ten'], "Mã NV": r['ma_nv'],
                    "Phòng ban": r.get('phong_ban_lam_viec') or '',
                    "Vai trò": NHAN_VAI_TRO.get(r.get('vai_tro'), r.get('vai_tro')),
                } for r in dang_giu_quyen])
                st.dataframe(df_pq, width='stretch', hide_index=True)
            else:
                st.warning("⚠️ Chưa có ai được cấp vai trò quản trị nào — hãy chỉ định ít nhất 1 Admin bên dưới.")

            st.divider()
            st.markdown("##### ✏️ Thay đổi vai trò cho 1 nhân viên")
            tuy_chon_pq = {f"{r['ho_ten']} ({r['ma_nv']}) — hiện tại: {NHAN_VAI_TRO.get(r.get('vai_tro') or 'nhan_vien')}": r
                            for r in ds_nv_pq}
            chon_pq = st.selectbox("Chọn nhân viên:", ["-- Chọn --"] + list(tuy_chon_pq.keys()), key="chon_nv_phan_quyen")
            if chon_pq != "-- Chọn --":
                nv_pq = tuy_chon_pq[chon_pq]
                vai_tro_hien_tai = nv_pq.get('vai_tro') or 'nhan_vien'
                idx_mac_dinh = [k for k, _ in VAI_TRO_LUA_CHON].index(vai_tro_hien_tai) if vai_tro_hien_tai in NHAN_VAI_TRO else 0
                vai_tro_moi_label = st.selectbox("Vai trò mới:", [v for _, v in VAI_TRO_LUA_CHON],
                                                  index=idx_mac_dinh, key=f"vt_moi_{nv_pq['id']}")
                vai_tro_moi = [k for k, v in VAI_TRO_LUA_CHON if v == vai_tro_moi_label][0]

                if not nv_pq.get('dien_thoai') and vai_tro_moi != 'nhan_vien':
                    st.error("❌ Nhân viên này chưa có số điện thoại trong hồ sơ — cần có SĐT để đăng nhập trước khi cấp quyền.")
                elif st.button("💾 Lưu vai trò", key=f"btn_luu_vt_{nv_pq['id']}", type="primary", disabled=not can_edit()):
                    # Chặn tự hạ quyền / hạ quyền người khác nếu đó là Admin CUỐI CÙNG còn lại
                    if vai_tro_hien_tai == 'admin' and vai_tro_moi != 'admin' and so_admin_hien_tai <= 1:
                        st.error("❌ Không thể thực hiện: đây là Admin CUỐI CÙNG của công ty. "
                                 "Hãy chỉ định 1 Admin khác trước khi đổi vai trò người này.")
                    else:
                        try:
                            db_pq2 = st.session_state.db_engine.get_connection()
                            c_pq2 = db_pq2.cursor()
                            c_pq2.execute("UPDATE nhan_vien SET vai_tro=%s WHERE id=%s", (vai_tro_moi, nv_pq['id']))
                            db_pq2.commit()
                            db_pq2.close()
                            st.success(f"✅ Đã đặt vai trò của **{nv_pq['ho_ten']}** thành **{vai_tro_moi_label}**.")
                            st.cache_data.clear()
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Không thể lưu vai trò: {e}")

elif menu == "🖼️ Tạo ảnh thẻ NV":
    photo_card_gender.render()

elif menu == "🔍 Audit Dashboard":
    st.title("🔍 Audit lệch số liệu Dashboard")
    st.caption(
        "Công cụ nội bộ (chỉ Admin) để kiểm tra các biểu đồ trong 📊 Dashboard có "
        "cùng Tổng với nhau không, và tìm nguyên nhân nếu bị lệch (LIMIT, filter thiếu "
        "đồng bộ, dữ liệu NULL/rỗng...). Dùng ngay kết nối DB hiện tại của công ty bạn "
        "đang đăng nhập — không cần cấu hình gì thêm."
    )
    if st.session_state.role != "admin":
        st.error("❌ Chỉ Admin mới được dùng công cụ này!")
    else:
        db_a = st.session_state.db_engine.get_connection()
        c_a = db_a.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

        # Tiêu chuẩn lọc nhân sự — PHẢI giống hệt biến DK_CHUAN_NV dùng ở 📊 Dashboard.
        # Nếu sửa điều kiện lọc ở Dashboard, nhớ sửa lại y hệt ở đây.
        DK_CHUAN_NV_AUDIT = "trang_thai IN ('DANG_LAM', 'THU_VIEC') AND so_hdld IS NOT NULL AND so_hdld != ''"

        c_a.execute(f"SELECT COUNT(*) AS t FROM nhan_vien WHERE {DK_CHUAN_NV_AUDIT}")
        TONG_CHUAN = c_a.fetchone()['t']

        st.subheader("0️⃣ Tổng chuẩn (dùng để đối chiếu mọi biểu đồ)")
        st.success(f"**Tổng chuẩn = {TONG_CHUAN} nhân viên**")
        st.caption("Điều kiện: `" + DK_CHUAN_NV_AUDIT + "`")
        st.divider()

        # 1) So sánh Tổng từng biểu đồ với Tổng chuẩn
        st.subheader("1️⃣ So sánh Tổng từng biểu đồ với Tổng chuẩn")
        checks_a = []
        c_a.execute(f"SELECT COUNT(*) t FROM nhan_vien WHERE {DK_CHUAN_NV_AUDIT}")
        checks_a.append(("Cơ cấu theo Phòng ban", c_a.fetchone()['t'], "không filter thêm"))
        c_a.execute(f"SELECT COUNT(*) t FROM nhan_vien WHERE {DK_CHUAN_NV_AUDIT}")
        checks_a.append(("Cơ cấu theo Giới tính", c_a.fetchone()['t'], "không filter thêm"))
        c_a.execute(f"SELECT COUNT(*) t FROM nhan_vien WHERE {DK_CHUAN_NV_AUDIT}")
        checks_a.append(("Cơ cấu theo Trình độ học vấn", c_a.fetchone()['t'], "không filter thêm"))
        c_a.execute(f"""
            SELECT COUNT(*) t FROM nhan_vien
            WHERE {DK_CHUAN_NV_AUDIT} AND chuc_danh_nghe IS NOT NULL AND chuc_danh_nghe != ''
        """)
        checks_a.append(("Cơ cấu theo Chức danh (SQL đầy đủ)", c_a.fetchone()['t'], "filter thêm: chuc_danh_nghe IS NOT NULL AND != ''"))
        c_a.execute(f"SELECT COUNT(*) t FROM nhan_vien WHERE {DK_CHUAN_NV_AUDIT} AND ngay_sinh IS NOT NULL")
        checks_a.append(("Cơ cấu theo Độ tuổi", c_a.fetchone()['t'], "filter thêm: ngay_sinh IS NOT NULL (chủ ý)"))
        c_a.execute(f"""
            SELECT COUNT(*) t FROM nhan_vien
            WHERE ngay_vao_lam >= (CURRENT_DATE - INTERVAL '6 months') AND {DK_CHUAN_NV_AUDIT}
        """)
        checks_a.append(("Xu hướng tuyển dụng 6 tháng", c_a.fetchone()['t'], "filter thêm: chỉ 6 tháng gần nhất (chủ ý)"))

        rows_out_a = []
        for ten_bd, tong_bd, ghi_chu in checks_a:
            lech = tong_bd - TONG_CHUAN
            if lech == 0:
                trang_thai = "✅ Khớp"
            elif "chủ ý" in ghi_chu:
                trang_thai = f"ℹ️ Lệch {lech:+d} — chủ ý (xem ghi chú)"
            else:
                trang_thai = f"🚨 LỆCH {lech:+d} — CẦN KIỂM TRA"
            rows_out_a.append({
                "Biểu đồ": ten_bd, "Tổng": tong_bd, "Tổng chuẩn": TONG_CHUAN,
                "Chênh lệch": lech, "Trạng thái": trang_thai, "Ghi chú": ghi_chu,
            })
        st.dataframe(pd.DataFrame(rows_out_a), width='stretch', hide_index=True)
        st.divider()

        # 2) Chi tiết toàn bộ chức danh — lộ ra những chức danh sẽ bị LIMIT cắt mất
        st.subheader("2️⃣ Chi tiết toàn bộ Chức danh (không LIMIT)")
        c_a.execute(f"""
            SELECT chuc_danh_nghe, COUNT(*) as so_luong
            FROM nhan_vien WHERE {DK_CHUAN_NV_AUDIT}
            AND chuc_danh_nghe IS NOT NULL AND chuc_danh_nghe != ''
            GROUP BY chuc_danh_nghe
            ORDER BY so_luong DESC
        """)
        df_full_role_a = pd.DataFrame(c_a.fetchall())
        if not df_full_role_a.empty:
            df_full_role_a.index = range(1, len(df_full_role_a) + 1)
            df_full_role_a['Sẽ bị LIMIT 10 cắt mất?'] = [
                "🚨 CÓ" if i > 10 else "✅ Không" for i in df_full_role_a.index
            ]
            st.dataframe(df_full_role_a, width='stretch')
            so_bi_cat = (df_full_role_a.index > 10).sum()
            if so_bi_cat > 0:
                nguoi_bi_cat = df_full_role_a[df_full_role_a.index > 10]['so_luong'].sum()
                st.warning(f"⚠️ {so_bi_cat} chức danh ({nguoi_bi_cat} người) sẽ bị mất nếu SQL có LIMIT 10.")
            else:
                st.success("✅ Tổng số chức danh ≤ 10, không có nguy cơ bị LIMIT cắt mất dữ liệu.")
                st.cache_data.clear()
        else:
            st.info("Không có dữ liệu chức danh.")
        st.divider()

        # 3) Nhân viên bị loại khỏi Tổng chuẩn kèm lý do
        st.subheader("3️⃣ Nhân viên KHÔNG nằm trong Tổng chuẩn (kèm lý do)")
        c_a.execute(f"""
            SELECT ma_nv, ho_ten, trang_thai, so_hdld,
                CASE
                    WHEN trang_thai NOT IN ('DANG_LAM','THU_VIEC') THEN 'Trạng thái không phải Đang làm/Thử việc: ' || COALESCE(trang_thai,'(rỗng)')
                    WHEN so_hdld IS NULL OR so_hdld = '' THEN 'Chưa có số HĐLĐ (hồ sơ chưa hoàn thiện)'
                    ELSE 'Không rõ'
                END as ly_do
            FROM nhan_vien
            WHERE NOT ({DK_CHUAN_NV_AUDIT})
            ORDER BY trang_thai, ho_ten
        """)
        loai_tru_a = c_a.fetchall()
        if loai_tru_a:
            st.dataframe(pd.DataFrame(loai_tru_a), width='stretch', hide_index=True)
            st.caption(f"Tổng cộng {len(loai_tru_a)} nhân viên bị loại khỏi Tổng chuẩn vì lý do trên.")
        else:
            st.info("Không có nhân viên nào bị loại — mọi bản ghi đều đạt Tổng chuẩn.")
        st.divider()

        # 4) Cảnh báo NULL/rỗng ở các trường dùng để group biểu đồ
        st.subheader("4️⃣ Kiểm tra NULL/rỗng ở các trường dùng để nhóm biểu đồ")
        truong_can_kiem_tra_a = {
            "gioi_tinh": "Giới tính", "trinh_do": "Trình độ",
            "phong_ban_lam_viec": "Phòng ban", "ngay_sinh": "Ngày sinh",
            "chuc_danh_nghe": "Chức danh",
        }
        rows_null_a = []
        chi_tiet_null_a = {}
        for cot, nhan in truong_can_kiem_tra_a.items():
            c_a.execute(f"""
                SELECT ma_nv, ho_ten, trang_thai
                FROM nhan_vien
                WHERE {DK_CHUAN_NV_AUDIT} AND ({cot} IS NULL OR {cot}::text = '')
                ORDER BY ho_ten
            """)
            ds_nv_thieu = c_a.fetchall()
            so_luong_null = len(ds_nv_thieu)
            chi_tiet_null_a[nhan] = ds_nv_thieu
            rows_null_a.append({
                "Trường": nhan, "Số nhân viên NULL/rỗng": so_luong_null,
                "Trạng thái": "✅ Không có" if so_luong_null == 0 else f"⚠️ Có {so_luong_null} người thiếu dữ liệu"
            })
        st.dataframe(pd.DataFrame(rows_null_a), width='stretch', hide_index=True)

        # Liệt kê CỤ THỂ tên/mã nhân viên cho từng trường bị thiếu, để không phải tự viết SQL
        for nhan, ds_nv_thieu in chi_tiet_null_a.items():
            if ds_nv_thieu:
                with st.expander(f"👤 Danh sách nhân viên thiếu '{nhan}' ({len(ds_nv_thieu)} người)"):
                    st.dataframe(pd.DataFrame(ds_nv_thieu), width='stretch', hide_index=True)
        db_a.close()

elif menu == "📘 Hướng dẫn sử dụng":
    st.title("📘 Hướng dẫn sử dụng HRM Master")
    st.caption("Tổng quan các chức năng chính của hệ thống - dành cho người dùng mới.")

    st.markdown("""
### 📊 Dashboard
Bức tranh tổng quan về nhân sự: tổng số nhân viên, cơ cấu theo phòng ban, độ tuổi, giới tính, xu hướng tuyển dụng...
giúp Ban điều hành nắm tình hình chỉ trong vài giây, không cần chờ báo cáo tổng hợp thủ công.

### ✅ Nhân viên
Quản lý toàn bộ hồ sơ nhân viên: Thêm nhân viên mới, cập nhật thông tin, tra cứu nhanh, in Hợp đồng lao động/Hợp đồng thử việc,
ra các Quyết định nhân sự (bổ nhiệm, điều chuyển, chấm dứt HĐLĐ...), xem cơ cấu nhân sự theo Phòng.

🎉 **Đặc biệt: Gửi lời chúc sinh nhật tự động** - hệ thống tự nhắc và hỗ trợ gửi lời chúc mừng sinh nhật đến từng
CBCNV. Đây là một chi tiết nhỏ nhưng có sức nặng lớn: nó giúp gắn kết giữa Ban điều hành với người lao động,
khiến nhân viên cảm thấy được quan tâm như một cá nhân chứ không chỉ là một con số trên bảng lương - góp phần
xây dựng văn hoá doanh nghiệp gắn bó, nhân văn.

### 📋 BHXH / 📋 Báo cáo định kỳ
Theo dõi tình hình đóng BHXH, tự tạo báo cáo tăng/giảm D02-LT, dự toán số tiền phải đóng theo kỳ — giảm tối đa
thao tác thủ công so với việc tự tổng hợp trên Excel. HRM Master sẽ luôn update các mẫu báo cáo theo kịp các văn bản quy định mới nhất về nghiệp vụ hành chính nhân sự.

### 🕒 Chấm công / 💰 Tính thu nhập
Quản lý chấm công theo ca, tự động tính lương, phụ cấp, các khoản khấu trừ theo đúng quy định hiện hành và sẽ tùy chỉnh đúng với chính sách mà Quý doanh nghiệp đang áp dụng.

### 📄 Quản lý Công văn & HĐ kinh tế
Lưu trữ, tra cứu công văn đến/đi và hợp đồng kinh tế tập trung, khoa hoạc - tránh thất lạc, dễ dàng tìm lại khi cần đối chiếu.

### ⏰ Báo cáo tự động & Nhắc hạn — không lo bị "miss" deadline
Hệ thống có các loại **báo cáo tự động** (tăng/giảm nhân sự, BHXH, hợp đồng...) giúp tiết kiệm thời gian tổng hợp
thủ công, đồng thời có **thông báo nhắc nhở các mốc quan trọng sắp đến hạn** (hết hạn HĐLĐ, hết hạn thử việc...),
giúp bộ phận Nhân sự chủ động xử lý trước hạn, tránh bỏ sót ảnh hưởng đến quyền lợi người lao động và rủi ro pháp lý
cho doanh nghiệp.

### 📊 Menu Báo cáo
Nơi tập trung liệt kê và chạy tất cả các loại báo cáo nhân sự sẵn có, xuất trực tiếp ra file để gửi cho Ban giám đốc,
cơ quan BHXH, hoặc lưu trữ nội bộ mà không cần thao tác qua nhiều màn hình.

### 🤖 Chatbot Giải đáp
Trợ lý AI trả lời nhanh các câu hỏi về BHXH, BHYT, thuế TNCN, thai sản, thất nghiệp, hợp đồng lao động — có trích dẫn
điều luật cụ thể, giúp CBCNV và bộ phận Nhân sự tự tra cứu quyền lợi mà không cần chờ hỏi trực tiếp.

### 💬 Chat nội bộ
Kênh trao đổi nội bộ ngay trong app — không cần chuyển qua ứng dụng nhắn tin khác.
""")

    st.info("💡 Có thắc mắc trong quá trình sử dụng, hãy dùng ngay mục **🤖 Chatbot Giải đáp** hoặc liên hệ bộ phận Nhân sự / IT để được hỗ trợ.")


#===== Hàm xử lý chính ===== 
def main():
    """Giữ tương thích với `if __name__ == '__main__': main()` ở cuối file.
    Landing Page đã bị bỏ (mỗi tenant có domain riêng, vào thẳng màn hình đăng nhập
    ở luồng phía trên) — hàm này không còn logic để chạy, chỉ giữ lại cho an toàn."""
    pass

def reset_ui_and_cache():
    """Reset toàn bộ cache và session state để refresh UI"""
    st.cache_data.clear()
    st.cache_resource.clear()
    
    # Giữ lại các session state quan trọng
    keep_keys = ['logged_in', 'role', 'username', 'language']
    for key in list(st.session_state.keys()):
        if key not in keep_keys:
            del st.session_state[key]
    
    st.rerun()

# Chạy ứng dụng
if __name__ == "__main__":
    main()