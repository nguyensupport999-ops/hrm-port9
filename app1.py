import streamlit as st
import psycopg2
import psycopg2.extras
from config import DB_CONFIG, COMPANY_CONFIG
from datetime import datetime, date, timedelta
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
import urllib.parse
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import requests
from openpyxl.styles import Font, Alignment
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit as st
from PIL import Image
import os
import qrcode
from io import BytesIO
from dotenv import load_dotenv
import os

load_dotenv()

# === FORCE DEBUG ===
import sys
print("=" * 50)
print("ENVIRONMENT VARIABLES AT START:")
print(f"DB_HOST: {os.getenv('DB_HOST')}")
print(f"DB_USER: {os.getenv('DB_USER')}")
print("=" * 50)
sys.stdout.flush()

# Chèn logo vào sidebar
logo_path = "logo_cty.png"
if os.path.exists(logo_path):
    with st.sidebar:
        st.image(logo_path, use_container_width=True)
        st.divider()

st.set_page_config(
    page_title="HRM-Port", 
    page_icon="🏗️", 
    layout="wide"
)

def force_center(p):
    pPr = p._p.get_or_add_pPr()
    for jc in pPr.findall(qn('w:jc')):
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

st.set_page_config(page_title="HRM-Port", page_icon="🏗️", layout="wide")

st.markdown("""
<style>
    [data-testid="stDataFrame"] > div {
        overflow-x: auto !important;
    }
    [data-testid="stDataFrame"] table {
        min-width: 2000px !important;
        width: max-content !important;
    }
</style>
""", unsafe_allow_html=True)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def get_connection():
    import psycopg2
    return psycopg2.connect(
        host="aws-1-ap-northeast-1.pooler.supabase.com",
        port=5432,
        user="postgres.ioesyihbsdxmxrotdetx",
        password="Xbr2w6bo1s5JY4Vq",  # Thay bằng mật khẩu thật
        database="postgres"
    )

def format_date(d):
    if d is None or pd.isna(d): return ''
    try: return d.strftime('%d/%m/%Y') if hasattr(d,'strftime') else str(d)[:10]
    except: return str(d)

def parse_date(s):
    if not s or s.strip()=='': return None
    try: return datetime.strptime(s.strip(),'%d/%m/%Y').date()
    except: return None

def tao_noi_dung_zalo(nv):
    ZC = COMPANY_CONFIG
    return f"""Gửi anh/chị: {nv.get('ho_ten','')},

Thông tin đã cập nhật:
- Họ tên: {nv.get('ho_ten','')}
- Ngày sinh: {format_date(nv.get('ngay_sinh'))}
- CCCD: {nv.get('so_cccd','')}
- Ngày cấp: {format_date(nv.get('Ngay_cap_CCCD'))}
- Thường trú: {nv.get('Thuong_tru','')}
- Số BHXH: {nv.get('ma_so_bhxh','')}
- TK NH: {nv.get('So_tai_khoan_NH','')}
- CN NH: {nv.get('Chi_nhanh_NH','')}

{ZC.get('loi_nhan_zalo','Vui lòng kiểm tra và phản hồi nếu có sai sót. Xin Cảm ơn!')}"""

def remove_table_border(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            b = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
            if b is not None: tcPr.remove(b)

# ============================================================
# HÀM TAO_HOP_DONG - GIỮ NGUYÊN KHÔNG SỬA
# ============================================================
def tao_hop_dong(nv):
    CC = COMPANY_CONFIG; doc = Document()
    s = doc.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
    s.paragraph_format.space_after=Pt(0); s.paragraph_format.space_before=Pt(0)
    sec = doc.sections[0]; sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(3.5); sec.right_margin=Cm(2)
    def add_p(text='', bold=False, size=Pt(13)):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if bold and p.runs:
            p.runs[0].bold = True
        if p.runs:
            p.runs[0].font.size = size
        return p
    def al(label,value):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(5))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r=p.add_run(f'{label}'); r.font.size=Pt(13)
        r=p.add_run('\t: '); r.font.size=Pt(13)
        r=p.add_run(f'{value}'); r.font.size=Pt(13)
    ht=doc.add_table(rows=4,cols=2); ht.alignment=WD_TABLE_ALIGNMENT.CENTER; ht.autofit=False; remove_table_border(ht)
    for row in ht.rows: row.cells[0].width=Cm(6); row.cells[1].width=Cm(11)
    c=ht.rows[0].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('CÔNG TY CỔ PHẦN'); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[0].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM'); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[1].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('CẢNG HÒN LA'); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[1].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Độc lập - Tự do - Hạnh phúc'); r.bold=True; r.italic=True; r.font.size=Pt(13)
    c=ht.rows[2].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('─'*12); r.font.size=Pt(9)
    c=ht.rows[2].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('─'*20); r.font.size=Pt(9)
    c=ht.rows[3].cells[1]; p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after=Pt(20)
    ngay_ky = nv.get("Ngay_ky_HD")
    ngay_vao = nv.get("Ngay_vao_lam")
    nk = ngay_ky if ngay_ky else ngay_vao
    
    ns = 'Quảng Trị, ngày ... tháng ... năm ......'
    if nk and hasattr(nk, 'day'):
        ns = f'Quảng Trị, ngày {nk.day} tháng {nk.month:02d} năm {nk.year}'
    run = p.add_run(ns)
    run.font.size = Pt(13)
    run.italic = True
    c=ht.rows[3].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'Số: {nv.get("So_HDLD","...")}'); r.italic=True; r.font.size=Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run('HỢP ĐỒNG LAO ĐỘNG')
    r.bold = True
    r.font.size = Pt(18)
    force_center(p)  
    p2 = doc.add_paragraph('- Căn cứ thông tư 10/2020/TT-LĐTBXH ngày 12/11/2020 hướng dẫn thi hành một số điều của Bộ luật Lao động số 45/2019/QH14 ngày 20/11/2019 về nội dung của hợp đồng lao động;')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2 = doc.add_paragraph('- Căn cứ nhu cầu sử dụng lao động trong đơn vị.')
    doc.add_paragraph('Chúng tôi gồm:')
    p=doc.add_paragraph(); r=p.add_run(f'BÊN A: {CC["ten_cong_ty"]} (Người sử dụng LĐ)'); r.bold=True
    al('Đại diện',f"Ông {CC['dai_dien']}"); al('Chức vụ',CC['chuc_vu']); al('Mã số thuế',CC['ma_so_thue'])
    al('Điện thoại',CC['dien_thoai_cty']); al('Địa chỉ',CC['dia_chi']); doc.add_paragraph()
    p=doc.add_paragraph(); r=p.add_run('BÊN B: (Người lao động)'); r.bold=True
    sk=nv.get('So_tai_khoan_NH','')
    if nv.get('Chi_nhanh_NH'): sk+=f' - {nv.get("Chi_nhanh_NH")}'
    gt = nv.get('gioi_tinh','')
    xung_ho = 'Ông' if gt == 'Nam' else ('Bà' if gt == 'Nữ' else 'Ông/Bà')
    al(xung_ho, nv.get('ho_ten',''))
    al('Ngày sinh',format_date(nv.get('ngay_sinh')))
    al('Số CMND/CCCD',nv.get('so_cccd','')); al('Ngày cấp',format_date(nv.get('Ngay_cap_CCCD')))
    al('Nơi cấp',nv.get('Noi_cap_CCCD','')); al('Số TKNH',sk)
    al('Điện thoại',nv.get('dien_thoai','')); al('Thường trú',nv.get('Thuong_tru',''))
    doc.add_paragraph('Thoả thuận ký kết Hợp đồng lao động với những điều khoản dưới đây:')
    p=doc.add_paragraph(); r=p.add_run('Điều 1. Thời hạn và công việc hợp đồng:'); r.bold=True
    ngay_hieu_luc = nv.get("Ngay_ky_HD") or nv.get("Ngay_vao_lam")
    ns2 = '.../.../..........'
    if ngay_hieu_luc and hasattr(ngay_hieu_luc, 'day'):
        ns2 = f'{ngay_hieu_luc.day} tháng {ngay_hieu_luc.month:02d} năm {ngay_hieu_luc.year}'
    elif ngay_hieu_luc:
        ns2 = str(ngay_hieu_luc)
    add_p(f'-    Bên B làm việc theo chế độ hợp đồng lao động không xác định thời hạn;')
    add_p(f'-    Thời gian: Từ ngày {ns2};')
    add_p('-    Địa điểm làm việc: Tại Cảng tổng hợp quốc tế Hòn La và các địa điểm khác theo sự sắp xếp của Công ty;')
    add_p(f'-    Vị trí: {nv.get("Chuc_danh_nghe","")};')
    add_p('-    Công việc phải làm: Thực hiện công việc theo đúng chuyên môn dưới sự quản lý, điều hành của cấp trên;')
    add_p('-    Mức lương và phụ cấp: Theo thỏa thuận;')
    add_p('-    Hình thức trả lương: Tiền mặt hoặc chuyển khoản, theo lần chi trả;')
    add_p('-    Kỳ hạn trả lương: Theo quy định Công ty;')
    add_p('-    Chế độ nâng lương: Theo thỏa thuận.')
    p=doc.add_paragraph(); r=p.add_run('Điều 2. Chế độ làm việc:'); r.bold=True
    add_p('-    Thời gian làm việc: Theo tính chất công việc, do nhu cầu kinh doanh của Công ty nên thời gian làm việc của bên B là linh hoạt nhưng phải đảm bảo hoàn thành công việc được giao;')
    add_p('-    Thời gian nghỉ ngơi của người lao động: Theo thỏa thuận và phù hợp với quy định của pháp luật;')
    add_p('-    Ngoài giờ làm việc: Người lao động phải tự chịu trách nhiệm về các hoạt động cá nhân của mình.')
    p=doc.add_paragraph(); r=p.add_run('Điều 3. Nghĩa vụ, quyền lợi NLĐ:'); r.bold=True
    p=doc.add_paragraph(); r=p.add_run('1. Nghĩa vụ:'); r.bold=True
    add_p('-    Hoàn thành những công việc được giao và sẵn sàng chấp nhận mọi sự điều động khi có yêu cầu;')
    add_p('-    Chấp hành nghiêm túc nội quy, kỷ luật lao động, an toàn lao động và các quy định của Công ty và pháp luật của Nhà nước;')
    add_p('-    Người lao động có trách nhiệm tuân thủ đầy đủ quy định về an toàn lao động, quy trình vận hành thiết bị và hướng dẫn của Công ty. Trường hợp NLĐ cố ý vi phạm hoặc vi phạm nghiêm trọng quy định an toàn lao động gây thiệt hại thì phải chịu trách nhiệm theo quy định pháp luật và nội quy Công ty;')
    add_p('-    Bồi thường vi phạm vật chất : Phải bồi thường vật chất do cá nhân vi phạm quy định của Công ty về bảo quản trang thiết bị được giao.')
    p=doc.add_paragraph(); r=p.add_run('2. Quyền Lợi:'); r.bold=True
    add_p('-    Phương tiện đi lại: Tự túc;')
    add_p('-    Được Công ty đóng Bảo hiểm xã hội, bảo hiểm y tế, BHTN: theo chế độ hiện hành của Nhà nước và Quy định của Công ty;')
    add_p('-    Được Công ty cấp đầy đủ bảo hộ lao động theo đúng vị trí làm việc;')
    add_p('-    Được phân công công việc theo yêu cầu của Công ty phù hợp với khả năng và trình độ chuyên môn mà người lao động đáp ứng;')
    add_p('-    Các quyền lợi khác thực hiện theo quy định của Pháp luật Lao động như tạm dừng, chấm dứt hợp đồng.')
    p=doc.add_paragraph(); r=p.add_run('Điều 4. Nghĩa vụ, quyền hạn NSDLĐ:'); r.bold=True
    add_p('-    Bảo đảm việc làm và thực hiện đầy đủ những điều đã cam kết trong hợp đồng;')
    add_p('-    Thanh toán đầy đủ, đúng hạn các chế độ và quyền lợi cho người lao động theo hợp đồng;')
    add_p('-    Điều hành người lao động hoàn thành công việc theo hợp đồng;')
    add_p('-    Tạm hoãn, chấm dứt hợp đồng, kỷ luật người lao động theo quy định của pháp luật, và nội quy lao động của Công ty.')
    p=doc.add_paragraph(); r=p.add_run('Điều 5. Điều khoản chung:'); r.bold=True
    add_p('-    Những nội dung về quan hệ lao động không ghi trong hợp đồng này thì được áp dụng theo pháp luật lao động;')
    add_p('-    Những thoả thuận khác (nếu có): không;')
    add_p('-    Hợp đồng này có hiệu lực từ ngày ký và được làm thành 02 bản, Bên A giữ 01 bản, Bên B giữ 01 có giá trị pháp lý như nhau, để làm căn cứ thực hiện.')
    add_p('Bản HĐ này lập tại văn phòng Công ty CP Cảng Hòn La.'); doc.add_paragraph()
    ts=doc.add_table(rows=3,cols=2); ts.alignment=WD_TABLE_ALIGNMENT.CENTER; remove_table_border(ts)
    c=ts.rows[0].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run('NGƯỜI LAO ĐỘNG'); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[0].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run('NGƯỜI SỬ DỤNG LAO ĐỘNG'); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[1].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].add_run('').font.size=Pt(12); sp=c.add_paragraph(); sp.paragraph_format.space_after=Pt(60)
    c=ts.rows[1].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].add_run('').font.size=Pt(12); sp=c.add_paragraph(); sp.paragraph_format.space_after=Pt(60)
    c=ts.rows[2].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run(nv.get('ho_ten','').upper()); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[2].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run(CC['dai_dien'].upper()); r.bold=True; r.font.size=Pt(13)
    tf=tempfile.NamedTemporaryFile(delete=False,suffix='.docx'); doc.save(tf.name); return tf.name

# ============================================================
# HÀM TAO_HOP_DONG_THU_VIEC - GIỮ NGUYÊN KHÔNG SỬA
# ============================================================
def tao_hop_dong_thu_viec(nv):
    CC = COMPANY_CONFIG; doc = Document()
    s = doc.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
    s.paragraph_format.space_after=Pt(0); s.paragraph_format.space_before=Pt(0)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sec = doc.sections[0]; sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(3.5); sec.right_margin=Cm(2)
    def al(label,value):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(5))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r=p.add_run(f'{label}'); r.font.size=Pt(13)
        r=p.add_run('\t: '); r.font.size=Pt(13)
        r=p.add_run(f'{value}'); r.font.size=Pt(13)
    ht=doc.add_table(rows=4,cols=2); ht.alignment=WD_TABLE_ALIGNMENT.CENTER; ht.autofit=False; remove_table_border(ht)
    for row in ht.rows: row.cells[0].width=Cm(6); row.cells[1].width=Cm(11)
    c=ht.rows[0].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('CÔNG TY CỔ PHẦN'); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[0].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM'); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[1].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('CẢNG HÒN LA'); r.bold=True; r.font.size=Pt(13)
    c=ht.rows[1].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Độc lập - Tự do - Hạnh phúc'); r.bold=True; r.italic=True; r.font.size=Pt(13)
    c=ht.rows[2].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('─'*12); r.font.size=Pt(9)
    c=ht.rows[2].cells[1]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('─'*20); r.font.size=Pt(9)
    c=ht.rows[3].cells[0]; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'Số: {nv.get("So_HDLD",".../.../HĐTV-CHL")}'); r.italic=True; r.font.size=Pt(12)
    c=ht.rows[3].cells[1]; p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after=Pt(20)
    nk=nv.get("Ngay_vao_lam") or nv.get("Ngay_ky_HD")
    ns='Quảng Trị, ngày ... tháng ... năm ......'
    if nk and hasattr(nk,'day'): ns=f'Quảng Trị, ngày {nk.day} tháng {nk.month:02d} năm {nk.year}'
    run = p.add_run(ns)
    run.font.size = Pt(13)
    run.italic = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run('HỢP ĐỒNG THỬ VIỆC')
    r.bold = True
    r.font.size = Pt(18)
    force_center(p)
    p2 = doc.add_paragraph('- Căn cứ thông tư 10/2020/TT-LĐTBXH ngày 12/11/2020 hướng dẫn thi hành một số điều của Bộ luật Lao động số 45/2019/QH14 ngày 20/11/2019 về nội dung của hợp đồng lao động;')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph('- Căn cứ nhu cầu sử dụng lao động trong đơn vị.')
    doc.add_paragraph('Chúng tôi gồm:')
    p=doc.add_paragraph(); r=p.add_run(f'BÊN A: {CC["ten_cong_ty"]} (Người sử dụng LĐ)'); r.bold=True
    al('Đại diện',f"Ông {CC['dai_dien']}"); al('Chức vụ',CC['chuc_vu']); al('Mã số thuế',CC['ma_so_thue'])
    al('Điện thoại',CC['dien_thoai_cty']); al('Địa chỉ',CC['dia_chi'])
    p=doc.add_paragraph(); r=p.add_run('BÊN B: (Người lao động)'); r.bold=True
    sk=nv.get('So_tai_khoan_NH','')
    if nv.get('Chi_nhanh_NH'): sk+=f' - {nv.get("Chi_nhanh_NH")}'
    gt = nv.get('gioi_tinh','')
    xung_ho = 'Ông' if gt == 'Nam' else ('Bà' if gt == 'Nữ' else 'Ông/Bà')
    al(xung_ho, nv.get('ho_ten',''))
    al('Ngày sinh',format_date(nv.get('ngay_sinh')))
    al('Số CMND/CCCD',nv.get('so_cccd','')); al('Ngày cấp',format_date(nv.get('Ngay_cap_CCCD')))
    al('Nơi cấp',nv.get('Noi_cap_CCCD','')); al('Số TKNH',sk)
    al('Điện thoại',nv.get('dien_thoai','')); al('Thường trú',nv.get('Thuong_tru',''))
    doc.add_paragraph('Thoả thuận ký kết Hợp đồng Thử việc với những điều khoản dưới đây:')
    p=doc.add_paragraph(); r=p.add_run('Điều 1. Thời hạn và công việc hợp đồng:'); r.bold=True
    ns2='.../.../..........'
    if nk and hasattr(nk,'day'): ns2=f'{nk.day} tháng {nk.month} năm {nk.year}'
    elif nk: ns2=str(nk)
    p=doc.add_paragraph(f'-    Bên B làm việc theo chế độ hợp đồng thử việc, có thời hạn 01 tháng;')
    if nk and hasattr(nk,'day'):
        nkt=nk+timedelta(days=30)
        p=doc.add_paragraph(f'     + Bắt đầu: {nk.day:02d}/{nk.month:02d}/{nk.year}')
        p=doc.add_paragraph(f'     + Kết thúc: {nkt.day:02d}/{nkt.month:02d}/{nkt.year}')
    else: doc.add_paragraph('     + Bắt đầu: .../.../......'); doc.add_paragraph('     + Kết thúc: .../.../......')
    p=doc.add_paragraph('-    Địa điểm làm việc: Tại Cảng tổng hợp quốc tế Hòn La và các địa điểm khác theo sự sắp xếp của Công ty;')
    p=doc.add_paragraph(f'-    Vị trí: {nv.get("Chuc_danh_nghe","")};')
    p=doc.add_paragraph('-    Công việc phải làm: Thực hiện công việc theo đúng chuyên môn dưới sự quản lý, điều hành của cấp trên;')
    p=doc.add_paragraph('-    Mức lương và phụ cấp: Theo thỏa thuận;')
    p=doc.add_paragraph('-    Hình thức trả lương: Tiền mặt hoặc chuyển khoản, theo lần chi trả;')
    p=doc.add_paragraph('-    Kỳ hạn trả lương: Theo quy định Công ty;')
    p=doc.add_paragraph(); r=p.add_run('Điều 2. Chế độ làm việc:'); r.bold=True
    p=doc.add_paragraph('-    Thời gian làm việc: Theo tính chất công việc, do nhu cầu kinh doanh của Công ty nên thời gian làm việc của bên B là linh hoạt nhưng phải đảm bảo hoàn thành công việc được giao;')
    p=doc.add_paragraph('-    Thời gian nghỉ ngơi của người lao động: Theo thỏa thuận và phù hợp với quy định của pháp luật;')
    p=doc.add_paragraph('-    Ngoài giờ làm việc: Người lao động phải tự chịu trách nhiệm về các hoạt động cá nhân của mình.')
    p=doc.add_paragraph(); r=p.add_run('Điều 3. Nghĩa vụ, quyền lợi NLĐ:'); r.bold=True
    p=doc.add_paragraph(); r=p.add_run('1. Nghĩa vụ:'); r.bold=True
    p=doc.add_paragraph('-    Hoàn thành những công việc được giao và sẵn sàng chấp nhận mọi sự điều động khi có yêu cầu;')
    p=doc.add_paragraph('-    Chấp hành nghiêm túc nội quy, kỷ luật lao động, an toàn lao động và các quy định của Công ty và pháp luật của Nhà nước;')
    p=doc.add_paragraph('-    Người lao động có trách nhiệm tuân thủ đầy đủ quy định về an toàn lao động, quy trình vận hành thiết bị và hướng dẫn của Công ty. Trường hợp NLĐ cố ý vi phạm hoặc vi phạm nghiêm trọng quy định an toàn lao động gây thiệt hại thì phải chịu trách nhiệm theo quy định pháp luật và nội quy Công ty;')
    p=doc.add_paragraph('-    Bồi thường vi phạm vật chất : Phải bồi thường vật chất do cá nhân vi phạm quy định của Công ty về bảo quản trang thiết bị được giao.')
    p=doc.add_paragraph(); r=p.add_run('2. Quyền Lợi:'); r.bold=True
    p=doc.add_paragraph('-    Phương tiện đi lại: Tự túc;')
    p=doc.add_paragraph('-    Được Công ty cấp đầy đủ bảo hộ lao động theo đúng vị trí làm việc;')
    p=doc.add_paragraph('-    Được phân công công việc theo yêu cầu của Công ty phù hợp với khả năng và trình độ chuyên môn mà người lao động đáp ứng;')
    p=doc.add_paragraph('-    Các quyền lợi khác thực hiện theo quy định của Pháp luật Lao động như tạm dừng, chấm dứt hợp đồng.')
    p=doc.add_paragraph(); r=p.add_run('Điều 4. Nghĩa vụ, quyền hạn NSDLĐ:'); r.bold=True
    p=doc.add_paragraph('-    Bảo đảm việc làm và thực hiện đầy đủ những điều đã cam kết trong hợp đồng;')
    p=doc.add_paragraph('-    Thanh toán đầy đủ, đúng hạn các chế độ và quyền lợi cho người lao động theo hợp đồng;')
    p=doc.add_paragraph('-    Điều hành người lao động hoàn thành công việc theo hợp đồng;')
    p=doc.add_paragraph('-    Tạm hoãn, chấm dứt hợp đồng theo quy định của pháp luật, và nội quy lao động của Công ty;')
    p=doc.add_paragraph(); r=p.add_run('Điều 5. Điều khoản chung:'); r.bold=True
    p=doc.add_paragraph('-    Những nội dung về quan hệ lao động không ghi trong hợp đồng này thì được áp dụng theo pháp luật lao động;')
    p=doc.add_paragraph('-    Những thoả thuận khác (nếu có): không;')
    p=doc.add_paragraph('-    Hợp đồng này có hiệu lực từ ngày ký và được làm thành 02 bản, Bên A giữ 01 bản, Bên B giữ 01 có giá trị pháp lý như nhau, để làm căn cứ thực hiện.'); doc.add_paragraph()
    p=doc.add_paragraph('Bản HĐ này lập tại văn phòng Công ty CP Cảng Hòn La.'); doc.add_paragraph()
    ts=doc.add_table(rows=3,cols=2); ts.alignment=WD_TABLE_ALIGNMENT.CENTER; remove_table_border(ts)
    c=ts.rows[0].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run('NGƯỜI LAO ĐỘNG'); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[0].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER    
    r=c.paragraphs[0].add_run('NGƯỜI SỬ DỤNG LAO ĐỘNG'); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[1].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].add_run('').font.size=Pt(12); sp=c.add_paragraph(); sp.paragraph_format.space_after=Pt(60)
    c=ts.rows[1].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c.paragraphs[0].add_run('').font.size=Pt(12); sp=c.add_paragraph(); sp.paragraph_format.space_after=Pt(60)
    c=ts.rows[2].cells[0]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run(nv.get('ho_ten','').upper()); r.bold=True; r.font.size=Pt(13)
    c=ts.rows[2].cells[1]; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.paragraphs[0].add_run(CC['dai_dien'].upper()); r.bold=True; r.font.size=Pt(13)
    tf=tempfile.NamedTemporaryFile(delete=False,suffix='.docx'); doc.save(tf.name); return tf.name

def tao_bao_cao_tang_giam(ds_tang, ds_giam, tu_ngay, den_ngay):
    """Tạo file Word báo cáo tăng/giảm nhân sự theo mẫu"""
    from docx.shared import Inches
    from docx.enum.section import WD_ORIENTATION
    
    CC = COMPANY_CONFIG
    doc = Document()
    
    # ===== CẤU HÌNH KHỔ GIẤY A4 NẰM NGANG =====
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    # Hàm format_date an toàn hơn
    def safe_format_date(value):
        if value is None or value == '':
            return ''
        try:
            if hasattr(value, 'strftime'):
                return value.strftime('%d/%m/%Y')
            if isinstance(value, str):
                # Nếu đã là string, thử parse
                from datetime import datetime
                for fmt in ['%Y-%m-%d', '%d/%m/%Y', '%Y%m%d']:
                    try:
                        d = datetime.strptime(value, fmt).date()
                        return d.strftime('%d/%m/%Y')
                    except:
                        continue
                return value[:10] if len(value) >= 10 else value
            return str(value)
        except:
            return ''
    
    # ===== HEADER =====
    header_table = doc.add_table(rows=3, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_border(header_table)
    
    cell_left = header_table.rows[0].cells[0]
    cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_left.paragraphs[0].add_run(CC['ten_cong_ty'])
    run.bold = True
    run.font.size = Pt(12)
    
    cell_right = header_table.rows[0].cells[1]
    cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_right.paragraphs[0].add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
    run.bold = True
    run.font.size = Pt(12)
    
    cell_left2 = header_table.rows[1].cells[0]
    cell_left2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_left2.paragraphs[0].add_run('Phòng Hành chính Nhân sự')
    run.font.size = Pt(11)
    
    cell_right2 = header_table.rows[1].cells[1]
    cell_right2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_right2.paragraphs[0].add_run('Độc lập - Tự do - Hạnh phúc')
    run.font.size = Pt(11)
    run.bold = True
    
    cell_left3 = header_table.rows[2].cells[0]
    cell_left3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_left3.paragraphs[0].add_run('─' * 12)
    run.font.size = Pt(9)
    
    cell_right3 = header_table.rows[2].cells[1]
    cell_right3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_right3.paragraphs[0].add_run('─' * 20)
    run.font.size = Pt(9)
    
    # ===== TIÊU ĐỀ BÁO CÁO =====
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BÁO CÁO TĂNG/GIẢM NHÂN SỰ')
    run.bold = True
    run.font.size = Pt(15)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'(Từ ngày {tu_ngay.strftime("%d/%m/%Y")} đến ngày {den_ngay.strftime("%d/%m/%Y")})')
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
      
        # ===== BẢNG CHÍNH =====
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Header
    headers = ['STT', 'Họ và tên', 'Ngày sinh', 'Số HĐ', 'Ngày ký HĐ', 'Chức danh', 'Tình trạng']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    # ===== PHÂN LOẠI VÀ GHÉP DỮ LIỆU =====
    all_items = []
    
    # Thêm nhân viên tăng
    for nv in ds_tang:
        loai_hd = nv.get('loai_hop_dong', '')
        if loai_hd in ['Không xác định thời hạn', 'Xác định thời hạn']:
            nhom = 'Hợp đồng không xác định thời hạn'
        else:
            nhom = 'Hợp đồng thử việc'
        
        ngay_ky = nv.get('Ngay_ky_HD')
        if ngay_ky is None or ngay_ky == '':
            ngay_ky = nv.get('ngay_vao_lam')
        
        all_items.append({
            'nhom': nhom,
            'tinh_trang': 'Tăng',
            'ho_ten': nv.get('ho_ten', ''),
            'ngay_sinh': nv.get('ngay_sinh'),
            'so_hd': nv.get('so_hdld', ''),
            'ngay_ky': ngay_ky,
            'chuc_danh': nv.get('chuc_danh_nghe', '')
        })
    
    # Thêm nhân viên giảm
    for nv in ds_giam:
        loai_hd = nv.get('loai_hop_dong', '')
        if loai_hd in ['Không xác định thời hạn', 'Xác định thời hạn']:
            nhom = 'Hợp đồng không xác định thời hạn'
        else:
            nhom = 'Hợp đồng thử việc'
        
        ngay_ky = nv.get('Ngay_ky_HD')
        if ngay_ky is None or ngay_ky == '':
            ngay_ky = nv.get('ngay_vao_lam')
        
        all_items.append({
            'nhom': nhom,
            'tinh_trang': 'Giảm',
            'ho_ten': nv.get('ho_ten', ''),
            'ngay_sinh': nv.get('ngay_sinh'),
            'so_hd': nv.get('so_hdld', ''),
            'ngay_ky': ngay_ky,
            'chuc_danh': nv.get('chuc_danh_nghe', '')
        })
    
    # Hàm merge cells
    def merge_row_cells(row):
        if len(row.cells) > 1:
            first_cell = row.cells[0]
            for cell in row.cells[1:]:
                first_cell.merge(cell)
    
    # Xử lý theo từng nhóm
    if all_items:
        # Nhóm HĐKXĐTH trước
        nhom_kxdt = [item for item in all_items if item['nhom'] == 'Hợp đồng không xác định thời hạn']
        nhom_tv = [item for item in all_items if item['nhom'] == 'Hợp đồng thử việc']
        
        # Xử lý nhóm HĐKXĐTH
        if nhom_kxdt:
            row_title = table.add_row()
            row_title.cells[0].text = 'Hợp đồng không xác định thời hạn'
            for run in row_title.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
            row_title.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            merge_row_cells(row_title)
            
            for idx, item in enumerate(nhom_kxdt, 1):
                row = table.add_row()
                row.cells[0].text = str(idx)
                row.cells[1].text = item['ho_ten']
                row.cells[2].text = safe_format_date(item['ngay_sinh'])
                row.cells[3].text = item['so_hd']
                row.cells[4].text = safe_format_date(item['ngay_ky'])
                row.cells[5].text = item['chuc_danh']
                row.cells[6].text = item['tinh_trang']
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in p.runs:
                            run.font.size = Pt(9)
                row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Xử lý nhóm HĐ thử việc
        if nhom_tv:
            row_title = table.add_row()
            row_title.cells[0].text = 'Hợp đồng thử việc'
            for run in row_title.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
            row_title.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            merge_row_cells(row_title)
            
            for idx, item in enumerate(nhom_tv, 1):
                row = table.add_row()
                row.cells[0].text = str(idx)
                row.cells[1].text = item['ho_ten']
                row.cells[2].text = safe_format_date(item['ngay_sinh'])
                row.cells[3].text = item['so_hd']
                row.cells[4].text = safe_format_date(item['ngay_ky'])
                row.cells[5].text = item['chuc_danh']
                row.cells[6].text = item['tinh_trang']
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in p.runs:
                            run.font.size = Pt(9)
                row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Không có dữ liệu
        row_empty = table.add_row()
        for cell in row_empty.cells:
            cell.text = ''
    
    # ===== SET LẠI ĐỘ RỘNG CỘT SAU KHI ĐÃ THÊM TẤT CẢ CÁC DÒNG =====
    # Quan trọng: Phải set lại độ rộng sau khi đã merge cells
    col_widths = [Inches(0.25), Inches(2.2), Inches(0.8), Inches(1.5), Inches(1.0), Inches(2.0), Inches(0.8)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    
    # Ép cột STT nhỏ hơn nữa
    table.columns[0].width = Inches(0.25)
    
    # Canh giữa cột STT và Tình trạng cho tất cả các dòng
    for row in table.rows:
        if len(row.cells) == 7:  # Dòng dữ liệu bình thường
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif len(row.cells) == 1:  # Dòng đã merge (tiêu đề nhóm)
            # Set lại width cho ô đã merge bằng tổng độ rộng các cột
            row.cells[0].width = sum(col_widths)
    
    # ===== CHỮ KÝ =====
    doc.add_paragraph()
    
    # Bảng 3 cột để căn chỉnh ngày tháng bên phải
    sign_intro = doc.add_table(rows=1, cols=3)
    remove_table_border(sign_intro)
    sign_intro.columns[0].width = Inches(2.5)
    sign_intro.columns[1].width = Inches(2)
    sign_intro.columns[2].width = Inches(3)
    
    # Ô bên trái để trống
    cell_date_left = sign_intro.rows[0].cells[0]
    cell_date_left.text = ''
    
    # Ô giữa để trống (tạo khoảng cách)
    cell_date_middle = sign_intro.rows[0].cells[1]
    cell_date_middle.text = ''
    
    # Ô bên phải chứa ngày tháng
    cell_date_right = sign_intro.rows[0].cells[2]
    cell_date_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ngay_hien_tai = date.today()
    run = cell_date_right.paragraphs[0].add_run(f'Quảng Trị, ngày {ngay_hien_tai.day} tháng {ngay_hien_tai.month} năm {ngay_hien_tai.year}')
    run.font.size = Pt(11)
    run.italic = True
    
    # Bảng chữ ký - 4 dòng 3 cột
    sign_table = doc.add_table(rows=5, cols=3)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_border(sign_table)
    
    sign_table.columns[0].width = Inches(2.5)
    sign_table.columns[1].width = Inches(1.5)
    sign_table.columns[2].width = Inches(3)
    
    # Dòng 1: Chức danh (Người lập / GIÁM ĐỐC) - cột giữa để trống
    cell_left_title = sign_table.rows[0].cells[0]
    cell_left_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_left_title.paragraphs[0].add_run('Người lập')
    run.font.size = Pt(11)
    
    sign_table.rows[0].cells[1].paragraphs[0].text = ''  # Cột giữa để trống
    
    cell_right_title = sign_table.rows[0].cells[2]
    cell_right_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_right_title.paragraphs[0].add_run(CC.get('chuc_vu', 'GIÁM ĐỐC'))
    run.font.size = Pt(11)
    run.bold = True
    
    # Dòng 2: trống (khoảng cách để ký)
    sign_table.rows[1].cells[0].paragraphs[0].text = ''
    sign_table.rows[1].cells[1].paragraphs[0].text = ''
    sign_table.rows[1].cells[2].paragraphs[0].text = ''
    sign_table.rows[1].cells[0].paragraphs[0].paragraph_format.space_after = Pt(30)
    sign_table.rows[1].cells[1].paragraphs[0].paragraph_format.space_after = Pt(30)
    sign_table.rows[1].cells[2].paragraphs[0].paragraph_format.space_after = Pt(30)
 
    
    # Dòng 3: Họ tên (Nguyễn Văn A / Tên Giám đốc) - cột giữa để trống
    cell_left_name = sign_table.rows[2].cells[0]
    cell_left_name.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_left_name.paragraphs[0].add_run('Nguyễn Văn Tuyến')
    run.font.size = Pt(11)
    
    sign_table.rows[2].cells[1].paragraphs[0].text = ''  # Cột giữa để trống
    
    cell_right_name = sign_table.rows[2].cells[2]
    cell_right_name.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_right_name.paragraphs[0].add_run(CC.get('dai_dien', 'GIÁM ĐỐC').upper())
    run.font.size = Pt(11)
    run.bold = True
    
    # Dòng 4: trống (tạo khoảng cách cuối trang)
    sign_table.rows[3].cells[0].paragraphs[0].text = ''
    sign_table.rows[3].cells[1].paragraphs[0].text = ''
    sign_table.rows[3].cells[2].paragraphs[0].text = ''
    sign_table.rows[3].cells[0].paragraphs[0].paragraph_format.space_after = Pt(10)
    sign_table.rows[3].cells[1].paragraphs[0].paragraph_format.space_after = Pt(10)
    sign_table.rows[3].cells[2].paragraphs[0].paragraph_format.space_after = Pt(10)
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name

def gui_email(loai,ds,file=None):
    from config import EMAIL_CONFIG as EC
    try:
        msg=MIMEMultipart(); msg['From']=EC['email']; msg['To']=EC['nguoi_nhan']
        tn=f"{datetime.now().month:02d}/{datetime.now().year}"
        msg['Subject']=f"[HRM-Port] Báo cáo {loai} lao động tháng {tn}"
        nd=f"<h3>BÁO CÁO {loai.upper()} LĐ</h3><p>Tháng: <b>{tn}</b></p><p>SL: <b>{len(ds)}</b></p><hr><ul>"
        for nv in ds[:10]: nd+=f"<li>{nv.get('ho_ten','')} - {nv.get('chuc_danh_nghe','')}</li>"
        if len(ds)>10: nd+=f"<li>... và {len(ds)-10} người khác</li>"
        nd+="</ul><p><b>File Excel đính kèm.</b></p>"
        msg.attach(MIMEText(nd,'html','utf-8'))
        if file and os.path.exists(file):
            with open(file,'rb') as f:
                part=MIMEBase('application','octet-stream'); part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition',f'attachment; filename="{os.path.basename(file)}"')
                msg.attach(part)
        srv=smtplib.SMTP(EC['smtp_server'],EC['smtp_port']); srv.starttls()
        srv.login(EC['email'],EC['password']); srv.send_message(msg); srv.quit()
        return True
    except Exception as e: st.error(f"Lỗi email: {e}"); return False

def gui_telegram(msg):
    from config import TELEGRAM_CONFIG as TC
    try:
        url=f"https://api.telegram.org/bot{TC['bot_token']}/sendMessage"
        r=requests.post(url,data={"chat_id":TC['chat_id'],"text":msg,"parse_mode":"HTML"},timeout=10)
        return r.status_code==200
    except: return False

# ========== SIDEBAR + LOGIN ==========
st.sidebar.title("🏗️ HRM-Port")
st.sidebar.caption("Quản lý nhân sự cảng biển")
if 'logged_in' not in st.session_state: st.session_state.logged_in=False; st.session_state.role=None; st.session_state.username=None
if not st.session_state.logged_in:
    st.sidebar.subheader("🔐 Đăng nhập")
    u=st.sidebar.text_input("Tài khoản"); p=st.sidebar.text_input("Mật khẩu",type="password")
    c1,c2=st.sidebar.columns(2)
    with c1:
        if st.button("Đăng nhập",use_container_width=True):
            from config import USERS
            if u in USERS and USERS[u]["password"]==p: st.session_state.logged_in=True; st.session_state.role=USERS[u]["role"]; st.session_state.username=u; st.rerun()
            else: st.sidebar.error("❌ Sai!")
    with c2:
        if st.button("👁️ Xem",use_container_width=True): st.session_state.logged_in=True; st.session_state.role="viewer"; st.session_state.username="guest"; st.rerun()
    st.stop()
menu_options = ["📊 Dashboard","👤 Ứng viên","✅ Nhân viên","📁 Upload hồ sơ","⚙️ Danh mục","📋 BHXH","📋 Báo cáo 01/PLI"] if st.session_state.role=="admin" else ["📊 Dashboard","✅ Nhân viên"]
menu=st.sidebar.radio("📋 Menu",menu_options)
st.sidebar.divider(); st.sidebar.caption(f"👤 {st.session_state.username} ({st.session_state.role})")
if st.sidebar.button("🚪 Đăng xuất",use_container_width=True): st.session_state.logged_in=False; st.session_state.role=None; st.session_state.username=None; st.rerun()

# ========== DASHBOARD ==========
if menu=="📊 Dashboard":
    st.title("📊 Dashboard"); db=get_connection(); c=db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    c.execute("SELECT COUNT(*) t FROM ung_vien"); tuv=c.fetchone()['t']
    c.execute("SELECT COUNT(*) t FROM nhan_vien WHERE Trang_thai IN ('DANG_LAM','THU_VIEC')"); tnv=c.fetchone()['t']
    c.execute("SELECT COUNT(*) t FROM ung_vien WHERE Trang_thai='CHO_DUYET'"); cd=c.fetchone()['t']
    c.execute("SELECT COUNT(*) t FROM ung_vien WHERE Trang_thai='TU_CHOI'"); tc=c.fetchone()['t']
    c.execute("SELECT COUNT(*) t FROM ung_vien WHERE Trang_thai='DA_NHAN_VIEC'"); dn=c.fetchone()['t']
    cl1,cl2,cl3,cl4,cl5=st.columns(5)
    cl1.metric("Tổng UV",tuv); cl2.metric("Nhân viên",tnv); cl3.metric("Chờ duyệt",cd); cl4.metric("Đã nhận",dn); cl5.metric("Từ chối",tc)
    st.divider(); st.subheader("📌 Thông báo")
    c.execute("SELECT Ho_ten FROM nhan_vien WHERE DATE(Ngay_vao_lam)=CURRENT_DATE"); hn=c.fetchall()
    c.execute("SELECT Ho_ten FROM nhan_vien WHERE DATE(Ngay_vao_lam)=CURRENT_DATE - INTERVAL '1 day'"); hq=c.fetchall()
    if hn: st.success(f"🟢 Hôm nay có thêm: **{', '.join([x['ho_ten'] for x in hn])}**")
    if hq: st.info(f"🔵 Hôm qua có thêm: **{', '.join([x['ho_ten'] for x in hq])}**")
    if st.session_state.role == "admin":
        c.execute("""SELECT STT, Ma_NV, Ho_ten, Ngay_vao_lam, 
            EXTRACT(DAY FROM AGE(Ngay_vao_lam + INTERVAL '30 days', CURRENT_DATE)) as ngay_con_lai
            FROM nhan_vien 
            WHERE Trang_thai = 'THU_VIEC' 
            AND EXTRACT(DAY FROM AGE(Ngay_vao_lam + INTERVAL '30 days', CURRENT_DATE)) <= 5
            AND EXTRACT(DAY FROM AGE(Ngay_vao_lam + INTERVAL '30 days', CURRENT_DATE)) >= 0 
            AND EXTRACT(DAY FROM AGE(CURRENT_DATE, Ngay_vao_lam)) >= 25 
            ORDER BY ngay_con_lai ASC""")
        tv_sap_het = c.fetchall()
        for x in tv_sap_het:
            if x['ngay_con_lai'] == 0: st.error(f"⚠️ **{x.get('ma_nv','')} {x['ho_ten']}** - HÔM NAY LÀ NGÀY CUỐI HỢP ĐỒNG THỬ VIỆC!")
            else: st.warning(f"⚠️ **{x.get('ma_nv','')} {x['ho_ten']}** còn **{x['ngay_con_lai']}** ngày sẽ kết thúc hợp đồng thử việc!")
#    c.execute("""
 #       SELECT STT, Ma_NV, Ho_ten, Ngay_sinh, Dien_thoai, Gioi_tinh,
 #       CASE 
 #           WHEN DATE(EXTRACT(YEAR FROM CURRENT_DATE) || '-' || EXTRACT(MONTH FROM Ngay_sinh) || '-' || EXTRACT(DAY FROM Ngay_sinh)) >= CURRENT_DATE
 #           THEN DATE(EXTRACT(YEAR FROM CURRENT_DATE) || '-' || EXTRACT(MONTH FROM Ngay_sinh) || '-' || EXTRACT(DAY FROM Ngay_sinh)) - CURRENT_DATE
 #           ELSE DATE(EXTRACT(YEAR FROM CURRENT_DATE) + 1 || '-' || EXTRACT(MONTH FROM Ngay_sinh) || '-' || EXTRACT(DAY FROM Ngay_sinh)) - CURRENT_DATE
 #       END as ngay_den_sn
 #       FROM nhan_vien 
  #      WHERE Trang_thai IN ('DANG_LAM','THU_VIEC') AND Ngay_sinh IS NOT NULL
  #      ORDER BY ngay_den_sn ASC
  #  """)
  #  sn_list = c.fetchall()
  #  sn_list = c.fetchall()
  #  for x in sn_list:
  #      today = date.today(); sn_date = x['ngay_sinh']
  #      this_year_sn = date(today.year, sn_date.month, sn_date.day)
  #      if this_year_sn < today: this_year_sn = date(today.year + 1, sn_date.month, sn_date.day)
  #      days_left = (this_year_sn - today).days
  #      ma_nv = x.get('ma_nv', '') or ''; ho_ten = x.get('ho_ten', '')
  #      gioi_tinh = x.get('gioi_tinh', '')
   #     if gioi_tinh == 'Nam': xung_ho = 'Anh'
   #     elif gioi_tinh == 'Nữ': xung_ho = 'Chị'
   #     else: xung_ho = 'Anh/Chị'
   #     if days_left == 0:
    #        st.balloons(); st.success(f"🎂🎉 **CHÚC MỪNG SINH NHẬT {ma_nv} {ho_ten}** 🎉🎂")
   #         if st.session_state.role == "admin":
   #             phone = x.get('dien_thoai', '')
   #             if phone:
   #                 phone = phone.replace('+84', '0').replace(' ', '').strip()
   #                 loi_chuc = f"""🎂🎉 CÔNG TY CỔ PHẦN CẢNG HÒN LA\n\nThân gửi {xung_ho}: {ho_ten}\n\nNhân dịp sinh nhật của {xung_ho}, Ban Lãnh đạo cùng toàn thể CB-CNV Công ty Cổ phần Cảng Hòn La xin gửi đến {xung_ho} lời chúc mừng tốt đẹp nhất!\n\nChúc {xung_ho} thật nhiều sức khỏe, hạnh phúc và thành công trong công việc cũng như cuộc sống.\n\nTrân trọng!\nCÔNG TY CỔ PHẦN CẢNG HÒN LA"""
   #                 st.code(loi_chuc); st.markdown(f"[👉 MỞ ZALO CHAT](https://zalo.me/{phone})")
   #     elif days_left <= 3: st.info(f"🎂 **{ma_nv} {ho_ten}** còn **{days_left}** ngày sẽ đến sinh nhật!")
    c.execute("SELECT Chuc_danh_nghe,COUNT(*) t FROM nhan_vien WHERE Trang_thai='DANG_LAM' GROUP BY Chuc_danh_nghe ORDER BY t DESC")
    data=c.fetchall(); db.close()
    if data: st.divider(); st.subheader("📈 Phân bố"); df=pd.DataFrame(data); df.columns=['Chức vụ','SL']; st.bar_chart(df.set_index('Chức vụ'))
    st.divider()
    if st.button("💾 BACKUP DỮ LIỆU NGAY", use_container_width=True):
        from backup_nv import backup_nhan_vien
        backup_nhan_vien(); st.success("✅ Đã backup! Kiểm tra thư mục D:\\HRM_Port\\backup")

# ========== ỨNG VIÊN ==========
elif menu=="👤 Ứng viên" and st.session_state.role=="admin":
    st.title("👤 Ứng viên")
    su = st.text_input("🔍 Tìm kiếm", key="suv")
    
    # Lấy danh mục vị trí dự tuyển
    db_f = get_connection()
    c_f = db_f.cursor()
    c_f.execute("SELECT Ten_vi_tri FROM vi_tri_cong_tac ORDER BY Ten_vi_tri")
    ds_vi_tri = [row[0] for row in c_f.fetchall()]
    c_f.execute("SELECT DISTINCT Vi_tri_du_tuyen FROM ung_vien WHERE Vi_tri_du_tuyen IS NOT NULL AND Vi_tri_du_tuyen != '' ORDER BY Vi_tri_du_tuyen")
    for row in c_f.fetchall():
        if row[0] not in ds_vi_tri:
            ds_vi_tri.append(row[0])
    db_f.close()
    
    col_f1, col_f2 = st.columns([2, 1])
    with col_f1:
        filter_vi_tri = st.selectbox("🔍 Lọc Vị trí dự tuyển:", ["Tất cả"] + ds_vi_tri)
    
    # ===== THÊM ỨNG VIÊN MỚI =====
    with st.expander("➕ THÊM ỨNG VIÊN MỚI", expanded=False):
        with st.form("add_uv_form"):
            db_f = get_connection()
            c_f = db_f.cursor()
            c_f.execute("SELECT Ten_vi_tri FROM vi_tri_cong_tac ORDER BY Ten_vi_tri")
            ds_vt_uv = [row[0] for row in c_f.fetchall()]
            db_f.close()
            
            col1, col2, col3 = st.columns(3)
            with col1:
                ho_ten_uv = st.text_input("Họ và tên *")
                vi_tri_uv = st.selectbox("Vị trí dự tuyển", [""] + ds_vt_uv)
                dien_thoai_uv = st.text_input("SĐT")
            with col2:
                ngay_sinh_uv = st.text_input("Ngày sinh (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                gioi_tinh_uv = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"])
            with col3:
                ngay_vao_lam_uv = st.text_input("Ngày vào làm (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                ghi_chu_uv = st.text_area("Ghi chú")
            
            submitted = st.form_submit_button("💾 LƯU", use_container_width=True)
            
            if submitted:
                if ho_ten_uv:
                    ngay_loi = []
                    if ngay_sinh_uv and not parse_date(ngay_sinh_uv): 
                        ngay_loi.append("Ngày sinh")
                    if ngay_vao_lam_uv and not parse_date(ngay_vao_lam_uv): 
                        ngay_loi.append("Ngày vào làm")
                    
                    if ngay_loi:
                        st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                    else:
                        try:
                            db = get_connection()
                            c = db.cursor()
                            
                            # Chèn trước, không có Ma_UV
                            c.execute("""INSERT INTO ung_vien (Ho_ten, Vi_tri_du_tuyen, Dien_thoai, 
                                Ngay_sinh, Gioi_tinh, Ngay_vao_lam, Luong_bao_hiem, Trang_thai)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, 'CHO_DUYET')""",
                                (ho_ten_uv, vi_tri_uv, dien_thoai_uv, parse_date(ngay_sinh_uv),
                                 gioi_tinh_uv, parse_date(ngay_vao_lam_uv), ghi_chu_uv))
                            
                            # Lấy Id vừa tạo
                            new_id = c.lastrowid
                            ma_uv = f"UV{new_id:04d}"
                            
                            # Cập nhật Ma_UV
                            c.execute("UPDATE ung_vien SET Ma_UV = %s WHERE Id = %s", (ma_uv, new_id))
                            
                            db.commit()
                            db.close()
                            
                            st.session_state['show_zalo_invite'] = True
                            st.session_state['last_uv'] = {
                                'ho_ten': ho_ten_uv,
                                'ma_uv': ma_uv,
                                'gioi_tinh': gioi_tinh_uv,
                                'dien_thoai': dien_thoai_uv
                            }
                            
                            st.success(f"✅ Đã thêm ứng viên: {ho_ten_uv} (Mã: {ma_uv})")
                            st.rerun()
                                
                        except Exception as e:
                            st.error(f"❌ Lỗi khi thêm ứng viên: {e}")
                else:
                    st.error("Họ tên không được để trống!")
    
    # Hiển thị phần mời Zalo group SAU KHI form đóng (đã sửa xưng hô)
    if st.session_state.get('show_zalo_invite', False):
        uv = st.session_state['last_uv']
        
        # Xác định xưng hô theo giới tính - ĐÃ SỬA
        gt = uv['gioi_tinh']
        if gt == 'Nam':
            xung_ho = 'Anh'
        elif gt == 'Nữ':
            xung_ho = 'Chị'
        else:
            xung_ho = 'Anh/Chị'
        
        ZALO_GROUP_LINK = COMPANY_CONFIG.get('zalo_group_link', '')
        ZALO_GROUP_NAME = COMPANY_CONFIG.get('zalo_group_name', 'Group Nhân sự')
        
        if ZALO_GROUP_LINK:
            st.markdown("---")
            st.markdown("### 📱 MỜI ỨNG VIÊN THAM GIA ZALO GROUP")
            st.markdown(f"**Hãy gửi lời mời sau cho {xung_ho} {uv['ho_ten']}:**")
            
            col_qr, col_info = st.columns([1, 2])
            
            with col_qr:
                try:
                    import qrcode
                    from io import BytesIO
                    qr = qrcode.QRCode(box_size=4, border=2)
                    qr.add_data(ZALO_GROUP_LINK)
                    qr.make(fit=True)
                    qr_img = qr.make_image(fill_color="black", back_color="white")
                    buf = BytesIO()
                    qr_img.save(buf, format='PNG')
                    st.image(buf, caption=f'QR: {ZALO_GROUP_NAME}', width=150)
                except Exception as e:
                    st.warning(f"Không thể tạo QR code: {e}")
            
            with col_info:
                st.markdown(f"**🔗 Link mời group:**")
                st.code(ZALO_GROUP_LINK, language="text")
                
                # Nội dung tin nhắn với xưng hô đúng
                noi_dung_moi = f"""Xin chào {xung_ho} {uv['ho_ten']},
Chào mừng {xung_ho} đã nộp hồ sơ gia nhập Công ty cổ phần Cảng Hòn La.
{xung_ho} vui lòng tham gia group Zalo của công ty để nhận thông báo và cập nhật công việc:
{ZALO_GROUP_LINK}
Xin cảm ơn!"""
                
                st.text_area("📝 Nội dung tin nhắn (copy để gửi):", noi_dung_moi, height=150, key="zalo_content")
                
                if uv['dien_thoai']:
                    phone = ''.join(filter(str.isdigit, uv['dien_thoai']))
                    if phone.startswith('84'):
                        phone = '0' + phone[2:]
                    elif not phone.startswith('0'):
                        phone = '0' + phone
                    zalo_link = f"https://zalo.me/{phone}"
                    st.markdown(f"[👉 MỞ ZALO CỦA {xung_ho} {uv['ho_ten']}]({zalo_link})")
            
            col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
            with col_btn2:
                if st.button("✅ ĐÃ GỬI LỜI MỜI", use_container_width=True):
                    st.session_state['show_zalo_invite'] = False
                    st.rerun()
            
            st.markdown("---")
            st.info("💡 **Hướng dẫn:** Copy nội dung tin nhắn → Mở Zalo → Dán và gửi cho ứng viên")
        else:
            st.info("ℹ️ Chưa cấu hình link Zalo group. Vui lòng thêm 'zalo_group_link' vào COMPANY_CONFIG")
            if st.button("✅ ĐÓNG", use_container_width=True):
                st.session_state['show_zalo_invite'] = False
                st.rerun()
    
    st.divider()
    
    # ===== HIỂN THỊ BẢNG ỨNG VIÊN =====
    t1, t2, t3, t4 = st.tabs(["📋 Tất cả", "⏳ Chờ duyệt", "✅ Đã nhận", "❌ Từ chối"])
    tm = {"📋 Tất cả": "", "⏳ Chờ duyệt": "CHO_DUYET", "✅ Đã nhận": "DA_NHAN_VIEC", "❌ Từ chối": "TU_CHOI"}
    
    for tn, tab in zip(tm.keys(), [t1, t2, t3, t4]):
        with tab:
            tt = tm[tn]
            db = get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            sql = "SELECT * FROM ung_vien WHERE 1=1"
            params = []
            if tt:
                sql += " AND Trang_thai = %s"
                params.append(tt)
            if su:
                sql += " AND (Ho_ten LIKE %s OR Dien_thoai LIKE %s)"
                params.extend([f'%{su}%', f'%{su}%'])
            if filter_vi_tri != "Tất cả":
                sql += " AND Vi_tri_du_tuyen = %s"
                params.append(filter_vi_tri)
            sql += " ORDER BY Id ASC"
            c.execute(sql, tuple(params))
            ds = c.fetchall()
            db.close()
            
            if ds:
                # Thêm cột STT vào dataframe
                for i, row in enumerate(ds, 1):
                    row['STT'] = i
                
                df = pd.DataFrame(ds)
                for col in df.columns:
                    if 'Ngay' in col and col != 'STT':
                        df[col] = df[col].apply(format_date)
                
                # Thêm cột checkbox
                if 'selected' not in df.columns:
                    df.insert(0, 'selected', False)
                
                # Hiển thị STT trong danh sách cột
                display_cols = ['selected', 'STT', 'ho_ten', 'Vi_tri_du_tuyen', 'dien_thoai', 'ngay_vao_lam', 'Luong_bao_hiem', 'ngay_sinh']
                available_cols = [c for c in display_cols if c in df.columns]
                df_show = df[available_cols]
                
                col_map = {
                    'selected': 'Chọn',
                    'STT': 'STT',
                    'ho_ten': 'Họ tên',
                    'Vi_tri_du_tuyen': 'Vị trí dự tuyển',
                    'dien_thoai': 'SĐT',
                    'ngay_vao_lam': 'Ngày vào làm',
                    'Luong_bao_hiem': 'Ghi chú',
                    'ngay_sinh': 'Ngày sinh',
                }
                df_show.rename(columns=col_map, inplace=True)
                
                st.caption(f"📌 {len(ds)} kết quả. Tick chọn 1 ứng viên để sửa hoặc chuyển sang thử việc.")
                
                # Kiểm tra số lượng checkbox được chọn
                selected_rows = []
                edited_df = st.data_editor(
                    df_show,
                    column_config={
                        "Chọn": st.column_config.CheckboxColumn("Chọn", default=False)
                    },
                    disabled=[col for col in df_show.columns if col != 'Chọn'],
                    hide_index=True,
                    height=400,
                    key=f"uv_editor_{tn}"
                )
                
                if edited_df is not None and 'Chọn' in edited_df.columns:
                    selected_rows = edited_df[edited_df['Chọn'] == True]
                    
                    # Nếu chọn nhiều hơn 1, báo lỗi và bỏ chọn
                    if len(selected_rows) > 1:
                        st.error("⚠️ Chỉ được chọn 1 ứng viên! Vui lòng bỏ chọn bớt.")
                    
                    elif len(selected_rows) == 1:
                        selected_idx = selected_rows.index[0]
                        selected_nv = df.iloc[selected_idx]
                        selected_name = selected_nv['ho_ten']
                        selected_stt = selected_nv['STT']
                        
                        # Tạo 2 cột cho các nút chức năng
                        col_btn1, col_btn2 = st.columns(2)
                        
                        with col_btn1:
                            if st.button(f"✏️ SỬA '{selected_name}'", key=f"edit_sel_{tn}"):
                                st.session_state['edit_uv_name'] = selected_name
                                st.session_state['edit_uv_id'] = int(selected_nv['id'])
                                st.rerun()
                        
                        # Nút chuyển sang thử việc (CHỈ HIỂN THỊ Ở TAB CHỜ DUYỆT)
                        if tn == "⏳ Chờ duyệt":
                            with col_btn2:
                                if st.button(f"✅ CHUYỂN '{selected_name}' SANG THỬ VIỆC", type="primary", key=f"chuyen_uv_{tn}"):
                                    try:
                                        db = get_connection()
                                        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                                        
                                        # Chuyển đổi Id sang int Python thuần
                                        uv_id = int(selected_nv['id'])  # <-- THÊM DÒNG NÀY ĐỂ CHUYỂN ĐỔI
                                        
                                        # Lấy thông tin ứng viên
                                        c.execute("SELECT * FROM ung_vien WHERE Id = %s", (uv_id,))
                                        uv = c.fetchone()
                                        
                                        if uv:
                                            # Tương tự, chuyển đổi các giá trị khác nếu cần
                                            stt_moi = int(c.fetchone()['COALESCE(MAX(STT),0)+1']) if c.fetchone() else 1
                                            
                                            # Hoặc dùng cách an toàn hơn:
                                            c.execute("SELECT COALESCE(MAX(STT),0)+1 as next_stt FROM nhan_vien")
                                            result = c.fetchone()
                                            stt_moi = int(result['next_stt']) if result else 1
                                            
                                            ma_nv = f"NV{stt_moi:03d}"
                                            nhl = uv.get('ngay_vao_lam') or date.today()
                                            if isinstance(nhl, str):
                                                from datetime import datetime
                                                nhl = datetime.strptime(nhl, '%Y-%m-%d').date()
                                            
                                            # Đếm số hợp đồng thử việc
                                            c.execute("SELECT COUNT(*) as tv_count FROM nhan_vien WHERE So_HDLD LIKE '%/HĐTV-CHL'")
                                            tv_result = c.fetchone()
                                            tv_cnt = int(tv_result['tv_count']) + 1
                                            
                                            so_hdtv = f"{tv_cnt:02d}/{nhl.year}/HĐTV-CHL"
                                            
                                            c.execute("""INSERT INTO nhan_vien (STT, Ma_NV, So_HDLD, Ho_ten, Chuc_danh_nghe, Dien_thoai,
                                                Ngay_sinh, Gioi_tinh, Ngay_vao_lam, Noi_lam_viec, Loai_hop_dong, Trang_thai, Trang_thai_BHXH, Ngay_ky_HD)
                                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'Thử việc', 'THU_VIEC', 'CHUA_DONG', %s)""",
                                                (stt_moi, ma_nv, so_hdtv, uv['ho_ten'], uv['Vi_tri_du_tuyen'], uv['dien_thoai'],
                                                 uv['ngay_sinh'], uv['gioi_tinh'], nhl, 'Cảng THQT Hòn La', nhl))
                                            
                                            c.execute("UPDATE ung_vien SET Trang_thai='DA_NHAN_VIEC', Ma_NV=%s WHERE Id=%s", (ma_nv, uv_id))
                                            db.commit()
                                            
                                            st.success(f"✅ Đã chuyển {uv['ho_ten']} → {ma_nv} ({so_hdtv})")
                                            st.rerun()
                                        else:
                                            st.error("Không tìm thấy ứng viên!")
                                        db.close()
                                    except Exception as e:
                                        st.error(f"❌ Lỗi: {e}")
                                        
                                        if uv:
                                            c.execute("SELECT COALESCE(MAX(STT),0)+1 FROM nhan_vien")
                                            stt_moi = c.fetchone()['COALESCE(MAX(STT),0)+1']
                                            ma_nv = f"NV{stt_moi:03d}"
                                            nhl = uv.get('ngay_vao_lam') or date.today()
                                            if isinstance(nhl, str):
                                                from datetime import datetime
                                                nhl = datetime.strptime(nhl, '%Y-%m-%d').date()
                                            
                                            c.execute("SELECT COUNT(*) FROM nhan_vien WHERE So_HDLD LIKE '%/HĐTV-CHL'")
                                            tv_cnt = c.fetchone()[0] + 1
                                            so_hdtv = f"{tv_cnt:02d}/{nhl.year}/HĐTV-CHL"
                                            
                                            c.execute("""INSERT INTO nhan_vien (STT, Ma_NV, So_HDLD, Ho_ten, Chuc_danh_nghe, Dien_thoai,
                                                Ngay_sinh, Gioi_tinh, Ngay_vao_lam, Noi_lam_viec, Loai_hop_dong, Trang_thai, Trang_thai_BHXH, Ngay_ky_HD)
                                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'Thử việc', 'THU_VIEC', 'CHUA_DONG', %s)""",
                                                (stt_moi, ma_nv, so_hdtv, uv['ho_ten'], uv['Vi_tri_du_tuyen'], uv['dien_thoai'],
                                                 uv['ngay_sinh'], uv['gioi_tinh'], nhl, 'Cảng THQT Hòn La', nhl))
                                            
                                            c.execute("UPDATE ung_vien SET Trang_thai='DA_NHAN_VIEC', Ma_NV=%s WHERE Id=%s", (ma_nv, uv['id']))
                                            db.commit()
                                            
                                            st.success(f"✅ Đã chuyển {uv['ho_ten']} → {ma_nv} ({so_hdtv})")
                                            st.rerun()
                                        else:
                                            st.error("Không tìm thấy ứng viên!")
                                        db.close()
                                    except Exception as e:
                                        st.error(f"❌ Lỗi: {e}")
            else:
                st.info("Không có dữ liệu")
    
    # ===== FORM SỬA ỨNG VIÊN =====
    if 'edit_uv_name' in st.session_state:
        st.divider()
        st.subheader(f"✏️ Sửa: {st.session_state['edit_uv_name']}")
        
        db = get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if 'edit_uv_id' in st.session_state:
            c.execute("SELECT * FROM ung_vien WHERE Id = %s", (int(st.session_state['edit_uv_id']),))
        else:
            c.execute("SELECT * FROM ung_vien WHERE Ho_ten = %s", (st.session_state['edit_uv_name'],))
        uv_data = c.fetchone()
        db.close()
        
        if uv_data:
            with st.form("edit_uv_direct"):
                col1, col2, col3 = st.columns(3)
                with col1:
                    ho_ten_e = st.text_input("Họ và tên *", value=uv_data['ho_ten'] or '')
                    vi_tri_e = st.selectbox("Vị trí dự tuyển", [""] + ds_vi_tri,
                        index=([""] + ds_vi_tri).index(uv_data['Vi_tri_du_tuyen']) if uv_data['Vi_tri_du_tuyen'] in ds_vi_tri else 0)
                    dien_thoai_e = st.text_input("SĐT", value=uv_data['dien_thoai'] or '')
                with col2:
                    ngay_sinh_e = st.text_input("Ngày sinh (dd/mm/yyyy)", 
                        value=format_date(uv_data['ngay_sinh']), placeholder="dd/mm/yyyy", max_chars=10)
                    gioi_tinh_e = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"],
                        index=["", "Nam", "Nữ", "Khác"].index(uv_data['gioi_tinh']) if uv_data['gioi_tinh'] in ["Nam", "Nữ", "Khác"] else 0)
                with col3:
                    ngay_vao_lam_e = st.text_input("Ngày vào làm (dd/mm/yyyy)", 
                        value=format_date(uv_data['ngay_vao_lam']), placeholder="dd/mm/yyyy", max_chars=10)
                    ghi_chu_e = st.text_area("Ghi chú", value=uv_data['Luong_bao_hiem'] or '')
                    trang_thai_e = st.selectbox("Trạng thái", 
                        ["CHO_DUYET", "TU_CHOI", "DA_NHAN_VIEC"],
                        index=["CHO_DUYET", "TU_CHOI", "DA_NHAN_VIEC"].index(uv_data['trang_thai']) if uv_data['trang_thai'] in ["CHO_DUYET", "TU_CHOI", "DA_NHAN_VIEC"] else 0)
                
                col_save, col_del, col_cancel = st.columns(3)
                with col_save:
                    if st.form_submit_button("💾 CẬP NHẬT"):
                        ngay_loi = []
                        if ngay_sinh_e and not parse_date(ngay_sinh_e): 
                            ngay_loi.append("Ngày sinh")
                        if ngay_vao_lam_e and not parse_date(ngay_vao_lam_e): 
                            ngay_loi.append("Ngày vào làm")
                        if ngay_loi:
                            st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                        else:
                            db = get_connection()
                            c = db.cursor()
                            c.execute("""UPDATE ung_vien SET Ho_ten=%s, Vi_tri_du_tuyen=%s, Dien_thoai=%s,
                                Ngay_sinh=%s, Gioi_tinh=%s, Ngay_vao_lam=%s, Luong_bao_hiem=%s, Trang_thai=%s
                                WHERE Id=%s""",
                                (ho_ten_e, vi_tri_e, dien_thoai_e, parse_date(ngay_sinh_e), gioi_tinh_e,
                                 parse_date(ngay_vao_lam_e), ghi_chu_e, trang_thai_e, uv_data['id']))
                            db.commit()
                            db.close()
                            st.success("✅ Đã cập nhật!")
                            del st.session_state['edit_uv_name']
                            if 'edit_uv_id' in st.session_state:
                                del st.session_state['edit_uv_id']
                            st.rerun()
                with col_del:
                    if st.form_submit_button("🗑️ XÓA"):
                        db = get_connection()
                        c = db.cursor()
                        c.execute("DELETE FROM ung_vien WHERE Id = %s", (uv_data['id'],))
                        db.commit()
                        db.close()
                        st.success("🗑️ Đã xóa!")
                        del st.session_state['edit_uv_name']
                        if 'edit_uv_id' in st.session_state:
                            del st.session_state['edit_uv_id']
                        st.rerun()
                with col_cancel:
                    if st.form_submit_button("❌ HỦY"):
                        del st.session_state['edit_uv_name']
                        if 'edit_uv_id' in st.session_state:
                            del st.session_state['edit_uv_id']
                        st.rerun()
    
    # ===== QUẢN LÝ DANH MỤC VỊ TRÍ DỰ TUYỂN =====
    st.divider()
    with st.expander("⚙️ Quản lý danh mục Vị trí dự tuyển", expanded=False):
        with st.form("add_vi_tri_uv"):
            ten_vt_moi = st.text_input("Tên vị trí dự tuyển mới *")
            if st.form_submit_button("➕ Thêm"):
                if ten_vt_moi:
                    db = get_connection()
                    c = db.cursor()
                    c.execute("SELECT COUNT(*) FROM vi_tri_cong_tac WHERE Ten_vi_tri = %s", (ten_vt_moi,))
                    if c.fetchone()[0] == 0:
                        c.execute("INSERT INTO vi_tri_cong_tac (Ten_vi_tri) VALUES (%s)", (ten_vt_moi,))
                        db.commit()
                        st.success(f"✅ Đã thêm: {ten_vt_moi}")
                        st.rerun()
                    else:
                        st.warning("Vị trí này đã tồn tại!")
                    db.close()
                else:
                    st.error("Tên không được để trống!")
        
        db = get_connection()
        c = db.cursor()
        c.execute("SELECT Id, Ten_vi_tri FROM vi_tri_cong_tac ORDER BY Ten_vi_tri")
        ds_vt = c.fetchall()
        db.close()
        if ds_vt:
            st.caption("📋 Danh sách vị trí dự tuyển:")
            for row in ds_vt:
                st.write(f"- {row[1]}")

# ========== NHÂN VIÊN ==========
elif menu == "✅ Nhân viên":
    st.title("✅ Quản lý nhân viên")
    
    # Tạo 3 tab: Đang làm việc ; Đã nghỉ việc và Lịch sử công tác.
    tab_dang_lam, tab_da_nghi, tab_qtct = st.tabs(["📌 ĐANG LÀM VIỆC", "📋 ĐÃ NGHỈ VIỆC", "📜 LỊCH SỬ CÔNG TÁC"])
    
    # ==================== TAB 1: ĐANG LÀM VIỆC ====================
    with tab_dang_lam:
        st.caption("👥 Danh sách nhân viên đang làm việc (bao gồm thử việc)")
        sn = st.text_input("🔍 Tìm kiếm", key="snv_dang_lam")
        
        # ===== THÊM NHÂN VIÊN MỚI (chỉ admin) =====
        if st.session_state.role == "admin":
            with st.expander("➕ THÊM NHÂN VIÊN MỚI", expanded=False):
                # ... (giữ nguyên phần thêm nhân viên cũ)
                with st.form("add_nv"):
                    db = get_connection()
                    c = db.cursor()
                    c.execute("SELECT DISTINCT Ten_vi_tri FROM vi_tri_cong_tac ORDER BY Ten_vi_tri")
                    dcv = [row[0] for row in c.fetchall()]
                    db.close()
                    st.caption("📝 Thông tin cá nhân")
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        htn = st.text_input("Họ và tên *")
                        nsn = st.text_input("Ngày sinh (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                        gtn = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"])
                        qtn = st.text_input("Quốc tịch", value="Việt Nam")
                        dtn = st.text_input("Dân tộc", value="Kinh")
                    with c2:
                        scc = st.text_input("CCCD")
                        ncc = st.text_input("Ngày cấp CCCD (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                        ncc2 = st.text_input("Nơi cấp CCCD")
                        nqn = st.text_input("Nguyên quán")
                        ttn = st.text_input("Thường trú")
                    with c3:
                        dtn2 = st.text_input("SĐT")
                        emn = st.text_input("Email")
                        cdn = st.selectbox("Chức danh", [""] + dcv)
                        pbn = st.text_input("Phòng ban")
                        nlv = st.text_input("Nơi làm việc", value="Cảng THQT Hòn La")
                    st.divider()
                    st.caption("💼 Hợp đồng & BHXH")
                    c4, c5, c6 = st.columns(3)
                    with c4:
                        lhd = st.selectbox("Loại HĐ *", ["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"])
                        nvl = st.text_input("Ngày vào làm (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                        nkt = st.text_input("Ngày kết thúc", placeholder="dd/mm/yyyy", max_chars=10)
                        mbh = st.text_input("Mã BHXH")
                        tbd = st.text_input("Bắt đầu BH (dd/mm/yyyy)", placeholder="dd/mm/yyyy", max_chars=10)
                    with c5:
                        lbh = st.text_input("Lương BH")
                        hsl = st.text_input("Hệ số lương")
                        pcv = st.text_input("PC chức vụ")
                        ptv = st.text_input("PC TNVK (%)")
                        ptn = st.text_input("PC TNN (%)")
                    with c6:
                        mhb = st.selectbox("Mức hưởng BHYT", ["80%", "95%", "100%"])
                        tld = st.text_input("Tỷ lệ đóng (%)")
                        mtd = st.text_input("Mức tiền đóng")
                        ptd = st.selectbox("PT đóng", ["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"])
                        nbh = st.selectbox("Nhóm BHXH", ["", "Văn phòng", "Lao động trực tiếp"])
                    st.divider()
                    st.caption("🏦 Ngân hàng & KCB")
                    c7, c8, c9 = st.columns(3)
                    with c7:
                        stk = st.text_input("STK")
                        cnh = st.text_input("Chi nhánh NH")
                        tkb = st.text_input("Tỉnh KCB")
                        nkb = st.text_input("Nơi KCB")
                    with c8:
                        ths = st.text_input("Tỉnh/TP nhận HS")
                        phs = st.text_input("Phường/Xã nhận HS")
                        dhs = st.text_area("Địa chỉ nhận HS", height=100)
                    with c9:
                        dks = st.selectbox("ĐK nhận sổ", ["Có", "Không"])
                        hso = st.selectbox("Hồ sơ", ["", "Đã có HS", "Chưa có"])
                    
                    if st.form_submit_button("💾 LƯU"):
                        if htn:
                            ngay_loi = []
                            if nsn and not parse_date(nsn):
                                ngay_loi.append("Ngày sinh")
                            if ncc and not parse_date(ncc):
                                ngay_loi.append("Ngày cấp CCCD")
                            if nvl and not parse_date(nvl):
                                ngay_loi.append("Ngày vào làm")
                            if nkt and not parse_date(nkt):
                                ngay_loi.append("Ngày kết thúc")
                            if tbd and not parse_date(tbd):
                                ngay_loi.append("Bắt đầu BH")
                            if ngay_loi:
                                st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                            else:
                                try:
                                    db = get_connection()
                                    c = db.cursor()
                                    c.execute("SELECT COALESCE(MAX(STT),0)+1 FROM nhan_vien")
                                    stt_moi = c.fetchone()[0]
                                    ma_nv = f"NV{stt_moi:03d}"
                                    nhl = parse_date(nvl) or date.today()
                                    if lhd == "Thử việc":
                                        ttnv, ttbh, tbd_val = 'THU_VIEC', 'CHUA_DONG', None
                                        c.execute("SELECT COUNT(*) FROM nhan_vien WHERE So_HDLD LIKE '%/HĐTV-CHL'")
                                        tv_cnt = c.fetchone()[0] + 1
                                        so_hd = f"{tv_cnt:02d}/{nhl.year}/HĐTV-CHL"
                                    else:
                                        ttnv, ttbh = 'DANG_LAM', 'DANG_DONG'
                                        tbd_val = parse_date(tbd) or parse_date(nvl)
                                        c.execute("SELECT COALESCE(MAX(CAST(SUBSTRING_INDEX(So_HDLD,'/',1) AS UNSIGNED)),0)+1 FROM nhan_vien WHERE So_HDLD LIKE '%/HĐLĐ-CHL'")
                                        so_hd = f"{int(c.fetchone()[0]):02d}/{nhl.year}/HĐLĐ-CHL"
                                    c.execute("""INSERT INTO nhan_vien (STT, Ma_NV, So_HDLD, Ho_ten, Chuc_danh_nghe, Vi_tri_Id, Ngay_sinh, Gioi_tinh,
                                        Tinh_trang_hon_nhan, So_CCCD, Ngay_cap_CCCD, Noi_cap_CCCD, Nguyen_quan, Thuong_tru,
                                        Dien_thoai, Email, Email_lien_he, Ho_so, Luong_bao_hiem, Ma_so_BHXH, Ngay_vao_lam,
                                        Noi_lam_viec, So_tai_khoan_NH, Chi_nhanh_NH, Ngay_ky_HD, Thoi_han_HD, Loai_hop_dong,
                                        Nhom_BHXH, Thang_bat_dau_BH, Thang_ket_thuc_BH, Ghi_chu, Trang_thai, Trang_thai_BHXH,
                                        phong_ban_lam_viec, Ngay_ket_thuc, Quoc_tich, Dan_toc, He_so_luong, Phu_cap_chuc_vu,
                                        Phu_cap_TNVK, Phu_cap_TNN, Muc_huong_BHYT, Ty_le_dong, Muc_tien_dong, Phuong_thuc_dong,
                                        Tinh_nhan_HS, Phuong_nhan_HS, Dia_chi_nhan_HS, Tinh_KCB, Noi_dang_ky_KCB, Dang_ky_nhan_so)
                                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                                          (stt_moi, ma_nv, so_hd, htn, cdn, None, parse_date(nsn), gtn,
                                           None, scc, parse_date(ncc), ncc2, nqn, ttn, dtn2, emn, emn, '', lbh, mbh,
                                           parse_date(nvl), nlv, stk, cnh, parse_date(nvl), None, lhd,
                                           nbh, tbd_val, None, None, ttnv, ttbh, pbn, parse_date(nkt), qtn, dtn, hsl, pcv,
                                           ptv, ptn,
                                           mhb, tld, mtd, ptd, ths, phs, dhs, tkb, nkb, dks))
                                    db.commit()
                                    db.close()
                                    st.success(f"✅ Đã lưu nhân viên mới thành công! {htn} - {ma_nv}")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Lỗi: {e}")
                        else:
                            st.error("Họ tên không được để trống!")
            st.divider()
        
        # ===== LỌC CHỨC DANH & LOẠI HĐ =====
        db_f = get_connection()
        c_f = db_f.cursor()
        c_f.execute("SELECT DISTINCT Chuc_danh_nghe FROM nhan_vien WHERE Trang_thai IN ('DANG_LAM','THU_VIEC') AND Chuc_danh_nghe IS NOT NULL AND Chuc_danh_nghe != '' ORDER BY Chuc_danh_nghe")
        ds_chuc_danh = [row[0] for row in c_f.fetchall()]
        c_f.execute("SELECT DISTINCT Loai_hop_dong FROM nhan_vien WHERE Trang_thai IN ('DANG_LAM','THU_VIEC') AND Loai_hop_dong IS NOT NULL AND Loai_hop_dong != '' ORDER BY Loai_hop_dong")
        ds_loai_hd = [row[0] for row in c_f.fetchall()]
        db_f.close()
        
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            filter_chuc_danh = st.selectbox("🔍 Lọc Chức danh:", ["Tất cả"] + ds_chuc_danh, key="filter_cd_danglam")
        with col_f2:
            filter_loai_hd = st.selectbox("🔍 Lọc Loại HĐ:", ["Tất cả"] + ds_loai_hd, key="filter_lhd_danglam")
        
        db = get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        sql = "SELECT * FROM nhan_vien WHERE Trang_thai IN ('DANG_LAM','THU_VIEC')"
        params = []
        if sn:
            sql += " AND (Ho_ten LIKE %s OR Dien_thoai LIKE %s OR So_CCCD LIKE %s OR Ma_NV LIKE %s)"
            params.extend([f'%{sn}%'] * 4)
        if filter_chuc_danh != "Tất cả":
            sql += " AND Chuc_danh_nghe = %s"
            params.append(filter_chuc_danh)
        if filter_loai_hd != "Tất cả":
            sql += " AND Loai_hop_dong = %s"
            params.append(filter_loai_hd)
        sql += " ORDER BY Id DESC"
        c.execute(sql, tuple(params))
        ds = c.fetchall()
        db.close()
        
        if ds:
            df = pd.DataFrame(ds)
            for col in df.columns:
                if 'Ngay' in col:
                    df[col] = df[col].apply(format_date)
            
            if 'selected' not in df.columns:
                df.insert(0, 'selected', False)
            
            display_cols = ['selected', 'ma_nv', 'ho_ten', 'ngay_sinh', 'gioi_tinh', 'so_hdld', 'so_cccd', 'dien_thoai',
                            'Thuong_tru', 'chuc_danh_nghe', 'loai_hop_dong', 'ngay_vao_lam', 'ma_so_bhxh', 'Thang_bat_dau_BH']
            available_cols = [c for c in display_cols if c in df.columns]
            df_show = df[available_cols]
            
            col_map = {
                'selected': 'Chọn',
                'ma_nv': 'Mã NV',
                'ho_ten': 'Họ và tên',
                'ngay_sinh': 'Ngày sinh',
                'gioi_tinh': 'Giới tính',
                'so_hdld': 'Số HĐLĐ',
                'so_cccd': 'CCCD',
                'dien_thoai': 'SĐT',
                'Thuong_tru': 'Thường trú',
                'chuc_danh_nghe': 'Chức danh',
                'loai_hop_dong': 'Loại HĐ',
                'ngay_vao_lam': 'Ngày vào làm',
                'ma_so_bhxh': 'Mã số BHXH',
                'Thang_bat_dau_BH': 'Bắt đầu BH',
            }
            df_show.rename(columns=col_map, inplace=True)
            
            st.caption(f"📌 {len(ds)} kết quả. Tick chọn 1 nhân viên để thao tác.")
            
            edited_df = st.data_editor(
                df_show,
                column_config={
                    "Chọn": st.column_config.CheckboxColumn("Chọn để sửa", default=False)
                },
                disabled=[col for col in df_show.columns if col != 'Chọn'],
                hide_index=True,
                height=400,
                key="nv_editor_danglam"
            )
            
            selected_nv = None
            if edited_df is not None and 'Chọn' in edited_df.columns:
                selected_rows = edited_df[edited_df['Chọn'] == True]
                if len(selected_rows) > 0:
                    selected_idx = selected_rows.index[0]
                    selected_nv = df.iloc[selected_idx]
                    nv_id_key = selected_nv['id']
                    
                    col_btn1, col_btn2, col_btn3, col_btn4, col_btn5, col_btn6 = st.columns(6)
                    
                    with col_btn1:
                        if st.button(f"✏️ SỬA '{selected_nv['ho_ten']}'", key=f"edit_nv_btn_{nv_id_key}", use_container_width=True):
                            st.session_state['selected_nv_id'] = int(selected_nv['id'])
                            st.rerun()
                    
                    with col_btn2:
                        trang_thai_nv = selected_nv.get('trang_thai', '')
                        if trang_thai_nv == 'DANG_LAM':
                            if st.button(f"🖨️ IN HĐLĐ - {selected_nv['ho_ten']}", key=f"print_hdld_btn_{nv_id_key}", use_container_width=True):
                                db = get_connection()
                                c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                                c.execute("SELECT * FROM nhan_vien WHERE Id = %s", (int(selected_nv['id']),))
                                nv_full = c.fetchone()
                                db.close()
                                if nv_full:
                                    fp = tao_hop_dong(nv_full)
                                    with open(fp, "rb") as f:
                                        st.download_button(
                                            label="📥 TẢI HĐLĐ",
                                            data=f,
                                            file_name=f"HDLD_{nv_full['ho_ten']}_{datetime.now().strftime('%Y%m%d')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"download_hdld_{nv_id_key}"
                                        )
                                else:
                                    st.error("Không tìm thấy thông tin nhân viên!")
                        elif trang_thai_nv == 'THU_VIEC':
                            if st.button(f"🖨️ IN HĐTV - {selected_nv['ho_ten']}", key=f"print_hdtv_btn_{nv_id_key}", use_container_width=True):
                                db = get_connection()
                                c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                                c.execute("SELECT * FROM nhan_vien WHERE Id = %s", (int(selected_nv['id']),))
                                nv_full = c.fetchone()
                                db.close()
                                if nv_full:
                                    fp = tao_hop_dong_thu_viec(nv_full)
                                    with open(fp, "rb") as f:
                                        st.download_button(
                                            label="📥 TẢI HĐTV",
                                            data=f,
                                            file_name=f"HDTV_{nv_full['ho_ten']}_{datetime.now().strftime('%Y%m%d')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"download_hdtv_{nv_id_key}"
                                        )
                                else:
                                    st.error("Không tìm thấy thông tin nhân viên!")
                        else:
                            st.button(f"📄 {selected_nv['ho_ten']} (Không thể in HĐ)", disabled=True, use_container_width=True, key=f"disabled_btn_{nv_id_key}")
                    
                    with col_btn3:
                        if st.button(f"📱 GỬI ZALO - {selected_nv['ho_ten']}", key=f"zalo_btn_{nv_id_key}", use_container_width=True):
                            db = get_connection()
                            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                            c.execute("SELECT * FROM nhan_vien WHERE Id = %s", (int(selected_nv['id']),))
                            nv_full = c.fetchone()
                            db.close()
                            if nv_full:
                                ph = nv_full.get('dien_thoai', '')
                                if ph:
                                    ph = ph.replace('+84', '0').replace(' ', '').strip()
                                    st.code(tao_noi_dung_zalo(nv_full))
                                    st.markdown(f"[👉 MỞ ZALO](https://zalo.me/{ph})")
                                else:
                                    st.error("Chưa có SĐT!")
                            else:
                                st.error("Không tìm thấy thông tin nhân viên!")
                    
                    with col_btn4:
                        ma_bhxh = selected_nv.get('ma_so_bhxh', '')
                        chua_co_bhxh = not bool(ma_bhxh and str(ma_bhxh).strip())
                        if chua_co_bhxh:
                            if st.button(f"🏠 NHẬP THÔNG TIN HỘ GIA ĐÌNH - {selected_nv['ho_ten']}", key=f"bhxh_family_btn_{nv_id_key}", use_container_width=True, type="primary"):
                                st.session_state['bhxh_family_nv_id'] = int(selected_nv['id'])
                                st.session_state['bhxh_family_nv_name'] = selected_nv['ho_ten']
                                st.rerun()
                        else:
                            st.button(f"✅ ĐÃ CÓ BHXH - {selected_nv['ho_ten']}", disabled=True, use_container_width=True, key=f"has_bhxh_btn_{nv_id_key}")
                    
                    with col_btn5:
                        trang_thai_nv = selected_nv.get('trang_thai', '')
                        if trang_thai_nv == 'THU_VIEC':
                            # Kiểm tra xem đã có form chuyển đổi đang mở chưa
                            if f'convert_open_{nv_id_key}' not in st.session_state:
                                st.session_state[f'convert_open_{nv_id_key}'] = False
                            
                            if not st.session_state[f'convert_open_{nv_id_key}']:
                                if st.button(f"🔄 CHUYỂN HĐLĐ KHÔNG XĐTH - {selected_nv['ho_ten']}", 
                                            key=f"convert_hdld_btn_{nv_id_key}", 
                                            use_container_width=True, type="primary"):
                                    st.session_state[f'convert_open_{nv_id_key}'] = True
                                    st.rerun()
                            else:
                                # Hiển thị form chi tiết
                                st.markdown("---")
                                st.markdown("### 📝 CHUYỂN ĐỔI HỢP ĐỒNG LAO ĐỘNG")
                                st.caption("Vui lòng nhập đầy đủ thông tin cho quyết định chuyển đổi")
                                
                                # Lấy thông tin nhân viên
                                db_temp = get_connection()
                                c_temp = db_temp.cursor(dictionary=True)
                                c_temp.execute("SELECT * FROM nhan_vien WHERE Id = %s", (int(selected_nv['id']),))
                                nv_data = c_temp.fetchone()
                                db_temp.close()
                                
                                if nv_data:
                                    # Ngày quyết định (mặc định là hôm nay)
                                    ngay_quyet_dinh = st.date_input(
                                        "📅 Ngày quyết định:", 
                                        value=date.today(),
                                        key=f"ngay_qd_{nv_id_key}"
                                    )
                                    
                                    # Ngày hiệu lực (mặc định là ngày quyết định)
                                    ngay_hieu_luc = st.date_input(
                                        "📅 Ngày hiệu lực (bắt đầu HĐLĐ):", 
                                        value=ngay_quyet_dinh,
                                        key=f"ngay_hl_{nv_id_key}"
                                    )
                                    
                                    # Thông tin hợp đồng mới - TỰ ĐỘNG SINH SỐ
                                    current_year = datetime.now().year
                                    
                                    # Lấy số thứ tự lớn nhất của HĐLĐ trong năm
                                    db_temp2 = get_connection()
                                    c_temp2 = db_temp2.cursor()
                                    c_temp2.execute("""
                                        SELECT COALESCE(MAX(CAST(SUBSTRING_INDEX(So_HDLD, '/', 1) AS UNSIGNED)), 0) as max_stt
                                        FROM nhan_vien 
                                        WHERE So_HDLD LIKE %s AND Trang_thai IN ('DANG_LAM', 'THU_VIEC')
                                    """, (f'%/{current_year}/HĐLĐ-CHL',))
                                    result = c_temp2.fetchone()
                                    max_stt = result[0] if result else 0
                                    db_temp2.close()
                                    
                                    next_stt = max_stt + 1
                                    stt_str = str(next_stt).zfill(2)
                                    so_hd_moi = f"{stt_str}/{current_year}/HĐLĐ-CHL"
                                    
                                    st.info(f"📄 **Số HĐLĐ mới:** {so_hd_moi} (tự động sinh)")
                                    
                                    ngay_bat_dau_bh = st.date_input(
                                        "📅 Bắt đầu đóng BHXH:", 
                                        value=ngay_hieu_luc,
                                        key=f"ngay_bhxh_{nv_id_key}"
                                    )
                                    st.caption("⚠️ Đây là ngày bắt đầu tính đóng BHXH")
                                    
                                    # Lý do/Hiệu lực
                                    ly_do_chuyen = st.text_area(
                                        "📝 Lý do/ Nội dung quyết định:", 
                                        value="Hoàn thành thời gian thử việc, chuyển sang hợp đồng lao động không xác định thời hạn",
                                        key=f"ly_do_{nv_id_key}",
                                        height=80
                                    )
                                    
                                    # Nút xác nhận
                                    col_confirm1, col_confirm2, col_confirm3 = st.columns([1, 2, 1])
                                    with col_confirm2:
                                        if st.button("✅ XÁC NHẬN CHUYỂN ĐỔI", key=f"confirm_convert_{nv_id_key}", use_container_width=True, type="primary"):
                                            try:
                                                db = get_connection()
                                                c = db.cursor()
                                                
                                                # 1. Cập nhật thông tin nhân viên
                                                c.execute("""
                                                    UPDATE nhan_vien SET 
                                                        Trang_thai = 'DANG_LAM',
                                                        Loai_hop_dong = 'Không xác định thời hạn',
                                                        So_HDLD = %s,
                                                        Ngay_ky_HD = %s,
                                                        Ngay_chinh_thuc = %s,
                                                        Thang_bat_dau_BH = %s,
                                                        Trang_thai_BHXH = 'DANG_DONG',
                                                        Ngay_ket_thuc = NULL
                                                    WHERE Id = %s
                                                """, (so_hd_moi, ngay_quyet_dinh, ngay_hieu_luc, ngay_bat_dau_bh, int(selected_nv['id'])))
                                                
                                                # 2. Thêm vào bảng quyet_dinh_nhan_su
                                                c.execute("""
                                                    INSERT INTO quyet_dinh_nhan_su (
                                                        nhan_vien_Id, Loai_quyet_dinh, Ngay_quyet_dinh, Ngay_hieu_luc,
                                                        Noi_dung, So_quyet_dinh, Loai_hop_dong_cu, Loai_hop_dong_moi,
                                                        He_so_luong_cu, He_so_luong_moi
                                                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                                """, (
                                                    int(selected_nv['id']),
                                                    'CHINH_THUC',
                                                    ngay_quyet_dinh,
                                                    ngay_hieu_luc,
                                                    ly_do_chuyen,
                                                    f"QD{ngay_quyet_dinh.strftime('%Y%m%d')}_{selected_nv['ma_nv']}",
                                                    nv_data.get('loai_hop_dong', 'Thử việc'),
                                                    'Không xác định thời hạn',
                                                    nv_data.get('He_so_luong', 0),
                                                    nv_data.get('He_so_luong', 0)
                                                ))
                                                
                                                # 3. Cập nhật lịch sử công tác
                                                # Kết thúc giai đoạn thử việc
                                                c.execute("""
                                                    UPDATE lich_su_cong_tac 
                                                    SET Den_ngay = %s 
                                                    WHERE nhan_vien_Id = %s AND Den_ngay IS NULL
                                                """, (ngay_hieu_luc - timedelta(days=1), int(selected_nv['id'])))
                                                
                                                # Thêm giai đoạn mới
                                                c.execute("""
                                                    INSERT INTO lich_su_cong_tac (
                                                        nhan_vien_Id, Tu_ngay, Chuc_danh, phong_ban, 
                                                        Noi_lam_viec, Loai_hop_dong, He_so_luong
                                                    ) VALUES (%s, %s, %s, %s, %s, %s, %s)
                                                """, (
                                                    int(selected_nv['id']),
                                                    ngay_hieu_luc,
                                                    nv_data.get('chuc_danh_nghe', ''),
                                                    nv_data.get('phong_ban_lam_viec', ''),
                                                    nv_data.get('Noi_lam_viec', 'Cảng THQT Hòn La'),
                                                    'Không xác định thời hạn',
                                                    nv_data.get('He_so_luong', 0)
                                                ))
                                                
                                                db.commit()
                                                db.close()
                                                
                                                st.success(f"✅ Đã chuyển {nv_data['ho_ten']} sang HĐLĐ không xác định thời hạn!")
                                                st.info(f"📄 Số HĐLĐ mới: {so_hd_moi}")
                                                st.info(f"📅 Ngày hiệu lực: {ngay_hieu_luc.strftime('%d/%m/%Y')}")
                                                st.info(f"💰 Bắt đầu đóng BHXH từ: {ngay_bat_dau_bh.strftime('%d/%m/%Y')}")
                                                
                                                # Reset session state
                                                st.session_state[f'convert_open_{nv_id_key}'] = False
                                                st.rerun()
                                            except Exception as e:
                                                st.error(f"❌ Lỗi: {str(e)}")
                                    
                                    # Nút hủy
                                    if st.button("❌ HỦY", key=f"cancel_convert_{nv_id_key}", use_container_width=True):
                                        st.session_state[f'convert_open_{nv_id_key}'] = False
                                        st.rerun()
                        else:
                            st.button(f"✅ ĐÃ LÀ HĐLĐ", disabled=True, use_container_width=True, key=f"already_hdld_btn_{nv_id_key}")
                    
                    with col_btn6:
                        st.write("")
                    
                    st.divider()
            
            # Form sửa nhân viên (giữ nguyên code cũ)
            if 'selected_nv_id' in st.session_state:
                nid = int(st.session_state['selected_nv_id'])
                db = get_connection()
                c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                c.execute("SELECT * FROM nhan_vien WHERE Id=%s", (nid,))
                nd = c.fetchone()
                db.close()
                
                if nd:
                    st.subheader(f"✏️ Cập nhật: {nd.get('ho_ten', '')} ({nd.get('ma_nv', '')})")
                    with st.form("edit_nv"):
                        # ... (giữ nguyên form sửa nhân viên cũ)
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            hnv = st.text_input("Họ tên *", value=nd.get('ho_ten', ''))
                            nsnv = st.text_input("Ngày sinh (dd/mm/yyyy)", value=format_date(nd.get('ngay_sinh')), placeholder="dd/mm/yyyy", max_chars=10)
                            gtnv = st.selectbox("Giới tính", ["", "Nam", "Nữ", "Khác"], index=["", "Nam", "Nữ", "Khác"].index(nd.get('gioi_tinh', '')) if nd.get('gioi_tinh') in ["Nam", "Nữ", "Khác"] else 0)
                            qtnv = st.text_input("Quốc tịch", value=nd.get('Quoc_tich', 'Việt Nam'))
                            dtnv = st.text_input("Dân tộc", value=nd.get('Dan_toc', 'Kinh'))
                        with col2:
                            sccv = st.text_input("CCCD", value=nd.get('so_cccd', ''))
                            nccv = st.text_input("Ngày cấp CCCD (dd/mm/yyyy)", value=format_date(nd.get('Ngay_cap_CCCD')), placeholder="dd/mm/yyyy", max_chars=10)
                            ncv = st.text_input("Nơi cấp CCCD", value=nd.get('Noi_cap_CCCD', ''))
                            nqnv = st.text_input("Nguyên quán", value=nd.get('Nguyen_quan', ''))
                            ttnv = st.text_input("Thường trú", value=nd.get('Thuong_tru', ''))
                        with col3:
                            dtnv2 = st.text_input("SĐT", value=nd.get('dien_thoai', ''))
                            emnv = st.text_input("Email", value=nd.get('Email_lien_he', ''))
                            cdnv = st.text_input("Chức danh", value=nd.get('chuc_danh_nghe', ''))
                            pbnv = st.text_input("Phòng ban", value=nd.get('phong_ban_lam_viec', ''))
                            nlv2 = st.text_input("Nơi làm việc", value=nd.get('Noi_lam_viec', 'Cảng THQT Hòn La'))
                        
                        st.divider()
                        st.caption("💼 Hợp đồng & BHXH")
                        col4, col5, col6 = st.columns(3)
                        with col4:
                            lhdv = st.selectbox("Loại HĐ", ["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"], index=["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"].index(nd.get('loai_hop_dong', 'Thử việc')) if nd.get('loai_hop_dong') in ["Thử việc", "Xác định thời hạn", "Không xác định thời hạn"] else 0)
                            nvlv = st.text_input("Ngày vào làm (dd/mm/yyyy)", value=format_date(nd.get('ngay_vao_lam')), placeholder="dd/mm/yyyy", max_chars=10)
                            nktv = st.text_input("Ngày kết thúc (dd/mm/yyyy)", value=format_date(nd.get('ngay_ket_thuc')), placeholder="dd/mm/yyyy", max_chars=10)
                            mbhv = st.text_input("Mã BHXH", value=nd.get('ma_so_bhxh', ''))
                            tbdv = st.text_input("Bắt đầu BH (dd/mm/yyyy)", value=format_date(nd.get('Thang_bat_dau_BH')), placeholder="dd/mm/yyyy", max_chars=10)
                        with col5:
                            lbhv = st.text_input("Lương BH", value=nd.get('Luong_bao_hiem', ''))
                            hslv = st.text_input("Hệ số lương", value=str(nd.get('He_so_luong', '')))
                            pcvv = st.text_input("PC chức vụ", value=str(nd.get('Phu_cap_chuc_vu', '')))
                            ptvv = st.text_input("PC TNVK (%)", value=str(nd.get('Phu_cap_TNVK', '')))
                            ptnv = st.text_input("PC TNN (%)", value=str(nd.get('Phu_cap_TNN', '')))
                        with col6:
                            mhbv = st.selectbox("Mức hưởng BHYT", ["80%", "95%", "100%"], index=["80%", "95%", "100%"].index(nd.get('Muc_huong_BHYT', '80%')) if nd.get('Muc_huong_BHYT') in ["80%", "95%", "100%"] else 0)
                            tldv = st.text_input("Tỷ lệ đóng (%)", value=str(nd.get('Ty_le_dong', '')))
                            mtdv = st.text_input("Mức tiền đóng", value=str(nd.get('Muc_tien_dong', '')))
                            ptdv = st.selectbox("PT đóng", ["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"], index=["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"].index(nd.get('Phuong_thuc_dong', 'Hàng tháng')) if nd.get('Phuong_thuc_dong') in ["Hàng tháng", "3 tháng", "6 tháng", "12 tháng"] else 0)
                            nbhv = st.selectbox("Nhóm BHXH", ["", "Văn phòng", "Lao động trực tiếp"], index=["", "Văn phòng", "Lao động trực tiếp"].index(nd.get('Nhom_BHXH', '')) if nd.get('Nhom_BHXH') in ["Văn phòng", "Lao động trực tiếp"] else 0)
                        
                        st.divider()
                        st.caption("🏦 Ngân hàng & KCB")
                        col7, col8, col9 = st.columns(3)
                        with col7:
                            stkv = st.text_input("STK", value=nd.get('So_tai_khoan_NH', ''))
                            cnhv = st.text_input("Chi nhánh NH", value=nd.get('Chi_nhanh_NH', ''))
                            tkbv = st.text_input("Tỉnh KCB", value=nd.get('Tinh_KCB', ''))
                            nkbv = st.text_input("Nơi KCB", value=nd.get('Noi_dang_ky_KCB', ''))
                        with col8:
                            thsv = st.text_input("Tỉnh/TP nhận HS", value=nd.get('Tinh_nhan_HS', ''))
                            phsv = st.text_input("Phường/Xã nhận HS", value=nd.get('Phuong_nhan_HS', ''))
                            dhsv = st.text_area("Địa chỉ nhận HS", value=nd.get('Dia_chi_nhan_HS', ''), height=100)
                        with col9:
                            dksv = st.selectbox("ĐK nhận sổ", ["Có", "Không"], index=["Có", "Không"].index(nd.get('Dang_ky_nhan_so', 'Có')) if nd.get('Dang_ky_nhan_so') in ["Có", "Không"] else 0)
                            hsov = st.selectbox("Hồ sơ", ["", "Đã có HS", "Chưa có"], index=["", "Đã có HS", "Chưa có"].index(nd.get('Ho_so', '')) if nd.get('Ho_so') in ["Đã có HS", "Chưa có"] else 0)
                        
                        col_save, col_quit, col_delete = st.columns(3)
                        with col_save:
                            if st.form_submit_button("💾 CẬP NHẬT"):
                                if hnv:
                                    ngay_loi = []
                                    if nsnv and not parse_date(nsnv):
                                        ngay_loi.append("Ngày sinh")
                                    if nccv and not parse_date(nccv):
                                        ngay_loi.append("Ngày cấp CCCD")
                                    if nvlv and not parse_date(nvlv):
                                        ngay_loi.append("Ngày vào làm")
                                    if nktv and not parse_date(nktv):
                                        ngay_loi.append("Ngày kết thúc")
                                    if tbdv and not parse_date(tbdv):
                                        ngay_loi.append("Bắt đầu BH")
                                    if ngay_loi:
                                        st.error(f"Sai định dạng dd/mm/yyyy: {', '.join(ngay_loi)}")
                                    else:
                                        try:
                                            db = get_connection()
                                            c = db.cursor()
                                            nhl = parse_date(nvlv) or date.today()
                                            ngay_bat_dau_bh = parse_date(tbdv) if tbdv and tbdv.strip() else None
                                            if lhdv == "Thử việc":
                                                tt_nv, tt_bh, tbd_val = 'THU_VIEC', 'CHUA_DONG', None
                                            else:
                                                tt_nv, tt_bh = 'DANG_LAM', 'DANG_DONG'
                                                if ngay_bat_dau_bh:
                                                    tbd_val = ngay_bat_dau_bh
                                                else:
                                                    tbd_val = parse_date(nvlv)
                                            c.execute("""UPDATE nhan_vien SET Ho_ten=%s,Chuc_danh_nghe=%s,Ngay_sinh=%s,Gioi_tinh=%s,
                                                So_CCCD=%s,Ngay_cap_CCCD=%s,Noi_cap_CCCD=%s,Nguyen_quan=%s,Thuong_tru=%s,Dien_thoai=%s,
                                                Email=%s,Email_lien_he=%s,Ho_so=%s,Luong_bao_hiem=%s,Ma_so_BHXH=%s,Ngay_vao_lam=%s,Noi_lam_viec=%s,
                                                So_tai_khoan_NH=%s,Chi_nhanh_NH=%s,Ngay_ky_HD=%s,Loai_hop_dong=%s,Nhom_BHXH=%s,
                                                Thang_bat_dau_BH=%s,Trang_thai=%s,Trang_thai_BHXH=%s,phong_ban_lam_viec=%s,
                                                Ngay_ket_thuc=%s,Quoc_tich=%s,Dan_toc=%s,He_so_luong=%s,Phu_cap_chuc_vu=%s,
                                                Phu_cap_TNVK=%s,Phu_cap_TNN=%s,Muc_huong_BHYT=%s,Ty_le_dong=%s,Muc_tien_dong=%s,
                                                Phuong_thuc_dong=%s,Tinh_nhan_HS=%s,Phuong_nhan_HS=%s,Dia_chi_nhan_HS=%s,
                                                Tinh_KCB=%s,Noi_dang_ky_KCB=%s,Dang_ky_nhan_so=%s WHERE Id=%s""",
                                                  (hnv, cdnv, parse_date(nsnv), gtnv, sccv, parse_date(nccv), ncv, nqnv, ttnv, dtnv2,
                                                   emnv, emnv, hsov, lbhv, mbhv, parse_date(nvlv), nlv2, stkv, cnhv, parse_date(nvlv), lhdv,
                                                   nbhv, tbd_val, tt_nv, tt_bh, pbnv, parse_date(nktv), qtnv, dtnv, hslv, pcvv, ptvv, ptnv,
                                                   mhbv, tldv, mtdv, ptdv, thsv, phsv, dhsv, tkbv, nkbv, dksv, nid))
                                            db.commit()
                                            db.close()
                                            st.success(f"✅ Đã cập nhật: {hnv}")
                                            del st.session_state['selected_nv_id']
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"❌ Lỗi: {e}")
                                else:
                                    st.error("Họ tên không được để trống!")
                        with col_quit:
                            # Sử dụng session state để quản lý trạng thái hiển thị form nghỉ việc
                            if f'nghi_viec_open_{nid}' not in st.session_state:
                                st.session_state[f'nghi_viec_open_{nid}'] = False
                            
                            if not st.session_state[f'nghi_viec_open_{nid}']:
                                if st.form_submit_button("🚫 NGHỈ VIỆC", use_container_width=True, type="secondary"):
                                    st.session_state[f'nghi_viec_open_{nid}'] = True
                                    st.rerun()
                            else:
                                st.markdown("---")
                                st.markdown("### 📝 XÁC NHẬN NGHỈ VIỆC")
                                
                                # Lấy ngày nghỉ mặc định là hôm nay
                                default_ngay_nghi = date.today()
                                
                                # Kiểm tra nếu đã có Ngay_ket_thuc trong database thì lấy làm mặc định
                                ngay_ket_thuc_hien_tai = nd.get('ngay_ket_thuc')
                                if ngay_ket_thuc_hien_tai and ngay_ket_thuc_hien_tai != '':
                                    try:
                                        if hasattr(ngay_ket_thuc_hien_tai, 'strftime'):
                                            default_ngay_nghi = ngay_ket_thuc_hien_tai
                                    except:
                                        pass
                                
                                # Input ngày nghỉ việc
                                ngay_nghi = st.date_input(
                                    "📅 Ngày quyết định nghỉ việc (Ngày kết thúc HĐLĐ):", 
                                    value=default_ngay_nghi,
                                    key=f"ngay_nghi_{nid}"
                                )
                                
                                # Input lý do nghỉ việc
                                ly_do_nghi = st.text_area(
                                    "📝 Lý do nghỉ việc:", 
                                    value=nd.get('Ly_do_nghi', '') if 'Ly_do_nghi' in nd else '',
                                    placeholder="VD: Xin nghỉ theo nguyện vọng cá nhân, Chuyển công tác, Hết hạn hợp đồng...",
                                    key=f"ly_do_nghi_{nid}",
                                    height=100
                                )
                                
                                col_xac_nhan, col_huy = st.columns(2)
                                
                                with col_xac_nhan:
                                    if st.form_submit_button("✅ XÁC NHẬN NGHỈ VIỆC", use_container_width=True, type="primary"):
                                        try:
                                            db = get_connection()
                                            c = db.cursor()
                                            
                                            # Kiểm tra xem cột Ly_do_nghi có tồn tại không
                                            c.execute("SHOW COLUMNS FROM nhan_vien LIKE 'Ly_do_nghi'")
                                            has_ly_do_column = c.fetchone() is not None
                                            
                                            if has_ly_do_column:
                                                c.execute("""
                                                    UPDATE nhan_vien 
                                                    SET Trang_thai = 'NGHI_VIEC', 
                                                        Ngay_ket_thuc = %s,
                                                        Ly_do_nghi = %s
                                                    WHERE Id = %s
                                                """, (ngay_nghi, ly_do_nghi if ly_do_nghi else None, nid))
                                            else:
                                                # Nếu chưa có cột Ly_do_nghi, chỉ cập nhật Ngay_ket_thuc
                                                c.execute("""
                                                    UPDATE nhan_vien 
                                                    SET Trang_thai = 'NGHI_VIEC', 
                                                        Ngay_ket_thuc = %s
                                                    WHERE Id = %s
                                                """, (ngay_nghi, nid))
                                                if ly_do_nghi:
                                                    st.info("💡 Lưu ý: Cột lưu lý do nghỉ việc chưa có trong database. Chỉ có ngày nghỉ được lưu.")
                                            
                                            db.commit()
                                            db.close()
                                            
                                            # Xóa session state
                                            st.session_state[f'nghi_viec_open_{nid}'] = False
                                            if 'selected_nv_id' in st.session_state:
                                                del st.session_state['selected_nv_id']
                                            
                                            st.success(f"✅ Đã cập nhật nghỉ việc cho {nd.get('ho_ten', '')}!")
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"❌ Lỗi khi cập nhật nghỉ việc: {e}")
                                
                                with col_huy:
                                    if st.form_submit_button("❌ HỦY", use_container_width=True):
                                        st.session_state[f'nghi_viec_open_{nid}'] = False
                                        st.rerun()                                  
                        with col_delete:
                            if st.form_submit_button("🗑️ XÓA"):
                                db = get_connection()
                                c = db.cursor()
                                c.execute("DELETE FROM ho_so_nhan_vien WHERE nhan_vien_Id=%s", (nid,))
                                c.execute("DELETE FROM nhan_vien WHERE Id=%s", (nid,))
                                db.commit()
                                db.close()
                                st.success("🗑️ Đã xóa!")
                                del st.session_state['selected_nv_id']
                                st.rerun()
                    
                    if st.button("❌ HỦY SỬA", use_container_width=True):
                        del st.session_state['selected_nv_id']
                        st.rerun()
            
            # Form nhập thông tin hộ gia đình (giữ nguyên)
            if 'bhxh_family_nv_id' in st.session_state:
                # ... (giữ nguyên code form hộ gia đình cũ)
                nv_id = st.session_state['bhxh_family_nv_id']
                nv_name = st.session_state['bhxh_family_nv_name']
                st.divider()
                st.subheader(f"🏠 NHẬP THÔNG TIN HỘ GIA ĐÌNH CHO: {nv_name}")
                st.caption("Vui lòng nhập đầy đủ thông tin chủ hộ và các thành viên trong hộ gia đình")
                
                db = get_connection()
                c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                c.execute("SELECT * FROM nhan_vien WHERE Id = %s", (nv_id,))
                nv_data = c.fetchone()
                db.close()
                
                if 'bhxh_family_members' not in st.session_state:
                    db_temp = get_connection()
                    c_temp = db_temp.cursor(dictionary=True)
                    c_temp.execute("SELECT * FROM phu_luc_gia_dinh WHERE nhan_vien_Id = %s", (nv_id,))
                    existing_members = c_temp.fetchall()
                    db_temp.close()
                    st.session_state['bhxh_family_members'] = []
                    for tv in existing_members:
                        st.session_state['bhxh_family_members'].append({
                            'ho_ten': tv['ho_ten'], 'ngay_sinh': tv['ngay_sinh'], 'gioi_tinh': tv['gioi_tinh'],
                            'quoc_tich': tv['Quoc_tich'], 'dan_toc': tv['Dan_toc'], 'quan_he': tv['Quan_he_voi_chu_ho'],
                            'tinh': tv['Tinh_thanh_pho'], 'phuong_xa': tv['Phuong_xa']
                        })
                
                db_temp = get_connection()
                c_temp = db_temp.cursor()
                c_temp.execute("SELECT Ma_tinh, Ten_tinh FROM danh_muc_tinh ORDER BY Ten_tinh")
                ds_tinh = c_temp.fetchall()
                db_temp.close()
                tinh_options = {ten: ma for ma, ten in ds_tinh}
                
                if st.session_state.bhxh_family_members:
                    st.markdown("**Danh sách thành viên đã thêm:**")
                    tv_data = []
                    for i, tv in enumerate(st.session_state.bhxh_family_members):
                        tv_data.append({"STT": i+1, "Họ và tên": tv['ho_ten'], "Ngày sinh": format_date(tv['ngay_sinh']),
                                        "Giới tính": tv['gioi_tinh'], "Quốc tịch": tv['quoc_tich'], "Dân tộc": tv['dan_toc'],
                                        "Quan hệ chủ hộ": tv['quan_he'], "Tỉnh/TP": tv['tinh'], "Phường/Xã": tv['phuong_xa']})
                    df_tv = pd.DataFrame(tv_data)
                    st.dataframe(df_tv, use_container_width=True, hide_index=True)
                    col_del1, col_del2, col_del3 = st.columns([1,1,1])
                    with col_del2:
                        tv_to_delete = st.number_input("Nhập STT thành viên cần xóa:", min_value=1, max_value=len(st.session_state.bhxh_family_members), step=1, key="tv_delete_family")
                        if st.button("🗑️ Xóa thành viên", key="btn_del_tv_family"):
                            st.session_state.bhxh_family_members.pop(tv_to_delete - 1)
                            st.rerun()
                
                with st.form(key=f"bhxh_family_form_{nv_id}"):
                    st.markdown("**I. THÔNG TIN CHỦ HỘ:**")
                    col1, col2 = st.columns(2)
                    with col1:
                        ho_ten_chu_ho = st.text_input("Họ và tên chủ hộ", value=nv_data.get('Ho_ten_chu_ho', '') if nv_data else '')
                        so_cccd_chu_ho = st.text_input("Số CCCD chủ hộ", value=nv_data.get('So_CCCD_chu_ho', '') if nv_data else '')
                        tinh_chu_ho_index = 0
                        tinh_chu_ho_current = nv_data.get('Tinh_thanh_pho_chu_ho', '') if nv_data else ''
                        if tinh_chu_ho_current in tinh_options:
                            tinh_chu_ho_index = list(tinh_options.keys()).index(tinh_chu_ho_current) + 1
                        tinh_chu_ho = st.selectbox("Tỉnh/Thành phố (chủ hộ)", options=[""] + list(tinh_options.keys()), index=tinh_chu_ho_index)
                    with col2:
                        phuong_xa_options = []
                        phuong_xa_current = nv_data.get('Phuong_xa_chu_ho', '') if nv_data else ''
                        if tinh_chu_ho and tinh_chu_ho != "":
                            ma_tinh = tinh_options.get(tinh_chu_ho)
                            db_temp2 = get_connection()
                            c_temp2 = db_temp2.cursor()
                            c_temp2.execute("SELECT Ten_xa FROM danh_muc_phuong_xa WHERE Ma_tinh = %s ORDER BY Ten_xa", (ma_tinh,))
                            phuong_xa_options = [row[0] for row in c_temp2.fetchall()]
                            db_temp2.close()
                        phuong_xa_index = 0
                        if phuong_xa_current in phuong_xa_options:
                            phuong_xa_index = phuong_xa_options.index(phuong_xa_current) + 1
                        phuong_xa_chu_ho = st.selectbox("Phường/Xã (chủ hộ)", options=[""] + phuong_xa_options, index=phuong_xa_index)
                    
                    st.markdown("**Thông tin thường trú:**")
                    col_tt1, col_tt2 = st.columns(2)
                    with col_tt1:
                        tinh_thuong_tru_index = 0
                        tinh_thuong_tru_current = nv_data.get('Tinh_thanh_pho_thuong_tru', '') if nv_data else ''
                        if tinh_thuong_tru_current in tinh_options:
                            tinh_thuong_tru_index = list(tinh_options.keys()).index(tinh_thuong_tru_current) + 1
                        tinh_thuong_tru = st.selectbox("Tỉnh/Thành phố thường trú", options=[""] + list(tinh_options.keys()), index=tinh_thuong_tru_index)
                        ma_tinh_thuong_tru = tinh_options.get(tinh_thuong_tru, "") if tinh_thuong_tru else ""
                    with col_tt2:
                        phuong_xa_tt_options = []
                        phuong_xa_tt_current = nv_data.get('Phuong_xa_thuong_tru', '') if nv_data else ''
                        if tinh_thuong_tru and tinh_thuong_tru != "":
                            ma_tinh_tt = tinh_options.get(tinh_thuong_tru)
                            db_temp3 = get_connection()
                            c_temp3 = db_temp3.cursor()
                            c_temp3.execute("SELECT Ten_xa, Ma_xa FROM danh_muc_phuong_xa WHERE Ma_tinh = %s ORDER BY Ten_xa", (ma_tinh_tt,))
                            phuong_xa_tt_options = c_temp3.fetchall()
                            db_temp3.close()
                        phuong_xa_tt_index = 0
                        ma_phuong_xa_thuong_tru = ""
                        for i, px in enumerate(phuong_xa_tt_options):
                            if px[0] == phuong_xa_tt_current:
                                phuong_xa_tt_index = i + 1
                                ma_phuong_xa_thuong_tru = px[1]
                                break
                        phuong_xa_display_list = [""] + [px[0] for px in phuong_xa_tt_options]
                        phuong_xa_thuong_tru = st.selectbox("Phường/Xã thường trú", options=phuong_xa_display_list, index=phuong_xa_tt_index)
                        for px in phuong_xa_tt_options:
                            if px[0] == phuong_xa_thuong_tru:
                                ma_phuong_xa_thuong_tru = px[1]
                                break
                    
                    st.markdown("**II. THÊM THÀNH VIÊN MỚI:**")
                    st.caption("Điền thông tin vào các cột bên dưới, sau đó bấm '➕ Thêm thành viên'")
                    col_tv1, col_tv2, col_tv3, col_tv4, col_tv5, col_tv6, col_tv7, col_tv8 = st.columns([2,1.3,1,1,1,1.5,1.8,1.8])
                    with col_tv1:
                        ho_ten_tv = st.text_input("Họ và tên", key="tv_ho_ten_family", placeholder="Nguyễn Văn A")
                    with col_tv2:
                        ngay_sinh_tv = st.text_input("Ngày sinh", key="tv_ngay_sinh_family", placeholder="dd/mm/yyyy")
                    with col_tv3:
                        gioi_tinh_tv = st.selectbox("Giới tính", ["Nam", "Nữ"], key="tv_gioi_tinh_family")
                    with col_tv4:
                        quoc_tich_tv = st.text_input("Quốc tịch", value="Việt Nam", key="tv_quoc_tich_family")
                    with col_tv5:
                        dan_toc_tv = st.text_input("Dân tộc", value="Kinh", key="tv_dan_toc_family")
                    with col_tv6:
                        quan_he_tv = st.selectbox("Quan hệ chủ hộ", ["", "Vợ", "Chồng", "Con", "Bố", "Mẹ", "Anh", "Chị", "Em", "Ông", "Bà", "Khác"], key="tv_quan_he_family")
                    with col_tv7:
                        tinh_tv = st.selectbox("Tỉnh/Thành phố", options=[""] + list(tinh_options.keys()), key="tv_tinh_family")
                    with col_tv8:
                        phuong_xa_tv_options = []
                        if tinh_tv and tinh_tv != "":
                            ma_tinh_tv = tinh_options.get(tinh_tv)
                            db_temp4 = get_connection()
                            c_temp4 = db_temp4.cursor()
                            c_temp4.execute("SELECT Ten_xa FROM danh_muc_phuong_xa WHERE Ma_tinh = %s ORDER BY Ten_xa", (ma_tinh_tv,))
                            phuong_xa_tv_options = [row[0] for row in c_temp4.fetchall()]
                            db_temp4.close()
                        phuong_xa_tv = st.selectbox("Phường/Xã", options=[""] + phuong_xa_tv_options, key="tv_phuong_xa_family")
                    
                    col_btn_add1, col_btn_add2, col_btn_add3 = st.columns([1,1,1])
                    with col_btn_add2:
                        if st.form_submit_button("➕ Thêm thành viên vào danh sách", use_container_width=True):
                            if ho_ten_tv:
                                st.session_state.bhxh_family_members.append({
                                    'ho_ten': ho_ten_tv, 'ngay_sinh': parse_date(ngay_sinh_tv), 'gioi_tinh': gioi_tinh_tv,
                                    'quoc_tich': quoc_tich_tv, 'dan_toc': dan_toc_tv, 'quan_he': quan_he_tv,
                                    'tinh': tinh_tv, 'phuong_xa': phuong_xa_tv
                                })
                                st.rerun()
                            else:
                                st.error("Vui lòng nhập họ tên thành viên")
                    
                    st.markdown("---")
                    col_save1, col_save2, col_save3 = st.columns([1,2,1])
                    with col_save2:
                        if st.form_submit_button("💾 LƯU THÔNG TIN CHỦ HỘ", use_container_width=True, type="primary"):
                            try:
                                db = get_connection()
                                c = db.cursor()
                                c.execute("""UPDATE nhan_vien SET Ho_ten_chu_ho=%s, So_CCCD_chu_ho=%s, Tinh_thanh_pho_chu_ho=%s, Phuong_xa_chu_ho=%s,
                                    Tinh_thanh_pho_thuong_tru=%s, Ma_tinh_thuong_tru=%s, Phuong_xa_thuong_tru=%s, Ma_phuong_xa_thuong_tru=%s WHERE Id=%s""",
                                    (ho_ten_chu_ho, so_cccd_chu_ho, tinh_chu_ho, phuong_xa_chu_ho, tinh_thuong_tru, ma_tinh_thuong_tru, phuong_xa_thuong_tru, ma_phuong_xa_thuong_tru, nv_id))
                                c.execute("DELETE FROM phu_luc_gia_dinh WHERE nhan_vien_Id = %s", (nv_id,))
                                for tv in st.session_state.bhxh_family_members:
                                    c.execute("""INSERT INTO phu_luc_gia_dinh (nhan_vien_Id, Ho_ten, Ngay_sinh, Gioi_tinh, Quoc_tich, Dan_toc, Quan_he_voi_chu_ho, Tinh_thanh_pho, Phuong_xa) 
                                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                                        (nv_id, tv['ho_ten'], tv['ngay_sinh'], tv['gioi_tinh'], tv['quoc_tich'], tv['dan_toc'], tv['quan_he'], tv['tinh'], tv['phuong_xa']))
                                db.commit()
                                db.close()
                                del st.session_state['bhxh_family_nv_id']
                                del st.session_state['bhxh_family_nv_name']
                                del st.session_state['bhxh_family_members']
                                st.success(f"✅ Đã lưu thông tin hộ gia đình cho nhân viên {nv_name}")
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi khi lưu: {e}")
                
                col_cancel1, col_cancel2, col_cancel3 = st.columns([1,2,1])
                with col_cancel2:
                    if st.button("❌ HỦY BỎ", use_container_width=True):
                        del st.session_state['bhxh_family_nv_id']
                        del st.session_state['bhxh_family_nv_name']
                        if 'bhxh_family_members' in st.session_state:
                            del st.session_state['bhxh_family_members']
                        st.rerun()
        
        else:
            st.info("Không có nhân viên nào đang làm việc")
    
    # ==================== TAB 2: ĐÃ NGHỈ VIỆC ====================
    with tab_da_nghi:
        st.caption("📋 Danh sách nhân viên đã nghỉ việc (có thông tin ngày nghỉ)")
        
        # Bộ lọc và tìm kiếm
        col_filter1, col_filter2, col_filter3 = st.columns([2, 2, 1])
        with col_filter1:
            search_nghi = st.text_input("🔍 Tìm kiếm (Tên, Mã NV, SĐT, CCCD)", key="search_da_nghi")
        with col_filter2:
            # Lọc theo năm nghỉ
            db_temp = get_connection()
            c_temp = db_temp.cursor()
            c_temp.execute("SELECT DISTINCT YEAR(Ngay_ket_thuc) as nam FROM nhan_vien WHERE Trang_thai='NGHI_VIEC' AND Ngay_ket_thuc IS NOT NULL ORDER BY nam DESC")
            years = [row[0] for row in c_temp.fetchall()]
            db_temp.close()
            filter_nam = st.selectbox("📅 Lọc theo năm nghỉ:", ["Tất cả"] + [str(y) for y in years] if years else ["Tất cả"])
        
        # Truy vấn danh sách nhân viên đã nghỉ
        db = get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        sql = """
            SELECT Id, Ma_NV, Ho_ten, Ngay_sinh, Gioi_tinh, So_CCCD, Dien_thoai, 
                   Chuc_danh_nghe, Loai_hop_dong, So_HDLD, Ngay_vao_lam, Ngay_ket_thuc,
                   Ma_so_BHXH, Thang_bat_dau_BH, Ly_do_nghi
            FROM nhan_vien 
            WHERE Trang_thai = 'NGHI_VIEC'
        """
        params = []
        
        if search_nghi:
            sql += " AND (Ho_ten LIKE %s OR Ma_NV LIKE %s OR Dien_thoai LIKE %s OR So_CCCD LIKE %s)"
            params.extend([f'%{search_nghi}%'] * 4)
        
        if filter_nam != "Tất cả" and filter_nam.isdigit():
            sql += " AND YEAR(Ngay_ket_thuc) = %s"
            params.append(int(filter_nam))
        
        sql += " ORDER BY Ngay_ket_thuc DESC, Id DESC"
        c.execute(sql, tuple(params))
        ds_nghi = c.fetchall()
        db.close()
        
        if ds_nghi:
            # Chuẩn bị dataframe
            df_nghi = pd.DataFrame(ds_nghi)
            for col in df_nghi.columns:
                if 'Ngay' in col:
                    df_nghi[col] = df_nghi[col].apply(format_date)
            
            # Chọn cột hiển thị
            display_cols_nghi = ['ma_nv', 'ho_ten', 'ngay_sinh', 'gioi_tinh', 'chuc_danh_nghe', 
                                 'so_hdld', 'ngay_vao_lam', 'ngay_ket_thuc', 'dien_thoai', 'ma_so_bhxh']
            available_cols_nghi = [c for c in display_cols_nghi if c in df_nghi.columns]
            df_show_nghi = df_nghi[available_cols_nghi]
            
            col_map_nghi = {
                'ma_nv': 'Mã NV',
                'ho_ten': 'Họ và tên',
                'ngay_sinh': 'Ngày sinh',
                'gioi_tinh': 'Giới tính',
                'chuc_danh_nghe': 'Chức danh',
                'so_hdld': 'Số HĐLĐ',
                'ngay_vao_lam': 'Ngày vào làm',
                'ngay_ket_thuc': '📅 Ngày nghỉ việc',
                'dien_thoai': 'SĐT',
                'ma_so_bhxh': 'Mã BHXH',
            }
            df_show_nghi.rename(columns=col_map_nghi, inplace=True)
            
            st.caption(f"📌 Tổng số: **{len(ds_nghi)}** nhân viên đã nghỉ việc")
            
            # Hiển thị bảng
            st.dataframe(df_show_nghi, use_container_width=True, hide_index=True, height=400)
            
            # Chi tiết và khôi phục
            st.divider()
            st.subheader("🔍 Xem chi tiết / Khôi phục nhân viên")
            
            # Chọn nhân viên để xem chi tiết
            nv_options = {f"{nv['ma_nv']} - {nv['ho_ten']} (Nghỉ: {format_date(nv.get('ngay_ket_thuc'))})": nv['id'] for nv in ds_nghi}
            selected_nghi_name = st.selectbox("Chọn nhân viên đã nghỉ:", list(nv_options.keys()))
            selected_nghi_id = nv_options[selected_nghi_name]
            
            # Lấy chi tiết nhân viên
            db = get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c.execute("SELECT * FROM nhan_vien WHERE Id = %s", (selected_nghi_id,))
            nv_nghi_detail = c.fetchone()
            db.close()
            
            if nv_nghi_detail:
                col_detail1, col_detail2 = st.columns(2)
                with col_detail1:
                    st.markdown("**📋 Thông tin cá nhân**")
                    st.write(f"- **Mã NV:** {nv_nghi_detail.get('ma_nv', '')}")
                    st.write(f"- **Họ tên:** {nv_nghi_detail.get('ho_ten', '')}")
                    st.write(f"- **Ngày sinh:** {format_date(nv_nghi_detail.get('ngay_sinh'))}")
                    st.write(f"- **Giới tính:** {nv_nghi_detail.get('gioi_tinh', '')}")
                    st.write(f"- **CCCD:** {nv_nghi_detail.get('so_cccd', '')}")
                    st.write(f"- **SĐT:** {nv_nghi_detail.get('dien_thoai', '')}")
                    st.write(f"- **Chức danh:** {nv_nghi_detail.get('chuc_danh_nghe', '')}")
                
                with col_detail2:
                    st.markdown("**📅 Thông tin hợp đồng & nghỉ việc**")
                    st.write(f"- **Số HĐLĐ:** {nv_nghi_detail.get('so_hdld', '')}")
                    st.write(f"- **Loại HĐ:** {nv_nghi_detail.get('loai_hop_dong', '')}")
                    st.write(f"- **Ngày vào làm:** {format_date(nv_nghi_detail.get('ngay_vao_lam'))}")
                    st.write(f"- **📅 Ngày nghỉ việc:** **{format_date(nv_nghi_detail.get('ngay_ket_thuc'))}**")
                    st.write(f"- **Mã BHXH:** {nv_nghi_detail.get('ma_so_bhxh', '')}")
                    st.write(f"- **Lý do nghỉ:** {nv_nghi_detail.get('Ly_do_nghi', 'Chưa có thông tin')}")
                
                # Nút khôi phục (chỉ admin)
                if st.session_state.role == "admin":
                    st.divider()
                    col_restore1, col_restore2, col_restore3 = st.columns([1, 2, 1])
                    with col_restore2:
                        if st.button(f"🔄 KHÔI PHỤC NHÂN VIÊN - {nv_nghi_detail['ho_ten']}", 
                                     use_container_width=True, type="primary"):
                            try:
                                db = get_connection()
                                c = db.cursor()
                                # Cập nhật trạng thái về DANG_LAM hoặc THU_VIEC tùy theo loại hợp đồng
                                loai_hd = nv_nghi_detail.get('loai_hop_dong', '')
                                if loai_hd == 'Thử việc':
                                    trang_thai_moi = 'THU_VIEC'
                                else:
                                    trang_thai_moi = 'DANG_LAM'
                                c.execute("""
                                    UPDATE nhan_vien 
                                    SET Trang_thai = %s, 
                                        Ngay_ket_thuc = NULL
                                    WHERE Id = %s
                                """, (trang_thai_moi, selected_nghi_id))
                                db.commit()
                                db.close()
                                st.success(f"✅ Đã khôi phục nhân viên {nv_nghi_detail['ho_ten']}!")
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Lỗi khi khôi phục: {e}")
        else:
            st.info("📭 Không có nhân viên nào đã nghỉ việc")
    
    # ==================== TAB 3: LỊCH SỬ CÔNG TÁC ====================
    with tab_qtct:
        st.caption("📜 Lịch sử công tác và quyết định nhân sự")
           
        # Chọn nhân viên để xem lịch sử
        db = get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT Id, Ma_NV, Ho_ten FROM nhan_vien ORDER BY Id DESC")
        all_nv = c.fetchall()
        db.close()
        
        if all_nv:
            nv_options = {f"{x['ma_nv']} - {x['ho_ten']}": x['id'] for x in all_nv}
            selected_nv_history = st.selectbox("🔍 Chọn nhân viên:", list(nv_options.keys()), key="history_nv")
            nv_id_history = nv_options[selected_nv_history]
            
            # Lấy thông tin nhân viên
            db = get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c.execute("SELECT * FROM nhan_vien WHERE Id = %s", (nv_id_history,))
            nv_current = c.fetchone()
            
            # Hiển thị thông tin hiện tại
            st.markdown(f"""
            ### 📌 Thông tin hiện tại của {nv_current['ho_ten']} ({nv_current['ma_nv']})
            | Trường | Giá trị |
            |--------|---------|
            | Trạng thái | {'🟢 Đang làm' if nv_current['trang_thai'] == 'DANG_LAM' else '🔵 Thử việc' if nv_current['trang_thai'] == 'THU_VIEC' else '🔴 Đã nghỉ'} |
            | Loại hợp đồng | {nv_current['loai_hop_dong']} |
            | Ngày vào làm | {format_date(nv_current['ngay_vao_lam'])} |
            | Ngày chính thức | {format_date(nv_current['Ngay_chinh_thuc']) or 'Chưa có'} |
            | Chức danh | {nv_current['chuc_danh_nghe']} |
            | Phòng ban | {nv_current['Phong_ban_lam_viec']} |
            """)
            
            # Lấy danh sách quyết định
            c.execute("""
                SELECT * FROM quyet_dinh_nhan_su 
                WHERE nhan_vien_Id = %s 
                ORDER BY Ngay_quyet_dinh DESC
            """, (nv_id_history,))
            quyet_dinh_list = c.fetchall()
            
            if quyet_dinh_list:
                st.markdown("### 📋 Các quyết định nhân sự")
                
                qd_data = []
                for i, qd in enumerate(quyet_dinh_list, 1):
                    loai_qd_map = {
                        'THU_VIEC': '📝 Quyết định thử việc',
                        'CHINH_THUC': '✅ Quyết định chính thức',
                        'DIEU_CHUYEN': '🔄 Quyết định điều chuyển',
                        'BO_NHIEM': '⭐ Quyết định bổ nhiệm',
                        'TANG_LUONG': '💰 Quyết định tăng lương',
                        'NGHI_VIEC': '🚫 Quyết định nghỉ việc'
                    }
                    qd_data.append({
                        "STT": i,
                        "Loại quyết định": loai_qd_map.get(qd['Loai_quyet_dinh'], qd['Loai_quyet_dinh']),
                        "Số quyết định": qd['So_quyet_dinh'] or '...',
                        "Ngày quyết định": format_date(qd['Ngay_quyet_dinh']),
                        "Ngày hiệu lực": format_date(qd['Ngay_hieu_luc']),
                        "Nội dung": qd['Noi_dung'][:50] + "..." if qd['Noi_dung'] and len(qd['Noi_dung']) > 50 else qd['Noi_dung']
                    })
                
                df_qd = pd.DataFrame(qd_data)
                st.dataframe(df_qd, use_container_width=True, hide_index=True)
                
                # Xem chi tiết quyết định
                with st.expander("🔍 Xem chi tiết quyết định"):
                    qd_options = {f"{format_date(qd['Ngay_quyet_dinh'])} - {qd['Loai_quyet_dinh']}": qd for qd in quyet_dinh_list}
                    selected_qd_name = st.selectbox("Chọn quyết định:", list(qd_options.keys()), key="qd_detail")
                    selected_qd = qd_options[selected_qd_name]
                    
                    st.markdown(f"""
                    **📄 Chi tiết quyết định:**
                    - **Số quyết định:** {selected_qd.get('So_quyet_dinh', '...')}
                    - **Ngày quyết định:** {format_date(selected_qd.get('Ngay_quyet_dinh'))}
                    - **Ngày hiệu lực:** {format_date(selected_qd.get('Ngay_hieu_luc'))}
                    - **Loại quyết định:** {selected_qd.get('Loai_quyet_dinh')}
                    - **Nội dung:** {selected_qd.get('Noi_dung', '...')}
                    - **Người ký:** {selected_qd.get('Nguoi_ky', COMPANY_CONFIG.get('dai_dien', 'GIÁM ĐỐC'))}
                    
                    **📊 Thay đổi:**
                    - Chức danh: {selected_qd.get('Chuc_danh_cu', '...')} → {selected_qd.get('Chuc_danh_moi', '...')}
                    - Phòng ban: {selected_qd.get('phong_ban_cu', '...')} → {selected_qd.get('phong_ban_moi', '...')}
                    - Loại HĐ: {selected_qd.get('Loai_hop_dong_cu', '...')} → {selected_qd.get('Loai_hop_dong_moi', '...')}
                    """)
            else:
                st.info("📭 Nhân viên này chưa có quyết định nào.")
            
            # Lấy lịch sử công tác
            c.execute("""
                SELECT * FROM lich_su_cong_tac 
                WHERE nhan_vien_Id = %s 
                ORDER BY Tu_ngay ASC
            """, (nv_id_history,))
            lich_su_list = c.fetchall()
            
            if lich_su_list:
                st.markdown("### 📅 Lịch sử công tác")
                
                ls_data = []
                for i, ls in enumerate(lich_su_list, 1):
                    ls_data.append({
                        "STT": i,
                        "Từ ngày": format_date(ls['Tu_ngay']),
                        "Đến ngày": format_date(ls['Den_ngay']) if ls['Den_ngay'] else "Đang làm",
                        "Chức danh": ls['Chuc_danh'] or '',
                        "Phòng ban": ls['Phong_ban'] or '',
                        "Loại HĐ": ls['loai_hop_dong'] or '',
                        "Hệ số lương": ls['He_so_luong'] or ''
                    })
                
                df_ls = pd.DataFrame(ls_data)
                st.dataframe(
                df_ls, 
                use_container_width=True, 
                hide_index=True,
                height=400,
                column_config={
                    "Từ ngày": st.column_config.DateColumn("Từ ngày", format="DD/MM/YYYY"),
                    "Đến ngày": st.column_config.DateColumn("Đến ngày", format="DD/MM/YYYY"),
                    "Loại HĐ": st.column_config.TextColumn("Loại HĐ"),
                    "Hệ số lương": st.column_config.NumberColumn("Hệ số lương", format="%.2f"),
                }
            )
            else:
                # Nếu chưa có lịch sử, tạo từ dữ liệu hiện tại
                st.info("📭 Chưa có lịch sử công tác. Đang tạo từ dữ liệu hiện tại...")
                
                # Xác định loại hợp đồng đúng theo trạng thái
                loai_hd_dung = nv_current['loai_hop_dong']
                if nv_current['trang_thai'] == 'THU_VIEC':
                    loai_hd_dung = 'Thử việc'
                
                c.execute("""
                    INSERT INTO lich_su_cong_tac (nhan_vien_Id, Tu_ngay, Chuc_danh, phong_ban, Noi_lam_viec, Loai_hop_dong, He_so_luong)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (nv_id_history, nv_current['ngay_vao_lam'], nv_current['chuc_danh_nghe'], 
                      nv_current['phong_ban_lam_viec'], nv_current['Noi_lam_viec'], 
                      loai_hd_dung, nv_current['He_so_luong']))
                db.commit()
                st.rerun()
            
            db.close()
        else:
            st.info("⚠️ Chưa có nhân viên nào trong hệ thống!")
    
    # ===== BÁO CÁO TĂNG/GIẢM NHÂN SỰ (giữ nguyên) =====
    st.divider()
    st.subheader("📊 Báo cáo tăng/giảm nhân sự trong kỳ")
    
    col_from, col_to, col_btn = st.columns([2, 2, 1])
    with col_from:
        tu_ngay_bc = st.date_input("Từ ngày:", value=date.today().replace(day=1), key="bc_tu")
    with col_to:
        den_ngay_bc = st.date_input("Đến ngày:", value=date.today(), key="bc_den")
    with col_btn:
        xuat_bc = st.button("📄 XUẤT BÁO CÁO WORD", use_container_width=True)
    
    if xuat_bc:
        db = get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("""
            SELECT Ho_ten, Chuc_danh_nghe, phong_ban_lam_viec, Loai_hop_dong, Ngay_vao_lam,
                   Ngay_sinh, So_HDLD, Ngay_ky_HD
            FROM nhan_vien 
            WHERE Trang_thai IN ('DANG_LAM', 'THU_VIEC')
            AND Ngay_vao_lam BETWEEN %s AND %s
            ORDER BY Ngay_vao_lam ASC
        """, (tu_ngay_bc, den_ngay_bc))
        tang_list = c.fetchall()
        c.execute("""
            SELECT Ho_ten, Chuc_danh_nghe, phong_ban_lam_viec, Loai_hop_dong, Ngay_vao_lam, Ngay_ket_thuc,
                   Ngay_sinh, So_HDLD, Ngay_ky_HD
            FROM nhan_vien 
            WHERE Trang_thai = 'NGHI_VIEC'
            AND Ngay_ket_thuc BETWEEN %s AND %s
            ORDER BY Ngay_ket_thuc ASC
        """, (tu_ngay_bc, den_ngay_bc))
        giam_list = c.fetchall()
        db.close()
        if tang_list or giam_list:
            file_path = tao_bao_cao_tang_giam(tang_list, giam_list, tu_ngay_bc, den_ngay_bc)
            with open(file_path, "rb") as f:
                st.download_button(
                    label="📥 TẢI FILE BÁO CÁO (Word)",
                    data=f,
                    file_name=f"Bao_cao_tang_giam_{tu_ngay_bc.strftime('%d%m%Y')}_{den_ngay_bc.strftime('%d%m%Y')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.info("Không có biến động nhân sự trong kỳ.")
        
        

# ========== UPLOAD ==========
elif menu=="📁 Upload hồ sơ" and st.session_state.role=="admin":
    st.title("📁 Quản lý hồ sơ nhân viên")
    
    # Tạo 2 tab: Upload và Danh sách hồ sơ
    tab_upload, tab_list = st.tabs(["📤 UPLOAD HỒ SƠ", "📋 DANH SÁCH HỒ SƠ"])
    
    # ===== TAB 1: UPLOAD HỒ SƠ =====
    with tab_upload:
        db = get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT Id, Ma_NV, Ho_ten FROM nhan_vien WHERE Trang_thai IN ('DANG_LAM','THU_VIEC') ORDER BY Id DESC")
        nvl = c.fetchall()
        db.close()
        
        if nvl:
            nd = {f"{x['ma_nv']} - {x['ho_ten']}": x['id'] for x in nvl}
            cn = st.selectbox("📌 Chọn nhân viên:", list(nd.keys()))
            
            col1, col2 = st.columns(2)
            with col1:
                lh = st.selectbox("📂 Loại hồ sơ:", ["BANG_CAP", "CHUNG_CHI", "CCCD", "HOP_DONG", "SO_YEU_LY_LICH", "KHAC"])
            with col2:
                st.markdown("💡 **Hướng dẫn:**")
                st.caption("- BANG_CAP: Bằng cấp, chứng chỉ")
                st.caption("- CHUNG_CHI: Chứng chỉ nghề")
                st.caption("- CCCD: Căn cước công dân")
                st.caption("- HOP_DONG: Hợp đồng lao động")
                st.caption("- SO_YEU_LY_LICH: Sơ yếu lý lịch")
            
            fl = st.file_uploader("📎 Chọn file:", type=['pdf', 'jpg', 'png', 'jpeg', 'doc', 'docx'])
            
            # Hiển thị thông tin file đã chọn
            if fl:
                st.info(f"📄 Tên file: {fl.name} | 📏 Kích thước: {fl.size/1024:.1f} KB")
            
            if fl and st.button("📤 UPLOAD", type="primary", use_container_width=True):
                nid = nd[cn]
                # Tạo tên file duy nhất: ID_NgayGio_Loai_TenGoc
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                fn = f"{nid}_{timestamp}_{lh}_{fl.name}"
                fp = os.path.join(UPLOAD_FOLDER, fn)
                
                # Lưu file
                with open(fp, "wb") as f:
                    f.write(fl.getbuffer())
                
                # Lưu vào database
                db = get_connection()
                c = db.cursor()
                c.execute("""
                    INSERT INTO ho_so_nhan_vien (nhan_vien_Id, Loai_ho_so, Ten_file, Duong_dan_file, Ngay_upload) 
                    VALUES (%s, %s, %s, %s, CURRENT_DATE)
                """, (nid, lh, fl.name, fp))
                db.commit()
                db.close()
                
                st.success(f"✅ Đã upload thành công!\n📁 Lưu tại: {fp}")
                st.rerun()
        else:
            st.info("⚠️ Chưa có nhân viên nào trong hệ thống!")
    
    # ===== TAB 2: DANH SÁCH HỒ SƠ =====
    with tab_list:
        st.subheader("📋 Danh sách hồ sơ đã upload")
        
        # Chọn nhân viên để xem hồ sơ
        db = get_connection()
        c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT Id, Ma_NV, Ho_ten FROM nhan_vien ORDER BY Id DESC")
        nvl = c.fetchall()
        db.close()
        
        if nvl:
            nd = {f"{x['ma_nv']} - {x['ho_ten']}": x['id'] for x in nvl}
            selected_nv = st.selectbox("🔍 Chọn nhân viên để xem hồ sơ:", list(nd.keys()), key="view_hoso")
            nv_id = nd[selected_nv]
            
            # Lấy danh sách hồ sơ
            db = get_connection()
            c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            c.execute("""
                SELECT Id, Loai_ho_so, Ten_file, Duong_dan_file, Ngay_upload 
                FROM ho_so_nhan_vien 
                WHERE nhan_vien_Id = %s 
                ORDER BY Ngay_upload DESC, Id DESC
            """, (nv_id,))
            hs_list = c.fetchall()
            db.close()
            
            if hs_list:
                st.caption(f"📌 Tổng số: **{len(hs_list)}** hồ sơ")
                
                # Hiển thị danh sách dạng bảng
                hs_data = []
                for i, hs in enumerate(hs_list, 1):
                    hs_data.append({
                        "STT": i,
                        "Loại hồ sơ": hs['Loai_ho_so'],
                        "Tên file gốc": hs['Ten_file'],
                        "Ngày upload": format_date(hs['Ngay_upload']),
                        "ID": hs['id'],
                        "Đường dẫn": hs['Duong_dan_file']
                    })
                
                df_hs = pd.DataFrame(hs_data)
                st.dataframe(df_hs[['STT', 'Loại hồ sơ', 'Tên file gốc', 'Ngày upload']], 
                           use_container_width=True, hide_index=True)
                
                # Chi tiết và tải xuống
                st.divider()
                st.subheader("📄 Xem chi tiết & Tải xuống")
                
                hs_options = {f"{hs['Loai_ho_so']} - {hs['Ten_file']} (Ngày: {format_date(hs['Ngay_upload'])})": hs for hs in hs_list}
                selected_hs_name = st.selectbox("Chọn hồ sơ:", list(hs_options.keys()))
                selected_hs = hs_options[selected_hs_name]
                
                col_info, col_download = st.columns([2, 1])
                with col_info:
                    st.markdown(f"""
                    **📋 Thông tin hồ sơ:**
                    - **Loại:** {selected_hs['Loai_ho_so']}
                    - **Tên file gốc:** {selected_hs['Ten_file']}
                    - **Đường dẫn:** `{selected_hs['Duong_dan_file']}`
                    - **Ngày upload:** {format_date(selected_hs['Ngay_upload'])}
                    """)
                
                with col_download:
                    # Kiểm tra file có tồn tại không
                    if os.path.exists(selected_hs['Duong_dan_file']):
                        with open(selected_hs['Duong_dan_file'], "rb") as f:
                            st.download_button(
                                label="📥 TẢI HỒ SƠ",
                                data=f,
                                file_name=selected_hs['Ten_file'],
                                mime="application/octet-stream",
                                use_container_width=True
                            )
                    else:
                        st.error("❌ File không tồn tại trên máy chủ!")
                
                # Nút xóa hồ sơ
                st.divider()
                col_del1, col_del2, col_del3 = st.columns([1, 2, 1])
                with col_del2:
                    if st.button("🗑️ XÓA HỒ SƠ NÀY", use_container_width=True, type="secondary"):
                        try:
                            # Xóa file vật lý
                            if os.path.exists(selected_hs['Duong_dan_file']):
                                os.remove(selected_hs['Duong_dan_file'])
                            
                            # Xóa record trong database
                            db = get_connection()
                            c = db.cursor()
                            c.execute("DELETE FROM ho_so_nhan_vien WHERE Id = %s", (selected_hs['id'],))
                            db.commit()
                            db.close()
                            
                            st.success(f"✅ Đã xóa hồ sơ: {selected_hs['Ten_file']}")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Lỗi khi xóa: {e}")
            else:
                st.info(f"📭 Nhân viên này chưa có hồ sơ nào được upload.")
        else:
            st.info("⚠️ Chưa có nhân viên nào trong hệ thống!")

# ========== DANH MỤC CHỨC DANH ==========
elif menu == "⚙️ Danh mục" and st.session_state.role == "admin":
    st.title("⚙️ Quản lý danh mục Chức danh")
    with st.expander("➕ Thêm chức danh mới", expanded=False):
        with st.form("add_chuc_danh"):
            ten_moi = st.text_input("Tên chức danh *"); mo_ta = st.text_area("Mô tả")
            if st.form_submit_button("💾 LƯU"):
                if ten_moi:
                    db = get_connection(); c = db.cursor()
                    c.execute("SELECT COALESCE(MIN(t1.Id + 1), 1) FROM vi_tri_cong_tac t1 LEFT JOIN vi_tri_cong_tac t2 ON t1.Id + 1 = t2.Id WHERE t2.Id IS NULL AND t1.Id >= 1"); id_trong = c.fetchone()[0]
                    c.execute("SELECT COALESCE(MAX(Id),0) FROM vi_tri_cong_tac"); id_max = c.fetchone()[0]; id_moi = id_trong if id_trong <= id_max + 1 else id_max + 1
                    c.execute("INSERT INTO vi_tri_cong_tac (Id, Ten_vi_tri, Ghi_chu) VALUES (%s, %s, %s)", (id_moi, ten_moi, mo_ta))
                    db.commit(); db.close(); st.success(f"✅ Đã thêm: {ten_moi}"); st.rerun()
                else: st.error("Tên chức danh không được để trống!")
    st.subheader("📋 Danh sách chức danh")
    db = get_connection(); c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    c.execute("SELECT Id, Ten_vi_tri, Ghi_chu FROM vi_tri_cong_tac ORDER BY Id"); ds = c.fetchall(); db.close()
    if ds:
        df = pd.DataFrame(ds); df.columns = ['id', 'Tên chức danh', 'Ghi chú']; st.dataframe(df, use_container_width=True, hide_index=True)
        st.divider(); cdx = st.number_input("Nhập ID cần xóa:", min_value=1, step=1)
        if st.button("🗑️ XÓA", key="del_cd"):
            db = get_connection(); c = db.cursor()
            c.execute("DELETE FROM vi_tri_cong_tac WHERE Id=%s", (cdx,)); db.commit(); db.close(); st.success("🗑️ Đã xóa!"); st.rerun()
    else: st.info("Chưa có chức danh nào")

# ========== BHXH ==========
elif menu=="📋 BHXH" and st.session_state.role=="admin":
    st.title("📋 BHXH")
    t1,t2=st.tabs(["📊 Tổng quan","📝 Báo cáo"])
    with t1:
        db=get_connection(); c=db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        c.execute("SELECT COUNT(*) t FROM nhan_vien WHERE Trang_thai='DANG_LAM'"); tong=c.fetchone()['t']
        c.execute("SELECT COUNT(*) t FROM nhan_vien WHERE Trang_thai='DANG_LAM' AND Trang_thai_BHXH='DANG_DONG'"); dang=c.fetchone()['t']
        c.execute("SELECT COUNT(*) t FROM nhan_vien WHERE Trang_thai='DANG_LAM' AND Trang_thai_BHXH='CHUA_DONG'"); chua=c.fetchone()['t']
        db.close(); cl1,cl2,cl3=st.columns(3)
        cl1.metric("Tổng LĐ",tong); cl2.metric("Đang đóng",dang); cl3.metric("Chưa đóng",chua)
    with t2:
        st.subheader("📝 Báo cáo tăng/giảm lao động (Mẫu D02-LT)")
        col_from, col_to = st.columns(2)
        with col_from: 
            tu_ngay = st.date_input("Từ ngày:", value=date(date.today().year, 1, 1), key="d02_tu")
        with col_to: 
            den_ngay = st.date_input("Đến ngày:", value=date.today(), key="d02_den")
        
        db = get_connection(); c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        # Lấy danh sách lao động tăng
        c.execute("""SELECT STT, Ma_NV, Ho_ten, Ma_so_BHXH, 'Tăng lao động' as Loai_PA, 
                    CASE WHEN Gioi_tinh='Nam' THEN 1 ELSE 0 END as Gioi_tinh_ma, 
                    Ngay_sinh, So_CCCD, Chuc_danh_nghe, phong_ban_lam_viec, Noi_lam_viec, 
                    Luong_bao_hiem, COALESCE(Thang_bat_dau_BH, Ngay_vao_lam) as Ngay_bat_dau, 
                    Loai_hop_dong, Quoc_tich, Dan_toc, Dien_thoai, Email_lien_he, 
                    Tinh_nhan_HS, Phuong_nhan_HS, Dia_chi_nhan_HS, Tinh_KCB, Noi_dang_ky_KCB, 
                    Dang_ky_nhan_so, He_so_luong, Phu_cap_chuc_vu, Phu_cap_TNVK, Phu_cap_TNN, 
                    Muc_huong_BHYT, Ty_le_dong, Muc_tien_dong, Phuong_thuc_dong, Trang_thai_BHXH,
                    So_HDLD, Ngay_ky_HD, Ngay_ket_thuc, Ngay_vao_lam, Thuong_tru
            FROM nhan_vien WHERE Trang_thai = 'DANG_LAM' 
            AND COALESCE(Thang_bat_dau_BH, Ngay_vao_lam) BETWEEN %s AND %s 
            ORDER BY COALESCE(Thang_bat_dau_BH, Ngay_vao_lam) ASC""", (tu_ngay, den_ngay))
        tang = c.fetchall()
        
        # Lấy danh sách lao động giảm
        c.execute("""SELECT STT, Ma_NV, Ho_ten, Ma_so_BHXH, 'Giảm lao động' as Loai_PA, 
                    CASE WHEN Gioi_tinh='Nam' THEN 1 ELSE 0 END as Gioi_tinh_ma, 
                    Ngay_sinh, So_CCCD, Chuc_danh_nghe, phong_ban_lam_viec, Noi_lam_viec, 
                    Luong_bao_hiem, Thang_ket_thuc_BH as Ngay_ket_thuc_BH, Loai_hop_dong, 
                    Quoc_tich, Dan_toc, Dien_thoai, Email_lien_he, Tinh_nhan_HS, Phuong_nhan_HS, 
                    Dia_chi_nhan_HS, Tinh_KCB, Noi_dang_ky_KCB, Dang_ky_nhan_so, He_so_luong, 
                    Phu_cap_chuc_vu, Phu_cap_TNVK, Phu_cap_TNN, Muc_huong_BHYT, Ty_le_dong, 
                    Muc_tien_dong, Phuong_thuc_dong, So_HDLD, Ngay_ky_HD, Ngay_vao_lam, Thuong_tru
            FROM nhan_vien WHERE Trang_thai = 'NGHI_VIEC' 
            AND Thang_ket_thuc_BH BETWEEN %s AND %s 
            ORDER BY Thang_ket_thuc_BH ASC""", (tu_ngay, den_ngay))
        giam = c.fetchall()
        db.close()
        
        col_t, col_g = st.columns(2)
        with col_t:
            st.subheader(f"🟢 TĂNG ({len(tang)})")
            if tang:
                df_t = pd.DataFrame(tang)
                for col in df_t.columns:
                    if 'Ngay' in col: df_t[col] = df_t[col].apply(format_date)
                st.dataframe(df_t, use_container_width=True, hide_index=True, height=400)
            else: st.info("Không có LĐ tăng")
        with col_g:
            st.subheader(f"🔴 GIẢM ({len(giam)})")
            if giam:
                df_g = pd.DataFrame(giam)
                for col in df_g.columns:
                    if 'Ngay' in col: df_g[col] = df_g[col].apply(format_date)
                st.dataframe(df_g, use_container_width=True, hide_index=True, height=400)
            else: st.info("Không có LĐ giảm")
        
        if tang or giam:
            st.divider()
            if st.button("📥 XUẤT EXCEL D02-LT (MẪU ĐẦY ĐỦ)", type="primary", use_container_width=True):
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment, Border, Side
                from openpyxl.utils import get_column_letter
                
                # Định nghĩa border
                thin_border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
                
                ten_cong_ty = COMPANY_CONFIG.get("ten_cong_ty", "CÔNG TY CỔ PHẦN CẢNG HÒN LA")
                
                # Tạo workbook cho sheet tổng hợp D02-LT
                wb = Workbook()
                ws = wb.active
                ws.title = "D02-LT"
                
                # ===== HEADER =====
                # Dòng 1: Table 1 merge A-F, Table 2 merge T-Z
                ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
                ws['A1'] = ten_cong_ty
                ws['A1'].font = Font(bold=True, size=13, name='Times New Roman')
                ws['A1'].alignment = Alignment(horizontal='center')
                
                ws.merge_cells(start_row=1, start_column=20, end_row=1, end_column=26)
                ws['DD1'] = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
                ws['DD1'].font = Font(bold=True, size=13, name='Times New Roman')
                ws['DD1'].alignment = Alignment(horizontal='center')
                
                # Dòng 2: Table 1 merge A-F, Table 2 merge T-Z
                ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=6)
                ws['A2'] = f"Số: 01/BC-D02LT-{datetime.now().year}/CHL"
                ws['A2'].font = Font(size=12, name='Times New Roman')
                ws['A2'].alignment = Alignment(horizontal='center')
                
                ws.merge_cells(start_row=2, start_column=20, end_row=2, end_column=26)
                ws['DD2'] = "Độc lập - Tự do - Hạnh phúc"
                ws['DD2'].font = Font(italic=True, size=12, name='Times New Roman')
                ws['DD2'].alignment = Alignment(horizontal='center')
                
                # Dòng 3: merge T-Z, canh phải
                ws.merge_cells(start_row=3, start_column=20, end_row=3, end_column=26)
                ws['DD3'] = f"Quảng Trị, ngày {date.today().day} tháng {date.today().month} năm {date.today().year}"
                ws['DD3'].font = Font(italic=True, size=12, name='Times New Roman')
                ws['DD3'].alignment = Alignment(horizontal='right')
                
                # Dòng 5: Tiêu đề báo cáo
                ws.merge_cells('A5:DH5')
                ws['A5'] = "DANH SÁCH LAO ĐỘNG THAM GIA BHXH; BHYT, BHTN, BHTNLĐ, BNN (Mẫu D02-LT TK1)"
                ws['A5'].font = Font(bold=True, size=13, name='Times New Roman')
                ws['A5'].alignment = Alignment(horizontal='center')
                
                ws.merge_cells('A6:DH6')
                ws['A6'] = f"(Từ ngày {tu_ngay.strftime('%d/%m/%Y')} đến ngày {den_ngay.strftime('%d/%m/%Y')})"
                ws['A6'].font = Font(size=12, name='Times New Roman')
                ws['A6'].alignment = Alignment(horizontal='center')
                
                # ===== HEADER BẢNG =====
                header_row = 8
                
                # Độ rộng các cột (A đến DH = 114 cột)
                for col in range(1, 115):
                    ws.column_dimensions[get_column_letter(col)].width = 12
                
                # Các cột đặc biệt cần rộng hơn
                special_cols = {'A': 6, 'B': 25, 'C': 18, 'D': 15, 'E': 12, 'F': 14, 'G': 15, 'H': 10, 
                               'I': 20, 'J': 30, 'K': 20, 'L': 20, 'M': 15, 'N': 15, 'O': 18,
                               'P': 12, 'Q': 12, 'R': 15, 'S': 15, 'T': 18, 'U': 10, 'V': 15, 
                               'W': 15, 'X': 12, 'Y': 12, 'Z': 15, 'AA': 15, 'AB': 15, 'AC': 12,
                               'AD': 15, 'AE': 18, 'AF': 12, 'AG': 12, 'AH': 10, 'AI': 18, 'AJ': 15,
                               'AK': 15, 'AL': 15, 'AM': 15, 'AN': 15, 'AO': 15, 'AP': 15, 'AQ': 12,
                               'AR': 12, 'AS': 15, 'AT': 10, 'AU': 15, 'AV': 10, 'AW': 18, 'AX': 25,
                               'AY': 18, 'AZ': 10, 'BA': 18, 'BB': 10, 'BC': 30, 'BD': 18, 'BE': 10,
                               'BF': 18, 'BG': 10, 'BH': 30, 'BI': 18, 'BJ': 10, 'BK': 25, 'BL': 15,
                               'BM': 18, 'BN': 18, 'BO': 10, 'BP': 18, 'BQ': 10, 'BR': 30, 'BS': 15,
                               'BT': 15, 'BU': 20, 'BV': 20, 'BW': 20, 'BX': 15, 'BY': 25, 'BZ': 20,
                               'CA': 15, 'CB': 15, 'CC': 15, 'CD': 18, 'CE': 10, 'CF': 18, 'CG': 10,
                               'CH': 15, 'CI': 30, 'CJ': 18, 'CK': 10, 'CL': 18, 'CM': 10, 'CN': 30,
                               'CO': 15, 'CP': 25, 'CQ': 18, 'CR': 14, 'CS': 15, 'CT': 10, 'CU': 15,
                               'CV': 10, 'CW': 15, 'CX': 10, 'CY': 20, 'CZ': 18, 'DA': 10, 'DB': 18,
                               'DC': 10, 'DD': 18, 'DE': 10, 'DF': 30, 'DG': 15, 'DH': 20}
                
                for col, width in special_cols.items():
                    ws.column_dimensions[col].width = width
                
                # Hàm tạo header cell
                def set_header_cell(row, col, value, merge_h=False, merge_v=False, merge_end_col=None):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.font = Font(bold=True, size=10, name='Times New Roman')
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = thin_border
                    return cell
                
                # Dòng 8 và 9: Header chi tiết
                # Merge cột A đến L từ dòng 8 đến dòng 9 (header_row = 8, header_row+1 = 9)
  
                # Cột A (1): STT
                ws.merge_cells(start_row=header_row, start_column=1, end_row=header_row+1, end_column=1)
                ws.cell(row=header_row, column=1, value="STT")
                ws.cell(row=header_row, column=1).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=1).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=1).border = thin_border

                # Cột B (2): Họ và tên
                ws.merge_cells(start_row=header_row, start_column=2, end_row=header_row+1, end_column=2)
                ws.cell(row=header_row, column=2, value="Họ và tên")
                ws.cell(row=header_row, column=2).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=2).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=2).border = thin_border

                # Cột C (3): Mã số BHXH
                ws.merge_cells(start_row=header_row, start_column=3, end_row=header_row+1, end_column=3)
                ws.cell(row=header_row, column=3, value="Mã số BHXH")
                ws.cell(row=header_row, column=3).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=3).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=3).border = thin_border

                # Cột D (4): Loại phương án
                ws.merge_cells(start_row=header_row, start_column=4, end_row=header_row+1, end_column=4)
                ws.cell(row=header_row, column=4, value="Loại phương án")
                ws.cell(row=header_row, column=4).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=4).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=4).border = thin_border

                # Cột E (5): Mã loại PA
                ws.merge_cells(start_row=header_row, start_column=5, end_row=header_row+1, end_column=5)
                ws.cell(row=header_row, column=5, value="Mã loại PA")
                ws.cell(row=header_row, column=5).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=5).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=5).border = thin_border

                # Cột F (6): Loại ngày sinh
                ws.merge_cells(start_row=header_row, start_column=6, end_row=header_row+1, end_column=6)
                ws.cell(row=header_row, column=6, value="Loại ngày sinh")
                ws.cell(row=header_row, column=6).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=6).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=6).border = thin_border

                # Cột G (7): Ngày sinh
                ws.merge_cells(start_row=header_row, start_column=7, end_row=header_row+1, end_column=7)
                ws.cell(row=header_row, column=7, value="Ngày sinh")
                ws.cell(row=header_row, column=7).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=7).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=7).border = thin_border

                # Cột H (8): Giới tính
                ws.merge_cells(start_row=header_row, start_column=8, end_row=header_row+1, end_column=8)
                ws.cell(row=header_row, column=8, value="Giới tính")
                ws.cell(row=header_row, column=8).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=8).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=8).border = thin_border

                # Cột I (9): Số CCCD/Hộ chiếu
                ws.merge_cells(start_row=header_row, start_column=9, end_row=header_row+1, end_column=9)
                ws.cell(row=header_row, column=9, value="Số CCCD/Hộ chiếu")
                ws.cell(row=header_row, column=9).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=9).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=9).border = thin_border

                # Cột J (10): Cấp bậc, chức vụ, chức danh nghề nghiệp
                ws.merge_cells(start_row=header_row, start_column=10, end_row=header_row+1, end_column=10)
                ws.cell(row=header_row, column=10, value="Cấp bậc, chức vụ, chức danh nghề nghiệp")
                ws.cell(row=header_row, column=10).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=10).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=10).border = thin_border

                # Cột K (11): Phòng ban làm việc
                ws.merge_cells(start_row=header_row, start_column=11, end_row=header_row+1, end_column=11)
                ws.cell(row=header_row, column=11, value="Phòng ban làm việc")
                ws.cell(row=header_row, column=11).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=11).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=11).border = thin_border

                # Cột L (12): Nơi làm việc
                ws.merge_cells(start_row=header_row, start_column=12, end_row=header_row+1, end_column=12)
                ws.cell(row=header_row, column=12, value="Nơi làm việc")
                ws.cell(row=header_row, column=12).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=12).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=12).border = thin_border  
                
                # M2:O2 = "Tiền đồng"
                ws.merge_cells(start_row=header_row, start_column=13, end_row=header_row, end_column=15)
                ws.cell(row=header_row, column=13, value="Tiền đồng")
                ws.cell(row=header_row, column=13).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=13).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 13, "Mức lương")
                set_header_cell(header_row+1, 14, "Phụ cấp lương")
                set_header_cell(header_row+1, 15, "Các khoản bổ sung")
                
                # P2:S2 = "Hệ số lương"
                ws.merge_cells(start_row=header_row, start_column=16, end_row=header_row, end_column=19)
                ws.cell(row=header_row, column=16, value="Hệ số lương")
                ws.cell(row=header_row, column=16).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=16).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 16, "Hệ số")
                set_header_cell(header_row+1, 17, "Phụ cấp CV")
                set_header_cell(header_row+1, 18, "Phụ cấp TNVK(%)")
                set_header_cell(header_row+1, 19, "Phụ cấp TN nghề(%)")
                
                # T - W
               
                # Cột T (20): Phương án điều chỉnh
                ws.merge_cells(start_row=header_row, start_column=20, end_row=header_row+1, end_column=20)
                ws.cell(row=header_row, column=20, value="Phương án điều chỉnh")
                ws.cell(row=header_row, column=20).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=20).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=20).border = thin_border

                # Cột U (21): Mã PA
                ws.merge_cells(start_row=header_row, start_column=21, end_row=header_row+1, end_column=21)
                ws.cell(row=header_row, column=21, value="Mã PA")
                ws.cell(row=header_row, column=21).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=21).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=21).border = thin_border

                # Cột V (22): Tháng/năm bắt đầu
                ws.merge_cells(start_row=header_row, start_column=22, end_row=header_row+1, end_column=22)
                ws.cell(row=header_row, column=22, value="Tháng/năm bắt đầu")
                ws.cell(row=header_row, column=22).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=22).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=22).border = thin_border

                # Cột W (23): Tháng/năm kết thúc
                ws.merge_cells(start_row=header_row, start_column=23, end_row=header_row+1, end_column=23)
                ws.cell(row=header_row, column=23, value="Tháng/năm kết thúc")
                ws.cell(row=header_row, column=23).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=23).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=23).border = thin_border
                
                # X2:Y2 = "Nghỉ ốm đau/thai sản/không lương"
                ws.merge_cells(start_row=header_row, start_column=24, end_row=header_row, end_column=25)
                ws.cell(row=header_row, column=24, value="Nghỉ ốm đau/thai sản/không lương")
                ws.cell(row=header_row, column=24).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=24).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 24, "Từ ngày")
                set_header_cell(header_row+1, 25, "Đến ngày")
                
                # ===== MERGE CỘT Z ĐẾN AK (CỘT 26 ĐẾN 37) =====
                # Mỗi cột merge riêng lẻ từ dòng 8 xuống dòng 9

                # Cột Z (26): Ghi chú
                ws.merge_cells(start_row=header_row, start_column=26, end_row=header_row+1, end_column=26)
                ws.cell(row=header_row, column=26, value="Ghi chú")
                ws.cell(row=header_row, column=26).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=26).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=26).border = thin_border

                # Cột AA (27): Số sổ BHXH
                ws.merge_cells(start_row=header_row, start_column=27, end_row=header_row+1, end_column=27)
                ws.cell(row=header_row, column=27, value="Số sổ BHXH")
                ws.cell(row=header_row, column=27).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=27).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=27).border = thin_border

                # Cột AB (28): Mức hưởng BHYT
                ws.merge_cells(start_row=header_row, start_column=28, end_row=header_row+1, end_column=28)
                ws.cell(row=header_row, column=28, value="Mức hưởng BHYT")
                ws.cell(row=header_row, column=28).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=28).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=28).border = thin_border

                # Cột AC (29): Tỷ lệ đóng(%)
                ws.merge_cells(start_row=header_row, start_column=29, end_row=header_row+1, end_column=29)
                ws.cell(row=header_row, column=29, value="Tỷ lệ đóng(%)")
                ws.cell(row=header_row, column=29).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=29).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=29).border = thin_border

                # Cột AD (30): Mã vùng sinh sống
                ws.merge_cells(start_row=header_row, start_column=30, end_row=header_row+1, end_column=30)
                ws.cell(row=header_row, column=30, value="Mã vùng sinh sống")
                ws.cell(row=header_row, column=30).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=30).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=30).border = thin_border

                # Cột AE (31): Mã vùng lương tối thiểu
                ws.merge_cells(start_row=header_row, start_column=31, end_row=header_row+1, end_column=31)
                ws.cell(row=header_row, column=31, value="Mã vùng lương tối thiểu")
                ws.cell(row=header_row, column=31).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=31).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=31).border = thin_border

                # Cột AF (32): Có giảm chết
                ws.merge_cells(start_row=header_row, start_column=32, end_row=header_row+1, end_column=32)
                ws.cell(row=header_row, column=32, value="Có giảm chết")
                ws.cell(row=header_row, column=32).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=32).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=32).border = thin_border

                # Cột AG (33): Ngày chết
                ws.merge_cells(start_row=header_row, start_column=33, end_row=header_row+1, end_column=33)
                ws.cell(row=header_row, column=33, value="Ngày chết")
                ws.cell(row=header_row, column=33).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=33).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=33).border = thin_border

                # Cột AH (34): Tính lãi
                ws.merge_cells(start_row=header_row, start_column=34, end_row=header_row+1, end_column=34)
                ws.cell(row=header_row, column=34, value="Tính lãi")
                ws.cell(row=header_row, column=34).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=34).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=34).border = thin_border

                # Cột AI (35): Nhóm vị trí việc làm
                ws.merge_cells(start_row=header_row, start_column=35, end_row=header_row+1, end_column=35)
                ws.cell(row=header_row, column=35, value="Nhóm vị trí việc làm")
                ws.cell(row=header_row, column=35).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=35).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=35).border = thin_border

                # Cột AJ (36): Ngày bắt đầu giữ vị trí
                ws.merge_cells(start_row=header_row, start_column=36, end_row=header_row+1, end_column=36)
                ws.cell(row=header_row, column=36, value="Ngày bắt đầu giữ vị trí")
                ws.cell(row=header_row, column=36).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=36).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=36).border = thin_border

                # Cột AK (37): Ngày kết thúc giữ vị trí
                ws.merge_cells(start_row=header_row, start_column=37, end_row=header_row+1, end_column=37)
                ws.cell(row=header_row, column=37, value="Ngày kết thúc giữ vị trí")
                ws.cell(row=header_row, column=37).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=37).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=37).border = thin_border
                
                # AL2:AN2 = "Hợp đồng lao động"
                ws.merge_cells(start_row=header_row, start_column=38, end_row=header_row, end_column=40)
                ws.cell(row=header_row, column=38, value="Hợp đồng lao động")
                ws.cell(row=header_row, column=38).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=38).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 38, "Loại HĐLĐ")
                set_header_cell(header_row+1, 39, "Hiệu lực từ ngày")
                set_header_cell(header_row+1, 40, "Hiệu lực đến ngày")
                
                # AO2:AP2 = "Ngành nghề nặng nhọc, độc hại"
                ws.merge_cells(start_row=header_row, start_column=41, end_row=header_row, end_column=42)
                ws.cell(row=header_row, column=41, value="Ngành nghề nặng nhọc, độc hại")
                ws.cell(row=header_row, column=41).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=41).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 41, "Ngày bắt đầu")
                set_header_cell(header_row+1, 42, "Ngày kết thúc")
                
                # AQ2:AR2 = "Hợp đồng lao động"
                ws.merge_cells(start_row=header_row, start_column=43, end_row=header_row, end_column=44)
                ws.cell(row=header_row, column=43, value="Hợp đồng lao động")
                ws.cell(row=header_row, column=43).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=43).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 43, "Số")
                set_header_cell(header_row+1, 44, "Ngày ký")
                
                # ===== MERGE CỘT AS ĐẾN AX (CỘT 45 ĐẾN 50) =====
                # Cột AS (45): Quốc tịch
                ws.merge_cells(start_row=header_row, start_column=45, end_row=header_row+1, end_column=45)
                ws.cell(row=header_row, column=45, value="Quốc tịch")
                ws.cell(row=header_row, column=45).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=45).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=45).border = thin_border

                # Cột AT (46): Mã QT
                ws.merge_cells(start_row=header_row, start_column=46, end_row=header_row+1, end_column=46)
                ws.cell(row=header_row, column=46, value="Mã QT")
                ws.cell(row=header_row, column=46).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=46).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=46).border = thin_border

                # Cột AU (47): Dân tộc
                ws.merge_cells(start_row=header_row, start_column=47, end_row=header_row+1, end_column=47)
                ws.cell(row=header_row, column=47, value="Dân tộc")
                ws.cell(row=header_row, column=47).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=47).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=47).border = thin_border

                # Cột AV (48): Mã DT
                ws.merge_cells(start_row=header_row, start_column=48, end_row=header_row+1, end_column=48)
                ws.cell(row=header_row, column=48, value="Mã DT")
                ws.cell(row=header_row, column=48).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=48).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=48).border = thin_border

                # Cột AW (49): Điện thoại liên hệ
                ws.merge_cells(start_row=header_row, start_column=49, end_row=header_row+1, end_column=49)
                ws.cell(row=header_row, column=49, value="Điện thoại liên hệ")
                ws.cell(row=header_row, column=49).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=49).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=49).border = thin_border

                # Cột AX (50): Email liên hệ
                ws.merge_cells(start_row=header_row, start_column=50, end_row=header_row+1, end_column=50)
                ws.cell(row=header_row, column=50, value="Email liên hệ")
                ws.cell(row=header_row, column=50).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=50).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=50).border = thin_border
                
                # AY2:BC2 = "Thông tin đăng ký khai sinh"
                ws.merge_cells(start_row=header_row, start_column=51, end_row=header_row, end_column=55)
                ws.cell(row=header_row, column=51, value="Thông tin đăng ký khai sinh")
                ws.cell(row=header_row, column=51).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=51).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 51, "Tỉnh/thành phố")
                set_header_cell(header_row+1, 52, "Mã tỉnh")
                set_header_cell(header_row+1, 53, "Phường/Xã")
                set_header_cell(header_row+1, 54, "Mã xã")
                set_header_cell(header_row+1, 55, "Địa chỉ khai sinh")
                
                # BD2:BH2 = "Thông tin địa chỉ nhận hồ sơ"
                ws.merge_cells(start_row=header_row, start_column=56, end_row=header_row, end_column=60)
                ws.cell(row=header_row, column=56, value="Thông tin địa chỉ nhận hồ sơ")
                ws.cell(row=header_row, column=56).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=56).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 56, "Tỉnh/thành phố")
                set_header_cell(header_row+1, 57, "Mã tỉnh")
                set_header_cell(header_row+1, 58, "Phường/Xã")
                set_header_cell(header_row+1, 59, "Mã xã")
                set_header_cell(header_row+1, 60, "Địa chỉ nhận hồ sơ")
                
                # BI2:BL2 = "Nơi đăng ký khám chữa bệnh ban đầu"
                ws.merge_cells(start_row=header_row, start_column=61, end_row=header_row, end_column=64)
                ws.cell(row=header_row, column=61, value="Nơi đăng ký khám chữa bệnh ban đầu")
                ws.cell(row=header_row, column=61).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=61).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 61, "Tỉnh nơi KCB")
                set_header_cell(header_row+1, 62, "Mã tỉnh")
                set_header_cell(header_row+1, 63, "Nơi đăng ký KCB")
                set_header_cell(header_row+1, 64, "Mã BV")
                
                # ===== MERGE CỘT BM ĐẾN BS (CỘT 65 ĐẾN 71) =====
                # Cột BM (65): Đăng ký nhận sổ và thẻ
                ws.merge_cells(start_row=header_row, start_column=65, end_row=header_row+1, end_column=65)
                ws.cell(row=header_row, column=65, value="Đăng ký nhận sổ và thẻ")
                ws.cell(row=header_row, column=65).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=65).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=65).border = thin_border

                # Cột BN (66): Mức tiền đóng
                ws.merge_cells(start_row=header_row, start_column=66, end_row=header_row+1, end_column=66)
                ws.cell(row=header_row, column=66, value="Mức tiền đóng")
                ws.cell(row=header_row, column=66).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=66).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=66).border = thin_border

                # Cột BO (67): Phương thức đóng
                ws.merge_cells(start_row=header_row, start_column=67, end_row=header_row+1, end_column=67)
                ws.cell(row=header_row, column=67, value="Phương thức đóng")
                ws.cell(row=header_row, column=67).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=67).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=67).border = thin_border

                # Cột BP (68): Nội dung thay đổi
                ws.merge_cells(start_row=header_row, start_column=68, end_row=header_row+1, end_column=68)
                ws.cell(row=header_row, column=68, value="Nội dung thay đổi")
                ws.cell(row=header_row, column=68).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=68).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=68).border = thin_border

                # Cột BQ (69): Hồ sơ kèm theo
                ws.merge_cells(start_row=header_row, start_column=69, end_row=header_row+1, end_column=69)
                ws.cell(row=header_row, column=69, value="Hồ sơ kèm theo")
                ws.cell(row=header_row, column=69).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=69).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=69).border = thin_border

                # Cột BR (70): Họ tên người giám hộ
                ws.merge_cells(start_row=header_row, start_column=70, end_row=header_row+1, end_column=70)
                ws.cell(row=header_row, column=70, value="Họ tên người giám hộ")
                ws.cell(row=header_row, column=70).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=70).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=70).border = thin_border

                # Cột BS (71): Mã số hộ gia đình
                ws.merge_cells(start_row=header_row, start_column=71, end_row=header_row+1, end_column=71)
                ws.cell(row=header_row, column=71, value="Mã số hộ gia đình")
                ws.cell(row=header_row, column=71).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=71).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                ws.cell(row=header_row, column=71).border = thin_border
                
                # BT2:BX2 = "Địa chỉ nhận sổ, thẻ"  (BN2:BR2)
                ws.merge_cells(start_row=header_row, start_column=72, end_row=header_row, end_column=76)
                ws.cell(row=header_row, column=72, value="Địa chỉ nhận sổ, thẻ")
                ws.cell(row=header_row, column=72).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=72).alignment = Alignment(horizontal='center')
                
                set_header_cell(header_row+1, 72, "Tỉnh/thành phố")
                set_header_cell(header_row+1, 73, "Mã tỉnh")
                set_header_cell(header_row+1, 74, "Phường/Xã")
                set_header_cell(header_row+1, 75, "Mã xã")
                set_header_cell(header_row+1, 76, "Địa chỉ nhận sổ, thẻ")
                
                # BY2:CO2 = "Thông tin chủ hộ"
                ws.merge_cells(start_row=header_row, start_column=77, end_row=header_row, end_column=93)
                ws.cell(row=header_row, column=77, value="Thông tin chủ hộ")
                ws.cell(row=header_row, column=77).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=77).alignment = Alignment(horizontal='center')
                
                for col, val in [(77, 'Họ tên chủ hộ'), (78, 'Số CCCD/Hộ chiếu'), (79, 'Điện thoại'),
                                 (80, 'Loại giấy tờ'), (81, 'Số giấy tờ'), (82, 'Tỉnh/thành phố'),
                                 (83, 'Mã tỉnh'), (84, 'Phường/Xã'), (85, 'Mã xã'), (86, 'Tổ/Thôn/Xóm'),
                                 (87, 'Địa chỉ hộ khẩu'), (88, 'Tỉnh/Thành phố thường trú'), (89, 'Mã Tỉnh'),
                                 (90, 'Phường/Xã thường trú'), (91, 'Mã xã'), (92, 'Địa chỉ thường trú'),
                                 (93, 'Mã số hộ gia đình')]:
                    set_header_cell(header_row+1, col, val)
                
                # CP8:DH8 = "Phụ lục gia đình"
                ws.merge_cells(start_row=header_row, start_column=94, end_row=header_row, end_column=112)
                ws.cell(row=header_row, column=94, value="Phụ lục gia đình")
                ws.cell(row=header_row, column=94).font = Font(bold=True, size=10, name='Times New Roman')
                ws.cell(row=header_row, column=94).alignment = Alignment(horizontal='center')
                
                for col, val in [(94, 'Họ và tên'), (95, 'Mã số BHXH'), (96, 'Loại ngày sinh'),
                                 (97, 'Ngày sinh'), (98, 'Giới tính'), (99, 'Quốc tịch'), (100, 'Mã QT'),
                                 (101, 'Dân tộc'), (102, 'Mã DT'), (103, 'Số CCCD/Hộ chiếu'),
                                 (104, 'Mối quan hệ với chủ hộ'), (105, 'Mã MQH'), (106, 'Tỉnh/thành phố'),
                                 (107, 'Mã tỉnh'), (108, 'Phường/Xã'), (109, 'Mã xã'), (110, 'Địa chỉ khai sinh'),
                                 (111, 'Người tham gia'), (112, 'Ghi chú')]:
                    set_header_cell(header_row+1, col, val)
                
                # Dòng 10: Đánh số thứ tự cột (1) đến (112)
                stt_row = header_row + 2
                for col in range(1, 113):
                    cell = ws.cell(row=stt_row, column=col, value=f"({col})")
                    cell.font = Font(size=9, name='Times New Roman')
                    cell.alignment = Alignment(horizontal='center')
                    cell.border = thin_border
                
                # ===== SET ALL BORDERS CHO VÙNG A8:DH10 =====
                for row in range(header_row, stt_row + 1):
                    for col in range(1, 115):
                        cell = ws.cell(row=row, column=col)
                        cell.border = thin_border
                
                # ===== ĐIỀN DỮ LIỆU =====
                data_row = stt_row + 1
                
                # Hàm ánh xạ dữ liệu nhân viên vào row
                def fill_employee_data(row, nv, loai_pa, idx):
                    # Lấy mã số BHXH
                    ma_bhxh = nv.get('ma_so_bhxh', '')
                    has_bhxh = bool(ma_bhxh and str(ma_bhxh).strip())
                    
                    # Màu đỏ cho text
                    red_font = Font(color='FF0000', size=10, name='Times New Roman')
                    normal_font = Font(size=10, name='Times New Roman')
                    
                    # Hàm lấy giá trị hoặc "Bổ sung thông tin" màu đỏ
                    def get_value_or_warning(value, is_required=True):
                        if is_required and (value is None or str(value).strip() == ''):
                            return "Bổ sung thông tin", red_font
                        return value, normal_font
                    
                    # A: STT (luôn có, không cần kiểm tra)
                    ws.cell(row=row, column=1, value=idx)
                    
                    # B: Họ và tên
                    val, font = get_value_or_warning(nv.get('ho_ten', ''), has_bhxh)
                    cell = ws.cell(row=row, column=2, value=val)
                    cell.font = font
                    
                    # C: Mã BHXH (bỏ qua kiểm tra vì đang dùng để xét điều kiện)
                    ws.cell(row=row, column=3, value=ma_bhxh)
                    
                    # D: Loại phương án
                    val, font = get_value_or_warning(loai_pa, has_bhxh)
                    cell = ws.cell(row=row, column=4, value=val)
                    cell.font = font
                    
                    # E: Mã loại PA (tự động, không cần kiểm tra)
                    ws.cell(row=row, column=5, value="01" if loai_pa == "Tăng lao động" else "02")
                    
                    # F: Loại ngày sinh (tự động)
                    ws.cell(row=row, column=6, value="01")
                    
                    # G: Ngày sinh
                    val, font = get_value_or_warning(format_date(nv.get('ngay_sinh')), has_bhxh)
                    cell = ws.cell(row=row, column=7, value=val)
                    cell.font = font
                    
                    # Cột H (8): Giới tính - Chuyển đổi từ số sang chữ
                    gt_value = nv.get('Gioi_tinh_ma', 0)

                    if gt_value == 1:
                        gt_display = 'Nam'
                    elif gt_value == 2:
                        gt_display = 'Nữ'
                    else:
                        gt_display = ''

                    cell = ws.cell(row=row, column=8, value=gt_display)
                    cell.font = Font(size=10, name='Times New Roman')
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell.border = thin_border
                    
                    
                    # I: Số CCCD
                    val, font = get_value_or_warning(nv.get('so_cccd', ''), has_bhxh)
                    cell = ws.cell(row=row, column=9, value=val)
                    cell.font = font
                    
                    # J: Chức danh
                    val, font = get_value_or_warning(nv.get('chuc_danh_nghe', ''), has_bhxh)
                    cell = ws.cell(row=row, column=10, value=val)
                    cell.font = font
                    
                    # K: Phòng ban (không kiểm tra)
                    ws.cell(row=row, column=11, value=nv.get('phong_ban_lam_viec', ''))
                    
                    # L: Nơi làm việc
                    val, font = get_value_or_warning(nv.get('Noi_lam_viec', 'Cảng THQT Hòn La'), has_bhxh)
                    cell = ws.cell(row=row, column=12, value=val)
                    cell.font = font
                    
                    # M: Mức lương
                    val, font = get_value_or_warning(nv.get('Luong_bao_hiem', ''), has_bhxh)
                    cell = ws.cell(row=row, column=13, value=val)
                    cell.font = font
                    
                    # N: Phụ cấp lương (không kiểm tra)
                    ws.cell(row=row, column=14, value='')
                    
                    # O: Các khoản bổ sung (không kiểm tra)
                    ws.cell(row=row, column=15, value='')
                    
                    # P: Hệ số
                    val, font = get_value_or_warning(nv.get('He_so_luong', ''), has_bhxh)
                    cell = ws.cell(row=row, column=16, value=val)
                    cell.font = font
                    
                    # Q: Phụ cấp CV (không kiểm tra)
                    ws.cell(row=row, column=17, value=nv.get('Phu_cap_chuc_vu', ''))
                    
                    # R: Phụ cấp TNVK (không kiểm tra)
                    ws.cell(row=row, column=18, value=f"{nv.get('Phu_cap_TNVK', '')}%" if nv.get('Phu_cap_TNVK') else '')
                    
                    # S: Phụ cấp TN nghề (không kiểm tra)
                    ws.cell(row=row, column=19, value=f"{nv.get('Phu_cap_TNN', '')}%" if nv.get('Phu_cap_TNN') else '')
                    
                    # T: Phương án điều chỉnh
                    val, font = get_value_or_warning('', has_bhxh)
                    cell = ws.cell(row=row, column=20, value=val)
                    cell.font = font
                    
                    # U: Mã PA (không kiểm tra)
                    ws.cell(row=row, column=21, value='')
                    
                    # V: Tháng/năm bắt đầu (không kiểm tra)
                    if loai_pa == "Tăng lao động":
                        ws.cell(row=row, column=22, value=format_date(nv.get('Ngay_bat_dau', nv.get('Thang_bat_dau_BH'))))
                    else:
                        ws.cell(row=row, column=22, value=format_date(nv.get('ngay_vao_lam')))
                    
                    # W: Tháng/năm kết thúc (không kiểm tra)
                    if loai_pa == "Giảm lao động":
                        ws.cell(row=row, column=23, value=format_date(nv.get('Ngay_ket_thuc_BH')))
                    else:
                        ws.cell(row=row, column=23, value='')
                    
                    # X: Từ ngày (nghỉ) - không kiểm tra
                    ws.cell(row=row, column=24, value='')
                    
                    # Y: Đến ngày (nghỉ) - không kiểm tra
                    ws.cell(row=row, column=25, value='')
                    
                    # Z: Ghi chú (không kiểm tra)
                    ws.cell(row=row, column=26, value='')
                    
                    # AA: Số sổ BHXH (không kiểm tra)
                    ws.cell(row=row, column=27, value=ma_bhxh)
                    
                    # AB: Mức hưởng BHYT (không kiểm tra theo yêu cầu, nhưng có trong danh sách AC)
                    ws.cell(row=row, column=28, value=nv.get('Muc_huong_BHYT', '80%'))
                    
                    # AC: Tỷ lệ đóng
                    val, font = get_value_or_warning(nv.get('Ty_le_dong', ''), has_bhxh)
                    cell = ws.cell(row=row, column=29, value=val)
                    cell.font = font
                    
                    # AD: Mã vùng sinh sống (không kiểm tra)
                    ws.cell(row=row, column=30, value='')
                    
                    # AE: Mã vùng lương tối thiểu - tự động gán "Vùng III"
                    ws.cell(row=row, column=31, value="Vùng III")
                    ws.cell(row=row, column=31).font = normal_font
                    
                    # AF: Có giảm chết (không kiểm tra)
                    ws.cell(row=row, column=32, value='')
                    
                    # AG: Ngày chết (không kiểm tra)
                    ws.cell(row=row, column=33, value='')
                    
                    # AH: Tính lãi (không kiểm tra)
                    ws.cell(row=row, column=34, value='')
                    
                    # AI: Nhóm vị trí việc làm (không kiểm tra)
                    ws.cell(row=row, column=35, value='')
                    
                    # AJ: Ngày bắt đầu giữ vị trí (không kiểm tra)
                    ws.cell(row=row, column=36, value=format_date(nv.get('ngay_vao_lam')))
                    
                    # AK: Ngày kết thúc giữ vị trí (không kiểm tra)
                    ws.cell(row=row, column=37, value=format_date(nv.get('ngay_ket_thuc')) if nv.get('ngay_ket_thuc') else '')
                    
                    # AL: Loại HĐLĐ (không kiểm tra)
                    ws.cell(row=row, column=38, value=nv.get('loai_hop_dong', ''))
                    
                    # AM: Hiệu lực từ ngày (không kiểm tra)
                    ws.cell(row=row, column=39, value=format_date(nv.get('Ngay_ky_HD') or nv.get('ngay_vao_lam')))
                    
                    # AN: Hiệu lực đến ngày (không kiểm tra)
                    ws.cell(row=row, column=40, value=format_date(nv.get('ngay_ket_thuc')) if nv.get('ngay_ket_thuc') else '')
                    
                    # AO: Ngày bắt đầu (nặng nhọc) - không kiểm tra
                    ws.cell(row=row, column=41, value='')
                    
                    # AP: Ngày kết thúc (nặng nhọc) - không kiểm tra
                    ws.cell(row=row, column=42, value='')
                    
                    # AQ: Số HĐLĐ
                    val, font = get_value_or_warning(nv.get('so_hdld', ''), has_bhxh)
                    cell = ws.cell(row=row, column=43, value=val)
                    cell.font = font
                    
                    # AR: Ngày ký HĐLĐ
                    val, font = get_value_or_warning(format_date(nv.get('Ngay_ky_HD')), has_bhxh)
                    cell = ws.cell(row=row, column=44, value=val)
                    cell.font = font
                    
                    # AS: Quốc tịch
                    val, font = get_value_or_warning(nv.get('Quoc_tich', 'Việt Nam'), has_bhxh)
                    cell = ws.cell(row=row, column=45, value=val)
                    cell.font = font
                    
                    # AT: Mã QT (không kiểm tra - tự động)
                    ws.cell(row=row, column=46, value='VN')
                    
                    # AU: Dân tộc
                    val, font = get_value_or_warning(nv.get('Dan_toc', 'Kinh'), has_bhxh)
                    cell = ws.cell(row=row, column=47, value=val)
                    cell.font = font
                    
                    # AV: Mã DT (không kiểm tra - tự động)
                    ws.cell(row=row, column=48, value='01')
                    
                    # AW: Điện thoại
                    val, font = get_value_or_warning(nv.get('dien_thoai', ''), has_bhxh)
                    cell = ws.cell(row=row, column=49, value=val)
                    cell.font = font
                    
                    # AX: Email (không kiểm tra)
                    ws.cell(row=row, column=50, value=nv.get('Email_lien_he', ''))
                    
                    # AY: Tỉnh/thành phố (khai sinh) - lấy từ địa chỉ thường trú
                    dia_chi = nv.get('Thuong_tru', '')
                    # Hàm trích xuất tỉnh/thành phố từ địa chỉ (giả sử phần cuối là tỉnh)
                    def extract_tinh_from_address(address):
                        if not address:
                            return ''
                        # Tách địa chỉ bằng dấu phẩy hoặc khoảng trắng, lấy phần cuối
                        parts = address.split(',')
                        if len(parts) > 0:
                            last_part = parts[-1].strip()
                            # Loại bỏ các từ như "Thành phố", "Tỉnh" nếu có
                            last_part = last_part.replace('Thành phố', '').replace('Tỉnh', '').strip()
                            return last_part
                        return address[:30] if len(address) > 30 else address

                    val, font = get_value_or_warning(extract_tinh_from_address(dia_chi), has_bhxh)
                    cell = ws.cell(row=row, column=51, value=val)
                    cell.font = font
                    
                    # AZ: Mã tỉnh (khai sinh) - không kiểm tra
                    ws.cell(row=row, column=52, value='')
                    
                    # BA (cột 53): Phường/Xã - Lấy từ Phuong_nhan_HS (giống như cột BF)
                    phuong_xa = nv.get('Phuong_nhan_HS', '')
                    if has_bhxh and not phuong_xa:
                        cell = ws.cell(row=row, column=53, value="Bổ sung thông tin")
                        cell.font = red_font
                    else:
                        cell = ws.cell(row=row, column=53, value=phuong_xa)
                        cell.font = Font(size=10, name='Times New Roman')
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell.border = thin_border
                    
                    # BB: Mã xã (khai sinh) - không kiểm tra
                    ws.cell(row=row, column=54, value='')
                    
                    # BC: Địa chỉ khai sinh (không kiểm tra)
                    ws.cell(row=row, column=55, value=nv.get('Thuong_tru', ''))
                    
                    # BD: Tỉnh/thành phố (nhận hồ sơ)
                    val, font = get_value_or_warning(nv.get('Tinh_nhan_HS', ''), has_bhxh)
                    cell = ws.cell(row=row, column=56, value=val)
                    cell.font = font
                    
                    # BE: Mã tỉnh (nhận hồ sơ) - không kiểm tra
                    ws.cell(row=row, column=57, value='')
                    
                    # BF: Phường/Xã (nhận hồ sơ)
                    val, font = get_value_or_warning(nv.get('Phuong_nhan_HS', ''), has_bhxh)
                    cell = ws.cell(row=row, column=58, value=val)
                    cell.font = font
                    
                    # BG: Mã xã (nhận hồ sơ) - không kiểm tra
                    ws.cell(row=row, column=59, value='')
                    
                    # BH: Địa chỉ nhận hồ sơ
                    val, font = get_value_or_warning(nv.get('Dia_chi_nhan_HS', ''), has_bhxh)
                    cell = ws.cell(row=row, column=60, value=val)
                    cell.font = font
                    
                    # BI: Tỉnh nơi KCB
                    val, font = get_value_or_warning(nv.get('Tinh_KCB', ''), has_bhxh)
                    cell = ws.cell(row=row, column=61, value=val)
                    cell.font = font
                    
                    # BJ: Mã tỉnh (KCB) - không kiểm tra
                    ws.cell(row=row, column=62, value='')
                    
                    # BK: Nơi đăng ký KCB
                    val, font = get_value_or_warning(nv.get('Noi_dang_ky_KCB', ''), has_bhxh)
                    cell = ws.cell(row=row, column=63, value=val)
                    cell.font = font
                    
                    # BL: Mã BV (không kiểm tra)
                    ws.cell(row=row, column=64, value='')
                    
                    # BM: Đăng ký nhận sổ (không kiểm tra)
                    ws.cell(row=row, column=65, value=nv.get('Dang_ky_nhan_so', 'Có'))
                    
                    # ===== TRƯỜNG HỢP CHƯA CÓ MÃ BHXH: KIỂM TRA CÁC CỘT THEO YÊU CẦU =====
                    if not has_bhxh:
                        # BY (77): Họ tên chủ hộ
                        val, font = get_value_or_warning(nv.get('Ho_ten_chu_ho', ''), True)
                        cell = ws.cell(row=row, column=77, value=val)
                        cell.font = font
                        
                        # BZ (78): Số CCCD chủ hộ
                        val, font = get_value_or_warning(nv.get('So_CCCD_chu_ho', ''), True)
                        cell = ws.cell(row=row, column=78, value=val)
                        cell.font = font
                        
                        # CD (82): Tỉnh/thành phố (chủ hộ)
                        val, font = get_value_or_warning(nv.get('Tinh_thanh_pho_chu_ho', ''), True)
                        cell = ws.cell(row=row, column=82, value=val)
                        cell.font = font
                        
                        # CF (84): Phường/Xã (chủ hộ)
                        val, font = get_value_or_warning(nv.get('Phuong_xa_chu_ho', ''), True)
                        cell = ws.cell(row=row, column=84, value=val)
                        cell.font = font
                        
                        # CJ (88): Tỉnh/Thành phố thường trú
                        val, font = get_value_or_warning(nv.get('Tinh_thanh_pho_thuong_tru', ''), True)
                        cell = ws.cell(row=row, column=88, value=val)
                        cell.font = font
                        
                        # CK (89): Mã Tỉnh
                        val, font = get_value_or_warning(nv.get('Ma_tinh_thuong_tru', ''), True)
                        cell = ws.cell(row=row, column=89, value=val)
                        cell.font = font
                        
                        # CL (90): Phường/Xã thường trú
                        val, font = get_value_or_warning(nv.get('Phuong_xa_thuong_tru', ''), True)
                        cell = ws.cell(row=row, column=90, value=val)
                        cell.font = font
                        
                        # CM (91): Mã xã
                        val, font = get_value_or_warning(nv.get('Ma_phuong_xa_thuong_tru', ''), True)
                        cell = ws.cell(row=row, column=91, value=val)
                        cell.font = font
                        
                        # Lấy thông tin thành viên đầu tiên từ bảng phu_luc_gia_dinh
                        db_temp = get_connection()
                        c_temp = db_temp.cursor(dictionary=True)
                        c_temp.execute("SELECT * FROM phu_luc_gia_dinh WHERE nhan_vien_Id = %s LIMIT 1", (nv.get('id'),))
                        tv = c_temp.fetchone()
                        db_temp.close()
                        
                        if tv:
                            # CP (94): Họ và tên
                            val, font = get_value_or_warning(tv.get('ho_ten', ''), True)
                            cell = ws.cell(row=row, column=94, value=val)
                            cell.font = font
                            
                            # CS (97): Ngày sinh
                            val, font = get_value_or_warning(format_date(tv.get('ngay_sinh')), True)
                            cell = ws.cell(row=row, column=97, value=val)
                            cell.font = font
                            
                            # CT (98): Giới tính
                            val, font = get_value_or_warning(tv.get('gioi_tinh', ''), True)
                            cell = ws.cell(row=row, column=98, value=val)
                            cell.font = font
                            
                            # CU (99): Quốc tịch
                            val, font = get_value_or_warning(tv.get('Quoc_tich', ''), True)
                            cell = ws.cell(row=row, column=99, value=val)
                            cell.font = font
                            
                            # CW (101): Dân tộc
                            val, font = get_value_or_warning(tv.get('Dan_toc', ''), True)
                            cell = ws.cell(row=row, column=101, value=val)
                            cell.font = font
                            
                            # CZ (104): Mối quan hệ với chủ hộ
                            val, font = get_value_or_warning(tv.get('Quan_he_voi_chu_ho', ''), True)
                            cell = ws.cell(row=row, column=104, value=val)
                            cell.font = font
                            
                            # DB (106): Tỉnh/thành phố
                            val, font = get_value_or_warning(tv.get('Tinh_thanh_pho', ''), True)
                            cell = ws.cell(row=row, column=106, value=val)
                            cell.font = font
                            
                            # DD (108): Phường/Xã
                            val, font = get_value_or_warning(tv.get('Phuong_xa', ''), True)
                            cell = ws.cell(row=row, column=108, value=val)
                            cell.font = font
                        else:
                            # Nếu không có thành viên, hiển thị "Bổ sung thông tin" cho các cột bắt buộc
                            for col in [94, 97, 98, 99, 101, 104, 106, 108]:
                                cell = ws.cell(row=row, column=col, value="Bổ sung thông tin")
                                cell.font = red_font
                    else:
                        # Nếu có BHXH, các cột này để trống
                        for col in [77, 78, 82, 84, 88, 89, 90, 91, 94, 97, 98, 99, 101, 104, 106, 108]:
                            ws.cell(row=row, column=col, value='')
                    
                    # BS (71): Mã số hộ gia đình (không kiểm tra theo yêu cầu)
                    ws.cell(row=row, column=71, value='')
                    
                    # Các cột còn lại (BT đến DH) để trống
                    for col in range(72, 77):  # BT-BR (72-76)
                        ws.cell(row=row, column=col, value='')
                    for col in range(78, 82):  # BZ-CA (78-81)
                        ws.cell(row=row, column=col, value='')
                    for col in range(83, 84):  # CE (83)
                        ws.cell(row=row, column=col, value='')
                    for col in range(85, 88):  # CG-CH (85-87)
                        ws.cell(row=row, column=col, value='')
                    for col in range(92, 94):  # CN-CO (92-93)
                        ws.cell(row=row, column=col, value='')
                    for col in range(95, 97):  # CQ-CR (95-96)
                        ws.cell(row=row, column=col, value='')
                    for col in range(100, 101):  # CV (100)
                        ws.cell(row=row, column=col, value='')
                    for col in range(102, 104):  # CX-CY (102-103)
                        ws.cell(row=row, column=col, value='')
                    for col in range(105, 106):  # DA (105)
                        ws.cell(row=row, column=col, value='')
                    for col in range(107, 108):  # DC (107)
                        ws.cell(row=row, column=col, value='')
                    for col in range(109, 115):  # DE-DH (109-114)
                        ws.cell(row=row, column=col, value='')
                    
                    # Định dạng border cho dòng dữ liệu
                    for col in range(1, 115):
                        cell = ws.cell(row=row, column=col)
                        cell.border = thin_border
                        if col in [1, 3, 5, 6, 7, 8, 9, 16, 17, 18, 19, 22, 23, 24, 25, 29, 36, 37, 39, 40, 44, 46, 48, 65]:
                            if cell.font != red_font:
                                cell.alignment = Alignment(horizontal='center', vertical='center')
                        else:
                            if cell.font != red_font:
                                cell.alignment = Alignment(horizontal='left', vertical='center')
                
                # Điền dữ liệu tăng
                idx = 1
                for nv in tang:
                    fill_employee_data(data_row + idx - 1, nv, "Tăng lao động", idx)
                    idx += 1
                
                # Điền dữ liệu giảm
                for nv in giam:
                    fill_employee_data(data_row + idx - 1, nv, "Giảm lao động", idx)
                    idx += 1
                
                # Lưu file
                filename = f"D02-LT_{tu_ngay.strftime('%d%m%Y')}_{den_ngay.strftime('%d%m%Y')}.xlsx"
                wb.save(filename)
                
                with open(filename, "rb") as f:
                    st.download_button(
                        label="📥 TẢI FILE EXCEL D02-LT",
                        data=f,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                
                st.success(f"✅ Đã xuất báo cáo D02-LT với {len(tang) + len(giam)} lao động (Tăng: {len(tang)}, Giảm: {len(giam)})")

# ========== BÁO CÁO TÌNH HÌNH SỬ DỤNG LAO ĐỘNG MẪU 01/PLI (EXCEL) ==========
elif menu == "📋 Báo cáo 01/PLI" and st.session_state.role == "admin":
    st.title("📋 Báo cáo tình hình sử dụng lao động")
    st.caption("Theo mẫu 01/PLI Phụ lục I - Nghị định 145/2020/NĐ-CP (sửa đổi bởi Nghị định 35/2022/NĐ-CP)")
    
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    # Định nghĩa border
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Chọn kỳ báo cáo
    col1, col2 = st.columns(2)
    with col1:
        tu_ngay = st.date_input("📅 Từ ngày:", value=date(date.today().year, 1, 1), key="pli_tu")
    with col2:
        den_ngay = st.date_input("📅 Đến ngày:", value=date.today(), key="pli_den")
    
    # Lấy dữ liệu
    db = get_connection()
    c = db.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    c.execute("""
        SELECT 
            nv.STT, nv.Ma_NV, nv.Ho_ten, nv.Ma_so_BHXH, nv.Ngay_sinh, nv.Gioi_tinh,
            nv.So_CCCD, nv.Chuc_danh_nghe, nv.Luong_bao_hiem, nv.He_so_luong,
            nv.Phu_cap_chuc_vu, nv.Phu_cap_TNVK, nv.Phu_cap_TNN, nv.Loai_hop_dong,
            nv.Ngay_vao_lam, nv.Ngay_ky_HD, nv.Ngay_ket_thuc, nv.Thang_bat_dau_BH,
            nv.Thang_ket_thuc_BH, nv.So_HDLD
        FROM nhan_vien nv
        WHERE nv.Trang_thai IN ('DANG_LAM', 'THU_VIEC')
        AND nv.Ngay_vao_lam <= %s
        AND (nv.Ngay_ket_thuc IS NULL OR nv.Ngay_ket_thuc > %s)
        ORDER BY nv.STT ASC
    """, (den_ngay, den_ngay))
    ds_lao_dong = c.fetchall()
    db.close()
    
    st.info(f"📊 Tổng số lao động đang làm việc: **{len(ds_lao_dong)}** người")
    
    if ds_lao_dong:
        if st.button("📥 XUẤT EXCEL MẪU 01/PLI", type="primary", use_container_width=True):
            wb = Workbook()
            ws = wb.active
            ws.title = "BC_Tinh_hinh_su_dung_LD"
            
            ten_cong_ty = COMPANY_CONFIG.get("ten_cong_ty", "CÔNG TY CỔ PHẦN CẢNG HÒN LA")
            dia_chi = COMPANY_CONFIG.get("dia_chi", "")
            ma_so_thue = COMPANY_CONFIG.get("ma_so_thue", "")
            dien_thoai_cty = COMPANY_CONFIG.get("dien_thoai_cty", "")
            
            # ===== HEADER =====
            # Dòng 1: Xóa dòng 1 cũ, thay bằng 2 table
            # Table 1: merge A-F
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
            ws['A1'] = ten_cong_ty
            ws['A1'].font = Font(bold=True, size=13, name='Times New Roman')
            ws['A1'].alignment = Alignment(horizontal='center')
            
            # Table 2: merge T-Z (cột 20-26)
            ws.merge_cells(start_row=1, start_column=20, end_row=1, end_column=26)
            ws['T1'] = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
            ws['T1'].font = Font(bold=True, size=13, name='Times New Roman')
            ws['T1'].alignment = Alignment(horizontal='center')
            
            # Dòng 2: Chia làm 2 table
            # Table 1: merge A-F
            ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=6)
            ws['A2'] = f"Số: 01/BC-01PLI-{datetime.now().year}/CHL"
            ws['A2'].font = Font(size=12, name='Times New Roman')
            ws['A2'].alignment = Alignment(horizontal='center')
            
            # Table 2: merge T-Z
            ws.merge_cells(start_row=2, start_column=20, end_row=2, end_column=26)
            ws['T2'] = "Độc lập - Tự do - Hạnh phúc"
            ws['T2'].font = Font(italic=True, size=12, name='Times New Roman')
            ws['T2'].alignment = Alignment(horizontal='center')
            
            # Dòng 3: merge T-Z, canh phải
            ws.merge_cells(start_row=3, start_column=20, end_row=3, end_column=26)
            ws['T3'] = f"Quảng Trị, ngày {date.today().day} tháng {date.today().month} năm {date.today().year}"
            ws['T3'].font = Font(italic=True, size=12, name='Times New Roman')
            ws['T3'].alignment = Alignment(horizontal='right')
            
            # Dòng 5: Tiêu đề báo cáo
            ws.merge_cells('A5:AA5')
            ws['A5'] = "BÁO CÁO TÌNH HÌNH SỬ DỤNG LAO ĐỘNG"
            ws['A5'].font = Font(bold=True, size=13, name='Times New Roman')
            ws['A5'].alignment = Alignment(horizontal='center')
            
            ws.merge_cells('A6:AA6')
            ws['A6'] = f"(Từ ngày {tu_ngay.strftime('%d/%m/%Y')} đến ngày {den_ngay.strftime('%d/%m/%Y')})"
            ws['A6'].font = Font(size=12, name='Times New Roman')
            ws['A6'].alignment = Alignment(horizontal='center')
            
            # Dòng 8: Kính gửi
            ws.merge_cells('A8:AA8')
            ws['A8'] = "Kính gửi: SỞ NỘI VỤ TỈNH QUẢNG TRỊ"
            ws['A8'].font = Font(bold=True, size=11, name='Times New Roman')
            ws['A8'].alignment = Alignment(horizontal='left')
            
            # Dòng 10-13: Thông tin doanh nghiệp
            ws['A10'] = "1. Thông tin chung về doanh nghiệp:"
            ws['A10'].font = Font(bold=True, size=11, name='Times New Roman')
            
            row_info = 11
            for label in [f"- Tên doanh nghiệp: {ten_cong_ty}", f"- Địa chỉ: {dia_chi}", 
                         f"- Mã số thuế: {ma_so_thue}", f"- Điện thoại: {dien_thoai_cty}"]:
                ws[f'A{row_info}'] = label
                ws[f'A{row_info}'].font = Font(size=11, name='Times New Roman')
                row_info += 1
            
            ws[f'A{row_info + 1}'] = "2. Thông tin tình hình sử dụng lao động của đơn vị:"
            ws[f'A{row_info + 1}'].font = Font(bold=True, size=11, name='Times New Roman')
            
            # ===== BẢNG CHÍNH =====
            # Bắt đầu bảng tại dòng 18 (row 18)
            header_row = 18
            
            # Độ rộng cột (A đến AA = 27 cột)
            col_widths = [5, 25, 18, 15, 8, 18, 25, 12, 18, 18, 12, 15, 
                         12, 12, 12, 12, 15, 12, 12, 18, 18, 18, 18, 18, 18, 18, 20]
            for i, w in enumerate(col_widths, 1):
                ws.column_dimensions[get_column_letter(i)].width = w
            
            # ===== DÒNG 21: Đánh số thứ tự cột từ (1) đến (27) =====
            stt_row = header_row + 3  # Dòng 21
            for col in range(1, 28):
                ws.cell(row=stt_row, column=col, value=f"({col})")
                ws.cell(row=stt_row, column=col).font = Font(size=9, name='Times New Roman')
                ws.cell(row=stt_row, column=col).alignment = Alignment(horizontal='center')
                ws.cell(row=stt_row, column=col).border = thin_border
            
            # ===== MERGE CÁC Ô THEO YÊU CẦU (bắt đầu từ dòng 18) =====
            
            # Cột A (1): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=1, end_row=header_row+2, end_column=1)
            ws.cell(row=header_row, column=1, value="STT")
            
            # Cột B (2): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=2, end_row=header_row+2, end_column=2)
            ws.cell(row=header_row, column=2, value="Họ và tên")
            
            # Cột C (3): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=3, end_row=header_row+2, end_column=3)
            ws.cell(row=header_row, column=3, value="Mã số BHXH")
            
            # Cột D (4): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=4, end_row=header_row+2, end_column=4)
            ws.cell(row=header_row, column=4, value="Ngày sinh")
            
            # Cột E (5): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=5, end_row=header_row+2, end_column=5)
            ws.cell(row=header_row, column=5, value="Giới tính")
            
            # Cột F (6): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=6, end_row=header_row+2, end_column=6)
            ws.cell(row=header_row, column=6, value="Số CCCD/Hộ chiếu")
            
            # Cột G (7): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=7, end_row=header_row+2, end_column=7)
            ws.cell(row=header_row, column=7, value="Chức danh nghề, vị trí, công việc")
            
            # Dòng 18: merge cột H->K (8->11)
            ws.merge_cells(start_row=header_row, start_column=8, end_row=header_row, end_column=11)
            ws.cell(row=header_row, column=8, value="Vị trí việc làm (2)")
            
            # Dòng 18: merge cột L->Q (12->17)
            ws.merge_cells(start_row=header_row, start_column=12, end_row=header_row, end_column=17)
            ws.cell(row=header_row, column=12, value="Tiền lương")
            
            # Dòng 18: merge cột T->X (20->24)
            ws.merge_cells(start_row=header_row, start_column=20, end_row=header_row, end_column=24)
            ws.cell(row=header_row, column=20, value="Loại và hiệu lực hợp đồng")
            
            # Merge dòng 18 & 19: cột R->S (18->19)
            ws.merge_cells(start_row=header_row, start_column=18, end_row=header_row+1, end_column=19)
            ws.cell(row=header_row, column=18, value="Ngành nghề nặng nhọc, độc hại")
            
            # Dòng 19: merge cột M->Q (13->17)
            ws.merge_cells(start_row=header_row+1, start_column=13, end_row=header_row+1, end_column=17)
            ws.cell(row=header_row+1, column=13, value="Phụ cấp")
            
            # Dòng 19: merge cột U->V (21->22)
            ws.merge_cells(start_row=header_row+1, start_column=21, end_row=header_row+1, end_column=22)
            ws.cell(row=header_row+1, column=21, value="Hiệu lực HĐLĐ xác định thời hạn")
            
            # Dòng 19: merge cột W->X (23->24)
            ws.merge_cells(start_row=header_row+1, start_column=23, end_row=header_row+1, end_column=24)
            cell = ws.cell(row=header_row+1, column=23, value="Hiệu lực HĐLĐ khác (dưới 1 tháng, thử việc)")
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Merge dòng 19-20: cột H (8)
            ws.merge_cells(start_row=header_row+1, start_column=8, end_row=header_row+2, end_column=8)
            ws.cell(row=header_row+1, column=8, value="Nhà quản lý")
            
            # Merge dòng 19-20: cột I (9)
            ws.merge_cells(start_row=header_row+1, start_column=9, end_row=header_row+2, end_column=9)
            ws.cell(row=header_row+1, column=9, value="Chuyên môn kỹ thuật bậc cao")
            
            # Merge dòng 19-20: cột J (10)
            ws.merge_cells(start_row=header_row+1, start_column=10, end_row=header_row+2, end_column=10)
            ws.cell(row=header_row+1, column=10, value="Chuyên môn kỹ thuật bậc trung")
            
            # Merge dòng 19-20: cột K (11)
            ws.merge_cells(start_row=header_row+1, start_column=11, end_row=header_row+2, end_column=11)
            ws.cell(row=header_row+1, column=11, value="Khác")
            
            # Merge dòng 19-20: cột L (12)
            ws.merge_cells(start_row=header_row+1, start_column=12, end_row=header_row+2, end_column=12)
            ws.cell(row=header_row+1, column=12, value="Mức lương/Hệ số lương")
            
            # Merge dòng 19-20: cột T (20)
            ws.merge_cells(start_row=header_row+1, start_column=20, end_row=header_row+2, end_column=20)
            ws.cell(row=header_row+1, column=20, value="Ngày bắt đầu HĐLĐ không xác định thời hạn")
            
            # Cột Y (25): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=25, end_row=header_row+2, end_column=25)
            ws.cell(row=header_row, column=25, value="Thời điểm bắt đầu đóng BHXH")
            
            # Cột Z (26): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=26, end_row=header_row+2, end_column=26)
            ws.cell(row=header_row, column=26, value="Thời điểm kết thúc đóng BHXH")
            
            # Cột AA (27): merge dòng 18-20
            ws.merge_cells(start_row=header_row, start_column=27, end_row=header_row+2, end_column=27)
            ws.cell(row=header_row, column=27, value="Ghi chú")
            
                        # ===== BORDER CHO VÙNG A18:AA20 =====
            for row in range(header_row, header_row + 3):
                for col in range(1, 28):
                    cell = ws.cell(row=row, column=col)
                    cell.border = thin_border
            
            # Dòng 20: các cột chi tiết (viết thường, không in hoa)
            ws.cell(row=header_row+2, column=13, value="Phụ cấp chức vụ")
            ws.cell(row=header_row+2, column=14, value="Phụ cấp thâm niên VK(%)")
            ws.cell(row=header_row+2, column=15, value="Phụ cấp thâm niên nghề (%)")
            ws.cell(row=header_row+2, column=16, value="Phụ cấp thâm niên nghề (%)")
            ws.cell(row=header_row+2, column=17, value="Các khoản bổ sung")
            ws.cell(row=header_row+2, column=18, value="Ngày bắt đầu")
            ws.cell(row=header_row+2, column=19, value="Ngày kết thúc")
            ws.cell(row=header_row+2, column=21, value="Ngày bắt đầu")
            ws.cell(row=header_row+2, column=22, value="Ngày kết thúc")
            ws.cell(row=header_row+2, column=23, value="Ngày bắt đầu")
            ws.cell(row=header_row+2, column=24, value="Ngày kết thúc")
            
            # Định dạng font và alignment cho header (bỏ chữ in hoa, chuyển sang viết thường có dấu)
            for row in range(header_row, header_row + 3):
                for col in range(1, 28):
                    cell = ws.cell(row=row, column=col)
                    if cell.value:
                        # Chuyển chữ in hoa thành viết thường (có dấu)
                        cell_value = cell.value
                        cell.value = cell_value
                        cell.font = Font(bold=True, size=10, name='Times New Roman')
                        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                        cell.border = thin_border
            
            # Định dạng cho dòng 21 (số thứ tự cột)
            for col in range(1, 28):
                cell = ws.cell(row=stt_row, column=col)
                cell.font = Font(size=9, name='Times New Roman')
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            
            # ===== ĐIỀN DỮ LIỆU (từ dòng 22 trở đi) =====
            data_row = stt_row + 1
            for idx, nv in enumerate(ds_lao_dong, 1):
                row = data_row + idx - 1
                
                # Cột A: STT
                ws.cell(row=row, column=1, value=idx)
                
                # Cột B: Họ tên
                ws.cell(row=row, column=2, value=nv.get('ho_ten', ''))
                
                # Cột C: Mã BHXH
                ws.cell(row=row, column=3, value=nv.get('ma_so_bhxh', ''))
                
                # Cột D: Ngày sinh
                ws.cell(row=row, column=4, value=format_date(nv.get('ngay_sinh')))
                
                # Cột E: Giới tính
                gt = nv.get('gioi_tinh', '')
                ws.cell(row=row, column=5, value='Nam' if gt == 'Nam' else 'Nữ' if gt == 'Nữ' else '')
                
                # Cột F: CCCD
                ws.cell(row=row, column=6, value=nv.get('so_cccd', ''))
                
                # Cột G: Chức danh
                ws.cell(row=row, column=7, value=nv.get('chuc_danh_nghe', ''))
                
                # Cột H-K: Phân loại
                # ===== PHÂN LOẠI THEO LOGIC MỚI =====
                cd = (nv.get('chuc_danh_nghe') or '').lower()

                # Cột (8) - Nhà quản lý: Giám đốc, Quản lý, Trưởng phòng, Phó phòng
                is_quan_ly = any(x in cd for x in ['giám đốc', 'trưởng phòng', 'phó'])
                ws.cell(row=row, column=8, value='x' if is_quan_ly else '')

                # Cột (9) - Chuyên môn kỹ thuật bậc cao: có từ 'kỹ thuật'/'kĩ thuật' và chứa 'cao' hoặc 'chính'
                is_chuyen_mon_cao = (any(x in cd for x in ['kỹ thuật', 'kĩ thuật']) and 
                                     any(x in cd for x in ['cao', 'chính']))
                ws.cell(row=row, column=9, value='x' if is_chuyen_mon_cao else '')

                # Cột (11) - Khác: Lao động phổ thông (các trường hợp còn lại sau khi loại trừ)
                # Cột (10) - Chuyên môn kỹ thuật bậc trung: tất cả trường hợp còn lại (không phải quản lý, không phải bậc cao, không phải khác)

                # Xác định xem có thuộc nhóm quản lý hoặc chuyên môn cao không
                is_quan_ly_or_cao = is_quan_ly or is_chuyen_mon_cao

                # Cột (11) - Khác: lao động phổ thông (có thể nhận diện qua từ khóa 'phổ thông', 'lao động', hoặc mặc định)
                is_khac = any(x in cd for x in ['phổ thông', 'lao động', 'tạp vụ', 'bảo vệ', 'tạp vụ'])
                ws.cell(row=row, column=11, value='x' if is_khac else '')

                # Cột (10) - Chuyên môn kỹ thuật bậc trung: những trường hợp còn lại
                # (không phải quản lý, không phải chuyên môn cao, không phải khác)
                is_trung = (not is_quan_ly) and (not is_chuyen_mon_cao) and (not is_khac)
                ws.cell(row=row, column=10, value='x' if is_trung else '')
                
                # Cột L: Mức lương/Hệ số
                luong = nv.get('Luong_bao_hiem', '')
                heso = nv.get('He_so_luong', '')
                ws.cell(row=row, column=12, value=f"Hệ số: {heso}" if heso and str(heso).strip() else str(luong) if luong else '')
                
                # Cột M: Phụ cấp chức vụ
                ws.cell(row=row, column=13, value=str(nv.get('Phu_cap_chuc_vu', '')) if nv.get('Phu_cap_chuc_vu') else '')
                
                # Cột N: Phụ cấp thâm niên VK
                ws.cell(row=row, column=14, value=f"{nv.get('Phu_cap_TNVK', '')}%" if nv.get('Phu_cap_TNVK') else '')
                
                # Cột O: Phụ cấp thâm niên nghề
                ws.cell(row=row, column=15, value=f"{nv.get('Phu_cap_TNN', '')}%" if nv.get('Phu_cap_TNN') else '')
                
                # Cột P: Phụ cấp thâm niên nghề (cột thứ 2) - để trống
                ws.cell(row=row, column=16, value='')
                
                # Cột Q: Các khoản bổ sung
                ws.cell(row=row, column=17, value='')
                
                # Cột R-S: Nặng nhọc (để trống)
                ws.cell(row=row, column=18, value='')
                ws.cell(row=row, column=19, value='')
                
                # Cột T: Ngày bắt đầu HĐLĐ không xác định thời hạn
                loai_hd = nv.get('loai_hop_dong', '')
                ngay_bd = nv.get('Ngay_ky_HD') or nv.get('ngay_vao_lam')
                ngay_kt = nv.get('ngay_ket_thuc')
                ws.cell(row=row, column=20, value=format_date(ngay_bd) if loai_hd == 'Không xác định thời hạn' else '')
                
                # Cột U-V: HĐ xác định thời hạn
                if loai_hd == 'Xác định thời hạn':
                    ws.cell(row=row, column=21, value=format_date(ngay_bd))
                    ws.cell(row=row, column=22, value=format_date(ngay_kt) if ngay_kt else '')
                else:
                    ws.cell(row=row, column=21, value='')
                    ws.cell(row=row, column=22, value='')
                
                # Cột W-X: HĐ thử việc/dưới 1 tháng
                if loai_hd == 'Thử việc':
                    ws.cell(row=row, column=23, value=format_date(ngay_bd))
                    ws.cell(row=row, column=24, value=format_date(ngay_kt) if ngay_kt else '')
                else:
                    ws.cell(row=row, column=23, value='')
                    ws.cell(row=row, column=24, value='')
                
                # Cột Y: Bắt đầu BHXH
                ws.cell(row=row, column=25, value=format_date(nv.get('Thang_bat_dau_BH')))
                
                # Cột Z: Kết thúc BHXH
                ws.cell(row=row, column=26, value=format_date(nv.get('Thang_ket_thuc_BH')))
                
                # Cột AA: Ghi chú
                ws.cell(row=row, column=27, value=nv.get('so_hdld', ''))
                
                # Định dạng border cho dòng dữ liệu
                for col in range(1, 28):
                    cell = ws.cell(row=row, column=col)
                    cell.border = thin_border
                    cell.font = Font(size=10, name='Times New Roman')
                    if col in [1, 3, 4, 5, 6, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26]:
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                    else:
                        cell.alignment = Alignment(horizontal='left', vertical='center')
            
            # Dòng tổng cộng
            total_row = data_row + len(ds_lao_dong)
            ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=2)
            ws.cell(row=total_row, column=1, value=f"Tổng cộng: {len(ds_lao_dong)} người")
            ws.cell(row=total_row, column=1).font = Font(bold=True, size=10, name='Times New Roman')
            ws.cell(row=total_row, column=1).border = thin_border
            
            # ===== CHỮ KÝ =====
            sign_row = total_row + 3
            ws.merge_cells(start_row=sign_row, start_column=23, end_row=sign_row, end_column=27)
            ws.cell(row=sign_row, column=23, value="ĐẠI DIỆN DOANH NGHIỆP")
            ws.cell(row=sign_row, column=23).font = Font(bold=True, size=11, name='Times New Roman')
            ws.cell(row=sign_row, column=23).alignment = Alignment(horizontal='center')
            
            ws.merge_cells(start_row=sign_row+1, start_column=23, end_row=sign_row+1, end_column=27)
            ws.cell(row=sign_row+1, column=23, value="(Ký, đóng dấu, ghi rõ họ tên)")
            ws.cell(row=sign_row+1, column=23).font = Font(size=10, name='Times New Roman')
            ws.cell(row=sign_row+1, column=23).alignment = Alignment(horizontal='center')
            
            ws.merge_cells(start_row=sign_row+2, start_column=23, end_row=sign_row+2, end_column=27)
            ws.cell(row=sign_row+2, column=23, value=COMPANY_CONFIG.get('dai_dien', 'GIÁM ĐỐC').upper())
            ws.cell(row=sign_row+2, column=23).font = Font(bold=True, size=11, name='Times New Roman')
            ws.cell(row=sign_row+2, column=23).alignment = Alignment(horizontal='center')
            
            # Lưu file
            filename = f"Bao_cao_01_PLI_{tu_ngay.strftime('%d%m%Y')}_{den_ngay.strftime('%d%m%Y')}.xlsx"
            wb.save(filename)
            
            with open(filename, "rb") as f:
                st.download_button(
                    label="📥 TẢI FILE EXCEL",
                    data=f,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            
            st.success(f"✅ Đã xuất báo cáo với {len(ds_lao_dong)} lao động")
    else:
        st.warning("⚠️ Không có lao động nào đang làm việc trong kỳ báo cáo!")
            
st.sidebar.divider()
st.sidebar.caption("© 2026 HRM-Port | Cảng biển quốc tế Hòn La")